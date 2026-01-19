#!/usr/bin/env python
"""Test if add_heading is working correctly"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a simple test document
doc = Document()

# Test adding a heading
heading = doc.add_heading('RESUME', level=1)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Apply formatting
for run in heading.runs:
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)

# Add a paragraph for comparison
para = doc.add_paragraph('This is a paragraph')

# Save and check
doc.save('test_heading_simple.docx')

# Now re-open and check
doc2 = Document('test_heading_simple.docx')
for i, para in enumerate(doc2.paragraphs):
    if 'RESUME' in para.text:
        print(f"Para {i}: {para.text}")
        print(f"  Style: {para.style.name}")
        print(f"  Alignment: {para.alignment}")
        print(f"  Runs: {[run.text for run in para.runs]}")
        for j, run in enumerate(para.runs):
            print(f"    Run {j}: {run.text}")
            print(f"      Bold: {run.bold}")
            print(f"      Font: {run.font.name}")
            print(f"      Size: {run.font.size}")
