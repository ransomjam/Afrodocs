#!/usr/bin/env python3
"""
Test cover page generation with new university data structure.
Tests both University of Bamenda and University of Buea templates.
"""

import os
import json
from datetime import datetime
from docx import Document

# Add backend to path
import sys
backend_path = r"c:\Users\user\Desktop\PATTERN\pattern-formatter\backend"
sys.path.insert(0, backend_path)

from coverpage_generator import generate_cover_page

def test_cover_page_generation():
    """Test generating cover pages for both universities and document types."""
    
    test_cases = [
        {
            "name": "Bamenda - Assignment",
            "data": {
                "university": "Bamenda",
                "documentType": "Assignment",
                "studentName": "John Doe",
                "studentId": "MB23001234",
                "courseCode": "CS201",
                "courseTitle": "Database Systems",
                "institution": "College of Technology",
                "faculty": "College of Technology",
                "department": "Computer Engineering",
                "level": "300 Level",
                "instructor": "Dr. Nkengasong",
                "date": "2025-01-15",
                "title": "Assignment 1: Database Design"
            }
        },
        {
            "name": "Bamenda - Dissertation",
            "data": {
                "university": "Bamenda",
                "documentType": "Dissertation",
                "studentName": "Jane Smith",
                "studentId": "MB21005678",
                "institution": "Faculty of Science",
                "faculty": "Faculty of Science",
                "department": "Mathematics",
                "level": "Masters",
                "supervisor": "Prof. Chinua",
                "coSupervisor": "Dr. Achebe",
                "date": "2025-01-15",
                "title": "Advanced Cryptographic Methods"
            }
        },
        {
            "name": "Buea - Assignment",
            "data": {
                "university": "Buea",
                "documentType": "Assignment",
                "studentName": "Alice Johnson",
                "studentId": "UB24002345",
                "courseCode": "ENG301",
                "courseTitle": "English Literature",
                "institution": "Faculty of Arts (FA)",
                "faculty": "Faculty of Arts (FA)",
                "department": "English (Language)",
                "level": "300 Level",
                "instructor": "Prof. Ngwane",
                "date": "2025-01-15",
                "title": "Literary Analysis Essay"
            }
        },
        {
            "name": "Buea - Dissertation",
            "data": {
                "university": "Buea",
                "documentType": "Dissertation",
                "studentName": "Robert Brown",
                "studentId": "UB22003456",
                "institution": "Faculty of Engineering and Technology (FET)",
                "faculty": "Faculty of Engineering and Technology (FET)",
                "department": "Computer Engineering",
                "level": "Masters",
                "supervisor": "Prof. Obase",
                "coSupervisor": "Dr. Eneha",
                "date": "2025-01-15",
                "title": "Machine Learning for Climate Modeling"
            }
        }
    ]
    
    results = []
    
    for test_case in test_cases:
        print(f"\n{'='*70}")
        print(f"Testing: {test_case['name']}")
        print(f"{'='*70}")
        
        try:
            output_path, error = generate_cover_page(test_case['data'])
            
            if error:
                results.append({
                    "test": test_case['name'],
                    "status": "FAILED",
                    "error": error
                })
                print(f"❌ ERROR: {error}")
            else:
                # Verify file exists
                if os.path.exists(output_path):
                    file_size = os.path.getsize(output_path)
                    
                    # Try to open and verify content
                    doc = Document(output_path)
                    paragraph_count = len(doc.paragraphs)
                    
                    # Check that placeholder markers are NOT in the document
                    full_text = "\n".join([p.text for p in doc.paragraphs])
                    has_placeholders = "{{" in full_text and "}}" in full_text
                    
                    if has_placeholders:
                        results.append({
                            "test": test_case['name'],
                            "status": "WARNING",
                            "message": "Document still contains placeholder markers",
                            "output": output_path,
                            "size": file_size
                        })
                        print(f"⚠️  WARNING: Document still has placeholders!")
                    else:
                        results.append({
                            "test": test_case['name'],
                            "status": "PASSED",
                            "output": output_path,
                            "size": file_size,
                            "paragraphs": paragraph_count
                        })
                        print(f"✅ PASSED: Generated successfully")
                        print(f"   Output: {output_path}")
                        print(f"   Size: {file_size} bytes")
                        print(f"   Paragraphs: {paragraph_count}")
                else:
                    results.append({
                        "test": test_case['name'],
                        "status": "FAILED",
                        "error": "Output file not created"
                    })
                    print(f"❌ ERROR: Output file not created at {output_path}")
                    
        except Exception as e:
            results.append({
                "test": test_case['name'],
                "status": "FAILED",
                "error": str(e)
            })
            print(f"❌ EXCEPTION: {str(e)}")
    
    # Summary
    print(f"\n{'='*70}")
    print("TEST SUMMARY")
    print(f"{'='*70}")
    
    passed = sum(1 for r in results if r['status'] == 'PASSED')
    failed = sum(1 for r in results if r['status'] == 'FAILED')
    warnings = sum(1 for r in results if r['status'] == 'WARNING')
    
    for r in results:
        if r['status'] == 'PASSED':
            print(f"✅ {r['test']}")
        elif r['status'] == 'FAILED':
            print(f"❌ {r['test']}: {r.get('error', 'Unknown error')}")
        else:
            print(f"⚠️  {r['test']}: {r.get('message', 'Warning')}")
    
    print(f"\nTotal: {len(results)} | Passed: {passed} | Failed: {failed} | Warnings: {warnings}")
    
    return passed, failed, warnings

if __name__ == "__main__":
    passed, failed, warnings = test_cover_page_generation()
    print(f"\n{'='*70}")
    print(f"Result: {'✅ ALL TESTS PASSED' if failed == 0 and warnings == 0 else '❌ SOME TESTS FAILED' if failed > 0 else '⚠️  WARNINGS PRESENT'}")
    print(f"{'='*70}")
