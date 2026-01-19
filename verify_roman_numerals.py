#!/usr/bin/env python3
"""
Verify Roman numeral page numbering in generated documents
"""

import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'pattern-formatter', 'backend'))

from docx import Document
from docx.oxml.ns import qn

def extract_footer_fields(doc_path):
    """Extract page numbering info from document footers"""
    doc = Document(doc_path)
    
    print("Analyzing footers in: {}".format(doc_path))
    print("")
    
    for i, section in enumerate(doc.sections):
        print("Section {}:".format(i))
        
        # Get section numbering format
        section_pr = section._sectPr
        pgNumType = section_pr.find(qn('w:pgNumType'))
        if pgNumType is not None:
            fmt = pgNumType.get(qn('w:fmt'))
            start = pgNumType.get(qn('w:start'))
            print("  Numbering format: fmt={}, start={}".format(fmt, start))
        else:
            print("  No explicit numbering format found")
        
        # Get footer paragraph contents
        footer = section.footer
        if footer.paragraphs:
            print("  Footer content:")
            for p_idx, para in enumerate(footer.paragraphs):
                # Get all runs and fields
                for run_idx, run in enumerate(para.runs):
                    text = run.text if run.text else "[empty]"
                    print("    Run {}: '{}'".format(run_idx, text))
                
                # Check for field codes
                for elem in para._element.iter():
                    if 'fldChar' in elem.tag:
                        fldCharType = elem.get(qn('w:fldCharType'))
                        print("    Field: fldCharType={}".format(fldCharType))
                    elif 'instrText' in elem.tag:
                        instr = elem.text if elem.text else ""
                        print("    Instruction: {}".format(instr))
        else:
            print("  No footer paragraphs")
        print("")

if __name__ == '__main__':
    extract_footer_fields('test_roman_numerals.docx')
