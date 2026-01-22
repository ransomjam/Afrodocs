#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Check paragraph styles to understand why they're being justified.
"""

import os
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def check_styles():
    """Check styles of paragraphs."""
    outputs_dir = Path(__file__).parent / 'pattern-formatter' / 'backend' / 'outputs'
    
    # Get most recent docx file
    docx_files = sorted(outputs_dir.glob('*.docx'), key=os.path.getmtime, reverse=True)
    if not docx_files:
        print("No documents found")
        return
    
    recent_doc = docx_files[0]
    doc = Document(str(recent_doc))
    
    print(f"\nDocument: {recent_doc.name}\n")
    print(f"{'#':<4} {'Align':<12} {'Style':<25} {'Content':<45}")
    print("-" * 90)
    
    for idx in range(min(25, len(doc.paragraphs))):
        para = doc.paragraphs[idx]
        align = "None" if para.alignment is None else str(para.alignment).split('.')[-1]
        style = para.style.name
        text = para.text[:40].replace('\n', ' ') if para.text else "(empty)"
        
        # Mark suspicious ones
        marker = " <-- ISSUE" if style in ['Normal', 'List Number', 'List Bullet'] and idx < 23 else ""
        
        print(f"{idx:<4} {align:<12} {style:<25} {text:<45}{marker}")

if __name__ == '__main__':
    check_styles()
