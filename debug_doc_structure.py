#!/usr/bin/env python3
"""
Debug: Extract all content from sample documents to see actual structure
"""

from docx import Document

print("="*70)
print("EXTRACTING FULL DOCUMENT CONTENT")
print("="*70)

doc1_path = r"c:\Users\user\Desktop\PATTERN\Samples\numbering and bulleting sample 1.docx"

try:
    doc = Document(doc1_path)
    
    print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
    print("\nFirst 50 paragraphs (raw):")
    print("-" * 70)
    
    for i, para in enumerate(doc.paragraphs[:50]):
        text = para.text
        if text or i < 20:  # Show first 20 even if empty
            print(f"Para {i}: '{text}' (level:{para.style.name})")
    
    print("\n" + "="*70)
    print("LOOKING FOR CHAPTER MARKERS")
    print("="*70)
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.lower()
        if 'chapter' in text or 'introduction' in text or '1.' in text[:5]:
            print(f"Para {i}: {para.text}")
    
except Exception as e:
    print(f"Error: {e}")
    import traceback
    traceback.print_exc()
