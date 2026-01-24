"""Simple test to verify the page break fix by processing a document through the Flask app"""
import os
import sys
import io

# Add backend to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'pattern-formatter', 'backend'))

from docx import Document

# Test documents
input_file = "samples/new page problems/document with new page problem input.docx"
output_file = "test_fixed_output.docx"

print("="*80)
print("TESTING PAGE BREAK FIX - SIMPLE VERSION")
print("="*80)

# Read input
print(f"\n1. Reading input: {input_file}")
with open(input_file, 'rb') as f:
    input_content = f.read()

input_doc = Document(io.BytesIO(input_content))
text_lines = []
for para in input_doc.paragraphs:
    text = para.text.strip()
    if text:
        text_lines.append(text)

input_text = '\n'.join(text_lines)
print(f"   Extracted {len(text_lines)} paragraphs")

# Import the processing classes
from pattern_formatter_backend import DocumentProcessor, WordGenerator, PatternEngine

# Create instances
print("\n2. Creating processor and generator...")
processor = DocumentProcessor()
generator = WordGenerator()

# Analyze document
print("\n3. Analyzing document...")
analyzed_lines = processor.analyze_document(input_text)
print(f"   Analyzed {len(analyzed_lines)} lines")

# Structure into sections
print("\n4. Structuring into sections...")
sections = generator.structure_sections(analyzed_lines)
print(f"   Created {len(sections)} sections")

# Check what sections are getting page breaks
print("\n5. Checking sections for page break flags...")
for i, section in enumerate(sections):
    heading = section.get('heading', '')
    if not heading:
        continue
    
    upper_heading = heading.upper()
    if 'RECOMMENDATION' in upper_heading or 'CONCLUSION' in upper_heading or '5.1' in heading or '5.2' in heading:
        print(f"   [{i}] '{heading}'")
        print(f"       level: {section.get('level')}")
        print(f"       type: {section.get('type')}")
        print(f"       use_page_break_before: {section.get('use_page_break_before')}")
        print(f"       needs_page_break: {section.get('needs_page_break')}")

# Generate output
print(f"\n6. Generating output document: {output_file}")
generator.generate(sections, output_path=output_file)

# Analyze output
print("\n7. Analyzing output document...")
output_doc = Document(output_file)

# Check heading styles
print("\n--- HEADING STYLES IN NEW OUTPUT ---")
for i in range(1, 5):
    try:
        style = output_doc.styles[f'Heading {i}']
        pf = style.paragraph_format
        print(f"Heading {i}: page_break_before={pf.page_break_before}, keep_with_next={pf.keep_with_next}, keep_together={pf.keep_together}")
    except KeyError:
        print(f"Heading {i}: NOT FOUND")

# Check for page_break_before=True on numbered subsections
print("\n--- CHECKING NUMBERED SUBSECTIONS ---")
import re
issues_found = False
for i, para in enumerate(output_doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    
    # Check if numbered subsection
    if re.match(r'^\d+\.\d+', text):
        pf = para.paragraph_format
        style_name = para.style.name if para.style else "None"
        
        if pf.page_break_before:
            print(f"[{i}] ISSUE: '{text[:50]}...' has page_break_before=True!")
            issues_found = True

if not issues_found:
    print("✓ No numbered subsections have page_break_before=True!")

print("\n" + "="*80)
print("TEST COMPLETE")
print("="*80)
