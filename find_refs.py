#!/usr/bin/env python
# -*- coding: utf-8 -*-

from docx import Document

docx_path = "samples/bulleting and numbering sample 2.docx"
doc = Document(docx_path)

print("Looking for REFERENCES/BIBLIOGRAPHY section...\n")

# Find where references section starts
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    if 'REFERENCE' in text.upper() or 'BIBLIOGRAPHY' in text.upper():
        print(f"Found section heading at line {i}: {text}\n")
        
        # Show next 10 paragraphs
        print("Next paragraphs in this section:")
        for j in range(i, min(i+10, len(doc.paragraphs))):
            p = doc.paragraphs[j]
            left_indent = p.paragraph_format.left_indent.inches if p.paragraph_format.left_indent else 0
            first_line_indent = p.paragraph_format.first_line_indent.inches if p.paragraph_format.first_line_indent else 0
            ptext = p.text[:50].strip()
            
            print(f"  {j}: LEFT={left_indent:.2f}, FIRST={first_line_indent:.2f} | {ptext}")
        
        print()
