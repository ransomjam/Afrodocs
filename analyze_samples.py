#!/usr/bin/env python
# -*- coding: utf-8 -*-

from docx import Document

files = [
    "samples/numbering and bulleting sample 1.docx",
    "samples/bulleting and numbering sample 2.docx"
]

for docx_path in files:
    print(f"\n{'='*80}")
    print(f"ANALYZING: {docx_path}")
    print(f"{'='*80}\n")
    
    doc = Document(docx_path)
    print(f"Total paragraphs: {len(doc.paragraphs)}\n")
    
    issues = []
    
    # Check for formatting issues
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        if not text:
            continue
        
        # Get formatting info
        has_bold = any(run.bold for run in para.runs)
        left_indent = para.paragraph_format.left_indent.inches if para.paragraph_format.left_indent else 0
        first_line_indent = para.paragraph_format.first_line_indent.inches if para.paragraph_format.first_line_indent else 0
        
        # Check for hanging indent issue
        if left_indent > 0.1 and first_line_indent < -0.1:
            issues.append({
                'line': i,
                'issue': f'HANGING INDENT: left={left_indent:.2f}, first_line={first_line_indent:.2f}',
                'text': text[:60]
            })
        
        # Check for unnecessary indentation on non-list items
        if left_indent > 0.1 and not (text[0].isdigit() or text[0] in '•-*○■'):
            issues.append({
                'line': i,
                'issue': f'LEFTINDENT on non-list: {left_indent:.2f}"',
                'text': text[:60]
            })
        
        # Check for list items not properly formatted
        if (text[0].isdigit() or text[0] in '•-*○■'):
            # This is a list item
            if len(text.split()) > 20 and not has_bold:
                # Substantial item should be bold if it's a title
                if ':' in text[:50]:  # Has a colon = title
                    issues.append({
                        'line': i,
                        'issue': f'TITLE NOT BOLD: {len(text.split())} words, has colon',
                        'text': text[:60]
                    })
    
    if not issues:
        print("✓ NO FORMATTING ISSUES FOUND\n")
    else:
        print(f"Found {len(issues)} potential issues:\n")
        for issue in issues[:20]:  # Show first 20
            print(f"Line {issue['line']}: {issue['issue']}")
            print(f"  > {issue['text']}\n")

print("\n" + "="*80)
print("ANALYSIS COMPLETE")
