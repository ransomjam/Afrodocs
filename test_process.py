#!/usr/bin/env python

import sys
sys.path.insert(0, 'pattern-formatter/backend')

from pattern_formatter_backend import DocumentProcessor, WordGenerator
import logging

logging.basicConfig(level=logging.WARNING)

# Read the test file
with open('test_formatting.txt', 'r') as f:
    content = f.read()

# Process it
processor = DocumentProcessor()
analyzed = processor.process_text(content)

print('=== DOCUMENT ANALYSIS ===')
print(f'Total analyzed lines: {len(analyzed)}')

# Generate Word document
generator = WordGenerator()
output_path = 'test_formatting_output.docx'
generator.generate(
    analyzed,
    output_path,
    font_size=12,
    line_spacing=1.5
)
print(f'\nGenerated: {output_path}')

# Read back and verify formatting
from docx import Document
doc = Document(output_path)

print('\n=== OUTPUT VERIFICATION ===')
print(f'Generated paragraphs: {len(doc.paragraphs)}')
print()

# Look for numbered/bulleted items to verify formatting
print('=== NUMBERED/BULLETED ITEMS FOUND ===')
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text and (text[0].isdigit() or text[0] in '•-*○'):
        has_bold = any(run.bold for run in para.runs)
        words = len(text.split())
        print(f'{i}: [BOLD={has_bold}] [{words} words] {text[:70]}')
