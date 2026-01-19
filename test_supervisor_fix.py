"""Test supervisor fields display in cover pages"""
import os
import sys
from docx import Document

sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'pattern-formatter', 'backend'))

from coverpage_generator import generate_cover_page

# Test data with supervisor info
test_data = {
    'university': 'Bamenda',
    'documentType': 'Dissertation',
    'studentName': 'John Doe',
    'studentId': 'MB26TEST001',
    'institution': 'College of Technology',
    'department': 'Electrical and Electronic Engineering',
    'faculty': 'Faculty of Engineering and Technology',
    'level': '400 Level',
    'supervisor': 'Prof. Dr. Emmanuel Tanyi',  # Try both keys
    'academicSupervisor': 'Prof. Dr. Emmanuel Tanyi',
    'fieldSupervisor': 'Dr. Bernard Ngu',
    'title': 'Advanced Wireless Communication Systems',
    'monthYear': 'January 2026'
}

# Generate
result = generate_cover_page(test_data)
print(f"\nGenerated: {result}")

# Extract document path
doc_path = result[0] if isinstance(result, tuple) else result
if doc_path is None:
    error = result[1] if isinstance(result, tuple) else None
    print(f"ERROR: {error}")
    sys.exit(1)

# Extract and check
if os.path.exists(doc_path):
    doc = Document(doc_path)
    
    print("\n=== DISSERTATION CONTENT ===")
    
    # Check paragraphs
    all_text = []
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            all_text.append(para.text)
            print(f"{i}: {para.text[:100]}")
    
    # Check textboxes
    print("\n=== TEXTBOX CONTENT ===")
    for shape in doc.inline_shapes:
        if hasattr(shape, 'textbox'):
            print(f"Shape: {shape.name}")
            for para in shape.textbox.paragraphs:
                if para.text.strip():
                    print(f"  {para.text[:100]}")
    
    # Find supervisor mentions
    print("\n=== SUPERVISOR SEARCH ===")
    full_content = '\n'.join(all_text)
    if 'Emmanuel' in full_content:
        print("[OK] Academic Supervisor FOUND")
    else:
        print("[MISSING] Academic Supervisor NOT found")
    
    if 'Bernard' in full_content:
        print("[OK] Field Supervisor FOUND")
    else:
        print("[MISSING] Field Supervisor NOT found")
    
    # Check for placeholder markers
    if '{{' in full_content:
        markers = [p for p in full_content.split() if '{{' in p]
        print(f"\n[WARNING] REMAINING PLACEHOLDERS: {markers}")
    else:
        print("\n[OK] No placeholder markers remaining")
