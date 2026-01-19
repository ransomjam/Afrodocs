#!/usr/bin/env python
"""
Extract faculties and departments from University Word documents
"""

import os
from docx import Document

def extract_from_document(doc_path):
    """Extract schools/faculties and departments from document"""
    print(f"\n{'='*70}")
    print(f"Extracting from: {os.path.basename(doc_path)}")
    print('='*70)
    
    doc = Document(doc_path)
    
    # Extract all text
    all_text = []
    for p in doc.paragraphs:
        if p.text.strip():
            all_text.append(p.text)
    
    # Also check tables
    if doc.tables:
        print(f"Found {len(doc.tables)} table(s)")
        for table_idx, table in enumerate(doc.tables):
            print(f"\nTable {table_idx + 1}:")
            for row_idx, row in enumerate(table.rows):
                row_data = [cell.text.strip() for cell in row.cells]
                print(f"  Row {row_idx}: {row_data}")
                for cell_text in row_data:
                    if cell_text and cell_text not in all_text:
                        all_text.append(cell_text)
    
    print("\nExtracted content:")
    print("-" * 70)
    for line in all_text[:50]:  # Show first 50 lines
        print(f"  {line}")
    
    return all_text

# Extract from both universities
bamenda_doc = r"c:\Users\user\Desktop\PATTERN\pattern-formatter\Cover Pages\Cover Pages _ University of Bamenda\The University of Bamenda _ Schools-Faculties-Departments.docx"
buea_doc = r"c:\Users\user\Desktop\PATTERN\pattern-formatter\Cover Pages\Cover Page _ University of Buea\University of Buea _ Schools-Faculties-Departments.docx"

bamenda_data = extract_from_document(bamenda_doc)
buea_data = extract_from_document(buea_doc)

print("\n" + "="*70)
print("SUMMARY")
print("="*70)
print(f"Bamenda data lines: {len(bamenda_data)}")
print(f"Buea data lines: {len(buea_data)}")
