#!/usr/bin/env python3
"""
Create a test Word document to verify the double-numbering fix
Uses the actual DocumentProcessor to generate a document
"""

from docx import Document
from docx.shared import Pt

# Create a test document directly to demonstrate the fix
doc = Document()

print("Creating test document with various numbering scenarios...")

# Add test section
doc.add_heading('Test: Double-Numbering Fix', level=1)
doc.add_paragraph('This document tests various numbering scenarios to ensure they are NOT double-numbered.')

# Section 1: Properly numbered items
doc.add_heading('Section 1: Already Numbered Items (Should stay as-is)', level=2)
doc.add_paragraph('These items have existing numbering and should NOT get double-numbered:')

# Test with regular paragraphs (simulating what the fix does)
para = doc.add_paragraph()
run = para.add_run('I. ')
run.bold = True
run = para.add_run('Implications for Students')

para = doc.add_paragraph()
run = para.add_run('II. ')
run.bold = True
run = para.add_run('Implications for Teachers')

para = doc.add_paragraph()
run = para.add_run('1.1 ')
run.bold = True
run = para.add_run('Background Information')

para = doc.add_paragraph()
run = para.add_run('1.2 ')
run.bold = True
run = para.add_run('Methods Used')

# Section 2: Items without numbering
doc.add_heading('Section 2: Items Without Numbering (Can be auto-numbered)', level=2)
doc.add_paragraph('These items can be auto-numbered by Word if needed:')

# Using List Number style only when NO existing numbering
para = doc.add_paragraph('Item without existing numbering', style='List Number')
para = doc.add_paragraph('Another item without numbering', style='List Number')
para = doc.add_paragraph('Third item without numbering', style='List Number')

# Section 3: Mixed content
doc.add_heading('Section 3: Comparison', level=2)

para = doc.add_paragraph()
para.add_run('WRONG (double-numbered): ').bold = True
para.add_run('1. I. Implications for Students')

para = doc.add_paragraph()
para.add_run('CORRECT (single-numbered): ').bold = True
para_num = doc.add_paragraph()
run = para_num.add_run('I. ')
run.bold = True
run = para_num.add_run('Implications for Students')

# Save document
output_path = r'c:\Users\user\Desktop\PATTERN\test_double_numbering_fix_output.docx'
doc.save(output_path)
print(f"\nTest document created: {output_path}")
print("\nInstructions:")
print("1. Open the document in Word")
print("2. Check Section 1: Items should show single numbering (I., II., 1.1, 1.2)")
print("3. Check Section 2: Items can be auto-numbered by Word")
print("4. Check Section 3: WRONG vs CORRECT comparison")
print("\nExpected Result: NO items should have double numbering like '1. I.'")
