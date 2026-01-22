#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Check if coverpage positioning was preserved in the most recent document.
This verifies that our fix is working by checking the properties.
"""

import os
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def analyze_document_positioning():
    """Analyze the most recent document for positioning preservation."""
    
    outputs_dir = Path(__file__).parent / 'pattern-formatter' / 'backend' / 'outputs'
    
    # Get most recent docx file
    docx_files = sorted(outputs_dir.glob('*.docx'), key=os.path.getmtime, reverse=True)
    if not docx_files:
        print("No documents found")
        return False
    
    recent_doc = docx_files[0]
    doc = Document(str(recent_doc))
    
    print(f"\n{'='*90}")
    print(f"Analyzing: {recent_doc.name}")
    print(f"{'='*90}\n")
    
    # Find section break
    section_break_idx = None
    for idx, para in enumerate(doc.paragraphs):
        pPr = para._element.get_or_add_pPr()
        sectPr = pPr.find(qn('w:sectPr'))
        if sectPr is not None:
            section_break_idx = idx
            break
    
    if section_break_idx is None:
        print("No section break found - this may not be a merged document")
        return False
    
    print(f"Section break at paragraph {section_break_idx}")
    print(f"Coverpage: paragraphs 0-{section_break_idx}")
    print(f"Body: paragraphs {section_break_idx+1}-{len(doc.paragraphs)-1}\n")
    
    # Load template for comparison
    template_path = Path(__file__).parent / 'pattern-formatter' / 'backend' / 'coverpage_template' / 'dissertation_coverpage_template.docx'
    template_doc = Document(str(template_path))
    
    print(f"Template has {len(template_doc.paragraphs)} paragraphs")
    print(f"Merged coverpage has {section_break_idx} paragraphs\n")
    
    # Compare alignments
    print(f"{'Para':<5} {'Template':<15} {'Merged':<15} {'Status':<15}")
    print("-" * 55)
    
    compare_limit = min(section_break_idx + 1, 25, len(template_doc.paragraphs))
    
    preserved_count = 0
    changed_count = 0
    
    for idx in range(compare_limit):
        t_para = template_doc.paragraphs[idx]
        m_para = doc.paragraphs[idx]
        
        t_align = "None" if t_para.alignment is None else str(t_para.alignment).split('.')[-1]
        m_align = "None" if m_para.alignment is None else str(m_para.alignment).split('.')[-1]
        
        if t_para.alignment == m_para.alignment:
            status = "PRESERVED"
            preserved_count += 1
        else:
            status = "CHANGED"
            changed_count += 1
        
        print(f"{idx:<5} {t_align:<15} {m_align:<15} {status:<15}")
    
    print(f"\n{'='*90}")
    print(f"Result: {preserved_count} preserved, {changed_count} changed")
    
    if changed_count == 0:
        print(f"SUCCESS: Positioning is preserved!")
        return True
    else:
        print(f"ISSUE: {changed_count} paragraphs had changed alignment")
        return False

if __name__ == '__main__':
    analyze_document_positioning()
