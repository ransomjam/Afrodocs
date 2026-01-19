#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Quick test for textbox field replacement fix.
"""

import os
import sys
import io
from datetime import datetime
from docx import Document

# Fix encoding for Windows console
if sys.stdout.encoding != 'utf-8':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

# Add backend to path
backend_path = r"c:\Users\user\Desktop\PATTERN\pattern-formatter\backend"
sys.path.insert(0, backend_path)

# Import without starting Flask
from coverpage_generator import generate_cover_page

def check_document_content(doc_path):
    """Check if document contains expected values in textboxes."""
    doc = Document(doc_path)
    
    # Get all text including textboxes
    full_text = []
    
    # Paragraphs
    for p in doc.paragraphs:
        full_text.append(p.text)
    
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    full_text.append(p.text)
    
    # Textboxes
    from docx.oxml.ns import qn
    if doc.element.body is not None:
        for txbx in doc.element.body.iter(qn('w:txbxContent')):
            for p in txbx.iter(qn('w:p')):
                para_text = ""
                for r in p.iter(qn('w:r')):
                    for t in r.iter(qn('w:t')):
                        if t.text:
                            para_text += t.text
                if para_text.strip():
                    full_text.append(para_text)
    
    return full_text

def test_textbox_replacement():
    """Test textbox field replacement."""
    
    print("\n" + "="*70)
    print("TEST: Textbox Field Replacement")
    print("="*70)
    
    test_data = {
        "university": "Bamenda",
        "documentType": "Dissertation",
        "studentName": "JOHN TESTING DOE",
        "studentId": "MB26TEST001",
        "institution": "Faculty of Science",
        "faculty": "Faculty of Science",
        "department": "Physics Department",
        "level": "Masters Level",
        "supervisor": "DR. TEST SUPERVISOR",
        "coSupervisor": "DR. TEST CO-SUPERVISOR",
        "fieldSupervisor": "DR. FIELD SUPERVISOR",
        "date": "2026-01-15",
        "title": "Testing Textbox Replacements"
    }
    
    try:
        print("\nGenerating cover page...")
        output_path, error = generate_cover_page(test_data)
        
        if error:
            print(f"❌ ERROR: {error}")
            return False
        
        print(f"✅ Cover page generated")
        print(f"   Output: {os.path.basename(output_path)}")
        
        # Check content
        content = check_document_content(output_path)
        
        print("\n" + "-"*70)
        print("EXTRACTED DOCUMENT CONTENT (Paragraphs + Textboxes):")
        print("-"*70)
        for i, text in enumerate(content):
            if text.strip():
                print(f"{i}: {text[:100]}")
        
        # Check for expected values
        full_doc_text = "\n".join(content)
        
        checks = [
            ("JOHN TESTING DOE", "Student Name"),
            ("MB26TEST001", "Student ID"),
            ("Physics Department", "Department"),
            ("Masters Level", "Level"),
            ("DR. TEST SUPERVISOR", "Supervisor"),
            ("DR. TEST CO-SUPERVISOR", "Co-Supervisor")
        ]
        
        print("\n" + "-"*70)
        print("VALUE VERIFICATION:")
        print("-"*70)
        
        all_found = True
        for value, label in checks:
            if value in full_doc_text:
                print(f"✅ {label}: FOUND - '{value}'")
            else:
                print(f"❌ {label}: NOT FOUND - '{value}'")
                all_found = False
        
        # Check for placeholder markers
        if "{{" in full_doc_text and "}}" in full_doc_text:
            print(f"\n❌ WARNING: Placeholder markers still present in document")
            all_found = False
        else:
            print(f"\n✅ No placeholder markers remaining")
        
        print("\n" + "="*70)
        if all_found:
            print("✅ TEXTBOX REPLACEMENT TEST PASSED")
        else:
            print("❌ TEXTBOX REPLACEMENT TEST FAILED")
        print("="*70)
        
        return all_found
        
    except Exception as e:
        print(f"❌ EXCEPTION: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    success = test_textbox_replacement()
    sys.exit(0 if success else 1)
