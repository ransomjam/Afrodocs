"""Deep analysis of new page problem documents to find root cause"""
import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

samples_dir = "samples/new page problems"

def analyze_document(filepath, label):
    """Analyze a document for page break settings on headings"""
    print(f"\n{'='*80}")
    print(f"ANALYZING: {label}")
    print(f"File: {filepath}")
    print('='*80)
    
    try:
        doc = Document(filepath)
    except Exception as e:
        print(f"ERROR opening document: {e}")
        return
    
    # First, check document styles
    print("\n--- HEADING STYLES CONFIGURATION ---")
    for i in range(1, 5):
        try:
            style = doc.styles[f'Heading {i}']
            pf = style.paragraph_format
            print(f"Heading {i} Style:")
            print(f"  page_break_before: {pf.page_break_before}")
            print(f"  keep_with_next: {pf.keep_with_next}")
            print(f"  keep_together: {pf.keep_together}")
        except KeyError:
            print(f"Heading {i} Style: NOT FOUND")
    
    # Now check actual paragraphs
    print("\n--- PARAGRAPHS WITH page_break_before=True ---")
    found_any = False
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        pf = para.paragraph_format
        if pf.page_break_before:
            found_any = True
            style_name = para.style.name if para.style else "None"
            print(f"[{i}] '{text[:60]}...' (style: {style_name})")
            print(f"     page_break_before: {pf.page_break_before}")
    
    if not found_any:
        print("  (No paragraphs with page_break_before=True)")
    
    # Check for headings that look like numbered subsections
    print("\n--- NUMBERED SUBSECTION HEADINGS ---")
    subsection_pattern = re.compile(r'^\d+\.\d+')
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        
        if subsection_pattern.match(text):
            pf = para.paragraph_format
            style_name = para.style.name if para.style else "None"
            print(f"[{i}] '{text[:60]}...' (style: {style_name})")
            print(f"     page_break_before: {pf.page_break_before}")
            print(f"     keep_with_next: {pf.keep_with_next}")
    
    # Check for headings containing RECOMMENDATIONS or CONCLUSION
    print("\n--- HEADINGS WITH 'RECOMMENDATIONS' OR 'CONCLUSION' ---")
    keywords = ['RECOMMENDATION', 'CONCLUSION', 'SUMMARY']
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().upper()
        if not text:
            continue
        
        if any(kw in text for kw in keywords):
            pf = para.paragraph_format
            style_name = para.style.name if para.style else "None"
            print(f"[{i}] '{para.text.strip()[:60]}...' (style: {style_name})")
            print(f"     page_break_before: {pf.page_break_before}")
            print(f"     keep_with_next: {pf.keep_with_next}")

# Analyze both input and output documents
files = [
    ("document with new page problem input.docx", "INPUT 1"),
    ("document with new page problem output 1.docx", "OUTPUT 1"),
    ("document with new page problem input 2.docx", "INPUT 2"),
    ("document with new page problem output 2.docx", "OUTPUT 2"),
]

for filename, label in files:
    filepath = os.path.join(samples_dir, filename)
    if os.path.exists(filepath):
        analyze_document(filepath, label)
    else:
        print(f"\nFile not found: {filepath}")

print("\n" + "="*80)
print("ANALYSIS COMPLETE")
print("="*80)
