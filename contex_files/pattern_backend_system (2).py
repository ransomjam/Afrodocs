# pattern_formatter_backend.py
# Ultra-Precise Pattern-Based Academic Document Formatter
# NO AI - 100% Rule-Based - Lightning Fast

from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import re
import os
import uuid
from datetime import datetime

app = Flask(__name__)
CORS(app)

# Configuration
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


class PatternEngine:
    """Ultra-fast pattern matching engine for document analysis"""
    
    def __init__(self):
        self.patterns = self._initialize_patterns()
        
    def _initialize_patterns(self):
        """Initialize all recognition patterns"""
        return {
            # Heading Level 1 Patterns
            'heading_1': [
                re.compile(r'^([A-Z\s]{3,50})$'),  # ALL CAPS
                re.compile(r'^(CHAPTER\s+\d+.*)$', re.IGNORECASE),
                re.compile(r'^(PART\s+[IVX]+.*)$', re.IGNORECASE),
                re.compile(r'^(\d+\.\s+[A-Z][A-Z\s]+)$'),  # "1. INTRODUCTION"
                re.compile(r'^(ACKNOWLEDGEMENT|ABSTRACT|INTRODUCTION|CONCLUSION|REFERENCES|BIBLIOGRAPHY|APPENDIX)S?$', re.IGNORECASE),
            ],
            
            # Heading Level 2 Patterns
            'heading_2': [
                re.compile(r'^([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,5})$'),  # Title Case
                re.compile(r'^\d+\.\d+\s+([A-Z].{3,80})$'),  # "1.1 Background"
                re.compile(r'^([A-Z][a-z]+\s+and\s+[A-Z][a-z]+)$'),  # "Methods and Results"
            ],
            
            # Heading Level 3 Patterns
            'heading_3': [
                re.compile(r'^\d+\.\d+\.\d+\s+(.+)$'),  # "1.1.1 Details"
                re.compile(r'^([a-z]\)\s+.+)$'),  # "a) Subsection"
                re.compile(r'^([A-Z][a-z]+:)\s*$'),  # "Definition:"
            ],
            
            # Reference Patterns (APA, MLA, Chicago)
            'reference': [
                re.compile(r'^([A-Z][a-z]+,?\s+[A-Z]\..*\(\d{4}\))'),  # "Smith, J. (2024)"
                re.compile(r'^([A-Z][a-z]+\s+et\s+al\..*\d{4})'),  # "Smith et al. 2024"
                re.compile(r'^\[\d+\]'),  # "[1] Reference"
                re.compile(r'^([A-Z][a-z]+.*Retrieved from)'),  # Web reference
                re.compile(r'^([A-Z][a-z]+.*https?://)'),  # URL reference
                re.compile(r'^([A-Z][a-z]+,\s+[A-Z]\.\s+\(\d{4}\)\.)'),  # APA format
            ],
            
            # List Patterns
            'bullet_list': [
                re.compile(r'^[•●○▪▫■□]\s+(.+)$'),
                re.compile(r'^[-–—]\s+(.+)$'),
                re.compile(r'^\*\s+(.+)$'),
            ],
            
            'numbered_list': [
                re.compile(r'^(\d+[\.)]\s+.+)$'),
                re.compile(r'^([a-z][\.)]\s+.+)$'),
                re.compile(r'^([ivxlcdm]+[\.)]\s+.+)$', re.IGNORECASE),
                re.compile(r'^\(\d+\)\s+.+$'),
            ],
            
            # Table Patterns
            'table_marker': [
                re.compile(r'^\[TABLE\s+START\]', re.IGNORECASE),
                re.compile(r'^\[TABLE\s+END\]', re.IGNORECASE),
                re.compile(r'^Table\s+\d+', re.IGNORECASE),
                re.compile(r'^TABLE\s+\d+', re.IGNORECASE),
            ],
            
            'table_row': [
                re.compile(r'^\|(.+\|)+$'),  # Markdown table row
            ],
            
            # Definition/Key Term Patterns
            'definition': [
                re.compile(r'^(Definition|Objective|Task|Goal|Purpose|Aim|Method|Result|Conclusion|Note|Important|Key Point|Summary|Overview|Background):\s*(.+)?$', re.IGNORECASE),
            ],
            
            # Figure/Image Caption
            'figure': [
                re.compile(r'^Figure\s+\d+', re.IGNORECASE),
                re.compile(r'^Fig\.\s+\d+', re.IGNORECASE),
            ],
            
            # Section Keywords
            'section_abstract': re.compile(r'^(abstract|executive summary)$', re.IGNORECASE),
            'section_intro': re.compile(r'^(introduction|background|overview|motivation)$', re.IGNORECASE),
            'section_methods': re.compile(r'^(method|methodology|approach|procedure|materials)$', re.IGNORECASE),
            'section_results': re.compile(r'^(results|findings|outcomes|data)$', re.IGNORECASE),
            'section_discussion': re.compile(r'^(discussion|analysis|interpretation)$', re.IGNORECASE),
            'section_conclusion': re.compile(r'^(conclusion|summary|final remarks|future work)$', re.IGNORECASE),
            'section_references': re.compile(r'^(references|bibliography|works cited|citations)$', re.IGNORECASE),
        }
    
    def analyze_line(self, line, line_num, prev_line='', next_line='', context=None):
        """Analyze a single line with multiple pattern checks"""
        trimmed = line.strip()
        
        if not trimmed:
            return {'type': 'empty', 'content': '', 'line_num': line_num}
        
        analysis = {
            'line_num': line_num,
            'type': 'paragraph',
            'content': trimmed,
            'original': line,
            'level': 0,
            'confidence': 0.0,
        }
        
        # Get line characteristics
        length = len(trimmed)
        is_short = length < 100
        is_very_short = length < 50
        is_all_caps = trimmed == trimmed.upper() and any(c.isalpha() for c in trimmed)
        is_title_case = re.match(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*$', trimmed) is not None
        
        # Priority 1: Check for heading patterns
        if is_short:
            # H1 detection
            for pattern in self.patterns['heading_1']:
                if pattern.match(trimmed):
                    analysis['type'] = 'heading'
                    analysis['level'] = 1
                    analysis['confidence'] = 0.95
                    return analysis
            
            # H2 detection
            for pattern in self.patterns['heading_2']:
                if pattern.match(trimmed):
                    analysis['type'] = 'heading'
                    analysis['level'] = 2
                    analysis['confidence'] = 0.90
                    return analysis
            
            # H3 detection
            for pattern in self.patterns['heading_3']:
                if pattern.match(trimmed):
                    analysis['type'] = 'heading'
                    analysis['level'] = 3
                    analysis['confidence'] = 0.85
                    return analysis
            
            # Heuristic heading detection
            if is_all_caps and is_very_short:
                analysis['type'] = 'heading'
                analysis['level'] = 1
                analysis['confidence'] = 0.80
                return analysis
            
            if is_title_case and is_very_short:
                analysis['type'] = 'heading'
                analysis['level'] = 2
                analysis['confidence'] = 0.75
                return analysis
        
        # Priority 2: Check for reference patterns
        for pattern in self.patterns['reference']:
            if pattern.match(trimmed):
                analysis['type'] = 'reference'
                analysis['confidence'] = 0.90
                return analysis
        
        # Priority 3: Check for list patterns
        for pattern in self.patterns['bullet_list']:
            match = pattern.match(trimmed)
            if match:
                analysis['type'] = 'bullet_list'
                analysis['content'] = match.group(1) if match.lastindex else trimmed
                analysis['confidence'] = 0.95
                return analysis
        
        for pattern in self.patterns['numbered_list']:
            if pattern.match(trimmed):
                analysis['type'] = 'numbered_list'
                analysis['confidence'] = 0.95
                return analysis
        
        # Priority 4: Check for table patterns
        for pattern in self.patterns['table_marker']:
            if pattern.match(trimmed):
                if 'START' in trimmed.upper():
                    analysis['type'] = 'table_start'
                elif 'END' in trimmed.upper():
                    analysis['type'] = 'table_end'
                else:
                    analysis['type'] = 'table_caption'
                analysis['confidence'] = 1.0
                return analysis
        
        for pattern in self.patterns['table_row']:
            if pattern.match(trimmed):
                analysis['type'] = 'table_row'
                cells = [c.strip() for c in trimmed.split('|') if c.strip()]
                analysis['cells'] = cells
                analysis['confidence'] = 1.0
                return analysis
        
        # Priority 5: Check for definition patterns
        for pattern in self.patterns['definition']:
            match = pattern.match(trimmed)
            if match:
                analysis['type'] = 'definition'
                analysis['term'] = match.group(1)
                analysis['definition'] = match.group(2) if match.lastindex > 1 else ''
                analysis['confidence'] = 0.90
                return analysis
        
        # Priority 6: Check for figure captions
        for pattern in self.patterns['figure']:
            if pattern.match(trimmed):
                analysis['type'] = 'figure'
                analysis['confidence'] = 0.95
                return analysis
        
        # Default: paragraph
        analysis['confidence'] = 0.70
        return analysis


class DocumentProcessor:
    """Process documents using pattern engine"""
    
    def __init__(self):
        self.engine = PatternEngine()
        
    def process_docx(self, file_path):
        """Process Word document line by line"""
        doc = Document(file_path)
        
        # Extract all text with original formatting info
        lines = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                lines.append({
                    'text': text,
                    'style': para.style.name,
                    'bold': any(run.bold for run in para.runs if run.bold),
                    'font_size': para.runs[0].font.size.pt if para.runs and para.runs[0].font.size else 12,
                })
        
        return self.process_lines(lines)
    
    def process_text(self, text):
        """Process plain text"""
        lines = [{'text': line, 'style': 'Normal', 'bold': False, 'font_size': 12} 
                 for line in text.split('\n')]
        return self.process_lines(lines)
    
    def process_lines(self, lines):
        """Core line-by-line processing"""
        analyzed = []
        stats = {
            'total_lines': len(lines),
            'headings': 0,
            'paragraphs': 0,
            'references': 0,
            'tables': 0,
            'lists': 0,
            'definitions': 0,
        }
        
        # Analyze each line
        for i, line_data in enumerate(lines):
            prev_line = lines[i-1]['text'] if i > 0 else ''
            next_line = lines[i+1]['text'] if i < len(lines) - 1 else ''
            
            analysis = self.engine.analyze_line(
                line_data['text'], 
                i, 
                prev_line, 
                next_line
            )
            
            # Enhance with original formatting
            analysis['original_style'] = line_data['style']
            analysis['original_bold'] = line_data['bold']
            analysis['original_font_size'] = line_data['font_size']
            
            analyzed.append(analysis)
            
            # Update stats
            if analysis['type'] == 'heading':
                stats['headings'] += 1
            elif analysis['type'] == 'paragraph':
                stats['paragraphs'] += 1
            elif analysis['type'] == 'reference':
                stats['references'] += 1
            elif analysis['type'] in ['table_start', 'table_caption']:
                stats['tables'] += 1
            elif 'list' in analysis['type']:
                stats['lists'] += 1
            elif analysis['type'] == 'definition':
                stats['definitions'] += 1
        
        # Structure the document
        structured = self._structure_document(analyzed)
        
        return {
            'analyzed': analyzed,
            'structured': structured,
            'stats': stats,
        }
    
    def _structure_document(self, analyzed):
        """Group lines into logical sections"""
        sections = []
        current_section = None
        current_list = None
        current_table = None
        in_references = False
        
        for line in analyzed:
            if line['type'] == 'empty':
                continue
            
            # Detect reference section
            if line['type'] == 'heading' and line['level'] == 1:
                if 'reference' in line['content'].lower() or 'bibliography' in line['content'].lower():
                    in_references = True
            
            # Handle headings
            if line['type'] == 'heading':
                if current_section:
                    sections.append(current_section)
                
                current_section = {
                    'type': 'section',
                    'heading': line['content'],
                    'level': line['level'],
                    'content': [],
                }
                current_list = None
                current_table = None
                continue
            
            # Initialize first section if no heading found
            if current_section is None:
                current_section = {
                    'type': 'section',
                    'heading': 'Introduction',
                    'level': 1,
                    'content': [],
                }
            
            # Handle references
            if in_references and line['type'] == 'reference':
                current_section['content'].append({
                    'type': 'reference',
                    'text': line['content'],
                })
                continue
            
            # Handle lists
            if 'list' in line['type']:
                if not current_list or current_list['type'] != line['type']:
                    current_list = {
                        'type': line['type'],
                        'items': [],
                    }
                    current_section['content'].append(current_list)
                
                current_list['items'].append(line['content'])
                continue
            else:
                current_list = None
            
            # Handle tables
            if line['type'] == 'table_start':
                current_table = {
                    'type': 'table',
                    'caption': '',
                    'rows': [],
                }
                continue
            
            if line['type'] == 'table_caption' and current_table:
                current_table['caption'] = line['content']
                continue
            
            if line['type'] == 'table_row' and current_table:
                current_table['rows'].append(line['cells'])
                continue
            
            if line['type'] == 'table_end' and current_table:
                current_section['content'].append(current_table)
                current_table = None
                continue
            
            # Handle definitions
            if line['type'] == 'definition':
                current_section['content'].append({
                    'type': 'definition',
                    'term': line['term'],
                    'definition': line['definition'],
                })
                continue
            
            # Handle figures
            if line['type'] == 'figure':
                current_section['content'].append({
                    'type': 'figure',
                    'caption': line['content'],
                })
                continue
            
            # Handle paragraphs
            if line['type'] == 'paragraph':
                current_section['content'].append({
                    'type': 'paragraph',
                    'text': line['content'],
                })
                continue
        
        # Add final section
        if current_section:
            sections.append(current_section)
        
        return sections


class WordGenerator:
    """Generate formatted Word documents"""
    
    def __init__(self):
        self.doc = None
        
    def generate(self, structured_data, output_path):
        """Generate Word document from structured data"""
        self.doc = Document()
        self._setup_styles()
        
        # Add title if exists
        if structured_data and len(structured_data) > 0:
            first_section = structured_data[0]
            if first_section['level'] == 1:
                self._add_title(first_section['heading'])
                structured_data = structured_data[1:]  # Skip first heading
        
        # Add TOC
        self._add_toc()
        
        # Add all sections
        for section in structured_data:
            self._add_section(section)
        
        # Save document
        self.doc.save(output_path)
        return output_path
    
    def _setup_styles(self):
        """Configure document styles"""
        styles = self.doc.styles
        
        # Normal style
        normal = styles['Normal']
        normal.font.name = 'Times New Roman'
        normal.font.size = Pt(12)
        normal.paragraph_format.line_spacing = 1.5
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal.paragraph_format.space_after = Pt(6)
        
        # Heading styles
        heading_sizes = {1: 16, 2: 14, 3: 13}
        for level, size in heading_sizes.items():
            heading = styles[f'Heading {level}']
            heading.font.name = 'Times New Roman'
            heading.font.bold = True
            heading.font.size = Pt(size)
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)
    
    def _add_title(self, title_text):
        """Add document title"""
        title = self.doc.add_heading(title_text, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph()  # Spacing
    
    def _add_toc(self):
        """Add Table of Contents"""
        self.doc.add_heading('Table of Contents', level=1)
        
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        
        # Add TOC field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        self.doc.add_page_break()
    
    def _add_section(self, section):
        """Add a document section"""
        # Add heading
        self.doc.add_heading(section['heading'], level=section['level'])
        
        # Add content
        for item in section['content']:
            if item['type'] == 'paragraph':
                para = self.doc.add_paragraph(item['text'])
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            elif item['type'] == 'definition':
                para = self.doc.add_paragraph()
                run = para.add_run(f"{item['term']}: ")
                run.bold = True
                para.add_run(item['definition'])
            
            elif item['type'] == 'bullet_list':
                for list_item in item['items']:
                    self.doc.add_paragraph(list_item, style='List Bullet')
            
            elif item['type'] == 'numbered_list':
                for list_item in item['items']:
                    self.doc.add_paragraph(list_item, style='List Number')
            
            elif item['type'] == 'table':
                self._add_table(item)
            
            elif item['type'] == 'figure':
                para = self.doc.add_paragraph(item['caption'])
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.runs[0].italic = True
            
            elif item['type'] == 'reference':
                para = self.doc.add_paragraph(item['text'])
                para.paragraph_format.left_indent = Inches(0.5)
                para.paragraph_format.first_line_indent = Inches(-0.5)
    
    def _add_table(self, table_data):
        """Add a table"""
        if table_data.get('caption'):
            caption = self.doc.add_paragraph(table_data['caption'])
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.runs[0].bold = True
        
        if table_data.get('rows') and len(table_data['rows']) > 0:
            num_cols = len(table_data['rows'][0])
            table = self.doc.add_table(rows=len(table_data['rows']), cols=num_cols)
            table.style = 'Light Grid Accent 1'
            
            # Fill table
            for row_idx, row_data in enumerate(table_data['rows']):
                for col_idx, cell_text in enumerate(row_data):
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = str(cell_text)
                    
                    # Bold first row (headers)
                    if row_idx == 0:
                        cell.paragraphs[0].runs[0].bold = True
        
        self.doc.add_paragraph()  # Spacing


# Flask Routes
@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({'status': 'healthy', 'timestamp': datetime.now().isoformat()})


@app.route('/upload', methods=['POST'])
def upload_document():
    """Upload and process document"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Empty filename'}), 400
    
    # Generate unique ID
    job_id = str(uuid.uuid4())
    
    # Save uploaded file
    file_ext = os.path.splitext(file.filename)[1]
    input_path = os.path.join(UPLOAD_FOLDER, f"{job_id}{file_ext}")
    file.save(input_path)
    
    try:
        # Process document
        processor = DocumentProcessor()
        
        if file_ext == '.docx':
            result = processor.process_docx(input_path)
        else:
            # Read as text
            with open(input_path, 'r', encoding='utf-8') as f:
                text = f.read()
            result = processor.process_text(text)
        
        # Generate formatted Word document
        output_path = os.path.join(OUTPUT_FOLDER, f"{job_id}_formatted.docx")
        generator = WordGenerator()
        generator.generate(result['structured'], output_path)
        
        return jsonify({
            'job_id': job_id,
            'stats': result['stats'],
            'download_url': f'/download/{job_id}',
            'status': 'complete',
        })
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        # Clean up input file
        if os.path.exists(input_path):
            os.remove(input_path)


@app.route('/download/<job_id>', methods=['GET'])
def download_document(job_id):
    """Download formatted document"""
    output_path = os.path.join(OUTPUT_FOLDER, f"{job_id}_formatted.docx")
    
    if not os.path.exists(output_path):
        return jsonify({'error': 'Document not found'}), 404
    
    return send_file(
        output_path,
        as_attachment=True,
        download_name='formatted_document.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


if __name__ == '__main__':
    print("🚀 Pattern-Based Document Formatter Backend")
    print("=" * 50)
    print("✓ No AI dependencies")
    print("✓ Ultra-fast pattern matching")
    print("✓ 100% reliable processing")
    print("=" * 50)
    app.run(debug=True, host='0.0.0.0', port=5000)
