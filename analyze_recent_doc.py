#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Analyze most recent document with coverpage merge for positioning preservation.
"""

import os
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def analyze_recent_doc():
    """Analyze the most recent document."""
    outputs_dir = Path(__file__).parent / 'pattern-formatter' / 'backend' / 'outputs'
    
    # Get most recent docx file
    docx_files = sorted(outputs_dir.glob('*.docx'), key=os.path.getmtime, reverse=True)
    if not docx_files:
        print("No documents found")
        return
    
    recent_doc = docx_files[0]
    print(f"\nAnalyzing: {recent_doc.name}")
    print(f"Modified: {Path(recent_doc).stat().st_mtime}")
    
    doc = Document(str(recent_doc))
    
    print(f"\n{'='*80}")
    print(f"DOCUMENT ANALYSIS: {recent_doc.name}")
    print(f"{'='*80}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total sections: {len(doc.sections)}")
    print(f"Total tables: {len(doc.tables)}")
    
    # Find section break
    section_break_idx = None
    for idx, para in enumerate(doc.paragraphs):
        pPr = para._element.get_or_add_pPr()
        sectPr = pPr.find(qn('w:sectPr'))
        if sectPr is not None:
            section_break_idx = idx
            break
    
    if section_break_idx is not None:
        print(f"\n✓ Section break found at paragraph {section_break_idx}")
        print(f"  Coverpage: paragraphs 0-{section_break_idx}")
        print(f"  Body: paragraphs {section_break_idx+1}-{len(doc.paragraphs)-1}")
    else:
        print(f"\n! No section break found - document may not have merged coverpage")
    
    # Analysis table
    print(f"\n{'COVERPAGE AREA':<40} (First 15 paragraphs)")
    print(f"{'#':<4} {'Align':<12} {'LS':<8} {'Content':<50}")
    print("-" * 80)
    
    for idx in range(min(15, len(doc.paragraphs))):
        para = doc.paragraphs[idx]
        align = "None" if para.alignment is None else str(para.alignment).split('.')[-1]
        ls = para.paragraph_format.line_spacing
        text = para.text[:47].replace('\n', ' ') if para.text else "(empty)"
        print(f"{idx:<4} {align:<12} {str(ls):<8} {text:<50}")
    
    # Show body area if section break exists
    if section_break_idx and section_break_idx + 1 < len(doc.paragraphs):
        print(f"\n{'BODY AREA':<40} (Paragraphs {section_break_idx+1} to {min(section_break_idx+15, len(doc.paragraphs)-1)})")
        print(f"{'#':<4} {'Align':<12} {'LS':<8} {'Content':<50}")
        print("-" * 80)
        
        for idx in range(section_break_idx + 1, min(section_break_idx + 11, len(doc.paragraphs))):
            para = doc.paragraphs[idx]
            align = "None" if para.alignment is None else str(para.alignment).split('.')[-1]
            ls = para.paragraph_format.line_spacing
            text = para.text[:47].replace('\n', ' ') if para.text else "(empty)"
            print(f"{idx:<4} {align:<12} {str(ls):<8} {text:<50}")
    
    # Statistics
    print(f"\n{'='*80}")
    print(f"FORMATTING STATISTICS")
    print(f"{'='*80}")
    
    if section_break_idx:
        # Coverpage justification
        cover_just = sum(1 for i in range(section_break_idx) 
                        if i < len(doc.paragraphs) and doc.paragraphs[i].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY)
        cover_total = section_break_idx
        
        # Body justification
        body_just = sum(1 for i in range(section_break_idx + 1, len(doc.paragraphs))
                       if doc.paragraphs[i].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY)
        body_total = len(doc.paragraphs) - section_break_idx - 1
        
        print(f"\nCoverpage ({cover_total} paragraphs):")
        print(f"  Justified: {cover_just}/{cover_total} ({100*cover_just//max(1,cover_total)}%)")
        if cover_just == 0:
            print(f"  ✓ GOOD: Coverpage preserved (no forced justification)")
        else:
            print(f"  ✗ ISSUE: {cover_just} paragraphs were changed to JUSTIFY")
        
        print(f"\nBody ({body_total} paragraphs):")
        print(f"  Justified: {body_just}/{body_total} ({100*body_just//max(1,body_total)}%)")
        if body_just > body_total // 2:
            print(f"  ✓ GOOD: Body has consistent justification")
        else:
            print(f"  ✗ WARNING: Body has low justification rate")
        
        # Line spacing check
        cover_ls_15 = sum(1 for i in range(section_break_idx) 
                         if i < len(doc.paragraphs) and doc.paragraphs[i].paragraph_format.line_spacing == 1.5)
        body_ls_15 = sum(1 for i in range(section_break_idx + 1, len(doc.paragraphs))
                        if doc.paragraphs[i].paragraph_format.line_spacing == 1.5)
        
        print(f"\nLine spacing 1.5:")
        print(f"  Coverpage: {cover_ls_15}/{cover_total}")
        print(f"  Body: {body_ls_15}/{body_total}")
    
    else:
        # No section break found
        just_count = sum(1 for para in doc.paragraphs if para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY)
        print(f"\nTotal justified paragraphs: {just_count}/{len(doc.paragraphs)}")
        
        ls_15_count = sum(1 for para in doc.paragraphs if para.paragraph_format.line_spacing == 1.5)
        print(f"Total with LS=1.5: {ls_15_count}/{len(doc.paragraphs)}")

if __name__ == '__main__':
    analyze_recent_doc()
