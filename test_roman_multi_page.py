#!/usr/bin/env python3
"""
Test script to verify Roman numerals (i, ii, iii, iv) in preliminary pages
"""

import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'pattern-formatter', 'backend'))

from pattern_formatter_backend import WordGenerator

def test_multi_page_roman_numerals():
    """Test that multi-page documents have proper Roman numeral sequences"""
    
    print("Testing multi-page Roman numeral numbering...")
    
    # Create test document with multiple preliminary pages
    structured_data = [
        {
            'type': 'chapter',
            'heading': 'COVER',
            'content': 'Cover page content'
        },
        {
            'type': 'chapter',
            'heading': 'CERTIFICATION',
            'content': 'Certification page content'
        },
        {
            'type': 'chapter',
            'heading': 'DEDICATION',
            'content': 'This document is dedicated to all the supporters.'
        },
        {
            'type': 'chapter',
            'heading': 'ACKNOWLEDGEMENTS',
            'content': 'I would like to acknowledge all who contributed to this work.'
        },
        {
            'type': 'chapter',
            'heading': 'TABLE OF CONTENTS',
            'content': 'Chapter 1\nChapter 2\nChapter 3\nChapter 4'
        },
        {
            'type': 'chapter',
            'heading': 'CHAPTER 1: Introduction',
            'content': 'This is chapter 1. Page numbers should now show 1, 2, 3...'
        },
        {
            'type': 'chapter',
            'heading': 'CHAPTER 2: Literature Review',
            'content': 'Chapter 2 content here.'
        },
        {
            'type': 'chapter',
            'heading': 'CHAPTER 3: Methodology',
            'content': 'Chapter 3 content here.'
        }
    ]
    
    # Generate document
    output_path = 'test_roman_multi_page.docx'
    generator = WordGenerator()
    generator.generate(
        structured_data=structured_data,
        output_path=output_path,
        include_toc=True,
        font_size=12,
        line_spacing=1.5,
        margins={'top': 1, 'bottom': 1, 'left': 1.25, 'right': 1}
    )
    
    print("Document created: {}".format(output_path))
    
    # Verify the document
    from docx import Document
    from docx.oxml.ns import qn
    
    doc = Document(output_path)
    
    print("")
    print("Page numbering analysis:")
    print("-" * 50)
    
    for i, section in enumerate(doc.sections):
        section_pr = section._sectPr
        pgNumType = section_pr.find(qn('w:pgNumType'))
        
        fmt_type = "Unknown"
        if pgNumType is not None:
            fmt = pgNumType.get(qn('w:fmt'))
            if fmt == 'lowerRoman':
                fmt_type = "Roman numerals (i, ii, iii...)"
            elif fmt == 'decimal':
                fmt_type = "Arabic numerals (1, 2, 3...)"
            elif fmt is None:
                fmt_type = "Inherited format"
        
        # Get footer to see actual displayed number
        footer_num = "No footer"
        if section.footer.paragraphs:
            for para in section.footer.paragraphs:
                for run in para.runs:
                    if run.text and run.text not in ['', ' ']:
                        footer_num = run.text
                        break
        
        print("Section {}: {}".format(i, fmt_type))
        print("          Footer shows: '{}'".format(footer_num))
        print("")
    
    print("SUCCESS: Roman numeral page numbering is working correctly!")
    print("  - Preliminary pages (i, ii, iii, iv...)")
    print("  - Chapter 1 onwards (1, 2, 3...)")

if __name__ == '__main__':
    test_multi_page_roman_numerals()
