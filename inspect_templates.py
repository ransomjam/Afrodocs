"""Inspect both dissertation templates for structure and formatting issues"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

templates = {
    'Bamenda': 'pattern-formatter/Cover Pages/Cover Pages _ University of Bamenda/Dissertation Cover Page Template.docx',
    'Buea': 'pattern-formatter/Cover Pages/Cover Page _ University of Buea/Dissertation Cover Page Template.docx'
}

for uni, path in templates.items():
    print(f"\n{'='*60}")
    print(f"UNIVERSITY OF {uni.upper()} - DISSERTATION TEMPLATE")
    print(f"{'='*60}")
    
    doc = Document(path)
    
    # Inspect structure
    print(f"\nParagraphs: {len(doc.paragraphs)}")
    print(f"Sections: {len(doc.sections)}")
    print(f"Inline shapes (textboxes): {len(doc.inline_shapes)}")
    
    # Check margins
    section = doc.sections[0]
    print(f"\nMargins:")
    print(f"  Top: {section.top_margin}")
    print(f"  Bottom: {section.bottom_margin}")
    print(f"  Left: {section.left_margin}")
    print(f"  Right: {section.right_margin}")
    
    # Check font properties in first few paragraphs
    print(f"\nFirst 15 paragraphs (fonts & styles):")
    for i, para in enumerate(doc.paragraphs[:15]):
        if para.text.strip():
            # Get run info
            fonts = set()
            for run in para.runs:
                if run.font.name:
                    fonts.add(run.font.name)
                else:
                    fonts.add('DEFAULT')
            
            font_info = ', '.join(fonts) if fonts else 'No formatting'
            alignment = para.alignment
            print(f"  {i}: [{font_info}] {para.text[:60]}")
    
    # Check textboxes
    print(f"\nTextbox content:")
    if doc.inline_shapes:
        for shape_idx, shape in enumerate(doc.inline_shapes):
            print(f"  Shape {shape_idx} ({shape.type}):")
            try:
                if hasattr(shape, 'textbox'):
                    for para_idx, para in enumerate(shape.textbox.paragraphs):
                        if para.text.strip():
                            fonts = set()
                            for run in para.runs:
                                fonts.add(run.font.name or 'DEFAULT')
                            font_info = ', '.join(fonts)
                            print(f"    Para {para_idx} [{font_info}]: {para.text[:60]}")
            except:
                pass
