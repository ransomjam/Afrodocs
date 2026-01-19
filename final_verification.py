#!/usr/bin/env python3
"""
Final verification: Test improved patterns on both sample documents
"""

import re
from docx import Document

class TextFormatterWithRegex:
    """IMPROVED patterns that match actual document structure"""
    
    def __init__(self):
        self.patterns = [
            {
                'regex': r'^(\d+\.\d+)(\s+)([A-Z][A-Za-z\s\-:]*?)(?=\n|$)',
                'replacement': r'**\1\2\3**',
                'name': 'Numbered sections (1.1, 2.2, etc)',
                'flags': re.MULTILINE
            },
            {
                'regex': r'^(\d+\.\s+)([A-Z][A-Za-z\s\-:]*?)(?=\n|$)',
                'replacement': r'**\1\2**',
                'name': 'Simple numbered items (1., 2., 3.)',
                'flags': re.MULTILINE
            },
            {
                'regex': r'^([IVX]+\.\s+)([A-Z][A-Za-z\s\-:]*?)(?=\n|$)',
                'replacement': r'**\1\2**',
                'name': 'Roman numerals (I., II., III.)',
                'flags': re.MULTILINE
            },
            {
                'regex': r'^(\s*[-•]\s+)([A-Z][A-Za-z\s\-:]*?)(?=\n|$)',
                'replacement': r'\1**\2**',
                'name': 'Bulleted items',
                'flags': re.MULTILINE
            },
            {
                'regex': r'^([A-Z][A-Z\s]+):(?=\s|$)',
                'replacement': r'**\1:**',
                'name': 'Section headers (CAPITALS with colon)',
                'flags': re.MULTILINE
            },
        ]
    
    def format_text(self, text):
        if not text:
            return text
        
        for pattern_config in self.patterns:
            try:
                regex = pattern_config['regex']
                replacement = pattern_config['replacement']
                flags = pattern_config['flags']
                text = re.sub(regex, replacement, text, flags=flags)
            except Exception as e:
                continue
        
        return text


def extract_chapter_one(doc_path):
    doc = Document(doc_path)
    chapter_start = None
    chapter_end = None
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().upper()
        if 'CHAPTER' in text and ('ONE' in text or text.endswith('1')):
            chapter_start = i
        if chapter_start is not None and chapter_start != i:
            if 'CHAPTER' in text and ('TWO' in text or text.endswith('2')):
                chapter_end = i
                break
    
    if chapter_start is None:
        return ""
    
    chapter_end = chapter_end or len(doc.paragraphs)
    text_parts = []
    
    for i in range(chapter_start, chapter_end):
        text_parts.append(doc.paragraphs[i].text)
    
    return '\n'.join(text_parts)


print("="*70)
print("FINAL VERIFICATION: IMPROVED PATTERNS ON SAMPLE DOCUMENTS")
print("="*70)

formatter = TextFormatterWithRegex()

# Test Sample 2
print("\n[SAMPLE 2] bulleting and numbering sample 2.docx - CHAPTER ONE")
print("-" * 70)

doc2_path = r"c:\Users\user\Desktop\PATTERN\Samples\bulleting and numbering sample 2.docx"
chapter_text = extract_chapter_one(doc2_path)

formatted = formatter.format_text(chapter_text)

before_lines = chapter_text.split('\n')
after_lines = formatted.split('\n')

changes = []
for i, (before, after) in enumerate(zip(before_lines, after_lines)):
    if before != after and before.strip():
        changes.append((i, before, after))

print(f"\nTotal lines changed: {len(changes)}")
print("\nFirst 15 changes:")
for idx, (line_no, before, after) in enumerate(changes[:15], 1):
    print(f"\n  {idx}. Line {line_no}:")
    print(f"     Before: {before}")
    print(f"     After:  {after}")

# Count by type
bold_count = formatted.count('**') // 2
print(f"\nTotal formatted items (bold markers): {bold_count}")

print("\n" + "="*70)
print("FORMATTER EFFECTIVENESS:")
if len(changes) > 0:
    print(f"✅ SUCCESS - {len(changes)} lines formatted")
else:
    print("❌ FAILED - No formatting applied")
print("="*70)
