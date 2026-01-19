from docx import Document

doc_path = r'pattern-formatter\Cover Pages\Cover Pages _ University of Bamenda\The University of Bamenda _ Schools-Faculties-Departments.docx'
doc = Document(doc_path)

print('=== UNIVERSITY OF BAMENDA - FACULTIES, SCHOOLS AND DEPARTMENTS ===\n')

# Extract tables
for table_idx, table in enumerate(doc.tables):
    print(f'Table {table_idx + 1}:')
    for row_idx, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        print('  |  '.join(cells))
    print()

# Also print paragraphs with content
print('\n=== PARAGRAPHS ===')
for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)
