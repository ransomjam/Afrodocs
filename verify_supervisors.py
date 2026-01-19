"""Deep verify supervisor fields are in the generated document"""
import os
import sys
from docx import Document

sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'pattern-formatter', 'backend'))
from coverpage_generator import generate_cover_page

# Test data
test_data = {
    'university': 'Bamenda',
    'documentType': 'Dissertation',
    'studentName': 'John Doe',
    'studentId': 'MB26TEST001',
    'institution': 'College of Technology',
    'department': 'Electrical and Electronic Engineering',
    'faculty': 'Faculty of Engineering and Technology',
    'level': '400 Level',
    'supervisor': 'Prof. Dr. Emmanuel Tanyi',
    'academicSupervisor': 'Prof. Dr. Emmanuel Tanyi',
    'fieldSupervisor': 'Dr. Bernard Ngu',
    'title': 'Advanced Wireless Communication Systems',
    'monthYear': 'January 2026'
}

result = generate_cover_page(test_data)
doc_path = result[0] if isinstance(result, tuple) else result

if doc_path and os.path.exists(doc_path):
    print(f"Document: {doc_path}\n")
    doc = Document(doc_path)
    
    # Extract ALL text including from textboxes using direct XML inspection
    print("=== FULL DOCUMENT ANALYSIS ===\n")
    
    # Check all paragraphs
    print("PARAGRAPHS:")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"  {i}: {para.text[:80]}")
    
    # Check textboxes the right way
    print("\nTEXTBOXES (via XML):")
    if doc.element.body is not None:
        textbox_count = 0
        for txbx in doc.element.body.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}txbxContent'):
            textbox_count += 1
            for p in txbx.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
                para_text = ""
                for r in p.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
                    for t in r.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                        if t.text:
                            para_text += t.text
                if para_text.strip():
                    print(f"  TB{textbox_count}: {para_text[:80]}")
    
    # Search for supervisor names
    print("\n=== SUPERVISOR VERIFICATION ===")
    all_doc_text = " ".join([p.text for p in doc.paragraphs])
    
    # Also get textbox text
    if doc.element.body is not None:
        for txbx in doc.element.body.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}txbxContent'):
            for p in txbx.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p'):
                para_text = ""
                for r in p.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r'):
                    for t in r.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                        if t.text:
                            para_text += t.text
                all_doc_text += " " + para_text
    
    if 'Emmanuel' in all_doc_text:
        print("[FOUND] Academic Supervisor: Emmanuel Tanyi")
    else:
        print("[NOT FOUND] Academic Supervisor")
    
    if 'Bernard' in all_doc_text:
        print("[FOUND] Field Supervisor: Bernard Ngu")
    else:
        print("[NOT FOUND] Field Supervisor")
    
    # Check for leftover placeholders
    if '{{' in all_doc_text:
        print("[WARNING] Placeholder markers still present")
    else:
        print("[OK] No placeholder markers")
else:
    print(f"ERROR: Document not found at {doc_path}")
