#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Test for textbox field replacement - all document types.
"""

import os
import sys
import io
from datetime import datetime
from docx import Document

# Fix encoding for Windows console
if sys.stdout.encoding != 'utf-8':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

backend_path = r"c:\Users\user\Desktop\PATTERN\pattern-formatter\backend"
sys.path.insert(0, backend_path)

from coverpage_generator import generate_cover_page

def check_document_for_values(doc_path, expected_values):
    """Check if document contains all expected values."""
    doc = Document(doc_path)
    
    # Get all text including textboxes
    full_text = []
    
    # Paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            full_text.append(p.text)
    
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        full_text.append(p.text)
    
    # Textboxes
    from docx.oxml.ns import qn
    if doc.element.body is not None:
        for txbx in doc.element.body.iter(qn('w:txbxContent')):
            for p in txbx.iter(qn('w:p')):
                para_text = ""
                for r in p.iter(qn('w:r')):
                    for t in r.iter(qn('w:t')):
                        if t.text:
                            para_text += t.text
                if para_text.strip():
                    full_text.append(para_text)
    
    full_doc_text = "\n".join(full_text)
    
    results = {}
    for value, label in expected_values.items():
        results[label] = value in full_doc_text
    
    # Check for placeholder markers
    results['No Placeholders'] = not ("{{" in full_doc_text and "}}" in full_doc_text)
    
    return results

def test_all_document_types():
    """Test all document types with textbox fields."""
    
    print("\n" + "="*70)
    print("TEXTBOX FIELD REPLACEMENT TEST - ALL DOCUMENT TYPES")
    print("="*70)
    
    test_cases = [
        {
            "name": "Assignment",
            "data": {
                "university": "Bamenda",
                "documentType": "Assignment",
                "studentName": "TEST STUDENT 1",
                "studentId": "MB26001",
                "courseCode": "CS101",
                "courseTitle": "Test Course",
                "institution": "College of Technology",
                "faculty": "College of Technology",
                "department": "Computer Engineering",
                "level": "100 Level",
                "instructor": "DR TEST",
                "date": "2026-01-15",
                "title": "Test Assignment 1"
            },
            "expected": {
                "TEST STUDENT 1": "Student Name",
                "MB26001": "Student ID",
                "CS101": "Course Code",
                "Test Course": "Course Title",
                "Computer Engineering": "Department"
            }
        },
        {
            "name": "Dissertation",
            "data": {
                "university": "Bamenda",
                "documentType": "Dissertation",
                "studentName": "TEST STUDENT 2",
                "studentId": "MB26002",
                "institution": "Faculty of Science",
                "faculty": "Faculty of Science",
                "department": "Mathematics",
                "level": "Masters",
                "supervisor": "PROF TEST",
                "date": "2026-01-15",
                "title": "Test Dissertation"
            },
            "expected": {
                "TEST STUDENT 2": "Student Name",
                "MB26002": "Student ID",
                "Mathematics": "Department",
                "Faculty of Science": "Faculty",
                "Master of Science": "Degree"
            }
        },
        {
            "name": "Internship Report",
            "data": {
                "university": "Buea",
                "documentType": "Internship Report",
                "studentName": "TEST STUDENT 3",
                "studentId": "UB26003",
                "institution": "College of Technology (COT)",
                "faculty": "College of Technology (COT)",
                "department": "Computer Engineering",
                "level": "400 Level",
                "date": "2026-01-15",
                "title": "Test Internship Report"
            },
            "expected": {
                "TEST STUDENT 3": "Student Name",
                "UB26003": "Student ID",
                "Computer Engineering": "Department",
                "College of Technology (COT)": "Faculty"
            }
        }
    ]
    
    results = []
    
    for test_case in test_cases:
        print(f"\n{'-'*70}")
        print(f"Testing: {test_case['name']}")
        print(f"{'-'*70}")
        
        try:
            output_path, error = generate_cover_page(test_case['data'])
            
            if error:
                print(f"ERROR: {error}")
                results.append({"name": test_case['name'], "status": "FAILED", "error": error})
                continue
            
            # Check content
            values_found = check_document_for_values(output_path, test_case['expected'])
            
            all_found = all(values_found.values())
            
            # Print results
            for check, found in values_found.items():
                status = "OK" if found else "MISSING"
                print(f"  {status}: {check}")
            
            if all_found:
                print(f"✓ {test_case['name']}: PASSED")
                results.append({"name": test_case['name'], "status": "PASSED"})
            else:
                print(f"✗ {test_case['name']}: FAILED")
                results.append({"name": test_case['name'], "status": "FAILED"})
                
        except Exception as e:
            print(f"ERROR: {str(e)}")
            results.append({"name": test_case['name'], "status": "FAILED", "error": str(e)})
    
    # Summary
    print(f"\n{'='*70}")
    print("SUMMARY")
    print(f"{'='*70}")
    
    for r in results:
        status_symbol = "PASS" if r['status'] == 'PASSED' else "FAIL"
        print(f"  [{status_symbol}] {r['name']}")
    
    passed = sum(1 for r in results if r['status'] == 'PASSED')
    total = len(results)
    
    print(f"\nResult: {passed}/{total} tests passed")
    
    if passed == total:
        print("\n*** ALL TESTS PASSED - Textbox fields are working! ***\n")
        return True
    else:
        print(f"\n*** {total - passed} tests failed ***\n")
        return False

if __name__ == "__main__":
    success = test_all_document_types()
    sys.exit(0 if success else 1)
