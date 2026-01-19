#!/usr/bin/env python3
"""
Extract CHAPTER ONE content with full detail to analyze formatting needs
"""

from docx import Document

def extract_and_show_chapter_one(doc_path, name):
    print(f"\n{'='*70}")
    print(f"{name}")
    print('='*70)
    
    try:
        doc = Document(doc_path)
        
        # Find Chapter One
        chapter_start = None
        chapter_end = None
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip().upper()
            
            if 'CHAPTER' in text and ('ONE' in text or text.endswith('1')):
                chapter_start = i
                print(f"\nChapter One starts at paragraph {i}")
            
            if chapter_start is not None and chapter_start != i:
                if 'CHAPTER' in text and ('TWO' in text or text.endswith('2')):
                    chapter_end = i
                    print(f"Chapter One ends at paragraph {i}")
                    break
        
        if chapter_start is None:
            print("Chapter One not found!")
            return
        
        chapter_end = chapter_end or len(doc.paragraphs)
        
        print(f"\n\nCHAPTER ONE CONTENT (paragraphs {chapter_start} to {chapter_end}):")
        print("-" * 70)
        
        for i in range(chapter_start, min(chapter_end, chapter_start + 100)):
            para = doc.paragraphs[i]
            text = para.text
            
            if text.strip():  # Only show non-empty
                # Show paragraph detail
                indent = len(text) - len(text.lstrip())
                print(f"{i:3d} | {text}")
    
    except Exception as e:
        print(f"Error: {e}")
        import traceback
        traceback.print_exc()

# Test both documents
extract_and_show_chapter_one(r"c:\Users\user\Desktop\PATTERN\Samples\bulleting and numbering sample 2.docx", 
                            "SAMPLE 2: bulleting and numbering sample 2.docx")
