#!/usr/bin/env python
# -*- coding: utf-8 -*-

from docx import Document

docx_path = "samples/bulleting and numbering sample 2.docx"
doc = Document(docx_path)

print("Looking for ACKNOWLEDGEMENTS section...\n")

# Find where ACKNOWLEDGEMENTS section is
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    
    if 'ACKNOWLEDGEMENT' in text.upper():
        print(f"Found section heading at line {i}: {text}\n")
        
        # Show context
        print("Context (5 before, 10 after):")
        for j in range(max(0, i-5), min(i+10, len(doc.paragraphs))):
            p = doc.paragraphs[j]
            left_indent = p.paragraph_format.left_indent.inches if p.paragraph_format.left_indent else 0
            first_line_indent = p.paragraph_format.first_line_indent.inches if p.paragraph_format.first_line_indent else 0
            ptext = p.text[:40].strip()
            
            marker = " <--" if j == i else ""
            print(f"  {j}: LEFT={left_indent:.2f}, FIRST={first_line_indent:.2f} | {ptext}{marker}")
        
        print()
