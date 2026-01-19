#!/usr/bin/env python3
"""
Extract data from Word documents including tables
"""

from docx import Document
from pathlib import Path

def extract_all_data_from_docx(docx_path):
    """Extract all data including tables from Word document"""
    doc = Document(docx_path)
    
    print(f"\n{'='*80}")
    print(f"Full Analysis: {Path(docx_path).name}")
    print(f"{'='*80}")
    
    print(f"\nTotal Paragraphs: {len(doc.paragraphs)}")
    print(f"Total Tables: {len(doc.tables)}")
    
    # Print all paragraphs
    print("\n--- PARAGRAPHS ---")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"[{i}] {para.text[:100]}")
    
    # Print all tables
    if doc.tables:
        print("\n--- TABLES ---")
        for t_idx, table in enumerate(doc.tables):
            print(f"\nTable {t_idx}:")
            print(f"Rows: {len(table.rows)}, Columns: {len(table.columns)}")
            for r_idx, row in enumerate(table.rows[:10]):  # First 10 rows
                cells_text = [cell.text.strip()[:30] for cell in row.cells]
                print(f"  Row {r_idx}: {cells_text}")
            if len(table.rows) > 10:
                print(f"  ... ({len(table.rows) - 10} more rows)")

# Read BUST
bust_doc = r"c:\Users\user\Desktop\PATTERN\pattern-formatter\Cover Pages\Cover Page _ BUST\BUST _ Schools-Faculties-Departments.docx"
extract_all_data_from_docx(bust_doc)

# Read Catholic University
catholic_doc = r"c:\Users\user\Desktop\PATTERN\pattern-formatter\Cover Pages\Cover Page _ Catholic University\Catholic University Of Cameroon, Bamenda _ Schools-Faculties-Departments.docx"
extract_all_data_from_docx(catholic_doc)
