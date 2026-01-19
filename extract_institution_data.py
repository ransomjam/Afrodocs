#!/usr/bin/env python3
"""
Extract institutions data from Word documents with tables
"""

from docx import Document
import json
from pathlib import Path

def extract_institution_from_docx(docx_path, institution_id, institution_name):
    """Extract institution data from Word document with table"""
    doc = Document(docx_path)
    
    if not doc.tables:
        print(f"❌ No tables found in {Path(docx_path).name}")
        return None
    
    table = doc.tables[0]
    
    faculties = {}
    current_faculty = None
    
    # Skip header row
    for row_idx in range(1, len(table.rows)):
        row = table.rows[row_idx]
        faculty_cell = row.cells[0].text.strip()
        department_cell = row.cells[1].text.strip()
        
        # If faculty cell is not empty, it's a faculty name
        if faculty_cell and not department_cell:
            current_faculty = faculty_cell
            faculties[current_faculty] = []
        # If only department cell has content, it's a department
        elif department_cell and current_faculty:
            faculties[current_faculty].append(department_cell)
    
    # Convert to institution format
    institution = {
        "id": institution_id,
        "name": institution_name,
        "short": "".join([word[0] for word in institution_name.split()]).upper()[:3],
        "logo": f"{institution_id}_logo.png",
        "faculties": [
            {
                "name": faculty_name,
                "departments": departments
            }
            for faculty_name, departments in faculties.items()
            if faculty_name and departments  # Only include faculties with departments
        ]
    }
    
    return institution

# Extract BUST data
bust_doc = r"c:\Users\user\Desktop\PATTERN\pattern-formatter\Cover Pages\Cover Page _ BUST\BUST _ Schools-Faculties-Departments.docx"
bust_inst = extract_institution_from_docx(bust_doc, "bust", "Bamenda University of Science and Technology")

# Extract Catholic University data
catholic_doc = r"c:\Users\user\Desktop\PATTERN\pattern-formatter\Cover Pages\Cover Page _ Catholic University\Catholic University Of Cameroon, Bamenda _ Schools-Faculties-Departments.docx"
catholic_inst = extract_institution_from_docx(catholic_doc, "cucb", "The Catholic University of Cameroon, Bamenda")

print(f"\n{'='*80}")
print("BUST Institution Data")
print(f"{'='*80}")
if bust_inst:
    print(f"ID: {bust_inst['id']}")
    print(f"Name: {bust_inst['name']}")
    print(f"Faculties: {len(bust_inst['faculties'])}")
    total_depts = sum(len(f['departments']) for f in bust_inst['faculties'])
    print(f"Total Departments: {total_depts}")
    print("\nFaculties:")
    for fac in bust_inst['faculties']:
        print(f"  - {fac['name']}: {len(fac['departments'])} departments")

print(f"\n{'='*80}")
print("Catholic University Institution Data")
print(f"{'='*80}")
if catholic_inst:
    print(f"ID: {catholic_inst['id']}")
    print(f"Name: {catholic_inst['name']}")
    print(f"Faculties: {len(catholic_inst['faculties'])}")
    total_depts = sum(len(f['departments']) for f in catholic_inst['faculties'])
    print(f"Total Departments: {total_depts}")
    print("\nFaculties:")
    for fac in catholic_inst['faculties']:
        print(f"  - {fac['name']}: {len(fac['departments'])} departments")

# Save as JSON for verification
extracted_data = {
    "bust": bust_inst,
    "catholic": catholic_inst
}

with open('new_institutions_data.json', 'w', encoding='utf-8') as f:
    json.dump(extracted_data, f, ensure_ascii=False, indent=2)

print(f"\n✅ Extracted data saved to new_institutions_data.json")

# Print JSON for reference
print(f"\n{'='*80}")
print("JSON Output for BUST (first 50 lines):")
print(f"{'='*80}")
if bust_inst:
    bust_json = json.dumps(bust_inst, ensure_ascii=False, indent=2)
    print("\n".join(bust_json.split("\n")[:50]))
