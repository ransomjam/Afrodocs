#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
TEST: Conservative Bullet Rendering

Tests the new bullet restrictions:
1. Bullets only used if 4+ items in section
2. Bullets only used if all items are short (<30 words)
3. No bullets if any item has colon (would be bolded)
4. No bullets if any item is multiline
5. No bullets on bolded text
"""

import sys
sys.path.insert(0, 'pattern-formatter/backend')

from pattern_formatter_backend import DocumentProcessor, WordGenerator
import tempfile
from docx import Document

# Test 1: Section with 3 items (should NOT use bullets)
print("="*80)
print("TEST 1: Section with 3 items (should convert to bold)")
print("="*80)

test1_content = """# Introduction

- Item 1
- Item 2
- Item 3
"""

processor = DocumentProcessor()
sections1, _ = processor.process_text(test1_content)

# Check if bullet_list type is used
has_bullets = any(
    any(item.get('type') == 'bullet_list' for item in section.get('content', []))
    for section in sections1.values()
)

print(f"Has bullet_list: {has_bullets}")
print(f"Expected: False (only 3 items, need 4+)")
print(f"Result: {'✅ PASS' if not has_bullets else '❌ FAIL'}\n")


# Test 2: Section with 4 short items (should use bullets)
print("="*80)
print("TEST 2: Section with 4 short items (should use bullets)")
print("="*80)

test2_content = """# Features

- Quick
- Easy
- Simple
- Fast
"""

sections2, _ = processor.process_text(test2_content)

has_bullets2 = any(
    any(item.get('type') == 'bullet_list' for item in section.get('content', []))
    for section in sections2.values()
)

print(f"Has bullet_list: {has_bullets2}")
print(f"Expected: True (4 items, all short)")
print(f"Result: {'✅ PASS' if has_bullets2 else '❌ FAIL'}\n")


# Test 3: Section with 4 items but one has colon (should NOT use bullets)
print("="*80)
print("TEST 3: Section with 4 items but one has colon (should use bold instead)")
print("="*80)

test3_content = """# Options

- Option A
- Option B: Description of this option
- Option C
- Option D
"""

sections3, _ = processor.process_text(test3_content)

has_bullets3 = any(
    any(item.get('type') == 'bullet_list' for item in section.get('content', []))
    for section in sections3.values()
)

print(f"Has bullet_list: {has_bullets3}")
print(f"Expected: False (one item has colon, would be bolded)")
print(f"Result: {'✅ PASS' if not has_bullets3 else '❌ FAIL'}\n")


# Test 4: Section with 4 long items (should NOT use bullets)
print("="*80)
print("TEST 4: Section with 4 items, some long (should use bold instead)")
print("="*80)

test4_content = """# Requirements

- This is a short one
- This is another requirement that is much longer with many words to exceed the 30 word threshold we have set
- A short one
- Another requirement that is quite lengthy and contains substantial information that requires more detailed explanation
"""

sections4, _ = processor.process_text(test4_content)

has_bullets4 = any(
    any(item.get('type') == 'bullet_list' for item in section.get('content', []))
    for section in sections4.values()
)

print(f"Has bullet_list: {has_bullets4}")
print(f"Expected: False (items exceed 30 words)")
print(f"Result: {'✅ PASS' if not has_bullets4 else '❌ FAIL'}\n")


# Test 5: Generate actual document and verify formatting
print("="*80)
print("TEST 5: Generate actual Word document and verify no inappropriate bullets")
print("="*80)

test5_content = """# Section A

- Short
- Items
- Only

# Section B

- This is much longer and would exceed the word limit we have set
- Another long one with lots of content

# Section C

- Four short ones
- That can be
- Represented as
- Bullet points
"""

sections5, _ = processor.process_text(test5_content)

# Generate Word document
generator = WordGenerator()
with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
    tmp_path = tmp.name

try:
    generator.generate(
        sections5,
        tmp_path,
        font_size=12,
        line_spacing=1.5
    )
    
    # Read back and check
    doc = Document(tmp_path)
    
    print(f"Generated paragraphs: {len(doc.paragraphs)}")
    
    # Count bullet paragraphs
    bullet_count = 0
    for para in doc.paragraphs:
        if para.text.strip().startswith('■') or para.text.strip().startswith('•'):
            bullet_count += 1
    
    print(f"Paragraphs with bullets: {bullet_count}")
    print(f"Expected: 4 (only Section C should have bullets)")
    
    if bullet_count == 4:
        print(f"Result: ✅ PASS")
    else:
        print(f"Result: ⚠️  Generated {bullet_count} bullets instead of 4")
        print("\nBullet items found:")
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text.startswith('■') or text.startswith('•'):
                print(f"  {i}: {text[:60]}")
    
    import os
    os.unlink(tmp_path)

except Exception as e:
    print(f"❌ Error generating document: {e}")
    import traceback
    traceback.print_exc()

print("\n" + "="*80)
print("TEST COMPLETE")
print("="*80)
