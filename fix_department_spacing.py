#!/usr/bin/env python3
"""
Fix DEPARTMENT OF spacing in all cover page templates
Ensure there's a space between "DEPARTMENT OF" and the placeholder
"""

from docx import Document
import os

templates_dir = r'c:\Users\user\Desktop\Afrodocs_dev\Afrodocs\pattern-formatter\Cover Pages'
count = 0

# Walk through all templates
for root, dirs, files in os.walk(templates_dir):
    for file in files:
        if 'Template.docx' in file and not file.startswith('sample'):
            path = os.path.join(root, file)
            doc = Document(path)
            modified = False
            
            # Check each paragraph
            for para in doc.paragraphs:
                if 'DEPARTMENT OF{{DEPARTMENT}}' in para.text:
                    print(f'Found issue in: {os.path.basename(root)} / {file}')
                    print(f'  Before: {repr(para.text)}')
                    
                    # Get formatting info from existing runs
                    original_format = {}
                    if para.runs:
                        orig_run = para.runs[0]
                        original_format = {
                            'bold': orig_run.bold,
                            'italic': orig_run.italic,
                            'font_name': orig_run.font.name,
                            'font_size': orig_run.font.size,
                        }
                    
                    # Clear paragraph
                    for run in para.runs:
                        run._element.getparent().remove(run._element)
                    
                    # Add corrected text with proper spacing
                    new_run = para.add_run('DEPARTMENT OF {{DEPARTMENT}}')
                    
                    # Apply original formatting
                    new_run.bold = original_format.get('bold', False)
                    new_run.italic = original_format.get('italic', False)
                    if original_format.get('font_name'):
                        new_run.font.name = original_format['font_name']
                    if original_format.get('font_size'):
                        new_run.font.size = original_format['font_size']
                    
                    print(f'  After: {repr(para.text)}')
                    modified = True
                    count += 1
            
            if modified:
                doc.save(path)
                print(f'  ✓ Saved')

print(f'\n✓ Total templates fixed: {count}')
