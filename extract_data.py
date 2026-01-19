#!/usr/bin/env python
"""
Extract faculties and departments from University Word documents
Standalone script - doesn't import Flask
"""

import os
import sys
from docx import Document

def extract_from_document(doc_path):
    """Extract schools/faculties and departments from document"""
    print(f"\n{'='*70}")
    print(f"File: {os.path.basename(doc_path)}")
    print('='*70)
    
    try:
        doc = Document(doc_path)
    except Exception as e:
        print(f"Error opening document: {e}")
        return []
    
    # Extract from paragraphs and tables
    content = []
    
    print("\nParagraphs:")
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            print(f"  {text}")
            content.append(text)
    
    print(f"\nTables found: {len(doc.tables)}")
    for t_idx, table in enumerate(doc.tables):
        print(f"\nTable {t_idx + 1}:")
        for r_idx, row in enumerate(table.rows):
            row_texts = [cell.text.strip() for cell in row.cells]
            print(f"  Row {r_idx}: {row_texts}")
            for cell_text in row_texts:
                if cell_text and cell_text not in content:
                    content.append(cell_text)
    
    return content

if __name__ == '__main__':
    bamenda_doc = r"c:\Users\user\Desktop\PATTERN\pattern-formatter\Cover Pages\Cover Pages _ University of Bamenda\The University of Bamenda _ Schools-Faculties-Departments.docx"
    buea_doc = r"c:\Users\user\Desktop\PATTERN\pattern-formatter\Cover Pages\Cover Page _ University of Buea\University of Buea _ Schools-Faculties-Departments.docx"
    
    print("\n" + "="*70)
    print("EXTRACTING UNIVERSITY DATA")
    print("="*70)
    
    bamenda_data = extract_from_document(bamenda_doc)
    buea_data = extract_from_document(buea_doc)
    
    print("\n" + "="*70)
    print("SUMMARY")
    print("="*70)
    print(f"Bamenda entries: {len(bamenda_data)}")
    print(f"Buea entries: {len(buea_data)}")
