#!/usr/bin/env python3
"""
Find numbered/bulleted content in sample documents
"""

from docx import Document

def analyze_document(doc_path, name):
    print(f"\n{'='*70}")
    print(f"{name}")
    print('='*70)
    
    try:
        doc = Document(doc_path)
        
        print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
        print("\nLooking for numbered/bulleted content...")
        print("-" * 70)
        
        found_content = False
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            
            # Look for numbered items
            if text and (text[0].isdigit() or text.startswith('-') or text.startswith('•')):
                print(f"Para {i}: {text}")
                found_content = True
            
            # Look for section headers with numbers
            if text and any(x in text.lower() for x in ['chapter', 'section', 'introduction', 'methodology']):
                print(f"Para {i}: {text}")
                found_content = True
        
        if not found_content:
            print("\nNo obvious numbered/bulleted content found. Showing all content:")
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if text:
                    print(f"Para {i}: {text[:100]}")
                    
    except Exception as e:
        print(f"Error: {e}")

# Test both documents
analyze_document(r"c:\Users\user\Desktop\PATTERN\Samples\numbering and bulleting sample 1.docx", 
                 "SAMPLE 1: numbering and bulleting sample 1.docx")

analyze_document(r"c:\Users\user\Desktop\PATTERN\Samples\bulleting and numbering sample 2.docx", 
                 "SAMPLE 2: bulleting and numbering sample 2.docx")
