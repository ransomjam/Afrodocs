#!/usr/bin/env python3
"""
Extract all text from Word documents to see structure
"""

from docx import Document
from pathlib import Path

def read_full_docx(docx_path):
    """Read all content from Word document"""
    doc = Document(docx_path)
    
    print(f"\n{'='*80}")
    print(f"Content of: {Path(docx_path).name}")
    print(f"{'='*80}\n")
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        style = para.style.name if para.style else "None"
        # Show indent level
        indent = len(text) - len(text.lstrip())
        print(f"[{i:3d}] Indent:{indent:2d} | Style:{style:15s} | {text}")
    
    print(f"\n{'='*80}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"{'='*80}\n")

# Read BUST
bust_doc = r"c:\Users\user\Desktop\PATTERN\pattern-formatter\Cover Pages\Cover Page _ BUST\BUST _ Schools-Faculties-Departments.docx"
read_full_docx(bust_doc)

# Read Catholic University
catholic_doc = r"c:\Users\user\Desktop\PATTERN\pattern-formatter\Cover Pages\Cover Page _ Catholic University\Catholic University Of Cameroon, Bamenda _ Schools-Faculties-Departments.docx"
read_full_docx(catholic_doc)
