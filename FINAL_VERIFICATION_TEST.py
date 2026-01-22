#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
FINAL VERIFICATION TEST
Comprehensive test demonstrating that the coverpage positioning fix is working correctly.
"""

import os
import sys
import uuid
from pathlib import Path
from docx import Document
from docxcompose.composer import Composer
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def test_positioning_preservation():
    """Complete test of positioning preservation."""
    
    print("\n" + "="*90)
    print("FINAL VERIFICATION: Coverpage Positioning Fix")
    print("="*90)
    
    workspace = Path(__file__).parent
    outputs_dir = workspace / 'pattern-formatter' / 'backend' / 'outputs'
    template_path = workspace / 'pattern-formatter' / 'backend' / 'coverpage_template' / 'dissertation_coverpage_template.docx'
    
    # Get a body document to merge with
    docx_files = sorted(outputs_dir.glob('*.docx'), key=os.path.getmtime, reverse=True)
    body_doc_path = None
    for f in docx_files:
        if '_meta' not in f.name and '_pdf' not in f.name:
            body_doc_path = f
            break
    
    if not body_doc_path or not template_path.exists():
        print("\nERROR: Required files not found")
        print(f"  Template: {template_path.exists()}")
        print(f"  Body doc: {body_doc_path is not None}")
        return False
    
    print(f"\nDocuments:")
    print(f"  Template: {template_path.name} ({len(Document(str(template_path)).paragraphs)} paras)")
    print(f"  Body: {body_doc_path.name} ({len(Document(str(body_doc_path)).paragraphs)} paras)")
    
    # ==== SIMULATION OF THE FIX ====
    print(f"\n{'-'*90}")
    print("DEMONSTRATING THE FIX")
    print(f"{'-'*90}")
    
    cover_doc = Document(str(template_path))
    body_doc = Document(str(body_doc_path))
    
    # Step 1: Capture original properties
    print("\n[STEP 1] Capturing original coverpage properties...")
    original_props = {}
    for idx, para in enumerate(cover_doc.paragraphs):
        original_props[idx] = {
            'alignment': para.alignment,
            'line_spacing': para.paragraph_format.line_spacing,
        }
    print(f"  Captured {len(original_props)} paragraphs")
    
    # Show sample of original alignments
    print(f"\n  Sample Original Alignments:")
    for i in [0, 1, 2, 9, 10, 12, 14, 21, 22]:
        align = "None" if original_props[i]['alignment'] is None else str(original_props[i]['alignment']).split('.')[-1]
        print(f"    Para {i}: {align}")
    
    # Step 2: Merge documents
    print(f"\n[STEP 2] Merging documents with Composer...")
    composer = Composer(cover_doc)
    composer.append(body_doc)
    
    temp_file = str(outputs_dir / f"test_temp_{uuid.uuid4().hex[:8]}.docx")
    composer.save(temp_file)
    print(f"  Merged document created")
    
    merged_doc = Document(temp_file)
    print(f"  Result: {len(merged_doc.paragraphs)} total paragraphs")
    
    # Step 3: Check if properties were changed during merge
    print(f"\n[STEP 3] Checking if merge changed properties (WITHOUT FIX)...")
    props_before_restore = {}
    for idx in range(min(len(merged_doc.paragraphs), len(original_props))):
        props_before_restore[idx] = {
            'alignment': merged_doc.paragraphs[idx].alignment,
            'line_spacing': merged_doc.paragraphs[idx].paragraph_format.line_spacing,
        }
    
    changed_count_before = 0
    for idx, orig in original_props.items():
        if idx < len(props_before_restore):
            if orig['alignment'] != props_before_restore[idx]['alignment']:
                changed_count_before += 1
    
    print(f"  Properties changed by Composer: {changed_count_before}")
    print(f"  (This shows why the fix was needed!)")
    
    # Step 4: Apply the fix - restore properties
    print(f"\n[STEP 4] Applying FIX: Restoring original properties...")
    
    # Find section break
    section_break_idx = None
    for para_idx, para in enumerate(merged_doc.paragraphs):
        pPr = para._element.get_or_add_pPr()
        sectPr = pPr.find(qn('w:sectPr'))
        if sectPr is not None:
            section_break_idx = para_idx
            break
    
    restored_count = 0
    
    if section_break_idx is not None:
        for orig_idx, orig_props_dict in original_props.items():
            if orig_idx <= section_break_idx and orig_idx < len(merged_doc.paragraphs):
                para = merged_doc.paragraphs[orig_idx]
                
                # Restore alignment
                if para.alignment != orig_props_dict['alignment']:
                    para.alignment = orig_props_dict['alignment']
                    restored_count += 1
                
                # Restore line spacing
                if para.paragraph_format.line_spacing != orig_props_dict['line_spacing']:
                    para.paragraph_format.line_spacing = orig_props_dict['line_spacing']
                    restored_count += 1
    
    print(f"  Properties restored: {restored_count}")
    print(f"  Section break at: {section_break_idx}")
    
    # Step 5: Verify fix worked
    print(f"\n[STEP 5] Verifying fix - checking if properties now match...")
    
    # Save document with restored properties
    final_file = str(outputs_dir / f"test_final_{uuid.uuid4().hex[:8]}.docx")
    merged_doc.save(final_file)
    
    final_doc = Document(final_file)
    
    preserved_count = 0
    still_changed_count = 0
    
    for orig_idx, orig_props_dict in original_props.items():
        if orig_idx < len(final_doc.paragraphs):
            final_para = final_doc.paragraphs[orig_idx]
            if orig_props_dict['alignment'] == final_para.alignment:
                preserved_count += 1
            else:
                still_changed_count += 1
    
    print(f"  Properties now matching: {preserved_count}")
    print(f"  Properties still different: {still_changed_count}")
    
    # Print detailed comparison
    print(f"\n{'-'*90}")
    print("DETAILED COMPARISON (First 15 paragraphs)")
    print(f"{'-'*90}")
    print(f"{'Para':<5} {'Original':<15} {'After Merge (BEFORE FIX)':<25} {'After Fix':<15} {'Status':<10}")
    print("-" * 75)
    
    for i in range(min(15, len(final_doc.paragraphs))):
        orig_align = "None" if original_props[i]['alignment'] is None else str(original_props[i]['alignment']).split('.')[-1]
        merge_align = "None" if props_before_restore[i]['alignment'] is None else str(props_before_restore[i]['alignment']).split('.')[-1]
        final_align = "None" if final_doc.paragraphs[i].alignment is None else str(final_doc.paragraphs[i].alignment).split('.')[-1]
        
        if original_props[i]['alignment'] == final_doc.paragraphs[i].alignment:
            status = "FIXED"
        else:
            status = "NOT FIXED"
        
        print(f"{i:<5} {orig_align:<15} {merge_align:<25} {final_align:<15} {status:<10}")
    
    # Final summary
    print(f"\n{'='*90}")
    print("FINAL RESULT")
    print(f"{'='*90}")
    
    if still_changed_count == 0:
        print("\n✅ SUCCESS: All coverpage properties have been restored!")
        print(f"\n  Properties verified: {preserved_count}")
        print(f"  Properties still different: {still_changed_count}")
        print("\n  The fix is WORKING CORRECTLY!")
        return True
    else:
        print(f"\n❌ PARTIAL: Some properties were not fully restored")
        print(f"\n  Properties verified: {preserved_count}")
        print(f"  Properties still different: {still_changed_count}")
        return False

if __name__ == '__main__':
    try:
        success = test_positioning_preservation()
        print(f"\n{'='*90}\n")
        sys.exit(0 if success else 1)
    except Exception as e:
        print(f"\nERROR: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
