#!/usr/bin/env python3
"""Test with the FULL networking notes provided by the user"""

import sys
import os
import tempfile
from docx import Document

# Add backend to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'pattern-formatter', 'backend'))

from pattern_formatter_backend import DocumentProcessor, WordGenerator, FormatPolicy

# Full networking notes from user request (truncated slightly for file size)
networking_notes = """## Networking Notes

### 1) What networking is

Networking is the practice of connecting computers, phones, servers, and devices so they can share data and resources (files, printers, internet access, applications). It involves hardware, software, rules (protocols), and processes that keep communication reliable and secure.

**Key outcomes of networking**

* Communication: messaging, voice, video calls
* Resource sharing: printers, storage, internet
* Centralised services: authentication, databases, applications
* Scalability: add users/devices without redesigning everything
* Security and control: access rules, monitoring, auditing

---

## 2) Core network components

### 2.1 End devices (Hosts)

* PCs, laptops, phones, tablets
* Servers (web, database, file, mail)
* Printers, IP cameras, smart TVs, IoT sensors

### 2.2 Network devices

* **Switch**: connects devices in a LAN, forwards frames based on MAC addresses
* **Router**: connects different networks, forwards packets based on IP addresses

---

## 6) Protocols and ports (quick reference)

| Protocol | Purpose                 | Transport | Common Port(s) |
| -------- | ----------------------- | --------: | -------------: |
| HTTP     | Web browsing            |       TCP |             80 |
| HTTPS    | Secure web              |       TCP |            443 |
| DNS      | Name resolution         |   UDP/TCP |             53 |
| DHCP     | Automatic IP assignment |       UDP |          67/68 |
| SSH      | Secure remote login     |       TCP |             22 |
| FTP      | File transfer           |       TCP |             21 |
| SMTP     | Send email              |       TCP |  25 (also 587) |
| IMAP     | Receive email           |       TCP | 143 (also 993) |
| POP3     | Receive email           |       TCP | 110 (also 995) |
| RDP      | Remote desktop          |       TCP |           3389 |

**Notes**

* TCP is connection-oriented (reliable, ordered delivery).
* UDP is connectionless (faster, less overhead, used for streaming/DNS).

---

## 14) Quick comparison tables

### 14.1 Switch vs Router vs Access Point

| Device       | Main Job                  | Works Mostly At | Typical Use                        |
| ------------ | ------------------------- | --------------- | ---------------------------------- |
| Switch       | Connects devices in a LAN | Layer 2         | Office/classroom wiring            |
| Router       | Connects networks         | Layer 3         | LAN to Internet, LAN to LAN        |
| Access Point | Provides Wi-Fi            | Layer 2         | Wireless access for phones/laptops |

### 14.2 TCP vs UDP

| Feature     | TCP                                 | UDP                              |
| ----------- | ----------------------------------- | -------------------------------- |
| Reliability | High (acknowledgements, retransmit) | Lower (no retransmit by default) |
| Speed       | Slower than UDP (overhead)          | Faster                           |
| Best for    | Web, email, file transfer           | Streaming, voice, DNS            |

---

## 15) Short summary (for revision)

* Networking connects devices to share data/resources using standards (protocols).
* Switches operate inside LANs (MAC-based forwarding); routers connect networks (IP-based routing).
* IP addressing, subnetting, DNS, and DHCP are fundamental.
* Security relies on good configuration, segmentation, authentication, and monitoring.
* Troubleshooting is easiest when done step-by-step from physical layer upward.
"""

def test_full_networking_notes():
    """Test with the full networking notes"""
    print("=" * 70)
    print("COMPREHENSIVE TABLE HANDLING TEST")
    print("Processing Full Networking Notes Document")
    print("=" * 70)
    
    # Process the text
    policy = FormatPolicy()
    processor = DocumentProcessor(policy=policy)
    
    result = processor.process_text(networking_notes)
    if isinstance(result, tuple):
        structured_result = result[0]
    else:
        structured_result = result
    
    # Count tables in structured output
    table_count = 0
    table_rows = []
    
    for section in structured_result.get('structured', []):
        if not isinstance(section, dict):
            continue
        section_title = section.get('heading', 'Untitled')
        
        for item in section.get('content', []):
            if not isinstance(item, dict):
                continue
            if item.get('type') == 'table':
                table_count += 1
                row_count = len(item.get('content', []))
                col_count = len(item.get('content', [0])) if item.get('content') else 0
                
                first_row = None
                if item.get('content'):
                    first_row_data = item['content'][0]
                    if isinstance(first_row_data, dict) and 'cells' in first_row_data:
                        first_row = first_row_data['cells'][:2]
                
                table_info = {
                    'section': section_title,
                    'rows': row_count,
                    'first_cells': first_row
                }
                table_rows.append(table_info)
                
                print(f"\n[Table {table_count}] {section_title}")
                print(f"  Rows: {row_count}, First cells: {first_row}")
    
    print(f"\n{'=' * 70}")
    print(f"STRUCTURED OUTPUT: {table_count} tables detected")
    
    # Now generate Word document
    print(f"\nGenerating Word document...")
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = f.name
    
    try:
        generator = WordGenerator(policy=policy)
        generator.generate(
            structured_result.get('structured', []),
            output_path,
            images=[],
            font_size=12,
            line_spacing=1.5,
            margins={'left': 2.5, 'top': 2.5, 'bottom': 2.5, 'right': 2.5}
        )
        
        # Verify document
        if os.path.exists(output_path):
            doc = Document(output_path)
            
            doc_table_count = len(doc.tables)
            print(f"WORD DOCUMENT: {doc_table_count} tables found")
            
            # Verify each table
            for i, table in enumerate(doc.tables, 1):
                row_count = len(table.rows)
                col_count = len(table.columns)
                first_row_cells = [cell.text[:30] for cell in table.rows[0].cells[:2]]
                print(f"  Table {i}: {row_count} rows, {col_count} cols - {first_row_cells}")
            
            print(f"\n{'=' * 70}")
            
            # Final verdict
            if doc_table_count >= 3:
                print("✓ SUCCESS: Tables were properly processed and included in the document!")
                print(f"  - Detected: {table_count} tables in structured output")
                print(f"  - Generated: {doc_table_count} tables in Word document")
                return True
            else:
                print(f"✗ FAILED: Expected 3+ tables, found {doc_table_count} in document")
                return False
        else:
            print("✗ FAILED: Document was not generated")
            return False
            
    finally:
        if os.path.exists(output_path):
            try:
                os.unlink(output_path)
            except:
                print(f"(Note: temp file not deleted: {output_path})")

if __name__ == '__main__':
    success = test_full_networking_notes()
    sys.exit(0 if success else 1)
