#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Check the original coverpage template to see what alignments should be preserved.
"""

from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def check_template():
    """Check template alignments."""
    template_path = Path(__file__).parent / 'pattern-formatter' / 'backend' / 'coverpage_template' / 'dissertation_coverpage_template.docx'
    
    if not template_path.exists():
        print(f"Template not found: {template_path}")
        return
    
    doc = Document(str(template_path))
    
    print(f"\nOriginal Template: {template_path.name}\n")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total sections: {len(doc.sections)}\n")
    
    print(f"{'#':<4} {'Align':<12} {'Style':<25} {'LS':<6} {'Content':<40}")
    print("-" * 90)
    
    for idx in range(min(25, len(doc.paragraphs))):
        para = doc.paragraphs[idx]
        align = "None" if para.alignment is None else str(para.alignment).split('.')[-1]
        style = para.style.name
        ls = para.paragraph_format.line_spacing
        text = para.text[:37].replace('\n', ' ') if para.text else "(empty)"
        
        print(f"{idx:<4} {align:<12} {style:<25} {str(ls):<6} {text:<40}")

if __name__ == '__main__':
    check_template()
