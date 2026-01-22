#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Test to verify that coverpage positioning is properly preserved after merge.
"""

import sys
import os
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def compare_alignments(template_path, merged_path, max_paras=25):
    """Compare template and merged document alignments."""
    template_doc = Document(str(template_path))
    merged_doc = Document(str(merged_path))
    
    print(f"\n{'='*90}")
    print(f"COMPARING: Template vs Merged Document")
    print(f"{'='*90}")
    
    # Find section break in merged doc
    section_break_idx = None
    for para_idx, para in enumerate(merged_doc.paragraphs):
        pPr = para._element.get_or_add_pPr()
        sectPr = pPr.find(qn('w:sectPr'))
        if sectPr is not None:
            section_break_idx = para_idx
            break
    
    print(f"\nTemplate: {len(template_doc.paragraphs)} paragraphs, 1 section")
    print(f"Merged: {len(merged_doc.paragraphs)} paragraphs, {len(merged_doc.sections)} sections")
    if section_break_idx is not None:
        print(f"Section break at paragraph {section_break_idx}")
    
    # Limit to coverpage size (min of template and section break)
    compare_limit = min(max_paras, len(template_doc.paragraphs))
    if section_break_idx is not None:
        compare_limit = min(compare_limit, section_break_idx + 1)
    
    print(f"\n{'Para':<5} {'Template Align':<15} {'Template LS':<10} {'Merged Align':<15} {'Merged LS':<10} {'Status':<10}")
    print("-" * 90)
    
    matches = 0
    mismatches = 0
    
    for idx in range(compare_limit):
        t_para = template_doc.paragraphs[idx]
        m_para = merged_doc.paragraphs[idx] if idx < len(merged_doc.paragraphs) else None
        
        if m_para is None:
            print(f"{idx:<5} (not in merged document)")
            continue
        
        t_align = "None" if t_para.alignment is None else str(t_para.alignment).split('.')[-1]
        t_ls = t_para.paragraph_format.line_spacing or "default"
        
        m_align = "None" if m_para.alignment is None else str(m_para.alignment).split('.')[-1]
        m_ls = m_para.paragraph_format.line_spacing or "default"
        
        # Check if they match
        align_match = t_para.alignment == m_para.alignment
        ls_match = t_para.paragraph_format.line_spacing == m_para.paragraph_format.line_spacing
        
        status = "✓ OK" if (align_match and ls_match) else "✗ DIFF"
        
        if align_match and ls_match:
            matches += 1
        else:
            mismatches += 1
        
        print(f"{idx:<5} {t_align:<15} {str(t_ls):<10} {m_align:<15} {str(m_ls):<10} {status:<10}")
    
    print(f"\n{'='*90}")
    print(f"RESULT: {matches} matching, {mismatches} mismatched")
    
    if mismatches == 0:
        print(f"✓ SUCCESS: All coverpage properties preserved!")
    else:
        print(f"✗ ISSUE: {mismatches} coverpage properties were modified during merge")
    
    return mismatches == 0

def main():
    """Main test."""
    workspace = Path(__file__).parent
    
    # Get template
    template_path = workspace / 'pattern-formatter' / 'backend' / 'coverpage_template' / 'dissertation_coverpage_template.docx'
    
    # Get most recent merged document
    outputs_dir = workspace / 'pattern-formatter' / 'backend' / 'outputs'
    docx_files = sorted(outputs_dir.glob('*.docx'), key=os.path.getmtime, reverse=True)
    
    if not docx_files:
        print("No documents found in outputs folder")
        return False
    
    recent_doc = docx_files[0]
    
    print(f"\nTesting document: {recent_doc.name}")
    
    if not template_path.exists():
        print(f"Template not found: {template_path}")
        return False
    
    success = compare_alignments(str(template_path), str(recent_doc))
    
    return success

if __name__ == '__main__':
    success = main()
    sys.exit(0 if success else 1)
