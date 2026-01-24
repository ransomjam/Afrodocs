"""Test script to verify the keep_with_next/keep_together fix for page breaks"""
import os
import sys

# Add backend to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'pattern-formatter', 'backend'))

from docx import Document
import io

# Use the pattern formatter
from pattern_formatter_backend import WordGenerator, DocumentProcessor

# Test with the problematic input document
input_file = "samples/new page problems/document with new page problem input.docx"
output_file = "test_pagebreak_fix_output.docx"

print("="*80)
print("TESTING PAGE BREAK FIX")
print("="*80)

# Read the input document
print(f"\n1. Reading input document: {input_file}")
with open(input_file, 'rb') as f:
    file_content = f.read()

# Extract text from input
print("\n2. Extracting text from input document...")
input_doc = Document(io.BytesIO(file_content))
text_lines = []
for para in input_doc.paragraphs:
    text = para.text.strip()
    if text:
        text_lines.append(text)

input_text = '\n'.join(text_lines)

# Create processor and generator
print("\n3. Processing document with formatter...")
processor = DocumentProcessor()
generator = WordGenerator(
    template_path="pattern-formatter/assets/uniba_template.docx",
    output_path=output_file
)

# Parse and format
analyzed_lines = processor.analyze_document(input_text)
sections = generator.structure_sections(analyzed_lines)
generator.generate(sections)

print(f"\n3. Output saved to: {output_file}")

# Now analyze the output
print("\n4. Analyzing output document for page break issues...")
output_doc = Document(output_file)

# Check heading styles
print("\n--- HEADING STYLES IN OUTPUT ---")
for i in range(1, 5):
    try:
        style = output_doc.styles[f'Heading {i}']
        pf = style.paragraph_format
        print(f"Heading {i} Style:")
        print(f"  page_break_before: {pf.page_break_before}")
        print(f"  keep_with_next: {pf.keep_with_next}")
        print(f"  keep_together: {pf.keep_together}")
    except KeyError:
        print(f"Heading {i} Style: NOT FOUND")

# Check for problematic paragraphs
print("\n--- CHECKING FOR page_break_before=True ---")
found_issues = False
for i, para in enumerate(output_doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    
    pf = para.paragraph_format
    if pf.page_break_before:
        style_name = para.style.name if para.style else "None"
        print(f"[{i}] '{text[:60]}...' (style: {style_name})")
        print(f"     page_break_before: {pf.page_break_before}")
        found_issues = True

if not found_issues:
    print("  ✓ No paragraphs with page_break_before=True found!")

# Check for keep_with_next=True on headings
print("\n--- CHECKING FOR keep_with_next=True ON HEADINGS ---")
heading_issues = False
for i, para in enumerate(output_doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    
    style_name = para.style.name if para.style else "None"
    if 'Heading' in style_name:
        pf = para.paragraph_format
        if pf.keep_with_next:
            print(f"[{i}] '{text[:60]}...' (style: {style_name})")
            print(f"     keep_with_next: {pf.keep_with_next}")
            heading_issues = True

if not heading_issues:
    print("  ✓ No heading paragraphs with keep_with_next=True found!")

print("\n" + "="*80)
print("TEST COMPLETE")
print("="*80)
