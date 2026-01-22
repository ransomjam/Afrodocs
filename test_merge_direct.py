#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Direct test of the coverpage merge positioning fix.
Simulates the exact merge logic from the backend.
"""

import os
import sys
import uuid
from pathlib import Path

# Add backend to path
sys.path.insert(0, str(Path(__file__).parent / 'pattern-formatter' / 'backend'))

from docx import Document
from docxcompose.composer import Composer
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import logging

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def simulate_merge_with_fix(cover_path, body_path, output_path):
    """Simulate the exact merge logic with the positioning fix."""
    
    print(f"\n{'='*80}")
    print(f"SIMULATING MERGE WITH POSITIONING FIX")
    print(f"{'='*80}\n")
    
    # Step 1: Load documents
    print("[Step 1] Loading documents...")
    cover_doc = Document(str(cover_path))
    processed_doc = Document(str(body_path))
    
    print(f"  Cover: {len(cover_doc.paragraphs)} paragraphs")
    print(f"  Body: {len(processed_doc.paragraphs)} paragraphs")
    
    # Step 2: SAVE coverpage properties BEFORE merge (KEY FIX)
    print("\n[Step 2] Saving coverpage properties BEFORE merge...")
    coverpage_original_props = []
    for idx, para in enumerate(cover_doc.paragraphs):
        coverpage_original_props.append({
            'idx': idx,
            'alignment': para.alignment,
            'line_spacing': para.paragraph_format.line_spacing,
        })
    
    print(f"  Saved {len(coverpage_original_props)} paragraph properties")
    
    # Step 3: Merge
    print("\n[Step 3] Performing merge...")
    composer = Composer(cover_doc)
    composer.append(processed_doc)
    
    # Save merged document
    temp_output = str(output_path).replace('.docx', '_temp.docx')
    composer.save(temp_output)
    print(f"  Merged document saved to temp")
    
    # Step 4: Load merged and RESTORE properties (KEY FIX)
    print("\n[Step 4] Loading merged document and restoring properties...")
    merged_doc = Document(temp_output)
    
    print(f"  Merged doc has {len(merged_doc.paragraphs)} paragraphs")
    
    # Find section break
    section_break_para_idx = None
    for para_idx, para in enumerate(merged_doc.paragraphs):
        pPr = para._element.get_or_add_pPr()
        sectPr = pPr.find(qn('w:sectPr'))
        if sectPr is not None:
            section_break_para_idx = para_idx
            break
    
    print(f"  Section break at paragraph {section_break_para_idx}")
    
    # RESTORE coverpage alignments
    if section_break_para_idx is not None:
        print(f"  Restoring coverpage properties...")
        
        restored_count = 0
        for orig_props in coverpage_original_props:
            orig_idx = orig_props['idx']
            if orig_idx <= section_break_para_idx and orig_idx < len(merged_doc.paragraphs):
                para = merged_doc.paragraphs[orig_idx]
                
                # Restore alignment if changed
                if para.alignment != orig_props['alignment']:
                    para.alignment = orig_props['alignment']
                    restored_count += 1
                
                # Restore line spacing if changed
                if para.paragraph_format.line_spacing != orig_props['line_spacing']:
                    para.paragraph_format.line_spacing = orig_props['line_spacing']
                    restored_count += 1
        
        print(f"  Restored {restored_count} properties")
    
    # Step 5: Save final document
    print("\n[Step 5] Saving final document...")
    merged_doc.save(str(output_path))
    print(f"  Saved to: {output_path}")
    
    # Step 6: Verify
    print("\n[Step 6] Verifying positioning preservation...")
    template_doc = Document(str(cover_path))
    final_doc = Document(str(output_path))
    
    preserved = 0
    changed = 0
    
    print(f"\n{'Para':<5} {'Original':<15} {'Final':<15} {'Status':<15}")
    print("-" * 55)
    
    for idx in range(min(25, len(template_doc.paragraphs), section_break_para_idx + 1 if section_break_para_idx else 0)):
        orig_align = "None" if template_doc.paragraphs[idx].alignment is None else str(template_doc.paragraphs[idx].alignment).split('.')[-1]
        final_align = "None" if final_doc.paragraphs[idx].alignment is None else str(final_doc.paragraphs[idx].alignment).split('.')[-1]
        
        if template_doc.paragraphs[idx].alignment == final_doc.paragraphs[idx].alignment:
            status = "PRESERVED"
            preserved += 1
        else:
            status = "CHANGED"
            changed += 1
        
        print(f"{idx:<5} {orig_align:<15} {final_align:<15} {status:<15}")
    
    print(f"\n{'='*80}")
    print(f"Result: {preserved} preserved, {changed} changed")
    
    if changed == 0:
        print("SUCCESS: All properties preserved!")
        return True
    else:
        print(f"ISSUE: {changed} properties were modified")
        return False

def main():
    """Main test."""
    workspace = Path(__file__).parent
    outputs_dir = workspace / 'pattern-formatter' / 'backend' / 'outputs'
    template_dir = workspace / 'pattern-formatter' / 'backend' / 'coverpage_template'
    
    # Use an existing body document
    docx_files = sorted(outputs_dir.glob('*.docx'), key=os.path.getmtime, reverse=True)
    if not docx_files:
        print("No documents found in outputs")
        return False
    
    # Use first non-meta file
    body_doc = None
    for f in docx_files:
        if '_meta' not in f.name and '_pdf' not in f.name:
            body_doc = f
            break
    
    if not body_doc:
        print("No valid body document found")
        return False
    
    # Use template as coverpage
    cover_doc = template_dir / 'dissertation_coverpage_template.docx'
    
    if not cover_doc.exists():
        print(f"Template not found: {cover_doc}")
        return False
    
    # Output path
    test_id = str(uuid.uuid4())[:8]
    output_path = outputs_dir / f"test_merge_{test_id}.docx"
    
    print(f"Using:")
    print(f"  Template: {cover_doc.name}")
    print(f"  Body: {body_doc.name}")
    print(f"  Output: {output_path.name}")
    
    success = simulate_merge_with_fix(str(cover_doc), str(body_doc), str(output_path))
    
    return success

if __name__ == '__main__':
    try:
        success = main()
        sys.exit(0 if success else 1)
    except Exception as e:
        print(f"\nERROR: {str(e)}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
