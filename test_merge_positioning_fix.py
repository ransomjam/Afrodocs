#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Test the coverpage merge with positioning preservation fix.
Simulates what happens when user adds coverpage to a formatted document.
"""

import sys
import os
import json
import uuid
from pathlib import Path

# Add backend path
sys.path.insert(0, str(Path(__file__).parent / 'pattern-formatter' / 'backend'))

from pattern_formatter_backend import WordGenerator, generate_cover_page, OUTPUT_FOLDER
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def test_merge_with_positioning():
    """Test that coverpage positioning is preserved during merge."""
    
    print(f"\n{'='*80}")
    print(f"TEST: Coverpage Merge with Positioning Preservation")
    print(f"{'='*80}\n")
    
    try:
        # Step 1: Generate a simple body document
        print("[1/4] Generating formatted body document...")
        gen = WordGenerator()
        body_data = {
            'full_name': 'Test Student',
            'degree_type': 'Master',
            'institution': 'Test University',
            'department': 'Computer Science',
            'field_of_study': 'AI',
            'supervisor': 'Prof. Test',
            'chapters': [
                {
                    'heading': 'Introduction',
                    'content': 'This is test content. ' * 20,
                    'subsections': []
                }
            ]
        }
        
        job_id = str(uuid.uuid4())
        body_output = os.path.join(OUTPUT_FOLDER, f"{job_id}_formatted.docx")
        
        # This will call the generate method which formats the document
        gen.generate(
            body_data,
            body_output,
            include_toc=False,
            font_size=12,
            line_spacing=1.5
        )
        
        if not os.path.exists(body_output):
            print(f"✗ FAILED: Body document not created")
            return False
        
        print(f"✓ Body document created: {body_output}")
        
        # Step 2: Load template
        print("\n[2/4] Loading coverpage template...")
        template_path = str(Path(__file__).parent / 'pattern-formatter' / 'backend' / 'coverpage_template' / 'dissertation_coverpage_template.docx')
        template_doc = Document(template_path)
        
        # Save original template properties
        template_props = {}
        for idx, para in enumerate(template_doc.paragraphs):
            template_props[idx] = {
                'alignment': para.alignment,
                'line_spacing': para.paragraph_format.line_spacing,
                'text': para.text[:50] if para.text else ""
            }
        
        print(f"✓ Template loaded with {len(template_doc.paragraphs)} paragraphs")
        
        # Step 3: Generate coverpage
        print("\n[3/4] Generating coverpage...")
        coverpage_data = {
            'institution': 'Test University',
            'department': 'Computer Science',
            'course_title': 'Test Course',
            'course_code': 'CS101',
            'level': 'Master',
            'student_name': 'Test Student',
            'student_id': '12345',
            'supervisor': 'Prof. Test',
            'submission_date': '2024',
        }
        
        coverpage_output, error = generate_cover_page(coverpage_data)
        if error:
            print(f"✗ FAILED: {error}")
            return False
        
        print(f"✓ Coverpage generated")
        
        # Step 4: Merge (this triggers our new fix)
        print("\n[4/4] Merging coverpage with body document...")
        
        # Load and merge
        cover_doc = Document(coverpage_output)
        processed_doc = Document(body_output)
        
        # Save original coverpage props from merged doc (simulating what will happen)
        original_cover_props = []
        for idx, para in enumerate(cover_doc.paragraphs):
            original_cover_props.append({
                'idx': idx,
                'alignment': para.alignment,
                'line_spacing': para.paragraph_format.line_spacing,
            })
        
        # Import Composer
        from docxcompose.composer import Composer
        
        # Merge
        composer = Composer(cover_doc)
        composer.append(processed_doc)
        
        merge_output = os.path.join(OUTPUT_FOLDER, f"{job_id}_merged_test.docx")
        composer.save(merge_output)
        
        # Load merged document
        merged_doc = Document(merge_output)
        
        print(f"✓ Merge complete: {len(merged_doc.paragraphs)} paragraphs")
        
        # Step 5: Check positioning preservation
        print("\n" + "="*80)
        print("ANALYSIS: Checking if coverpage positioning is preserved")
        print("="*80)
        
        # Find section break
        section_break_idx = None
        for para_idx, para in enumerate(merged_doc.paragraphs):
            pPr = para._element.get_or_add_pPr()
            sectPr = pPr.find(qn('w:sectPr'))
            if sectPr is not None:
                section_break_idx = para_idx
                break
        
        if section_break_idx is None:
            print("! No section break found in merged document")
            check_limit = min(15, len(cover_doc.paragraphs))
        else:
            check_limit = min(15, section_break_idx)
        
        print(f"\nComparing first {check_limit} paragraphs:")
        print(f"{'Para':<5} {'Original Align':<16} {'Original LS':<12} {'Merged Align':<16} {'Merged LS':<12} {'Status':<10}")
        print("-" * 90)
        
        preserved = 0
        changed = 0
        
        for idx in range(min(check_limit, len(original_cover_props), len(merged_doc.paragraphs))):
            orig_props = original_cover_props[idx]
            merged_para = merged_doc.paragraphs[idx]
            
            orig_align = "None" if orig_props['alignment'] is None else str(orig_props['alignment']).split('.')[-1]
            orig_ls = orig_props['line_spacing'] or "default"
            
            merged_align = "None" if merged_para.alignment is None else str(merged_para.alignment).split('.')[-1]
            merged_ls = merged_para.paragraph_format.line_spacing or "default"
            
            # Check if properties match
            align_preserved = orig_props['alignment'] == merged_para.alignment
            ls_preserved = orig_props['line_spacing'] == merged_para.paragraph_format.line_spacing
            
            if align_preserved and ls_preserved:
                status = "✓ OK"
                preserved += 1
            else:
                status = "✗ CHANGED"
                changed += 1
            
            print(f"{idx:<5} {orig_align:<16} {str(orig_ls):<12} {merged_align:<16} {str(merged_ls):<12} {status:<10}")
        
        print(f"\n{'='*80}")
        print(f"RESULT: {preserved} preserved, {changed} changed")
        
        if changed == 0:
            print(f"✓ SUCCESS: All coverpage properties were preserved!")
            return True
        else:
            print(f"✗ ISSUE: {changed} coverpage properties were modified")
            return False
            
    except Exception as e:
        print(f"✗ ERROR: {str(e)}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == '__main__':
    try:
        success = test_merge_with_positioning()
        sys.exit(0 if success else 1)
    except Exception as e:
        print(f"Fatal error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
