#!/usr/bin/env python3
"""
IMPROVED regex patterns that match actual document structure
"""

import re

class ImprovedTextFormatter:
    """Improved patterns for real document content"""
    
    def __init__(self):
        self.patterns = [
            # Pattern 1: Numbered sections at start of line with optional colon (1.1, 2.1, 1.2, etc.)
            {
                'name': 'Numbered sections (1.1, 2.2, etc)',
                'regex': r'^(\d+\.\d+)(\s+)([A-Z][A-Za-z\s\-:]*?)(?=\n|$)',
                'replacement': r'**\1\2\3**',
                'flags': re.MULTILINE
            },
            
            # Pattern 2: Simple numbered items (1., 2., 3.) at line start with title
            {
                'name': 'Simple numbered items (1., 2., 3.)',
                'regex': r'^(\d+\.\s+)([A-Z][A-Za-z\s\-:]*?)(?=\n|$)',
                'replacement': r'**\1\2**',
                'flags': re.MULTILINE
            },
            
            # Pattern 3: Roman numerals with title (I., II., III.)
            {
                'name': 'Roman numerals (I., II., III.)',
                'regex': r'^([IVX]+\.\s+)([A-Z][A-Za-z\s\-:]*?)(?=\n|$)',
                'replacement': r'**\1\2**',
                'flags': re.MULTILINE
            },
            
            # Pattern 4: Bulleted items with text
            {
                'name': 'Bulleted items',
                'regex': r'^(\s*[-•]\s+)([A-Z][A-Za-z\s\-:]*?)(?=\n|$)',
                'replacement': r'\1**\2**',
                'flags': re.MULTILINE
            },
            
            # Pattern 5: Section headers in capitals followed by colon (METHODOLOGY:, RESULTS:)
            {
                'name': 'Section headers (CAPITALS with colon)',
                'regex': r'^([A-Z][A-Z\s]+):(?=\s|$)',
                'replacement': r'**\1:**',
                'flags': re.MULTILINE
            },
        ]
    
    def format_text(self, text):
        if not text:
            return text
        
        changes = []
        for pattern_config in self.patterns:
            try:
                regex = pattern_config['regex']
                replacement = pattern_config['replacement']
                flags = pattern_config['flags']
                
                # Track what gets replaced
                matches = list(re.finditer(regex, text, flags))
                if matches:
                    changes.append((pattern_config['name'], len(matches)))
                
                text = re.sub(regex, replacement, text, flags=flags)
                
            except Exception as e:
                print(f"Error in pattern '{pattern_config['name']}': {e}")
                continue
        
        return text, changes


# Test on actual chapter one content
from docx import Document

def extract_chapter_one(doc_path):
    doc = Document(doc_path)
    chapter_start = None
    chapter_end = None
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip().upper()
        if 'CHAPTER' in text and ('ONE' in text or text.endswith('1')):
            chapter_start = i
        if chapter_start is not None and chapter_start != i:
            if 'CHAPTER' in text and ('TWO' in text or text.endswith('2')):
                chapter_end = i
                break
    
    if chapter_start is None:
        return ""
    
    chapter_end = chapter_end or len(doc.paragraphs)
    text_parts = []
    
    for i in range(chapter_start, chapter_end):
        text_parts.append(doc.paragraphs[i].text)
    
    return '\n'.join(text_parts)


print("="*70)
print("TESTING IMPROVED PATTERNS ON REAL DOCUMENT")
print("="*70)

formatter = ImprovedTextFormatter()

doc_path = r"c:\Users\user\Desktop\PATTERN\Samples\bulleting and numbering sample 2.docx"
chapter_text = extract_chapter_one(doc_path)

print("\nBEFORE (first 1000 chars):")
print("-" * 70)
print(chapter_text[:1000])

formatted_text, changes = formatter.format_text(chapter_text)

print("\n\nAFTER (first 1000 chars):")
print("-" * 70)
print(formatted_text[:1000])

print("\n\nPattern Matches Found:")
print("-" * 70)
for name, count in changes:
    print(f"  {name}: {count} matches")

print("\n\nSample Changes:")
print("-" * 70)
before_lines = chapter_text.split('\n')
after_lines = formatted_text.split('\n')

change_count = 0
for i, (before, after) in enumerate(zip(before_lines, after_lines)):
    if before != after and before.strip():
        print(f"\n  Line {i}:")
        print(f"    Before: {before}")
        print(f"    After:  {after}")
        change_count += 1
        if change_count >= 20:
            break

print(f"\n\nTotal lines changed: {change_count}")
