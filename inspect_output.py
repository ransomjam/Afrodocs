#!/usr/bin/env python

from docx import Document

output_path = 'test_formatting_output.docx'

try:
    doc = Document(output_path)
    print(f'✓ Generated file exists: {output_path}')
    print(f'Total paragraphs: {len(doc.paragraphs)}\n')
    
    # Look for numbered/bulleted items to verify formatting
    print('=== NUMBERED/BULLETED ITEMS AND THEIR FORMATTING ===\n')
    
    count = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        
        # Check if it looks like a list item (starts with number, bullet, or similar)
        if text and (text[0].isdigit() or text[0] in '•-*○■'):
            has_bold = any(run.bold for run in para.runs)
            words = len(text.split())
            indent = len(para.paragraph_format.left_indent.inches) if para.paragraph_format.left_indent else 0
            
            print(f'Item {i}:')
            print(f'  Text: {text[:80]}{"..." if len(text) > 80 else ""}')
            print(f'  Bold: {has_bold}')
            print(f'  Words: {words}')
            print(f'  Left Indent: {indent:.2f} inches')
            print()
            
            count += 1
            
    if count == 0:
        print('No numbered/bulleted items found. Showing first 10 paragraphs:\n')
        for i, para in enumerate(doc.paragraphs[:10]):
            text = para.text.strip()
            if text:
                has_bold = any(run.bold for run in para.runs)
                indent = len(para.paragraph_format.left_indent.inches) if para.paragraph_format.left_indent else 0
                print(f'{i}: [BOLD={has_bold}] [INDENT={indent:.2f}"] {text[:70]}')
            
except Exception as e:
    print(f'✗ Error: {e}')
