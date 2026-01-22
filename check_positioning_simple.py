#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Simple test to verify positioning is preserved when merging coverpage and body document.
"""

import sys
import os
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def check_positioning(doc_path, title):
    """Check positioning properties of document."""
    doc = Document(doc_path)
    
    print(f"\n{'='*70}")
    print(f"{title}")
    print(f"{'='*70}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total sections: {len(doc.sections)}")
    
    # Find section break
    section_break_idx = None
    for idx, para in enumerate(doc.paragraphs):
        pPr = para._element.get_or_add_pPr()
        from docx.oxml.ns import qn
        sectPr = pPr.find(qn('w:sectPr'))
        if sectPr is not None:
            section_break_idx = idx
            print(f"✓ Section break found at paragraph {idx}")
            break
    
    if section_break_idx is None:
        print(f"! No section break found")
    
    # Show first 15 paragraphs (coverpage area)
    print(f"\nCOVERPAGE AREA (first 15 paragraphs):")
    print(f"{'#':<4} {'Align':<10} {'LS':<8} {'Text':<40}")
    print("-" * 72)
    
    for idx in range(min(15, len(doc.paragraphs))):
        para = doc.paragraphs[idx]
        align = "None" if para.alignment is None else str(para.alignment).split('.')[-1]
        ls = para.paragraph_format.line_spacing
        text = para.text[:35] if para.text else "(empty)"
        print(f"{idx:<4} {align:<10} {str(ls):<8} {text:<40}")
    
    # If we found a section break, show body area
    if section_break_idx:
        body_start = section_break_idx + 1
        body_end = min(body_start + 10, len(doc.paragraphs))
        
        print(f"\nBODY AREA (paragraphs {body_start}-{body_end}):")
        print(f"{'#':<4} {'Align':<10} {'LS':<8} {'Text':<40}")
        print("-" * 72)
        
        for idx in range(body_start, body_end):
            para = doc.paragraphs[idx]
            align = "None" if para.alignment is None else str(para.alignment).split('.')[-1]
            ls = para.paragraph_format.line_spacing
            text = para.text[:35] if para.text else "(empty)"
            print(f"{idx:<4} {align:<10} {str(ls):<8} {text:<40}")
    
    # Show statistics
    print(f"\nSTATISTICS:")
    
    if section_break_idx:
        # Count justified paragraphs in coverpage vs body
        cover_just = sum(1 for i in range(section_break_idx) if doc.paragraphs[i].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY)
        body_just = sum(1 for i in range(section_break_idx + 1, len(doc.paragraphs)) if doc.paragraphs[i].alignment == WD_ALIGN_PARAGRAPH.JUSTIFY)
        
        print(f"  Coverpage (0-{section_break_idx-1}): {cover_just} justified out of {section_break_idx}")
        print(f"  Body ({section_break_idx+1}-{len(doc.paragraphs)-1}): {body_just} justified out of {len(doc.paragraphs) - section_break_idx - 1}")
        
        if cover_just == 0:
            print(f"  ✓ GOOD: Coverpage preserved (no forced justification)")
        else:
            print(f"  ✗ ISSUE: Coverpage has {cover_just} justified paragraphs (should be preserved as-is)")

def main():
    """Main test function."""
    workspace = Path(__file__).parent
    test_output_dir = workspace / 'pattern-formatter' / 'backend' / 'outputs'
    
    # Check what documents we have
    test_docs = list(test_output_dir.glob('*.docx'))
    
    if not test_docs:
        print(f"No test documents found in {test_output_dir}")
        print(f"Please generate a document first through the Flask app")
        return
    
    print(f"\nFound {len(test_docs)} test documents")
    
    # Look for merged and body documents
    for doc_path in sorted(test_docs):
        if 'merged' in doc_path.name or 'body' in doc_path.name or 'output' in doc_path.name:
            check_positioning(str(doc_path), f"Document: {doc_path.name}")

if __name__ == '__main__':
    main()
