import requests
from docx import Document
import os
from datetime import datetime
import time

API_BASE = 'http://localhost:5000'
session = requests.Session()

# Login
print("Logging in...")
login_response = session.post(f'{API_BASE}/api/auth/login', json={
    'username': 'admin',
    'password': 'admin@secure123'
})

if login_response.status_code != 200:
    print(f"Login failed!")
    exit(1)

print("✓ Logged in\n")

# Test data for each institution
test_cases = [
    {
        'name': 'Bamenda (UBA)',
        'institution': 'uba',
        'faculty': 'College of Technology',
        'department': 'Computer Engineering',
        'document_type': 'Assignment'
    },
    {
        'name': 'Buea (UB)',
        'institution': 'ub',
        'faculty': 'Faculty of Science',
        'department': 'Department of Geology',
        'document_type': 'Thesis'
    },
    {
        'name': 'NPUI',
        'institution': 'npui',
        'faculty': 'School of Applied Sciences',
        'department': 'Information Technology',
        'document_type': 'Research Proposal'
    }
]

# Files are saved to backend/outputs/ not outputs/Cover Pages/
output_dir = r"c:\Users\user\Desktop\PATTERN\pattern-formatter\backend\outputs"

print("=" * 80)
print("GENERATING AND VERIFYING TEMPLATES")
print("=" * 80)

for test in test_cases:
    print(f"\n{test['name']} ({test['institution']})")
    print(f"Document Type: {test['document_type']}")
    
    payload = {
        'institution': test['institution'],
        'faculty': test['faculty'],
        'department': test['department'],
        'documentType': test['document_type'],
        'studentName': 'Test Student',
        'studentId': 'TEST2026',
        'title': f'Verify {test["name"]}',
        'supervisor': 'Dr. Smith' if test['document_type'] in ['Thesis', 'Research Proposal'] else None,
        'instructor': 'Prof. Johnson' if test['document_type'] == 'Assignment' else None
    }
    
    response = session.post(f'{API_BASE}/api/coverpage/generate', json=payload)
    
    if response.status_code == 200:
        data = response.json()
        if data.get('success'):
            job_id = data.get('job_id')
            print(f"  ✓ Generated successfully")
            print(f"  Job ID: {job_id}")
            
            # Find the generated file
            time.sleep(0.3)  # Wait for file to be written
            expected_file = os.path.join(output_dir, f"{job_id}_formatted.docx")
            
            if os.path.exists(expected_file):
                try:
                    doc = Document(expected_file)
                    
                    # Get all text to check for institution markers
                    all_text = '\n'.join([p.text for p in doc.paragraphs])
                    
                    # Check for key indicators
                    if 'Bamenda' in all_text or 'University of Bamenda' in all_text:
                        print(f"  ✓ Template: Bamenda (UBA)")
                    elif 'Buea' in all_text or 'University of Buea' in all_text:
                        print(f"  ✓ Template: Buea (UB)")
                    elif 'National University Institute' in all_text or 'NPUI' in all_text:
                        print(f"  ✓ Template: NPUI")
                    else:
                        print(f"  ℹ Template: Could not identify from content")
                    
                    # Show paragraph count
                    print(f"  ℹ Paragraphs: {len(doc.paragraphs)}")
                    
                except Exception as e:
                    print(f"  ✗ Error reading file: {str(e)}")
            else:
                print(f"  ✗ File not found at: {expected_file}")
                print(f"  Checking directory...")
                if os.path.exists(output_dir):
                    files = os.listdir(output_dir)
                    recent = sorted(files, key=lambda f: os.path.getmtime(os.path.join(output_dir, f)), reverse=True)[:3]
                    print(f"  Recent files: {recent}")
        else:
            print(f"  ✗ Generation failed: {data.get('error')}")
    else:
        print(f"  ✗ API error: {response.status_code}")

print("\n" + "=" * 80)
print("VERIFICATION COMPLETE!")
print("=" * 80)
