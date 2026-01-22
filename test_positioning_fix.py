#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Test to verify that positioning properties are preserved correctly after merge.
Tests that coverpage paragraphs keep their original alignment/spacing, 
while body paragraphs get justified formatting.
"""

import sys
import os
from pathlib import Path

# Add path to imports
sys.path.insert(0, str(Path(__file__).parent / 'pattern-formatter' / 'backend'))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def analyze_document_positioning(doc_path, title):
    """Analyze positioning properties of a document."""
    doc = Document(doc_path)
    
    print(f"\n{'='*60}")
    print(f"ANALYSIS: {title}")
    print(f"{'='*60}")
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total sections: {len(doc.sections)}")
    
    # Check for section breaks to identify cover/body boundary
    section_break_para = None
    for idx, para in enumerate(doc.paragraphs):
        pPr = para._element.get_or_add_pPr()
        from docx.oxml.ns import qn
        sectPr = pPr.find(qn('w:sectPr'))
        if sectPr is not None:
            section_break_para = idx
            print(f"\n✓ Section break found at paragraph {idx}")
            break
    
    # Analyze first 10 paragraphs (should mostly be coverpage)
    print(f"\nCOVERPAGE AREA (first 10 paragraphs):")
    print(f"{'Para':<5} {'Align':<10} {'LineSpacing':<12} {'Style':<25}")
    print("-" * 52)
    
    for idx in range(min(10, len(doc.paragraphs))):
        para = doc.paragraphs[idx]
        align = str(para.alignment) if para.alignment is not None else "None"
        ls = para.paragraph_format.line_spacing if para.paragraph_format.line_spacing else "default"
        print(f"{idx:<5} {align:<10} {str(ls):<12} {para.style.name:<25}")
    
    # Analyze paragraphs after section break (should be body)
    if section_break_para:
        start_idx = section_break_para + 1
        end_idx = min(start_idx + 10, len(doc.paragraphs))
        
        print(f"\nBODY AREA (paragraphs {start_idx}-{end_idx}):")
        print(f"{'Para':<5} {'Align':<10} {'LineSpacing':<12} {'Style':<25}")
        print("-" * 52)
        
        for idx in range(start_idx, end_idx):
            para = doc.paragraphs[idx]
            align = str(para.alignment) if para.alignment is not None else "None"
            ls = para.paragraph_format.line_spacing if para.paragraph_format.line_spacing else "default"
            print(f"{idx:<5} {align:<10} {str(ls):<12} {para.style.name:<25}")
    
    # Check formatting consistency
    print(f"\nFORMATTING CHECK:")
    
    # Count alignments in different regions
    if section_break_para:
        cover_count_justify = sum(1 for para in doc.paragraphs[:section_break_para] if para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY)
        body_count_justify = sum(1 for para in doc.paragraphs[section_break_para+1:] if para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY)
        
        print(f"  Coverpage justified: {cover_count_justify}/{section_break_para}")
        print(f"  Body justified: {body_count_justify}/{len(doc.paragraphs) - section_break_para - 1}")
        
        if cover_count_justify == 0:
            print("  ✓ GOOD: Coverpage has no forced justification (preserves original)")
        else:
            print(f"  ✗ ISSUE: Coverpage has {cover_count_justify} justified paragraphs (should be preserved)")
            
        if body_count_justify > 0:
            print(f"  ✓ GOOD: Body has {body_count_justify} justified paragraphs")
        else:
            print("  ! WARNING: Body has no justified paragraphs")

def main():
    """Test positioning preservation."""
    workspace_root = Path(__file__).parent
    test_dir = workspace_root / 'test_output'
    test_dir.mkdir(exist_ok=True)
    
    # Create test document with coverpage
    print("\n" + "="*60)
    print("POSITIONING PRESERVATION TEST")
    print("="*60)
    
    # Check if we can test the actual merge
    try:
        from pattern_formatter_backend import WordGenerator
        from flask import Flask
        from flask_login import LoginManager
        import sqlite3
        
        app = Flask(__name__)
        app.config['SECRET_KEY'] = 'test-key'
        login_manager = LoginManager()
        login_manager.init_app(app)
        
        # Simple user for testing
        class TestUser:
            id = 'test_user'
            is_authenticated = True
            is_active = True
            is_anonymous = False
            
            def get_id(self):
                return self.id
        
        @login_manager.user_loader
        def load_user(user_id):
            return TestUser()
        
        with app.app_context():
            print("\n[Test] Initializing WordGenerator...")
            gen = WordGenerator()
            
            # Create simple test data
            test_data = {
                'full_name': 'Test Student',
                'degree_type': 'Master',
                'institution': 'Test University',
                'department': 'Computer Science',
                'field_of_study': 'AI',
                'supervisor': 'Prof. Test',
                'chapters': [
                    {
                        'title': 'Introduction',
                        'content': 'This is a test chapter. ' * 50
                    }
                ]
            }
            
            # Generate body document
            print("[Test] Generating formatted document...")
            body_path = test_dir / 'test_body.docx'
            gen.generate_word_document(test_data, str(body_path))
            
            # Generate with coverpage
            print("[Test] Adding coverpage to document...")
            merged_path = test_dir / 'test_merged.docx'
            gen.add_coverpage_to_document(str(body_path), str(merged_path))
            
            # Analyze both documents
            analyze_document_positioning(str(body_path), "Formatted Body Document")
            analyze_document_positioning(str(merged_path), "Merged Document (Body + Coverpage)")
            
            print("\n" + "="*60)
            print("TEST COMPLETE")
            print("="*60)
            
    except Exception as e:
        print(f"Error during test: {e}")
        import traceback
        traceback.print_exc()

if __name__ == '__main__':
    main()
