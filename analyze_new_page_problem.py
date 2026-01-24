"""
Analyze the new page problem documents to understand what's causing headings to jump to new pages.
"""

import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'pattern-formatter', 'backend'))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def analyze_document(filepath):
    """Analyze a Word document for page break issues"""
    print(f"\n{'='*70}")
    print(f"ANALYZING: {os.path.basename(filepath)}")
    print(f"{'='*70}")
    
    doc = Document(filepath)
    
    print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
    
    # Look for headings and their formatting properties
    print("\n--- HEADINGS AND PAGE BREAK PROPERTIES ---")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
            
        style_name = para.style.name if para.style else 'None'
        
        # Check if it looks like a heading
        is_heading = (
            'Heading' in style_name or
            text.startswith('CHAPTER') or
            (len(text) < 100 and text.isupper()) or
            any(c.isdigit() and '.' in text[:10] for c in text[:5])
        )
        
        if is_heading or 'Heading' in style_name:
            pf = para.paragraph_format
            
            # Get page break properties
            page_break_before = pf.page_break_before
            keep_with_next = pf.keep_with_next
            keep_together = pf.keep_together
            widow_control = pf.widow_control
            
            # Check spacing
            space_before = pf.space_before
            space_after = pf.space_after
            
            # Check for explicit page break in runs
            has_run_page_break = False
            for run in para.runs:
                # Check XML for page break
                xml = run._element.xml
                if '<w:br w:type="page"' in xml:
                    has_run_page_break = True
                    break
            
            print(f"\n[{i}] Style: {style_name}")
            print(f"    Text: {text[:60]}{'...' if len(text) > 60 else ''}")
            print(f"    page_break_before: {page_break_before}")
            print(f"    keep_with_next: {keep_with_next}")
            print(f"    keep_together: {keep_together}")
            print(f"    widow_control: {widow_control}")
            print(f"    space_before: {space_before}")
            print(f"    space_after: {space_after}")
            print(f"    has_run_page_break: {has_run_page_break}")
    
    # Check document styles
    print("\n--- DOCUMENT HEADING STYLES ---")
    for style_name in ['Heading 1', 'Heading 2', 'Heading 3', 'Heading 4']:
        try:
            style = doc.styles[style_name]
            pf = style.paragraph_format
            print(f"\n{style_name}:")
            print(f"    page_break_before: {pf.page_break_before}")
            print(f"    keep_with_next: {pf.keep_with_next}")
            print(f"    keep_together: {pf.keep_together}")
        except KeyError:
            print(f"\n{style_name}: Not found in document")

def main():
    samples_dir = os.path.join(os.path.dirname(__file__), 'samples', 'new page problems')
    
    # Analyze all documents in the folder
    for filename in os.listdir(samples_dir):
        if filename.endswith('.docx') and not filename.startswith('~$'):
            filepath = os.path.join(samples_dir, filename)
            try:
                analyze_document(filepath)
            except Exception as e:
                print(f"Error analyzing {filename}: {e}")

if __name__ == '__main__':
    main()
