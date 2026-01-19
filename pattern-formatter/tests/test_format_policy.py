import os
import re
import sys
import tempfile
import types
import unittest

from docx import Document

sys.modules.setdefault('requests', types.ModuleType('requests'))
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'backend')))

from pattern_formatter_backend import DocumentProcessor, WordGenerator, FormatPolicy


class FormatPolicyTests(unittest.TestCase):
    def _generate_doc(self, text):
        policy = FormatPolicy()
        processor = DocumentProcessor(policy=policy)
        result, _ = processor.process_text(text)

        with tempfile.TemporaryDirectory() as tmpdir:
            output_path = os.path.join(tmpdir, 'output.docx')
            generator = WordGenerator(policy=policy)
            generator.generate(
                result['structured'],
                output_path,
                include_toc=False,
                font_size=12,
                line_spacing=1.5,
                margins=3.0,
            )
            doc = Document(output_path)
        return result, doc

    def test_headings_without_numbering_remain_unnumbered(self):
        text = "INTRODUCTION\n\nBACKGROUND"
        policy = FormatPolicy()
        processor = DocumentProcessor(policy=policy)
        result, _ = processor.process_text(text)

        headings = [section.get('heading', '') for section in result['structured']]
        for heading in headings:
            if heading:
                self.assertIsNone(re.match(r'^\d', heading))

    def test_headings_with_numbering_preserve_numbers(self):
        text = "2.1 RESEARCH DESIGN\n\n2.2 DATA COLLECTION"
        policy = FormatPolicy()
        processor = DocumentProcessor(policy=policy)
        result, _ = processor.process_text(text)

        headings = [section.get('heading', '') for section in result['structured']]
        self.assertTrue(any(h.startswith('2.1') for h in headings))
        self.assertTrue(any(h.startswith('2.2') for h in headings))

    def test_numbered_lists_do_not_double_or_renumber(self):
        text = "1. First item\n1. Second item"
        _, doc = self._generate_doc(text)

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        self.assertIn("1. First item", paragraphs)
        self.assertIn("1. Second item", paragraphs)
        self.assertFalse(any("1. 1." in p for p in paragraphs))

    def test_label_bold_applies_only_to_label(self):
        text = "Definition: Something important"
        _, doc = self._generate_doc(text)

        target_para = None
        for para in doc.paragraphs:
            if para.text.strip().startswith("Definition:"):
                target_para = para
                break

        self.assertIsNotNone(target_para)
        self.assertGreaterEqual(len(target_para.runs), 2)
        self.assertTrue(target_para.runs[0].text.startswith("Definition:"))
        self.assertTrue(target_para.runs[0].bold)
        self.assertFalse(target_para.runs[1].bold)


if __name__ == "__main__":
    unittest.main()
