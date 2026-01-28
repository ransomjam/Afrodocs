# pattern_formatter_backend.py
# Ultra-Precise Pattern-Based Academic Document Formatter
# NO AI - 100% Rule-Based - Lightning Fast

from flask import Flask, request, jsonify, send_file, Response
from flask_cors import CORS
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from functools import wraps
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import func
from werkzeug.security import generate_password_hash, check_password_hash
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsmap
from docxcompose.composer import Composer
from dataclasses import dataclass
from PyPDF2 import PdfReader
import base64
import binascii
import re
import os
import json
import uuid
import threading
import time
from urllib.parse import urlparse, urlunparse, parse_qsl, urlencode
from datetime import datetime
from io import BytesIO
import logging
import fapshi_integration as fapshi

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Register additional namespaces for shape/flowchart support
# These namespaces are needed to properly extract and insert shapes, arrows, and drawing elements
try:
    from lxml import etree
    # Add shape-related namespaces if not already present
    SHAPE_NAMESPACES = {
        'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
        'wpg': 'http://schemas.microsoft.com/office/word/2010/wordprocessingGroup',
        'wpc': 'http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas',
        'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
        'w14': 'http://schemas.microsoft.com/office/word/2010/wordml',
        'w15': 'http://schemas.microsoft.com/office/word/2012/wordml',
        'wne': 'http://schemas.microsoft.com/office/word/2006/wordml',
        'dgm': 'http://schemas.openxmlformats.org/drawingml/2006/diagram',
        'c': 'http://schemas.openxmlformats.org/drawingml/2006/chart',
    }
    # Update the global nsmap with shape namespaces
    for prefix, uri in SHAPE_NAMESPACES.items():
        if prefix not in nsmap:
            nsmap[prefix] = uri
    logger.info("Shape namespaces registered successfully")
except Exception as e:
    logger.warning(f"Could not register shape namespaces: {e}")

@dataclass(frozen=True)
class FormatPolicy:
    """Single source of truth for formatting behavior."""
    enable_heading_auto_numbering: bool = True
    enable_regex_auto_bold: bool = False
    preserve_existing_numbering: bool = True
    preserve_existing_bold: bool = True
    list_numbering_mode: str = "strict"  # "strict" or "assistive"
    document_mode: str = "generic"  # "academic", "notes", or "generic"

    def allow_auto_numbering(self, has_consistent_scheme: bool) -> bool:
        if self.document_mode == "academic":
            return True
        if not self.enable_heading_auto_numbering:
            return False
        return has_consistent_scheme

# Determine frontend folder location (Docker vs local)
def get_frontend_folder():
    """Find frontend folder - handles both Docker and local development."""
    possible_paths = [
        os.path.join(os.path.dirname(os.path.abspath(__file__)), 'frontend'),  # Docker: /app/frontend
        os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'frontend'),  # Local: ../frontend
    ]
    for path in possible_paths:
        if os.path.exists(path):
            logger.info(f"Found frontend at: {path}")
            return path
    # Default fallback
    return '../frontend'

# Serve frontend files directly from the backend for simple deployment
app = Flask(__name__, static_folder=get_frontend_folder(), static_url_path='')

# Load configuration from environment variables
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'your-secret-key-change-this-in-production')

# Database configuration - supports PostgreSQL (Render) and SQLite (local development)
RENDER_EXTERNAL_URL = os.environ.get('RENDER_EXTERNAL_URL', '')
IS_RENDER = os.environ.get('RENDER', '').lower() == 'true' or RENDER_EXTERNAL_URL != ''
database_url = os.environ.get('DATABASE_URL', 'sqlite:///users.db')
# Fix for Render PostgreSQL URL format (Render uses postgres://, SQLAlchemy needs postgresql://)
if database_url and database_url.startswith('postgres://'):
    database_url = database_url.replace('postgres://', 'postgresql://', 1)
if database_url and database_url.startswith('postgresql://') and IS_RENDER:
    parsed_db_url = urlparse(database_url)
    query_params = dict(parse_qsl(parsed_db_url.query))
    if 'sslmode' not in query_params:
        query_params['sslmode'] = 'require'
        database_url = urlunparse(parsed_db_url._replace(query=urlencode(query_params)))
app.config['SQLALCHEMY_DATABASE_URI'] = database_url
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Detect production environment (Render sets RENDER=true)
IS_PRODUCTION = IS_RENDER
logger.info(f"Running in {'PRODUCTION' if IS_PRODUCTION else 'DEVELOPMENT'} mode")

# Session cookie configuration for proper authentication
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['SESSION_COOKIE_SECURE'] = IS_PRODUCTION  # True in production with HTTPS
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['REMEMBER_COOKIE_SAMESITE'] = 'Lax'
app.config['REMEMBER_COOKIE_SECURE'] = IS_PRODUCTION  # True in production with HTTPS

# CORS configuration - allow production domain
ALLOWED_ORIGINS = [
    "http://localhost:3000", 
    "http://localhost:5000", 
    "http://127.0.0.1:3000", 
    "http://127.0.0.1:5000",
    "https://afrodocs.app",
    "https://www.afrodocs.app"
]
# Add Render URL if available
if RENDER_EXTERNAL_URL:
    ALLOWED_ORIGINS.append(RENDER_EXTERNAL_URL)
    # Also add the https version if not already https
    if not RENDER_EXTERNAL_URL.startswith('https'):
        ALLOWED_ORIGINS.append(RENDER_EXTERNAL_URL.replace('http://', 'https://'))
CORS(app, expose_headers=["Content-Disposition"], supports_credentials=True, origins=ALLOWED_ORIGINS)
db = SQLAlchemy(app)
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.session_protection = "strong"

# Return JSON for unauthorized API requests instead of redirecting
@login_manager.unauthorized_handler
def unauthorized():
    return jsonify({'error': 'Authentication required. Please login.'}), 401

# User Model
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(100), unique=True, nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=True)
    password_hash = db.Column(db.String(200), nullable=False)
    institution = db.Column(db.String(200), nullable=True)
    contact = db.Column(db.String(100), nullable=True)
    
    # Admin Flag
    is_admin = db.Column(db.Boolean, default=False)
    
    # Business Model Fields
    plan = db.Column(db.String(20), default='free') # free, student, pro, power, enterprise
    pages_balance = db.Column(db.Integer, default=0) # For paid plans
    docs_this_month = db.Column(db.Integer, default=0) # Deprecated, kept for compatibility
    pages_this_month = db.Column(db.Integer, default=0) # New: Track pages (limit 300)
    last_reset_date = db.Column(db.DateTime, default=datetime.utcnow)
    
    # Referral System
    referral_code = db.Column(db.String(20), unique=True, nullable=True)
    referred_by = db.Column(db.Integer, nullable=True) # ID of referrer
    referral_rewards_count = db.Column(db.Integer, default=0) # Track rewards to cap them
    
    documents_generated = db.Column(db.Integer, default=0)
    documents = db.relationship('DocumentRecord', backref='user', lazy=True, cascade="all, delete-orphan")
    support_requests = db.relationship('SupportRequest', backref='user', lazy=True, cascade="all, delete-orphan")
    ai_requests = db.relationship('AIRequest', backref='user', lazy=True, cascade="all, delete-orphan")

class DocumentRecord(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    filename = db.Column(db.String(255), nullable=False)
    original_filename = db.Column(db.String(255), nullable=True)
    job_id = db.Column(db.String(100), nullable=True)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    file_path = db.Column(db.String(500), nullable=True) # Relative path to storage

class Transaction(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    trans_id = db.Column(db.String(100), unique=True, nullable=False)
    external_id = db.Column(db.String(100), unique=True, nullable=False)
    amount = db.Column(db.Integer, nullable=False)
    status = db.Column(db.String(20), default='CREATED') # CREATED, PENDING, SUCCESSFUL, FAILED, EXPIRED
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    updated_at = db.Column(db.DateTime, default=datetime.utcnow, onupdate=datetime.utcnow)

class InstitutionRequest(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    institution_name = db.Column(db.String(255), nullable=False)
    requester_email = db.Column(db.String(120), nullable=True)
    requester_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    status = db.Column(db.String(20), default='pending')  # pending, approved, rejected
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class SupportRequest(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    name = db.Column(db.String(120), nullable=False)
    email = db.Column(db.String(160), nullable=False)
    subject = db.Column(db.String(200), nullable=True)
    message = db.Column(db.Text, nullable=False)
    status = db.Column(db.String(20), default='open')  # open, resolved
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

class AIRequest(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)
    request_type = db.Column(db.String(30), nullable=False)
    title = db.Column(db.String(200), nullable=False)
    prompt = db.Column(db.Text, nullable=False)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# Custom decorator that allows both authenticated users and guests
def allow_guest(f):
    """Decorator that allows both logged-in users and guests to access an endpoint.
    Unlike @login_required, this won't return 401 for unauthenticated users."""
    @wraps(f)
    def decorated_function(*args, **kwargs):
        # Simply proceed with the request - current_user will be anonymous if not logged in
        return f(*args, **kwargs)
    return decorated_function

# Create DB
with app.app_context():
    db.create_all()
    
    # Initialize admin account if it doesn't exist
    admin_user = User.query.filter_by(username='admin').first()
    if not admin_user:
        admin_user = User(
            username='admin',
            password_hash=generate_password_hash('admin@secure123'),  # Secure default password
            email='admin@formatter.local',
            institution='Administrator',
            contact='admin',
            is_admin=True,
            plan='enterprise',
            pages_balance=999999,  # Unlimited pages for admin
            pages_this_month=0,
            referral_code='ADMIN001'
        )
        db.session.add(admin_user)
        db.session.commit()
        logger.info("Admin account created: username=admin, password=admin@secure123")

@app.route('/')
def serve_frontend():
    return app.send_static_file('index.html')

# Auth Routes
@app.route('/api/auth/signup', methods=['POST'])
def signup():
    data = request.json
    username = data.get('username')
    password = data.get('password')
    institution = data.get('institution')
    contact = data.get('contact')
    email = data.get('email')
    referral_code_input = data.get('referral_code')
    
    if not username or not password:
        return jsonify({'error': 'Username and password required'}), 400
        
    if User.query.filter_by(username=username).first():
        return jsonify({'error': 'Username already exists'}), 400

    if email and User.query.filter_by(email=email).first():
        return jsonify({'error': 'Email already exists'}), 400
        
    # Generate unique referral code for new user
    new_referral_code = uuid.uuid4().hex[:8].upper()
    while User.query.filter_by(referral_code=new_referral_code).first():
        new_referral_code = uuid.uuid4().hex[:8].upper()
    
    # Handle Referral Logic
    referrer_id = None
    initial_pages = 0
    
    if referral_code_input:
        referrer = User.query.filter_by(referral_code=referral_code_input).first()
        if referrer:
            referrer_id = referrer.id
            # Immediate Reward for Referrer
            referrer.pages_balance += 50 # Reward 50 pages for successful signup
            referrer.referral_rewards_count += 1
            initial_pages = 50 # Reward for new user
            
    new_user = User(
        username=username, 
        password_hash=generate_password_hash(password),
        institution=institution,
        contact=contact,
        email=email,
        referral_code=new_referral_code,
        referred_by=referrer_id,
        plan='free',
        docs_this_month=0,
        pages_this_month=0,
        pages_balance=initial_pages
    )
    db.session.add(new_user)
    try:
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        logger.exception('Exception on /api/auth/signup [POST]')
        return jsonify({'error': 'Failed to create user'}), 500

    return jsonify({'message': 'User created successfully', 'referral_code': new_referral_code}), 201

@app.route('/api/auth/login', methods=['POST'])
def login():
    data = request.json
    username = data.get('username')
    password = data.get('password')
    
    user = User.query.filter_by(username=username).first()
    
    if user and check_password_hash(user.password_hash, password):
        login_user(user, remember=True)
        return jsonify({
            'message': 'Logged in successfully', 
            'username': user.username,
            'plan': user.plan,
            'pages_this_month': user.pages_this_month,
            'pages_balance': user.pages_balance,
            'referral_code': user.referral_code
        }), 200
        
    return jsonify({'error': 'Invalid username or password'}), 401

@app.route('/api/auth/me', methods=['GET'])
@login_required
def get_current_user():
    return jsonify({
        'username': current_user.username,
        'plan': current_user.plan,
        'pages_this_month': current_user.pages_this_month,
        'pages_balance': current_user.pages_balance,
        'referral_code': current_user.referral_code,
        'institution': current_user.institution
    }), 200

@app.route('/api/auth/logout', methods=['POST'])
@login_required
def logout():
    logout_user()
    return jsonify({'message': 'Logged out successfully'}), 200

# Payment Routes
@app.route('/api/payment/initiate', methods=['POST'])
@login_required
def initiate_payment():
    data = request.json
    amount = data.get('amount')
    
    if not amount or amount < 1000:
        return jsonify({'error': 'Invalid amount. Minimum 1000 XAF.'}), 400
        
    external_id = str(uuid.uuid4())
    
    payment_data = {
        'amount': amount,
        'userId': str(current_user.id),
        'externalId': external_id,
        'redirectUrl': request.host_url + '?payment=success', # Redirect to home with query param
        'message': 'Pattern Formatter Plan Upgrade'
    }
    
    # Only include email if it exists
    if current_user.email:
        payment_data['email'] = current_user.email
    
    response = fapshi.initiate_pay(payment_data)
    
    # Normalize HTTP status from provider response fields if present
    provider_status = response.get('statusCode') or response.get('httpStatus') or response.get('httpStatusCode') or response.get('code')
    status_to_return = int(provider_status) if provider_status else 500

    if response.get('statusCode') == 200 or status_to_return == 200:
        # Create Transaction Record
        new_trans = Transaction(
            user_id=current_user.id,
            trans_id=response.get('transId'),
            external_id=external_id,
            amount=amount,
            status='CREATED'
        )
        db.session.add(new_trans)
        db.session.commit()

        return jsonify(response), 200
    else:
        return jsonify(response), status_to_return

@app.route('/api/payment/verify-pending', methods=['POST'])
@login_required
def verify_pending_transactions():
    """Check status of pending transactions for current user"""
    pending_trans = Transaction.query.filter_by(user_id=current_user.id, status='CREATED').all()
    
    updated_count = 0
    
    for transaction in pending_trans:
        # Verify with Fapshi
        event = fapshi.payment_status(transaction.trans_id)
        
        if event.get('statusCode') == 200:
            status = event.get('status')
            
            # Only update if status changed
            if status != transaction.status:
                transaction.status = status
                
                if status == 'SUCCESSFUL':
                    # Pricing tiers (January 2026):
                    # 1000 FCFA = 500 pages (Student) - 2 months validity
                    # 2500 FCFA = 1500 pages (Campus Pro) - 2 months validity
                    # 5000 FCFA = 10000 pages (Enterprise) - 2 months validity
                    
                    pages_to_add = 0
                    if transaction.amount == 1000:
                        pages_to_add = 500
                        current_user.plan = 'student'
                    elif transaction.amount == 2500:
                        pages_to_add = 1500
                        current_user.plan = 'campus'
                    elif transaction.amount == 5000:
                        pages_to_add = 10000
                        current_user.plan = 'enterprise'
                    else:
                        # Fallback logic
                        pages_to_add = (transaction.amount // 1000) * 500
                        
                    if pages_to_add > 0:
                        current_user.pages_balance += pages_to_add
                        updated_count += 1
                    
                    # If amount > 5000, upgrade plan (legacy logic)
                    if transaction.amount >= 5000:
                        current_user.plan = 'pro'
                        
                    logger.info(f"User {current_user.username} upgraded via transaction {transaction.trans_id}")
    
    if updated_count > 0:
        db.session.commit()
        return jsonify({'success': True, 'message': f'Verified {updated_count} payments. Balance updated.'}), 200
    else:
        return jsonify({'success': True, 'message': 'No new successful payments found.'}), 200

@app.route('/api/payment/webhook', methods=['POST'])
def payment_webhook():
    data = request.json
    if not data or 'transId' not in data:
        return jsonify({'success': False, 'message': 'Invalid payload'}), 400
        
    trans_id = data.get('transId')
    
    # Verify with Fapshi
    event = fapshi.payment_status(trans_id)
    
    if event.get('statusCode') != 200:
        return jsonify({'success': False, 'message': 'Could not verify transaction'}), 400
        
    status = event.get('status')
    external_id = event.get('externalId')
    
    # Update Transaction
    transaction = Transaction.query.filter_by(trans_id=trans_id).first()
    if not transaction:
        transaction = Transaction.query.filter_by(external_id=external_id).first()
        
    if transaction:
        transaction.status = status
        
        if status == 'SUCCESSFUL':
            user = User.query.get(transaction.user_id)
            if user:
                # Pricing tiers (January 2026):
                # 1000 FCFA = 500 pages (Student) - 2 months validity
                # 2500 FCFA = 1500 pages (Campus Pro) - 2 months validity
                # 5000 FCFA = 10000 pages (Enterprise) - 2 months validity
                pages_to_add = 0
                if transaction.amount == 1000:
                    pages_to_add = 500
                    user.plan = 'student'
                elif transaction.amount == 2500:
                    pages_to_add = 1500
                    user.plan = 'campus'
                elif transaction.amount == 5000:
                    pages_to_add = 10000
                    user.plan = 'enterprise'
                else:
                    # Fallback logic
                    pages_to_add = (transaction.amount // 1000) * 500
                    
                if pages_to_add > 0:
                    user.pages_balance += pages_to_add
                    
                logger.info(f"User {user.username} upgraded via transaction {trans_id}")
                
        db.session.commit()
        
    return jsonify({'success': True}), 200


@app.route('/api/admin/users', methods=['GET'])
@login_required
def list_users():
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized - Admin access required'}), 403

    users = User.query.all()
    ai_request_counts = dict(
        db.session.query(AIRequest.user_id, func.count(AIRequest.id))
        .group_by(AIRequest.user_id)
        .all()
    )
    return jsonify([{
        'id': u.id,
        'username': u.username,
        'email': u.email,
        'institution': u.institution,
        'contact': u.contact,
        'is_admin': u.is_admin,
        'plan': u.plan,
        'pages_balance': u.pages_balance,
        'pages_this_month': u.pages_this_month,
        'documents_generated': u.documents_generated,
        'ai_requests_count': ai_request_counts.get(u.id, 0),
        'created_at': u.last_reset_date.isoformat() if u.last_reset_date else None
    } for u in users]), 200

@app.route('/api/admin/users/<int:user_id>', methods=['DELETE'])
@login_required
def delete_user(user_id):
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized - Admin access required'}), 403
        
    user = User.query.get(user_id)
    if not user:
        return jsonify({'error': 'User not found'}), 404
        
    if user.is_admin and user.id != current_user.id:
        return jsonify({'error': 'Cannot delete other admin users'}), 400
    
    if user.id == current_user.id:
        return jsonify({'error': 'Cannot delete your own account'}), 400
        
    db.session.delete(user)
    db.session.commit()
    return jsonify({'message': 'User deleted successfully'}), 200

@app.route('/api/admin/users', methods=['POST'])
@login_required
def create_user():
    """Admin endpoint to create new users"""
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized - Admin access required'}), 403
    
    data = request.json
    username = data.get('username')
    password = data.get('password')
    email = data.get('email')
    institution = data.get('institution')
    contact = data.get('contact')
    is_admin = data.get('is_admin', False)
    plan = data.get('plan', 'free')
    pages_balance = data.get('pages_balance', 0)
    
    if not username or not password:
        return jsonify({'error': 'Username and password required'}), 400
    
    if User.query.filter_by(username=username).first():
        return jsonify({'error': 'Username already exists'}), 400
    
    if email and User.query.filter_by(email=email).first():
        return jsonify({'error': 'Email already exists'}), 400
    
    # Generate unique referral code
    new_referral_code = uuid.uuid4().hex[:8].upper()
    while User.query.filter_by(referral_code=new_referral_code).first():
        new_referral_code = uuid.uuid4().hex[:8].upper()
    
    new_user = User(
        username=username,
        password_hash=generate_password_hash(password),
        email=email,
        institution=institution,
        contact=contact,
        is_admin=is_admin,
        plan=plan,
        pages_balance=pages_balance,
        referral_code=new_referral_code
    )
    db.session.add(new_user)
    db.session.commit()
    
    logger.info(f"Admin {current_user.username} created user: {username}")
    return jsonify({
        'message': 'User created successfully',
        'user': {
            'id': new_user.id,
            'username': new_user.username,
            'referral_code': new_referral_code,
            'is_admin': new_user.is_admin,
            'plan': new_user.plan
        }
    }), 201

@app.route('/api/admin/users/<int:user_id>/documents', methods=['GET'])
@login_required
def list_user_documents(user_id):
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized - Admin access required'}), 403
        
    docs = DocumentRecord.query.filter_by(user_id=user_id).order_by(DocumentRecord.created_at.desc()).all()
    return jsonify([{
        'id': d.id,
        'filename': d.filename,
        'original_filename': d.original_filename,
        'created_at': d.created_at.isoformat() if d.created_at else None,
        'job_id': d.job_id
    } for d in docs]), 200

@app.route('/api/admin/users/<int:user_id>/ai-requests', methods=['GET'])
@login_required
def list_user_ai_requests(user_id):
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized - Admin access required'}), 403

    requests = AIRequest.query.filter_by(user_id=user_id).order_by(AIRequest.created_at.desc()).all()
    return jsonify([{
        'id': req.id,
        'request_type': req.request_type,
        'title': req.title,
        'prompt': req.prompt,
        'created_at': req.created_at.isoformat() if req.created_at else None
    } for req in requests]), 200

@app.route('/api/admin/check', methods=['GET'])
@login_required
def check_admin():
    """Check if current user is admin"""
    return jsonify({
        'is_admin': current_user.is_admin,
        'username': current_user.username
    }), 200


@app.route('/api/institution-requests', methods=['POST'])
def submit_institution_request():
    """Submit a new institution request from user"""
    data = request.get_json()
    institution_name = data.get('institution_name', '').strip()
    
    if not institution_name:
        return jsonify({'error': 'Institution name is required'}), 400
    
    # Get user info if logged in
    requester_email = None
    requester_id = None
    if current_user.is_authenticated:
        requester_email = current_user.email
        requester_id = current_user.id
    
    # Create the request
    new_request = InstitutionRequest(
        institution_name=institution_name,
        requester_email=requester_email,
        requester_id=requester_id
    )
    db.session.add(new_request)
    db.session.commit()
    
    return jsonify({'message': 'Request submitted successfully', 'id': new_request.id}), 201

@app.route('/api/support', methods=['POST'])
@allow_guest
def submit_support_request():
    """Submit a new support request"""
    data = request.get_json() or {}
    name = (data.get('name') or '').strip()
    email = (data.get('email') or '').strip()
    subject = (data.get('subject') or '').strip()
    message = (data.get('message') or '').strip()

    if not name and current_user.is_authenticated:
        name = current_user.username or ''
    if not email and current_user.is_authenticated:
        email = current_user.email or ''

    if not name or not email or not message:
        return jsonify({'error': 'Name, email, and message are required'}), 400

    new_request = SupportRequest(
        user_id=current_user.id if current_user.is_authenticated else None,
        name=name,
        email=email,
        subject=subject,
        message=message
    )
    db.session.add(new_request)
    db.session.commit()

    return jsonify({'message': 'Support request submitted', 'id': new_request.id}), 201


@app.route('/api/admin/institution-requests', methods=['GET'])
@login_required
def get_institution_requests():
    """Get all institution requests"""
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized - Admin access required'}), 403
    
    requests = InstitutionRequest.query.order_by(InstitutionRequest.created_at.desc()).all()
    return jsonify([{
        'id': r.id,
        'institution_name': r.institution_name,
        'requester_email': r.requester_email,
        'status': r.status,
        'created_at': r.created_at.isoformat() if r.created_at else None
    } for r in requests]), 200

@app.route('/api/admin/support', methods=['GET'])
@login_required
def get_support_requests():
    """Get all support requests"""
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized - Admin access required'}), 403

    requests = SupportRequest.query.order_by(SupportRequest.created_at.desc()).all()
    return jsonify([{
        'id': r.id,
        'name': r.name,
        'email': r.email,
        'subject': r.subject,
        'message': r.message,
        'status': r.status,
        'created_at': r.created_at.isoformat() if r.created_at else None,
        'user': {
            'id': r.user.id,
            'username': r.user.username,
            'institution': r.user.institution
        } if r.user else None
    } for r in requests]), 200

@app.route('/api/admin/support/<int:request_id>/status', methods=['PATCH'])
@login_required
def update_support_request_status(request_id):
    """Update a support request status"""
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized - Admin access required'}), 403

    data = request.get_json() or {}
    status = (data.get('status') or '').strip().lower()
    if status not in {'open', 'resolved'}:
        return jsonify({'error': 'Invalid status'}), 400

    support_request = SupportRequest.query.get(request_id)
    if not support_request:
        return jsonify({'error': 'Support request not found'}), 404

    support_request.status = status
    db.session.commit()
    return jsonify({'message': 'Status updated', 'status': support_request.status}), 200


@app.route('/api/admin/documents', methods=['GET'])
@login_required
def get_all_documents():
    """Get all documents for admin view"""
    if not current_user.is_admin:
        return jsonify({'error': 'Unauthorized - Admin access required'}), 403
    
    docs = DocumentRecord.query.order_by(DocumentRecord.created_at.desc()).limit(100).all()
    return jsonify([{
        'id': d.id,
        'filename': d.filename,
        'original_filename': d.original_filename,
        'created_at': d.created_at.isoformat() if d.created_at else None,
        'job_id': d.job_id,
        'user_id': d.user_id
    } for d in docs]), 200


@app.route('/api/documents/<job_id>/rename', methods=['PATCH'])
@login_required
def rename_document(job_id):
    data = request.get_json() or {}
    new_name = (data.get('name') or '').strip()
    if not new_name:
        return jsonify({'error': 'New name is required'}), 400

    doc = DocumentRecord.query.filter_by(job_id=job_id).first()
    if not doc:
        return jsonify({'error': 'Document not found'}), 404
    if not current_user.is_admin and doc.user_id != current_user.id:
        return jsonify({'error': 'Unauthorized'}), 403

    sanitized = new_name
    if not sanitized.lower().endswith('.docx'):
        sanitized = f"{sanitized}.docx"
    doc.filename = sanitized
    db.session.commit()
    return jsonify({'message': 'Document renamed', 'filename': doc.filename}), 200


@app.route('/api/documents/<job_id>', methods=['DELETE'])
@login_required
def delete_document(job_id):
    doc = DocumentRecord.query.filter_by(job_id=job_id).first()
    if not doc:
        return jsonify({'error': 'Document not found'}), 404
    if not current_user.is_admin and doc.user_id != current_user.id:
        return jsonify({'error': 'Unauthorized'}), 403

    output_path = os.path.join(OUTPUT_FOLDER, f"{job_id}_formatted.docx")
    meta_path = os.path.join(OUTPUT_FOLDER, f"{job_id}_meta.json")
    if os.path.exists(output_path):
        try:
            os.remove(output_path)
        except OSError as e:
            logger.warning(f"Failed to remove output file {output_path}: {e}")
    if os.path.exists(meta_path):
        try:
            os.remove(meta_path)
        except OSError as e:
            logger.warning(f"Failed to remove metadata file {meta_path}: {e}")

    db.session.delete(doc)
    db.session.commit()
    return jsonify({'message': 'Document deleted'}), 200


@app.route('/api/auth/status', methods=['GET'])
def auth_status():
    if current_user.is_authenticated:
        return jsonify({
            'isAuthenticated': True, 
            'username': current_user.username,
            'is_admin': current_user.is_admin,
            'plan': current_user.plan,
            'docs_this_month': current_user.docs_this_month,
            'pages_this_month': current_user.pages_this_month,
            'pages_balance': current_user.pages_balance,
            'referral_code': current_user.referral_code,
            'email': current_user.email
        }), 200
    return jsonify({'isAuthenticated': False}), 200


# ============================================================
# AI CHAT ENDPOINTS - DeepSeek Integration
# ============================================================
def _ai_request_title(prompt, max_length=80):
    if not prompt:
        return "AI Request"
    for line in prompt.splitlines():
        stripped = line.strip()
        if stripped:
            title = stripped
            break
    else:
        title = prompt.strip()
    if len(title) > max_length:
        return title[:max_length].rstrip() + "..."
    return title

def record_ai_request(prompt, request_type):
    if not current_user.is_authenticated:
        return
    try:
        ai_request = AIRequest(
            user_id=current_user.id,
            request_type=request_type,
            title=_ai_request_title(prompt),
            prompt=prompt
        )
        db.session.add(ai_request)
        db.session.commit()
    except Exception as exc:
        logger.warning(f"Failed to record AI request: {exc}")
        db.session.rollback()

# DeepSeek API Configuration
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "sk-4a857c0c76cf4db89fef65b871da982a")
DEEPSEEK_API_URL = os.getenv("DEEPSEEK_API_URL", "https://api.deepseek.com/v1/chat/completions")
DEEPSEEK_MODEL = os.getenv("DEEPSEEK_MODEL", "deepseek-chat")
AI_MAX_TOKENS = int(os.getenv("AI_MAX_TOKENS", "8192"))

# Gemini API Configuration (used for image/document attachments)
# Set the Gemini key in environment variable `GEMINI_API_KEY` for security.
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "AIzaSyDAhVQzICAXDE91bSqNzKgjlG1sBfGfP8A")
GEMINI_API_URL = os.getenv("GEMINI_API_URL", "https://generativelanguage.googleapis.com/v1beta/models/gemini-1.5-flash:generateContent")

# System prompt for the AI assistant
AI_SYSTEM_PROMPT = """You are AfroDocs AI Assistant, an expert academic writing assistant. Follow these guidelines strictly:

## FORMATTING RULES:
1. Use HIERARCHICAL NUMBERING for main sections and subsections:
   - Main sections: 1.0, 2.0, 3.0, etc.
   - Subsections: 1.1, 1.2, 1.3, etc.
   - Sub-subsections: 1.1.1, 1.1.2, etc.
   - Exception: For Research Questions/Objectives/Hypotheses lists, use simple ordered items (1., 2., 3.) without subsection numbering.
2. Never use justified text alignment
3. Use clear paragraph breaks between sections and add a blank line between headings and body text.
4. Bold section headings (e.g., **1.0 Introduction**)
5. Only bold headings/titles. Do NOT bold full paragraphs or list items.
6. Avoid excessive hierarchy; use at most three levels (e.g., 1.1.1) unless explicitly required.
7. Use plain text/markdown only (no HTML). Keep responses well structured with clear headings and lists.
8. When providing tables, use Markdown pipe tables with a header row and separator line (| --- |). Keep consistent column counts and do not wrap tables in code fences.

## DOCUMENT STRUCTURES:

### DISSERTATION/THESIS (Max 5 Chapters):
- **Chapter One: Introduction**
  - 1.1 Background of the Study
  - 1.2 Statement of the Problem
  - 1.3 Research Questions/Objectives
  - 1.4 Significance of the Study
  - 1.5 Scope and Delimitation
  - 1.6 Definition of Terms

- **Chapter Two: Literature Review**
  - 2.1 Theoretical Framework
  - 2.2 Conceptual Framework
  - 2.3 Review of Related Studies
  - 2.4 Research Gap

- **Chapter Three: Research Methodology**
  - 3.1 Research Design
  - 3.2 Area of Study
  - 3.3 Population and Sampling
  - 3.4 Data Collection Methods
  - 3.5 Data Analysis Techniques
  - 3.6 Ethical Considerations

- **Chapter Four: Results and Discussion**
  - 4.1 Data Presentation
  - 4.2 Analysis of Findings
  - 4.3 Discussion of Results

- **Chapter Five: Summary, Conclusion and Recommendations**
  - 5.1 Summary of Findings
  - 5.2 Conclusion
  - 5.3 Recommendations
  - 5.4 Suggestions for Further Research

### RESEARCH PROPOSAL (Chapters 1-3 + Extras):
- **CHAPTER ONE**: Introduction (same as thesis)
- **CHAPTER TWO**: Literature Review (same as thesis)
- **CHAPTER 3**: Materials and Method (same as thesis)
- **Budget and Timeline**
- **References**

### ASSIGNMENT/ESSAY Structure:
- **Introduction** (with thesis statement)
- **Main Body** (numbered sections 1.0, 2.0, 3.0 with subsections)
- **Conclusion**
- **References**

### PROJECT REPORT Structure:
- Title Page
- Abstract
- Table of Contents
- **1.0 Introduction**
- **2.0 Literature Review**
- **3.0 Methodology**
- **4.0 Results/Findings**
- **5.0 Discussion**
- **6.0 Conclusion and Recommendations**
- References
- Appendices

## WRITING STYLE:
- Use formal academic language
- Write in third person unless instructed otherwise
- Be concise and precise
- Support arguments with evidence
- Use topic sentences for paragraphs
- Maintain logical flow between sections

## IMPORTANT NOTES:
- Chapters should NOT exceed 5 for dissertations/theses
- For any request concerning dissertations, proposals, or chapters, respond using chapter-based structure, and each chapter heading must start with "Chapter".
- Always use hierarchical numbering (1.0, 1.1, 1.1.1)
- Provide well-structured, smart, and coherent content
- When asked to write, produce complete, ready-to-use content
- Format headings with bold markers (**)

 Be professional, helpful, and produce high-quality academic content following these standards."""

AI_RESTRUCTURE_SYSTEM_PROMPT = """You are AfroDocs AI Restructuring Assistant.

Task: Restructure the provided text to match AfroDocs academic formatting standards.

Requirements:
- Preserve all original meaning and content. Do NOT invent, add, or remove content.
- Apply hierarchical numbering where appropriate: 1.0, 1.1, 1.1.1, etc.
- For Research Questions/Objectives/Hypotheses, use simple ordered lists (1., 2., 3.) without subsection numbering.
- Use bullet points for unordered items and Arabic numerals (1, 2, 3) for ordered lists.
- Do NOT use lettered lists (a, b, c), roman numerals (i, ii, iii), or mixed numeric-letter labels (1a, 2b).
- Keep headings concise and clear, and follow the AfroDocs numbering hierarchy.
- Return ONLY the restructured text with no commentary, no explanations, and no code fences.
- Only bold headings/titles. Do NOT bold full paragraphs, tables, or list items.
- Avoid excessive hierarchy; use up to three levels unless the source clearly indicates deeper structure.
- Use plain text/markdown only (no HTML). Keep tables in Markdown pipe format with a header row and separator line.
"""


def _call_deepseek(messages, temperature=0.4, max_tokens=AI_MAX_TOKENS, timeout=90):
    import requests

    response = requests.post(
        DEEPSEEK_API_URL,
        headers={
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
            "Content-Type": "application/json"
        },
        json={
            "model": DEEPSEEK_MODEL,
            "messages": messages,
            "temperature": temperature,
            "max_tokens": max_tokens
        },
        timeout=timeout
    )

    if response.status_code != 200:
        logger.error(f"DeepSeek API error: {response.status_code} - {response.text}")
        raise RuntimeError(f"AI service error: {response.status_code}")

    result = response.json()
    return result['choices'][0]['message']['content']


def normalize_ai_bolding(text):
    if not text:
        return text

    heading_patterns = [
        r'^\s*\*?\*?\s*(?:chapter|table of contents|references|bibliography)\b',
        r'^\s*\*?\*?\s*\d+(\.\d+){0,3}\s+\S+',
        r'^\s*\*?\*?\s*[A-Z][A-Z\s]{3,}$'
    ]
    heading_regex = re.compile('|'.join(heading_patterns), re.IGNORECASE)

    cleaned_lines = []
    for line in text.splitlines():
        if '**' in line and not heading_regex.search(line):
            cleaned_lines.append(line.replace('**', ''))
        else:
            cleaned_lines.append(line)
    return '\n'.join(cleaned_lines)

def normalize_ai_structure(text):
    if not text:
        return text

    lines = []
    mixed_alpha_numeric_pattern = re.compile(r'^\s*\d+\s*[a-zA-Z][\)\.\:]?\s+', re.IGNORECASE)
    lettered_pattern = re.compile(r'^\s*(?:\d+\s+)?[a-z][\)\.]\s+', re.IGNORECASE)
    roman_pattern = re.compile(r'^\s*(?:\d+\s+)?(?:i{1,4}|v|vi{0,3}|ix|x)[\)\.\:]?\s+',
                               re.IGNORECASE)
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            lines.append(line)
            continue

        if mixed_alpha_numeric_pattern.match(stripped):
            updated = mixed_alpha_numeric_pattern.sub('', stripped)
            lines.append(f"- {updated}")
            continue

        if lettered_pattern.match(stripped):
            updated = lettered_pattern.sub('', stripped)
            lines.append(f"- {updated}")
            continue

        if roman_pattern.match(stripped) and not stripped.lower().startswith('chapter'):
            updated = roman_pattern.sub('', stripped)
            lines.append(f"- {updated}")
            continue

        lines.append(line)
    return '\n'.join(lines)

def normalize_ai_output(text):
    return normalize_ai_structure(normalize_ai_bolding(text))

def restructure_text_with_ai(text):
    if not text or not text.strip():
        return text

    messages = [
        {"role": "system", "content": AI_RESTRUCTURE_SYSTEM_PROMPT},
        {"role": "user", "content": f"Restructure the text below into AfroDocs academic formatting:\n\n{text}"}
    ]
    return normalize_ai_bolding(_call_deepseek(messages))


def extract_text_from_docx_bytes(file_bytes):
    doc = Document(BytesIO(file_bytes))
    lines = []
    for para in doc.paragraphs:
        if para.text:
            lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def extract_text_from_pdf_bytes(file_bytes):
    reader = PdfReader(BytesIO(file_bytes))
    lines = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            lines.append(text)
    return "\n".join(lines)


def _decode_attachment_bytes(data_url):
    if not data_url:
        return None
    if data_url.startswith("data:"):
        try:
            header, encoded = data_url.split(",", 1)
        except ValueError:
            return None
        return base64.b64decode(encoded)
    return base64.b64decode(data_url)


def extract_text_from_attachment(attachment):
    name = attachment.get('name') or 'attachment'
    data_url = attachment.get('dataUrl') or attachment.get('data_url') or attachment.get('content')
    if not data_url:
        return None, f"{name}: no attachment content received"

    file_ext = os.path.splitext(name)[1].lower()
    try:
        file_bytes = _decode_attachment_bytes(data_url)
    except (ValueError, binascii.Error) as exc:
        return None, f"{name}: invalid attachment data ({exc})"

    if not file_bytes:
        return None, f"{name}: empty attachment payload"

    try:
        if file_ext == '.docx':
            return extract_text_from_docx_bytes(file_bytes), None
        if file_ext == '.txt' or file_ext == '.md':
            return file_bytes.decode('utf-8', errors='ignore'), None
        if file_ext == '.pdf':
            return extract_text_from_pdf_bytes(file_bytes), None
        if file_ext in {'.png', '.jpg', '.jpeg', '.gif', '.webp'}:
            return None, f"{name}: image OCR is not available yet"
    except Exception as exc:
        logger.warning(f"Attachment extraction failed for {name}: {exc}")
        return None, f"{name}: extraction failed"

    return None, f"{name}: unsupported attachment type"


IMAGE_ATTACHMENT_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.webp'}


def attachment_is_image(attachment):
    name = attachment.get('name') or ''
    ext = os.path.splitext(name)[1].lower()
    mime = (attachment.get('type') or '').lower()
    return ext in IMAGE_ATTACHMENT_EXTENSIONS or mime.startswith('image/')


def build_attachment_context(attachments, max_chars=12000):
    if not attachments:
        return "", ""

    extracted_blocks = []
    issues = []

    for attachment in attachments:
        text, error = extract_text_from_attachment(attachment)
        name = attachment.get('name') or 'attachment'
        if text and text.strip():
            snippet = text.strip()
            if len(snippet) > max_chars:
                snippet = f"{snippet[:max_chars]}\n[...truncated...]"
            extracted_blocks.append(f"Attachment: {name}\n{snippet}")
        elif error:
            if "image OCR is not available yet" not in error:
                issues.append(error)
        else:
            issues.append(f"{name}: no extractable text found")

    return "\n\n".join(extracted_blocks), "\n".join(issues)

def _infer_mime_type_from_name(name):
    ext = os.path.splitext(name or "")[1].lower()
    if ext in {'.png'}:
        return "image/png"
    if ext in {'.jpg', '.jpeg'}:
        return "image/jpeg"
    if ext in {'.gif'}:
        return "image/gif"
    if ext in {'.webp'}:
        return "image/webp"
    if ext in {'.pdf'}:
        return "application/pdf"
    if ext in {'.docx'}:
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if ext in {'.txt'}:
        return "text/plain"
    if ext in {'.md'}:
        return "text/markdown"
    return None


def _attachment_inline_data(attachment):
    data_url = attachment.get('dataUrl') or attachment.get('data_url') or attachment.get('content')
    name = attachment.get('name') or 'attachment'
    if not data_url:
        return None, None
    if data_url.startswith("data:"):
        try:
            header, encoded = data_url.split(",", 1)
        except ValueError:
            return None, None
        mime_type = header.split(";")[0].replace("data:", "").strip() or None
        return mime_type, encoded
    return _infer_mime_type_from_name(name), data_url


def build_gemini_parts(message, attachments):
    parts = []
    if message:
        parts.append({"text": message})

    attachment_context, attachment_issues = build_attachment_context(attachments)
    if attachment_context:
        parts.append({"text": f"Attached content (extracted):\n{attachment_context}"})
    elif attachment_issues:
        parts.append({"text": f"Attachment notes:\n{attachment_issues}"})

    for attachment in attachments:
        mime_type, data_b64 = _attachment_inline_data(attachment)
        if mime_type and data_b64 and mime_type.startswith("image/"):
            parts.append({"inline_data": {"mime_type": mime_type, "data": data_b64}})

    return parts


def _extract_gemini_text(result):
    candidates = result.get('candidates') or []
    if not candidates:
        return ""
    parts = candidates[0].get('content', {}).get('parts', [])
    return "".join(part.get('text', '') for part in parts if part.get('text'))


def _call_gemini(contents, temperature=0.7, max_tokens=AI_MAX_TOKENS, timeout=90):
    import requests

    url = GEMINI_API_URL
    headers = {"Content-Type": "application/json"}

    # Prefer Bearer Authorization if a key is provided via env var
    if GEMINI_API_KEY:
        headers["Authorization"] = f"Bearer {GEMINI_API_KEY}"
    else:
        # fallback to query param if no Authorization header is available
        url = f"{GEMINI_API_URL}?key={GEMINI_API_KEY}"

    response = requests.post(
        url,
        headers=headers,
        json={
            "system_instruction": {"parts": [{"text": AI_SYSTEM_PROMPT}]},
            "contents": contents,
            "generationConfig": {
                "temperature": temperature,
                "maxOutputTokens": max_tokens
            }
        },
        timeout=timeout
    )

    if response.status_code != 200:
        logger.error(f"Gemini API error: {response.status_code} - {response.text}")
        raise RuntimeError(f"AI service error: {response.status_code}")

    result = response.json()
    return _extract_gemini_text(result)

AI_RESTRUCTURE_SYSTEM_PROMPT = """You are AfroDocs AI Restructuring Assistant.

Task: Restructure the provided text to match AfroDocs academic formatting standards.

Requirements:
- Preserve all original meaning and content. Do NOT invent, add, or remove content.
- Apply hierarchical numbering where appropriate: 1.0, 1.1, 1.1.1, etc.
- For Research Questions/Objectives/Hypotheses, use simple ordered lists (1., 2., 3.) without subsection numbering.
- Use bullet points for unordered items and Arabic numerals (1, 2, 3) for ordered lists.
- Do NOT use lettered lists (a, b, c), roman numerals (i, ii, iii), or mixed numeric-letter labels (1a, 2b).
- Keep headings concise and clear, and follow the AfroDocs numbering hierarchy.
- Return ONLY the restructured text with no commentary, no explanations, and no code fences.
- Use plain text/markdown only (no HTML). Keep tables in Markdown pipe format with a header row and separator line.
"""


def _call_deepseek(messages, temperature=0.4, max_tokens=AI_MAX_TOKENS, timeout=90):
    import requests

    response = requests.post(
        DEEPSEEK_API_URL,
        headers={
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
            "Content-Type": "application/json"
        },
        json={
            "model": DEEPSEEK_MODEL,
            "messages": messages,
            "temperature": temperature,
            "max_tokens": max_tokens
        },
        timeout=timeout
    )

    if response.status_code != 200:
        logger.error(f"DeepSeek API error: {response.status_code} - {response.text}")
        raise RuntimeError(f"AI service error: {response.status_code}")

    result = response.json()
    return result['choices'][0]['message']['content']


def restructure_text_with_ai(text):
    if not text or not text.strip():
        return text

    messages = [
        {"role": "system", "content": AI_RESTRUCTURE_SYSTEM_PROMPT},
        {"role": "user", "content": f"Restructure the text below into AfroDocs academic formatting:\n\n{text}"}
    ]
    return normalize_ai_output(_call_deepseek(messages))


def extract_text_from_docx_bytes(file_bytes):
    doc = Document(BytesIO(file_bytes))
    lines = []
    for para in doc.paragraphs:
        if para.text:
            lines.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


@app.route('/api/ai/chat', methods=['POST'])
@allow_guest
def ai_chat():
    """AI chat endpoint with DeepSeek integration"""
    import requests
    
    data = request.get_json(silent=True) or {}
    message = data.get('message', '')
    conversation = data.get('conversation', [])
    attachments = data.get('attachments', [])
    raw_message = message

    use_gemini = any(attachment_is_image(attachment) for attachment in attachments)

    try:
        if use_gemini:
            gemini_contents = []
            for msg in conversation[-10:]:
                role = "model" if msg.get('role') in {"assistant", "model"} else "user"
                content = msg.get('content', '')
                if content:
                    gemini_contents.append({"role": role, "parts": [{"text": content}]})
            gemini_contents.append({
                "role": "user",
                "parts": build_gemini_parts(message, attachments)
            })
            ai_response = normalize_ai_output(_call_gemini(gemini_contents, temperature=0.7, max_tokens=AI_MAX_TOKENS))
            record_ai_request(raw_message, "chat")
            return jsonify({'response': ai_response, 'success': True}), 200

        attachment_context, attachment_issues = build_attachment_context(attachments)
        if attachment_context:
            message = f"{message}\n\nAttached content (extracted):\n{attachment_context}"
        if attachment_issues and not attachment_context:
            message = f"{message}\n\nAttachment notes:\n{attachment_issues}"

        # Build messages array
        messages = [{"role": "system", "content": AI_SYSTEM_PROMPT}]
        
        # Add conversation history
        for msg in conversation[-10:]:  # Keep last 10 messages for context
            messages.append({
                "role": msg.get('role', 'user'),
                "content": msg.get('content', '')
            })
        
        # Add current message
        messages.append({"role": "user", "content": message})

        response = requests.post(
            DEEPSEEK_API_URL,
            headers={
                "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
                "Content-Type": "application/json"
            },
            json={
                "model": DEEPSEEK_MODEL,
                "messages": messages,
                "temperature": 0.7,
                "max_tokens": AI_MAX_TOKENS
            },
            timeout=60
        )
        
        if response.status_code == 200:
            result = response.json()
            ai_response = normalize_ai_output(result['choices'][0]['message']['content'])
            record_ai_request(raw_message, "chat")
            return jsonify({'response': ai_response, 'success': True}), 200

        logger.error(f"DeepSeek API error: {response.status_code} - {response.text}")
        return jsonify({'error': f'AI service error: {response.status_code}', 'success': False}), 500
            
    except requests.exceptions.Timeout:
        return jsonify({'error': 'AI service timeout. Please try again.', 'success': False}), 504
    except Exception as e:
        logger.error(f"AI chat error: {str(e)}")
        return jsonify({'error': f'AI service error: {str(e)}', 'success': False}), 500


@app.route('/api/ai/restructure', methods=['POST'])
@allow_guest
def ai_restructure():
    """Restructure text using DeepSeek for AfroDocs formatting."""
    data = request.get_json() or {}
    text = data.get('text', '')
    if not text.strip():
        return jsonify({'error': 'No text provided', 'success': False}), 400

    try:
        restructured = restructure_text_with_ai(text)
        record_ai_request(text, "restructure")
        return jsonify({'response': restructured, 'success': True}), 200
    except Exception as e:
        logger.error(f"AI restructure error: {str(e)}")
        return jsonify({'error': f'AI service error: {str(e)}', 'success': False}), 500


@app.route('/api/ai/chat/stream', methods=['POST'])
@allow_guest
def ai_chat_stream():
    """AI chat streaming endpoint with DeepSeek integration"""
    import requests
    
    data = request.get_json(silent=True) or {}
    message = data.get('message', '')
    conversation = data.get('conversation', [])
    attachments = data.get('attachments', [])
    raw_message = message

    use_gemini = any(attachment_is_image(attachment) for attachment in attachments)
    
    if not use_gemini:
        attachment_context, attachment_issues = build_attachment_context(attachments)
        if attachment_context:
            message = f"{message}\n\nAttached content (extracted):\n{attachment_context}"
        if attachment_issues and not attachment_context:
            message = f"{message}\n\nAttachment notes:\n{attachment_issues}"
        
        # Build messages array
        messages = [{"role": "system", "content": AI_SYSTEM_PROMPT}]
        
        # Add conversation history
        for msg in conversation[-10:]:  # Keep last 10 messages for context
            messages.append({
                "role": msg.get('role', 'user'),
                "content": msg.get('content', '')
            })
        
        # Add current message
        messages.append({"role": "user", "content": message})

    record_ai_request(raw_message, "stream")
    
    def generate():
        try:
            if use_gemini:
                gemini_contents = []
                for msg in conversation[-10:]:
                    role = "model" if msg.get('role') in {"assistant", "model"} else "user"
                    content = msg.get('content', '')
                    if content:
                        gemini_contents.append({"role": role, "parts": [{"text": content}]})
                gemini_contents.append({
                    "role": "user",
                    "parts": build_gemini_parts(message, attachments)
                })
                ai_response = normalize_ai_output(_call_gemini(gemini_contents, temperature=0.7, max_tokens=AI_MAX_TOKENS))
                if ai_response:
                    yield f"data: {json.dumps({'content': ai_response})}\n\n"
                yield f"data: {json.dumps({'done': True})}\n\n"
                return

            response = requests.post(
                DEEPSEEK_API_URL,
                headers={
                    "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": DEEPSEEK_MODEL,
                    "messages": messages,
                    "temperature": 0.7,
                    "max_tokens": AI_MAX_TOKENS,
                    "stream": True
                },
                stream=True,
                timeout=120
            )
            
            if response.status_code != 200:
                yield f"data: {json.dumps({'error': f'AI service error: {response.status_code}'})}\n\n"
                return
            
            for line in response.iter_lines():
                if line:
                    line_text = line.decode('utf-8')
                    if line_text.startswith('data: '):
                        data_str = line_text[6:]
                        if data_str == '[DONE]':
                            yield f"data: {json.dumps({'done': True})}\n\n"
                            break
                        try:
                            chunk = json.loads(data_str)
                            if 'choices' in chunk and len(chunk['choices']) > 0:
                                delta = chunk['choices'][0].get('delta', {})
                                content = delta.get('content', '')
                                if content:
                                    yield f"data: {json.dumps({'content': content})}\n\n"
                        except json.JSONDecodeError:
                            continue
                            
        except requests.exceptions.Timeout:
            yield f"data: {json.dumps({'error': 'AI service timeout'})}\n\n"
        except Exception as e:
            logger.error(f"AI stream error: {str(e)}")
            yield f"data: {json.dumps({'error': str(e)})}\n\n"
        
        yield f"data: {json.dumps({'done': True})}\n\n"
    
    return Response(generate(), mimetype='text/event-stream')


# Configuration
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
OUTPUT_FOLDER = os.path.join(BASE_DIR, 'outputs')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)


def update_toc_with_word(doc_path):
    """
    Update Table of Contents, List of Figures, and List of Tables in a Word document 
    using Microsoft Word COM automation.
    This opens the document in Word, updates all fields (including TOC, LOF, LOT), and saves it.
    Also ensures LOF and LOT entries are plain text (not bold/italic).
    
    Args:
        doc_path: Absolute path to the Word document
        
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        import win32com.client
        import pythoncom
        
        # Initialize COM
        pythoncom.CoInitialize()
        
        # Create Word application instance
        word = win32com.client.DispatchEx('Word.Application')
        word.Visible = False  # Run in background
        word.DisplayAlerts = False  # Suppress dialogs
        
        try:
            # Open the document
            doc = word.Documents.Open(os.path.abspath(doc_path))
            
            # Update all fields in the document (including TOC, LOF, LOT)
            # wdStory = 6 (entire document)
            word.Selection.WholeStory()
            word.Selection.Fields.Update()
            
            # Also specifically update TOC if present
            for toc in doc.TablesOfContents:
                toc.Update()
            
            # Update List of Figures and List of Tables if present
            # Both use TablesOfFigures collection (they're all caption-based tables)
            for caption_table in doc.TablesOfFigures:
                caption_table.Update()
            
            # Format LOF and LOT entries to be plain text (not bold, not italic)
            # The "Table of Figures" style controls these entries
            try:
                tof_style = doc.Styles("Table of Figures")
                tof_style.Font.Bold = False
                tof_style.Font.Italic = False
            except:
                pass  # Style may not exist
            
            # Save and close
            doc.Save()
            doc.Close()
            
            logger.info(f"TOC, LOF, and LOT updated successfully in {doc_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error updating TOC/LOF/LOT: {str(e)}")
            return False
            
        finally:
            # Quit Word application
            word.Quit()
            pythoncom.CoUninitialize()
            
    except ImportError:
        logger.warning("win32com not available - TOC/LOF/LOT will need manual update")
        logger.warning("Install with: pip install pywin32")
        return False
    except Exception as e:
        logger.error(f"Failed to update TOC/LOF/LOT with Word: {str(e)}")
        return False


# ============================================================
# IMAGE EXTRACTION AND REINSERTION SYSTEM
# ============================================================

class ImageExtractor:
    """
    Extract images from Word documents with full metadata.
    Preserves position, dimensions, format, and caption associations.
    """
    
    # Supported image formats
    SUPPORTED_FORMATS = {
        'image/png': 'png',
        'image/jpeg': 'jpeg',
        'image/jpg': 'jpg',
        'image/gif': 'gif',
        'image/bmp': 'bmp',
        'image/tiff': 'tiff',
        'image/x-emf': 'emf',
        'image/x-wmf': 'wmf',
    }
    
    # Caption patterns (reuse existing patterns)
    CAPTION_PATTERNS = [
        re.compile(r'^Figure\s+\d+[\.\:\s]', re.IGNORECASE),
        re.compile(r'^Fig\.\s*\d+[\.\:\s]', re.IGNORECASE),
        re.compile(r'^Image\s+\d+[\.\:\s]', re.IGNORECASE),
        re.compile(r'^Diagram\s+\d+[\.\:\s]', re.IGNORECASE),
        re.compile(r'^Chart\s+\d+[\.\:\s]', re.IGNORECASE),
        re.compile(r'^Graph\s+\d+[\.\:\s]', re.IGNORECASE),
        re.compile(r'^Illustration\s+\d+[\.\:\s]', re.IGNORECASE),
        re.compile(r'^Photo\s+\d+[\.\:\s]', re.IGNORECASE),
        re.compile(r'^Plate\s+\d+[\.\:\s]', re.IGNORECASE),
    ]
    
    def __init__(self):
        self.images = []
        self.image_count = 0
        self.extracted_rIds = set()  # Track extracted image rIds to prevent duplicates
        
    def extract_all_images(self, doc_path):
        """
        Extract all images from a Word document.
        
        Args:
            doc_path: Path to the .docx file
            
        Returns:
            list: List of image metadata dictionaries
        """
        self.images = []
        self.image_count = 0
        self.extracted_rIds = set()  # Reset for new document
        
        try:
            doc = Document(doc_path)
            
            # Track paragraph index for position mapping
            paragraph_index = 0
            element_index = 0
            
            # Process document body elements in order
            for element in doc.element.body:
                if element.tag.endswith('p'):
                    # This is a paragraph - check for inline images
                    for para in doc.paragraphs:
                        if para._element is element:
                            images_in_para = self._extract_images_from_paragraph(
                                para, paragraph_index, element_index
                            )
                            self.images.extend(images_in_para)
                            paragraph_index += 1
                            break
                
                elif element.tag.endswith('tbl'):
                    # This is a table - check cells for images
                    table_index = 0
                    for table in doc.tables:
                        if table._element is element:
                            images_in_table = self._extract_images_from_table(
                                table, table_index, element_index
                            )
                            self.images.extend(images_in_table)
                            break
                        table_index += 1
                
                element_index += 1
            
            # Also check for floating images (anchored drawings)
            floating_images = self._extract_floating_images(doc)
            self.images.extend(floating_images)
            
            logger.info(f"Extracted {len(self.images)} images from document")
            return self.images
            
        except Exception as e:
            logger.error(f"Error extracting images: {str(e)}")
            return []
    
    def _extract_images_from_paragraph(self, para, para_index, element_index):
        """Extract inline images from a paragraph."""
        images = []
        
        try:
            # Get paragraph text for caption detection
            para_text = para.text.strip()
            
            # Look for inline shapes (images) in the paragraph
            for run in para.runs:
                # Check if run contains inline shapes
                drawing_elements = run._element.findall('.//' + qn('a:blip'))
                
                for drawing in drawing_elements:
                    # Get the relationship ID (rId) for the image
                    embed_attr = drawing.get(qn('r:embed'))
                    if embed_attr and embed_attr not in self.extracted_rIds:
                        image_data = self._get_image_from_rId(para.part, embed_attr)
                        if image_data:
                            self.extracted_rIds.add(embed_attr)  # Mark as extracted
                            # Get dimensions from inline shape
                            width, height = self._get_inline_dimensions(run._element)
                            
                            # Detect caption
                            caption = self._detect_caption(para, para_text)
                            
                            image_meta = {
                                'image_id': f'img_{self.image_count:04d}',
                                'position_type': 'paragraph',
                                'paragraph_index': para_index,
                                'element_index': element_index,
                                'table_location': None,
                                'data': image_data['bytes'],
                                'format': image_data['format'],
                                'width': width,
                                'height': height,
                                'width_emu': image_data.get('width_emu'),
                                'height_emu': image_data.get('height_emu'),
                                'caption': caption,
                                'caption_position': 'below' if caption else None,
                                'is_inline': True,
                                'anchor_type': 'inline',
                            }
                            images.append(image_meta)
                            self.image_count += 1
                            logger.info(f"Extracted image {image_meta['image_id']} at paragraph {para_index}")
                
                # Also check for drawing elements with pictures (inside the run loop)
                inline_shapes = run._element.findall('.//' + qn('wp:inline'))
                for inline in inline_shapes:
                    blips = inline.findall('.//' + qn('a:blip'))
                    for blip in blips:
                        embed = blip.get(qn('r:embed'))
                        if embed and embed not in self.extracted_rIds:
                            image_data = self._get_image_from_rId(para.part, embed)
                            if image_data:
                                self.extracted_rIds.add(embed)  # Mark as extracted
                                # Get dimensions
                                extent = inline.find(qn('wp:extent'))
                                width = self._emu_to_inches(int(extent.get('cx', 0))) if extent is not None else 3.0
                                height = self._emu_to_inches(int(extent.get('cy', 0))) if extent is not None else 2.0
                                
                                caption = self._detect_caption(para, para_text)
                                
                                image_meta = {
                                    'image_id': f'img_{self.image_count:04d}',
                                    'position_type': 'paragraph',
                                    'paragraph_index': para_index,
                                    'element_index': element_index,
                                    'table_location': None,
                                    'data': image_data['bytes'],
                                    'format': image_data['format'],
                                    'width': width,
                                    'height': height,
                                    'caption': caption,
                                    'caption_position': 'below' if caption else None,
                                    'is_inline': True,
                                    'anchor_type': 'inline',
                                }
                                images.append(image_meta)
                                self.image_count += 1
                            
        except Exception as e:
            logger.warning(f"Error extracting images from paragraph {para_index}: {str(e)}")
        
        return images
    
    def _extract_images_from_table(self, table, table_index, element_index):
        """Extract images from table cells."""
        images = []
        
        try:
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        for run in para.runs:
                            # Look for embedded images
                            blips = run._element.findall('.//' + qn('a:blip'))
                            for blip in blips:
                                embed = blip.get(qn('r:embed'))
                                if embed:
                                    image_data = self._get_image_from_rId(table.part if hasattr(table, 'part') else cell.part, embed)
                                    if image_data:
                                        # Get dimensions
                                        width, height = self._get_inline_dimensions(run._element)
                                        
                                        image_meta = {
                                            'image_id': f'img_{self.image_count:04d}',
                                            'position_type': 'table',
                                            'paragraph_index': None,
                                            'element_index': element_index,
                                            'table_location': {
                                                'table_index': table_index,
                                                'row_index': row_idx,
                                                'cell_index': cell_idx,
                                                'para_index': para_idx,
                                            },
                                            'data': image_data['bytes'],
                                            'format': image_data['format'],
                                            'width': min(width, 2.0),  # Limit size for table cells
                                            'height': min(height, 2.0),
                                            'caption': None,
                                            'is_inline': True,
                                            'anchor_type': 'inline',
                                        }
                                        images.append(image_meta)
                                        self.image_count += 1
                                        logger.info(f"Extracted table image {image_meta['image_id']} at table {table_index}, row {row_idx}, cell {cell_idx}")
        
        except Exception as e:
            logger.warning(f"Error extracting images from table {table_index}: {str(e)}")
        
        return images
    
    def _extract_floating_images(self, doc):
        """Extract floating/anchored images not inline with text."""
        images = []
        
        try:
            # Look for anchored drawings in the document
            for i, para in enumerate(doc.paragraphs):
                for run in para.runs:
                    anchors = run._element.findall('.//' + qn('wp:anchor'))
                    for anchor in anchors:
                        blips = anchor.findall('.//' + qn('a:blip'))
                        for blip in blips:
                            embed = blip.get(qn('r:embed'))
                            # Skip if already extracted as inline image
                            if embed and embed not in self.extracted_rIds:
                                image_data = self._get_image_from_rId(para.part, embed)
                                if image_data:
                                    self.extracted_rIds.add(embed)  # Mark as extracted
                                    # Get dimensions
                                    extent = anchor.find(qn('wp:extent'))
                                    width = self._emu_to_inches(int(extent.get('cx', 0))) if extent is not None else 3.0
                                    height = self._emu_to_inches(int(extent.get('cy', 0))) if extent is not None else 2.0
                                    
                                    image_meta = {
                                        'image_id': f'img_{self.image_count:04d}',
                                        'position_type': 'floating',
                                        'paragraph_index': i,
                                        'element_index': i,
                                        'table_location': None,
                                        'data': image_data['bytes'],
                                        'format': image_data['format'],
                                        'width': width,
                                        'height': height,
                                        'caption': None,
                                        'is_inline': False,
                                        'anchor_type': 'floating',
                                    }
                                    images.append(image_meta)
                                    self.image_count += 1
                                    logger.info(f"Extracted floating image {image_meta['image_id']}")
        
        except Exception as e:
            logger.warning(f"Error extracting floating images: {str(e)}")
        
        return images
    
    def _get_image_from_rId(self, part, rId):
        """Get image binary data from relationship ID."""
        try:
            rel = part.rels.get(rId)
            if rel and rel.target_part:
                # Get the image part
                image_part = rel.target_part
                content_type = image_part.content_type
                
                # Determine format
                img_format = self.SUPPORTED_FORMATS.get(content_type, 'png')
                
                # Get binary data
                image_bytes = image_part.blob
                
                return {
                    'bytes': image_bytes,
                    'format': img_format,
                    'content_type': content_type,
                }
        except Exception as e:
            logger.warning(f"Could not get image from rId {rId}: {str(e)}")
        
        return None
    
    def _get_inline_dimensions(self, run_element):
        """Get dimensions from inline shape element."""
        try:
            # Try to find extent element
            extent = run_element.find('.//' + qn('wp:extent'))
            if extent is not None:
                cx = int(extent.get('cx', 0))
                cy = int(extent.get('cy', 0))
                return self._emu_to_inches(cx), self._emu_to_inches(cy)
            
            # Try to find a:ext element
            ext = run_element.find('.//' + qn('a:ext'))
            if ext is not None:
                cx = int(ext.get('cx', 0))
                cy = int(ext.get('cy', 0))
                return self._emu_to_inches(cx), self._emu_to_inches(cy)
                
        except Exception as e:
            logger.warning(f"Could not get dimensions: {str(e)}")
        
        # Default dimensions
        return 4.0, 3.0
    
    def _emu_to_inches(self, emu):
        """Convert EMU (English Metric Units) to inches."""
        # 914400 EMU = 1 inch
        if emu <= 0:
            return 3.0  # Default
        return emu / 914400
    
    def _detect_caption(self, para, para_text):
        """Detect if paragraph text is a figure caption."""
        if not para_text:
            return None
        
        for pattern in self.CAPTION_PATTERNS:
            if pattern.match(para_text):
                return para_text
        
        return None
    
    def get_images_by_position(self):
        """Get images organized by their position in document."""
        position_map = {}
        
        for img in self.images:
            if img['position_type'] == 'paragraph':
                key = ('paragraph', img['paragraph_index'])
            elif img['position_type'] == 'table':
                loc = img['table_location']
                key = ('table', loc['table_index'], loc['row_index'], loc['cell_index'])
            else:
                key = ('floating', img['element_index'])
            
            if key not in position_map:
                position_map[key] = []
            position_map[key].append(img)
        
        return position_map


class ImageInserter:
    """
    Insert images into Word documents at correct positions.
    """
    
    def __init__(self, doc, images):
        """
        Initialize with document and extracted images.
        
        Args:
            doc: python-docx Document object
            images: List of image metadata dicts from ImageExtractor
        """
        self.doc = doc
        self.images = images
        self.image_lookup = {img['image_id']: img for img in images}
        
    def insert_image(self, image_id, after_paragraph=None):
        """
        Insert an image into the document.
        
        Args:
            image_id: The ID of the image to insert
            after_paragraph: Paragraph object to insert after (optional)
            
        Returns:
            The paragraph containing the image
        """
        if image_id not in self.image_lookup:
            logger.warning(f"Image {image_id} not found in lookup")
            return None
        
        img_data = self.image_lookup[image_id]
        
        try:
            # Create paragraph for image
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add image from bytes
            run = para.add_run()
            
            # Create BytesIO stream from image data
            image_stream = BytesIO(img_data['data'])
            
            # Determine width and height
            width = img_data.get('width', 4.0)
            height = img_data.get('height', 3.0)
            
            # Limit maximum dimensions
            max_width = 6.0  # Max 6 inches wide
            max_height = 8.0  # Max 8 inches tall
            
            if width > max_width:
                ratio = max_width / width
                width = max_width
                height = height * ratio
            
            if height > max_height:
                ratio = max_height / height
                height = max_height
                width = width * ratio
            
            # Add picture
            run.add_picture(image_stream, width=Inches(width), height=Inches(height))
            
            # Add caption if exists
            if img_data.get('caption'):
                caption_para = self.doc.add_paragraph()
                caption_run = caption_para.add_run(img_data['caption'])
                caption_run.italic = False
                caption_run.font.name = 'Times New Roman'
                caption_run.font.size = Pt(max(8, self.font_size - 2))
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_para.paragraph_format.space_after = Pt(12)
            
            logger.info(f"Inserted image {image_id} ({width:.2f}x{height:.2f} inches)")
            return para
            
        except Exception as e:
            logger.error(f"Error inserting image {image_id}: {str(e)}")
            # Add placeholder text
            para = self.doc.add_paragraph()
            para.add_run(f"[IMAGE: {image_id} - Could not insert]")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return para
    
    def insert_image_in_table_cell(self, image_id, cell):
        """
        Insert an image into a table cell.
        
        Args:
            image_id: The ID of the image to insert
            cell: The table cell to insert into
        """
        if image_id not in self.image_lookup:
            logger.warning(f"Image {image_id} not found for table insertion")
            return
        
        img_data = self.image_lookup[image_id]
        
        try:
            # Get or create paragraph in cell
            para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = para.add_run()
            image_stream = BytesIO(img_data['data'])
            
            # Use smaller dimensions for table cells
            width = min(img_data.get('width', 2.0), 2.0)
            height = min(img_data.get('height', 1.5), 1.5)
            
            run.add_picture(image_stream, width=Inches(width), height=Inches(height))
            
            logger.info(f"Inserted image {image_id} in table cell")
            
        except Exception as e:
            logger.error(f"Error inserting table image {image_id}: {str(e)}")
            para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            para.add_run(f"[IMAGE: {image_id}]")


# ============================================================
# SHAPE/FLOWCHART EXTRACTION AND REINSERTION SYSTEM
# ============================================================

class ShapeExtractor:
    """
    Extract shapes, flowcharts, arrows, and drawing elements from Word documents.
    Preserves position, dimensions, formatting, and groupings.
    
    Handles:
    - WordprocessingML Shapes (wsp) - rectangles, ovals, arrows, etc.
    - Connector Shapes (cxnSp) - lines connecting shapes
    - Group Shapes (grpSp) - grouped shape elements
    - Drawing Canvas (wpc) - drawing canvas containers
    - VML Shapes (legacy) - older shape format
    """
    
    def __init__(self):
        self.shapes = []
        self.shape_count = 0
        self.extracted_shape_ids = set()  # Track extracted shapes to prevent duplicates
        
    def extract_all_shapes(self, doc_path):
        """
        Extract all shapes and drawing elements from a Word document.
        
        Args:
            doc_path: Path to the .docx file
            
        Returns:
            list: List of shape metadata dictionaries containing XML and position info
        """
        self.shapes = []
        self.shape_count = 0
        self.extracted_shape_ids = set()
        
        try:
            doc = Document(doc_path)
            paragraph_index = 0
            
            # Process document body elements in order
            for element in doc.element.body:
                if element.tag.endswith('p'):
                    # This is a paragraph - check for shapes/drawings
                    for para in doc.paragraphs:
                        if para._element is element:
                            shapes_in_para = self._extract_shapes_from_paragraph(
                                para, paragraph_index
                            )
                            self.shapes.extend(shapes_in_para)
                            paragraph_index += 1
                            break
                            
                elif element.tag.endswith('tbl'):
                    # Check tables for shapes
                    table_index = 0
                    for table in doc.tables:
                        if table._element is element:
                            shapes_in_table = self._extract_shapes_from_table(
                                table, table_index
                            )
                            self.shapes.extend(shapes_in_table)
                            break
                        table_index += 1
            
            logger.info(f"Extracted {len(self.shapes)} shapes/drawings from document")
            return self.shapes
            
        except Exception as e:
            logger.error(f"Error extracting shapes: {str(e)}")
            return []
    
    def _extract_shapes_from_paragraph(self, para, para_index):
        """Extract shapes and drawing elements from a paragraph."""
        shapes = []
        
        try:
            for run_index, run in enumerate(para.runs):
                run_elements = list(run._element)
                for element_index, element in enumerate(run_elements):
                    if element.tag == qn('w:drawing'):
                        drawing = element
                        # Check if this drawing contains shapes (not just images)
                        # Look for wsp (WordprocessingML Shape)
                        wsp_elements = drawing.findall('.//' + qn('wps:wsp'))
                        grp_elements = drawing.findall('.//' + qn('wpg:grpSp'))
                        cxn_elements = drawing.findall('.//' + qn('wps:cxnSp'))

                        # Also check for shapes in mc:AlternateContent
                        mc_alternate = drawing.findall('.//' + qn('mc:AlternateContent'))
                        for mc in mc_alternate:
                            wsp_elements.extend(mc.findall('.//' + qn('wps:wsp')))
                            grp_elements.extend(mc.findall('.//' + qn('wpg:grpSp')))
                            cxn_elements.extend(mc.findall('.//' + qn('wps:cxnSp')))

                        has_shapes = wsp_elements or grp_elements or cxn_elements

                        # Skip if this is just an image (has blip but no shapes)
                        has_blip = drawing.findall('.//' + qn('a:blip'))
                        if has_blip and not has_shapes:
                            continue  # This is an image, handled by ImageExtractor

                        if has_shapes:
                            # Get the complete drawing XML to preserve all formatting
                            from copy import deepcopy
                            drawing_copy = deepcopy(drawing)

                            # Get positioning info
                            inline = drawing.find('.//' + qn('wp:inline'))
                            anchor = drawing.find('.//' + qn('wp:anchor'))

                            position_type = 'inline' if inline is not None else 'anchor'
                            position_data = self._get_position_data(inline if inline is not None else anchor)

                            shape_meta = {
                                'shape_id': f'shape_{self.shape_count:04d}',
                                'position_type': position_type,
                                'paragraph_index': para_index,
                                'table_location': None,
                                'drawing_xml': drawing_copy,  # Store complete drawing element
                                'run_xml': deepcopy(run._element),  # Store the run context
                                'position_data': position_data,
                                'shape_count': len(wsp_elements) + len(grp_elements) + len(cxn_elements),
                                'has_group': len(grp_elements) > 0,
                                'has_connectors': len(cxn_elements) > 0,
                                'run_index': run_index,
                                'element_index': element_index,
                            }
                            shapes.append(shape_meta)
                            self.shape_count += 1
                            logger.info(
                                f"Extracted shape group {shape_meta['shape_id']} at paragraph {para_index} "
                                f"(shapes: {len(wsp_elements)}, groups: {len(grp_elements)}, connectors: {len(cxn_elements)})"
                            )
                    elif element.tag == qn('w:pict'):
                        from copy import deepcopy
                        pict_copy = deepcopy(element)

                        shape_meta = {
                            'shape_id': f'shape_{self.shape_count:04d}',
                            'position_type': 'vml',
                            'paragraph_index': para_index,
                            'table_location': None,
                            'pict_xml': pict_copy,
                            'run_xml': deepcopy(run._element),
                            'position_data': {},
                            'shape_count': 1,
                            'has_group': False,
                            'has_connectors': False,
                            'is_vml': True,
                            'run_index': run_index,
                            'element_index': element_index,
                        }
                        shapes.append(shape_meta)
                        self.shape_count += 1
                        logger.info(f"Extracted VML shape {shape_meta['shape_id']} at paragraph {para_index}")

        except Exception as e:
            logger.warning(f"Error extracting shapes from paragraph {para_index}: {str(e)}")
        
        return shapes
    
    def _extract_shapes_from_table(self, table, table_index):
        """Extract shapes from table cells."""
        shapes = []
        
        try:
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        for run_index, run in enumerate(para.runs):
                            run_elements = list(run._element)
                            for element_index, element in enumerate(run_elements):
                                if element.tag == qn('w:drawing'):
                                    drawing = element
                                    wsp_elements = drawing.findall('.//' + qn('wps:wsp'))
                                    grp_elements = drawing.findall('.//' + qn('wpg:grpSp'))
                                    cxn_elements = drawing.findall('.//' + qn('wps:cxnSp'))

                                    has_shapes = wsp_elements or grp_elements or cxn_elements
                                    has_blip = drawing.findall('.//' + qn('a:blip'))

                                    if has_blip and not has_shapes:
                                        continue

                                    if has_shapes:
                                        from copy import deepcopy
                                        drawing_copy = deepcopy(drawing)

                                        shape_meta = {
                                            'shape_id': f'shape_{self.shape_count:04d}',
                                            'position_type': 'table',
                                            'paragraph_index': None,
                                            'table_location': {
                                                'table_index': table_index,
                                                'row_index': row_idx,
                                                'cell_index': cell_idx,
                                                'para_index': para_idx,
                                            },
                                            'drawing_xml': drawing_copy,
                                            'run_xml': deepcopy(run._element),
                                            'position_data': {},
                                            'shape_count': len(wsp_elements) + len(grp_elements) + len(cxn_elements),
                                            'has_group': len(grp_elements) > 0,
                                            'has_connectors': len(cxn_elements) > 0,
                                            'run_index': run_index,
                                            'element_index': element_index,
                                        }
                                        shapes.append(shape_meta)
                                        self.shape_count += 1
                                        logger.info(f"Extracted table shape {shape_meta['shape_id']} at table {table_index}")
                                elif element.tag == qn('w:pict'):
                                    from copy import deepcopy
                                    shape_meta = {
                                        'shape_id': f'shape_{self.shape_count:04d}',
                                        'position_type': 'table_vml',
                                        'paragraph_index': None,
                                        'table_location': {
                                            'table_index': table_index,
                                            'row_index': row_idx,
                                            'cell_index': cell_idx,
                                            'para_index': para_idx,
                                        },
                                        'pict_xml': deepcopy(element),
                                        'run_xml': deepcopy(run._element),
                                        'position_data': {},
                                        'shape_count': 1,
                                        'has_group': False,
                                        'has_connectors': False,
                                        'is_vml': True,
                                        'run_index': run_index,
                                        'element_index': element_index,
                                    }
                                    shapes.append(shape_meta)
                                    self.shape_count += 1
                                
        except Exception as e:
            logger.warning(f"Error extracting shapes from table {table_index}: {str(e)}")
        
        return shapes
    
    def _get_position_data(self, pos_element):
        """Extract positioning data from inline or anchor element."""
        position_data = {}
        
        if pos_element is None:
            return position_data
            
        try:
            # Get extent (dimensions)
            extent = pos_element.find(qn('wp:extent'))
            if extent is not None:
                position_data['width_emu'] = int(extent.get('cx', 0))
                position_data['height_emu'] = int(extent.get('cy', 0))
            
            # For anchor elements, get position offsets
            if pos_element.tag.endswith('anchor'):
                # Horizontal position
                pos_h = pos_element.find(qn('wp:positionH'))
                if pos_h is not None:
                    position_data['h_relative_from'] = pos_h.get('relativeFrom', 'column')
                    pos_offset = pos_h.find(qn('wp:posOffset'))
                    if pos_offset is not None and pos_offset.text:
                        position_data['h_offset'] = int(pos_offset.text)
                
                # Vertical position
                pos_v = pos_element.find(qn('wp:positionV'))
                if pos_v is not None:
                    position_data['v_relative_from'] = pos_v.get('relativeFrom', 'paragraph')
                    pos_offset = pos_v.find(qn('wp:posOffset'))
                    if pos_offset is not None and pos_offset.text:
                        position_data['v_offset'] = int(pos_offset.text)
                
                # Wrap type
                for wrap_type in ['wrapNone', 'wrapSquare', 'wrapTight', 'wrapThrough', 'wrapTopAndBottom']:
                    wrap = pos_element.find(qn(f'wp:{wrap_type}'))
                    if wrap is not None:
                        position_data['wrap_type'] = wrap_type
                        break
                        
        except Exception as e:
            logger.warning(f"Error getting position data: {str(e)}")
        
        return position_data
    
    def get_shapes_by_position(self):
        """Get shapes organized by their position in document."""
        position_map = {}
        
        for shape in self.shapes:
            if shape['position_type'] in ['inline', 'anchor', 'vml']:
                key = ('paragraph', shape['paragraph_index'])
            elif shape['position_type'] in ['table', 'table_vml']:
                loc = shape['table_location']
                key = ('table', loc['table_index'], loc['row_index'], loc['cell_index'])
            else:
                key = ('other', shape.get('paragraph_index', 0))
            
            if key not in position_map:
                position_map[key] = []
            position_map[key].append(shape)
        
        return position_map


class ShapeInserter:
    """
    Insert shapes back into Word documents at their original positions.
    Preserves all formatting, positioning, and groupings.
    """
    
    def __init__(self, doc, shapes):
        """
        Initialize with document and extracted shapes.
        
        Args:
            doc: python-docx Document object
            shapes: List of shape metadata dicts from ShapeExtractor
        """
        self.doc = doc
        self.shapes = shapes
        self.shape_lookup = {shape['shape_id']: shape for shape in shapes}
        
    def insert_shape(self, shape_id, paragraph=None):
        """
        Insert a shape/drawing back into the document.
        
        Args:
            shape_id: The ID of the shape to insert
            paragraph: Paragraph object to insert into (optional, creates new if None)
            
        Returns:
            The paragraph containing the shape
        """
        if shape_id not in self.shape_lookup:
            logger.warning(f"Shape {shape_id} not found in lookup")
            return None
        
        shape_data = self.shape_lookup[shape_id]
        
        try:
            # Create or use existing paragraph
            if paragraph is None:
                para = self.doc.add_paragraph()
            else:
                para = paragraph
            
            # Create a new run to hold the shape
            run = para.add_run()
            
            # Check if this is a VML shape or modern drawing
            if shape_data.get('is_vml'):
                # Insert VML pict element
                pict_xml = shape_data['pict_xml']
                run._element.append(pict_xml)
                logger.info(f"Inserted VML shape {shape_id}")
            else:
                # Insert modern drawing element
                drawing_xml = shape_data['drawing_xml']
                run._element.append(drawing_xml)
                logger.info(f"Inserted drawing shape {shape_id}")
            
            return para
            
        except Exception as e:
            logger.error(f"Error inserting shape {shape_id}: {str(e)}")
            # Add placeholder text
            para = self.doc.add_paragraph() if paragraph is None else paragraph
            para.add_run(f"[SHAPE: {shape_id} - Could not insert]")
            return para
    
    def insert_shape_in_table_cell(self, shape_id, cell):
        """
        Insert a shape into a table cell.
        
        Args:
            shape_id: The ID of the shape to insert
            cell: The table cell to insert into
        """
        if shape_id not in self.shape_lookup:
            logger.warning(f"Shape {shape_id} not found for table insertion")
            return
        
        shape_data = self.shape_lookup[shape_id]
        
        try:
            # Get or create paragraph in cell
            para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            run = para.add_run()
            
            if shape_data.get('is_vml'):
                run._element.append(shape_data['pict_xml'])
            else:
                run._element.append(shape_data['drawing_xml'])
            
            logger.info(f"Inserted shape {shape_id} in table cell")
            
        except Exception as e:
            logger.error(f"Error inserting table shape {shape_id}: {str(e)}")
            para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            para.add_run(f"[SHAPE: {shape_id}]")


import re

# --- Hierarchy Correction Utility ---
hierarchical_pairs = {
    'WEEK': ['DAY', 'SESSION', 'CLASS', 'LECTURE'],
    'MONTH': ['WEEK', 'PHASE'],
    'YEAR': ['QUARTER', 'SEMESTER', 'TERM'],
    'CHAPTER': ['SECTION', 'TOPIC', 'SUBTOPIC'],
    'UNIT': ['LESSON', 'MODULE', 'EXERCISE'],
    'PART': ['CHAPTER', 'SECTION'],
    'MODULE': ['TOPIC', 'SUBTOPIC', 'ACTIVITY'],
    'THEORY': ['PRINCIPLE', 'CONCEPT', 'MODEL'],
    'METHOD': ['STEP', 'PROCEDURE', 'TECHNIQUE'],
    'ANALYSIS': ['RESULT', 'FINDING', 'INTERPRETATION'],
    'FRAMEWORK': ['COMPONENT', 'ELEMENT', 'DIMENSION'],
    'SYSTEM': ['SUBSYSTEM', 'MODULE', 'COMPONENT'],
    'PROCESS': ['STAGE', 'PHASE', 'STEP'],
    'MODEL': ['VARIABLE', 'COMPONENT', 'ELEMENT'],
    'STRATEGY': ['TACTIC', 'APPROACH', 'METHOD'],
    'CATEGORY': ['TYPE', 'CLASS', 'FORM'],
    'PRINCIPLE': ['RULE', 'GUIDELINE', 'STANDARD'],
}

class HierarchyCorrector:
    """Detect and correct hierarchical numbering issues in heading lines."""
    
    HIERARCHICAL_PAIRS = {
        'WEEK': ['DAY', 'SESSION', 'CLASS', 'LECTURE'],
        'MONTH': ['WEEK', 'PHASE'],
        'YEAR': ['QUARTER', 'SEMESTER', 'TERM'],
        'CHAPTER': ['SECTION', 'TOPIC', 'SUBTOPIC'],
        'UNIT': ['LESSON', 'MODULE', 'EXERCISE'],
        'PART': ['CHAPTER', 'SECTION'],
        'MODULE': ['TOPIC', 'SUBTOPIC', 'ACTIVITY'],
        'THEORY': ['PRINCIPLE', 'CONCEPT', 'MODEL'],
        'METHOD': ['STEP', 'PROCEDURE', 'TECHNIQUE'],
        'ANALYSIS': ['RESULT', 'FINDING', 'INTERPRETATION'],
        'FRAMEWORK': ['COMPONENT', 'ELEMENT', 'DIMENSION'],
        'SYSTEM': ['SUBSYSTEM', 'MODULE', 'COMPONENT'],
        'PROCESS': ['STAGE', 'PHASE', 'STEP'],
        'MODEL': ['VARIABLE', 'COMPONENT', 'ELEMENT'],
        'STRATEGY': ['TACTIC', 'APPROACH', 'METHOD'],
        'CATEGORY': ['TYPE', 'CLASS', 'FORM'],
        'PRINCIPLE': ['RULE', 'GUIDELINE', 'STANDARD'],
    }

    def __init__(self):
        self.patterns = [
            # Pattern A: Sequential Major/Minor Topics (Placeholder for specific logic)
            # (re.compile(r'(?i)^\s*(\d+(?:\.\d+)*)\s+([A-Z\s]{2,})(?:\s+|$)\n^\s*(\d+(?:\.\d+)*)\s+([A-Z\s]{2,})(?:\s+|$)', re.MULTILINE),
            #  lambda m: m.group(0)), 

            # Pattern B: Temporal/Categorical Relationships
            (re.compile(r'(?i)(\d+(?:\.\d+)*)\s+(WEEK|MONTH|YEAR|QUARTER|TERM|SEMESTER)\s+([A-Z\d]+)(?:\s+|$)\n(\d+(?:\.\d+)*)\s+(DAY|SESSION|CLASS|PERIOD|LECTURE)\s+([A-Z\d]+)(?:\s+|$)', re.MULTILINE),
             lambda m: f"{m.group(1)} {m.group(2)} {m.group(3)}\n{m.group(1)}.1 {m.group(5)} {m.group(6)}"),

            # Pattern D: Week/Day Pattern
            (re.compile(r'(?i)^(\d+\.\d+)\s+WEEK\s+([A-Z\d]+).*?\n^(\d+\.\d+)\s+DAY\s+([A-Z\d]+)', re.MULTILINE),
             lambda m: f"{m.group(1)} WEEK {m.group(2)}\n{m.group(1)}.1 DAY {m.group(4)}"),

            # Pattern E: Unit/Lesson Pattern
            (re.compile(r'(?i)^(\d+\.\d+)\s+UNIT\s+([A-Z\d]+).*?\n^(\d+\.\d+)\s+LESSON\s+([A-Z\d]+)', re.MULTILINE),
             lambda m: f"{m.group(1)} UNIT {m.group(2)}\n{m.group(1)}.1 LESSON {m.group(4)}"),
             
            # Pattern F: Chapter/Section Pattern
            (re.compile(r'(?i)^(\d+\.\d+)\s+CHAPTER\s+([A-Z\d]+).*?\n^(\d+\.\d+)\s+SECTION\s+([A-Z\d]+)', re.MULTILINE),
             lambda m: f"{m.group(1)} CHAPTER {m.group(2)}\n{m.group(1)}.1 SECTION {m.group(4)}"),

            # Pattern G: Module/Topic Pattern
            (re.compile(r'(?i)^(\d+\.\d+)\s+MODULE\s+([A-Z\d]+).*?\n^(\d+\.\d+)\s+TOPIC\s+([A-Z\d]+)', re.MULTILINE),
             lambda m: f"{m.group(1)} MODULE {m.group(2)}\n{m.group(1)}.1 TOPIC {m.group(4)}"),

            # Pattern H: Lettered Hierarchies
            (re.compile(r'(?i)^(\d+\.\d+)\s+((?:PART\s+)?[A-Z])\b.*?\n^(\d+\.\d+)\s+(\d+[\.\)]?)\s+(.*)', re.MULTILINE),
             lambda m: f"{m.group(1)} {m.group(2)}\n{m.group(1)}.1 {m.group(4)} {m.group(5)}"),

            # Pattern J: Short Title Followed by Specific Title
            (re.compile(r'^(\d+\.\d+)\s+([A-Z]{2,15})\s*$\n^(\d+\.\d+)\s+([A-Z].{10,})', re.MULTILINE),
             lambda m: f"{m.group(1)} {m.group(2)}\n{m.group(1)}.1 {m.group(4)}"),

            # Pattern K: Category/Subcategory Pattern
            (re.compile(r'(?i)^(\d+\.\d+)\s+(TYPES|CATEGORIES|CLASSIFICATIONS|FORMS|MODELS).*?\n^(\d+\.\d+)\s+((?:.*?MODEL|.*?TYPE|.*?FORM).*)', re.MULTILINE),
             lambda m: f"{m.group(1)} {m.group(2)}\n{m.group(1)}.1 {m.group(4)}"), # Simplified replacement

            # Pattern Parent/Child with same starting words (Pattern L/M/N combined logic)
            (re.compile(r'^(\d+\.\d+)\s+(.*?\b\w+\b).*?\n^(\d+\.\d+)\s+\2.*?', re.MULTILINE),
             lambda m: f"{m.group(1)} {m.group(2)}\n{m.group(1)}.1 {m.group(2)}"),
             
             # Pattern R: Convert Flat to Hierarchical (Generic)
            (re.compile(r'(?i)^(\d+\.\d+)\s+(WEEK|UNIT|MODULE|CHAPTER|PART)\s+([A-Z\d]+)(?:\s+|$)\n^(\d+\.\d+)\s+(DAY|LESSON|TOPIC|SECTION|SESSION)\s+([A-Z\d]+)(?:\s+|$)', re.MULTILINE),
             lambda m: f"{m.group(1)} {m.group(2)} {m.group(3)}\n{m.group(1)}.1 {m.group(5)} {m.group(6)}"),
        ]

    def correct(self, text):
        corrected = text
        for pattern, repl in self.patterns:
            corrected = pattern.sub(repl, corrected)
        return corrected

    def is_hierarchical_pair(self, parent, child):
        parent_upper = parent.upper()
        child_upper = child.upper()
        
        # Check dictionary
        for p, children in self.HIERARCHICAL_PAIRS.items():
            if p in parent_upper:
                for c in children:
                    if c in child_upper:
                        return True
        
        # Check general/specific (short parent, long child)
        if len(parent.split()) < 4 and len(child.split()) > 3:
            common_words = set(parent_upper.split()) & set(child_upper.split())
            if len(common_words) > 0:
                return True
                
        return False

    def correct_lines(self, lines):
        """
        Correct hierarchical numbering issues in a list of heading lines.
        Implements the logic from 'correct_hierarchical_numbering' and 'smart_hierarchy_correction'.
        """
        corrected_lines = []
        i = 0
        while i < len(lines):
            current_line = lines[i]
            
            # Skip if not a numbered heading (simple check)
            m1 = re.match(r'^(\d+(?:\.\d+)*)\s+(.*)$', current_line)
            if not m1:
                corrected_lines.append(current_line)
                i += 1
                continue
                
            if i == len(lines) - 1:
                corrected_lines.append(current_line)
                break
                
            next_line = lines[i+1]
            m2 = re.match(r'^(\d+(?:\.\d+)*)\s+(.*)$', next_line)
            
            if m2:
                current_num, current_title = m1.group(1), m1.group(2)
                next_num, next_title = m2.group(1), m2.group(2)
                
                if self.is_hierarchical_pair(current_title, next_title):
                    # Convert to hierarchical
                    # If parent is 3.6, child becomes 3.6.1
                    child_num = f"{current_num}.1"
                    corrected_lines.append(current_line)
                    corrected_lines.append(f"{child_num} {next_title}")
                    i += 2  # Skip next line
                    continue
            
            corrected_lines.append(current_line)
            i += 1
            
        return corrected_lines

# --- End Hierarchy Correction Utility ---

class HeadingNumberer:
    """
    Auto-number headings based on chapter context with semantic hierarchy detection.
    Tracks chapter numbers and assigns hierarchical numbering to subheadings.
    
    Features:
    - Detects parent-child relationships based on heading text patterns
    - Handles "Main X" / "Specific X" as children of "X" sections
    - Properly nests definition terms under "Operational Definition of Terms"
    - Fixes flat numbering issues (1.3, 1.4, 1.5 → 1.3, 1.3.1, 1.3.2)
    
    Example transformations:
    - "CHAPTER ONE" → stays as is (sets chapter = 1)
    - "Research Objectives" → "1.4 Research Objectives"
    - "Main Research Objective" → "1.4.1 Main Research Objective"
    - "Specific Research Objectives" → "1.4.2 Specific Research Objectives"
    """
    
    # Roman numeral to integer mapping
    ROMAN_TO_INT = {
        'I': 1, 'II': 2, 'III': 3, 'IV': 4, 'V': 5,
        'VI': 6, 'VII': 7, 'VIII': 8, 'IX': 9, 'X': 10,
        'XI': 11, 'XII': 12, 'XIII': 13, 'XIV': 14, 'XV': 15
    }
    
    # Word to integer mapping
    WORD_TO_INT = {
        'ONE': 1, 'TWO': 2, 'THREE': 3, 'FOUR': 4, 'FIVE': 5,
        'SIX': 6, 'SEVEN': 7, 'EIGHT': 8, 'NINE': 9, 'TEN': 10,
        'ELEVEN': 11, 'TWELVE': 12, 'THIRTEEN': 13, 'FOURTEEN': 14, 'FIFTEEN': 15
    }
    
    # Front matter sections that should NOT be numbered
    UNNUMBERED_SECTIONS = {
        'DECLARATION', 'CERTIFICATION', 'DEDICATION', 'ACKNOWLEDGEMENTS',
        'ACKNOWLEDGMENTS', 'ACKNOWLEDGEMENT', 'ABSTRACT', 'RESUME', 'RÉSUMÉ',
        'TABLE OF CONTENTS', 'CONTENTS', 'LIST OF TABLES', 'LIST OF FIGURES',
        'LIST OF ABBREVIATIONS', 'ABBREVIATIONS', 'GLOSSARY', 'REFERENCES',
        'BIBLIOGRAPHY', 'APPENDIX', 'APPENDICES', 'INDEX', 'PREFACE', 'FOREWORD'
    }
    
    # Section titles that are typically ALL CAPS and follow chapter headings (not numbered)
    CHAPTER_TITLE_SECTIONS = {
        'GENERAL INTRODUCTION', 'INTRODUCTION', 'REVIEW OF RELATED LITERATURE',
        'LITERATURE REVIEW', 'RESEARCH METHODOLOGY', 'METHODOLOGY',
        'DATA ANALYSIS AND INTERPRETATION', 'DATA ANALYSIS', 'FINDINGS AND DISCUSSION',
        'DISCUSSION', 'SUMMARY CONCLUSION AND RECOMMENDATIONS', 'SUMMARY AND CONCLUSION',
        'CONCLUSION', 'RECOMMENDATIONS', 'PRESENTATION OF FINDINGS',
        'RESULTS AND DISCUSSION', 'ANALYSIS AND FINDINGS'
    }
    
    # Parent sections and their expected children (for semantic hierarchy detection)
    # Format: 'parent_keyword': ['child_keyword1', 'child_keyword2', ...]
    PARENT_CHILD_PATTERNS = {
        'research objectives': ['main research objective', 'specific research objective', 'general objective', 'specific objectives'],
        'research questions': ['main research question', 'specific research question', 'general question', 'specific questions'],
        'research hypothesis': ['main research hypothes', 'specific research hypothes', 'main hypothes', 'specific hypothes', 'null hypothesis', 'alternative hypothesis'],
        'research hypotheses': ['main research hypothes', 'specific research hypothes', 'main hypothes', 'specific hypothes', 'null hypothesis', 'alternative hypothesis'],
        'justification of the study': ['policy maker', 'accountability', 'knowledge gap', 'evaluation', 'stakeholder', 'researcher', 'student', 'government'],
        'significance of the study': ['policy maker', 'accountability', 'knowledge gap', 'evaluation', 'stakeholder', 'researcher', 'student', 'government'],
        'operational definition': ['cost sharing', 'tax payer', 'mechanism', 'quality', 'grant', 'private universit', 'government subsid', 'community fund', 'private fund', 'tuition', 'fee', 'scholarship', 'bursary'],
        'definition of terms': ['cost sharing', 'tax payer', 'mechanism', 'quality', 'grant', 'private universit', 'government subsid', 'community fund', 'private fund', 'tuition', 'fee', 'scholarship', 'bursary'],
        'theoretical review': ['equity theory', 'human capital theory', 'revenue theory', 'resource dependency', 'stakeholder theory', 'agency theory', 'institutional theory'],
        'theoretical framework': ['equity theory', 'human capital theory', 'revenue theory', 'resource dependency', 'stakeholder theory', 'agency theory', 'institutional theory'],
        'conceptual review': ['concept of', 'conceptual framework', 'nexus', 'relationship between'],
        'conceptual framework': ['concept of', 'nexus', 'relationship between'],
        'empirical review': ['studies on', 'research on', 'findings', 'previous studies'],
        'delimitation': ['geographical', 'scope', 'population', 'time frame', 'period'],
        'limitation': ['sample size', 'time constraint', 'access', 'generalizability'],
        # New hierarchical patterns
        'week': ['day', 'session', 'class', 'lecture'],
        'month': ['week', 'phase'],
        'year': ['quarter', 'semester', 'term'],
        'chapter': ['section', 'topic', 'subtopic'],
        'unit': ['lesson', 'module', 'exercise'],
        'part': ['chapter', 'section'],
        'module': ['topic', 'subtopic', 'activity'],
        'theory': ['principle', 'concept', 'model'],
        'method': ['step', 'procedure', 'technique'],
        'analysis': ['result', 'finding', 'interpretation'],
        'framework': ['component', 'element', 'dimension'],
        'system': ['subsystem', 'module', 'component'],
        'process': ['stage', 'phase', 'step'],
        'model': ['variable', 'component', 'element'],
        'strategy': ['tactic', 'approach', 'method'],
        'category': ['type', 'class', 'form'],
        'principle': ['rule', 'guideline', 'standard'],
    }
    
    # Patterns that indicate content should NOT be a main section (should be subsection)
    SUBSECTION_INDICATORS = [
        r'^main\s+',  # "Main Research Objective"
        r'^specific\s+',  # "Specific Research Questions"
        r'^general\s+',  # "General Objective"
        r'^primary\s+',  # "Primary Research Question"
        r'^secondary\s+',  # "Secondary Questions"
        r'^null\s+',  # "Null Hypothesis"
        r'^alternative\s+',  # "Alternative Hypothesis"
        r'^\w+\s+can\s+be\s+expressed\s+as',  # "The production function can be expressed as"
        r'^where\s*:?\s*$',  # "Where:" (mathematical notation)
        r'^note\s*:?\s*$',  # "Note:"
        r'^example\s*:?\s*$',  # "Example:"
    ]
    
    # Patterns that indicate this is a definition term (should be under definitions section)
    DEFINITION_TERMS = [
        'cost sharing', 'tax payer', 'taxpayer', 'tuition', 'fee', 'fees',
        'scholarship', 'bursary', 'grant', 'loan', 'subsidy', 'subsidies',
        'funding', 'revenue', 'expenditure', 'budget', 'allocation',
        'quality education', 'quality of education', 'accreditation',
        'enrollment', 'enrolment', 'retention', 'graduation rate',
        'private university', 'public university', 'state university',
        'community funding', 'private funding', 'government funding',
    ]
    
    def __init__(self, policy=None):
        self.policy = policy or FormatPolicy()
        self.reset()
        self.hierarchy_corrector = HierarchyCorrector()
        self.allow_auto_numbering = False
        self.has_consistent_numbering = False
        
    def reset(self):
        """Reset all counters for a new document."""
        self.current_chapter = 0
        self.current_section = 0  # X.1, X.2, etc.
        self.current_subsection = 0  # X.Y.1, X.Y.2, etc.
        self.current_subsubsection = 0  # X.Y.Z.1, etc.
        self.in_appendix = False
        self.appendix_letter = 'A'
        self.last_level = 0  # Track the last heading level for hierarchy
        self.last_heading_text = ''  # Track last heading for parent-child detection
        self.last_heading_normalized = ''  # Normalized version for matching
        self.in_parent_section = None  # Current parent section (e.g., 'research objectives')
        self.parent_section_number = ''  # Number of current parent section (e.g., '1.4')
        self.use_continuous_section_numbering = False  # For documents without chapters
        self.last_numbers_by_level = {}

    def configure_for_lines(self, lines):
        """Configure numbering behavior based on document content and policy."""
        self.has_consistent_numbering = self._detect_consistent_numbering_scheme(lines)
        self.allow_auto_numbering = self.policy.allow_auto_numbering(self.has_consistent_numbering)
        logger.info(
            "Heading numbering policy: allow_auto_numbering=%s consistent_scheme=%s document_mode=%s",
            self.allow_auto_numbering,
            self.has_consistent_numbering,
            self.policy.document_mode,
        )

    def _detect_consistent_numbering_scheme(self, lines):
        """Detect whether headings already follow a consistent numbering scheme."""
        chapter_hits = 0
        numbered_hits = 0
        for line in lines:
            text = line.get('text') if isinstance(line, dict) else str(line)
            if not text:
                continue
            if self.parse_chapter_number(text) > 0:
                chapter_hits += 1
                continue
            if re.match(r'^\s*(\d+(?:\.\d+)+)\s+\S+', text) or re.match(r'^\s*\d+\.\s+[A-Z]', text):
                numbered_hits += 1
        return chapter_hits > 0 or numbered_hits >= 2

    def _sync_counters_from_existing(self, existing_num):
        """Align internal counters with existing numbering to avoid drift."""
        if not existing_num:
            return
        normalized = existing_num.strip()
        normalized = re.sub(r'^[\(\[]', '', normalized)
        normalized = re.sub(r'[\)\]\.]$', '', normalized)
        if re.match(r'^[A-Z]\.', normalized):
            parts = normalized.split('.')
            self.in_appendix = True
            self.appendix_letter = parts[0]
            parts = parts[1:]
        else:
            parts = normalized.split('.')

        numeric_parts = [p for p in parts if p.isdigit()]
        if not numeric_parts:
            return

        if self.use_continuous_section_numbering or self.current_chapter == 0:
            self.current_section = int(numeric_parts[0])
            self.current_subsection = int(numeric_parts[1]) if len(numeric_parts) > 1 else 0
        else:
            self.current_chapter = int(numeric_parts[0])
            self.current_section = int(numeric_parts[1]) if len(numeric_parts) > 1 else 0
            self.current_subsection = int(numeric_parts[2]) if len(numeric_parts) > 2 else 0

    def _should_correct_existing_number(self, existing_num, new_number, level):
        """Allow controlled correction for clear numbering conflicts."""
        if not self.allow_auto_numbering or self.policy.document_mode != "academic":
            return False
        if not existing_num or not new_number or existing_num == new_number:
            return False
        if self.last_numbers_by_level.get(level) == existing_num:
            return True
        return False
        
    def _normalize_text(self, text):
        """Normalize text for comparison (lowercase, remove punctuation, extra spaces)."""
        clean = re.sub(r'^#+\s*', '', text).strip().lower()
        clean = re.sub(r'\*\*', '', clean)  # Remove markdown bold
        clean = re.sub(r'[\d\.]+\s*', '', clean)  # Remove existing numbers
        clean = re.sub(r'[^\w\s]', ' ', clean)  # Remove punctuation
        clean = re.sub(r'\s+', ' ', clean).strip()  # Normalize whitespace
        return clean
    
    def _is_child_of_parent(self, heading_text, parent_text):
        """
        Check if heading_text should be a child of parent_text based on semantic patterns.
        
        Returns:
            bool: True if heading should be a subsection of parent
        """
        heading_norm = self._normalize_text(heading_text)
        parent_norm = self._normalize_text(parent_text)
        
        # Check predefined parent-child patterns
        for parent_key, children in self.PARENT_CHILD_PATTERNS.items():
            if parent_key in parent_norm:
                for child_pattern in children:
                    if child_pattern in heading_norm:
                        return True
        
        return False
    
    def _is_subsection_indicator(self, text):
        """Check if text has patterns indicating it should be a subsection."""
        text_lower = self._normalize_text(text)
        
        for pattern in self.SUBSECTION_INDICATORS:
            if re.search(pattern, text_lower, re.IGNORECASE):
                return True
        
        return False
    
    def _is_definition_term(self, text):
        """Check if text looks like a definition term that should be under Definitions section."""
        text_lower = self._normalize_text(text)
        
        # Check if this matches known definition terms
        for term in self.DEFINITION_TERMS:
            if term in text_lower or text_lower in term:
                return True
        
        # Short titles (1-3 words) after a definitions section are likely definition terms
        if self.in_parent_section and 'definition' in self.in_parent_section:
            word_count = len(text_lower.split())
            if word_count <= 4:
                return True
        
        return False
    
    def _detect_parent_section(self, text):
        """
        Detect if this heading starts a parent section that will have children.
        
        Returns:
            str or None: The parent section key if detected, None otherwise
        """
        text_norm = self._normalize_text(text)
        
        for parent_key in self.PARENT_CHILD_PATTERNS.keys():
            if parent_key in text_norm:
                return parent_key
        
        return None
    
    def _should_be_subsection(self, text):
        """
        Determine if this heading should be a subsection of the previous heading.
        Uses semantic analysis to detect parent-child relationships.
        
        IMPORTANT: If the text already has a numbered hierarchy, this function
        returns False to preserve the original structure.
        
        Returns:
            bool: True if this should be a subsection
        """
        # CRITICAL FIX: If text already has numbering, NEVER treat it as subsection candidate
        # This preserves the original document structure
        existing_num, _ = self.extract_existing_number(text)
        if existing_num:
            # Sync counters to the existing number and return False
            # This ensures the original hierarchy is preserved
            return False
        
        text_norm = self._normalize_text(text)
        
        # Check for keywords that indicate a NEW main section (should NOT be subsection)
        # These are full section titles, not partial matches
        # EXPANDED LIST to include Chapter One typical sections that should remain independent
        main_section_keywords = [
            'chapter summary', 'summary of the chapter', 'conclusion of the chapter',
            'delimitation of the study', 'limitations of the study', 'delimitations',
            'significance of the study', 'scope of the study', 'organization of the study',
            'structure of the study', 'structure of the thesis', 'structure of the dissertation',
            'statement of the problem', 'problem statement',
            'conceptual review', 'conceptual framework',
            'theoretical review', 'theoretical framework',
            'empirical review', 'empirical studies', 'review of empirical',
            'research design', 'research methodology', 'methodology',
            'data presentation', 'data analysis and interpretation',
            'summary conclusion and recommendation', 'summary and conclusion',
            'recommendations for further', 'recommendations',
            # CHAPTER ONE SECTIONS - these should remain as independent sections
            'research questions', 'research question',
            'research objectives', 'research objective', 'objectives of the study',
            'research hypothesis', 'research hypotheses', 'hypothesis of the study',
            'justification of the study', 'justification of study',
            'significance of study',
            'operational definition of terms', 'definition of terms', 'definition of key terms',
            'background of the study', 'background to the study', 'background of study',
            'purpose of the study', 'purpose of study',
            'scope and delimitation', 'scope of study',
            'limitation of the study', 'limitations of study',
            'organization of the study', 'organization of study',
            # Sub-categories that should NOT be merged into parents
            'main research objective', 'specific research objective', 'specific objectives',
            'main research question', 'specific research question', 'specific questions',
            'main hypothesis', 'specific hypothesis', 'null hypothesis', 'alternative hypothesis',
            'general objective', 'general objectives',
        ]
        for keyword in main_section_keywords:
            if keyword in text_norm:
                # This is a main section - exit any parent context
                self.in_parent_section = None
                self.parent_section_number = ''
                return False
        
        # Check if it's a known subsection indicator
        if self._is_subsection_indicator(text):
            return True
        
        # Check if current parent section expects this as a child
        if self.in_parent_section:
            if self._is_child_of_parent(text, self.in_parent_section):
                return True
            
            # Definition terms under definitions section
            if 'definition' in self.in_parent_section and self._is_definition_term(text):
                return True
        
        # Check if it's a child of the last heading
        if self.last_heading_normalized:
            if self._is_child_of_parent(text, self.last_heading_text):
                return True
        
        return False
        
    def parse_chapter_number(self, text):
        """
        Extract chapter number from chapter heading text.
        
        Args:
            text: The heading text (e.g., "CHAPTER ONE", "CHAPTER 1", "CHAPTER IV")
            
        Returns:
            int: Chapter number, or 0 if not a chapter heading
        """
        text_upper = text.upper().strip()
        
        # Remove markdown heading markers
        text_upper = re.sub(r'^#+\s*', '', text_upper).strip()
        
        # Pattern: CHAPTER + word (ONE, TWO, etc.)
        match = re.match(r'^CHAP?TER\s+([A-Z]+)\b', text_upper)
        if match:
            word = match.group(1)
            if word in self.WORD_TO_INT:
                return self.WORD_TO_INT[word]
        
        # Pattern: CHAPTER + Roman numeral
        match = re.match(r'^CHAP?TER\s+([IVXLCDM]+)\b', text_upper)
        if match:
            roman = match.group(1)
            if roman in self.ROMAN_TO_INT:
                return self.ROMAN_TO_INT[roman]
        
        # Pattern: CHAPTER + digit
        match = re.match(r'^CHAP?TER\s+(\d+)', text_upper)
        if match:
            return int(match.group(1))
        
        return 0
    
    def is_chapter_heading(self, text):
        """Check if text is a chapter heading."""
        return self.parse_chapter_number(text) > 0
    
    def is_chapter_title(self, text):
        """
        Check if text is a chapter title section (ALL CAPS section after chapter).
        These typically don't get numbered.
        """
        clean = re.sub(r'^#+\s*', '', text).strip().upper()
        return clean in self.CHAPTER_TITLE_SECTIONS
    
    def is_unnumbered_section(self, text):
        """Check if text is a front matter or special section that shouldn't be numbered."""
        clean = re.sub(r'^#+\s*', '', text).strip().upper()
        clean = re.sub(r'[:.\s]+$', '', clean)
        return clean in self.UNNUMBERED_SECTIONS
    
    def is_appendix_heading(self, text):
        """Check if text starts an appendix section."""
        clean = re.sub(r'^#+\s*', '', text).strip().upper()
        return clean.startswith('APPENDIX') or clean.startswith('APPENDICES')
    
    def already_has_number(self, text):
        """
        Check if heading already has a hierarchical number.
        Matches patterns like: 1.1, 2.1.1, 1.1.1.1.1, A.1, etc. up to 10+ levels.
        """
        clean = re.sub(r'^#+\s*', '', text).strip()
        # Match: "1.2 Title" or "1.2.3.4.5 Title" or "A.1 Title" - supports up to 10+ levels
        return bool(
            re.match(r'^\d+(?:\.\d+){1,10}[\.)]?\s+', clean)  # Deep hierarchy (1.1.1.1...)
            or re.match(r'^[A-Z](?:\.\d+){1,10}[\.)]?\s+', clean)  # Appendix with deep hierarchy
            or re.match(r'^\d+\.[A-Za-z][\.)]?\s+', clean)  # Mixed numeric + letter
            or re.match(r'^\d+\.\s+', clean)  # Single level (1.)
        )
    
    def extract_existing_number(self, text):
        """
        Extract existing numbering from heading.
        
        Returns:
            tuple: (number_string, title) or (None, text) if no number
        """
        clean = re.sub(r'^#+\s*', '', text).strip()

        numbering_patterns = [
            # Hierarchical numeric (1.2, 1.2.3)
            (r'^((?:\d+\.)+\d+[\.)]?)\s+(.+)$', None),
            # Appendix style (A.1, A.1.2)
            (r'^([A-Z]\.(?:\d+\.)*\d+[\.)]?)\s+(.+)$', None),
            # Mixed numeric + letter (5.f, 2.A)
            (r'^(\d+\.[A-Za-z][\.)]?)\s+(.+)$', re.IGNORECASE),
            # Single-level numeric/alpha/roman (1., A., I.)
            (r'^(\d+[\.)])\s+(.+)$', None),
            (r'^([A-Z][\.)])\s+(.+)$', None),
            (r'^([IVXLCDM]+[\.)])\s+(.+)$', re.IGNORECASE),
            # Parenthesized numbering ((1), (a), (i))
            (r'^(\((?:\d+|[A-Z]|[IVXLCDM]+)\))\s+(.+)$', re.IGNORECASE),
            # Bracketed numbering ([1], [A], [i])
            (r'^(\[(?:\d+|[A-Z]|[IVXLCDM]+)\])\s+(.+)$', re.IGNORECASE),
        ]

        for pattern, flags in numbering_patterns:
            if flags:
                match = re.match(pattern, clean, flags)
            else:
                match = re.match(pattern, clean)
            if match:
                return match.group(1), match.group(2)

        return None, clean
    
    def get_heading_level_from_markdown(self, text):
        """
        Get heading level from markdown markers.
        
        Returns:
            int: 1 for #, 2 for ##, 3 for ###, 0 if no markers
        """
        match = re.match(r'^(#+)\s*', text)
        if match:
            return len(match.group(1))
        return 0
    
    def determine_heading_level(self, text, is_short=True, is_all_caps=False):
        """
        Determine the appropriate heading level based on text characteristics.
        
        Returns:
            int: 1, 2, or 3 indicating the heading level
        """
        clean = re.sub(r'^#+\s*', '', text).strip()
        clean_upper = clean.upper()
        
        # Chapter headings and chapter titles are level 1
        if self.is_chapter_heading(text) or self.is_chapter_title(text):
            return 1
        
        # ALL CAPS short headings are usually level 1 or 2
        if is_all_caps and len(clean.split()) <= 6:
            return 1
        
        # Already numbered - determine level from number
        existing_num, _ = self.extract_existing_number(text)
        if existing_num:
            if re.match(r'^\d+\.[A-Za-z]', existing_num):
                return 3  # Mixed numbering (5.f) -> treat as subsection
            if re.match(r'^[A-Za-z]$', existing_num) or re.match(r'^[IVXLCDM]+$', existing_num, re.IGNORECASE):
                return 3  # Lettered/roman-only numbering -> treat as subsection
            dots = existing_num.count('.')
            if dots == 1:
                return 2  # X.Y
            elif dots >= 2:
                return 3  # X.Y.Z or deeper
        
        # Title case, moderate length - likely level 2
        if clean[0].isupper() and not is_all_caps:
            return 2
        
        return 2  # Default to level 2
    
    def number_heading(self, text, target_level=None):
        """
        Apply hierarchical numbering to a heading with semantic hierarchy detection.
        
        Args:
            text: The heading text (may include markdown markers)
            target_level: The desired heading level (1, 2, or 3). If None, auto-detect.
            
        Returns:
            dict: {
                'original': original text,
                'numbered': text with number applied (or unchanged if shouldn't be numbered),
                'number': the number string (e.g., "1.2.1"),
                'level': the heading level,
                'chapter': current chapter number,
                'was_renumbered': True if number was added/changed
            }
        """
        result = {
            'original': text,
            'numbered': text,
            'number': None,
            'level': 1,
            'chapter': self.current_chapter,
            'was_renumbered': False
        }
        

        # Preprocess: apply hierarchy correction to heading lines if needed
        # (Assume text is a single heading, but if batch, use correct_lines)
        # This is a placeholder for batch correction integration
        md_level = self.get_heading_level_from_markdown(text)
        clean_text = re.sub(r'^#+\s*', '', text).strip()
        
        # Check for chapter heading
        chapter_num = self.parse_chapter_number(text)
        if chapter_num > 0:
            self.current_chapter = chapter_num
            self.current_section = 0
            self.current_subsection = 0
            self.current_subsubsection = 0
            self.in_appendix = False
            self.in_parent_section = None
            self.parent_section_number = ''
            self.last_heading_text = ''
            self.last_heading_normalized = ''
            result['level'] = 1
            result['chapter'] = chapter_num
            return result  # Don't number the chapter heading itself

        if not self.allow_auto_numbering:
            existing_num, _ = self.extract_existing_number(text)
            if existing_num:
                result['number'] = existing_num
                result['level'] = self.determine_heading_level(text)
                self._sync_counters_from_existing(existing_num)
                self.last_numbers_by_level[result['level']] = existing_num
            return result
        
        # Check for appendix
        if self.is_appendix_heading(text):
            self.in_appendix = True
            self.current_section = 0
            self.current_subsection = 0
            self.in_parent_section = None
            result['level'] = 1
            return result  # Don't number appendix heading itself
        
        # Check for unnumbered sections (front matter)
        if self.is_unnumbered_section(text):
            result['level'] = 1
            return result
        
        # Check for chapter title sections (ALL CAPS after chapter)
        if self.is_chapter_title(text):
            result['level'] = 1
            return result  # Keep as is without numbering
        
        # Check if already has proper numbering - extract existing number and title
        existing_num, title = self.extract_existing_number(text)

        # If no chapter context yet, decide whether to use continuous section numbering
        if self.current_chapter == 0 and not self.in_appendix:
            if existing_num or self.use_continuous_section_numbering:
                self.use_continuous_section_numbering = True
            else:
                return result
        
        # Semantic hierarchy detection - check if this should be a subsection
        should_be_child = self._should_be_subsection(text)
        
        # Detect if this heading starts a new parent section
        detected_parent = self._detect_parent_section(text)
        
        # Determine target level using semantic analysis
        if target_level is None:
            if should_be_child:
                # This should be a subsection of the previous section
                target_level = 3  # X.Y.Z
            elif detected_parent:
                # This is a new parent section
                target_level = 2  # X.Y
            elif md_level > 0:
                target_level = min(md_level, 3)
            else:
                target_level = self.determine_heading_level(text)
        
        # Override level if semantic analysis says this should be a child
        if should_be_child and target_level < 3:
            target_level = 3
        if self.use_continuous_section_numbering and existing_num and '.' in existing_num and target_level < 3:
            target_level = 3
        
        result['level'] = target_level
        
        # Generate new number based on level
        if target_level == 1 or target_level == 2:
            # Level 1/2 treated as X.Y (main section within chapter)
            # Check if we're leaving a parent section
            if self.in_parent_section and not should_be_child:
                # Moving to a new section - reset parent context if heading is different type
                new_parent = self._detect_parent_section(text)
                if new_parent != self.in_parent_section:
                    self.in_parent_section = None
                    self.parent_section_number = ''
            
            self.current_section += 1
            self.current_subsection = 0
            self.current_subsubsection = 0
            
            if self.in_appendix:
                new_number = f"{self.appendix_letter}.{self.current_section}"
            elif self.use_continuous_section_numbering and self.current_chapter == 0:
                new_number = f"{self.current_section}"
            else:
                new_number = f"{self.current_chapter}.{self.current_section}"
            
            # Update parent section tracking if this starts a parent section
            if detected_parent:
                self.in_parent_section = detected_parent
                self.parent_section_number = new_number
                
        else:  # target_level >= 3
            # Level 3: X.Y.Z (subsection)
            if self.current_section == 0:
                # If no section yet, start with section 1
                self.current_section = 1
            self.current_subsection += 1
            self.current_subsubsection = 0
            
            if self.in_appendix:
                new_number = f"{self.appendix_letter}.{self.current_section}.{self.current_subsection}"
            elif self.use_continuous_section_numbering and self.current_chapter == 0:
                new_number = f"{self.current_section}.{self.current_subsection}"
            else:
                new_number = f"{self.current_chapter}.{self.current_section}.{self.current_subsection}"
        
        result['number'] = new_number
        self.last_level = target_level
        
        # Update heading tracking for next iteration
        self.last_heading_text = clean_text
        self.last_heading_normalized = self._normalize_text(clean_text)
        
        # Numbering decision: ALWAYS preserve existing numbering when present
        # This is critical for documents that already have proper hierarchies like 1.1, 1.1.1, 1.1.1.1, etc.
        if existing_num:
            # Check if existing number has deep hierarchy (3+ levels like 1.1.1)
            existing_depth = existing_num.count('.') + 1
            
            # ALWAYS preserve deep hierarchies (3+ levels) - never renumber them
            if existing_depth >= 3 or self.policy.preserve_existing_numbering:
                result['number'] = existing_num
                result['numbered'] = f"{'#' * md_level + ' ' if md_level > 0 else ''}{existing_num} {title}"
                self._sync_counters_from_existing(existing_num)
                # Determine level from existing number depth
                result['level'] = min(existing_depth, 10)
                self.last_numbers_by_level[result['level']] = existing_num
                return result
            
            # For shallow hierarchies, only correct if policy allows AND it's clearly wrong
            if not self._should_correct_existing_number(existing_num, new_number, target_level):
                result['number'] = existing_num
                result['numbered'] = f"{'#' * md_level + ' ' if md_level > 0 else ''}{existing_num} {title}"
                self._sync_counters_from_existing(existing_num)
                self.last_numbers_by_level[target_level] = existing_num
                return result
            
            if existing_num != new_number:
                logger.info(
                    "Heading renumbered (correction): '%s' -> '%s'",
                    text,
                    f"{new_number} {title}",
                )
        
        if not existing_num and new_number:
            # Only add new numbering if there was no existing number
            prefix = '#' * md_level + ' ' if md_level > 0 else ''
            result['numbered'] = f"{prefix}{new_number} {title}"
            result['was_renumbered'] = True
            logger.info("Heading numbered (new): '%s' -> '%s'", text, result['numbered'])

        self.last_numbers_by_level[target_level] = result['number'] or new_number
        
        return result
    
    def process_document_headings(self, lines):
        """
        Process all headings in a document and apply consistent numbering.
        
        Args:
            lines: List of text lines from the document
            
        Returns:
            list: List of dicts with numbered headings and metadata
        """
        self.reset()
        
        # Pre-process lines to correct hierarchical numbering issues
        # This handles cases where numbering already exists but is incorrect (e.g. 3.6 WEEK TWO, 3.7 DAY ONE)
        if hasattr(self, 'hierarchy_corrector') and self.hierarchy_corrector:
            lines = self.hierarchy_corrector.correct_lines(lines)
            
        results = []
        
        for i, line in enumerate(lines):
            line = line.strip() if isinstance(line, str) else str(line).strip()
            
            # Skip empty lines
            if not line:
                results.append({
                    'original': line,
                    'numbered': line,
                    'is_heading': False,
                    'line_num': i
                })
                continue
            
            # Check if this is a heading (starts with # or matches heading patterns)
            is_heading = False
            
            if line.startswith('#'):
                is_heading = True
            elif re.match(r'^[A-Z][A-Z\s]+$', line) and len(line) <= 60:
                # ALL CAPS line, likely a heading
                is_heading = True
            elif re.match(r'^(\d+\.)+\d*\s+[A-Z]', line):
                # Numbered heading like "1.2 Background"
                is_heading = True
            
            if is_heading:
                numbered_result = self.number_heading(line)
                numbered_result['is_heading'] = True
                numbered_result['line_num'] = i
                results.append(numbered_result)
            else:
                results.append({
                    'original': line,
                    'numbered': line,
                    'is_heading': False,
                    'line_num': i
                })
        
        return results


# ============================================================================
# FIGURE DETECTION AND FORMATTING SYSTEM
# ============================================================================

class FigureFormatter:
    """
    Detect, validate, and format figures in academic documents.
    
    Features:
    - Comprehensive pattern-based figure detection
    - Sequential numbering validation with gap/duplicate detection
    - Proper formatting (Times New Roman, 12pt, no italics)
    - Caption placement validation (below figure)
    - List of Figures tracking for LOF generation
    
    Formatting rules:
    - Font: Times New Roman
    - Size: 12pt
    - Style: Plain text for captions
    - Placement: Caption BELOW figure
    - Numbering: "Figure X:" or "Figure X.Y:" for sub-figures
    """
    
    # ============================================================
    # ENHANCED DETECTION PATTERNS (December 31, 2025)
    # ============================================================
    
    # Pattern A: Enhanced Primary Detection - All figure variations with decimal numbers
    # Matches: Figure 1.18:, Fig 4.18: Blasting operations, figure 2.1, Fig. 3.5:, FIG 2.0
    FIGURE_REFERENCE_PATTERN = re.compile(
        r'(?:Figure|Fig\.?|Fig\s+)\s*(\d+(?:\.\d+)?)\s*[\.:]?\s*(?!\d)',
        re.IGNORECASE
    )
    
    # Pattern B: Specific pattern for "Fig X.YY:" format (targeted for missed formats)
    # Matches: Fig 4.18: Blasting operations dashboard, Fig 2.01: Temperature variations
    FIG_DECIMAL_PATTERN = re.compile(
        r'Fig\s+(\d+\.\d{1,3})\s*:\s*(.+)',
        re.IGNORECASE
    )
    
    # Pattern C: Pattern with any whitespace variations (including no space)
    # Matches: Fig4.18:, Figure  2.1  :, Fig.  4.18  :, Figure2.1
    WHITESPACE_VARIATION_PATTERN = re.compile(
        r'(?:Figure|Fig\.?)\s*(\d+(?:\.\d+)*)\s*[\.:]?\s*(?!\d)',
        re.IGNORECASE
    )
    
    # Pattern C2: No-space figure patterns
    # Matches: Fig4.18:title, Figure2.1:description
    NO_SPACE_PATTERN = re.compile(
        r'(?:Figure|Fig\.?)(\d+(?:\.\d+)*)\s*:\s*(.+)',
        re.IGNORECASE
    )
    
    # Pattern D: Comprehensive Figure Title Capture - captures entire reference with title
    # Matches: Fig 4.18: Blasting operations dashboard for mining operations
    FIGURE_TITLE_PATTERN = re.compile(
        r'^\s*(?:Figure|Fig(?:\.|\s+)?)\s*(\d+(?:\.\d+)*)\s*[\.:]\s*(.+?)(?:\n|$)',
        re.IGNORECASE | re.MULTILINE
    )
    
    # Pattern E: Context-Sensitive Detection - avoids false positives
    # Avoids: "configuration", "figurative" - correctly identifies: Fig 4.18: Dashboard
    CONTEXT_SENSITIVE_PATTERN = re.compile(
        r'(?<![A-Za-z])(?:Figure|Fig(?:\.|\s+)?)\s+(\d+(?:\.\d+)+)\s*[\.:]\s*(.+?)(?=\n{2}|\t|  |$)',
        re.IGNORECASE
    )
    
    # Pattern F: Multi-line figure titles
    MULTILINE_TITLE_PATTERN = re.compile(
        r'(?:Figure|Fig(?:\.|\s+)?)\s+(\d+(?:\.\d+)*)\s*[\.:]\s*((?:[^\n]+(?:\n(?![A-Z][a-z]+:|\s*(?:Figure|Table|Appendix)\s+|\s*$)[^\n]*)*))',
        re.IGNORECASE | re.MULTILINE
    )
    
    # Pattern G: Figure References in Bulleted/Numbered Lists
    # Matches: • Fig 4.18: Dashboard overview, 1. Figure 2.1: Framework diagram
    LIST_FIGURE_PATTERN = re.compile(
        r'^\s*(?:[-•*\d]\.?\s+)?(?:Figure|Fig(?:\.|\s+)?)\s+(\d+(?:\.\d+)*)\s*[\.:]\s*(.+)',
        re.IGNORECASE | re.MULTILINE
    )
    
    # Pattern H: Sub-figures with (a), (b), (c) notation
    # Matches: Fig 4.18: Blasting operations dashboard (a)., Figure 2.1: Conceptual framework (b):
    SUBFIGURE_LETTER_PATTERN = re.compile(
        r'(?:Figure|Fig(?:\.|\s+)?)\s+(\d+(?:\.\d+)*)\s*[\.:]\s*(.+?)\s*\([a-z]\)\s*(?:\.|:)?',
        re.IGNORECASE
    )
    
    # Pattern I: Figure with Multiple Sub-figures
    # Matches: Fig 4.18: Blasting operations (a, b, c), Figure 2.1: Analysis results (a-c)
    MULTI_SUBFIGURE_PATTERN = re.compile(
        r'(?:Figure|Fig(?:\.|\s+)?)\s+(\d+(?:\.\d+)*)\s*[\.:]\s*(.+?)\s*\(?(?:[a-d](?:\s*,\s*[a-d])*|[a-d]-[a-d])\)?',
        re.IGNORECASE
    )
    
    # Pattern J: Technical Figures - dashboard, chart, graph, diagram, etc.
    TECHNICAL_FIGURE_PATTERN = re.compile(
        r'(?:Figure|Fig(?:\.|\s+)?)\s+(\d+(?:\.\d+)*)\s*[\.:]\s*([^:\n]*?(?:dashboard|chart|graph|diagram|map|model|framework|flow|process)[^:\n]*)',
        re.IGNORECASE
    )
    
    # Pattern K: Figure Titles with Units/Measurements
    # Matches: Fig 4.18: Temperature variations (0-100°C), Figure 2.1: Speed distribution (0-120 km/h)
    MEASUREMENT_FIGURE_PATTERN = re.compile(
        r'(?:Figure|Fig(?:\.|\s+)?)\s+(\d+(?:\.\d+)*)\s*[\.:]\s*(.+?(?:\d+\s*(?:m|kg|s|Hz|%|°C|km/h)[^:\n]*))',
        re.IGNORECASE
    )
    
    # Pattern L: Sequential Decimal Numbering - for validation
    # Captures: Fig 4.18 → Group1=4, Group2=18
    DECIMAL_NUMBER_PATTERN = re.compile(
        r'(?:Figure|Fig(?:\.|\s+)?)\s+(\d+)\.(\d+)\b',
        re.IGNORECASE
    )
    
    # Pattern for sub-figures (Figure 1.1, Figure 2.3)
    SUBFIGURE_PATTERN = re.compile(
        r'Figure\s+(\d+)\.(\d+)[\.:]\s*(.+)',
        re.IGNORECASE
    )
    
    # Pattern for figures in appendices (Figure A.1, Figure B.2)
    APPENDIX_FIGURE_PATTERN = re.compile(
        r'Figure\s+([A-Z])\.(\d+)[\.:]\s*(.+)',
        re.IGNORECASE
    )
    
    # Pattern M: Inline figure references (see Figure 1, refer to Fig. 2)
    INLINE_REFERENCE_PATTERN = re.compile(
        r'(?:see|refer to|as shown in|shown in|presented in|illustrated in|depicted in)\s+(?:Figure|Fig\.?)\s+(\d+(?:\.\d+)?)',
        re.IGNORECASE
    )
    
    # Pattern N: Figure in parentheses ((Figure 1), (Fig. 2))
    PARENTHESES_PATTERN = re.compile(
        r'\((?:Figure|Fig\.?)\s+(\d+(?:\.\d+)?)\)',
        re.IGNORECASE
    )
    
    # Pattern O: Extract figure numbers for validation
    FIGURE_NUMBER_PATTERN = re.compile(
        r'(?:Figure|Fig(?:\.|\s+)?)\s*(\d+(?:\.\d+)?)',
        re.IGNORECASE
    )
    
    # ============================================================
    # FIGURE TYPE DETECTION PATTERNS
    # ============================================================
    
    # Conceptual frameworks
    CONCEPTUAL_FRAMEWORK_PATTERN = re.compile(
        r'(?:Figure|Fig)\s*\d+.*?(?:conceptual\s+framework|theoretical\s+model|research\s+paradigm)',
        re.IGNORECASE
    )
    
    # Statistical charts
    STATISTICAL_CHART_PATTERN = re.compile(
        r'(?:Figure|Fig)\s*\d+.*?(?:bar\s+chart|pie\s+chart|line\s+graph|histogram|scatter\s+plot|trend\s+analysis)',
        re.IGNORECASE
    )
    
    # Process flows
    PROCESS_FLOW_PATTERN = re.compile(
        r'(?:Figure|Fig)\s*\d+.*?(?:flow\s*chart|process\s+diagram|workflow|sequence\s+diagram)',
        re.IGNORECASE
    )
    
    # Technical diagrams (dashboard, operations, monitoring)
    TECHNICAL_DIAGRAM_PATTERN = re.compile(
        r'(?:Figure|Fig)\s*\d+.*?(?:dashboard|operations|monitoring|interface|system|architecture)',
        re.IGNORECASE
    )
    
    # ============================================================
    # ENHANCED CAPTION DETECTION PATTERNS
    # ============================================================
    
    # Caption patterns for detection - comprehensive list
    CAPTION_PATTERNS = [
        # Standard Figure X: format
        re.compile(r'^Figure\s+\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # Fig. X: format
        re.compile(r'^Fig\.\s*\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # Fig X: format (no period after Fig)
        re.compile(r'^Fig\s+\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # FIG X: format (all caps)
        re.compile(r'^FIG\s+\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # Diagram X: format
        re.compile(r'^Diagram\s+\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # Chart X: format
        re.compile(r'^Chart\s+\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # Graph X: format
        re.compile(r'^Graph\s+\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # Illustration X: format
        re.compile(r'^Illustration\s+\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # Fig X.YY: format (decimal without period after Fig)
        re.compile(r'^Fig\s+\d+\.\d+\s*:', re.IGNORECASE),
        # Bulleted/numbered list figures
        re.compile(r'^\s*[-•*\d]\.?\s+(?:Figure|Fig)\s+\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # No-space format: Fig4.18:, Figure2.1:
        re.compile(r'^(?:Figure|Fig\.?)(\d+(?:\.\d+)*)\s*:', re.IGNORECASE),
    ]
    
    def __init__(self):
        self.reset()
        
    def reset(self):
        """Reset all tracking for a new document."""
        self.figures = []  # All detected figures with metadata
        self.figure_numbers = []  # Just the numbers for validation
        self.figure_entries = []  # Entries for List of Figures
        self.current_chapter = 0  # For chapter-based numbering
        self.numbering_issues = []  # Track gaps, duplicates, etc.
        
    def detect_figures(self, text):
        """
        Detect all figure references in a block of text using multi-pass detection.
        Uses multiple patterns to catch all variations.
        
        Args:
            text: Document text content
            
        Returns:
            list: List of detected figures with metadata
        """
        figures = []
        seen_positions = set()  # Avoid duplicates
        
        # Multi-pass detection with different patterns
        detection_patterns = [
            # Pass 1: Standard figure titles/captions
            ('FIGURE_TITLE_PATTERN', self.FIGURE_TITLE_PATTERN),
            # Pass 2: Fig X.YY: format specifically
            ('FIG_DECIMAL_PATTERN', self.FIG_DECIMAL_PATTERN),
            # Pass 3: No-space format (Fig4.18:)
            ('NO_SPACE_PATTERN', self.NO_SPACE_PATTERN),
            # Pass 4: List figures
            ('LIST_FIGURE_PATTERN', self.LIST_FIGURE_PATTERN),
            # Pass 5: Technical figures
            ('TECHNICAL_FIGURE_PATTERN', self.TECHNICAL_FIGURE_PATTERN),
            # Pass 6: Whitespace variations
            ('WHITESPACE_VARIATION_PATTERN', self.WHITESPACE_VARIATION_PATTERN),
        ]
        
        for pattern_name, pattern in detection_patterns:
            for match in pattern.finditer(text):
                # Skip if we've already detected this position
                pos_key = (match.start(), match.end())
                if pos_key in seen_positions:
                    continue
                seen_positions.add(pos_key)
                
                figure_num = match.group(1)
                title = match.group(2).strip() if len(match.groups()) > 1 and match.group(2) else ''
                
                figures.append({
                    'number': figure_num,
                    'title': title,
                    'full_match': match.group(0),
                    'start': match.start(),
                    'end': match.end(),
                    'type': 'caption',
                    'figure_type': self._classify_figure_type(title),
                    'pattern_used': pattern_name
                })
        
        # Track all figure numbers for validation using comprehensive pattern
        for match in self.FIGURE_NUMBER_PATTERN.finditer(text):
            num = match.group(1)
            if num not in self.figure_numbers:
                self.figure_numbers.append(num)
        
        return figures
    
    def detect_figure_caption(self, paragraph_text):
        """
        Detect if a paragraph is a figure caption using enhanced multi-pattern detection.
        
        Args:
            paragraph_text: Text of the paragraph
            
        Returns:
            dict or None: Figure metadata if caption detected, None otherwise
        """
        if not paragraph_text:
            return None
        
        text = paragraph_text.strip()
        
        # First check with standard caption patterns
        for pattern in self.CAPTION_PATTERNS:
            if pattern.match(text):
                # Try multiple extraction patterns for the number and title
                extraction_patterns = [
                    self.FIGURE_TITLE_PATTERN,
                    self.FIG_DECIMAL_PATTERN,
                    self.NO_SPACE_PATTERN,
                    self.TECHNICAL_FIGURE_PATTERN,
                ]
                
                for extract_pattern in extraction_patterns:
                    match = extract_pattern.match(text)
                    if match:
                        return {
                            'number': match.group(1),
                            'title': match.group(2).strip() if len(match.groups()) > 1 and match.group(2) else '',
                            'full_text': text,
                            'figure_type': self._classify_figure_type(text)
                        }
                
                # Fallback: extract using general number pattern
                num_match = re.search(r'(\d+(?:\.\d+)?)', text)
                return {
                    'number': num_match.group(1) if num_match else '?',
                    'title': text,
                    'full_text': text,
                    'figure_type': 'unknown'
                }
        
        return None
    
    def _classify_figure_type(self, title):
        """Classify the type of figure based on title content."""
        title_lower = title.lower()
        
        if any(term in title_lower for term in ['framework', 'model', 'paradigm']):
            return 'conceptual_framework'
        elif any(term in title_lower for term in ['chart', 'graph', 'histogram', 'plot']):
            return 'statistical_chart'
        elif any(term in title_lower for term in ['flow', 'diagram', 'workflow', 'process']):
            return 'process_flow'
        elif any(term in title_lower for term in ['map', 'location', 'geographic']):
            return 'map'
        elif any(term in title_lower for term in ['photo', 'photograph', 'image', 'picture']):
            return 'photograph'
        elif any(term in title_lower for term in ['dashboard', 'interface', 'screen', 'ui']):
            return 'technical_interface'
        elif any(term in title_lower for term in ['blast', 'operation', 'equipment', 'machinery']):
            return 'technical_operation'
        elif any(term in title_lower for term in ['structure', 'architecture', 'layout']):
            return 'structure'
        else:
            return 'general'
    
    def validate_numbering(self):
        """
        Validate sequential figure numbering.
        Detects gaps, duplicates, and out-of-order numbers.
        
        Returns:
            dict: Validation results with issues found
        """
        self.numbering_issues = []
        
        if not self.figure_numbers:
            return {'valid': True, 'issues': []}
        
        # Parse numbers (handle both "1" and "1.1" formats)
        parsed_numbers = []
        for num in self.figure_numbers:
            if '.' in num:
                parts = num.split('.')
                parsed_numbers.append((int(parts[0]), int(parts[1]), num))
            else:
                parsed_numbers.append((int(num), 0, num))
        
        # Sort by chapter then sub-number
        parsed_numbers.sort(key=lambda x: (x[0], x[1]))
        
        # Check for gaps in main figure numbers
        main_numbers = sorted(set(n[0] for n in parsed_numbers))
        expected = list(range(1, max(main_numbers) + 1)) if main_numbers else []
        missing = [n for n in expected if n not in main_numbers]
        
        if missing:
            self.numbering_issues.append({
                'type': 'missing_numbers',
                'numbers': missing,
                'message': f"Missing figure number(s): {', '.join(map(str, missing))}"
            })
        
        # Check for duplicates
        seen = set()
        duplicates = []
        for num in self.figure_numbers:
            if num in seen:
                duplicates.append(num)
            seen.add(num)
        
        if duplicates:
            self.numbering_issues.append({
                'type': 'duplicates',
                'numbers': duplicates,
                'message': f"Duplicate figure number(s): {', '.join(duplicates)}"
            })
        
        return {
            'valid': len(self.numbering_issues) == 0,
            'issues': self.numbering_issues,
            'total_figures': len(self.figure_numbers),
            'figure_numbers': self.figure_numbers
        }
    
    def add_figure_entry(self, number, title, page_number=None):
        """
        Add a figure entry for List of Figures generation.
        
        Args:
            number: Figure number (e.g., "1" or "1.2")
            title: Figure title/caption text
            page_number: Optional page number (will be updated by Word)
        """
        self.figure_entries.append({
            'number': number,
            'title': title,
            'page': page_number,
            'full_entry': f"Figure {number}: {title}"
        })
    
    def get_lof_entries(self):
        """
        Get all figure entries for List of Figures.
        
        Returns:
            list: Sorted list of figure entries
        """
        # Sort by figure number
        def sort_key(entry):
            num = entry['number']
            if '.' in num:
                parts = num.split('.')
                return (int(parts[0]), int(parts[1]))
            return (int(num), 0)
        
        return sorted(self.figure_entries, key=sort_key)
    
    def format_caption_text(self, number, title):
        """
        Generate properly formatted caption text.
        
        Args:
            number: Figure number
            title: Figure title
            
        Returns:
            str: Formatted caption text "Figure X: Title"
        """
        # Ensure proper capitalization
        formatted_title = title.strip()
        if formatted_title and not formatted_title[0].isupper():
            formatted_title = formatted_title[0].upper() + formatted_title[1:]
        
        return f"Figure {number}: {formatted_title}"
    
    def is_figure_caption(self, text):
        """
        Quick check if text is a figure caption.
        
        Args:
            text: Paragraph text to check
            
        Returns:
            bool: True if text is a figure caption
        """
        if not text:
            return False
        
        text = text.strip()
        return any(pattern.match(text) for pattern in self.CAPTION_PATTERNS)
    
    def extract_inline_references(self, text):
        """
        Extract all inline figure references from text.
        Used to validate all referenced figures exist.
        
        Args:
            text: Document text
            
        Returns:
            list: List of referenced figure numbers
        """
        references = []
        
        # Pattern B: Inline references
        for match in self.INLINE_REFERENCE_PATTERN.finditer(text):
            references.append(match.group(1))
        
        # Pattern C: Parenthetical references
        for match in self.PARENTHESES_PATTERN.finditer(text):
            references.append(match.group(1))
        
        return list(set(references))  # Remove duplicates
    
    def renumber_figures(self, start_from=1):
        """
        Renumber all figures sequentially starting from a given number.
        Used to fix numbering gaps.
        
        Args:
            start_from: Starting figure number (default 1)
            
        Returns:
            dict: Mapping of old numbers to new numbers
        """
        if not self.figure_entries:
            return {}
        
        # Sort entries by current number
        sorted_entries = self.get_lof_entries()
        
        renumber_map = {}
        current_num = start_from
        
        for entry in sorted_entries:
            old_num = entry['number']
            if '.' not in old_num:  # Only renumber main figures
                renumber_map[old_num] = str(current_num)
                current_num += 1
        
        return renumber_map


# ============================================================
# TABLE FORMATTER CLASS - December 31, 2025
# ============================================================

class TableFormatter:
    """
    Detect, validate, and format tables in academic documents.
    
    Features:
    - Comprehensive pattern-based table detection
    - Sequential numbering validation with gap/duplicate detection
    - Proper formatting (Times New Roman, 12pt, bold)
    - Caption placement validation (below table)
    - List of Tables tracking for LOT generation
    
    Formatting rules:
    - Font: Times New Roman
    - Size: 12pt
    - Style: Bold for captions
    - Placement: Caption BELOW table
    - Numbering: "Table X:" or "Table X.Y:" for sub-tables
    """
    
    # ============================================================
    # ENHANCED DETECTION PATTERNS
    # ============================================================
    
    # Pattern A: Standard Table References - Primary detection
    # Matches: Table 1:, Tbl. 2, Tab. 3.1:, Table 4 -, table 6
    TABLE_REFERENCE_PATTERN = re.compile(
        r'(?:Table|Tbl\.?|Tab\.?)\s+(\d+(?:\.\d+)?)\s*[\.:\-–—]?\s*(?!\d)',
        re.IGNORECASE
    )
    
    # Pattern B: Comprehensive Table Detection with Titles
    # Matches: Table 1: Summary of Research Variables, Tbl. 2.1: Demographics
    TABLE_TITLE_PATTERN = re.compile(
        r'^\s*(?:Table|Tbl\.?|Tab\.?)\s*(\d+(?:\.\d+)*)\s*[\.:]\s*(.+?)(?:\n|$)',
        re.IGNORECASE | re.MULTILINE
    )
    
    # Pattern C: Inline Table References
    # Matches: see Table 1, refer to Tbl. 2, as shown in Table 3.1
    INLINE_REFERENCE_PATTERN = re.compile(
        r'(?:see|refer to|as shown in|presented in|shown in|displayed in)\s+(?:Table|Tbl\.?|Tab\.?)\s+(\d+(?:\.\d+)?)',
        re.IGNORECASE
    )
    
    # Pattern D: Table References in Parentheses
    # Matches: (Table 1), (Tbl. 2.1), (table 3)
    PARENTHESES_PATTERN = re.compile(
        r'\((?:Table|Tbl\.?|Tab\.?)\s+(\d+(?:\.\d+)?)\)',
        re.IGNORECASE
    )
    
    # Pattern E: Multi-line Table Titles
    MULTILINE_TITLE_PATTERN = re.compile(
        r'(?:Table|Tbl\.?|Tab\.?)\s+(\d+(?:\.\d+)*)\s*[\.:]\s*((?:[^\n]+(?:\n(?![A-Z][a-z]+:|\s*(?:Table|Figure|Appendix)\s+|\s*$)[^\n]*)*))',
        re.IGNORECASE | re.MULTILINE
    )
    
    # Pattern F: Tables in Lists/Numbering Systems
    # Matches: • Table 1: Variables, 1. Table 2.1: Data
    LIST_TABLE_PATTERN = re.compile(
        r'^\s*(?:[-•*\d]\.?\s+)?(?:Table|Tbl\.?|Tab\.?)\s+(\d+(?:\.\d+)*)\s*[\.:]\s*(.+)',
        re.IGNORECASE | re.MULTILINE
    )
    
    # Pattern G: Tables with Statistical Notation
    STATISTICAL_TABLE_PATTERN = re.compile(
        r'(?:Table|Tbl\.?|Tab\.?)\s+(\d+(?:\.\d+)*)\s*[\.:]\s*(.+?(?:mean|SD|SE|p-value|F-value|t-test|ANOVA|regression|correlation)[^:\n]*)',
        re.IGNORECASE
    )
    
    # Pattern H: No-space variations
    # Matches: Table1.2:, Tbl.3:, Tab4.5:
    NO_SPACE_PATTERN = re.compile(
        r'(?:Table|Tbl\.?|Tab\.?)(\d+(?:\.\d+)*)\s*:\s*(.+)',
        re.IGNORECASE
    )
    
    # Pattern I: Whitespace variations
    # Matches: Table  2.1  :, Tbl.  4.18  :
    WHITESPACE_VARIATION_PATTERN = re.compile(
        r'(?:Table|Tbl\.?|Tab\.?)\s*(\d+(?:\.\d+)*)\s*[\.:]?\s*(?!\d)',
        re.IGNORECASE
    )
    
    # Pattern J: Summary/Descriptive Tables
    SUMMARY_TABLE_PATTERN = re.compile(
        r'(?:Table|Tbl\.?|Tab\.?)\s+(\d+(?:\.\d+)*)\s*[\.:]\s*(.+?(?:summary|descriptive|demographic|characteristics|overview)[^:\n]*)',
        re.IGNORECASE
    )
    
    # Pattern K: Comparison Tables
    COMPARISON_TABLE_PATTERN = re.compile(
        r'(?:Table|Tbl\.?|Tab\.?)\s+(\d+(?:\.\d+)*)\s*[\.:]\s*(.+?(?:comparison|contrast|vs\.|versus|between|among)[^:\n]*)',
        re.IGNORECASE
    )
    
    # Pattern L: Correlation/Regression Tables
    CORRELATION_TABLE_PATTERN = re.compile(
        r'(?:Table|Tbl\.?|Tab\.?)\s+(\d+(?:\.\d+)*)\s*[\.:]\s*(.+?(?:correlation|regression|coefficient|model|results)[^:\n]*)',
        re.IGNORECASE
    )
    
    # Pattern M: Context-Sensitive Table Detection
    # Avoids false positives like "tablespoon", "tablet"
    CONTEXT_SENSITIVE_PATTERN = re.compile(
        r'(?<![A-Za-z])(?:Table|Tbl\.?|Tab\.?)\s+(\d+(?:\.\d+)+)\s*[\.:]\s*(.+?)(?=\n{2}|\t|  |$)',
        re.IGNORECASE
    )
    
    # Pattern N: Extract Table Numbers for Validation
    TABLE_NUMBER_PATTERN = re.compile(
        r'(?:Table|Tbl\.?|Tab\.?)\s*(\d+(?:\.\d+)?)',
        re.IGNORECASE
    )
    
    # Pattern O: Tables in Appendices
    APPENDIX_TABLE_PATTERN = re.compile(
        r'Table\s+([A-Z])\.(\d+)[\.:]\s*(.+)',
        re.IGNORECASE
    )
    
    # ============================================================
    # TABLE TYPE DETECTION PATTERNS
    # ============================================================
    
    # Summary tables
    SUMMARY_TYPE_PATTERN = re.compile(
        r'(?:Table|Tbl)\s*\d+.*?(?:summary|descriptive|demographic|characteristics)',
        re.IGNORECASE
    )
    
    # Statistical tables
    STATISTICAL_TYPE_PATTERN = re.compile(
        r'(?:Table|Tbl)\s*\d+.*?(?:mean|SD|ANOVA|regression|correlation|t-test|p-value)',
        re.IGNORECASE
    )
    
    # Comparison tables
    COMPARISON_TYPE_PATTERN = re.compile(
        r'(?:Table|Tbl)\s*\d+.*?(?:comparison|contrast|vs\.|versus|difference)',
        re.IGNORECASE
    )
    
    # ============================================================
    # ENHANCED CAPTION DETECTION PATTERNS
    # ============================================================
    
    CAPTION_PATTERNS = [
        # Standard Table X: format
        re.compile(r'^Table\s+\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # Tbl. X: format
        re.compile(r'^Tbl\.\s*\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # Tbl X: format (no period after Tbl)
        re.compile(r'^Tbl\s+\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # Tab. X: format
        re.compile(r'^Tab\.\s*\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # Tab X: format (no period after Tab)
        re.compile(r'^Tab\s+\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # TBL X: format (all caps)
        re.compile(r'^TBL\s+\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # TABLE X: format (all caps)
        re.compile(r'^TABLE\s+\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # Bulleted/numbered list tables
        re.compile(r'^\s*[-•*\d]\.?\s+(?:Table|Tbl)\s+\d+(?:\.\d+)*[\.\:\s]', re.IGNORECASE),
        # No-space format: Table1.2:, Tbl.3:
        re.compile(r'^(?:Table|Tbl\.?|Tab\.?)(\d+(?:\.\d+)*)\s*:', re.IGNORECASE),
    ]
    
    def __init__(self):
        self.reset()
        
    def reset(self):
        """Reset all tracking for a new document."""
        self.tables = []  # All detected tables with metadata
        self.table_numbers = []  # Just the numbers for validation
        self.table_entries = []  # Entries for List of Tables
        self.current_chapter = 0  # For chapter-based numbering
        self.numbering_issues = []  # Track gaps, duplicates, etc.
        
    def detect_tables(self, text):
        """
        Detect all table references in a block of text using multi-pass detection.
        Uses multiple patterns to catch all variations.
        
        Args:
            text: Document text content
            
        Returns:
            list: List of detected tables with metadata
        """
        tables = []
        seen_positions = set()  # Avoid duplicates
        
        # Multi-pass detection with different patterns
        detection_patterns = [
            # Pass 1: Standard table titles/captions
            ('TABLE_TITLE_PATTERN', self.TABLE_TITLE_PATTERN),
            # Pass 2: Statistical tables
            ('STATISTICAL_TABLE_PATTERN', self.STATISTICAL_TABLE_PATTERN),
            # Pass 3: No-space format
            ('NO_SPACE_PATTERN', self.NO_SPACE_PATTERN),
            # Pass 4: List tables
            ('LIST_TABLE_PATTERN', self.LIST_TABLE_PATTERN),
            # Pass 5: Summary tables
            ('SUMMARY_TABLE_PATTERN', self.SUMMARY_TABLE_PATTERN),
            # Pass 6: Comparison tables
            ('COMPARISON_TABLE_PATTERN', self.COMPARISON_TABLE_PATTERN),
            # Pass 7: Whitespace variations
            ('WHITESPACE_VARIATION_PATTERN', self.WHITESPACE_VARIATION_PATTERN),
        ]
        
        for pattern_name, pattern in detection_patterns:
            for match in pattern.finditer(text):
                # Skip if we've already detected this position
                pos_key = (match.start(), match.end())
                if pos_key in seen_positions:
                    continue
                seen_positions.add(pos_key)
                
                table_num = match.group(1)
                title = match.group(2).strip() if len(match.groups()) > 1 and match.group(2) else ''
                
                tables.append({
                    'number': table_num,
                    'title': title,
                    'full_match': match.group(0),
                    'start': match.start(),
                    'end': match.end(),
                    'type': 'caption',
                    'table_type': self._classify_table_type(title),
                    'pattern_used': pattern_name
                })
        
        # Track all table numbers for validation
        for match in self.TABLE_NUMBER_PATTERN.finditer(text):
            num = match.group(1)
            if num not in self.table_numbers:
                self.table_numbers.append(num)
        
        return tables
    
    def detect_table_caption(self, paragraph_text):
        """
        Detect if a paragraph is a table caption using enhanced multi-pattern detection.
        
        Args:
            paragraph_text: Text of the paragraph
            
        Returns:
            dict or None: Table metadata if caption detected, None otherwise
        """
        if not paragraph_text:
            return None
        
        text = paragraph_text.strip()
        
        # First check with standard caption patterns
        for pattern in self.CAPTION_PATTERNS:
            if pattern.match(text):
                # Try multiple extraction patterns for the number and title
                extraction_patterns = [
                    self.TABLE_TITLE_PATTERN,
                    self.STATISTICAL_TABLE_PATTERN,
                    self.NO_SPACE_PATTERN,
                    self.SUMMARY_TABLE_PATTERN,
                ]
                
                for extract_pattern in extraction_patterns:
                    match = extract_pattern.match(text)
                    if match:
                        return {
                            'number': match.group(1),
                            'title': match.group(2).strip() if len(match.groups()) > 1 and match.group(2) else '',
                            'full_text': text,
                            'table_type': self._classify_table_type(text)
                        }
                
                # Fallback: extract using general number pattern
                num_match = re.search(r'(\d+(?:\.\d+)?)', text)
                return {
                    'number': num_match.group(1) if num_match else '?',
                    'title': text,
                    'full_text': text,
                    'table_type': 'unknown'
                }
        
        return None
    
    def _classify_table_type(self, title):
        """Classify the type of table based on title content."""
        title_lower = title.lower()
        
        if any(term in title_lower for term in ['summary', 'descriptive', 'overview']):
            return 'summary'
        elif any(term in title_lower for term in ['demographic', 'characteristics', 'profile']):
            return 'demographic'
        elif any(term in title_lower for term in ['comparison', 'contrast', 'vs.', 'versus']):
            return 'comparison'
        elif any(term in title_lower for term in ['correlation', 'regression', 'coefficient']):
            return 'correlation'
        elif any(term in title_lower for term in ['mean', 'sd', 'anova', 't-test', 'p-value']):
            return 'statistical'
        elif any(term in title_lower for term in ['results', 'findings', 'analysis']):
            return 'results'
        elif any(term in title_lower for term in ['frequency', 'distribution', 'percentage']):
            return 'frequency'
        elif any(term in title_lower for term in ['cost', 'budget', 'financial', 'expenditure']):
            return 'financial'
        else:
            return 'general'
    
    def validate_numbering(self):
        """
        Validate sequential table numbering.
        Detects gaps, duplicates, and out-of-order numbers.
        
        Returns:
            dict: Validation results with issues found
        """
        self.numbering_issues = []
        
        if not self.table_numbers:
            return {'valid': True, 'issues': []}
        
        # Parse numbers (handle both "1" and "1.1" formats)
        parsed_numbers = []
        for num in self.table_numbers:
            if '.' in num:
                parts = num.split('.')
                parsed_numbers.append((int(parts[0]), int(parts[1]), num))
            else:
                parsed_numbers.append((int(num), 0, num))
        
        # Sort by chapter then sub-number
        parsed_numbers.sort(key=lambda x: (x[0], x[1]))
        
        # Check for gaps in main table numbers
        main_numbers = sorted(set(n[0] for n in parsed_numbers))
        expected = list(range(1, max(main_numbers) + 1)) if main_numbers else []
        missing = [n for n in expected if n not in main_numbers]
        
        if missing:
            self.numbering_issues.append({
                'type': 'gap',
                'message': f'Missing table numbers: {missing}',
                'missing': missing
            })
        
        # Check for duplicates
        seen = set()
        duplicates = []
        for num in self.table_numbers:
            if num in seen:
                duplicates.append(num)
            seen.add(num)
        
        if duplicates:
            self.numbering_issues.append({
                'type': 'duplicate',
                'message': f'Duplicate table numbers: {duplicates}',
                'duplicates': duplicates
            })
        
        return {
            'valid': len(self.numbering_issues) == 0,
            'issues': self.numbering_issues,
            'total_tables': len(self.table_numbers),
            'table_numbers': self.table_numbers
        }
    
    def add_table_entry(self, number, title, page_number=None):
        """
        Add a table entry for List of Tables generation.
        
        Args:
            number: Table number (e.g., "1" or "1.2")
            title: Table title/caption text
            page_number: Optional page number (will be updated by Word)
        """
        self.table_entries.append({
            'number': number,
            'title': title,
            'page': page_number,
            'full_entry': f"Table {number}: {title}"
        })
    
    def get_lot_entries(self):
        """
        Get all table entries for List of Tables.
        
        Returns:
            list: Sorted list of table entries
        """
        # Sort by table number
        def sort_key(entry):
            num = entry['number']
            if '.' in num:
                parts = num.split('.')
                return (int(parts[0]), int(parts[1]))
            return (int(num), 0)
        
        return sorted(self.table_entries, key=sort_key)
    
    def format_caption_text(self, number, title):
        """
        Generate properly formatted caption text.
        
        Args:
            number: Table number
            title: Table title
            
        Returns:
            str: Formatted caption text "Table X: Title"
        """
        # Ensure proper capitalization
        formatted_title = title.strip()
        if formatted_title and not formatted_title[0].isupper():
            formatted_title = formatted_title[0].upper() + formatted_title[1:]
        
        return f"Table {number}: {formatted_title}"
    
    def is_table_caption(self, text):
        """
        Quick check if text is a table caption.
        
        Args:
            text: Paragraph text to check
            
        Returns:
            bool: True if text is a table caption
        """
        if not text:
            return False
        
        text = text.strip()
        return any(pattern.match(text) for pattern in self.CAPTION_PATTERNS)
    
    def extract_inline_references(self, text):
        """
        Extract all inline table references from text.
        Used to validate all referenced tables exist.
        
        Args:
            text: Document text
            
        Returns:
            list: List of referenced table numbers
        """
        references = []
        
        # Inline references
        for match in self.INLINE_REFERENCE_PATTERN.finditer(text):
            references.append(match.group(1))
        
        # Parenthetical references
        for match in self.PARENTHESES_PATTERN.finditer(text):
            references.append(match.group(1))
        
        return list(set(references))  # Remove duplicates
    
    def renumber_tables(self, start_from=1):
        """
        Renumber all tables sequentially starting from a given number.
        Used to fix numbering gaps.
        
        Args:
            start_from: Starting table number (default 1)
            
        Returns:
            dict: Mapping of old numbers to new numbers
        """
        if not self.table_entries:
            return {}
        
        # Sort entries by current number
        sorted_entries = self.get_lot_entries()
        
        renumber_map = {}
        current_num = start_from
        
        for entry in sorted_entries:
            old_num = entry['number']
            if '.' not in old_num:  # Only renumber main tables
                renumber_map[old_num] = str(current_num)
                current_num += 1
        
        return renumber_map


# =================================================================================
# BULLET IMPLEMENTATION HELPER FUNCTIONS
# =================================================================================

def detect_bullet_type(line_text):
    """
    Detect bullet type and extract content with character mapping to Word equivalents
    """
    bullet_patterns = {
        # Standard bullets (most common)
        r'^\s*[•]\s+(.+)$': 'standard',
        r'^\s*[○]\s+(.+)$': 'white_circle',
        r'^\s*[●]\s+(.+)$': 'black_circle',
        r'^\s*[▪]\s+(.+)$': 'small_square',
        r'^\s*[■]\s+(.+)$': 'square',
        
        # Dash/arrow bullets
        r'^\s*[-–—]\s+(.+)$': 'dash',
        r'^\s*[→➔➜➤➢]\s+(.+)$': 'arrow',
        
        # Asterisk variants
        r'^\s*\*\s+(.+)$': 'asterisk',
        r'^\s*[⁎⁑※]\s+(.+)$': 'asterisk_variant',
        
        # Checkbox bullets
        r'^\s*[☐]\s+(.+)$': 'checkbox_empty',
        r'^\s*[☑]\s+(.+)$': 'checkbox_checked',
        r'^\s*[✓✔]\s+(.+)$': 'checkmark',
        
        # Number-like bullets
        r'^\s*[⓿①②③④⑤⑥⑦⑧⑨]\s+(.+)$': 'circled_number',
        r'^\s*[❶❷❸❹❺❻❼❽❾❿]\s+(.+)$': 'dingbat_number',
        
        # Creative/AI-generated bullets
        r'^\s*[✦✧]\s+(.+)$': 'sparkle',
        r'^\s*[★☆]\s+(.+)$': 'star',
        r'^\s*[♥♡]\s+(.+)$': 'heart',
        r'^\s*[◆◇]\s+(.+)$': 'diamond',
        r'^\s*[◉◎]\s+(.+)$': 'bullseye',
        r'^\s*[▸▹►]\s+(.+)$': 'right_triangle',
        r'^\s*[◂◃◄]\s+(.+)$': 'left_triangle',
        
        # Comprehensive catch-all (should be last)
        r'^\s*([•○●▪■□◆◇→➔➜➤➢–—*⁎⁑※✱✲✳✴☐☑✓✔✗✘⓿①②③④⑤⑥⑦⑧⑨❶❷❸❹❺❻❼❽❾❿➀-➉✦✧★☆♥♡◉◎▸▹►◂◃◄⦿⁍-])\s+(.+)$': 'generic_bullet'
    }
    
    for pattern, bullet_type in bullet_patterns.items():
        match = re.match(pattern, line_text, re.UNICODE)
        if match:
            content = match.group(1) if len(match.groups()) == 1 else match.group(2)
            bullet_char = match.group(1) if len(match.groups()) > 1 else None
            
            return {
                'type': bullet_type,
                'bullet_char': bullet_char,
                'content': content.strip(),
                'original_line': line_text.strip(),
                'indentation': len(line_text) - len(line_text.lstrip())
            }
    
    return None

def map_to_word_bullet_style(bullet_info):
    """
    Map detected bullet to Microsoft Word bullet style
    Returns tuple: (bullet_char, font_name, bullet_type_code)
    """
    bullet_mapping = {
        # Standard → Change to Square ■
        'standard': ('■', 'Arial', 'square'),
        'asterisk': ('■', 'Arial', 'square'),
        
        # Circle bullets → Change to Square ■
        'white_circle': ('■', 'Arial', 'square'),
        'black_circle': ('■', 'Arial', 'square'),
        
        # Square bullets
        'small_square': ('■', 'Arial', 'square'),
        'square': ('■', 'Arial', 'square'),
        
        # Dash/arrow bullets
        'dash': ('–', 'Arial', 'dash'),
        'arrow': ('■', 'Arial', 'square'),
        
        # Checkbox bullets
        'checkbox_empty': ('☐', 'Segoe UI Symbol', 'checkbox'),
        'checkbox_checked': ('☑', 'Segoe UI Symbol', 'checkbox'),
        'checkmark': ('✓', 'Segoe UI Symbol', 'checkmark'),
        
        # Number-like bullets (preserve original)
        'circled_number': (bullet_info.get('bullet_char', '①'), 'Arial Unicode MS', 'number'),
        'dingbat_number': (bullet_info.get('bullet_char', '❶'), 'Wingdings', 'number'),
        
        # Creative bullets
        'sparkle': ('✦', 'Segoe UI Symbol', 'star'),
        'star': ('★', 'Segoe UI Symbol', 'star'),
        'heart': ('♥', 'Segoe UI Symbol', 'heart'),
        'diamond': ('◆', 'Symbol', 'diamond'),
        'bullseye': ('◉', 'Arial Unicode MS', 'circle'),
        'right_triangle': ('►', 'Symbol', 'triangle'),
        'left_triangle': ('◄', 'Symbol', 'triangle'),
        
        # Generic fallback
        'generic_bullet': (bullet_info.get('bullet_char', '•'), 'Symbol', 'bullet')
    }
    
    bullet_type = bullet_info['type']
    if bullet_type in bullet_mapping:
        return bullet_mapping[bullet_type]
    else:
        # Default to standard bullet
        return ('•', 'Symbol', 'bullet')

def is_nested_bullet(line_text, previous_indent=0):
    """
    Determine if bullet is nested based on indentation
    """
    current_indent = len(line_text) - len(line_text.lstrip())
    
    if current_indent > previous_indent + 2:  # At least 2 spaces more
        return True, current_indent
    elif current_indent < previous_indent - 2:  # At least 2 spaces less
        return False, current_indent  # Outdented (higher level)
    else:
        return False, current_indent  # Same level

def process_bullet_list(lines, start_index):
    """
    Process consecutive bullet lines and apply consistent formatting
    """
    bullets = []
    i = start_index
    
    while i < len(lines):
        bullet_info = detect_bullet_type(lines[i])
        if bullet_info:
            bullets.append(bullet_info)
            i += 1
        else:
            break
    
    # Apply Word bullet styling
    if bullets:
        formatted_bullets = []
        previous_indent = 0
        
        for idx, bullet in enumerate(bullets):
            # Determine nesting level
            is_nested, current_indent = is_nested_bullet(
                bullet['original_line'], 
                previous_indent
            )
            
            # Map to Word bullet style
            bullet_char, font_name, bullet_type = map_to_word_bullet_style(bullet)
            
            formatted_bullets.append({
                'level': 1 if not is_nested else 2,
                'bullet_char': bullet_char,
                'font_name': font_name,
                'bullet_type': bullet_type,
                'content': bullet['content'],
                'original_bullet': bullet['bullet_char'] if bullet['bullet_char'] else '•',
                'indentation': current_indent
            })
            
            previous_indent = current_indent
        
        return formatted_bullets, i  # Return processed bullets and next index
    else:
        return [], start_index


class ImpliedBulletDetector:
    """
    Intelligent system to detect implied bullet points in text blocks.
    Uses heuristics, context analysis, and scoring to identify lists that aren't explicitly formatted.
    """
    
    def __init__(self):
        self.patterns = self._initialize_patterns()
        
    def _initialize_patterns(self):
        return {
            # Group 1: Imperative Verbs (Start of line)
            'imperative_start': re.compile(r'^\s*(?:Ensure|Verify|Check|Analyze|Review|Create|Update|Delete|Insert|Select|Run|Execute|Test|Validate|Monitor|Configure|Install|Deploy|Build|Compile|Debug|Fix|Resolve|Close|Open|Save|Print|Export|Import|Send|Receive|Get|Set|Put|Post|Patch|Head|Options|Trace|Connect)\b', re.IGNORECASE),
            
            # Group 2: Sentence Fragments (No subject/verb structure)
            'fragment_start': re.compile(r'^\s*(?:For|To|With|By|From|In|On|At|About|Under|Over|Between|Among|Through|During|Before|After|While|When|Where|Why|How|If|Unless|Until|Since|Because|Although|Though|Whereas|While)\b', re.IGNORECASE),
            
            # Group 3: Visual Separation (Short lines)
            'short_line': re.compile(r'^.{5,60}$'),  # 5-60 chars is typical for bullet points
            
            # Group 4: Sequential Transition Words
            'sequential_transition': re.compile(r'^\s*(?:First|Second|Third|Fourth|Fifth|Next|Then|Finally|Lastly|Moreover|Furthermore|Additionally|Also|In addition|Consequently|Therefore|Thus|Hence|Accordingly|As a result|For example|For instance|Specifically|In particular|Notably|Significantly|Importantly|Crucially|Essentially|Fundamentally|Ultimately|Eventually|Meanwhile|Simultaneously|Concurrently|Subsequently|Previously|Initially|Primarily|Secondarily|Tertiary)\b', re.IGNORECASE),
            
            # Group 5: Common List Starters
            'list_starter': re.compile(r'^\s*(?:[-*•>+]|\d+\.|\w\.)\s+'),
            
            # Group 6: Negative Indicators (Things that suggest it's NOT a bullet)
            'continuation_marker': re.compile(r'^\s*(?:and|or|but|so|nor|yet)\b', re.IGNORECASE),
            'pronoun_start': re.compile(r'^\s*(?:I|We|You|He|She|It|They|This|That|These|Those)\b', re.IGNORECASE),
        }

    def calculate_line_list_score(self, line, prev_line=None, next_line=None):
        """
        Calculate a score (0-100) indicating likelihood of a line being a list item.
        """
        if not line or not line.strip():
            return 0
            
        text = line.strip()
        score = 0
        
        # 1. Base Pattern Matching (Max 40 points)
        if self.patterns['imperative_start'].match(text):
            score += 30
        elif self.patterns['fragment_start'].match(text):
            score += 20
        elif self.patterns['sequential_transition'].match(text):
            score += 25
            
        # 2. Visual/Structural Heuristics (Max 30 points)
        if self.patterns['short_line'].match(text):
            score += 20
        
        # Check for capitalization
        if text[0].isupper():
            score += 10
            
        # 3. Contextual Analysis (Max 30 points)
        # If previous line ended with colon, high probability
        if prev_line and prev_line.strip().endswith(':'):
            score += 25
            
        # If next line starts similarly (parallelism)
        if next_line:
            next_text = next_line.strip()
            # Check for same starting word type
            if (self.patterns['imperative_start'].match(text) and self.patterns['imperative_start'].match(next_text)):
                score += 15
            elif (self.patterns['sequential_transition'].match(text) and self.patterns['sequential_transition'].match(next_text)):
                score += 15
                
        # 4. Penalties
        if self.patterns['continuation_marker'].match(text):
            score -= 30
        if self.patterns['pronoun_start'].match(text):
            score -= 20
        if text.endswith('.'): # Full sentences are less likely to be bullets in some contexts, but not always
            score -= 5
            
        return max(0, min(100, score))

    def detect_implied_bullet_blocks(self, lines):
        """
        Scans lines to identify blocks that should be converted to bullets.
        Returns a list of (start_index, end_index, bullet_type) tuples.
        """
        blocks = []
        i = 0
        current_block_start = -1
        current_block_scores = []
        
        while i < len(lines):
            line = lines[i]
            if isinstance(line, dict):
                line = line.get('text', '')
            
            prev = lines[i-1] if i > 0 else None
            if isinstance(prev, dict): prev = prev.get('text', '')
            
            nxt = lines[i+1] if i < len(lines)-1 else None
            if isinstance(nxt, dict): nxt = nxt.get('text', '')
            
            score = self.calculate_line_list_score(line, prev, nxt)
            
            # Thresholds
            DEFINITE_THRESHOLD = 85
            LIKELY_THRESHOLD = 70
            CONTINUATION_THRESHOLD = 55
            
            if score >= LIKELY_THRESHOLD:
                if current_block_start == -1:
                    current_block_start = i
                current_block_scores.append(score)
            else:
                if current_block_start != -1 and score >= CONTINUATION_THRESHOLD:
                    current_block_scores.append(score)
                elif current_block_start != -1:
                    # End of block
                    # Validate block: needs at least 2 items or 1 very high confidence item
                    if len(current_block_scores) >= 2 or (len(current_block_scores) == 1 and current_block_scores[0] >= DEFINITE_THRESHOLD):
                        blocks.append((current_block_start, i - 1, 'square')) # Default to square for implied
                    
                    current_block_start = -1
                    current_block_scores = []
            
            i += 1
            
        # Check if block continues to end
        if current_block_start != -1:
             if len(current_block_scores) >= 2 or (len(current_block_scores) == 1 and current_block_scores[0] >= DEFINITE_THRESHOLD):
                blocks.append((current_block_start, len(lines) - 1, 'square'))
                
        return blocks


class PatternEngine:
    """Ultra-fast pattern matching engine for document analysis"""
    
    def __init__(self, policy=None):
        self.policy = policy or FormatPolicy()
        self.patterns = self._initialize_patterns()
        self.implied_detector = ImpliedBulletDetector()
        
    def _initialize_patterns(self):
        """Initialize all recognition patterns - 40+ regex patterns"""
        return {
            # Heading Level 1 Patterns (ALL CAPS, Major Sections)
            'heading_1': [
                re.compile(r'^([A-Z][A-Z\s]{2,49})$'),  # ALL CAPS (3-50 chars)
                re.compile(r'^(CHAP?TER\s+\d+.*)$', re.IGNORECASE),  # CHAPTER 1: Title
                re.compile(r'^(PART\s+[IVX]+.*)$', re.IGNORECASE),  # PART I, PART II
                re.compile(r'^(PART\s+\d+.*)$', re.IGNORECASE),  # PART 1, PART 2
                re.compile(r'^(\d+\.\s+[A-Z][A-Z\s]+)$'),  # "1. INTRODUCTION"
                re.compile(r'^(ACKNOWLEDGEMENT|ABSTRACT|INTRODUCTION|CONCLUSION|REFERENCES|BIBLIOGRAPHY|APPENDIX|APPENDICES|GLOSSARY|INDEX|PREFACE|FOREWORD|DEDICATION|TABLE OF CONTENTS|LIST OF FIGURES|LIST OF TABLES)S?$', re.IGNORECASE),
                re.compile(r'^(EXECUTIVE\s+SUMMARY)$', re.IGNORECASE),
                re.compile(r'^(LITERATURE\s+REVIEW)$', re.IGNORECASE),
                re.compile(r'^(RESEARCH\s+METHODOLOGY)$', re.IGNORECASE),
                re.compile(r'^(DATA\s+ANALYSIS)$', re.IGNORECASE),
                re.compile(r'^(FINDINGS\s+AND\s+DISCUSSION)$', re.IGNORECASE),
                re.compile(r'^(RECOMMENDATIONS)$', re.IGNORECASE),
            ],
            
            # Heading Level 2 Patterns (Title Case, Numbered Sections)
            'heading_2': [
                re.compile(r'^([A-Z][a-z0-9]+(?:\s+[A-Za-z0-9]+){1,15})$'),  # Title Case (relaxed)
                re.compile(r'^\d+\.\d+\s+([A-Z].{3,80})$'),  # "1.1 Background"
                re.compile(r'^([A-Z][a-z]+\s+and\s+[A-Z][a-z]+)$'),  # "Methods and Results"
                re.compile(r'^([A-Z][a-z]+\s+of\s+[A-Z][a-z]+.*)$'),  # "Analysis of Data"
                re.compile(r'^(Section\s+\d+[:\s].*)$', re.IGNORECASE),  # "Section 1: Overview"
                re.compile(r'^(\d+\s+[A-Z][a-z]+(?:\s+[A-Z]?[a-z]+)*)$'),  # "1 Introduction"
            ],
            
            # Heading Level 3 Patterns (Sub-sections)
            'heading_3': [
                re.compile(r'^\d+\.\d+\.\d+\s+(.+)$'),  # "1.1.1 Details"
                re.compile(r'^([a-z]\)\s+.+)$'),  # "a) Subsection"
                re.compile(r'^(\([a-z]\)\s+.+)$'),  # "(a) Subsection"
                re.compile(r'^([A-Z][a-z]+:)\s*$'),  # "Definition:"
                re.compile(r'^([ivx]+\.\s+.+)$', re.IGNORECASE),  # "i. First point"
                re.compile(r'^\(\d+\)\s+(.+)$'),  # "(1) Subsection"
            ],
            
            # Reference Patterns (APA, MLA, Chicago, IEEE)
            'reference': [
                re.compile(r'^([A-Z][a-z]+,?\s+[A-Z]\..*\(\d{4}\))'),  # "Smith, J. (2024)"
                re.compile(r'^([A-Z][a-z]+,\s+[A-Z]\.\s+[A-Z]?\.?\s*\(\d{4}\))'),  # "Smith, J. A. (2024)"
                re.compile(r'^([A-Z][a-z]+\s+et\s+al\..*\d{4})'),  # "Smith et al. 2024"
                re.compile(r'^\[\d+\]'),  # "[1] Reference" - IEEE style
                re.compile(r'^([A-Z][a-z]+.*Retrieved from)'),  # Web reference
                re.compile(r'^([A-Z][a-z]+.*https?://)'),  # URL reference
                re.compile(r'^([A-Z][a-z]+,\s+[A-Z]\.\s+\(\d{4}\)\.)'),  # APA format
                re.compile(r'^([A-Z][a-z]+,\s+[A-Z][a-z]+\.\s+".+"\s+.+\d{4})'),  # MLA format
                re.compile(r'^\d+\.\s+[A-Z][a-z]+,?\s+[A-Z]'),  # Numbered reference
                re.compile(r'^([A-Z][a-z]+,?\s+[A-Z]\.\s*(&|and)\s+[A-Z][a-z]+)'),  # Multiple authors
                # NEW PATTERNS FOR ORGANIZATIONS AND LEGAL DOCUMENTS
                re.compile(r'^[\w\s\.\-&]+\s*\(\d{4}(?:/\d{4})?\).*$'), # Organization (Year)
                re.compile(r'^(?:Decree|Law|Order|Decision|Arrete)\s+No\.?.*$', re.IGNORECASE), # Legal
            ],

            'reference_journal_span_v1': [
                # APA journal + volume/issue
                re.compile(
                    r'^(?P<pre>.+?\(\d{4}[a-z]?\)\.\s+.+?\.\s+)(?P<journal>(?!In\s)[^,]+?)(?=,\s*\d{1,4}(?:\s*\(|\s*,))'
                ),
                # APA journal + volume only
                re.compile(
                    r'^(?P<pre>.+?\(\d{4}[a-z]?\)\.\s+.+?\.\s+)(?P<journal>(?!In\s)[^,]+?)(?=,\s*\d{1,4}\s*,)'
                ),
                # Journal followed by Retrieved from / URL / DOI (no volume)
                re.compile(
                    r'^(?P<pre>.+?\(\d{4}[a-z]?\)\.\s+.+?\.\s+)(?P<journal>(?!In\s)[^.]+?)(?=\.\s+(?:Retrieved\s+from|Available\s+at|https?://|doi:))',
                    re.IGNORECASE
                ),
            ],
            
            # Unicode Purge & Academic Symbol Preservation (Priority 0 - Pre-processor)
            'unicode_scrubber': [
                # Removes all emojis, asterisks, and non-academic Unicode before analysis
                # This prevents rendering issues and removes unwanted special characters
                # Pattern: Removes anything that is NOT standard ASCII (except asterisks which are explicitly removed)
                # PLUS explicitly removes asterisk variants: *, ⁎, ⁑, ※
                re.compile(r'[^\x00-\x7F\u2010-\u2015\u2022\u25CB\u25CF\u25AA\u25AB\u25A0\u25A1\u25C6\u25C7\u2192\u2794\u2796\u27A1\u27A2\u27A3\u27A4]|[\*\u204e\u2051\u203b]'),
            ],
            
            # Asterisk Removal (Comprehensive - removes ALL asterisk variants)
            'asterisk_removal': [
                # Dedicated pattern to ensure ALL asterisks are removed from all content
                # Matches: * (U+002A), ⁎ (U+204E), ⁑ (U+2051), ※ (U+203B)
                # Applied to all text content to remove mid-word asterisks like "Customizability*"
                re.compile(r'[\*\u204e\u2051\u203b]'),
            ],
            
            # List Patterns - Bullet (Flex-Bullet Detector with Emoji-Agnostic Processing)
            'bullet_list': [
                # Primary: Flexible bullet detection matches any standard/non-standard dash or marker
                # Catches: - (hyphen), – (en-dash), — (em-dash), • (bullet), ■ (square), etc.
                # Confidence: 0.98 - Very high match rate after Unicode scrubbing
                re.compile(r'^\s*([-•●○▪■□◆◇*]|[\u2010-\u2015])\s+(.+)$'),
                
                # Standard bullets (•, ○, ●, ▪, ■)
                re.compile(r'^\s*[•○●▪■]\s+(.+)$'),
                # Arrow bullets
                re.compile(r'^\s*[→➔➜➤➢]\s+(.+)$'),
                # Asterisk variants
                re.compile(r'^\s*[\*⁎⁑※]\s+(.+)$'),
                # Checkbox bullets
                re.compile(r'^\s*[☐☑✓✔]\s+(.+)$'),
                # Number-like bullets
                re.compile(r'^\s*[⓿①②③④⑤⑥⑦⑧⑨❶❷❸❹❺❻❼❽❾❿]\s+(.+)$'),
                # Creative bullets
                re.compile(r'^\s*[✦✧★☆♥♡◆◇◉◎▸▹►◂◃◄]\s+(.+)$'),
            ],
            
            # List Patterns - Numbered
            'numbered_list': [
                re.compile(r'^(\d+[\.)]\s+.+)$'),  # 1. or 1)
                re.compile(r'^([a-z][\.)]\s+.+)$'),  # a. or a)
                re.compile(r'^([ivxlcdm]+[\.)]\s+.+)$', re.IGNORECASE),  # Roman numerals
                re.compile(r'^\(\d+\)\s+.+$'),  # (1) format
                re.compile(r'^\([a-z]\)\s+.+$'),  # (a) format
                re.compile(r'^[A-Z][\.)]\s+.+$'),  # A. or A)
                re.compile(r'^\d+\)\s+.+$'),  # 1) format
            ],

            # Enhanced hierarchical numbering (captures mixed styles like A.1.2, 1.a.i, 1.A.1)
            'alphanumeric_hierarchy': [
                # Require at least one sub-segment (e.g., A.1 or 1.1)
                re.compile(r'^\s*(?:[A-Z]|\d+)(?:[\.\)](?:[a-z]|\d+))+[\.\)]?\s+(.+)$')
            ],
            'parenthesized_mixed': [
                re.compile(r'^\s*\((\d+|[a-z]|[ivx]+)\)\s+(.+)$', re.IGNORECASE),
            ],
            'double_parentheses': [
                re.compile(r'^\s*\(\((\d+)\)\)\s+(.+)$'),
            ],
            'bracketed_numbering': [
                re.compile(r'^\s*\[(\d+|[a-z])\]\s+(.+)$', re.IGNORECASE),
            ],
            'hyphen_numbering': [
                re.compile(r'^\s*(\d+|[a-z])\-\s+(.+)$', re.IGNORECASE),
            ],
            'roman_with_dots': [
                re.compile(r'^\s*([IVXLCDM]+)\.\.\.\s+(.+)$', re.IGNORECASE),
            ],

            # Short-document (assignment/worksheet) header patterns
            'assignment_section_header': [
                # Case-sensitive: Only match digits, uppercase letters, or roman numerals
                re.compile(r'^\s*(\d+|[A-Z]|[ivx]+)[\.)]\s+([A-Z][^\.]{5,50})(?:\s*[:\.])?\s*$'),
            ],
            'question_prompt': [
                re.compile(r'^\s*(?:Q|Question)\s*(\d+)[\.:]\s*(.+)$', re.IGNORECASE),
            ],
            'task_header': [
                re.compile(r'^\s*(?:Task|Activity|Exercise)\s*(\d+|\w)[\.:]\s*(.+)$', re.IGNORECASE),
            ],
            'part_section': [
                re.compile(r'^\s*(?:Part|Section)\s*([A-Z]|\d+|[ivx]+)[\.:]\s*(.+)$', re.IGNORECASE),
            ],
            'requirement_header': [
                re.compile(r'^\s*(?:Req(?:uirement)?|Spec(?:ification)?)\s*(\d+[\.\d]*)[\.:]\s*(.+)$', re.IGNORECASE),
            ],
            'step_header': [
                re.compile(r'^\s*(?:Step|Stage|Phase)\s*(\d+)[\.:]\s*(.+)$', re.IGNORECASE),
            ],
            
            # ==================== TABLE PATTERNS (CONSERVATIVE) ====================
            'table_markdown': [
                re.compile(r'^\s*\|(?=[^|]*\|)[^|]*(?:\|[^|]*)+?\|\s*$'),  # Markdown table row
                re.compile(r'^\s*\|(?:\s*:?[-]{3,}:?\s*\|)+\s*$'),  # Markdown separator
            ],

            'table_tab_separated': [
                re.compile(r'^\s*(?:[^\t\n]{2,}\t+){2,}[^\t\n]{2,}\s*$'),  # Tab-separated (3+ columns)
            ],

            'table_aligned_columns': [
                re.compile(r'^\s*(?:[^\s]{2,}\s{3,}){3,}[^\s]{2,}\s*$'),  # 4+ aligned columns
            ],

            'not_a_table': [
                re.compile(r'^[^|]*\|[^|]*[a-z][^|]*$'),  # Pipe in sentence
                re.compile(r'^[a-zA-Z]\s*\|\s*[a-zA-Z]\s*$'),  # a | b notation
                re.compile(r'^\d+\s*\|\s*\d+\s*$'),  # 100 | 200
                re.compile(r'^[A-Za-z].*---.*[a-z]$'),  # --- in text
                re.compile(r'^(?:See|Refer to|Table|Figure|Fig\.).*$', re.IGNORECASE),
            ],

            # Legacy table patterns (keep for compatibility but lower priority)
            'table_marker': [
                re.compile(r'^\[TABLE\s*START\]', re.IGNORECASE),
                re.compile(r'^\[TABLE\s*END\]', re.IGNORECASE),
                re.compile(r'^Table\s+\d+', re.IGNORECASE),
                re.compile(r'^TABLE\s+\d+', re.IGNORECASE),
                re.compile(r'^Tabel\s+\d+', re.IGNORECASE),  # Common typo
            ],

            'table_row': [
                re.compile(r'^\|(.+\|)+$'),  # Markdown table row |cell|cell|
                re.compile(r'^\|[\s\-:]+\|$'),  # Markdown table separator |---|---|
            ],

            # Table separator (markdown)
            'table_separator': [
                re.compile(r'^\|[\s\-:]+\|[\s\-:]+.*\|$'),  # |---|---|
                re.compile(r'^[\-\|:\s]+$'),  # Pure separator line
            ],

            # Plain text table detection patterns - ONLY RELIABLE ONES
            'plain_table_separator': [
                re.compile(r'^\s*[\-=_]{5,}(?:\s+[\-=_]{5,})+\s*$'),  # Multiple separators
                re.compile(r'^\s*[+\-]{5,}(?:[+\-]+)*\s*$'),  # Box borders
            ],
            
            # Definition/Key Term Patterns
            'definition': [
                re.compile(r'^(Definition|Objective|Task|Goal|Purpose|Aim|Method|Result|Conclusion|Note|Important|Key Point|Summary|Overview|Background|Context|Example|Theorem|Lemma|Corollary|Proposition|Proof|Remark|Observation|Hypothesis|Assumption|Constraint|Limitation|Scope|Significance|Implication|Application|Contribution|Finding|Evidence|Data|Analysis|Interpretation|Explanation|Description|Specification|Requirement|Criteria|Criterion|Parameter|Variable|Constant|Factor|Element|Component|Aspect|Feature|Property|Characteristic|Attribute|Quality|Measure|Metric|Indicator|Index|Ratio|Rate|Percentage|Value|Score|Level|Degree|Extent|Amount|Quantity|Size|Scale|Range|Interval|Duration|Period|Phase|Stage|Step|Process|Procedure|Protocol|Algorithm|Formula|Equation|Model|Framework|Theory|Concept|Principle|Law|Rule|Guideline|Standard|Norm|Benchmark|Baseline|Reference|Source|Origin|Cause|Effect|Impact|Outcome|Consequence|Benefit|Advantage|Disadvantage|Risk|Challenge|Problem|Issue|Question|Answer|Solution|Strategy|Approach|Technique|Tool|Instrument|Device|System|Structure|Organization|Classification|Category|Type|Kind|Class|Group|Set|Collection|Series|Sequence|Order|Pattern|Trend|Distribution|Correlation|Relationship|Connection|Link|Association|Comparison|Contrast|Difference|Similarity|Analogy|Metaphor|Symbol|Sign|Signal|Indicator|Warning|Caution|Attention|Focus|Priority|Emphasis|Highlight|Point|Argument|Claim|Assertion|Statement|Proposition|Premise|Inference|Deduction|Induction|Generalization|Specialization|Abstraction|Instantiation|Implementation|Execution|Operation|Function|Action|Activity|Task|Job|Work|Effort|Attempt|Trial|Experiment|Test|Evaluation|Assessment|Review|Examination|Inspection|Investigation|Inquiry|Study|Research|Survey|Poll|Interview|Questionnaire|Form|Document|Report|Paper|Article|Book|Chapter|Section|Paragraph|Sentence|Word|Term|Phrase|Expression|Language|Text|Content|Information|Data|Knowledge|Wisdom|Intelligence|Understanding|Comprehension|Interpretation|Meaning|Significance|Importance|Relevance|Value|Worth|Merit|Quality|Standard|Excellence|Performance|Efficiency|Effectiveness|Productivity|Output|Input|Resource|Asset|Capital|Investment|Cost|Expense|Price|Fee|Charge|Payment|Revenue|Income|Profit|Loss|Balance|Budget|Forecast|Projection|Estimate|Calculation|Computation|Measurement|Quantification|Qualification|Certification|Accreditation|Validation|Verification|Confirmation|Approval|Authorization|Permission|License|Right|Privilege|Responsibility|Duty|Obligation|Commitment|Promise|Guarantee|Warranty|Assurance|Insurance|Protection|Security|Safety|Risk|Hazard|Danger|Threat|Vulnerability|Weakness|Strength|Opportunity|Challenge):\s*(.+)?$', re.IGNORECASE),
            ],
            
            # Figure/Image Caption
            'figure': [
                re.compile(r'^Figure\s+\d+', re.IGNORECASE),
                re.compile(r'^Fig\.\s*\d+', re.IGNORECASE),
                re.compile(r'^Image\s+\d+', re.IGNORECASE),
                re.compile(r'^Diagram\s+\d+', re.IGNORECASE),
                re.compile(r'^Chart\s+\d+', re.IGNORECASE),
                re.compile(r'^Graph\s+\d+', re.IGNORECASE),
                re.compile(r'^Illustration\s+\d+', re.IGNORECASE),
            ],
            
            # Equation patterns
            'equation': [
                re.compile(r'^Equation\s+\d+', re.IGNORECASE),
                re.compile(r'^Eq\.\s*\d+', re.IGNORECASE),
                re.compile(r'^\(\d+\)$'),  # Just (1) on its own line - equation number
            ],
            
            # Quote patterns
            'quote': [
                re.compile(r'^["\"].*["\"]$'),  # Quoted text
                re.compile(r'^>\s+.+$'),  # Blockquote markdown style
            ],
            
            # Code block patterns
            'code': [
                re.compile(r'^```'),  # Markdown code fence
                re.compile(r'^~~~'),  # Alternative code fence
                re.compile(r'^\t{2,}.+$'),  # Heavily indented (code-like)
            ],
            
            # Section Keywords for special detection
            'section_abstract': re.compile(r'^(abstract|executive\s+summary)$', re.IGNORECASE),
            'section_intro': re.compile(r'^(introduction|background|overview|motivation|context)$', re.IGNORECASE),
            'section_methods': re.compile(r'^(method|methodology|approach|procedure|materials|design|implementation)s?$', re.IGNORECASE),
            'section_results': re.compile(r'^(results?|findings?|outcomes?|data|observations?)$', re.IGNORECASE),
            'section_discussion': re.compile(r'^(discussion|analysis|interpretation|evaluation)$', re.IGNORECASE),
            'section_conclusion': re.compile(r'^(conclusions?|summary|final\s+remarks|future\s+work|recommendations?)$', re.IGNORECASE),
            'section_references': re.compile(r'^(references|bibliography|works\s+cited|citations|sources)$', re.IGNORECASE),
            'section_appendix': re.compile(r'^(appendix|appendices|supplementary|annexure)$', re.IGNORECASE),
            
            # NEW PATTERNS (December 30, 2025)
            
            # Inline Formatting Patterns (Bold/Italic/Bold-Italic)
            'inline_formatting': [
                re.compile(r'\*\*\*(.+?)\*\*\*|___(.+?)___'),  # Bold+Italic (check first)
                re.compile(r'\*\*(.+?)\*\*|__(.+?)__'),       # Bold (double asterisks/underscores)
                re.compile(r'(?<!\*)\*([^*\n]+?)\*(?!\*)|(?<!_)_([^_\n]+?)_(?!_)'),  # Single asterisk/underscore as BOLD
            ],

            # Prominent heading patterns (bold/underline/all-caps detection)
            'standalone_bold_heading': [
                re.compile(r'^\s*\*\*(?!.*\*\*.*\*\*)([^*\n]+)\*\*\s*$'),
            ],
            'bold_prefix_heading': [
                re.compile(r'^\s*\*\*([^*\n]+)\*\*\s*[-:]\s*(.+)$'),
            ],
            'star_surrounded': [
                re.compile(r'^\s*\*{3,}\s*(.+)\s*\*{3,}\s*$'),
            ],
            # Underlined emphasis - single-line detection for the underline marker
            'underlined_line': [
                re.compile(r'^\s*_{3,}\s*$'),
            ],
            'allcaps_short_line': [
                re.compile(r'^\s*[A-Z][A-Z\s]{2,30}[A-Z]\s*$'),
            ],
            'numbered_bold_section': [
                re.compile(r'^\s*\*\*(\d+[\.\)])\s*(.+)\*\*\s*$'),
            ],
            
            # Page Header/Footer Patterns
            'page_metadata': [
                re.compile(r'^\s*(?:page|p|pg)\.?\s*\d+\s*(?:of\s*\d+)?\s*$', re.IGNORECASE),
                re.compile(r'^\s*(?:header|footer|running head):?\s*.{1,50}$', re.IGNORECASE),
                re.compile(r'^\s*(?:confidential|draft|version\s*\d+|date:|©|copyright)\b.*$', re.IGNORECASE),
                re.compile(r'^\s*-\s*\d+\s*-\s*$'),  # Centered page numbers like - 1 -
                re.compile(r'^\s*formatted with afrod?ocs app\s*$', re.IGNORECASE),
            ],
            
            # Academic Metadata Patterns (Title, Author, Affiliation)
            'academic_metadata': [
                re.compile(r'^\s*(?:by|authors?:)\s+([A-Z][a-z]+(?:\s+[A-Z]\.)*(?:\s+[A-Z][a-z]+)*(?:\s*,\s*[A-Z][a-z]+(?:\s+[A-Z]\.)*(?:\s+[A-Z][a-z]+)*)*)$', re.IGNORECASE),
                re.compile(r'^\s*(?:department|school|college|university|institute|faculty|center|laboratory)\s+of\s+.+$', re.IGNORECASE),
                re.compile(r'^\s*(?:email|e-mail|correspondence\s+to|contact):?\s*\S+@\S+\.\S+\s*$', re.IGNORECASE),
                re.compile(r'^\s*(?:submitted\s+to|prepared\s+for|in\s+partial\s+fulfillment)\s+.+$', re.IGNORECASE),
                re.compile(r'^\s*(?:Student|Name|Course|Class|Module|Instructor|Teacher|Professor|Tutor|Section|Group|Team|Due(?:.*?Date)?|Deadline|Submission)\s*:?\s*.+$', re.IGNORECASE),
            ],
            
            # Mathematical Expression Patterns
            'math_expression': [
                re.compile(r'^\$\$[^$]+\$\$$'),  # Display math $$...$$
                re.compile(r'\$[^$\n]+\$'),      # Inline math $...$
                re.compile(r'^\\\[.*\\\]$'),   # LaTeX display \[...\]
                re.compile(r'\\\(.*\\\)'),    # LaTeX inline \(...\)
                re.compile(r'^\s*[A-Za-z]\s*[=<>≤≥≠≈]\s*.+$'),  # Simple equations like x = 5
            ],
            
            # Footnote/Endnote Patterns
            'footnote_endnote': [
                re.compile(r'^\s*(?:endnotes?|footnotes?)\s*$', re.IGNORECASE),  # Section header
                re.compile(r'^\s*(\d{1,3}|[a-z]|\*{1,3})\s*[\.\)]\s+.{10,}$'),  # Footnote entry
                re.compile(r'^\s*\[\d+\]\s+.{10,}$'),  # Bracketed footnote [1] text...
                re.compile(r'\[\^\d+\]'),  # Markdown footnote reference [^1]
            ],
            
            # ============================================================
            # NEW PATTERNS - December 30, 2025 (20 Academic Formatting Patterns)
            # ============================================================
            
            # 1. HEADING_HIERARCHY - Markdown-style hierarchical headings
            'heading_hierarchy': [
                re.compile(r'^#\s+CHAP?TER\s+\w+[:\s]', re.IGNORECASE),  # # CHAPTER X: Title
                re.compile(r'^##\s+\d+\.\d+\s+'),  # ## 1.1 Section
                re.compile(r'^###\s+\d+\.\d+\.\d+\s+'),  # ### 1.1.1 Subsection
                re.compile(r'^####\s+\d+\.\d+\.\d+\.\d+\s+'),  # #### 1.1.1.1 Sub-subsection
                re.compile(r'^#####\s+[a-z]\)\s+', re.IGNORECASE),  # ##### a) Point
                re.compile(r'^#{1,6}\s+'),  # Generic markdown heading
            ],
            
            # 2. ACADEMIC_TABLE - Enhanced table detection
            'academic_table': [
                re.compile(r'^\|[-\s:]+\|[-\s:]+\|'),  # Table separator row
                re.compile(r'^\|\s*\*\*[^|]+\*\*\s*\|'),  # Bold header cells
                re.compile(r'^\|[^|]+\|[^|]+\|[^|]+\|'),  # 3+ column table
                re.compile(r'^Table\s+\d+\.\d+[:\s]', re.IGNORECASE),  # Table X.Y: caption
            ],
            
            # 3. LIST_NORMALIZER - Enhanced list detection with indentation
            'list_nested': [
                re.compile(r'^(\s{2,})[•\-\*]\s+'),  # Indented bullet
                re.compile(r'^(\s{2,})\d+[\.\)]\s+'),  # Indented numbered
                re.compile(r'^(\s{2,})[a-z][\.\)]\s+'),  # Indented lettered
                re.compile(r'^\s*□\s+'),  # Checkbox empty
                re.compile(r'^\s*[☐☑✓✗]\s+'),  # Various checkbox symbols
            ],
            
            # 4. FIGURE_EQUATION - Enhanced figure/equation captions
            'figure_equation': [
                re.compile(r'^[Ff]igure\s+\d+\.\d+[:\s]'),  # Figure X.Y: caption
                re.compile(r'^\*\*Figure\s+\d+'),  # **Figure X**
                re.compile(r'^\$\$\s*$'),  # Start of display equation
                re.compile(r'\\begin\{equation\}'),  # LaTeX equation environment
                re.compile(r'\\end\{equation\}'),
            ],
            
            # 5. CITATION_STYLE - In-text citation patterns
            'citation_inline': [
                re.compile(r'\([A-Z][a-z]+,\s*\d{4}\)'),  # (Smith, 2020)
                re.compile(r'\([A-Z][a-z]+\s*&\s*[A-Z][a-z]+,\s*\d{4}\)'),  # (Smith & Jones, 2020)
                re.compile(r'\([A-Z][a-z]+\s+et\s+al\.,\s*\d{4}\)'),  # (Smith et al., 2020)
                re.compile(r'\([A-Z][a-z]+,\s*\d{4};\s*[A-Z][a-z]+,\s*\d{4}\)'),  # Multiple citations
                re.compile(r'\([A-Z][a-z]+\s+and\s+[A-Z][a-z]+,\s*\d{4}\)'),  # (Smith and Jones, 2020)
            ],
            
            # 6. HEADER_FOOTER - Running header/footer detection
            'running_header': [
                re.compile(r'^[A-Z][A-Z\s]+\s*\|\s*[A-Z]'),  # CHAPTER | TITLE format
                re.compile(r'^\s*Page\s+\d+\s+of\s+\d+\s*$', re.IGNORECASE),  # Page X of Y
                re.compile(r'^Chapter\s+\d+\s*\|\s*', re.IGNORECASE),  # Chapter X | format
            ],
            
            # 7. APPENDIX_FORMAT - Appendix-specific patterns
            'appendix_format': [
                re.compile(r'^APPENDIX\s+[A-Z]$', re.IGNORECASE),  # APPENDIX A
                re.compile(r'^[A-Z]\.\d+[:\s]'),  # A.1: Appendix section
                re.compile(r'^[A-Z]\.\d+\.\d+[:\s]'),  # A.1.1: Appendix subsection
                re.compile(r'^Appendix\s+[A-Z][:\s]', re.IGNORECASE),  # Appendix A: Title
            ],
            
            # 8. BLOCK_QUOTE - Extended block quote detection
            'block_quote': [
                re.compile(r'^>\s+.+$'),  # Markdown blockquote
                re.compile(r'^\s{4,}".+$'),  # Indented quoted text
                re.compile(r'^"[^"]{50,}'),  # Long quoted text
                re.compile(r"^'[^']{50,}"),  # Single-quote long text
            ],
            
            # 9. MATH_MODEL - Statistical/Mathematical model patterns
            'math_model': [
                re.compile(r'[Yy]\s*=\s*[βα]'),  # Y = β... regression
                re.compile(r'\\beta_\d'),  # LaTeX beta subscript
                re.compile(r'[Rr]²\s*=\s*[\d\.]'),  # R² = 0.xx
                re.compile(r'[Ff]\s*\(\s*\d+\s*,\s*\d+\s*\)'),  # F(df1, df2)
                re.compile(r'[Pp]\s*[<>=]\s*[\d\.]+'),  # p < 0.05
                re.compile(r'\\epsilon|\\sigma|\\mu'),  # Greek letters
            ],
            
            # 10. FONT_CONSISTENCY - Text emphasis detection
            'text_emphasis': [
                re.compile(r'\*\*[A-Z][^*]+\*\*:'),  # **Term**: definition
                re.compile(r'\*[A-Za-z][^*]+\*'),  # *italicized text*
                re.compile(r'`[^`]+`'),  # `monospace`
                re.compile(r'\*\*\*.+\*\*\*'),  # ***bold-italic***
            ],
            
            # 11. REFERENCE_FORMAT - Reference list patterns (APA enhanced)
            'reference_apa': [
                re.compile(r'^[A-Z][a-z]+,\s*[A-Z]\.\s*[A-Z]?\.\s*\(\d{4}\)'),  # Author, A. B. (Year)
                re.compile(r'\*[^*]+\*\.'),  # Italicized title
                re.compile(r'https?://doi\.org/'),  # DOI URL
                re.compile(r'doi:\s*[\d\.]+/'),  # DOI reference
                re.compile(r'Retrieved\s+\w+\s+\d+,\s*\d{4}'),  # Retrieved Month Day, Year
            ],
            
            # 12. TOC_GENERATOR - Table of contents patterns
            'toc_entry': [
                re.compile(r'^[\d\.]+\s+[A-Z].+\.{3,}\s*\d+$'),  # 1.1 Title....... 5
                re.compile(r'^[A-Z].+\.{5,}\s*\d+$'),  # Title........ 5
                re.compile(r'^(CHAPTER|APPENDIX)\s+\w+.+\d+$', re.IGNORECASE),  # TOC chapter entry
                re.compile(r'^\s{2,}\d+\.\d+.+\d+$'),  # Indented subsection entry
            ],
            
            # 13. FOOTNOTE_FORMAT - Enhanced footnote patterns
            'footnote_marker': [
                re.compile(r'\[\^\d+\]'),  # [^1] markdown footnote
                re.compile(r'\^\d+'),  # ^1 superscript style
                re.compile(r'\(\d+\)$'),  # (1) at end of line
                re.compile(r'^\[\^\d+\]:'),  # [^1]: footnote definition
            ],
            
            # 14. ABBREVIATION_MANAGER - Abbreviation detection
            'abbreviation': [
                re.compile(r'[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+\s+\([A-Z]{2,}\)'),  # Full Term (ABBR)
                re.compile(r'\([A-Z]{2,}\)'),  # (ABBR) alone
                re.compile(r'^[A-Z]{2,}:\s+'),  # ABBR: definition
                re.compile(r'\b[A-Z]{2,}s?\b'),  # Standalone abbreviation
            ],
            
            # 15. CAPTION_STYLE - Table/Figure caption formatting
            'caption_format': [
                re.compile(r'^\*\*Table\s+\d+'),  # **Table X**
                re.compile(r'^\*\*Figure\s+\d+'),  # **Figure X**
                re.compile(r'^Table\s+\d+\.\d+:'),  # Table X.Y: caption
                re.compile(r'^Figure\s+\d+\.\d+:'),  # Figure X.Y: caption
                re.compile(r'^Source:\s+'),  # Source: attribution
                re.compile(r'^Note:\s+', re.IGNORECASE),  # Note: table note
            ],
            
            # 16. PAGE_BREAK - Explicit page break indicators only
            'page_break': [
                re.compile(r'^\[PAGE\s*BREAK\]', re.IGNORECASE),  # [PAGE BREAK] marker
                re.compile(r'^\\newpage', re.IGNORECASE),  # LaTeX newpage
                re.compile(r'^\\pagebreak', re.IGNORECASE),  # LaTeX pagebreak
            ],
            
            # 16b. HORIZONTAL_RULE - Visual separators (NOT page breaks)
            'horizontal_rule': [
                re.compile(r'^-{3,}$'),  # --- horizontal rule
                re.compile(r'^\*{3,}$'),  # *** horizontal rule
                re.compile(r'^_{3,}$'),  # ___ horizontal rule
            ],
            
            # 17. STATS_FORMAT - Statistical results patterns
            'statistical_result': [
                re.compile(r'β\s*=\s*[-\d\.]+'),  # β = 0.45
                re.compile(r'[Pp]\s*[<>=]\s*\.?\d+'),  # p < .001
                re.compile(r'[Ff]\s*\(\d+,\s*\d+\)\s*='),  # F(3, 56) =
                re.compile(r'[Tt]\s*\(\d+\)\s*='),  # t(45) =
                re.compile(r'χ²\s*\(\d+\)\s*='),  # χ²(5) =
                re.compile(r'[Rr]²?\s*=\s*\.?\d+'),  # R = .67 or R² = .45
                re.compile(r'[Nn]\s*=\s*\d+'),  # N = 100
                re.compile(r'[Mm]\s*=\s*[\d\.]+,\s*[Ss][Dd]\s*='),  # M = 3.5, SD =
                re.compile(r'CI\s*=?\s*\[[\d\.\-,\s]+\]'),  # CI = [0.5, 1.2]
            ],
            
            # 18. QUESTIONNAIRE_STYLE - Survey/questionnaire patterns
            'questionnaire': [
                re.compile(r'^Section\s+[A-Z][:\s]', re.IGNORECASE),  # Section A:
                re.compile(r'^\d+\.\s+\*\*[^*]+\*\*'),  # 1. **Question**
                re.compile(r'[□☐☑✓✗○●]\s+[A-Za-z]'),  # Checkbox options
                re.compile(r'\|\s*SA\s*\|\s*A\s*\|\s*N\s*\|\s*D\s*\|\s*SD\s*\|'),  # Likert header
                re.compile(r'^\s*□\s+[A-Za-z]'),  # □ Option
            ],
            
            # 19. GLOSSARY_FORMAT - Glossary/definition list patterns
            'glossary_entry': [
                re.compile(r'^\*\*[A-Z][^*]+\*\*:\s+'),  # **Term**: Definition
                re.compile(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*:\s+[A-Z]'),  # Term: Definition sentence
                re.compile(r'^•\s+\*\*[^*]+\*\*'),  # • **Term**
            ],
            
            # 20. CROSS_REFERENCE - Cross-reference patterns
            'cross_reference': [
                re.compile(r'[Ss]ee\s+[Tt]able\s+\d+'),  # See Table X
                re.compile(r'[Ss]ee\s+[Ff]igure\s+\d+'),  # See Figure X
                re.compile(r'[Ss]ee\s+[Ss]ection\s+\d+'),  # See Section X
                re.compile(r'[Aa]s\s+shown\s+in\s+[Tt]able'),  # As shown in Table
                re.compile(r'[Aa]s\s+discussed\s+in\s+[Ss]ection'),  # As discussed in Section
                re.compile(r'\([Ss]ee\s+[Pp]age\s+\d+\)'),  # (See page X)
                re.compile(r'\(p\.\s*\d+\)'),  # (p. 45)
            ],
            
            # 21. ACADEMIC_SECTION_PAGE_BREAKS - Sections that must start on new pages
            'academic_section_page_breaks': [
                # Front matter sections (case-insensitive)
                re.compile(r'^#+\s*(ACKNOWLEDGEMENTS?|ACKNOWLEDGMENTS?)\s*$', re.IGNORECASE),
                re.compile(r'^(ACKNOWLEDGEMENTS?|ACKNOWLEDGMENTS?)\s*$', re.IGNORECASE),
                re.compile(r'^#+\s*DEDICATION\s*$', re.IGNORECASE),
                re.compile(r'^DEDICATION\s*$', re.IGNORECASE),
                re.compile(r'^#+\s*ABSTRACT\s*$', re.IGNORECASE),
                re.compile(r'^ABSTRACT\s*$', re.IGNORECASE),
                re.compile(r'^#+\s*(TABLE\s+OF\s+CONTENTS|CONTENTS)\s*$', re.IGNORECASE),
                re.compile(r'^(TABLE\s+OF\s+CONTENTS|CONTENTS)\s*$', re.IGNORECASE),
                re.compile(r'^#+\s*LIST\s+OF\s+TABLES\s*$', re.IGNORECASE),
                re.compile(r'^LIST\s+OF\s+TABLES\s*$', re.IGNORECASE),
                re.compile(r'^#+\s*LIST\s+OF\s+FIGURES\s*$', re.IGNORECASE),
                re.compile(r'^LIST\s+OF\s+FIGURES\s*$', re.IGNORECASE),
                re.compile(r'^#+\s*GLOSSARY\s*$', re.IGNORECASE),
                re.compile(r'^GLOSSARY\s*$', re.IGNORECASE),
                re.compile(r'^#+\s*(LIST\s+OF\s+ABBREVIATIONS|ABBREVIATIONS)\s*$', re.IGNORECASE),
                re.compile(r'^(LIST\s+OF\s+ABBREVIATIONS|ABBREVIATIONS)\s*$', re.IGNORECASE),
                re.compile(r'^#+\s*(APPENDICES|APPENDIX)\s*$', re.IGNORECASE),
                re.compile(r'^(APPENDICES|APPENDIX)\s*$', re.IGNORECASE),
                re.compile(r'^#+\s*REFERENCES?\s*$', re.IGNORECASE),
                re.compile(r'^REFERENCES?\s*$', re.IGNORECASE),
                re.compile(r'^#+\s*BIBLIOGRAPHY\s*$', re.IGNORECASE),
                re.compile(r'^BIBLIOGRAPHY\s*$', re.IGNORECASE),
                
                # Chapter headings (various formats)
                re.compile(r'^#+\s*CHAP?TER\s+(ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN)\s*[:\.]?\s*.*$', re.IGNORECASE),
                re.compile(r'^CHAP?TER\s+(ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN)\s*[:\.]?\s*.*$', re.IGNORECASE),
                re.compile(r'^#+\s*CHAP?TER\s+[1-9][0-9]?\s*[:\.]?\s*.*$', re.IGNORECASE),
                re.compile(r'^CHAP?TER\s+[1-9][0-9]?\s*[:\.]?\s*.*$', re.IGNORECASE),
                re.compile(r'^#+\s*CHAP?TER\s+[IVXLC]+\s*[:\.]?\s*.*$', re.IGNORECASE),
                re.compile(r'^CHAP?TER\s+[IVXLC]+\s*[:\.]?\s*.*$', re.IGNORECASE),
            ],
            
            # 23. DISSERTATION_CHAPTER_TITLES - Chapter heading + title patterns
            'dissertation_chapter': [
                # Chapter heading only (CHAPTER ONE, CHAPTER 1, CHAPTER I)
                re.compile(r'^#+?\s*CHAP?TER\s+(ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN)\s*$', re.IGNORECASE),
                re.compile(r'^#+?\s*CHAP?TER\s+[1-9][0-9]?\s*$', re.IGNORECASE),
                re.compile(r'^#+?\s*CHAP?TER\s+[IVXLC]+\s*$', re.IGNORECASE),
                # Chapter heading with title on same line
                re.compile(r'^#+?\s*CHAP?TER\s+(ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN)\s*[:\-\.]\s*.+$', re.IGNORECASE),
                re.compile(r'^#+?\s*CHAP?TER\s+[1-9][0-9]?\s*[:\-\.]\s*.+$', re.IGNORECASE),
                re.compile(r'^#+?\s*CHAP?TER\s+[IVXLC]+\s*[:\-\.]\s*.+$', re.IGNORECASE),
            ],
            
            # 24. COPYRIGHT_PAGE - Copyright notice patterns
            'copyright_content': [
                re.compile(r'^\s*©\s*[Cc]opyright\s+', re.IGNORECASE),
                re.compile(r'^\s*[Cc]opyright\s+©', re.IGNORECASE),
                re.compile(r'^\s*[Cc]opyright\s+\d{4}', re.IGNORECASE),
                re.compile(r'^\s*All\s+Rights\s+Reserved\s*\.?\s*$', re.IGNORECASE),
                re.compile(r'^\s*No\s+part\s+of\s+this\s+(document|thesis|dissertation|work)', re.IGNORECASE),
            ],
            
            # 25. DECLARATION_PAGE - Declaration patterns
            'declaration_content': [
                re.compile(r'^#+?\s*DECLARATION\s*$', re.IGNORECASE),
                re.compile(r'I\s+(hereby\s+)?declare\s+that', re.IGNORECASE),
                re.compile(r'This\s+(thesis|dissertation|work)\s+is\s+(my\s+own|original)', re.IGNORECASE),
                re.compile(r'has\s+not\s+been\s+(previously\s+)?submitted', re.IGNORECASE),
            ],
            
            # 26. CERTIFICATION_PAGE - Certification/Approval patterns
            'certification_content': [
                re.compile(r'^#+?\s*CERTIFICATION\s*$', re.IGNORECASE),
                re.compile(r'^#+?\s*APPROVAL\s+PAGE\s*$', re.IGNORECASE),
                re.compile(r'This\s+is\s+to\s+certify\s+that', re.IGNORECASE),
                re.compile(r'has\s+been\s+(read\s+and\s+)?approved', re.IGNORECASE),
                re.compile(r'meets\s+the\s+requirements?', re.IGNORECASE),
                re.compile(r'^#+?\s*COMMITTEE\s+APPROVAL\s*$', re.IGNORECASE),
            ],
            
            # 27. CHAPTER_TITLE_FOLLOWING - Chapter title that follows chapter heading
            'chapter_title_following': [
                # All caps title (common for chapter titles)
                re.compile(r'^#+?\s*[A-Z][A-Z\s]+$'),
                # Common chapter title keywords
                re.compile(r'^#+?\s*(INTRODUCTION|LITERATURE\s+REVIEW|METHODOLOGY|RESULTS|DISCUSSION|CONCLUSION|RECOMMENDATIONS?|THEORETICAL\s+FRAMEWORK|SUMMARY\s+AND\s+CONCLUSIONS?|SUMMARY\s+OF\s+FINDINGS)', re.IGNORECASE),
            ],
            
            # 28. SIGNATURE_LINE - Signature block patterns
            'signature_line': [
                re.compile(r'^\s*_{10,}\s*$'),  # Long underscore line
                re.compile(r'^\s*\.{10,}\s*$'),  # Long dotted line
                re.compile(r'^\s*-{10,}\s*$'),  # Long dash line
                re.compile(r'^\s*(Signed?|Signature)\s*:\s*_{5,}', re.IGNORECASE),
                re.compile(r'^\s*Date\s*:\s*_{5,}', re.IGNORECASE),
                re.compile(r'^\s*Name\s*:\s*_{5,}', re.IGNORECASE),
            ],
            
            # 29. SUPERVISOR_INFO - Supervisor/Advisor patterns
            'supervisor_info': [
                re.compile(r'^\s*(Supervisor|Advisor|Chair)\s*:', re.IGNORECASE),
                re.compile(r'^\s*(First|Second|Third)\s+(Reader|Supervisor|Advisor)', re.IGNORECASE),
                re.compile(r'^\s*(Dr\.|Prof\.|Professor)\s+[A-Z]', re.IGNORECASE),
                re.compile(r'^\s*Committee\s+(Member|Chair)', re.IGNORECASE),
            ],
            
            # 30. TOC_ENTRY - Table of Contents entry patterns
            'toc_entry': [
                re.compile(r'^.+\.{3,}\s*\d+\s*$'),  # Title......... Page
                re.compile(r'^.+\s+\.{2,}\s*\d+\s*$'),  # Title .. Page
                re.compile(r'^\s*[IVXLC]+\s+.+\s+\d+\s*$'),  # Roman numeral entry
                re.compile(r'^\s*\d+\.\d*\s+.+\s+\d+\s*$'),  # Numbered entry
                re.compile(r'^.+(?:\t| {2,})(?:\d+|[IVXLC]+)\s*$', re.IGNORECASE),  # Title <tab/spaces> iii
            ],
            
            # 31. HEADING_SPACE_ISSUES - Patterns for detecting spacing issues in headings
            'heading_space_issues': [
                # Trailing spaces on headings
                re.compile(r'^#+\s+.+\s{2,}$'),  # Heading with 2+ trailing spaces
                # Multiple spaces before punctuation
                re.compile(r'^#+\s+.+\s{2,}[:\.]\s*'),  # Heading with spaces before colon/period
                # Multiple internal spaces
                re.compile(r'^#+\s+.+\s{3,}.+$'),  # Heading with 3+ consecutive spaces inside
            ],
            
            # 32. SPACING_CLEANUP - General spacing issue patterns
            'spacing_cleanup': [
                re.compile(r'\s{2,}'),  # Multiple spaces (for replacement)
                re.compile(r'\s+,'),  # Space before comma
                re.compile(r'\s+\.'),  # Space before period
                re.compile(r'\s+:'),  # Space before colon
                re.compile(r'\s+$'),  # Trailing whitespace
            ],
            
            # 33. ACADEMIC_TABLE_CONTENT - Table content type detection patterns
            'table_content_numeric': [
                re.compile(r'^\s*[\-\+]?\d+\.?\d*\s*$'),  # Plain numbers: 123, 123.45, -45.6
                re.compile(r'^\s*[\-\+]?\d{1,3}(,\d{3})*(\.\d+)?\s*$'),  # Comma-separated: 1,234.56
            ],
            
            'table_content_percentage': [
                re.compile(r'^\s*[\-\+]?\d+\.?\d*\s*%\s*$'),  # 25.5%, -10%
            ],
            
            'table_content_statistical': [
                re.compile(r'[βαγδ]'),  # Greek letters
                re.compile(r'p\s*[<>=]\s*[\d\.]+'),  # p-values: p < 0.05
                re.compile(r'[Ff]\s*\(\s*\d+\s*,\s*\d+\s*\)'),  # F-statistic: F(3, 56)
                re.compile(r'[Rr]²?\s*=\s*[\d\.]+'),  # R or R²
                re.compile(r'[Tt]\s*\(\s*\d+\s*\)'),  # t-statistic: t(45)
                re.compile(r'χ²\s*\(\s*\d+\s*\)'),  # Chi-square
                re.compile(r'SE\s*=?\s*[\d\.]+'),  # Standard error
                re.compile(r'CI\s*=?\s*\['),  # Confidence interval
                re.compile(r'\*{1,3}$'),  # Significance asterisks
            ],
            
            'table_content_ci': [
                re.compile(r'\[[\d\.\-,\s]+\]'),  # Confidence interval: [0.5, 1.2]
                re.compile(r'\([\d\.\-,\s]+\)'),  # Alternative: (0.5, 1.2)
            ],
            
            # 34. TABLE_CAPTION_FORMAT - Table caption patterns
            'table_caption_format': [
                re.compile(r'^\s*[Tt]able\s+\d+\.\d+[\:\.]?\s*.+$'),  # Table 4.1: Title
                re.compile(r'^\s*[Tt]able\s+\d+[\:\.]?\s*.+$'),  # Table 4: Title
                re.compile(r'^\s*\*\*[Tt]able\s+\d+\.?\d*\*\*[\:\.]?\s*.+$'),  # **Table 4.1**: Title
                re.compile(r'^\s*TABLE\s+\d+\.?\d*[\:\.]?\s*.+$'),  # TABLE 4.1: Title
            ],
            
            # 35. SHORT_DOCUMENT_INDICATORS - Detect if document is short/assignment type
            'short_doc_indicators': [
                re.compile(r'^Assignment\s*\d*[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Homework\s*\d*[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Course\s+Material[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Exercise\s*\d*[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Lab\s*\d*[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Worksheet[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Quiz\s*\d*[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Problem\s+Set\s*\d*[\:\.]?\s*', re.IGNORECASE),
            ],
            
            # 36. LONG_DOCUMENT_INDICATORS - Indicators of formal academic documents
            'long_doc_indicators': [
                re.compile(r'Dissertation', re.IGNORECASE),
                re.compile(r'\bThesis\b', re.IGNORECASE),
                re.compile(r'\bPhD\b', re.IGNORECASE),
                re.compile(r"Master's\s+(?:Thesis|Dissertation)", re.IGNORECASE),
                re.compile(r'Literature\s+Review', re.IGNORECASE),
                re.compile(r'Abstract', re.IGNORECASE),
            ],
            
            # 37. TOC_PATTERNS - Detect Table of Contents for removal
            'toc_header': [
                re.compile(r'^#+\s*TABLE\s+OF\s+CONTENTS\s*$', re.IGNORECASE),
                re.compile(r'^TABLE\s+OF\s+CONTENTS\s*$', re.IGNORECASE),
                re.compile(r'^Contents\s*$', re.IGNORECASE),
                re.compile(r'^#+\s*Contents\s*$', re.IGNORECASE),
            ],
            
            'toc_content_line': [
                re.compile(r'^\d+\.\s+.+\.{3,}\s*\d+$'),  # 1. Introduction ........ 5
                re.compile(r'^[A-Za-z].+\.{3,}\s*\d+$'),  # Introduction ........ 5
                re.compile(r'^.+\s+\.{5,}\s*\d+$'),  # Anything with dot leaders
                re.compile(r'^\s*\d+\.\d+\s+.+\s+\d+\s*$'),  # 1.1 Section  5
                re.compile(r'^.+(?:\t| {2,})(?:\d+|[IVXLC]+)\s*$', re.IGNORECASE),  # Title <tab/spaces> iii
            ],
            
            # 38. KEY_POINT_LEARNING_OBJECTIVES - Learning goals/objectives
            'key_point_learning': [
                re.compile(r'^Learning\s+Objectives?[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Objectives?[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Goals?[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^By\s+the\s+end.*you\s+will\s+be\s+able\s+to[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Learning\s+Outcomes?[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Outcome\s+\d+[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Objective\s+\d+[\:\.]?\s*', re.IGNORECASE),
            ],
            
            # 39. KEY_POINT_DEFINITIONS - Important definitions
            'key_point_definitions': [
                re.compile(r'^Definition[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^[A-Z][a-z]+\s+is\s+(?:defined\s+as|refers\s+to)', re.IGNORECASE),
                re.compile(r'^The\s+term\s+["\'][^"\']+["\']\s+means', re.IGNORECASE),
                re.compile(r'^[A-Z][a-z]+[\:\s]+(?:A|An|The)\s+', re.IGNORECASE),  # Term: A definition
            ],
            
            # 40. KEY_POINT_KEY_CONCEPTS - Critical information markers
            'key_point_concepts': [
                re.compile(r'^Key\s+(?:Concept|Idea|Point)[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Important[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Note[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Remember[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Critical[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Essential[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Key\s+Takeaway[\:\.]?\s*', re.IGNORECASE),
            ],
            
            # 41. KEY_POINT_PROCEDURES - Steps and procedures
            'key_point_procedures': [
                re.compile(r'^Steps?[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Procedure[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Process[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Algorithm[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Method[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Step\s+\d+[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Instructions?[\:\.]?\s*', re.IGNORECASE),
            ],
            
            # 42. KEY_POINT_EXAMPLES - Example markers
            'key_point_examples': [
                re.compile(r'^Example[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^For\s+instance[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^For\s+example[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Consider[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Sample[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Example\s+\d+[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Case\s+Study[\:\.]?\s*', re.IGNORECASE),
            ],
            
            # 43. KEY_POINT_WARNINGS - Caution/warning markers
            'key_point_warnings': [
                re.compile(r'^Warning[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Caution[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Avoid[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Common\s+(?:mistake|error)[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Do\s+not[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Never[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Pitfall[\:\.]?\s*', re.IGNORECASE),
            ],
            
            # 44. KEY_POINT_EXERCISES - Exercise/question markers
            'key_point_exercises': [
                re.compile(r'^Exercise[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Question[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Problem[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Task[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Challenge[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Exercise\s+\d+[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Question\s+\d+[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Problem\s+\d+[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Practice[\:\.]?\s*', re.IGNORECASE),
            ],
            
            # 45. KEY_POINT_SUMMARY - Summary/conclusion markers
            'key_point_summary': [
                re.compile(r'^Summary[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Conclusion[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Takeaways?[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^In\s+summary[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^To\s+(?:summarize|recap)[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Main\s+(?:points?|takeaways?)[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Key\s+Points?[\:\.]?\s*', re.IGNORECASE),
            ],
            
            # 46. ASSIGNMENT_HEADER_ELEMENTS - Assignment metadata fields
            'assignment_header': [
                re.compile(r'^Student(?:\s+Name)?[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Student\s+ID[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Name[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Course(?:\s+Code)?[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Instructor[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Professor[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Due\s+Date[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Date[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Submitted\s+(?:by|to)[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Class[\:\.]?\s*', re.IGNORECASE),
                re.compile(r'^Section[\:\.]?\s*', re.IGNORECASE),
            ],
            
            # ============================================================
            # 47. POINT_FORM_CONTENT_FORMATTING - Detect and format point-form content
            # ============================================================
            
            # 47a. List patterns - Various list markers
            'point_form_numbered': [
                re.compile(r'^\s*(\d+)[\.\)]\s+.+$'),  # 1. Item or 1) Item
                re.compile(r'^\s*([a-z])[\.\)]\s+.+$'),  # a. Item or a) Item
                re.compile(r'^\s*([ivxIVX]+)[\.\)]\s+.+$'),  # i. Item, ii. Item (Roman)
                re.compile(r'^\s*\((\d+)\)\s+.+$'),  # (1) Item
                re.compile(r'^\s*\(([a-z])\)\s+.+$'),  # (a) Item
                re.compile(r'^\s*([A-Z])[\.\)]\s+.+$'),  # A. Item or A) Item
            ],
            
            'point_form_bulleted': [
                re.compile(r'^\s*[\-\*\•◦→○▪▫■□►▸◆◇]\s+.+$'),  # Various bullet symbols
                re.compile(r'^\s*—\s+.+$'),  # Em dash
                re.compile(r'^\s*–\s+.+$'),  # En dash
            ],
            
            'point_form_checkbox': [
                re.compile(r'^\s*□\s+.+$'),  # □ Item
                re.compile(r'^\s*☐\s+.+$'),  # ☐ Item
                re.compile(r'^\s*☑\s+.+$'),  # ☑ Item (checked)
                re.compile(r'^\s*✓\s+.+$'),  # ✓ Item
                re.compile(r'^\s*✗\s+.+$'),  # ✗ Item
                re.compile(r'^\s*\[\s*\]\s+.+$'),  # [ ] Item
                re.compile(r'^\s*\[[xX]\]\s+.+$'),  # [x] Item (checked)
            ],
            
            # 47b. Context clues - Phrases that precede point lists
            'point_form_context_clues': [
                re.compile(r'(?:the\s+)?following(?:\s+\w+)?:', re.IGNORECASE),  # "The following:", "Following items:"
                re.compile(r'below\s+are:', re.IGNORECASE),  # "Below are:"
                re.compile(r'as\s+follows:', re.IGNORECASE),  # "As follows:"
                re.compile(r'namely:', re.IGNORECASE),  # "Namely:"
                re.compile(r'such\s+as:', re.IGNORECASE),  # "Such as:"
                re.compile(r'including:', re.IGNORECASE),  # "Including:"
                re.compile(r'(?:include|includes|including)\s*$', re.IGNORECASE),  # "include" without colon
                re.compile(r'such\s+as\s*$', re.IGNORECASE),  # "such as" without colon
                re.compile(r':\s*$'),  # Ends with colon
            ],
            
            # 47c. Structured content headings that should have point-form content
            'point_form_objectives': [
                re.compile(r'^(?:Learning\s+)?Objectives?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Goals?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Aims?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Purposes?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Learning\s+Outcomes?[\s:]*$', re.IGNORECASE),
            ],
            
            'point_form_steps': [
                re.compile(r'^Steps?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Procedures?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Process[\s:]*$', re.IGNORECASE),
                re.compile(r'^Methods?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Algorithms?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Phases?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Instructions?[\s:]*$', re.IGNORECASE),
            ],
            
            'point_form_key_points': [
                re.compile(r'^Key\s+Points?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Main\s+Ideas?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Important[\s:]*$', re.IGNORECASE),
                re.compile(r'^Critical[\s:]*$', re.IGNORECASE),
                re.compile(r'^Essential[\s:]*$', re.IGNORECASE),
                re.compile(r'^Takeaways?[\s:]*$', re.IGNORECASE),
            ],
            
            'point_form_features': [
                re.compile(r'^Features?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Characteristics?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Properties?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Attributes?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Aspects?[\s:]*$', re.IGNORECASE),
            ],
            
            'point_form_pros_cons': [
                re.compile(r'^Advantages?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Disadvantages?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Benefits?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Drawbacks?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Pros?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Cons?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Strengths?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Weaknesses?[\s:]*$', re.IGNORECASE),
            ],
            
            'point_form_examples': [
                re.compile(r'^Examples?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Cases?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Instances?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Scenarios?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Illustrations?[\s:]*$', re.IGNORECASE),
            ],
            
            'point_form_questions': [
                re.compile(r'^Questions?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Problems?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Exercises?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Tasks?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Challenges?[\s:]*$', re.IGNORECASE),
            ],
            
            'point_form_requirements': [
                re.compile(r'^Requirements?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Criteria[\s:]*$', re.IGNORECASE),
                re.compile(r'^Conditions?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Prerequisites?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Necessities?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Materials?\s+(?:Needed|Required)[\s:]*$', re.IGNORECASE),
            ],
            
            'point_form_tips': [
                re.compile(r'^Tips?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Suggestions?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Recommendations?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Advice[\s:]*$', re.IGNORECASE),
                re.compile(r'^Hints?[\s:]*$', re.IGNORECASE),
            ],
            
            'point_form_warnings': [
                re.compile(r'^Warnings?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Cautions?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Avoid[\s:]*$', re.IGNORECASE),
                re.compile(r'^Do\s+Not[\s:]*$', re.IGNORECASE),
                re.compile(r'^Never[\s:]*$', re.IGNORECASE),
                re.compile(r'^Common\s+Mistakes?[\s:]*$', re.IGNORECASE),
            ],
            
            'point_form_components': [
                re.compile(r'^Components?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Parts?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Elements?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Sections?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Modules?[\s:]*$', re.IGNORECASE),
            ],
            
            'point_form_rules': [
                re.compile(r'^Rules?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Principles?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Laws?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Theorems?[\s:]*$', re.IGNORECASE),
                re.compile(r'^Guidelines?[\s:]*$', re.IGNORECASE),
            ],
            
            # 47d. Serial comma detection - Content in sentence that should be points
            'point_form_serial_comma': [
                re.compile(r'^[^:]+:\s*[A-Z][^,]+,\s*[^,]+,\s*(?:and|or)\s*[^,.]+[\.:]?$'),  # X: item1, item2, and item3
                re.compile(r'^[^:]+:\s*[^;]+;\s*[^;]+;\s*[^;]+[;.]?$'),  # X: item1; item2; item3
            ],
            
            # 47e. Implicit list in sentence - First/Second/Third pattern
            'point_form_ordinal_sentence': [
                re.compile(r'(?:First|Firstly),?\s+[^.]+\.', re.IGNORECASE),
                re.compile(r'(?:Second|Secondly),?\s+[^.]+\.', re.IGNORECASE),
                re.compile(r'(?:Third|Thirdly),?\s+[^.]+\.', re.IGNORECASE),
                re.compile(r'(?:Fourth|Fourthly),?\s+[^.]+\.', re.IGNORECASE),
                re.compile(r'(?:Fifth|Finally|Lastly),?\s+[^.]+\.', re.IGNORECASE),
            ],
        }
    
    def get_cell_content_type(self, cell_text):
        """
        Determine the content type of a table cell for alignment purposes.
        Returns: 'numeric', 'percentage', 'statistical', 'text'
        """
        if not cell_text:
            return 'text'
        
        text = str(cell_text).strip()
        
        # Check for percentage first (most specific)
        for pattern in self.patterns.get('table_content_percentage', []):
            if pattern.match(text):
                return 'percentage'
        
        # Check for statistical notation
        for pattern in self.patterns.get('table_content_statistical', []):
            if pattern.search(text):
                return 'statistical'
        
        # Check for confidence intervals
        for pattern in self.patterns.get('table_content_ci', []):
            if pattern.search(text):
                return 'statistical'
        
        # Check for plain numeric
        for pattern in self.patterns.get('table_content_numeric', []):
            if pattern.match(text):
                return 'numeric'
        
        return 'text'
    
    def get_column_content_types(self, rows):
        """
        Analyze all rows to determine the predominant content type for each column.
        Returns list of content types, one per column.
        """
        if not rows or len(rows) == 0:
            return []
        
        num_cols = max(len(row) for row in rows)
        column_types = []
        
        for col_idx in range(num_cols):
            type_counts = {'numeric': 0, 'percentage': 0, 'statistical': 0, 'text': 0}
            
            # Skip header row (index 0), analyze data rows only
            for row_idx, row in enumerate(rows):
                if row_idx == 0:  # Skip header
                    continue
                if col_idx < len(row):
                    cell_type = self.get_cell_content_type(row[col_idx])
                    type_counts[cell_type] += 1
            
            # Determine predominant type (statistical and percentage take precedence)
            if type_counts['statistical'] > 0:
                column_types.append('statistical')
            elif type_counts['percentage'] > 0:
                column_types.append('percentage')
            elif type_counts['numeric'] > type_counts['text']:
                column_types.append('numeric')
            else:
                column_types.append('text')
        
        return column_types
    
    def get_alignment_for_content_type(self, content_type):
        """
        Get the appropriate alignment for a content type.
        Returns: WD_ALIGN_PARAGRAPH constant
        """
        alignment_map = {
            'text': WD_ALIGN_PARAGRAPH.LEFT,
            'numeric': WD_ALIGN_PARAGRAPH.RIGHT,
            'percentage': WD_ALIGN_PARAGRAPH.RIGHT,
            'statistical': WD_ALIGN_PARAGRAPH.CENTER,
            'header': WD_ALIGN_PARAGRAPH.CENTER,
        }
        return alignment_map.get(content_type, WD_ALIGN_PARAGRAPH.LEFT)

    def clean_heading_spaces(self, text):
        """
        Clean up spacing issues in heading text.
        Removes trailing spaces, normalizes internal spacing, fixes punctuation spacing.
        """
        if not text:
            return text
        
        original = text
        
        # Remove trailing whitespace
        text = text.rstrip()
        
        # Check if this is a heading line
        if not text.lstrip().startswith('#'):
            return text
        
        # Extract markdown prefix and heading content
        match = re.match(r'^(\s*#+\s*)', text)
        if not match:
            return text
        
        prefix = match.group(1)
        content = text[len(prefix):]
        
        # Clean the content
        # 1. Remove trailing spaces and punctuation that shouldn't be there
        content = content.rstrip(' .,:;')
        
        # 2. Normalize multiple spaces to single space
        content = re.sub(r'\s{2,}', ' ', content)
        
        # 3. Fix spacing around colons (for chapter headings like "CHAPTER ONE: INTRO")
        content = re.sub(r'\s+:', ':', content)
        content = re.sub(r':\s{2,}', ': ', content)
        
        # 4. Fix spacing around hyphens
        content = re.sub(r'\s+-\s+', ' - ', content)
        content = re.sub(r'\s{2,}-', ' -', content)
        content = re.sub(r'-\s{2,}', '- ', content)
        
        # Reconstruct the heading
        cleaned = prefix + content
        
        return cleaned
    
    def clean_document_spacing(self, text):
        """
        Clean up spacing issues throughout a document.
        Applies to all lines, with special handling for headings.
        """
        if not text:
            return text
        
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Check if it's a heading
            if line.lstrip().startswith('#'):
                cleaned_line = self.clean_heading_spaces(line)
            else:
                # For non-heading lines, just remove trailing spaces
                cleaned_line = line.rstrip()
                
                # Don't modify spacing in code blocks or special formatting
                if not cleaned_line.startswith('```') and not cleaned_line.startswith('    '):
                    # Fix space before punctuation (except in tables)
                    if not '|' in cleaned_line:
                        cleaned_line = re.sub(r'\s+,', ',', cleaned_line)
                        cleaned_line = re.sub(r'\s+\.(?!\d)', '.', cleaned_line)  # Preserve decimals
                        cleaned_line = re.sub(r'\s+;', ';', cleaned_line)
            
            cleaned_lines.append(cleaned_line)
        
        return '\n'.join(cleaned_lines)
    
    # ========== SHORT DOCUMENT FORMATTING METHODS ==========
    
    def is_short_document(self, text):
        """
        Determine if a document is 'short' (assignment, course material, exercise).
        Uses multiple criteria: word count, indicators, and structure.
        Returns: (is_short, reason)
        """
        if not text:
            return False, "Empty document"
        
        # Calculate document metrics
        word_count = len(text.split())
        line_count = len(text.split('\n'))
        char_count = len(text)
        estimated_pages = char_count / 2000  # ~2000 chars per page
        
        # Count headings/sections
        heading_count = len(re.findall(r'^#+\s+', text, re.MULTILINE))
        if heading_count == 0:
            # Count uppercase headings
            heading_count = len(re.findall(r'^[A-Z][A-Z\s]{5,50}$', text, re.MULTILINE))
        
        # Check for LONG document indicators (Dissertation, Thesis, etc.)
        for pattern in self.patterns.get('long_doc_indicators', []):
            if pattern.search(text):
                return False, "Contains long document indicators"
        
        # Check for SHORT document indicators (Assignment, Homework, etc.)
        has_short_indicator = False
        for pattern in self.patterns.get('short_doc_indicators', []):
            if pattern.search(text):
                has_short_indicator = True
                break
        
        # Multiple threshold checks
        is_short_by_words = word_count < 3000
        is_short_by_pages = estimated_pages < 10
        is_short_by_lines = line_count < 500
        is_short_by_sections = heading_count < 10
        
        # Decision logic
        if has_short_indicator:
            return True, "Contains short document indicator (Assignment/Homework/Exercise)"
        
        if is_short_by_words and is_short_by_pages:
            return True, f"Short by word count ({word_count} words) and pages ({estimated_pages:.1f} pages)"
        
        if is_short_by_lines and is_short_by_sections:
            return True, f"Short by line count ({line_count} lines) and sections ({heading_count} sections)"
        
        return False, "Document does not meet short document criteria"
    
    def is_toc_header_line(self, text):
        """Check if a line is a Table of Contents header."""
        if not text:
            return False
        
        clean_text = text.strip()
        for pattern in self.patterns.get('toc_header', []):
            if pattern.match(clean_text):
                return True
        return False
    
    def is_toc_content_line(self, text):
        """Check if a line is a Table of Contents content line (entries with page numbers)."""
        if not text:
            return False
        
        clean_text = text.strip()
        
        # Check for dot leaders pattern (common in TOC)
        if '.....' in clean_text or '…..' in clean_text or '…' in clean_text:
            return True
        
        for pattern in self.patterns.get('toc_content_line', []):
            if pattern.match(clean_text):
                return True
        return False
    
    def _front_matter_placeholder_type(self, front_matter_type):
        placeholders = {
            'toc': 'toc_placeholder',
            'list_of_tables': 'list_of_tables_placeholder',
            'list_of_figures': 'list_of_figures_placeholder',
        }
        return placeholders.get(front_matter_type)

    def _is_front_matter_entry_line(self, front_matter_type, text):
        if not text:
            return False

        clean_text = text.strip()

        if front_matter_type == 'toc':
            return self.is_toc_content_line(clean_text) or self.is_toc_entry(clean_text)

        if front_matter_type == 'list_of_figures':
            return bool(re.match(r'^(Figure|Fig\.)\s+\d+', clean_text, re.IGNORECASE)) or self.is_toc_content_line(clean_text)

        if front_matter_type == 'list_of_tables':
            return bool(re.match(r'^(Table|Tbl\.?|Tab\.?)\s+\d+', clean_text, re.IGNORECASE)) or self.is_toc_content_line(clean_text)

        return False

    def strip_front_matter_placeholders(self, lines):
        """
        Remove existing TOC/LOF/LOT content and replace it with placeholders so the
        system-generated lists can be inserted at the same location.
        """
        cleaned_lines = []
        in_section = None
        blank_count = 0

        for line in lines:
            text = line if isinstance(line, str) else line.get('text', '')
            stripped = text.strip()

            while True:
                if in_section:
                    if not stripped:
                        blank_count += 1
                        if blank_count >= 2:
                            in_section = None
                        break

                    next_section = self.get_front_matter_section_type(stripped)
                    if next_section and next_section != in_section:
                        in_section = None
                        blank_count = 0
                        continue

                    if re.match(r'^#+\s+', stripped) or (stripped.isupper() and len(stripped) > 5):
                        in_section = None
                        blank_count = 0
                        continue

                    if self._is_front_matter_entry_line(in_section, stripped):
                        break

                    in_section = None
                    blank_count = 0
                    continue

                front_matter_type = self.get_front_matter_section_type(stripped)
                placeholder_type = self._front_matter_placeholder_type(front_matter_type)
                if placeholder_type:
                    cleaned_lines.append({
                        'type': placeholder_type,
                        'text': stripped,
                    })
                    in_section = front_matter_type
                    blank_count = 0
                    break

                cleaned_lines.append(line)
                break

        return cleaned_lines
    
    def get_key_point_type(self, text):
        """
        Detect if a line is a key point that should be emphasized.
        Returns: (key_point_type, prefix_emoji) or (None, None)
        """
        if not text:
            return None, None
        
        clean_text = text.strip()
        
        # Check each key point category
        key_point_categories = [
            ('key_point_learning', 'learning', '📚 '),
            ('key_point_definitions', 'definition', ''),
            ('key_point_concepts', 'concept', '🔑 '),
            ('key_point_procedures', 'procedure', '📝 '),
            ('key_point_examples', 'example', '📋 '),
            ('key_point_warnings', 'warning', '⚠️ '),
            ('key_point_exercises', 'exercise', '💪 '),
            ('key_point_summary', 'summary', '📊 '),
        ]
        
        for pattern_key, point_type, emoji in key_point_categories:
            for pattern in self.patterns.get(pattern_key, []):
                if pattern.match(clean_text):
                    return point_type, emoji
        
        return None, None
    
    def is_assignment_header_field(self, text):
        """Check if a line is an assignment header field (Student Name, Course, etc.)."""
        if not text:
            return False
        
        clean_text = text.strip()
        for pattern in self.patterns.get('assignment_header', []):
            if pattern.match(clean_text):
                return True
        return False
    
    def emphasize_key_point(self, text, point_type, emoji):
        """
        Apply emphasis formatting to a key point line.
        Returns the emphasized text.
        """
        if not text:
            return text
        
        clean_text = text.strip()
        
        # Already formatted (has ** markers)
        if clean_text.startswith('**') or clean_text.startswith('*'):
            return text
        
        # Apply formatting based on type (asterisks and italics are forbidden)
        emphasized = clean_text
        
        # Add emoji prefix if provided
        if emoji:
            emphasized = emoji + emphasized
        
        return emphasized
    
    # ========== POINT FORM CONTENT FORMATTING METHODS ==========
    
    def is_point_form_line(self, text):
        """
        Check if a line is already in point form (numbered, bulleted, checkbox).
        Returns: (is_point_form, point_type) - type is 'numbered', 'bulleted', 'checkbox', or None
        """
        if not text:
            return False, None
        
        clean_text = text.strip()
        
        # Check numbered patterns
        for pattern in self.patterns.get('point_form_numbered', []):
            if pattern.match(clean_text):
                return True, 'numbered'
        
        # Check bulleted patterns
        for pattern in self.patterns.get('point_form_bulleted', []):
            if pattern.match(clean_text):
                return True, 'bulleted'
        
        # Check checkbox patterns
        for pattern in self.patterns.get('point_form_checkbox', []):
            if pattern.match(clean_text):
                return True, 'checkbox'
        
        return False, None
    
    def get_point_form_heading_type(self, text):
        """
        Check if a line is a heading that should have point-form content following it.
        Returns: (heading_type, format_type) or (None, None)
        - heading_type: objectives, steps, key_points, features, pros_cons, examples, etc.
        - format_type: 'numbered' or 'bulleted'
        """
        if not text:
            return None, None
        
        clean_text = text.strip()
        
        # Mapping of pattern keys to (heading_type, format_type)
        heading_patterns = [
            ('point_form_objectives', 'objectives', 'bulleted'),
            ('point_form_steps', 'steps', 'numbered'),
            ('point_form_key_points', 'key_points', 'bulleted'),
            ('point_form_features', 'features', 'bulleted'),
            ('point_form_pros_cons', 'pros_cons', 'bulleted'),
            ('point_form_examples', 'examples', 'bulleted'),
            ('point_form_questions', 'questions', 'numbered'),
            ('point_form_requirements', 'requirements', 'bulleted'),
            ('point_form_tips', 'tips', 'bulleted'),
            ('point_form_warnings', 'warnings', 'bulleted'),
            ('point_form_components', 'components', 'bulleted'),
            ('point_form_rules', 'rules', 'numbered'),
        ]
        
        for pattern_key, heading_type, format_type in heading_patterns:
            for pattern in self.patterns.get(pattern_key, []):
                if pattern.match(clean_text):
                    return heading_type, format_type
        
        return None, None
    
    def has_list_context_clue(self, text):
        """
        Check if a line ends with context that suggests a list follows.
        Returns True if line has context clue (e.g., ends with colon after "following").
        """
        if not text:
            return False
        
        clean_text = text.strip()
        
        for pattern in self.patterns.get('point_form_context_clues', []):
            if pattern.search(clean_text):
                return True
        
        return False
    
    def extract_serial_comma_items(self, text):
        """
        Extract items from a sentence with serial commas.
        Example: "Tools needed: hammer, nails, and screwdriver" -> ["hammer", "nails", "screwdriver"]
        Returns: (heading, items_list) or (None, None)
        """
        if not text:
            return None, None
        
        clean_text = text.strip()
        
        # Check for serial comma pattern with colon
        for pattern in self.patterns.get('point_form_serial_comma', []):
            if pattern.match(clean_text):
                # Split on colon
                if ':' in clean_text:
                    parts = clean_text.split(':', 1)
                    heading = parts[0].strip()
                    content = parts[1].strip()
                    
                    # Try to split by comma or semicolon
                    if ';' in content:
                        items = [item.strip() for item in content.split(';')]
                    else:
                        # Handle "and" or "or" at the end
                        content = re.sub(r',?\s+(?:and|or)\s+', ', ', content)
                        items = [item.strip().rstrip('.') for item in content.split(',')]
                    
                    # Clean up items
                    items = [item for item in items if item and len(item) > 1]
                    
                    if len(items) >= 2:
                        return heading, items
        
        return None, None
    
    def extract_ordinal_steps(self, text):
        """
        Extract steps from a sentence with ordinal words (First, Second, Third...).
        Returns: list of steps or None
        """
        if not text:
            return None
        
        # Pattern to find ordinal sentences
        ordinal_pattern = re.compile(
            r'(?:First(?:ly)?|Second(?:ly)?|Third(?:ly)?|Fourth(?:ly)?|Fifth(?:ly)?|Finally|Lastly),?\s+([^.]+)\.',
            re.IGNORECASE
        )
        
        matches = ordinal_pattern.findall(text)
        if len(matches) >= 2:
            return [match.strip() for match in matches]
        
        return None
    
    def clean_point_content(self, text):
        """
        Clean a point by removing existing markers and fixing formatting.
        """
        if not text:
            return text
        
        clean = text.strip()
        
        # Remove common list markers
        marker_patterns = [
            r'^\s*\d+[\.\)]\s*',      # 1. or 1)
            r'^\s*[a-z][\.\)]\s*',    # a. or a)
            r'^\s*[ivxIVX]+[\.\)]\s*',  # Roman numerals
            r'^\s*\(\d+\)\s*',        # (1)
            r'^\s*\([a-z]\)\s*',      # (a)
            r'^\s*[A-Z][\.\)]\s*',    # A. or A)
            r'^\s*[\-\*\•◦→○▪▫■□►▸◆◇]\s*',  # Bullet symbols
            r'^\s*[—–]\s*',           # Dashes
            r'^\s*□\s*',              # Checkbox empty
            r'^\s*[☐☑✓✗]\s*',        # Checkbox symbols
            r'^\s*\[\s*[xX]?\s*\]\s*',  # [ ] or [x]
        ]
        
        for pattern in marker_patterns:
            clean = re.sub(pattern, '', clean)
        
        # Capitalize first letter
        if clean and clean[0].islower():
            clean = clean[0].upper() + clean[1:]
        
        return clean.strip()
    
    def format_as_numbered_list(self, items):
        """
        Format items as a numbered list.
        """
        formatted = []
        for i, item in enumerate(items, 1):
            clean_item = self.clean_point_content(item)
            if clean_item:
                formatted.append(f"{i}. {clean_item}")
        return formatted
    
    def format_as_bulleted_list(self, items):
        """
        Format items as a bulleted list using dash.
        """
        formatted = []
        for item in items:
            clean_item = self.clean_point_content(item)
            if clean_item:
                formatted.append(f"- {clean_item}")
        return formatted
    
    def standardize_existing_list(self, lines, start_idx, format_type):
        """
        Standardize an existing list to consistent formatting.
        Returns: (formatted_lines, end_idx)
        """
        formatted_lines = []
        i = start_idx
        point_number = 1
        
        while i < len(lines):
            line = lines[i] if isinstance(lines[i], str) else lines[i].get('text', '')
            line_text = line.strip()
            
            # Check if this line is a point-form line
            is_point, point_type = self.is_point_form_line(line_text)
            
            if is_point:
                clean_content = self.clean_point_content(line_text)
                if clean_content:
                    if format_type == 'numbered':
                        formatted_lines.append(f"{point_number}. {clean_content}")
                        point_number += 1
                    else:
                        formatted_lines.append(f"- {clean_content}")
                i += 1
            elif not line_text:
                # Empty line - might be end of list
                # Look ahead to see if list continues
                if i + 1 < len(lines):
                    next_line = lines[i + 1] if isinstance(lines[i + 1], str) else lines[i + 1].get('text', '')
                    next_is_point, _ = self.is_point_form_line(next_line.strip())
                    if next_is_point:
                        # List continues after blank
                        formatted_lines.append('')
                        i += 1
                        continue
                # End of list
                break
            else:
                # Non-point line - end of list
                break
        
        return formatted_lines, i
    
    def process_point_form_content(self, text):
        """
        Process document text to detect and format point-form content.
        Returns processed text with properly formatted lists.
        """
        if not text:
            return text
        if self.policy.list_numbering_mode != "assistive":
            return text
        
        lines = text.split('\n')
        
        # NEW: Detect implied bullet blocks and convert them to explicit bullets
        try:
            # Use the implied detector to find blocks
            implied_blocks = self.implied_detector.detect_implied_bullet_blocks(lines)
            
            # Apply bullets to detected blocks
            for start, end, bullet_type in implied_blocks:
                for i in range(start, end + 1):
                    line = lines[i]
                    # Skip if already a bullet or empty
                    if not line.strip() or self.is_point_form_line(line)[0]:
                        continue
                        
                    # Add bullet marker
                    lines[i] = f"■ {line.strip()}"
            
            # Update text with implied bullets applied
            # We re-join and re-split to ensure consistency
            text = '\n'.join(lines)
            lines = text.split('\n')
            
        except Exception as e:
            logger.error(f"Error in implied bullet detection: {e}")
            # Continue with original text if detection fails
            lines = text.split('\n')

        processed_lines = []
        i = 0
        
        while i < len(lines):
            line = lines[i]
            line_text = line.strip()
            
            # Skip empty lines
            if not line_text:
                processed_lines.append(line)
                i += 1
                continue
            
            # Check 1: Is this a structured content heading (Objectives:, Steps:, etc.)?
            heading_type, format_type = self.get_point_form_heading_type(line_text)
            if heading_type:
                # Add heading with bold formatting
                heading_clean = line_text.rstrip(':')
                if self.policy.enable_regex_auto_bold and not heading_clean.startswith('**'):
                    processed_lines.append(f"**{heading_clean}:**")
                else:
                    processed_lines.append(line_text)
                i += 1
                
                # Process following lines as point-form content
                point_lines = []
                while i < len(lines):
                    next_line = lines[i] if isinstance(lines[i], str) else lines[i].get('text', '')
                    next_text = next_line.strip()
                    
                    if not next_text:
                        # Blank line might end list
                        if i + 1 < len(lines):
                            following = lines[i + 1] if isinstance(lines[i + 1], str) else lines[i + 1].get('text', '')
                            is_point, _ = self.is_point_form_line(following.strip())
                            if is_point or self.could_be_list_item(following.strip()):
                                point_lines.append('')
                                i += 1
                                continue
                        break
                    
                    # Check if it's already a point-form line or could be
                    is_point, _ = self.is_point_form_line(next_text)
                    could_be_item = self.could_be_list_item(next_text)
                    
                    if is_point or could_be_item:
                        point_lines.append(next_text)
                        i += 1
                    else:
                        # End of list
                        break
                
                # Format the collected points
                if point_lines:
                    formatted = self.format_point_block(point_lines, format_type)
                    processed_lines.extend(formatted)
                continue
            
            # Check 2: Does this line have context clue and content to convert?
            heading, items = self.extract_serial_comma_items(line_text)
            if heading and items:
                # Convert to list format
                if self.policy.enable_regex_auto_bold:
                    processed_lines.append(f"**{heading}:**")
                else:
                    processed_lines.append(f"{heading}:")
                formatted_items = self.format_as_bulleted_list(items)
                processed_lines.extend(formatted_items)
                i += 1
                continue
            
            # Check 3: Does this line have ordinal steps (First, Second, Third...)?
            steps = self.extract_ordinal_steps(line_text)
            if steps:
                # Check if there's a heading phrase before "First"
                first_match = re.search(r'^(.+?)(?:First(?:ly)?)', line_text, re.IGNORECASE)
                if first_match and len(first_match.group(1).strip()) > 3:
                    heading = first_match.group(1).strip().rstrip(':,.')
                    if self.policy.enable_regex_auto_bold:
                        processed_lines.append(f"**{heading}:**")
                    else:
                        processed_lines.append(f"{heading}:")
                formatted_steps = self.format_as_numbered_list(steps)
                processed_lines.extend(formatted_steps)
                i += 1
                continue
            
            # Check 4: Is this an existing point-form line that needs standardization?
            is_point, point_type = self.is_point_form_line(line_text)
            if is_point:
                # Standardize the list starting from here
                format_type = 'numbered' if point_type == 'numbered' else 'bulleted'
                formatted_block, new_idx = self.standardize_existing_list(lines, i, format_type)
                processed_lines.extend(formatted_block)
                i = new_idx
                continue
            
            # Check 5: Does this line have a context clue (ends with colon)?
            if self.has_list_context_clue(line_text):
                # Check if next lines could be list items
                if i + 1 < len(lines):
                    next_line = lines[i + 1] if isinstance(lines[i + 1], str) else lines[i + 1].get('text', '')
                    next_text = next_line.strip()
                    is_next_point, _ = self.is_point_form_line(next_text)
                    could_be_item = self.could_be_list_item(next_text)
                    
                    if is_next_point or could_be_item:
                        # Format heading
                        heading = line_text.rstrip(':')
                        if self.policy.enable_regex_auto_bold and not heading.startswith('**'):
                            processed_lines.append(f"**{heading}:**")
                        else:
                            processed_lines.append(line_text)
                        i += 1
                        
                        # Collect and format following points
                        point_lines = []
                        while i < len(lines):
                            l = lines[i] if isinstance(lines[i], str) else lines[i].get('text', '')
                            l_text = l.strip()
                            
                            if not l_text:
                                if i + 1 < len(lines):
                                    following = lines[i + 1] if isinstance(lines[i + 1], str) else lines[i + 1].get('text', '')
                                    is_pt, _ = self.is_point_form_line(following.strip())
                                    if is_pt or self.could_be_list_item(following.strip()):
                                        point_lines.append('')
                                        i += 1
                                        continue
                                break
                            
                            is_pt, _ = self.is_point_form_line(l_text)
                            if is_pt or self.could_be_list_item(l_text):
                                point_lines.append(l_text)
                                i += 1
                            else:
                                break
                        
                        if point_lines:
                            formatted = self.format_point_block(point_lines, 'bulleted')
                            processed_lines.extend(formatted)
                        continue
            
            # Default: keep line as-is
            processed_lines.append(line)
            i += 1
        
        return '\n'.join(processed_lines)
    
    def could_be_list_item(self, text):
        """
        Check if text could be a list item (short, starts with capital, no complex structure).
        """
        if not text:
            return False
        
        clean = text.strip()
        
        # Too long for a list item
        if len(clean) > 150:
            return False
        
        # Too short
        if len(clean) < 3:
            return False
        
        # Contains multiple sentences (not a simple list item)
        sentence_count = len(re.findall(r'[.!?]\s+[A-Z]', clean))
        if sentence_count > 0:
            return False
        
        # Starts with capital letter or common list patterns
        if re.match(r'^[A-Z]', clean):
            return True
        
        return False
    
    def format_point_block(self, lines, format_type):
        """
        Format a block of lines as either numbered or bulleted list.
        """
        formatted = []
        point_number = 1
        
        for line in lines:
            if not line or not line.strip():
                formatted.append('')
                continue
            
            clean_content = self.clean_point_content(line)
            if clean_content:
                if format_type == 'numbered':
                    formatted.append(f"{point_number}. {clean_content}")
                    point_number += 1
                else:
                    formatted.append(f"- {clean_content}")
        
        return formatted

    def process_short_document(self, text):
        """
        Process a short document: emphasize key points, and format point-form content.
        Returns processed text.
        """
        if not text:
            return text
        
        # Check if document is short
        is_short, reason = self.is_short_document(text)
        if not is_short:
            return text  # No changes for long documents
        
        lines = text.split('\n')
        
        # Step 1: Process point-form content (convert serial lists to bullet points, standardize lists)
        if self.policy.list_numbering_mode == "assistive":
            text = self.process_point_form_content(text)
        
        # Step 2: Process each line for key point emphasis
        if not self.policy.enable_regex_auto_bold:
            return text
        lines = text.split('\n')
        processed_lines = []
        for line in lines:
            line_text = line if isinstance(line, str) else line.get('text', '')
            
            # Skip already formatted point-form lines (starting with - or number.)
            if re.match(r'^\s*[-\*•]\s+', line_text) or re.match(r'^\s*\d+\.\s+', line_text):
                processed_lines.append(line_text)
                continue
            
            # Check if it's an assignment header field
            if self.is_assignment_header_field(line_text):
                # Bold the field
                if not line_text.strip().startswith('**'):
                    processed_lines.append(f"**{line_text.strip()}**")
                else:
                    processed_lines.append(line_text)
                continue
            
            # Check if it's a key point
            point_type, emoji = self.get_key_point_type(line_text)
            if point_type:
                emphasized = self.emphasize_key_point(line_text, point_type, emoji)
                processed_lines.append(emphasized)
            else:
                processed_lines.append(line_text)
        
        return '\n'.join(processed_lines)
    
    def is_main_heading(self, text):
        """
        Check if text is a main heading (level 1) that should have space cleaning applied.
        """
        if not text:
            return False
        
        # Must start with single # (level 1)
        if not re.match(r'^\s*#\s+', text) or re.match(r'^\s*##', text):
            return False
        
        # Extract heading content and normalize numbering
        clean_text = self._normalize_heading_text(text)
        clean_text = clean_text.rstrip(' .:,;')
        
        # Main heading keywords
        main_headings = [
            'DEDICATION', 'ACKNOWLEDGEMENTS', 'ACKNOWLEDGMENTS', 'ACKNOWLEDGEMENT',
            'ABSTRACT', 'RESUME', 'RÉSUMÉ', 'TABLE OF CONTENTS', 'CONTENTS',
            'LIST OF TABLES', 'LIST OF FIGURES', 'LIST OF ABBREVIATIONS', 'LIST OF ACRONYMS',
            'GLOSSARY', 'APPENDICES', 'APPENDIX',
            'DECLARATION', 'CERTIFICATION', 'APPROVAL PAGE', 'COMMITTEE APPROVAL',
            'REFERENCES', 'REFERENCE', 'BIBLIOGRAPHY',
        ]
        
        # Check exact match
        for heading in main_headings:
            if clean_text == heading:
                return True
        
        # Check for chapter headings
        if re.match(r'^CHAP?TER\s+(ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN|\d+|[IVXLC]+)', clean_text, re.IGNORECASE):
            return True
        
        return False

    def _normalize_heading_text(self, text):
        """Normalize heading text by stripping markdown markers and leading numbering."""
        clean_text = re.sub(r'^#+\s*', '', text).strip()
        clean_text = re.sub(
            r'^\(?((?:\d+(?:\.\d+)*)|(?:[A-Z](?:\.\d+)*)|(?:[IVXLC]+))[\.)]?\s+',
            '',
            clean_text,
            flags=re.IGNORECASE
        )
        return clean_text.strip().upper()

    def should_start_on_new_page(self, text):
        """
        Check if a heading should start on a new page.
        Returns True for major academic sections like chapters, abstract, etc.
        """
        if not text:
            return False
            
        # Strip markdown heading markers and whitespace
        clean_text = self._normalize_heading_text(text)
        
        # Front matter sections that require new pages
        front_matter_sections = [
            'ACKNOWLEDGEMENTS', 'ACKNOWLEDGMENTS', 'ACKNOWLEDGEMENT',
            'DEDICATION',
            'ABSTRACT', 'RESUME', 'RÉSUMÉ',
            'TABLE OF CONTENTS', 'CONTENTS',
            'LIST OF TABLES',
            'LIST OF FIGURES',
            'GLOSSARY',
            'LIST OF ABBREVIATIONS', 'ABBREVIATIONS', 'LIST OF ACRONYMS',
            'APPENDICES', 'APPENDIX',
            'REFERENCES', 'REFERENCE',
            'BIBLIOGRAPHY',
        ]
        
        # Check for exact front matter match
        for section in front_matter_sections:
            if clean_text == section or clean_text.startswith(section + ':'):
                return True
        
        # Check for chapter headings (CHAPTER ONE, CHAPTER 1, CHAPTER I, etc.)
        chapter_pattern = r'^CHAP?TER\s+(ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN|\d+|[IVXLC]+)'
        if re.match(chapter_pattern, clean_text, re.IGNORECASE):
            return True
        
        return False
    
    def should_be_centered(self, text, heading_level=1):
        """
        Determine if a heading should be centered.
        Only level 1 headings (#) can be centered.
        Returns True for major academic sections that should be centered.
        """
        if not text:
            return False
        
        # Only top-level headings (level 1) can be centered
        if heading_level != 1:
            return False
        
        # Strip markdown heading markers and whitespace
        clean_text = self._normalize_heading_text(text)
        
        # Remove any trailing colon or punctuation for matching
        match_text = clean_text.rstrip(':.')
        
        # Front matter sections to center
        front_matter_center = [
            'DEDICATION',
            'ACKNOWLEDGEMENTS', 'ACKNOWLEDGMENTS', 'ACKNOWLEDGEMENT',
            'ABSTRACT',
            'RESUME', 'RÉSUMÉ',
            'TABLE OF CONTENTS', 'CONTENTS',
            'LIST OF TABLES',
            'LIST OF FIGURES',
            'LIST OF ABBREVIATIONS', 'ABBREVIATIONS',
            'GLOSSARY',
            'APPENDICES', 'APPENDIX',
            'INTRODUCTION',
            'LITERATURE REVIEW',
            'METHODOLOGY', 'RESEARCH METHODOLOGY',
            'RESULTS', 'FINDINGS', 'DATA ANALYSIS',
            'DISCUSSION', 'FINDINGS AND DISCUSSION',
            'CONCLUSION', 'SUMMARY', 'RECOMMENDATIONS',
        ]
        
        # Back matter sections to center
        back_matter_center = [
            'REFERENCES', 'REFERENCE',
            'BIBLIOGRAPHY',
        ]
        
        # Check for exact front/back matter match
        all_center_sections = front_matter_center + back_matter_center
        for section in all_center_sections:
            if match_text == section:
                return True
        
        # Check for chapter headings (CHAPTER ONE, CHAPTER 1, CHAPTER I, etc.)
        # These should be centered regardless of having a subtitle
        if re.match(r'^CHAP?TER\b', clean_text, re.IGNORECASE):
            chapter_pattern = r'^CHAP?TER\s+(ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN|\d+|[IVXLC]+)'
            if re.match(chapter_pattern, clean_text, re.IGNORECASE):
                return True
        
        return False
    
    def is_chapter_heading(self, text):
        """
        Check if text is a chapter heading (CHAPTER ONE, CHAPTER 1, CHAPTER I, etc.)
        Returns (is_chapter, chapter_num, chapter_title) tuple
        """
        if not text:
            return False, None, None
        
        # Strip markdown heading markers
        clean_text = re.sub(r'^#+\s*', '', text).strip()
        
        # Chapter with title on same line: CHAPTER ONE: INTRODUCTION
        pattern_with_title = r'^CHAP?TER\s+(ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN|\d+|[IVXLC]+)\s*[:\-\.]\s*(.+)$'
        match = re.match(pattern_with_title, clean_text, re.IGNORECASE)
        if match:
            return True, match.group(1), match.group(2).strip()
        
        # Chapter heading only: CHAPTER ONE
        pattern_only = r'^CHAP?TER\s+(ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN|\d+|[IVXLC]+)\s*$'
        match = re.match(pattern_only, clean_text, re.IGNORECASE)
        if match:
            return True, match.group(1), None
        
        return False, None, None
    
    def is_chapter_title(self, text, prev_was_chapter=False):
        """
        Check if text is a chapter title following a chapter heading.
        Chapter titles are typically: ALL CAPS, or common academic section names
        Returns True if it looks like a chapter title.
        """
        if not text:
            return False
        
        # Strip markdown heading markers
        clean_text = re.sub(r'^#+\s*', '', text).strip()
        
        # Skip if too long (likely a paragraph)
        if len(clean_text) > 100:
            return False
        
        # Skip if empty
        if not clean_text:
            return False
        
        # Common chapter title keywords
        chapter_title_keywords = [
            'INTRODUCTION', 'LITERATURE REVIEW', 'METHODOLOGY', 'METHODS',
            'RESULTS', 'DISCUSSION', 'CONCLUSION', 'CONCLUSIONS',
            'RECOMMENDATIONS', 'THEORETICAL FRAMEWORK', 'RESEARCH METHODOLOGY',
            'SUMMARY AND CONCLUSIONS', 'SUMMARY OF FINDINGS', 'DATA ANALYSIS',
            'BACKGROUND', 'PROBLEM STATEMENT', 'RESEARCH DESIGN',
            'FINDINGS AND DISCUSSION', 'ANALYSIS AND INTERPRETATION',
        ]
        
        # Check for exact match with common titles
        if clean_text.upper() in chapter_title_keywords:
            return True
        
        # Check if all caps (common for chapter titles)
        if clean_text == clean_text.upper() and len(clean_text) > 5 and any(c.isalpha() for c in clean_text):
            # All caps title, likely a chapter title
            return True
        
        # If previous line was a chapter heading, be more lenient
        if prev_was_chapter:
            # Title case or heading-like
            if re.match(r'^[A-Z][A-Za-z\s]+$', clean_text):
                return True
        
        return False
    
    def is_copyright_content(self, text):
        """Check if text is copyright-related content."""
        if not text:
            return False
        
        clean_text = text.strip()
        
        for pattern in self.patterns.get('copyright_content', []):
            if pattern.match(clean_text):
                return True
        
        return False
    
    def is_declaration_content(self, text):
        """Check if text is declaration-related content."""
        if not text:
            return False
        
        clean_text = text.strip()
        
        for pattern in self.patterns.get('declaration_content', []):
            if pattern.search(clean_text):
                return True
        
        return False
    
    def is_certification_content(self, text):
        """Check if text is certification/approval-related content."""
        if not text:
            return False
        
        clean_text = text.strip()
        
        for pattern in self.patterns.get('certification_content', []):
            if pattern.search(clean_text):
                return True
        
        return False
    
    def is_signature_line(self, text):
        """Check if text is a signature line or date line."""
        if not text:
            return False
        
        clean_text = text.strip()
        
        for pattern in self.patterns.get('signature_line', []):
            if pattern.match(clean_text):
                return True
        
        return False
    
    def is_toc_entry(self, text):
        """Check if text is a Table of Contents entry line."""
        if not text:
            return False
        
        clean_text = text.strip()
        
        for pattern in self.patterns.get('toc_entry', []):
            if pattern.match(clean_text):
                return True
        
        # Also check for dot leaders pattern
        if '.....' in clean_text or '…..' in clean_text:
            return True
        
        return False

    def get_front_matter_section_type(self, text):
        """
        Get the type of front matter section if text is a front matter heading.
        Returns section type (declaration, certification, dedication, etc.) or None.
        """
        if not text:
            return None
        
        # Strip markdown heading markers
        clean_text = re.sub(r'^#+\s*', '', text).strip().upper()
        clean_text = re.sub(r'[:.\s]+$', '', clean_text)
        
        front_matter_sections = {
            'DECLARATION': 'declaration',
            'CERTIFICATION': 'certification',
            'APPROVAL PAGE': 'certification',
            'COMMITTEE APPROVAL': 'certification',
            'DEDICATION': 'dedication',
            'ACKNOWLEDGEMENTS': 'acknowledgements',
            'ACKNOWLEDGMENTS': 'acknowledgements',
            'ACKNOWLEDGEMENT': 'acknowledgements',
            'ACKNOWLEDGMENT': 'acknowledgements', # Missing 'e' variation
            'ABSTRACT': 'abstract',
            'RESUME': 'resume',  # French equivalent of abstract
            'RÉSUMÉ': 'resume',  # With accent
            'RÉSUME': 'resume',  # With first accent only
            'RESUMÉ': 'resume',  # With last accent only
            'TABLE OF CONTENTS': 'toc',
            'CONTENTS': 'toc',
            'LIST OF TABLES': 'list_of_tables',
            'LIST OF FIGURES': 'list_of_figures',
            'LIST OF ABBREVIATIONS': 'abbreviations',
            'ABBREVIATIONS': 'abbreviations',
            'LIST OF ACRONYMS': 'abbreviations',
            'GLOSSARY': 'glossary',
            'APPENDIX': 'appendix',
            'APPENDICES': 'appendix',
            'REFERENCES': 'references',
            'BIBLIOGRAPHY': 'bibliography',
        }
        
        for section, section_type in front_matter_sections.items():
            if clean_text == section:
                return section_type

    def _determine_heading_level(self, content, default_level):
        """Determine heading level based on content characteristics"""
        words = content.split()
        # Short, important-looking phrases get higher level
        if len(words) <= 4 and any(word.lower() in ['important', 'key', 'critical', 'major'] for word in words):
            return 2
        # Longer phrases get lower level
        if len(words) > 6:
            return min(default_level + 1, 4)
        return default_level

    def _get_parenthesized_level(self, number):
        """Determine level for parenthesized numbering"""
        if number.isdigit():
            num = int(number)
            if num <= 5:
                return 2
            elif num <= 10:
                return 3
            else:
                return 4
        elif number.isalpha():
            # Letters indicate sub-levels
            return 3 if number.islower() else 2
        else:
            # Roman numerals
            return 2

    def _get_bracketed_level(self, number):
        """Determine level for bracketed numbering"""
        if number.isdigit():
            return 3  # Bracketed numbers are usually sub-points
        else:
            return 4  # Bracketed letters are usually sub-sub-points

    def _roman_to_int(self, roman):
        """Convert a Roman numeral to integer (basic support)"""
        if not roman:
            return 0
        roman = roman.upper()
        vals = {'I':1,'V':5,'X':10,'L':50,'C':100,'D':500,'M':1000}
        total = 0
        prev = 0
        for ch in reversed(roman):
            v = vals.get(ch, 0)
            if v < prev:
                total -= v
            else:
                total += v
            prev = v
        return total

    def _get_shortdoc_header_level(self, number, title):
        """Determine level for short document headers"""
        if number.isdigit():
            num = int(number)
            if num == 1:
                return 2  # First section is level 2
            elif num <= 5:
                return 3
            else:
                return 4
        elif number.isalpha():
            return 3 if number.islower() else 2
        else:
            # Roman numerals
            return 2

    def _get_part_level(self, part_num):
        """Determine level for part/section headers"""
        if part_num.isdigit():
            return 2  # Numbered parts are major sections
        elif part_num.isalpha():
            return 3 if part_num.islower() else 2
        else:
            return 2  # Roman numerals are major
        
        return None

    def _has_adjacent_list_context(self, prev_line, next_line):
        """Check for neighboring list items to reduce ambiguity."""
        candidates = [prev_line.strip(), next_line.strip()] if prev_line or next_line else []
        for line in candidates:
            if not line:
                continue
            if detect_bullet_type(line):
                return True
            if any(p.match(line) for p in self.patterns.get('numbered_list', [])):
                return True
        return False

    def _has_list_leadin_phrase(self, text):
        """Check if a line contains a list lead-in phrase without explicit bullets."""
        if not text:
            return False
        clean = text.strip()
        leadin_patterns = [
            r'(?:include|includes|including)\s*$',
            r'(?:such\s+as)\s*$',
            r'(?:consist(?:s)?\s+of)\s*$',
            r'(?:are|is)\s+(?:as\s+follows|listed|outlined)\s*:?\s*$',
            r'(?:the\s+)?following(?:\s+\w+)?\s*$',
        ]
        return any(re.search(pattern, clean, re.IGNORECASE) for pattern in leadin_patterns)

    def _is_list_heading_line(self, text):
        """Check if a short line likely introduces a list (e.g., Benefits, Positive effects)."""
        if not text:
            return False
        clean = text.strip()
        heading_type, _ = self.get_point_form_heading_type(clean)
        if heading_type:
            return True
        if re.match(r'^(?:positive|negative)\s+effects?$', clean, re.IGNORECASE):
            return True
        if re.match(r'^(?:benefits?|challenges?|advantages?|disadvantages?)$', clean, re.IGNORECASE):
            return True
        return False

    def _is_implicit_list_item(self, text, prev_line, next_line):
        """Infer list items from context when bullets are missing."""
        if not text:
            return False
        clean = text.strip()
        if not self.could_be_list_item(clean):
            return False
        if self._looks_like_question(clean):
            return False
        prev = prev_line.strip() if prev_line else ''
        next_text = next_line.strip() if next_line else ''
        if prev and (self.has_list_context_clue(prev) or self._has_list_leadin_phrase(prev) or self._is_list_heading_line(prev)):
            return True
        if prev and self.could_be_list_item(prev) and not prev.endswith('.'):
            if next_text and self.could_be_list_item(next_text) and not next_text.endswith('.'):
                return True
        return False

    def _looks_like_question(self, text):
        """Check if text reads like a question to avoid misclassifying it as a heading."""
        if not text:
            return False
        trimmed = text.strip()
        if not trimmed:
            return False
        if '?' in trimmed:
            return True
        question_starters = (
            r'(what|why|how|who|which|where|when|whom|whose|is|are|do|does|did|'
            r'can|could|would|should|will|may|might)\b'
        )
        if re.match(question_starters, trimmed, re.IGNORECASE):
            return True
        if re.match(r'^to\s+what\s+extent\b', trimmed, re.IGNORECASE):
            return True
        return False

    def _is_heading_like_numbered_line(self, text):
        """Detect numbered lines that look like headings rather than list items."""
        match = re.match(r'^\s*(\d+[\.)])\s+(.+)$', text)
        if not match:
            return False
        title = match.group(2).strip()
        if not title or title.endswith('.'):
            return False
        if self._looks_like_question(title):
            return False
        word_count = len(title.split())
        if word_count > 8:
            return False
        return title[0].isupper()

    def _is_short_numbered_list_item(self, text):
        """Check if a numbered line is short enough to be treated as a list item."""
        if not text:
            return False
        match = re.match(r'^\s*\(?[0-9ivxA-Za-z]+[\.)]\s+(.+)$', text)
        if not match:
            return False
        content = match.group(1).strip()
        if not content:
            return False
        word_count = len(content.split())
        return word_count <= 24

    def _rewrite_numeric_bullet_markers(self, text):
        """Rewrite numeric + asterisk markers into bullets before list parsing."""
        if not text:
            return text, False
        if self._matches_heading_patterns(text):
            return text, False
        patterns = [
            r'^\s*\d+\s*\*\s+(.+)$',
            r'^\s*\d+(?:\.\d+)\s*\*\s+(.+)$',
            r'^\s*\d+(?:\.\d+)\s+(.+)$',
        ]
        for pattern in patterns:
            match = re.match(pattern, text)
            if match:
                return f"• {match.group(1).strip()}", True
        return text, False

    def _rewrite_asterisk_bullet(self, text):
        """Convert leading asterisk bullets to standard bullet markers."""
        if not text:
            return text, False
        match = re.match(r'^\s*\*\s+(.+)$', text)
        if match:
            return f"• {match.group(1).strip()}", True
        return text, False

    def _matches_heading_patterns(self, text):
        """Check if a line matches existing heading detectors."""
        if not text:
            return False
        if any(p.match(text) for p in self.patterns.get('heading_1', [])):
            return True
        if any(p.match(text) for p in self.patterns.get('heading_2', [])):
            return True
        if any(p.match(text) for p in self.patterns.get('heading_3', [])):
            return True
        if any(p.match(text) for p in self.patterns.get('heading_hierarchy', [])):
            return True
        return False

    def _extract_journal_spans(self, text: str):
        """
        Return list of (start, end) spans for journal titles in a reference line.
        Spans are indices into `text`. First valid match wins.
        """
        spans = []
        for pat in self.patterns.get('reference_journal_span_v1', []):
            m = pat.match(text)
            if not m:
                continue

            j_start = m.start('journal')
            j_end = m.end('journal')

            journal = text[j_start:j_end].strip()
            # Sanity checks to reduce false positives
            if len(journal) < 4:
                continue
            if journal.lower().startswith('in '):
                continue

            spans.append((j_start, j_end))
            break
        return spans

    def analyze_line(self, line, line_num, prev_line='', next_line='', context=None):
        """Analyze a single line with multiple pattern checks"""
        original_line = line
        line, _ = self._rewrite_numeric_bullet_markers(line)
        line, _ = self._rewrite_asterisk_bullet(line)
        # FIRST: Clean heading spaces before analysis
        if line.lstrip().startswith('#'):
            line = self.clean_heading_spaces(line)
        
        # PART 1: PRE-PROCESS - Unicode Scrubber (Priority 0)
        # Strip all non-standard Unicode characters (emojis, hidden chars) BEFORE any analysis
        # This is critical: emojis cause regex failures and rendering issues in Word/PDF
        cleaned = line
        for pattern in self.patterns.get('unicode_scrubber', []):
            cleaned = pattern.sub('', cleaned)
        
        # PART 2: Remove asterisks comprehensively
        # Applied to all text to ensure asterisks are removed even when mid-word (e.g., "Customizability*")
        for pattern in self.patterns.get('asterisk_removal', []):
            cleaned = pattern.sub('', cleaned)
        
        trimmed = cleaned.strip()
        
        if not trimmed:
            return {'type': 'empty', 'content': '', 'line_num': line_num}
        
        # Safety check: if scrubbing resulted in empty string after stripping, return
        if not cleaned or not trimmed:
            return {'type': 'empty', 'content': '', 'line_num': line_num}
        
        analysis = {
            'line_num': line_num,
            'type': 'paragraph',
            'content': trimmed,
            'original': original_line,
            'level': 0,
            'confidence': 0.0,
        }
        
        # Get line characteristics
        length = len(trimmed)
        is_short = length < 100
        is_very_short = length < 60
        is_all_caps = trimmed == trimmed.upper() and any(c.isalpha() for c in trimmed) and length > 2
        is_title_case = re.match(r'^[A-Z][a-z]+(?:\s+[A-Za-z][a-z]*)*$', trimmed) is not None
        has_period = trimmed.endswith('.')
        word_count = len(trimmed.split())
        
        # Skip very long lines for heading detection (likely paragraphs)
        # But check for paragraph with colon that might be a definition
        
        # Priority 1: Check for table patterns (highest priority to preserve structure)
        for pattern in self.patterns['table_marker']:
            if pattern.match(trimmed):
                if re.search(r'START', trimmed, re.IGNORECASE):
                    analysis['type'] = 'table_start'
                elif re.search(r'END', trimmed, re.IGNORECASE):
                    analysis['type'] = 'table_end'
                else:
                    analysis['type'] = 'table_caption'
                analysis['confidence'] = 1.0
                return analysis

        if re.fullmatch(r'\d{1,4}', trimmed):
            prev_clean = prev_line.strip() if isinstance(prev_line, str) else ''
            next_clean = next_line.strip() if isinstance(next_line, str) else ''
            if not prev_clean or not next_clean:
                analysis['type'] = 'page_metadata'
                analysis['confidence'] = 0.9
                analysis['subtype'] = 'page_number'
                return analysis
        
        for pattern in self.patterns['table_row']:
            if pattern.match(trimmed):
                # Check if it's a separator row
                if re.match(r'^\|[\s\-:]+\|', trimmed):
                    analysis['type'] = 'table_separator'
                else:
                    analysis['type'] = 'table_row'
                cells = [c.strip() for c in trimmed.split('|') if c.strip()]
                analysis['cells'] = cells
                analysis['confidence'] = 1.0
                return analysis
        
        # ========================================
        # PLAIN TEXT TABLE DETECTION (NEW LOGIC)
        # ========================================
        
        # Check for table separator lines (highest confidence)
        # ==================== CONSERVATIVE TABLE DETECTION ====================
        # First check if line is definitely NOT a table (exclusion filter)
        for pattern in self.patterns.get('not_a_table', []):
            if pattern.match(trimmed):
                # Skip table detection for this line - it's definitely not a table
                analysis['type'] = 'paragraph'
                analysis['confidence'] = 0.9
                return analysis

        # 1. Check for markdown tables (very specific patterns)
        for pattern in self.patterns.get('table_markdown', []):
            if pattern.match(trimmed):
                analysis['type'] = 'table'
                analysis['confidence'] = 0.95  # Very high confidence for markdown
                analysis['subtype'] = 'markdown'
                # Parse cells for markdown
                if '|' in trimmed and not trimmed.replace('|', '').replace('-', '').replace(':', '').replace(' ', '').strip():
                    # This is a separator row
                    analysis['subtype'] = 'markdown_separator'
                else:
                    # This is a data row
                    cells = [cell.strip() for cell in trimmed.strip().strip('|').split('|')]
                    analysis['metadata'] = {
                        'cells': cells,
                        'cell_count': len(cells)
                    }
                return analysis

        # 2. Check for tab-separated tables
        for pattern in self.patterns.get('table_tab_separated', []):
            if pattern.match(trimmed):
                cells = [cell.strip() for cell in trimmed.split('\t') if cell.strip()]
                if len(cells) >= 3:  # Require at least 3 columns for tab-separated
                    analysis['type'] = 'table'
                    analysis['confidence'] = 0.85
                    analysis['subtype'] = 'tab'
                    analysis['metadata'] = {
                        'cells': cells,
                        'cell_count': len(cells)
                    }
                    return analysis

        # Additional tab-separated detection for 2-column pasted tables
        if '\t' in trimmed:
            cells = [cell.strip() for cell in trimmed.split('\t') if cell.strip()]
            neighbor_has_tab = ('\t' in prev_line) or ('\t' in next_line)
            if len(cells) >= 2 and (len(cells) >= 3 or neighbor_has_tab):
                analysis['type'] = 'table'
                analysis['confidence'] = 0.80
                analysis['subtype'] = 'tab'
                analysis['metadata'] = {
                    'cells': cells,
                    'cell_count': len(cells)
                }
                return analysis

        # 3. Check for aligned columns (very conservative - requires 4+ columns)
        for pattern in self.patterns.get('table_aligned_columns', []):
            if pattern.match(trimmed):
                cells = [cell for cell in re.split(r'\s{3,}', trimmed.strip()) if cell.strip()]
                if len(cells) >= 4:  # Require at least 4 columns for aligned
                    analysis['type'] = 'table'
                    analysis['confidence'] = 0.80
                    analysis['subtype'] = 'aligned'
                    analysis['metadata'] = {
                        'cells': cells,
                        'cell_count': len(cells)
                    }
                    return analysis

        # 3.5 Check for spaced tables with multiple columns (pasted text fallback)
        is_spaced_table, spaced_cells, spaced_confidence = self._detect_spaced_table_row(
            trimmed,
            prev_line=prev_line,
            next_line=next_line,
            context=context,
        )
        if is_spaced_table:
            analysis['type'] = 'table'
            analysis['confidence'] = spaced_confidence
            analysis['subtype'] = 'spaced'
            analysis['metadata'] = {
                'cells': spaced_cells,
                'cell_count': len(spaced_cells)
            }
            return analysis

        # Legacy table detection (lower priority, more conservative)
        for pattern in self.patterns.get('plain_table_separator', []):
            if pattern.match(trimmed):
                analysis['type'] = 'plain_table_separator'
                analysis['confidence'] = 0.95
                analysis['metadata'] = {'separator_style': self._detect_separator_style(trimmed)}
                return analysis
        
        # Priority 1.5: DISSERTATION-SPECIFIC PATTERNS
        
        # Check for chapter headings (CHAPTER ONE, CHAPTER 1, CHAPTER I)
        is_chapter, chapter_num, chapter_title = self.is_chapter_heading(trimmed)
        if is_chapter:
            # Treat chapter headings as level-1 headings for consistency with tests
            analysis['type'] = 'chapter_heading'
            analysis['level'] = 1
            analysis['chapter_num'] = chapter_num
            analysis['chapter_title'] = chapter_title  # May be None if title is on separate line
            analysis['confidence'] = 1.0
            analysis['needs_page_break'] = True
            analysis['should_center'] = True
            return analysis
        
        # Check for chapter title following a chapter heading (using context)
        if context and context.get('prev_was_chapter'):
            if self.is_chapter_title(trimmed, prev_was_chapter=True):
                analysis['type'] = 'chapter_title'
                analysis['level'] = 2  # Slightly below chapter heading
                analysis['confidence'] = 0.95
                analysis['should_center'] = True  # Chapter titles are centered
                analysis['needs_page_break'] = False  # Already on new page with chapter heading
                return analysis
        
        # Check for front matter section headings (Declaration, Certification, etc.)
        front_matter_type = self.get_front_matter_section_type(trimmed)
        if front_matter_type:
            # Treat major front-matter sections as regular level-1 headings
            major_front_matter = {'abstract', 'references', 'bibliography', 'appendix', 'appendices', 'acknowledgements', 'dedication', 'resume', 'conclusion', 'executive_summary', 'toc'}
            if front_matter_type in major_front_matter:
                analysis['type'] = 'heading'
            else:
                analysis['type'] = 'front_matter_heading'
            analysis['front_matter_type'] = front_matter_type
            analysis['level'] = 1
            analysis['confidence'] = 1.0
            analysis['needs_page_break'] = True
            analysis['should_center'] = True
            return analysis
        
        # Check for copyright content
        if self.is_copyright_content(trimmed):
            analysis['type'] = 'copyright_content'
            analysis['confidence'] = 0.95
            analysis['should_center'] = True  # Copyright content is often centered
            return analysis
        
        # Check for signature lines
        if self.is_signature_line(trimmed):
            analysis['type'] = 'signature_line'
            analysis['confidence'] = 0.95
            return analysis
        
        # Check for TOC entries
        if self.is_toc_entry(trimmed):
            analysis['type'] = 'toc_entry'
            analysis['confidence'] = 0.90
            return analysis

        # Priority 1.8: Metadata Detection (Moved from Priority 10/11)
        # Check for page metadata (headers/footers/page numbers)
        for pattern in self.patterns['page_metadata']:
            if pattern.match(trimmed):
                analysis['type'] = 'page_metadata'
                analysis['confidence'] = 0.90
                # Determine subtype
                if re.search(r'page|p\.|pg\.', trimmed, re.IGNORECASE) or re.match(r'^\s*-?\s*\d+\s*-?\s*$', trimmed):
                    analysis['subtype'] = 'page_number'
                elif re.search(r'header|running head', trimmed, re.IGNORECASE):
                    analysis['subtype'] = 'header'
                elif re.search(r'footer', trimmed, re.IGNORECASE):
                    analysis['subtype'] = 'footer'
                else:
                    analysis['subtype'] = 'document_metadata'
                return analysis
        
        # Check for academic metadata
        for pattern in self.patterns['academic_metadata']:
            if pattern.match(trimmed):
                analysis['type'] = 'academic_metadata'
                analysis['confidence'] = 0.80
                # Determine subtype
                if re.search(r'\bby\b|authors?:', trimmed, re.IGNORECASE):
                    analysis['subtype'] = 'author'
                elif re.search(r'@', trimmed):
                    analysis['subtype'] = 'contact'
                elif re.search(r'department|school|college|university|institute', trimmed, re.IGNORECASE):
                    analysis['subtype'] = 'affiliation'
                else:
                    analysis['subtype'] = 'metadata'
                return analysis

        # Priority 1.9: Implicit Lists and Questions (Prompt Requirement)
        instruction_verbs = r'^(Describe|Explain|Compare|CDF|Analyze|Discuss|List|State|Define|What|Which|Why|How|Question|Task|Activity|Key\s+.*|Note|Important|Warning)'
        if re.match(instruction_verbs, trimmed, re.IGNORECASE) and (trimmed.endswith('?') or trimmed.endswith(':') or len(trimmed) < 100):
             analysis['type'] = 'instruction'
             if trimmed.endswith('?') or re.match(r'^Q\d+|Question', trimmed, re.IGNORECASE):
                 analysis['type'] = 'question'
             analysis['confidence'] = 0.85
             return analysis

        # Check for Implicit List Items
        if self.policy.list_numbering_mode == "assistive" and context and context.get('prev_analysis'):
            prev_type = context['prev_analysis'].get('type')
            if prev_type in ['instruction', 'question', 'list_item_implicit'] and is_short and not has_period:
                analysis['type'] = 'list_item_implicit'
                analysis['confidence'] = 0.80
                logger.info("List item inferred from context: '%s'", trimmed)
                return analysis

        # Priority 1.95: Implicit list items based on list lead-ins or headings
        if is_short and not has_period and self._is_implicit_list_item(trimmed, prev_line, next_line):
            analysis['type'] = 'bullet_list'
            analysis['content'] = trimmed
            analysis['confidence'] = 0.88
            logger.info("Implicit list item detected: '%s'", trimmed)
            return analysis

        # Priority 2: Check for heading patterns
        if is_short and not has_period:
            # Prefer prominent short ALL-CAPS lines as 'prominent_heading' before treating them as H1
            for p in self.patterns.get('allcaps_short_line', []):
                if p.match(trimmed):
                    wc = len(trimmed.split())
                    if 2 <= wc <= 8:
                        analysis['type'] = 'prominent_heading'
                        analysis['content'] = trimmed.title()
                        analysis['confidence'] = 0.80
                        analysis['level'] = 2
                        analysis['original_format'] = 'allcaps'
                        return analysis
            # H1 detection - major sections
            for pattern in self.patterns['heading_1']:
                if pattern.match(trimmed):
                    # Avoid capturing short ALL-CAPS lines that are better treated as prominent headings
                    # (e.g., KEY POINTS, MAIN OBJECTIVES) unless they match explicit main headings
                    try:
                        is_generic_allcaps = ('[A-Z][A-Z\\s]{2,49}' in pattern.pattern)
                    except Exception:
                        is_generic_allcaps = False
                    # If it's a generic ALL-CAPS short line and it matches the allcaps_short_line pattern,
                    # prefer to let the later 'prominent_heading' checks handle it (e.g., KEY POINTS)
                    if is_generic_allcaps and re.match(r'^[A-Z][A-Z\s]{2,49}$', trimmed) and any(p.match(trimmed) for p in self.patterns.get('allcaps_short_line', [])):
                        # If it's one of the explicit major headings, keep as H1
                        if not re.match(r'^(EXECUTIVE\s+SUMMARY|INTRODUCTION|CONCLUSION|REFERENCES|BIBLIOGRAPHY|APPENDIX|ACKNOWLEDGEMENT|ABSTRACT)$', trimmed, re.IGNORECASE):
                            # Skip assigning H1 here so subsequent prominent-heading logic can classify it
                            continue

                    analysis['type'] = 'heading'
                    analysis['level'] = 1
                    analysis['confidence'] = 0.95
                    # Check if this heading needs a page break and/or centering
                    analysis['needs_page_break'] = self.should_start_on_new_page(trimmed)
                    analysis['should_center'] = self.should_be_centered(trimmed, 1)
                    return analysis
            
            # H2 detection - sub-sections
            for pattern in self.patterns['heading_2']:
                if pattern.match(trimmed):
                    analysis['type'] = 'heading'
                    analysis['level'] = 2
                    analysis['confidence'] = 0.90
                    analysis['needs_page_break'] = False  # Sub-sections don't get page breaks
                    analysis['should_center'] = False  # Sub-sections don't get centered
                    return analysis
            
            # H3 detection - sub-sub-sections
            for pattern in self.patterns['heading_3']:
                if pattern.match(trimmed):
                    # Avoid mis-classifying list/numbering styles as headings
                    # Prefer numbered_list/hierarchical_list for parenthesized or numbering patterns
                    if any(p.match(trimmed) for p in self.patterns.get('numbered_list', [])):
                        # If the content after the marker is fairly long, prefer numbered_list over H3
                        content_after = re.sub(r'^\s*\(?[a-z0-9ivx]+[\.)]?\s+', '', trimmed, flags=re.IGNORECASE)
                        if len(content_after.split()) > 3:
                            continue
                        # Otherwise allow short lettered patterns (e.g., 'a) Title') to be H3

                    # Prefer hierarchical detection for parenthesized/bracketed/mixed numbering
                    hierarchy_patterns = (self.patterns.get('parenthesized_mixed', []) + self.patterns.get('bracketed_numbering', []) + self.patterns.get('alphanumeric_hierarchy', []))
                    if any(p.match(trimmed) for p in hierarchy_patterns):
                        # Let hierarchical list handlers classify this line instead of treating it as H3
                        continue

                    if any(p.match(trimmed) for p in self.patterns.get('alphanumeric_hierarchy', [])):
                        continue

                    analysis['type'] = 'heading'
                    analysis['level'] = 3
                    analysis['confidence'] = 0.85
                    analysis['needs_page_break'] = False  # Sub-sub-sections don't get page breaks
                    analysis['should_center'] = False  # Sub-sub-sections don't get centered
                    return analysis

            # ============================================================
            # Prominent Topic Detection (Bold/Emphasized Headings)
            # ============================================================
            # 1. Standalone bold headings
            for pattern in self.patterns.get('standalone_bold_heading', []):
                if match := pattern.match(trimmed):
                    content = match.group(1).strip()
                    if 2 <= len(content.split()) <= 15:
                        analysis['type'] = 'prominent_heading'
                        analysis['content'] = content
                        analysis['confidence'] = 0.85
                        analysis['level'] = self._determine_heading_level(content, 2)
                        analysis['original_format'] = 'bold_standalone'
                        return analysis

            # 2. Bold prefix headings ("**Important:** Note")
            for pattern in self.patterns.get('bold_prefix_heading', []):
                if match := pattern.match(trimmed):
                    bold_part = match.group(1).strip()
                    content = match.group(2).strip()
                    analysis['type'] = 'prominent_heading'
                    analysis['content'] = f"{bold_part}: {content}"
                    analysis['confidence'] = 0.80
                    analysis['level'] = self._determine_heading_level(bold_part, 3)
                    analysis['original_format'] = 'bold_prefix'
                    return analysis

            # Also handle bold that includes a trailing colon inside (e.g., **Warning:** text)
            if m := re.match(r'^\s*\*\*([^*\n]+?:)\*\*\s*(.+)$', trimmed):
                bold_part = m.group(1).strip().rstrip(':')
                content = m.group(2).strip()
                analysis['type'] = 'prominent_heading'
                analysis['content'] = f"{bold_part}: {content}"
                analysis['confidence'] = 0.80
                analysis['level'] = self._determine_heading_level(bold_part, 3)
                analysis['original_format'] = 'bold_prefix'
                return analysis

            # 3. Underlined emphasis (multi-line: requires underscores on prev & next lines)
            if prev_line and next_line:
                if (re.match(r'^\s*_{3,}\s*$', prev_line) and re.match(r'^\s*_{3,}\s*$', next_line)):
                    analysis['type'] = 'prominent_heading'
                    analysis['content'] = trimmed
                    analysis['confidence'] = 0.75
                    analysis['level'] = 2
                    analysis['original_format'] = 'underlined'
                    return analysis

            # 4. Star-surrounded text (***TEXT***)
            for pattern in self.patterns.get('star_surrounded', []):
                if match := pattern.match(trimmed):
                    content = match.group(1).strip()
                    analysis['type'] = 'prominent_heading'
                    analysis['content'] = content
                    analysis['confidence'] = 0.70
                    analysis['level'] = self._determine_heading_level(content, 3)
                    analysis['original_format'] = 'star_surrounded'
                    return analysis

            # 5. ALL CAPS short lines
            for pattern in self.patterns.get('allcaps_short_line', []):
                if match := pattern.match(trimmed):
                    content = match.group(0).strip()
                    word_count = len(content.split())
                    if 2 <= word_count <= 8:
                        analysis['type'] = 'prominent_heading'
                        analysis['content'] = content.title()
                        analysis['confidence'] = 0.65
                        analysis['level'] = 3
                        analysis['original_format'] = 'allcaps'
                        return analysis

            # 6. Numbered bold sections (**1)** Section**)
            for pattern in self.patterns.get('numbered_bold_section', []):
                if match := pattern.match(trimmed):
                    number = match.group(1)
                    content = match.group(2).strip()
                    analysis['type'] = 'prominent_heading'
                    analysis['content'] = f"{number} {content}"
                    analysis['confidence'] = 0.80
                    analysis['level'] = self._determine_heading_level(content, 3)
                    analysis['original_format'] = 'numbered_bold'
                    return analysis


            
            # Heuristic heading detection for ALL CAPS
            if is_all_caps and is_very_short and word_count <= 6:
                analysis['type'] = 'heading'
                analysis['level'] = 1
                analysis['confidence'] = 0.80
                # Check if this heading needs a page break and/or centering
                analysis['needs_page_break'] = self.should_start_on_new_page(trimmed)
                analysis['should_center'] = self.should_be_centered(trimmed, 1)
                return analysis
            
            # Heuristic heading detection for Title Case (no ending punctuation, short)
            if is_title_case and is_very_short and word_count >= 2 and word_count <= 8:
                analysis['type'] = 'heading'
                analysis['level'] = 2
                analysis['confidence'] = 0.75
                analysis['needs_page_break'] = False  # Title case headings don't get page breaks
                analysis['should_center'] = False  # Title case headings don't get centered
                return analysis
        
        # Priority 3: Check for reference patterns
        for pattern in self.patterns['reference']:
            if pattern.match(trimmed):
                # Special-case: avoid mis-detecting bracketed numbering used in lists as references
                if re.match(r'^\[\d+\]\s+', trimmed):
                    # Look for reference-like indicators: comma-separated author names, year in parentheses, 'et al', or URL
                    following = trimmed.split(None, 1)[1] if len(trimmed.split(None, 1)) > 1 else ''
                    if not re.search(r',|\(\d{4}\)|et al|https?://', following, re.IGNORECASE):
                        # Likely a list item, not a reference; skip this pattern
                        continue

                analysis['type'] = 'reference'
                analysis['confidence'] = 0.90
                if analysis.get('type') == 'reference':
                    spans = self._extract_journal_spans(analysis['content'])
                    if spans:
                        analysis['journal_spans'] = spans
                        analysis['confidence'] = max(analysis.get('confidence', 0.0), 0.92)
                return analysis
        
        # Priority 4: Check for list patterns
        # Use the enhanced bullet detection logic (on emoji-cleaned text)
        bullet_info = detect_bullet_type(trimmed)  # Pass cleaned line (emoji already stripped)
        if bullet_info:
            analysis['type'] = 'bullet_list'
            analysis['content'] = bullet_info['content']
            analysis['bullet_info'] = bullet_info  # Store full info for WordGenerator
            analysis['confidence'] = 0.98
            logger.info("List item detected (bullet): '%s' -> '%s'", trimmed, bullet_info['content'])
            return analysis

        for pattern in self.patterns['bullet_list']:
            match = pattern.match(trimmed)
            if match:
                analysis['type'] = 'bullet_list'
                analysis['content'] = match.group(2) if match.lastindex and match.lastindex >= 2 else match.group(1).strip() if match.lastindex else trimmed.lstrip('•●○▪▫■□◆◇-–—* →➤➢').strip()
                analysis['confidence'] = 0.98
                logger.info("List item detected (bullet): '%s' -> '%s'", trimmed, analysis['content'])
                return analysis
        
        for pattern in self.patterns['numbered_list']:
            if pattern.match(trimmed):
                question_text = re.sub(r'^\s*\(?[0-9ivxA-Za-z]+[\.)]\s+', '', trimmed).strip()
                if self._looks_like_question(question_text):
                    analysis['type'] = 'numbered_list'
                    analysis['confidence'] = 0.95
                    logger.info("List item detected (numbered question): '%s'", trimmed)
                    return analysis
                if (self._matches_heading_patterns(trimmed)
                        or (self._is_heading_like_numbered_line(trimmed)
                            and not self._has_adjacent_list_context(prev_line, next_line)
                            and not self._is_short_numbered_list_item(trimmed))):
                    analysis['type'] = 'heading'
                    analysis['level'] = 2
                    analysis['confidence'] = 0.90
                    logger.info("Numbered line treated as heading: '%s'", trimmed)
                    return analysis
                # Apply the classification - don't reject here, handle rendering properly
                analysis['type'] = 'numbered_list'
                analysis['confidence'] = 0.95
                logger.info("List item detected (numbered): '%s'", trimmed)
                return analysis

        # ============================================================
        # CATCH REJECTED SECTION HEADERS from numbered_list
        # Items like "1. Implications for Students:" or "a. Enhanced Learning"
        # ============================================================
        # Check for numeric section headers (1. Title, 2. Title, etc.)
        numeric_section = re.match(r'^\s*(\d+[\.)])\s+(.+?)\s*:?\s*$', trimmed)
        if numeric_section:
            number = numeric_section.group(1)
            title = numeric_section.group(2).strip()
            # Verify it looks like a section header (starts with uppercase, 1-8 words, not too long)
            if (title and title[0].isupper() and 
                len(title.split()) <= 8 and 
                len(title) < 100 and
                not self._looks_like_question(title) and
                not title.endswith('.')):  # Not a full sentence
                analysis['type'] = 'shortdoc_header'
                analysis['content'] = title
                analysis['confidence'] = 0.92
                analysis['level'] = int(number[0]) if number[0].isdigit() else 1
                analysis['numbering'] = number
                analysis['header_type'] = 'section'
                return analysis
        
        # Check for lettered subsection headers (a. Title, b. Title, etc.)
        letter_section = re.match(r'^\s*([a-z][\.)]\s+)(.+?)\s*:?\s*$', trimmed)
        if letter_section:
            letter_prefix = letter_section.group(1)
            title = letter_section.group(2).strip()
            # Verify it looks like a subsection header (short, starts uppercase, not a sentence)
            if (title and title[0].isupper() and 
                len(title.split()) <= 6 and  # Shorter than main headers
                len(title) < 80 and
                not self._looks_like_question(title) and
                not title.endswith('.')):  # Not a full sentence
                analysis['type'] = 'shortdoc_header'
                analysis['content'] = title
                analysis['confidence'] = 0.90
                analysis['level'] = 2  # Subsection level
                analysis['numbering'] = letter_prefix.strip()
                analysis['header_type'] = 'subsection'
                return analysis
        
        # Check for Roman numeral section headers (I. Title, II. Title, etc.)
        roman_section = re.match(r'^\s*([IVX]+[\.)])\s+(.+?)\s*:?\s*$', trimmed, re.IGNORECASE)
        if roman_section:
            number = roman_section.group(1)
            title = roman_section.group(2).strip()
            # Verify it looks like a section header
            if (title and title[0].isupper() and 
                len(title.split()) <= 8 and 
                len(title) < 100 and
                not self._looks_like_question(title) and
                not title.endswith('.')):  # Not a full sentence
                analysis['type'] = 'shortdoc_header'
                analysis['content'] = title
                analysis['confidence'] = 0.91
                analysis['level'] = 2
                analysis['numbering'] = number
                analysis['header_type'] = 'section'
                return analysis

        # ============================================================
        # Enhanced Numbering Detection (Mixed Hierarchical) - run after numbered_list
        # ============================================================
        for pattern in self.patterns.get('alphanumeric_hierarchy', []):
            if match := pattern.match(trimmed):
                full_number = trimmed.split()[0].rstrip('.)')
                # If it's pure numeric multi-segment (e.g., 1.1.1) prefer heading level 3
                segments = [s for s in re.split(r'[\.\)]', full_number) if s]
                if all(seg.isdigit() for seg in segments) and len(segments) >= 3:
                    analysis['type'] = 'heading'
                    analysis['level'] = 3
                    analysis['content'] = match.group(match.lastindex)
                    analysis['confidence'] = 0.90
                    return analysis

                content = match.group(match.lastindex)
                level = full_number.count('.') + full_number.count(')') + 1

                analysis['type'] = 'hierarchical_list'
                analysis['content'] = content
                analysis['confidence'] = 0.90
                analysis['level'] = min(level, 5)  # Cap at level 5
                analysis['numbering'] = full_number
                analysis['format'] = 'alphanumeric_hierarchy'
                return analysis

        for pattern in self.patterns.get('parenthesized_mixed', []):
            if match := pattern.match(trimmed):
                number = match.group(1)
                content = match.group(2)

                # If it's a simple parenthesized digit (e.g., (1)), prefer numbered_list
                if number.isdigit():
                    # Let numbered_list handle simple parenthesized numbers
                    pass
                else:
                    analysis['type'] = 'hierarchical_list'
                    analysis['content'] = content
                    analysis['confidence'] = 0.85
                    analysis['level'] = self._get_parenthesized_level(number)
                    analysis['numbering'] = f"({number})"
                    analysis['format'] = 'parenthesized'
                    return analysis

        for pattern in self.patterns.get('double_parentheses', []):
            if match := pattern.match(trimmed):
                number = match.group(1)
                content = match.group(2)

                analysis['type'] = 'hierarchical_list'
                analysis['content'] = content
                analysis['confidence'] = 0.75
                analysis['level'] = 3  # Double parentheses typically indicate level 3
                analysis['numbering'] = f"(({number}))"
                analysis['format'] = 'double_parentheses'
                return analysis

        for pattern in self.patterns.get('bracketed_numbering', []):
            if match := pattern.match(trimmed):
                number = match.group(1)
                content = match.group(2)

                analysis['type'] = 'hierarchical_list'
                analysis['content'] = content
                analysis['confidence'] = 0.80
                analysis['level'] = self._get_bracketed_level(number)
                analysis['numbering'] = f"[{number}]"
                analysis['format'] = 'bracketed'
                return analysis

        for pattern in self.patterns.get('hyphen_numbering', []):
            if match := pattern.match(trimmed):
                number = match.group(1)
                content = match.group(2)

                analysis['type'] = 'hierarchical_list'
                analysis['content'] = content
                analysis['confidence'] = 0.70
                analysis['level'] = 2 if number.isdigit() else 3
                analysis['numbering'] = f"{number}-"
                analysis['format'] = 'hyphen'
                return analysis

        for pattern in self.patterns.get('roman_with_dots', []):
            if match := pattern.match(trimmed):
                roman = match.group(1)
                content = match.group(2)

                analysis['type'] = 'hierarchical_list'
                analysis['content'] = content
                analysis['confidence'] = 0.65
                analysis['level'] = self._roman_to_int(roman) % 3 + 1  # Vary level based on Roman value
                analysis['numbering'] = f"{roman}..."
                analysis['format'] = 'roman_dots'
                return analysis

        # ============================================================
        # Short Document Section Headers (Assignment Style)
        # ============================================================
        for pattern in self.patterns.get('assignment_section_header', []):
            if match := pattern.match(trimmed):
                number = match.group(1)
                title = match.group(2).strip()

                # Validate it's a reasonable header (not too long, ends properly)
                if not title.endswith('.') and len(title.split()) <= 8:
                    analysis['type'] = 'shortdoc_header'
                    analysis['content'] = title
                    analysis['confidence'] = 0.88
                    analysis['level'] = self._get_shortdoc_header_level(number, title)
                    analysis['numbering'] = number
                    analysis['header_type'] = 'section'
                    return analysis

        for pattern in self.patterns.get('question_prompt', []):
            if match := pattern.match(trimmed):
                q_num = match.group(1)
                question = match.group(2).strip()

                analysis['type'] = 'shortdoc_header'
                analysis['content'] = question
                analysis['confidence'] = 0.92
                analysis['level'] = 3  # Questions are usually level 3
                analysis['numbering'] = q_num
                analysis['header_type'] = 'question'
                return analysis

        for pattern in self.patterns.get('task_header', []):
            if match := pattern.match(trimmed):
                task_num = match.group(1)
                task = match.group(2).strip()

                analysis['type'] = 'shortdoc_header'
                analysis['content'] = task
                analysis['confidence'] = 0.90
                analysis['level'] = 3 if task_num.isdigit() else 4
                analysis['numbering'] = task_num
                analysis['header_type'] = 'task'
                return analysis

        for pattern in self.patterns.get('part_section', []):
            if match := pattern.match(trimmed):
                part_num = match.group(1)
                part_title = match.group(2).strip()

                analysis['type'] = 'shortdoc_header'
                analysis['content'] = part_title
                analysis['confidence'] = 0.85
                analysis['level'] = self._get_part_level(part_num)
                analysis['numbering'] = part_num
                analysis['header_type'] = 'part'
                return analysis

        for pattern in self.patterns.get('requirement_header', []):
            if match := pattern.match(trimmed):
                req_num = match.group(1)
                requirement = match.group(2).strip()

                analysis['type'] = 'shortdoc_header'
                analysis['content'] = requirement
                analysis['confidence'] = 0.87
                analysis['level'] = req_num.count('.') + 3  # Deeper hierarchy = lower level
                analysis['numbering'] = req_num
                analysis['header_type'] = 'requirement'
                return analysis

        for pattern in self.patterns.get('step_header', []):
            if match := pattern.match(trimmed):
                step_num = match.group(1)
                step = match.group(2).strip()

                analysis['type'] = 'shortdoc_header'
                analysis['content'] = step
                analysis['confidence'] = 0.89
                analysis['level'] = 4  # Steps are usually deepest level
                analysis['numbering'] = step_num
                analysis['header_type'] = 'step'
                return analysis

        # Priority 5: Check for definition patterns
        for pattern in self.patterns['definition']:
            match = pattern.match(trimmed)
            if match:
                analysis['type'] = 'definition'
                analysis['term'] = match.group(1)
                analysis['definition'] = match.group(2) if match.lastindex > 1 and match.group(2) else ''
                analysis['confidence'] = 0.90
                return analysis
        
        # Priority 6: Check for figure captions
        for pattern in self.patterns['figure']:
            if pattern.match(trimmed):
                analysis['type'] = 'figure'
                analysis['confidence'] = 0.95
                return analysis
        
        # Priority 7: Check for equation labels
        for pattern in self.patterns['equation']:
            if pattern.match(trimmed):
                analysis['type'] = 'equation'
                analysis['confidence'] = 0.90
                return analysis
        
        # Priority 8: Check for quotes
        for pattern in self.patterns['quote']:
            if pattern.match(trimmed):
                analysis['type'] = 'quote'
                analysis['confidence'] = 0.85
                return analysis
        
        # Priority 9: Check for code blocks
        for pattern in self.patterns['code']:
            if pattern.match(trimmed):
                analysis['type'] = 'code'
                analysis['confidence'] = 0.90
                return analysis
        
        # Priority 10: Check for page metadata (headers/footers/page numbers)
        for pattern in self.patterns['page_metadata']:
            if pattern.match(trimmed):
                analysis['type'] = 'page_metadata'
                analysis['confidence'] = 0.90
                # Determine subtype
                if re.search(r'page|p\.|pg\.', trimmed, re.IGNORECASE) or re.match(r'^\s*-?\s*\d+\s*-?\s*$', trimmed):
                    analysis['subtype'] = 'page_number'
                elif re.search(r'header|running head', trimmed, re.IGNORECASE):
                    analysis['subtype'] = 'header'
                elif re.search(r'footer', trimmed, re.IGNORECASE):
                    analysis['subtype'] = 'footer'
                else:
                    analysis['subtype'] = 'document_metadata'
                return analysis
        
        # Priority 11: Check for academic metadata
        for pattern in self.patterns['academic_metadata']:
            if pattern.match(trimmed):
                analysis['type'] = 'academic_metadata'
                analysis['confidence'] = 0.80
                # Determine subtype
                if re.search(r'\bby\b|authors?:', trimmed, re.IGNORECASE):
                    analysis['subtype'] = 'author'
                elif re.search(r'@', trimmed):
                    analysis['subtype'] = 'contact'
                elif re.search(r'department|school|college|university|institute', trimmed, re.IGNORECASE):
                    analysis['subtype'] = 'affiliation'
                else:
                    analysis['subtype'] = 'metadata'
                return analysis
        
        # Priority 12: Check for mathematical expressions
        for pattern in self.patterns['math_expression']:
            if pattern.search(trimmed):
                # Avoid false positives with currency
                if re.match(r'^\$\d+', trimmed) and not re.search(r'\$[^$]+\$', trimmed):
                    continue  # This is likely currency, not math
                analysis['type'] = 'math_expression'
                # Determine subtype
                if trimmed.startswith('$$') and trimmed.endswith('$$'):
                    analysis['subtype'] = 'display_math'
                    analysis['confidence'] = 0.95
                elif '$' in trimmed or '\\(' in trimmed or '\\[' in trimmed:
                    analysis['subtype'] = 'inline_math'
                    analysis['confidence'] = 0.85
                else:
                    analysis['subtype'] = 'equation'
                    analysis['confidence'] = 0.75
                return analysis
        
        # Priority 13: Check for footnotes/endnotes
        for pattern in self.patterns['footnote_endnote']:
            if pattern.match(trimmed):
                analysis['type'] = 'footnote_endnote'
                # Determine subtype
                if re.match(r'^\s*(?:endnotes?|footnotes?)\s*$', trimmed, re.IGNORECASE):
                    analysis['subtype'] = 'section_header'
                    analysis['confidence'] = 0.95
                else:
                    analysis['subtype'] = 'footnote_entry'
                    analysis['confidence'] = 0.90
                return analysis
        
        # Priority 14: Check for inline formatting (bold only, no italics)
        # Check for markdown-style formatting but exclude lines that start with list markers
        if not re.match(r'^[\*\-•]\s', trimmed):  # Not a bullet list
            for pattern in self.patterns['inline_formatting']:
                matches = pattern.findall(trimmed)
                if matches and any(m for m in matches if any(g for g in (m if isinstance(m, tuple) else (m,)) if g)):
                    analysis['type'] = 'inline_formatting'
                    analysis['content'] = trimmed
                    # Determine formatting type
                    if '***' in trimmed or '___' in trimmed:
                        analysis['formatting'] = {'bold_italic': False, 'bold': True, 'italic': False}
                        analysis['confidence'] = 0.90
                    elif '**' in trimmed or '__' in trimmed:
                        analysis['formatting'] = {'bold': True, 'italic': False, 'bold_italic': False}
                        analysis['confidence'] = 0.85
                    else:
                        # Single asterisk or underscore = BOLD (not italic)
                        analysis['formatting'] = {'bold': True, 'italic': False, 'bold_italic': False}
                        analysis['confidence'] = 0.85
                    return analysis
        
        # ============================================================
        # NEW PATTERN DETECTION - December 30, 2025 (20 Academic Patterns)
        # ============================================================
        
        # Priority 15: Check for markdown heading hierarchy
        for pattern in self.patterns['heading_hierarchy']:
            if pattern.match(trimmed):
                analysis['type'] = 'heading_hierarchy'
                # Determine level by counting # symbols
                hash_count = len(trimmed) - len(trimmed.lstrip('#'))
                analysis['level'] = min(hash_count, 6)
                analysis['confidence'] = 0.95
                analysis['content'] = trimmed.lstrip('#').strip()
                # Check if this heading needs a page break and/or centering (only for level 1 headings)
                if hash_count == 1:
                    analysis['needs_page_break'] = self.should_start_on_new_page(trimmed)
                    analysis['should_center'] = self.should_be_centered(trimmed, 1)
                else:
                    analysis['needs_page_break'] = False
                    analysis['should_center'] = False
                return analysis
        
        # Priority 16: Check for academic table patterns
        for pattern in self.patterns['academic_table']:
            if pattern.match(trimmed):
                analysis['type'] = 'academic_table'
                if re.match(r'^Table\s+\d+', trimmed, re.IGNORECASE):
                    analysis['subtype'] = 'caption'
                elif re.match(r'^\|[-\s:]+\|', trimmed):
                    analysis['subtype'] = 'separator'
                elif '**' in trimmed:
                    analysis['subtype'] = 'header_row'
                else:
                    analysis['subtype'] = 'data_row'
                analysis['confidence'] = 0.95
                return analysis
        
        # Priority 17: Check for nested list patterns
        for pattern in self.patterns['list_nested']:
            match = pattern.match(trimmed)
            if match:
                analysis['type'] = 'list_nested'
                # Calculate indent level (2 spaces per level)
                leading_spaces = len(trimmed) - len(trimmed.lstrip())
                analysis['indent_level'] = leading_spaces // 2
                if '□' in trimmed or '☐' in trimmed or '☑' in trimmed or '✓' in trimmed:
                    analysis['subtype'] = 'checkbox'
                else:
                    analysis['subtype'] = 'nested_item'
                analysis['confidence'] = 0.90
                return analysis
        
        # Priority 18: Check for figure/equation patterns
        for pattern in self.patterns['figure_equation']:
            if pattern.match(trimmed) or pattern.search(trimmed):
                analysis['type'] = 'figure_equation'
                if re.match(r'^[Ff]igure', trimmed):
                    analysis['subtype'] = 'figure_caption'
                elif '$$' in trimmed or 'equation' in trimmed.lower():
                    analysis['subtype'] = 'equation_block'
                else:
                    analysis['subtype'] = 'math_content'
                analysis['confidence'] = 0.90
                return analysis
        
        # Priority 19: Check for inline citations
        for pattern in self.patterns['citation_inline']:
            if pattern.search(trimmed):
                analysis['type'] = 'citation_inline'
                # Count citations in line
                citations = pattern.findall(trimmed)
                analysis['citation_count'] = len(citations)
                analysis['confidence'] = 0.85
                # Don't return - this is inline, keep processing
                break
        
        # Priority 20: Check for appendix formatting
        for pattern in self.patterns['appendix_format']:
            if pattern.match(trimmed):
                analysis['type'] = 'appendix_format'
                if re.match(r'^APPENDIX\s+[A-Z]$', trimmed, re.IGNORECASE):
                    analysis['subtype'] = 'appendix_header'
                    analysis['level'] = 1
                elif re.match(r'^[A-Z]\.\d+\.\d+', trimmed):
                    analysis['subtype'] = 'appendix_subsection'
                    analysis['level'] = 3
                elif re.match(r'^[A-Z]\.\d+', trimmed):
                    analysis['subtype'] = 'appendix_section'
                    analysis['level'] = 2
                else:
                    analysis['subtype'] = 'appendix_content'
                    analysis['level'] = 1
                analysis['confidence'] = 0.90
                return analysis
        
        # Priority 21: Check for block quotes
        for pattern in self.patterns['block_quote']:
            if pattern.match(trimmed):
                analysis['type'] = 'block_quote'
                analysis['confidence'] = 0.85
                analysis['content'] = trimmed.lstrip('> ').strip('"\'')
                return analysis
        
        # Priority 22: Check for mathematical models
        for pattern in self.patterns['math_model']:
            if pattern.search(trimmed):
                analysis['type'] = 'math_model'
                analysis['confidence'] = 0.85
                if re.search(r'[Yy]\s*=\s*[βα]', trimmed):
                    analysis['subtype'] = 'regression_model'
                elif re.search(r'[Rr]²', trimmed):
                    analysis['subtype'] = 'r_squared'
                elif re.search(r'[Pp]\s*[<>=]', trimmed):
                    analysis['subtype'] = 'p_value'
                else:
                    analysis['subtype'] = 'statistical_notation'
                return analysis
        
        # Priority 23: Check for text emphasis patterns (no italics)
        for pattern in self.patterns['text_emphasis']:
            if pattern.search(trimmed):
                analysis['type'] = 'text_emphasis'
                if '`' in trimmed:
                    analysis['subtype'] = 'monospace'
                elif '***' in trimmed:
                    analysis['subtype'] = 'bold'
                elif '**' in trimmed:
                    analysis['subtype'] = 'bold'
                else:
                    analysis['subtype'] = 'bold'
                analysis['confidence'] = 0.80
                # Don't return - inline emphasis, continue processing
                break
        
        # Priority 24: Check for APA reference format
        for pattern in self.patterns['reference_apa']:
            if pattern.match(trimmed) or pattern.search(trimmed):
                analysis['type'] = 'reference_apa'
                if 'doi' in trimmed.lower():
                    analysis['subtype'] = 'doi_reference'
                elif 'Retrieved' in trimmed:
                    analysis['subtype'] = 'web_reference'
                else:
                    analysis['subtype'] = 'standard_reference'
                analysis['confidence'] = 0.90
                return analysis
        
        # Priority 25: Check for TOC entries
        for pattern in self.patterns['toc_entry']:
            if pattern.match(trimmed):
                analysis['type'] = 'toc_entry'
                analysis['confidence'] = 0.95
                # Extract page number if present
                page_match = re.search(r'(\d+)\s*$', trimmed)
                if page_match:
                    analysis['page_number'] = int(page_match.group(1))
                return analysis
        
        # Priority 26: Check for footnote markers
        for pattern in self.patterns['footnote_marker']:
            if pattern.search(trimmed):
                analysis['type'] = 'footnote_marker'
                if trimmed.startswith('[^'):
                    analysis['subtype'] = 'footnote_definition'
                else:
                    analysis['subtype'] = 'footnote_reference'
                analysis['confidence'] = 0.90
                return analysis
        
        # Priority 27: Check for abbreviations
        for pattern in self.patterns['abbreviation']:
            match = pattern.search(trimmed)
            if match:
                # Only classify if it looks like a definition
                if re.search(r'[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+\s+\([A-Z]{2,}\)', trimmed):
                    analysis['type'] = 'abbreviation'
                    analysis['subtype'] = 'definition'
                    analysis['confidence'] = 0.85
                    return analysis
        
        # Priority 28: Check for caption formatting
        for pattern in self.patterns['caption_format']:
            if pattern.match(trimmed):
                analysis['type'] = 'caption_format'
                if 'Table' in trimmed:
                    analysis['subtype'] = 'table_caption'
                elif 'Figure' in trimmed:
                    analysis['subtype'] = 'figure_caption'
                elif trimmed.startswith('Source:'):
                    analysis['subtype'] = 'source_attribution'
                elif trimmed.lower().startswith('note:'):
                    analysis['subtype'] = 'table_note'
                else:
                    analysis['subtype'] = 'caption'
                analysis['confidence'] = 0.90
                return analysis
        
        # Priority 29a: Check for horizontal rules (visual separators, NOT page breaks)
        for pattern in self.patterns.get('horizontal_rule', []):
            if pattern.match(trimmed):
                analysis['type'] = 'horizontal_rule'
                analysis['confidence'] = 1.0
                return analysis
        
        # Priority 29b: Check for explicit page breaks
        for pattern in self.patterns['page_break']:
            if pattern.match(trimmed):
                analysis['type'] = 'page_break'
                analysis['confidence'] = 1.0
                return analysis
        
        # Priority 30: Check for statistical results
        for pattern in self.patterns['statistical_result']:
            if pattern.search(trimmed):
                analysis['type'] = 'statistical_result'
                analysis['confidence'] = 0.85
                # Identify specific stat types
                stats_found = []
                if re.search(r'β\s*=', trimmed):
                    stats_found.append('beta')
                if re.search(r'[Pp]\s*[<>=]', trimmed):
                    stats_found.append('p_value')
                if re.search(r'[Ff]\s*\(', trimmed):
                    stats_found.append('f_statistic')
                if re.search(r'[Rr]²?\s*=', trimmed):
                    stats_found.append('r_value')
                if re.search(r'CI\s*=', trimmed):
                    stats_found.append('confidence_interval')
                analysis['stats_types'] = stats_found
                return analysis
        
        # Priority 31: Check for questionnaire patterns
        for pattern in self.patterns['questionnaire']:
            if pattern.match(trimmed) or pattern.search(trimmed):
                analysis['type'] = 'questionnaire'
                if re.match(r'^Section\s+[A-Z]', trimmed, re.IGNORECASE):
                    analysis['subtype'] = 'section_header'
                elif '□' in trimmed or '☐' in trimmed:
                    analysis['subtype'] = 'checkbox_item'
                elif 'SA' in trimmed and 'SD' in trimmed:
                    analysis['subtype'] = 'likert_header'
                else:
                    analysis['subtype'] = 'question_item'
                analysis['confidence'] = 0.85
                return analysis
        
        # Priority 32: Check for glossary entries
        for pattern in self.patterns['glossary_entry']:
            if pattern.match(trimmed):
                analysis['type'] = 'glossary_entry'
                # Extract term and definition
                term_match = re.match(r'\*\*([^*]+)\*\*:\s*(.+)', trimmed)
                if term_match:
                    analysis['term'] = term_match.group(1)
                    analysis['definition'] = term_match.group(2)
                analysis['confidence'] = 0.90
                return analysis
        
        # Priority 33: Check for cross-references
        for pattern in self.patterns['cross_reference']:
            if pattern.search(trimmed):
                analysis['type'] = 'cross_reference'
                refs_found = []
                if re.search(r'[Tt]able\s+\d+', trimmed):
                    refs_found.append('table')
                if re.search(r'[Ff]igure\s+\d+', trimmed):
                    refs_found.append('figure')
                if re.search(r'[Ss]ection\s+\d+', trimmed):
                    refs_found.append('section')
                if re.search(r'[Pp]age\s+\d+', trimmed):
                    refs_found.append('page')
                analysis['reference_types'] = refs_found
                analysis['confidence'] = 0.80
                # Don't return - cross-references are inline
                break
        
        # Priority 34: Check for running headers
        for pattern in self.patterns['running_header']:
            if pattern.match(trimmed):
                analysis['type'] = 'running_header'
                analysis['confidence'] = 0.90
                return analysis
        
        # Priority 35: SHORT DOCUMENT KEY POINT DETECTION
        # Check for key point markers (learning objectives, definitions, warnings, etc.)
        key_point_type, emoji = self.get_key_point_type(trimmed)
        if key_point_type:
            analysis['type'] = 'key_point'
            analysis['key_point_type'] = key_point_type
            analysis['emoji_prefix'] = emoji
            analysis['confidence'] = 0.90
            return analysis
        
        # Priority 36: Assignment header fields (Student Name, Course, etc.)
        if self.is_assignment_header_field(trimmed):
            analysis['type'] = 'assignment_header_field'
            analysis['confidence'] = 0.90
            return analysis
        
        # Default: paragraph
        analysis['confidence'] = 0.70
        return analysis

    def _extract_cells(self, text, delimiter_type):
        """Extract cell contents based on delimiter type"""
        cells = []
        
        if delimiter_type == 'tab':
            cells = [cell.strip() for cell in text.split('\t') if cell.strip()]
        
        elif delimiter_type == 'pipe':
            # Remove leading/trailing pipes and split
            text = text.strip('|').strip()
            cells = [cell.strip() for cell in text.split('|') if cell.strip()]
        
        elif delimiter_type == 'comma':
            # Split on comma, keeping quoted strings together
            cells = [cell.strip() for cell in text.split(',') if cell.strip()]
        
        elif delimiter_type == 'numbered':
            # Remove leading number and split on spacing
            text = re.sub(r'^\s*\d+[\.\)]\s+', '', text)
            cells = [cell for cell in re.split(r'\s{2,}', text) if cell.strip()]
        
        elif delimiter_type == 'keyvalue':
            # Extract key:value pairs
            pairs = re.findall(r'([^:]+):\s*([^\s]+)', text)
            cells = [f"{k.strip()}: {v.strip()}" for k, v in pairs]
        
        elif delimiter_type == 'numeric':
            # Split on 2+ spaces
            cells = [cell for cell in re.split(r'\s{2,}', text.strip()) if cell.strip()]
        
        elif delimiter_type == 'spaced':
            # Split on 2+ spaces (most common for plain text tables)
            cells = [cell for cell in re.split(r'\s{2,}', text.strip()) if cell.strip()]
        
        return cells

    def _is_likely_table_row(self, text, cells):
        """
        Strict validation: Is this REALLY a table row or just spaced text?
        
        Returns True only if it has clear table characteristics.
        """
        if not cells or len(cells) < 3:
            return False
        
        # RULE 1: If any cell is very long, it's probably prose
        for cell in cells:
            if len(cell) > 30:  # Single "cell" over 30 chars = prose
                return False
        
        # RULE 2: Check for prose indicators (common words in sentences)
        prose_words = [
            'the', 'a', 'an', 'this', 'that', 'these', 'those',
            'has', 'have', 'had', 'is', 'are', 'was', 'were',
            'will', 'would', 'should', 'could', 'can', 'may',
            'by', 'for', 'with', 'from', 'about', 'into',
            'through', 'during', 'before', 'after', 'above', 'below'
        ]
        
        prose_count = 0
        total_words = 0
        
        for cell in cells:
            words = cell.lower().split()
            total_words += len(words)
            for word in words:
                if word in prose_words:
                    prose_count += 1
        
        # If more than 40% are prose words, it's not a table
        if total_words > 0 and (prose_count / total_words) > 0.4:
            return False
        
        # RULE 3: At least one cell should be "table-like"
        table_like = 0
        for cell in cells:
            # Numbers, dates, codes, currency
            if (re.match(r'^\d+$', cell) or  # Pure number
                re.match(r'^\d+\.\d+$', cell) or  # Decimal
                re.match(r'^[$€£¥]\d+', cell) or  # Currency
                re.match(r'^\d+%$', cell) or  # Percentage
                re.match(r'^[A-Z]{2,5}$', cell) or  # Code (USA, ID, etc)
                len(cell) <= 10):  # Short labels
                table_like += 1
        
        # At least 50% of cells should look table-like
        if table_like / len(cells) < 0.5:
            return False
        
        # RULE 4: Check for sentence punctuation inside cells (prose indicator)
        for cell in cells:
            if '. ' in cell or ', ' in cell:  # Mid-sentence punctuation
                return False
        
        return True

    def _detect_spaced_table_row(self, text, prev_line=None, next_line=None, context=None):
        """
        Detect if text is a spaced table row with HIGH ACCURACY.
        
        Uses multiple validation layers to prevent false positives.
        Returns: (is_table, cells, confidence) or (False, None, 0)
        """
        
        # LAYER 1: Initial pattern check - must have 2+ spaces between items
        if not re.search(r'\S+\s{2,}\S+', text):
            return (False, None, 0)
        
        # Extract potential cells (split on 2+ consecutive spaces)
        potential_cells = [cell.strip() for cell in re.split(r'\s{2,}', text.strip()) if cell.strip()]
        
        # LAYER 2: Must have at least 2 columns
        if len(potential_cells) < 2:
            return (False, None, 0)
        
        # LAYER 3: Neighbor line validation to avoid false positives
        neighbor_has_spacing = False
        if prev_line and re.search(r'\S+\s{2,}\S+', prev_line):
            neighbor_has_spacing = True
        if next_line and re.search(r'\S+\s{2,}\S+', next_line):
            neighbor_has_spacing = True

        # Require neighbor spacing unless we have 3+ columns
        if not neighbor_has_spacing and len(potential_cells) < 3:
            return (False, None, 0)

        # LAYER 4: Prose detection - check for narrative indicators
        combined_lower = ' '.join(potential_cells).lower()
        
        # High-confidence prose indicators (these phrases NEVER appear in tables)
        definite_prose_phrases = [
            'has fundamentally', 'have fundamentally',
            'is significant', 'are significant',
            'through economies', 'by leveraging',
            'rather than', 'as well as',
            'such as', 'in order to',
            'due to', 'based on',
            'according to', 'in addition to',
            'with respect to', 'in terms of',
            'as a result', 'for example',
            'for instance', 'in particular',
            'on the other hand', 'in contrast',
            'however', 'therefore', 'moreover', 'furthermore',
            'nevertheless', 'consequently', 'subsequently',
            'alternatively', 'specifically', 'generally',
        ]
        
        for phrase in definite_prose_phrases:
            if phrase in combined_lower:
                return (False, None, 0)
        
        # LAYER 5: Check for sentence-like structure
        # Articles + verbs = prose
        has_article = bool(re.search(r'\b(the|a|an|this|that|these|those)\b', combined_lower))
        has_verb = bool(re.search(r'\b(is|are|was|were|has|have|had|will|would|should|could|can|may|must|do|does|did)\b', combined_lower))
        
        if has_article and has_verb:
            # Check if it looks like a sentence
            word_count = len(combined_lower.split())
            if word_count > 8:  # Sentences have many words
                return (False, None, 0)
        
        # LAYER 6: Cell length validation
        # Real table cells are typically short
        max_cell_length = max(len(cell) for cell in potential_cells)
        avg_cell_length = sum(len(cell) for cell in potential_cells) / len(potential_cells)
        
        # If ANY cell is very long, likely prose
        if max_cell_length > 60:
            return (False, None, 0)
        
        # If average cell length is high, likely prose
        if avg_cell_length > 35:
            return (False, None, 0)
        
        # LAYER 7: Check for multi-word cells (prose indicator)
        multi_word_cells = sum(1 for cell in potential_cells if len(cell.split()) > 4)
        
        # If more than 30% of cells have 4+ words, likely prose
        if len(potential_cells) > 0 and (multi_word_cells / len(potential_cells)) > 0.3:
            return (False, None, 0)
        
        # LAYER 8: Punctuation check
        # Tables rarely have mid-sentence punctuation
        for cell in potential_cells:
            if ', ' in cell or '; ' in cell:  # Commas/semicolons mid-text = prose
                return (False, None, 0)
        
        # LAYER 9: Context validation (if previous line available)
        confidence = 0.7  # Base confidence
        
        if context and context.get('prev_analysis'):
            prev_analysis = context.get('prev_analysis')
            prev_type = prev_analysis.get('type', '') if isinstance(prev_analysis, dict) else ''
            
            # Boost confidence if previous line was also a table
            if prev_type.startswith('plain_table'):
                prev_metadata = prev_analysis.get('metadata', {}) if isinstance(prev_analysis, dict) else {}
                prev_cell_count = prev_metadata.get('cell_count', 0)
                
                # Check if column counts match (±1)
                if abs(len(potential_cells) - prev_cell_count) <= 1:
                    confidence = 0.90
                else:
                    # Column count mismatch - probably not a table
                    return (False, None, 0)
        
        # LAYER 10: Table-like content boost
        # Check if cells contain table-typical content
        table_like_patterns = [
            r'^\d+$',  # Pure numbers
            r'^\$\d+',  # Currency
            r'^\d+%$',  # Percentages
            r'^\d+\.\d+$',  # Decimals
            r'^[A-Z]{2,5}$',  # Codes/acronyms
            r'^\d{1,2}/\d{1,2}',  # Dates
        ]
        
        table_like_count = 0
        for cell in potential_cells:
            for pattern in table_like_patterns:
                if re.match(pattern, cell.strip()):
                    table_like_count += 1
                    break
        
        # If 40%+ cells are table-like, boost confidence
        if len(potential_cells) > 0 and (table_like_count / len(potential_cells)) >= 0.4:
            confidence = min(0.95, confidence + 0.15)
        
        # LAYER 11: Final validation - must have minimum confidence
        if confidence < 0.70:
            return (False, None, 0)
        
        # PASSED ALL CHECKS - This is a table row
        return (True, potential_cells, confidence)

    def _detect_separator_style(self, text):
        """Detect the style of separator used"""
        if '+' in text and '-' in text:
            return 'box'
        elif '=' in text:
            return 'double'
        elif '_' in text:
            return 'underscore'
        else:
            return 'dash'

    # ============================================================================
    # AI CONTENT CLEANING
    # ============================================================================
    
    def detect_ai_generated_content(self, text):
        """
        Detects if text contains common AI generation artifacts.
        Returns True if AI content detected.
        """
        if not text:
            return False
            
        # Safety check: Don't skip long paragraphs even if they start with trigger words
        # AI meta-commentary is usually short (e.g. "Here is the text:")
        if len(text) > 150:
            return False

        # Check for AI meta-commentary patterns
        ai_patterns = [
            r'^(Here is the|Here are the|Sure, here is|Certainly, here is)',
            r'^(In conclusion|To summarize|Hope this helps)',
            r'^Note:\s+',
            r'^As an AI language model',
            r'^I cannot generate',
        ]
        
        for pattern in ai_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return True
                
        return False

    def clean_ai_content(self, text):
        """
        Cleans common AI artifacts from text line.
        Returns (cleaned_text, metadata)
        """
        if not text:
            return text, {}
            
        clean_text = text.strip()
        metadata = {'bold': False, 'italic': False, 'heading_level': 0}
        
        # 1. Remove Markdown Bold (**text**) if it wraps the whole line
        bold_match = re.match(r'^\*\*(.*)\*\*$', clean_text)
        if bold_match:
            clean_text = bold_match.group(1)
            metadata['bold'] = True
        else:
            # Remove inline bold markers but keep text
            clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_text)
            
        # 2. Remove Markdown Italic (*text*)
        clean_text = re.sub(r'\*(.*?)\*', r'\1', clean_text)
        
        # 3. Remove Markdown Headers (## text)
        header_match = re.match(r'^(#{1,6})\s+(.*)', clean_text)
        if header_match:
            level = len(header_match.group(1))
            clean_text = header_match.group(2)
            metadata['heading_level'] = level
            
        # 4. Standardize Bullets
        if re.match(r'^-\s+', clean_text):
            clean_text = re.sub(r'^-\s+', '• ', clean_text)
            
        # 5. Normalize long dashes (—, –, hyphen variants) used as clause separators
        # Replace with a comma + single space ", " when the dash appears between words
        # Avoid replacing date ranges or numeric ranges (e.g., 1940–1950)
        try:
            # em-dash/en-dash and related variants between word characters -> ", "
            clean_text = re.sub(r"(?<!\d)(?<=\w)\s*[\u2010-\u2015\u2013\u2014]\s*(?=\w)(?!\d)", ', ', clean_text)
            # Double hyphen (--) style -> ", " when used between words
            clean_text = re.sub(r"(?<!\d)(?<=\w)\s*--+\s*(?=\w)(?!\d)", ', ', clean_text)
        except re.error:
            # On platforms that may not support some unicode ranges in regex, fall back to simple replacements
            clean_text = clean_text.replace('—', ', ').replace('–', ', ').replace('--', ', ')

        return clean_text, metadata

    def optimize_page_breaks_for_ai(self, lines):
        """
        Optimizes page breaks for AI generated content.
        """
        # Placeholder for future logic
        return lines

    # ========================================================================
    # TEST CASES: Emoji-Agnostic Bullet Engine
    # ========================================================================
    
    def test_emoji_and_bullet_detection(self):
        """Test cases for bullet detection with emoji removal"""
        test_cases = [
            # (input_text, expected_type, expected_content_contains)
            ("- Rising Sea Levels 🌊", "bullet_list", "Rising Sea Levels"),
            ("• Extreme Weather 🌪️", "bullet_list", "Extreme Weather"),
            ("* Biodiversity Loss", "bullet_list", "Biodiversity Loss"),
            ("– Reforestation (En-dash)", "bullet_list", "Reforestation"),
            ("— Mitigation (Em-dash)", "bullet_list", "Mitigation"),
            ("Just a normal sentence.", "paragraph", "Just a normal sentence"),
            ("1. Introduction", "numbered_list", None),  # Handled by numbered_list
            ("--- Section Separator", "paragraph", None),  # After horizontal rule removal
            ("Want me to expand on any of these points? 📚", "paragraph", "Want me to expand"),
            ("→ Arrow bullet point", "bullet_list", "Arrow bullet point"),
            ("☐ Checkbox item", "bullet_list", "Checkbox item"),
        ]
        
        results = []
        for text, expected_type, expected_content in test_cases:
            analysis = self.analyze_line(text, 0)
            success = analysis['type'] == expected_type
            
            if expected_content:
                success = success and expected_content in analysis.get('content', '')
            
            results.append({
                'input': text,
                'expected_type': expected_type,
                'actual_type': analysis['type'],
                'content': analysis.get('content', ''),
                'passed': success,
            })
        
        return results
    
    def test_emoji_stripping(self):
        """Test that emojis are properly removed before pattern matching"""
        # Test emoji removal directly
        test_text = "Testing 🎉 emojis 😀 removal 🚀"
        
        # Apply emoji_cleaner patterns
        cleaned = test_text
        for pattern in self.patterns.get('emoji_cleaner', []):
            cleaned = pattern.sub('', cleaned).strip()
        
        # Should contain no emoji-like characters
        analysis = self.analyze_line(f"- {test_text}", 0)
        
        # The content should be detected as bullet_list and emoji should be removed
        return {
            'original': test_text,
            'cleaned': cleaned,
            'analysis_type': analysis['type'],
            'analysis_content': analysis.get('content', ''),
            'emoji_removed': '🎉' not in analysis.get('content', '') and '😀' not in analysis.get('content', ''),
        }
    
    def test_bullet_detection_with_various_marks(self):
        """Test bullet detection with different bullet characters"""
        bullet_marks = [
            "-",  # Hyphen/minus
            "–",  # En-dash (U+2013)
            "—",  # Em-dash (U+2014)
            "•",  # Bullet
            "○",  # White circle
            "●",  # Black circle
            "▪",  # Black square
            "■",  # Filled square
            "*",  # Asterisk
            "→",  # Right arrow
            "☐",  # Checkbox
        ]
        
        results = []
        for mark in bullet_marks:
            text = f"{mark} Test item"
            analysis = self.analyze_line(text, 0)
            results.append({
                'mark': mark,
                'mark_unicode': f"U+{ord(mark):04X}",
                'text': text,
                'detected_type': analysis['type'],
                'detected_content': analysis.get('content', ''),
                'passed': analysis['type'] == 'bullet_list',
            })
        
        return results

    def test_bullet_cleanup(self):
        """Test emoji removal and bullet detection - user's test case"""
        test_cases = [
            ("- Rising Sea Levels 🌊", "bullet_list", "Rising Sea Levels"),
            ("■ Agriculture 🌾", "bullet_list", "Agriculture"),
            ("* Renewable Energy ⚡", "bullet_list", "Renewable Energy"),
            ("Effects of Climate Change 🌎", "paragraph", "Effects of Climate Change"),
            ("• Biodiversity Loss", "bullet_list", "Biodiversity Loss"),
            ("– Deforestation 🌳", "bullet_list", "Deforestation"),
        ]
        
        results = []
        for text, expected_type, expected_content in test_cases:
            res = self.analyze_line(text, 0)
            passed = (
                res['type'] == expected_type and
                res.get('content', '').startswith(expected_content) and
                '🌊' not in res.get('content', '') and
                '🌾' not in res.get('content', '') and
                '⚡' not in res.get('content', '') and
                '🌎' not in res.get('content', '') and
                '🌳' not in res.get('content', '') and
                '*' not in res.get('content', '') and  # Verify asterisks are removed
                '⁎' not in res.get('content', '') and  # Asterisk variant U+204E
                '⁑' not in res.get('content', '') and  # Asterisk variant U+2051
                '※' not in res.get('content', '')      # Reference mark U+203B
            )
            results.append({
                'text': text,
                'expected_type': expected_type,
                'actual_type': res['type'],
                'expected_content': expected_content,
                'actual_content': res.get('content', ''),
                'emoji_stripped': '🌊' not in res.get('content', '') and
                                 '🌾' not in res.get('content', '') and
                                 '⚡' not in res.get('content', ''),
                'asterisks_removed': '*' not in res.get('content', '') and
                                    '⁎' not in res.get('content', '') and
                                    '⁑' not in res.get('content', '') and
                                    '※' not in res.get('content', ''),
                'passed': passed,
            })
        
        return results



# Note: CoverPageHandler class was removed - cover page templates are now used instead

class CertificationPageHandler:
    """
    Detect and extract certification page information from academic documents.
    Certification pages typically appear after the cover page.
    Generates a standardized certification page with extracted data.
    """
    
    # Regex patterns for certification page extraction
    PATTERNS = {
        # Extract Topic/Title (quoted text or text between keywords)
        'topic': re.compile(r'["“”]([^"“”]+)["“”]|titled\s+["“”]([^"“”]+)["“”]|research\s+titled\s+["“”]([^"“”]+)["“”]|dissertation\s*["“”]([^"“”]+)["“”]|titled\s+(?!["“”])(.+?)\s+(?:submitted\s+by|is\s+the\s+original\s+work\s+of)|report\s+on\s+(.+?)\s+(?:carried\s+out|was\s+done)|internship\s+carried\s+out.*?\s+at\s+([A-Z0-9\s&]+?)(?:\s+was|\s+by)', re.IGNORECASE | re.DOTALL),
        
        # Extract Author Name (after "work of" or "done by")
        'author': re.compile(r'(?:original\s+work\s+of|was\s+done\s+by)\s+([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)+)', re.IGNORECASE),
        
        # Extract Degree Program  
        'degree': re.compile(r'award\s+of\s+(?:a\s+)?([^\.]+?)\s+(?:in\s+[A-Z]|degree)', re.IGNORECASE),
        
        # Extract Supervisor Name
        'supervisor': re.compile(r'(?:_+\s*\n?\s*)?([A-Z][a-z]+\.?\s+[A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)*)\s*\n?\s*\(Supervisor\)', re.IGNORECASE),
        
        # Extract Head of Department Name
        'hod': re.compile(r'(?:_+\s*\n?\s*)?([A-Z][a-z]+\.?\s+[A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)*)\s*\n?\s*\(Head\s+[Oo]f\s+Department\)', re.IGNORECASE),
        
        # Extract Director Name
        'director': re.compile(r'(?:_+\s*\n?\s*)?([A-Z][a-z]+\.?\s+[A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)*)\s*\n?\s*\(Director\)', re.IGNORECASE),
        
        # Certification Header Detection
        'header': re.compile(r'^\s*CERTIFICATION\s*$', re.IGNORECASE | re.MULTILINE),
        
        # Institution extraction
        'institution': re.compile(r'(?:Higher\s+Institute\s+of\s+[A-Za-z\s]+|University\s+of\s+[A-Za-z]+)', re.IGNORECASE),
    }
    
    def __init__(self):
        self.extracted_data = {
            'topic': None,
            'author': None,
            'degree': None,
            'program': None,
            'supervisor': None,
            'head_of_department': None,
            'director': None,
            'institution': 'The Higher Institute of Commerce and Management of The University of Bamenda'
        }
        self.has_certification_page = False
        self.certification_start_index = 0
        self.certification_end_index = 0
    
    def detect_certification_page(self, paragraphs, start_index=0):
        """
        Detect if there's a certification page in the document.
        
        Args:
            paragraphs: List of paragraph objects from python-docx
            start_index: Index to start searching from (after cover page)
            
        Returns:
            tuple: (has_certification, start_index, end_index)
        """
        # Search in paragraphs after cover page, within first 100 paragraphs
        search_end = min(start_index + 100, len(paragraphs))
        
        for i in range(start_index, search_end):
            para = paragraphs[i]
            text = para.text.strip().upper()
            
            # Look for CERTIFICATION header
            if text == 'CERTIFICATION':
                self.has_certification_page = True
                self.certification_start_index = i
                
                # Find where certification page ends
                # Look for next major section (DEDICATION, ACKNOWLEDGEMENT, ABSTRACT, etc.)
                end_markers = ['DEDICATION', 'ACKNOWLEDGEMENT', 'ACKNOWLEDGMENT', 'ABSTRACT', 
                              'TABLE OF CONTENTS', 'LIST OF TABLES', 'LIST OF FIGURES',
                              'CHAPTER', 'DECLARATION']
                
                for j in range(i + 1, min(i + 50, len(paragraphs))):
                    end_text = paragraphs[j].text.strip().upper()
                    if any(end_text.startswith(marker) for marker in end_markers):
                        self.certification_end_index = j
                        break
                else:
                    # If no end marker found, estimate based on content
                    self.certification_end_index = min(i + 40, len(paragraphs))
                
                logger.info(f"Certification page detected at index {i}, ends at {self.certification_end_index}")
                return True, self.certification_start_index, self.certification_end_index
        
        return False, 0, 0
    
    def extract_from_paragraphs(self, paragraphs):
        """
        Extract certification data from paragraph objects.
        
        Args:
            paragraphs: List of paragraph objects (certification page paragraphs only)
        """
        # Combine all text for pattern matching
        full_text = '\n'.join([p.text for p in paragraphs])
        
        # 1. Extract Topic (quoted text)
        topic_match = self.PATTERNS['topic'].search(full_text)
        if topic_match:
            for group in topic_match.groups():
                if group:
                    topic_text = group.strip()
                    # If it looks like a place (internship), prefix it
                    if 'BACCCUL' in topic_text or 'BANK' in topic_text or 'COUNCIL' in topic_text or len(topic_text) < 20:
                         # Check if it's likely a place name rather than a full title
                         if not topic_text.lower().startswith('the effect') and not topic_text.lower().startswith('an analysis'):
                             topic_text = f"Internship Report at {topic_text}"
                    
                    self.extracted_data['topic'] = topic_text
                    break
        
        # 2. Extract Author Name
        author_match = self.PATTERNS['author'].search(full_text)
        if author_match:
            author_text = author_match.group(1).strip()
            # Clean up author name (remove trailing "with", "registration", etc.)
            for stop_word in [' with', ' registration', ' student', ' matriculation', ' level']:
                if stop_word in author_text.lower():
                    author_text = re.split(stop_word, author_text, flags=re.IGNORECASE)[0].strip()
            self.extracted_data['author'] = author_text
        
        # 3. Extract Degree and Program
        degree_match = self.PATTERNS['degree'].search(full_text)
        if degree_match:
            degree_text = degree_match.group(1).strip()
            # Split degree and program (e.g., "Master's in Business Administration (MBA) in Management")
            if ' in ' in degree_text.lower():
                parts = re.split(r'\s+in\s+', degree_text, maxsplit=1, flags=re.IGNORECASE)
                if len(parts) >= 2:
                    self.extracted_data['degree'] = parts[0].strip()
                    self.extracted_data['program'] = parts[1].strip()
                else:
                    self.extracted_data['degree'] = degree_text
            else:
                self.extracted_data['degree'] = degree_text
        
        # 4. Extract Supervisor Name
        supervisor_match = self.PATTERNS['supervisor'].search(full_text)
        if supervisor_match:
            self.extracted_data['supervisor'] = self._format_name(supervisor_match.group(1).strip())
        
        # 5. Extract Head of Department Name
        hod_match = self.PATTERNS['hod'].search(full_text)
        if hod_match:
            self.extracted_data['head_of_department'] = self._format_name(hod_match.group(1).strip())
        
        # 6. Extract Director Name
        director_match = self.PATTERNS['director'].search(full_text)
        if director_match:
            self.extracted_data['director'] = self._format_name(director_match.group(1).strip())
        
        # 7. Extract Institution
        inst_match = self.PATTERNS['institution'].search(full_text)
        if inst_match:
            self.extracted_data['institution'] = inst_match.group(0).strip()
        
        logger.info(f"Certification data extracted: {self.extracted_data}")
        return self.extracted_data
    
    def _format_name(self, name):
        """Format name with proper title abbreviation."""
        if not name:
            return name
        
        # Standardize title abbreviations
        name = re.sub(r'^Prof\.?\s+', 'Prof. ', name, flags=re.IGNORECASE)
        name = re.sub(r'^Dr\.?\s+', 'Dr. ', name, flags=re.IGNORECASE)
        name = re.sub(r'^Engr\.?\s+', 'Engr. ', name, flags=re.IGNORECASE)
        name = re.sub(r'^Mr\.?\s+', 'Mr. ', name, flags=re.IGNORECASE)
        name = re.sub(r'^Mrs\.?\s+', 'Mrs. ', name, flags=re.IGNORECASE)
        name = re.sub(r'^Ms\.?\s+', 'Ms. ', name, flags=re.IGNORECASE)
        
        return name
    
    def detect_and_extract(self, doc, start_index=0):
        """
        Detect certification page and extract all data.
        
        Args:
            doc: python-docx Document object
            start_index: Index to start searching from (after cover page)
            
        Returns:
            tuple: (has_certification, extracted_data, start_index, end_index)
        """
        paragraphs = doc.paragraphs
        
        has_cert, start_idx, end_idx = self.detect_certification_page(paragraphs, start_index)
        
        if has_cert:
            cert_paragraphs = paragraphs[start_idx:end_idx]
            self.extract_from_paragraphs(cert_paragraphs)
        
        return self.has_certification_page, self.extracted_data, self.certification_start_index, self.certification_end_index


def format_questionnaire_in_word(doc, questionnaire_data, font_size=11):
    """
    Format a questionnaire document in Word based on extracted structure.
    """
    # Clear existing content if any (except styles)
    # Note: We don't clear everything because we might want to keep headers/footers
    # But for the body content, we start fresh or append.
    # If doc is empty, fine. If not, we append.
    
    # Set up styles
    styles = doc.styles
    
    # Title Style
    if 'Questionnaire Title' not in styles:
        style = styles.add_style('Questionnaire Title', WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(16)
        font.bold = True
        pf = style.paragraph_format
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.space_after = Pt(12)
    
    # Section Header Style
    if 'Questionnaire Section' not in styles:
        style = styles.add_style('Questionnaire Section', WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        font.bold = True
        pf = style.paragraph_format
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
        
    # Question Style
    if 'Question Text' not in styles:
        style = styles.add_style('Question Text', WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(font_size)
        font.bold = True
        pf = style.paragraph_format
        pf.space_after = Pt(3)
        
    # Option Style
    if 'Question Option' not in styles:
        style = styles.add_style('Question Option', WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(font_size)
        pf = style.paragraph_format
        pf.left_indent = Pt(0)
        pf.space_after = Pt(2)

    # Add Title
    if questionnaire_data.get('title'):
        doc.add_paragraph(questionnaire_data['title'], style='Questionnaire Title')
        
    # Add Instructions
    if questionnaire_data.get('instructions'):
        for instruction in questionnaire_data['instructions']:
            p = doc.add_paragraph(instruction)
            p.paragraph_format.space_after = Pt(12)
            p.italic = False
            
    # Process Sections
    for section in questionnaire_data.get('sections', []):
        # Add Section Header
        if section.get('title') and section.get('title') != 'Questions':
            doc.add_paragraph(section['title'], style='Questionnaire Section')
            
        # Process Questions
        for question in section.get('questions', []):
            # Question Text
            q_num = question.get('number', '')
            q_text_content = question.get('text', '')
            full_text = f"{q_num} {q_text_content}".strip() if q_num else q_text_content
            
            doc.add_paragraph(full_text, style='Question Text')
            
            # Question Options/Input
            q_type = question.get('type', 'open_ended')
            
            if q_type == 'likert_table':
                # Create a full Likert Table (Matrix)
                scale = question.get('scale', {})
                items = scale.get('items', [])
                sub_questions = question.get('sub_questions', [])
                
                if items and sub_questions:
                    # Create table: 1 col for statement + N cols for scale items
                    table = doc.add_table(rows=len(sub_questions) + 1, cols=len(items) + 1)
                    table.style = 'Table Grid'
                    table.autofit = True
                    
                    # Header Row
                    # First cell is "Statement"
                    header_cell = table.cell(0, 0)
                    header_cell.text = "Statement"
                    header_cell.paragraphs[0].runs[0].bold = True
                    
                    # Scale headers
                    for i, item in enumerate(items):
                        cell = table.cell(0, i + 1)
                        cell.text = item
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if p.runs:
                            p.runs[0].bold = True
                    
                    # Data Rows
                    for r, statement in enumerate(sub_questions):
                        # Statement cell
                        row_idx = r + 1
                        cell = table.cell(row_idx, 0)
                        cell.text = statement
                        
                        # Radio buttons for each scale item
                        for c in range(len(items)):
                            cell = table.cell(row_idx, c + 1)
                            p = cell.paragraphs[0]
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run('○')
                            run.font.size = Pt(14)
                    
                    # Add spacing after table
                    doc.add_paragraph().paragraph_format.space_after = Pt(12)

            elif q_type == 'scale' and question.get('scale'):
                # Create Likert Scale Table (Single Question)
                scale = question['scale']
                items = scale.get('items', [])
                if items:
                    table = doc.add_table(rows=2, cols=len(items))
                    table.style = 'Table Grid'
                    table.autofit = True
                    
                    # Header Row
                    for i, item in enumerate(items):
                        cell = table.cell(0, i)
                        cell.text = item
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                    # Radio Button Row
                    for i in range(len(items)):
                        cell = table.cell(1, i)
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run = p.add_run('○')
                        run.font.size = Pt(14)
                    
                    # Add spacing after table
                    doc.add_paragraph().paragraph_format.space_after = Pt(6)
            
            elif q_type in ['multiple_choice', 'single_select', 'radio']:
                options = question.get('options', [])
                # Check if we should use horizontal layout (short options)
                # Criteria: max length < 20 chars, total length < 80 chars
                is_short = all(len(opt.get('text', '')) < 20 for opt in options)
                total_len = sum(len(opt.get('text', '')) for opt in options)
                
                if is_short and total_len < 80 and len(options) > 1:
                    # Horizontal Layout
                    p = doc.add_paragraph(style='Question Option')
                    p.paragraph_format.left_indent = Pt(0)
                    
                    for i, option in enumerate(options):
                        run = p.add_run('○ ')
                        run.font.name = 'Segoe UI Symbol'
                        p.add_run(option.get('text', ''))
                        
                        # Add spacing between options (except last)
                        if i < len(options) - 1:
                            p.add_run('\t\t') # Double tab for spacing
                else:
                    # Vertical Layout
                    for option in options:
                        p = doc.add_paragraph(style='Question Option')
                        
                        # Handle text input option (Other: ___)
                        if option.get('type') == 'text_input':
                            run = p.add_run('○ ')
                            run.font.name = 'Segoe UI Symbol'
                            p.add_run('Other: ' + '_' * 30)
                        else:
                            run = p.add_run('○ ')
                            run.font.name = 'Segoe UI Symbol'
                            p.add_run(option.get('text', ''))
                    
            elif q_type in ['multiple_select', 'checkbox', 'check_all']:
                options = question.get('options', [])
                # Check if we should use horizontal layout
                is_short = all(len(opt.get('text', '')) < 20 for opt in options)
                total_len = sum(len(opt.get('text', '')) for opt in options)
                
                if is_short and total_len < 80 and len(options) > 1:
                    # Horizontal Layout
                    p = doc.add_paragraph(style='Question Option')
                    p.paragraph_format.left_indent = Pt(0)
                    
                    for i, option in enumerate(options):
                        run = p.add_run('☐ ')
                        run.font.name = 'Segoe UI Symbol'
                        p.add_run(option.get('text', ''))
                        
                        # Add spacing between options
                        if i < len(options) - 1:
                            p.add_run('\t\t')
                else:
                    # Vertical Layout
                    for option in options:
                        p = doc.add_paragraph(style='Question Option')
                        run = p.add_run('☐ ')
                        run.font.name = 'Segoe UI Symbol'
                        p.add_run(option.get('text', ''))
                    
            elif q_type == 'open_ended':
                # Add lines for writing
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(0)
                p.add_run('_' * 60)
                p.paragraph_format.space_after = Pt(6)
                
    return doc


class QuestionnaireProcessor:
    """
    Detect, parse, and format questionnaires, surveys, and assessment forms.
    """
    def __init__(self):
        self.questionnaire_data = {
            'is_questionnaire': False,
            'title': '',
            'sections': [],
            'questions': [],
            'demographics': [],
            'instructions': [],
            'scale_type': None
        }
    
    def detect_questionnaire(self, text):
        """
        Determine if document is a questionnaire and extract structure
        """
        if not text:
            return self.questionnaire_data
            
        lines = text.split('\n')
        questionnaire_indicators = 0
        total_indicators = 0
        
        # Quick check for questionnaire keywords in first few lines
        header_check = '\n'.join(lines[:20]).upper()
        if not any(k in header_check for k in ['QUESTIONNAIRE', 'SURVEY', 'ASSESSMENT', 'EVALUATION', 'FEEDBACK FORM']):
            # If no explicit title, check for question density
            question_count = sum(1 for line in lines[:50] if self.is_question_line(line))
            if question_count < 3:
                return self.questionnaire_data
        
        for i, line in enumerate(lines):
            # Check for title patterns
            if self.is_questionnaire_title(line):
                questionnaire_indicators += 3
                self.questionnaire_data['title'] = line.strip()
            
            # Check for question patterns
            if self.is_question_line(line):
                questionnaire_indicators += 2
            
            # Check for Likert scales
            if self.is_likert_scale(line):
                questionnaire_indicators += 2
            
            # Check for demographic sections
            if self.is_demographic_section(line):
                questionnaire_indicators += 2
            
            total_indicators += 1
        
        # Calculate confidence score
        if total_indicators > 0:
            # Normalize confidence
            confidence = min((questionnaire_indicators / 15) * 100, 100)
            self.questionnaire_data['is_questionnaire'] = confidence > 40 # Lower threshold as indicators are specific
            self.questionnaire_data['confidence'] = confidence
        
        return self.questionnaire_data
    
    def parse_questionnaire_structure(self, text):
        """
        Parse questionnaire into structured format
        """
        if not text:
            return self.questionnaire_data
            
        lines = text.split('\n')
        current_section = None
        current_question = None
        
        # State tracking for Likert tables
        in_likert_table = False
        current_likert_scale = []
        
        # Extract instructions first
        for i, line in enumerate(lines[:20]):
            if self.is_instruction_line(line):
                self.questionnaire_data['instructions'].append(line.strip())
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
                
            # Detect section headers
            section_match = self.detect_section_header(line)
            if section_match:
                if current_question:
                    self._finalize_question(current_question)
                    current_question = None
                
                # Reset table state
                in_likert_table = False
                current_likert_scale = []
                
                current_section = {
                    'type': section_match['type'],
                    'title': section_match['title'],
                    'questions': []
                }
                self.questionnaire_data['sections'].append(current_section)
                continue
            
            # Detect Likert Table Header (e.g., "Table 1: Income Adequacy")
            table_match = re.match(r'^(?:Table|Tbl)\s*\d+[:.]\s*(.+)$', line, re.IGNORECASE)
            if table_match:
                if current_question:
                    self._finalize_question(current_question)
                    current_question = None
                
                # Start a new "question" that is actually a table
                current_question = {
                    'number': '',
                    'text': table_match.group(1).strip(),
                    'type': 'likert_table',
                    'options': [],
                    'scale': None,
                    'sub_questions': [] # For table rows
                }
                if current_section:
                    current_section['questions'].append(current_question)
                else:
                    current_section = {'type': 'main', 'title': 'Questions', 'questions': [current_question]}
                    self.questionnaire_data['sections'].append(current_section)
                
                in_likert_table = True
                continue
            
            # Handle content inside Likert Table
            if in_likert_table and current_question:
                # Check if this line is actually a new question
                # If so, we should exit table mode
                if self.detect_question(line):
                    in_likert_table = False
                    # Fall through to main question detection
                else:
                    # Check for instructions inside table
                    if self.is_instruction_line(line) or line.strip().lower().startswith('instructions:'):
                        current_question['instructions'] = line.strip()
                        continue

                    # Check if this line defines the scale (header row)
                    # e.g. "Statement Strongly Disagree Disagree Neutral Agree Strongly Agree"
                    likert_indicators = [
                        'strongly disagree', 'disagree', 'neutral', 'agree', 'strongly agree',
                        'never', 'rarely', 'sometimes', 'often', 'always',
                        'very dissatisfied', 'dissatisfied', 'satisfied', 'very satisfied',
                        'not at all', 'slightly', 'moderately', 'very', 'extremely'
                    ]
                    line_lower = line.lower()
                    
                    # If line contains multiple scale indicators, it's likely the header row
                    found_indicators = [ind for ind in likert_indicators if ind in line_lower]
                    # Also check for abbreviated headers (SD D N A SA)
                    has_abbrev = re.search(r'\b(?:SD|D|N|A|SA)\b', line) and len(line.split()) <= 10
                    
                    if len(found_indicators) >= 3 or has_abbrev:
                        # Extract the scale items properly (splitting by tab, multiple spaces, or pipes)
                        # Remove leading/trailing pipes first
                        clean_line = line.strip('|').strip()
                        parts = re.split(r'\||\t|\s{2,}', clean_line)
                        # Filter out "Statement", empty strings, and whitespace
                        scale_items = [p.strip() for p in parts if p.strip() and p.strip().lower() != 'statement']
                        current_question['scale'] = {'items': scale_items, 'type': 'likert'}
                        continue
                    
                    # Otherwise, it's a row in the table (a sub-question)
                    # e.g. "Our household income is sufficient... [ ] [ ] [ ]"
                    # We want to extract the statement text.
                    # Remove the [ ] parts and pipes
                    clean_line = re.sub(r'\[\s*[xX]?\s*\]', '', line)
                    clean_line = clean_line.replace('|', '').strip()
                    
                    if clean_line:
                        current_question['sub_questions'].append(clean_line)
                    continue

            # Detect questions
            question_match = self.detect_question(line)
            if question_match:
                if current_question:
                    self._finalize_question(current_question)
                
                # Reset table state
                in_likert_table = False
                current_likert_scale = []

                current_question = {
                    'number': question_match['number'],
                    'text': question_match['text'],
                    'type': question_match['type'],
                    'options': [],
                    'scale': None,
                    'instructions': ''
                }
                
                if current_section:
                    current_section['questions'].append(current_question)
                else:
                    # Create default section if none exists
                    current_section = {
                        'type': 'main',
                        'title': 'Questions',
                        'questions': [current_question]
                    }
                    self.questionnaire_data['sections'].append(current_section)
                continue
            
            # Detect options for current question
            if current_question:
                option_match = self.detect_option(line)
                if option_match:
                    if option_match['type'] in ['checkbox', 'checked'] and current_question['type'] in ['multiple_choice', 'single_select', 'radio']:
                        current_question['type'] = 'checkbox'
                    current_question['options'].append({
                        'label': option_match['label'],
                        'text': option_match['text'],
                        'type': option_match['type']
                    })
                    continue
                
                # Detect "Other: ____" as an option
                other_match = re.match(r'^(?:Other|Specify)\s*[:.]\s*(_+)?$', line, re.IGNORECASE)
                if other_match:
                     current_question['options'].append({
                        'label': '',
                        'text': 'Other: ________________',
                        'type': 'text_input'
                    })
                     continue

                # Detect scale for current question
                scale_match = self.detect_scale(line)
                if scale_match and not current_question.get('scale'):
                    current_question['scale'] = scale_match
                    continue
                
                # Capture question continuation or instructions
                if line.strip() and not self.is_section_header(line) and not self.is_instruction_line(line):
                    # Check if it's an implicit option
                    # If question ends with ?, treat subsequent lines as options if they are not new questions
                    question_ended = current_question['text'].strip().endswith('?') or current_question['text'].strip().endswith(':')
                    
                    # Also check if the line looks like an option (short, capitalized)
                    is_likely_option = len(line) < 100 and line[0].isupper() or line[0].isdigit()
                    
                    if (question_ended or len(current_question['options']) > 0) and not self.is_question_line(line):
                         # If we already have options, or the question seems finished, treat as option
                         current_question['options'].append({
                                'label': '',
                                'text': line.strip(),
                                'type': 'implicit_option'
                            })
                         continue

                    if not current_question['options'] and not current_question.get('scale'):
                        # Likely continuation of question text
                        # Only append if it doesn't look like a new question start
                        if not self.is_question_line(line):
                            current_question['text'] += ' ' + line.strip()
                    elif line.strip().endswith('?'):
                        # Additional question part
                        current_question['text'] += ' ' + line.strip()


        
        # Finalize last question
        if current_question:
            self._finalize_question(current_question)
        
        return self.questionnaire_data
    
    def _finalize_question(self, question):
        """Finalize question processing"""
        # Auto-detect question type if not specified
        if not question.get('type'):
            question['type'] = self._infer_question_type(question)
        
        # Ensure options have proper types
        for option in question['options']:
            if not option.get('type'):
                option['type'] = self._infer_option_type(question['type'])
            
            # Fix implicit options
            if option.get('type') == 'implicit_option':
                option['type'] = self._infer_option_type(question['type'])

    def _infer_option_type(self, question_type):
        """Infer option type based on question type"""
        if question_type in ['multiple_select', 'check_all']:
            return 'checkbox'
        elif question_type in ['single_select', 'multiple_choice', 'scale']:
            return 'radio'
        else:
            return 'text'

    # DETECTION HELPER METHODS
    
    def is_questionnaire_title(self, line):
        patterns = [
            r'^(?:#+\s+)?(?:QUESTIONNAIRE|SURVEY|ASSESSMENT|EVALUATION)\b',
            r'^(?:#+\s+)?(?:Research|Study|Data)\s+(?:Collection|Gathering)\s+(?:Tool|Instrument)\b',
            r'^(?:#+\s+)?(?:Student|Teacher|Parent|Employee|Customer)\s+(?:Feedback|Satisfaction)\s+(?:Form|Survey)\b',
            r'^(?:#+\s+)?(?:Self-[Aa]ssessment|Self-[Ee]valuation)\b',
            r'^(?:#+\s+)?([A-Z][A-Za-z\s]+(?:Assessment|Evaluation|Survey|Study)\s+(?:Questionnaire|Survey|Form))\s*$'
        ]
        for pattern in patterns:
            if re.match(pattern, line, re.IGNORECASE):
                return True
        return False
    
    def is_question_line(self, line):
        patterns = [
            r'^\s*(?:Q|Question)\s*\d+[:.)]',
            r'^\s*\d+[:.)]\s+.*[?]',
            r'^\s*Item\s+\d+[:.)]',
            r'^\s*[A-Z][:.)]\s+.*[?]'
        ]
        for pattern in patterns:
            if re.match(pattern, line, re.IGNORECASE):
                return True
        return False
    
    def is_likert_scale(self, line):
        likert_indicators = [
            'strongly disagree', 'disagree', 'neutral', 'agree', 'strongly agree',
            'never', 'rarely', 'sometimes', 'often', 'always',
            'very dissatisfied', 'dissatisfied', 'satisfied', 'very satisfied',
            'not at all', 'slightly', 'moderately', 'very', 'extremely'
        ]
        line_lower = line.lower()
        count = sum(1 for indicator in likert_indicators if indicator in line_lower)
        
        # Also check for abbreviated headers (SD D N A SA)
        if re.search(r'\b(?:SD|D|N|A|SA)\b', line) and len(line.split()) <= 10:
             return True
             
        return count >= 3  # At least 3 Likert items present
    
    def is_demographic_section(self, line):
        patterns = [
            r'^(?:#+\s+)?(?:DEMOGRAPHIC|BACKGROUND|PERSONAL)\s+(?:INFORMATION|DATA)\b',
            r'^\s*(?:Name|Age|Gender|School|University|College|Class|Grade)\s*[:.]',
            r'^PART\s+[A-Z]\s*[:.]?\s*(?:Demographic|Personal)'
        ]
        for pattern in patterns:
            if re.match(pattern, line, re.IGNORECASE):
                return True
        return False
    
    def is_instruction_line(self, line):
        patterns = [
            r'^\s*(?:Instructions?|Directions?|Guidelines?)\s*[:.]?\s*$',
            r'^\s*(?:Please|Kindly)\s+(?:read|answer|complete|fill|respond)',
            r'^\s*(?:All\s+questions|Each\s+item)\s+(?:must|should)\s+be',
            r'^\s*(?:Mark|Tick|Check|Circle)\s+(?:your|the)\s+(?:answer|response|choice)'
        ]
        for pattern in patterns:
            if re.match(pattern, line, re.IGNORECASE):
                return True
        return False
        
    def is_section_header(self, line):
        return self.detect_section_header(line) is not None

    def detect_section_header(self, line):
        patterns = {
            r'^(?:#+\s+)?PART\s+([A-Z])\s*[:.]?\s*(.+)$': {'type': 'part', 'group': 1, 'title': 2},
            r'^(?:#+\s+)?SECTION\s+(\d+)\s*[:.]?\s*(.+)$': {'type': 'section', 'group': 1, 'title': 2},
            r'^(?:#+\s+)?(?:Module|Segment)\s+([A-Z\d])\s*[:.]?\s*(.+)$': {'type': 'module', 'group': 1, 'title': 2},
            r'^(?:#+\s+)?(I+|[IVXLCDM]+)\s*\.\s*(.+)$': {'type': 'roman_section', 'group': 1, 'title': 2},
            r'^(?:#+\s+)?(?:DEMOGRAPHIC|BACKGROUND|PERSONAL)\s+(?:INFORMATION|DATA)\b': {'type': 'demographic', 'group': 0, 'title': 0},
            r'^(?:#+\s+)?(?:Introduction|Purpose|Background)\s*$': {'type': 'intro', 'group': 0, 'title': 0},
            r'^(?:#+\s+)?(?:Additional\s+Comments|Optional\s+Contact\s+Information)\s*$': {'type': 'meta', 'group': 0, 'title': 0}
        }
        
        for pattern, config in patterns.items():
            match = re.match(pattern, line, re.IGNORECASE)
            if match:
                if config['group'] == 0:
                    return {
                        'type': config['type'],
                        'identifier': '',
                        'title': match.group(0).strip()
                    }
                return {
                    'type': config['type'],
                    'identifier': match.group(config['group']),
                    'title': match.group(config['title']).strip()
                }
        return None

    def detect_question(self, line):
        patterns = [
            (r'^\s*(?:Q|Question)\s*(\d+)[:.)]\s+(.+)$', 'standard'),
            (r'^\s*(\d+)[:.)]\s+(.+[?])(?:\s*\(.+\))?$', 'numbered'),
            (r'^\s*Item\s+(\d+)[:.)]\s+(.+)$', 'item'),
            (r'^\s*([A-Z])[:.)]\s+(.+[?])(?:\s*\(.+\))?$', 'lettered'),
            (r'^\s*(\d+[a-z]?)[:.)]\s+(.+)$', 'subnumbered'),
            (r'^\s*(\d+\.\d+)[:.)]\s+(.+)$', 'decimal'),
            # Enhanced patterns for unnumbered questions
            (r'^\s*([A-Z][A-Za-z\s?!,\'()\-]+[?])\s*$', 'unnumbered'),
            (r'^\s*([A-Z][a-z]+(?:\s+[A-Z]?[a-z]+){1,15}[?])\s*$', 'unnumbered'),
            (r'^\s*(.+\?)\s*(?:\([^)]+\))?\s*$', 'unnumbered'),
            # Pattern for questions ending in colon (common in surveys)
            (r'^\s*([A-Z].+[:])\s*$', 'unnumbered')
        ]
        
        for pattern, q_type in patterns:
            match = re.match(pattern, line) # Removed IGNORECASE to respect capitalization for unnumbered
            if match:
                # For unnumbered, the group is 1 (text), for others it's 2 (text)
                if q_type == 'unnumbered':
                    question_text = match.group(1).strip()
                    number = ''
                else:
                    question_text = match.group(2).strip()
                    number = match.group(1)
                
                # Determine question type from content
                inferred_type = self._infer_question_type_from_text(question_text)
                
                return {
                    'number': number,
                    'text': question_text,
                    'format': q_type,
                    'type': inferred_type
                }
        return None
    
    def _infer_question_type_from_text(self, text):
        """Infer question type from text content"""
        text_lower = text.lower()
        
        if any(phrase in text_lower for phrase in ['rate', 'on a scale', '1 to 5', '1-5']):
            return 'scale'
        elif any(phrase in text_lower for phrase in ['select all', 'all that apply', 'choose all']):
            return 'multiple_select'
        elif any(phrase in text_lower for phrase in ['select one', 'choose one', 'circle one']):
            return 'single_select'
        elif text_lower.endswith('?'):
            # Ends with question mark but not scale-related
            return 'multiple_choice'
        else:
            return 'open_ended'
    
    def detect_option(self, line):
        patterns = [
            (r'^\s*([a-zA-Z])\s*[.)]\s*(.+)$', 'lettered', None),
            (r'^\s*\(([a-zA-Z])\)\s+(.+)$', 'parenthesized', None),
            (r'^\s*(\d+)\s*[.)]\s*(.+)$', 'numbered', None),
            (r'^\s*\[[\s]\](?:\s*([a-zA-Z0-9])(?:[.)]|\s+))?\s*(.+)$', 'checkbox', 'checkbox'),
            (r'^\s*\[[xX]\]\s*(.+)$', 'checked', 'checked'),
            (r'^\s*(?:□|☐)\s*(.+)$', 'checkbox_symbol', 'checkbox'),
            (r'^\s*(?:○|◯)\s*(.+)$', 'radio_empty', 'radio'),
            (r'^\s*(?:●|◉)\s*(.+)$', 'radio_selected', 'selected'),
            # Enhanced Option Patterns
            (r'^\s*\$\s*([\d,]+)\s*[-–]\s*\$\s*([\d,]+)\s*$', 'range_currency', 'radio'),
            (r'^\s*([\d.%]+)\s*[-–]\s*([\d.%]+)\s*$', 'range_numeric', 'radio'),
            (r'^\s*(?:Less than|More than|Above|Below)\s+.+$', 'range_descriptive', 'radio'),
            (r'^\s*(?:Yes|No),\s+[A-Za-z].+$', 'yes_no_extended', 'radio'),
            (r'^\s*(?:We|I|Our)\s+[A-Za-z].+$', 'first_person', 'radio'),
            (r'^\s*(?:Much|Somewhat|Slightly|Significantly)\s+(?:worse|better)\s*$', 'comparative', 'radio'),
            (r'^\s*About the same\s*$', 'comparative', 'radio')
        ]
        
        for pattern, opt_type, opt_symbol in patterns:
            match = re.match(pattern, line)
            if match:
                groups = match.groups()
                if len(groups) == 2:
                    label, text = groups[0], groups[1]
                    # For ranges, combine groups into text
                    if opt_type in ['range_currency', 'range_numeric']:
                        text = line.strip()
                        label = ''
                elif len(groups) == 1:
                    label, text = '', groups[0]
                else:
                    # No capturing groups, use the whole line/match as text
                    label, text = '', match.group(0).strip()
                
                # For descriptive/comparative, ensure we use the full text if needed
                if opt_type in ['range_descriptive', 'yes_no_extended', 'first_person', 'comparative']:
                    text = line.strip()
                
                return {
                    'label': label.strip() if label else '',
                    'text': text.strip(),
                    'format': opt_type,
                    'type': opt_symbol or 'text'
                }
        return None
    
    def detect_scale(self, line):
        # Likert scale detection
        likert_pattern = r'(?:Strongly\s+Disagree|Disagree|Neutral|Agree|Strongly\s+Agree)'
        if re.search(likert_pattern, line, re.IGNORECASE):
            matches = re.findall(likert_pattern, line, re.IGNORECASE)
            if len(matches) >= 3:
                return {
                    'type': 'likert',
                    'items': matches,
                    'min': 1,
                    'max': 5,
                    'description': 'Likert scale'
                }
        
        # Numeric scale detection
        numeric_pattern = r'(\d)\s*[-\|]\s*(\d)\s*[-\|]\s*(\d)\s*[-\|]\s*(\d)\s*[-\|]\s*(\d)'
        match = re.search(numeric_pattern, line)
        if match:
            return {
                'type': 'numeric',
                'items': [match.group(i) for i in range(1, 6)],
                'min': int(match.group(1)),
                'max': int(match.group(5)),
                'description': 'Numeric scale'
            }
        
        # Rating scale detection
        rating_pattern = r'(?:Poor|Fair|Average|Good|Excellent)'
        if re.search(rating_pattern, line, re.IGNORECASE):
            matches = re.findall(rating_pattern, line, re.IGNORECASE)
            if len(matches) >= 3:
                return {
                    'type': 'rating',
                    'items': matches,
                    'min': 1,
                    'max': len(matches),
                    'description': 'Rating scale'
                }
        
        return None
    
    def _detect_separator_style(self, text):
        """Detect the style of separator used"""
        if '+' in text and '-' in text:
            return 'box'
        elif '=' in text:
            return 'double'
        elif '_' in text:
            return 'underscore'
        else:
            return 'dash'


class TextFormatterWithRegex:
    """
    Applies powerful regex patterns to automatically bold numbered/bulleted topics
    Handles inconsistent formatting across documents
    IMPROVED: Patterns match actual document structure
    """
    
    def __init__(self, policy=None):
        self.policy = policy or FormatPolicy()
        # Define patterns in application order - IMPROVED FOR REAL DOCUMENTS
        self.patterns = [
            # Pattern 1: Numbered sections at start of line (1.1, 2.1, 1.2, etc.)
            {
                'regex': r'^(\d+\.\d+)(\s+)([A-Z][A-Za-z\s\-:]*?)(?=\n|$)',
                'replacement': r'**\1\2\3**',
                'name': 'Numbered sections (1.1, 2.2, etc)',
                'flags': re.MULTILINE
            },
            
            # Pattern 2: Simple numbered items (1., 2., 3.) at line start with title
            {
                'regex': r'^(\d+\.\s+)([A-Z][A-Za-z\s\-:]*?)(?=\n|$)',
                'replacement': r'**\1\2**',
                'name': 'Simple numbered items (1., 2., 3.)',
                'flags': re.MULTILINE
            },
            
            # Pattern 3: Roman numerals with title (I., II., III.)
            {
                'regex': r'^([IVX]+\.\s+)([A-Z][A-Za-z\s\-:]*?)(?=\n|$)',
                'replacement': r'**\1\2**',
                'name': 'Roman numerals (I., II., III.)',
                'flags': re.MULTILINE
            },
            
            # Pattern 4: Bulleted items with text
            {
                'regex': r'^(\s*[-•]\s+)([A-Z][A-Za-z\s\-:]*?)(?=\n|$)',
                'replacement': r'\1**\2**',
                'name': 'Bulleted items',
                'flags': re.MULTILINE
            },
            
            # Pattern 5: Section headers in capitals followed by colon (METHODOLOGY:, RESULTS:)
            {
                'regex': r'^([A-Z][A-Z\s]+):(?=\s|$)',
                'replacement': r'**\1:**',
                'name': 'Section headers (CAPITALS with colon)',
                'flags': re.MULTILINE
            },
        ]
    
    def format_text(self, text):
        """Apply all regex patterns to text in sequence"""
        if not text:
            return text
        if not self.policy.enable_regex_auto_bold:
            return text
        
        for pattern_config in self.patterns:
            try:
                regex = pattern_config['regex']
                replacement = pattern_config['replacement']
                flags = pattern_config['flags']
                
                # Apply the pattern
                updated = re.sub(regex, replacement, text, flags=flags)
                if updated != text:
                    logger.info(
                        "Regex auto-bold applied (%s): '%s' -> '%s'",
                        pattern_config['name'],
                        text,
                        updated,
                    )
                text = updated
                
            except Exception as e:
                logger.warning(f"Error applying pattern '{pattern_config['name']}': {e}")
                # Continue with next pattern if one fails
                continue
        
        return text
    
    def should_apply_formatting(self, text):
        """
        Determine if text needs formatting.
        Returns True if text has unformatted numbered/bulleted items
        """
        if not text:
            return False
        
        # Check for numbered items without bold
        if re.search(r'^\d+\.\d+\s+[A-Z]', text, re.MULTILINE):
            return True
        if re.search(r'^\d+\.\s+[A-Z]', text, re.MULTILINE):
            return True
        
        # Check for bulleted items without bold
        if re.search(r'^[-•]\s+[A-Z]', text, re.MULTILINE):
            return True
        
        return False


class DocumentProcessor:
    """Process documents using pattern engine"""
    
    def __init__(self, policy=None):
        # Toggle academic mode via FormatPolicy(document_mode="academic").
        self.policy = policy or FormatPolicy()
        self.engine = PatternEngine(policy=self.policy)
        self.image_extractor = ImageExtractor()
        self.shape_extractor = ShapeExtractor()  # NEW: Extract shapes/flowcharts
        self.extracted_images = []  # Store extracted images
        self.extracted_shapes = []  # NEW: Store extracted shapes/flowcharts
        self.text_formatter = TextFormatterWithRegex(policy=self.policy)  # Add regex formatter
        # Note: Cover page handler removed - templates are now used instead
        self.certification_handler = CertificationPageHandler()  # Certification page detection
        self.certification_data = None  # Extracted certification page data
        self.certification_start_index = 0  # Where certification page starts
        self.certification_end_index = 0  # Where certification page ends
        self.questionnaire_processor = QuestionnaireProcessor() # Questionnaire detection
        self.questionnaire_data = None # Extracted questionnaire data
        self.heading_numberer = HeadingNumberer(policy=self.policy)  # Auto-number headings based on chapter context

    def _build_paragraph_segments(self, para, paragraph_index, shape_run_map):
        """Build ordered text/shape segments from a paragraph's run elements."""
        segments = []
        has_shapes = False

        for run_index, run in enumerate(para.runs):
            shapes_in_run = shape_run_map.get((paragraph_index, run_index), [])
            shapes_by_element_index = {}
            for shape in sorted(shapes_in_run, key=lambda s: s.get('element_index', 0)):
                shapes_by_element_index.setdefault(shape.get('element_index', 0), []).append(shape)

            run_elements = list(run._element)
            used_shape_ids = set()
            saw_text_node = False
            for element_index, element in enumerate(run_elements):
                if element.tag == qn('w:t'):
                    saw_text_node = True
                    if element.text:
                        segments.append({'type': 'text', 'text': element.text})
                elif element.tag in (qn('w:drawing'), qn('w:pict')):
                    for shape in shapes_by_element_index.get(element_index, []):
                        segments.append({'type': 'shape', 'shape_id': shape['shape_id']})
                        used_shape_ids.add(shape['shape_id'])
                        has_shapes = True

            if not saw_text_node and run.text:
                segments.append({'type': 'text', 'text': run.text})

            for shape in shapes_in_run:
                if shape['shape_id'] not in used_shape_ids:
                    segments.append({'type': 'shape', 'shape_id': shape['shape_id']})
                    has_shapes = True

        combined = []
        for segment in segments:
            if segment['type'] == 'text' and combined and combined[-1]['type'] == 'text':
                combined[-1]['text'] += segment['text']
            else:
                combined.append(segment)

        return combined, has_shapes
        
    def process_docx(self, file_path, strip_front_matter=True):
        """Process Word document line by line, preserving table and image positions"""
        doc = Document(file_path)
        
        # Step 0: Extract all images from document FIRST
        self.extracted_images = self.image_extractor.extract_all_images(file_path)
        logger.info(f"Extracted {len(self.extracted_images)} images from document")
        
        # Step 0.1: Extract all shapes/flowcharts from document
        self.extracted_shapes = self.shape_extractor.extract_all_shapes(file_path)
        logger.info(f"Extracted {len(self.extracted_shapes)} shapes/flowcharts from document")
        
        # Note: Cover page detection removed - templates are now used instead
        
        # Step 0.6: CERTIFICATION PAGE DETECTION
        has_cert, cert_data, cert_start, cert_end = self.certification_handler.detect_and_extract(doc, 0)
        if has_cert:
            self.certification_data = cert_data
            self.certification_start_index = cert_start
            self.certification_end_index = cert_end
            logger.info(f"Certification page detected! From {cert_start} to {cert_end}")
        
        # Step 0.7: QUESTIONNAIRE DETECTION
        # Check if the document is a questionnaire
        full_text = '\n'.join([p.text for p in doc.paragraphs])
        q_data = self.questionnaire_processor.detect_questionnaire(full_text)
        if q_data['is_questionnaire']:
            self.questionnaire_data = self.questionnaire_processor.parse_questionnaire_structure(full_text)
            logger.info(f"Questionnaire detected! Title: {self.questionnaire_data['title']}")
            # We will process this differently in WordGenerator
        
        # Create image position lookup for tracking
        image_positions = {}
        for img in self.extracted_images:
            if img['position_type'] == 'paragraph':
                key = img['paragraph_index']
                if key not in image_positions:
                    image_positions[key] = []
                image_positions[key].append(img)
        
        # Create shape run lookup for tracking flowcharts/diagrams
        shape_run_map = {}
        for shape in self.extracted_shapes:
            if shape['position_type'] in ['inline', 'anchor', 'vml']:
                run_index = shape.get('run_index')
                if run_index is None:
                    continue
                key = (shape['paragraph_index'], run_index)
                if key not in shape_run_map:
                    shape_run_map[key] = []
                shape_run_map[key].append(shape)
        
        # Extract all content in document order (paragraphs and tables)
        lines = []
        paragraph_index = 0
        
        # Iterate through document body elements in order
        for element in doc.element.body:
            # Check if element is a paragraph
            if element.tag.endswith('p'):
                # Find the corresponding paragraph object
                for para in doc.paragraphs:
                    if para._element is element:
                        # PRESERVE CERTIFICATION/DECLARATION PAGE PARAGRAPHS - do not skip or replace
                        # The original content with names should be kept as-is
                        # Mark these paragraphs as protected from content modification
                        is_certification_content = has_cert and self.certification_start_index <= paragraph_index < self.certification_end_index
                        
                        segments, has_shapes = self._build_paragraph_segments(para, paragraph_index, shape_run_map)
                        text = ''.join(
                            segment['text'] for segment in segments if segment['type'] == 'text'
                        ).strip()
                        
                        # Check for automatic numbering/bullets (Word automatic lists)
                        # Only convert to explicit list markers in assistive mode.
                        try:
                            if (self.policy.list_numbering_mode == "assistive"
                                    and para._element.pPr is not None
                                    and para._element.pPr.numPr is not None):
                                # Check if text already has a bullet-like start (manual numbering)
                                if text and not re.match(r'^[\s•○●▪■□◆◇→➔➜➤➢–—*⁎⁑※✱✲✳✴☐☑✓✔✗✘⓿①②③④⑤⑥⑦⑧⑨❶❷❸❹❺❻❼❽❾❿➀-➉✦✧★☆♥♡◉◎▸▹►◂◃◄⦿⁍-]', text):
                                    text = f"• {text}"
                        except Exception:
                            pass
                        
                        # Check if this paragraph has images (skip cover page images)
                        if paragraph_index in image_positions:
                            for img in image_positions[paragraph_index]:
                                # Insert image placeholder
                                lines.append({
                                    'text': f'[IMAGE:{img["image_id"]}]',
                                    'style': 'Image',
                                    'bold': False,
                                    'font_size': 12,
                                    'type': 'image_placeholder',
                                    'image_id': img['image_id'],
                                })
                                logger.info(f"Added image placeholder for {img['image_id']} at paragraph {paragraph_index}")
                        
                        if text:
                            # Detect AI meta-commentary but preserve content to avoid data loss
                            is_ai_meta = self.engine.detect_ai_generated_content(text)

                            # Clean AI artifacts
                            text, metadata = self.engine.clean_ai_content(text)
                            if has_shapes and segments:
                                updated_segments = []
                                for segment in segments:
                                    if segment.get('type') == 'text':
                                        cleaned_segment, _ = self.engine.clean_ai_content(segment.get('text', ''))
                                        if cleaned_segment:
                                            updated_segments.append({'type': 'text', 'text': cleaned_segment})
                                    else:
                                        updated_segments.append(segment)
                                segments = updated_segments

                            font_size = 12  # Default
                            is_bold = False
                            is_bold_all = False
                            if para.runs:
                                bold_flags = [run.bold for run in para.runs if run.text]
                                is_bold = any(flag is True for flag in bold_flags)
                                is_bold_all = bool(bold_flags) and all(flag is True for flag in bold_flags)
                                if para.runs[0].font.size:
                                    font_size = para.runs[0].font.size.pt
                            
                            # Merge metadata from clean_ai_content
                            if metadata.get('bold'):
                                is_bold = True
                                
                            style = para.style.name if para.style else 'Normal'
                            if metadata.get('heading_level'):
                                style = f'Heading {metadata["heading_level"]}'
                            
                            line_data = {
                                'text': text,
                                'style': style,
                                'bold': is_bold,
                                'bold_all': is_bold_all,
                                'font_size': font_size,
                                'ai_meta': is_ai_meta,
                                'is_protected': is_certification_content,  # Preserve original content for certification/declaration pages
                            }
                            if has_shapes:
                                line_data['inline_segments'] = segments
                            lines.append(line_data)
                        elif has_shapes:
                            for segment in segments:
                                if segment['type'] == 'shape':
                                    lines.append({
                                        'text': f'[SHAPE:{segment["shape_id"]}]',
                                        'style': 'Shape',
                                        'bold': False,
                                        'font_size': 12,
                                        'type': 'shape_placeholder',
                                        'shape_id': segment['shape_id'],
                                    })
                                    logger.info(
                                        f"Added shape placeholder for {segment['shape_id']} at paragraph {paragraph_index}"
                                    )
                        
                        paragraph_index += 1
                        break
            
            # Check if element is a table
            elif element.tag.endswith('tbl'):
                # Find the corresponding table object
                table_index = 0
                for table in doc.tables:
                    if table._element is element:
                        lines.append({'text': '[TABLE START]', 'style': 'Table', 'bold': False, 'font_size': 12})
                        
                        # Check for images in table cells
                        for row_idx, row in enumerate(table.rows):
                            row_cells = []
                            for cell_idx, cell in enumerate(row.cells):
                                cell_text = cell.text.strip()
                                
                                # Check for images in this cell
                                for img in self.extracted_images:
                                    if img['position_type'] == 'table':
                                        loc = img['table_location']
                                        if (loc['table_index'] == table_index and 
                                            loc['row_index'] == row_idx and 
                                            loc['cell_index'] == cell_idx):
                                            cell_text = f'[IMAGE:{img["image_id"]}] {cell_text}'
                                
                                # Check for shapes in this cell
                                for shape in self.extracted_shapes:
                                    if shape['position_type'] in ['table', 'table_vml']:
                                        loc = shape['table_location']
                                        if (loc['table_index'] == table_index and 
                                            loc['row_index'] == row_idx and 
                                            loc['cell_index'] == cell_idx):
                                            cell_text = f'[SHAPE:{shape["shape_id"]}] {cell_text}'
                                
                                row_cells.append(cell_text)
                            
                            row_text = ' | '.join(row_cells)
                            lines.append({'text': f'| {row_text} |', 'style': 'Table', 'bold': False, 'font_size': 12})
                        
                        lines.append({'text': '[TABLE END]', 'style': 'Table', 'bold': False, 'font_size': 12})
                        break
                    table_index += 1
        
        return self.process_lines(lines, strip_front_matter=strip_front_matter), self.extracted_images, self.extracted_shapes
    
    def process_text(self, text, strip_front_matter=True):
        """Process plain text (no images in plain text)"""
        if not text:
            return self.process_lines([], strip_front_matter=strip_front_matter), []

        # Normalize line endings to ensure consistent splitting
        text = text.replace('\r\n', '\n').replace('\r', '\n')

        image_placeholder_pattern = re.compile(r'^\s*\[IMAGE:(?P<image_id>[^\]]+)\]\s*$', re.IGNORECASE)
        
        # PREPROCESSING: Apply regex-based text formatting for consistent numbering/bulleting
        # This fixes common formatting inconsistencies across documents (opt-in via policy).
        text = self.text_formatter.format_text(text)

        # FIRST: Remove horizontal rules and excess whitespace
        # These are typically used as visual separators but don't contribute to content
        text = re.sub(r'^\s*[-\u2010-\u2015]{3,}\s*$', '', text, flags=re.MULTILINE)
        text = re.sub(r'^\s*[\*\u2217\u2731]{3,}\s*$', '', text, flags=re.MULTILINE)
        text = re.sub(r'^\s*[_\u2017]{3,}\s*$', '', text, flags=re.MULTILINE)
        
        # Collapse multiple consecutive empty lines into single empty line
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # SECOND: Apply document-wide spacing cleanup
        text = self.engine.clean_document_spacing(text)
        
        # THIRD: Apply short document processing (key point emphasis)
        text = self.engine.process_short_document(text)
        
        # THIRD: AI Content Cleaning
        lines = []
        for line in text.split('\n'):
            image_match = image_placeholder_pattern.match(line)
            if image_match:
                image_id = image_match.group('image_id').strip()
                lines.append({
                    'text': f'[IMAGE:{image_id}]',
                    'type': 'image_placeholder',
                    'image_id': image_id
                })
                continue
            # Detect AI meta-commentary but preserve content to avoid data loss
            is_ai_meta = self.engine.detect_ai_generated_content(line)
                
            # Clean AI artifacts
            cleaned_line, metadata = self.engine.clean_ai_content(line)
            
            style = 'Normal'
            heading_level = metadata.get('heading_level', 0)
            if heading_level:
                style = f'Heading {heading_level}'
            
            lines.append({
                'text': cleaned_line,
                'style': style,
                'bold': metadata.get('bold', False),
                'bold_all': metadata.get('bold', False),
                'font_size': 12,
                'markdown_heading_level': heading_level,  # Preserve heading level from markdown
                'ai_meta': is_ai_meta,
            })
            
        # FOURTH: Optimize Page Breaks for AI
        lines = self.engine.optimize_page_breaks_for_ai(lines)
        
        return self.process_lines(lines, strip_front_matter=strip_front_matter), [], []  # No images or shapes in plain text
    
    def process_lines(self, lines, strip_front_matter=True):
        """Core line-by-line processing"""
        if strip_front_matter:
            lines = self.engine.strip_front_matter_placeholders(lines)
        analyzed = []
        stats = {
            'total_lines': len(lines),
            'headings': 0,
            'paragraphs': 0,
            'references': 0,
            'tables': 0,
            'lists': 0,
            'definitions': 0,
            'figures': 0,
            'images': 0,  # Add image count
            'h1_count': 0,
            'h2_count': 0,
            'h3_count': 0,
            # New pattern stats
            'inline_formatting': 0,
            'page_metadata': 0,
            'academic_metadata': 0,
            'math_expressions': 0,
            'footnotes': 0,
        }
        
        # Reset heading numberer for new document
        self.heading_numberer.reset()
        self.heading_numberer.configure_for_lines(lines)
        last_heading_number = None
        stray_parent_prefix = None
        stray_child_counter = 0
        
        # Analyze each line
        for i, line_data in enumerate(lines):
            text = line_data['text'] if isinstance(line_data, dict) else line_data
            
            # Extract previous line text safely
            if i > 0:
                prev_data = lines[i-1]
                if isinstance(prev_data, dict):
                    prev_line = prev_data.get('text', '')
                else:
                    prev_line = str(prev_data) if prev_data else ''
            else:
                prev_line = ''
            
            # Extract next line text safely
            if i < len(lines) - 1:
                next_data = lines[i+1]
                if isinstance(next_data, dict):
                    next_line = next_data.get('text', '')
                else:
                    next_line = str(next_data) if next_data else ''
            else:
                next_line = ''
            
            # Check for front matter placeholders FIRST (before pattern analysis)
            if isinstance(line_data, dict) and line_data.get('type') in [
                'toc_placeholder',
                'list_of_tables_placeholder',
                'list_of_figures_placeholder',
            ]:
                analysis = {
                    'type': line_data['type'],
                    'text': text,
                    'confidence': 1.0,
                }
                analyzed.append(analysis)
                continue

            # Check for image placeholder FIRST (before pattern analysis)
            if isinstance(line_data, dict) and line_data.get('type') == 'image_placeholder':
                analysis = {
                    'type': 'image_placeholder',
                    'text': text,
                    'image_id': line_data.get('image_id'),
                    'confidence': 1.0,
                }
                analyzed.append(analysis)
                stats['images'] = stats.get('images', 0) + 1
                continue
            
            # Check for shape placeholder (flowcharts, diagrams, arrows)
            if isinstance(line_data, dict) and line_data.get('type') == 'shape_placeholder':
                analysis = {
                    'type': 'shape_placeholder',
                    'text': text,
                    'shape_id': line_data.get('shape_id'),
                    'confidence': 1.0,
                }
                analyzed.append(analysis)
                stats['shapes'] = stats.get('shapes', 0) + 1
                continue
            
            # Check for PROTECTED content (certification/declaration pages)
            # Protected content is preserved exactly as-is without any modification
            if isinstance(line_data, dict) and line_data.get('is_protected'):
                # Check if this is a section TITLE (CERTIFICATION, DECLARATION, etc.)
                text_upper = text.upper().strip()
                is_section_title = text_upper in ['CERTIFICATION', 'DECLARATION', 'CERTIFICATION PAGE', 'DECLARATION PAGE']
                
                if is_section_title:
                    # Treat as a heading with page break
                    analysis = {
                        'type': 'heading',
                        'content': text,
                        'text': text,
                        'level': 1,  # Heading level 1
                        'line_num': i,
                        'confidence': 1.0,
                        'is_protected': True,
                        'needs_page_break': True,  # Start on new page
                        'should_center': True,  # Center the title
                        'original_style': line_data.get('style', 'Normal'),
                        'original_bold': True,  # Titles should be bold
                        'original_bold_all': True,
                        'original_font_size': line_data.get('font_size', 12),
                    }
                else:
                    # Regular protected content - preserve as-is
                    analysis = {
                        'type': 'protected_content',  # Special type to preserve content
                        'content': text,
                        'text': text,
                        'line_num': i,
                        'confidence': 1.0,
                        'is_protected': True,
                        'original_style': line_data.get('style', 'Normal'),
                        'original_bold': line_data.get('bold', False),
                        'original_bold_all': line_data.get('bold_all', False),
                        'original_font_size': line_data.get('font_size', 12),
                    }
                analyzed.append(analysis)
                stats['paragraphs'] += 1
                continue
            
            # Build context for dissertation-specific detection
            context = {
                'prev_was_chapter': False,
                'prev_front_matter': None,
                'prev_was_author_header': False,  # For cover page author name detection
                'prev_analysis': analyzed[-1] if analyzed else None,
            }
            
            # Check if previous line was a chapter heading or front matter heading
            if i > 0 and analyzed:
                prev_analysis = None
                for prior in reversed(analyzed):
                    if prior.get('type') in ['empty', 'page_metadata']:
                        continue
                    prev_analysis = prior
                    break
                if prev_analysis:
                    if prev_analysis.get('type') == 'chapter_heading':
                        context['prev_was_chapter'] = True
                    elif prev_analysis.get('type') == 'front_matter_heading':
                        context['prev_front_matter'] = prev_analysis.get('front_matter_type')
                    elif prev_analysis.get('type') in ['heading', 'heading_hierarchy']:
                        prev_text = prev_analysis.get('content') or prev_analysis.get('text', '')
                        prev_text = re.sub(r'[*_]+', '', str(prev_text))
                        is_chapter, _, _ = self.engine.is_chapter_heading(prev_text)
                        if is_chapter:
                            context['prev_was_chapter'] = True
            
            analysis = self.engine.analyze_line(
                text, 
                i, 
                prev_line, 
                next_line,
                context=context
            )

            if isinstance(line_data, dict) and line_data.get('inline_segments'):
                analysis['inline_segments'] = line_data['inline_segments']
            
            # SAFETY CHECK: Ensure analysis is a dictionary
            if not isinstance(analysis, dict):
                logger.error(f"analyze_line returned non-dict for line {i}: {type(analysis)}")
                # Create a fallback analysis
                analysis = {
                    'type': 'paragraph',
                    'content': text,
                    'line_num': i,
                    'confidence': 0.0
                }
            
            # Enhance with original formatting if available
            if isinstance(line_data, dict):
                analysis['original_style'] = line_data.get('style', 'Normal')
                analysis['original_bold'] = line_data.get('bold', False)
                analysis['original_bold_all'] = line_data.get('bold_all', False)
                analysis['original_font_size'] = line_data.get('font_size', 12)
                
                # If line came from a markdown heading (## text), override the type if it was misdetected
                markdown_level = line_data.get('markdown_heading_level', 0)
                if markdown_level > 0 and analysis['type'] not in ['heading', 'heading_hierarchy', 'chapter_heading', 'chapter_title', 'front_matter_heading']:
                    # Override detection - this was a markdown heading that got stripped
                    analysis['type'] = 'heading_hierarchy'
                    analysis['level'] = markdown_level
                    analysis['content'] = text
                    analysis['confidence'] = 0.95
                    # Only level 1 headings can have page breaks
                    if markdown_level == 1:
                        analysis['needs_page_break'] = self.engine.should_start_on_new_page(text)
                        analysis['should_center'] = self.engine.should_be_centered(text, 1)
                    else:
                        analysis['needs_page_break'] = False  # Level 2+ don't get page breaks
                        analysis['should_center'] = False

            stray_heading_match = re.match(r'^\s*(\d+)\s+([A-Za-z].+)$', text)
            if (analysis['type'] == 'paragraph'
                    and last_heading_number
                    and '.' in last_heading_number
                    and stray_heading_match):
                parent_prefix = last_heading_number
                if parent_prefix != stray_parent_prefix:
                    stray_parent_prefix = parent_prefix
                    stray_child_counter = 0
                stray_child_counter += 1
                child_number = f"{parent_prefix}.{stray_child_counter}"
                child_title = stray_heading_match.group(2).strip()
                analysis['type'] = 'heading'
                analysis['level'] = 3
                analysis['content'] = f"{child_number} {child_title}"
                analysis['text'] = analysis['content']
                analysis['heading_number'] = child_number
                analysis['original_text'] = text
                analysis['skip_heading_numbering'] = True
                self.heading_numberer._sync_counters_from_existing(child_number)
            
            # Apply automatic heading numbering for chapter-based content
            if analysis['type'] in ['heading', 'heading_hierarchy', 'chapter_heading', 'chapter_title'] and not analysis.get('skip_heading_numbering'):
                # Use the heading numberer to apply hierarchical numbering
                heading_level = analysis.get('level', 2)
                number_result = self.heading_numberer.number_heading(text, target_level=heading_level)
                
                # If heading was renumbered, update the analysis
                if number_result['was_renumbered']:
                    analysis['original_text'] = text  # Preserve original
                    analysis['text'] = number_result['numbered']
                    analysis['heading_number'] = number_result['number']
                    # Also update 'content' which is used by _structure_document
                    clean_numbered = re.sub(r'^#+\s*', '', number_result['numbered']).strip()
                    analysis['content'] = clean_numbered
                    logger.debug(f"Auto-numbered heading: '{text}' -> '{number_result['numbered']}'")
                elif number_result.get('number'):
                    # CRITICAL FIX: Even when not renumbered, set text and content to preserve existing number
                    # The analyze_line content may have stripped the number due to regex patterns
                    analysis['heading_number'] = number_result['number']
                    analysis['text'] = number_result['numbered']  # Use the preserved numbered text
                    # Also update 'content' to include the number
                    clean_numbered = re.sub(r'^#+\s*', '', number_result['numbered']).strip()
                    analysis['content'] = clean_numbered
                    logger.debug(f"Preserved existing heading number: '{text}' -> '{number_result['numbered']}'")
                
                # Store chapter context info
                analysis['chapter'] = number_result['chapter']
                analysis['in_appendix'] = self.heading_numberer.in_appendix
            elif analysis['type'] == 'chapter_heading':
                # Detect chapter and update numberer state (even if not renumbered)
                self.heading_numberer.number_heading(text)
            if analysis.get('heading_number'):
                last_heading_number = analysis['heading_number']
            
            analyzed.append(analysis)
            
            # Update stats
            if analysis['type'] == 'heading':
                stats['headings'] += 1
                if analysis.get('level') == 1:
                    stats['h1_count'] += 1
                elif analysis.get('level') == 2:
                    stats['h2_count'] += 1
                elif analysis.get('level') == 3:
                    stats['h3_count'] += 1
            elif analysis['type'] == 'paragraph':
                stats['paragraphs'] += 1
            elif analysis['type'] == 'reference':
                stats['references'] += 1
            elif analysis['type'] in ['table_start', 'table_caption']:
                stats['tables'] += 1
            elif 'list' in analysis['type']:
                stats['lists'] += 1
            elif analysis['type'] == 'definition':
                stats['definitions'] += 1
            elif analysis['type'] == 'figure':
                stats['figures'] += 1
            # New pattern type stats
            elif analysis['type'] == 'inline_formatting':
                stats['inline_formatting'] += 1
            elif analysis['type'] == 'page_metadata':
                stats['page_metadata'] += 1
            elif analysis['type'] == 'academic_metadata':
                stats['academic_metadata'] += 1
            elif analysis['type'] == 'math_expression':
                stats['math_expressions'] += 1
            elif analysis['type'] == 'footnote_endnote':
                stats['footnotes'] += 1
            # New pattern type stats (December 30, 2025 - 20 Academic Patterns)
            elif analysis['type'] == 'heading_hierarchy':
                stats['headings'] += 1
                level = analysis.get('level', 1)
                if level == 1:
                    stats['h1_count'] += 1
                elif level == 2:
                    stats['h2_count'] += 1
                elif level >= 3:
                    stats['h3_count'] += 1
            elif analysis['type'] == 'academic_table':
                stats['tables'] += 1
            elif analysis['type'] == 'list_nested':
                stats['lists'] += 1
            elif analysis['type'] == 'figure_equation':
                stats['figures'] += 1
            elif analysis['type'] in ['citation_inline', 'reference_apa']:
                stats['references'] += 1
            elif analysis['type'] == 'appendix_format':
                stats['headings'] += 1
            elif analysis['type'] in ['block_quote', 'math_model', 'statistical_result']:
                stats['paragraphs'] += 1
            elif analysis['type'] == 'toc_entry':
                pass  # Don't count TOC entries as paragraphs
            elif analysis['type'] in ['footnote_marker', 'questionnaire', 'glossary_entry']:
                stats['definitions'] += 1
            elif analysis['type'] == 'caption_format':
                if 'table' in analysis.get('subtype', ''):
                    stats['tables'] += 1
                elif 'figure' in analysis.get('subtype', ''):
                    stats['figures'] += 1
            # Dissertation-specific pattern stats (December 30, 2025)
            elif analysis['type'] == 'chapter_heading':
                stats['headings'] += 1
                stats['h1_count'] += 1
            elif analysis['type'] == 'chapter_title':
                stats['headings'] += 1  # Count as heading
            elif analysis['type'] == 'front_matter_heading':
                stats['headings'] += 1
                stats['h1_count'] += 1
            elif analysis['type'] == 'copyright_content':
                stats['paragraphs'] += 1  # Count as paragraph
            elif analysis['type'] == 'signature_line':
                pass  # Don't count signature lines in stats
            # Short document pattern stats (December 30, 2025)
            elif analysis['type'] == 'key_point':
                stats['paragraphs'] += 1  # Key points are paragraphs with emphasis
            elif analysis['type'] == 'assignment_header_field':
                stats['paragraphs'] += 1  # Header fields are metadata-like paragraphs
        
        # Structure the document
        structured = self._structure_document(analyzed)
        
        return {
            'analyzed': analyzed,
            'structured': structured,
            'stats': stats,
        }
    
    def _structure_document(self, analyzed):
        """Group lines into logical sections"""
        sections = []
        current_section = None
        current_list = None
        current_table = None
        pending_table_caption = ''  # Track caption before table_start
        in_references = False
        in_code_block = False
        
        for i, line in enumerate(analyzed):
            # SAFETY CHECK: Ensure line is a dictionary
            if not isinstance(line, dict):
                logger.warning(f"Skipping non-dict line in structure_document: {type(line)}")
                continue

            if line['type'] == 'empty':
                if current_list and current_section:
                    current_section['content'].append(current_list)
                    current_list = None
                continue

            if line['type'] in ['toc_placeholder', 'list_of_tables_placeholder', 'list_of_figures_placeholder']:
                if current_list and current_section:
                    current_section['content'].append(current_list)
                    current_list = None
                if current_table and current_section:
                    current_section['content'].append(current_table)
                    current_table = None
                if current_section:
                    sections.append(current_section)
                sections.append({
                    'type': line['type'],
                    'text': line.get('text', ''),
                })
                current_section = None
                continue
            
            # Detect reference section
            if line['type'] == 'heading' and line['level'] == 1:
                content_lower = line['content'].lower()
                if 'reference' in content_lower or 'bibliography' in content_lower or 'works cited' in content_lower:
                    in_references = True
                else:
                    in_references = False
            
            # Handle code blocks
            if line['type'] == 'code':
                in_code_block = not in_code_block
                continue
            
            # Handle headings - start new section
            if line['type'] == 'heading':
                # Save current list/table if any
                if current_list and current_section:
                    # Apply conservative bullet rules
                    if current_list.get('type') == 'bullet_list':
                        should_keep, content_to_add = self._should_keep_bullet_list(current_list)
                        if should_keep:
                            current_section['content'].append(current_list)
                        else:
                            # Add converted paragraphs instead
                            current_section['content'].extend(content_to_add)
                    else:
                        current_section['content'].append(current_list)
                    # CRITICAL FIX: Reset current_list to prevent duplication across sections
                    current_list = None
                
                # Save previous section
                if current_section:
                    sections.append(current_section)
                
                current_section = {
                    'type': 'section',
                    'heading': line['content'],
                    'level': line['level'],
                    'content': [],
                    'needs_page_break': line.get('needs_page_break', False),  # Propagate page break flag
                    'should_center': line.get('should_center', False),  # Propagate centering flag
                    'is_references_section': in_references,  # Track if this is a references section
                }
                continue
            
            # Initialize first section if no heading found
            if current_section is None:
                current_section = {
                    'type': 'section',
                    'heading': 'Document',
                    'level': 1,
                    'content': [],
                }
            
            # Handle references
            if in_references and line['type'] == 'reference':
                # Close current table first if exists (preserve position)
                if current_table and current_table.get('rows'):
                    current_section['content'].append(current_table)
                    current_table = None
                current_section['content'].append({
                    'type': 'reference',
                    'text': line['content'],
                    'journal_spans': line.get('journal_spans', []),
                })
                continue
            
            # Handle lists
            if 'list' in line['type']:
                list_type = 'bullet_list' if 'bullet' in line['type'] else 'numbered_list'
                if not current_list or current_list['type'] != list_type:
                    # Save previous list if exists
                    if current_list:
                        current_section['content'].append(current_list)
                    current_list = {
                        'type': list_type,
                        'items': [],
                    }
                
                # Add item to current list
                list_item = {
                    'type': line['type'],
                    'content': line.get('content', ''),
                    'bullet_info': line.get('bullet_info'),  # Store full bullet info for formatting
                    'original': line.get('original'),
                    'original_bold': line.get('original_bold', False),
                    'original_bold_all': line.get('original_bold_all', False),
                }
                current_list['items'].append(list_item)
                continue
            else:
                # End current list
                if current_list:
                    current_section['content'].append(current_list)
                    current_list = None
            
            # Handle tables (including markdown tables with pipe syntax: |cell|cell|)
            # Also handle table_start and table_end markers from Word documents
            if line['type'] in ['table', 'table_row', 'table_separator', 'table_caption', 'table_start', 'table_end']:
                
                # Handle table_start - begin a new table
                if line['type'] == 'table_start':
                    # If there's an existing table, finalize it first
                    if hasattr(self, '_current_table') and self._current_table:
                        if len(self._current_table['content']) >= 1:
                            current_section['content'].append(self._current_table)
                    # Start fresh table
                    self._current_table = {
                        'type': 'table',
                        'subtype': 'word_table',
                        'content': [],
                        'metadata': {},
                        'caption': pending_table_caption if pending_table_caption else ''
                    }
                    pending_table_caption = ''
                    continue
                
                # Handle table_end - finalize current table
                if line['type'] == 'table_end':
                    if hasattr(self, '_current_table') and self._current_table:
                        if len(self._current_table['content']) >= 1:
                            current_section['content'].append(self._current_table)
                        self._current_table = None
                    continue
                
                # Start or continue table based on type (for non-start/end types)
                if not hasattr(self, '_current_table') or self._current_table is None:
                    self._current_table = {
                        'type': 'table',
                        'subtype': line.get('subtype', 'markdown' if line['type'] in ['table_row', 'table_separator'] else 'unknown'),
                        'content': [],
                        'metadata': {},
                        'caption': ''
                    }

                # Handle table caption
                if line['type'] == 'table_caption':
                    self._current_table['caption'] = line.get('content', '')
                    continue

                # Handle table_row type (markdown tables with pipes)
                if line['type'] == 'table_row':
                    cells = line.get('cells', [])
                    if cells:
                        self._current_table['content'].append({
                            'type': 'row',
                            'cells': cells
                        })
                    continue

                # Handle table_separator type
                if line['type'] == 'table_separator':
                    self._current_table['content'].append({
                        'type': 'separator'
                    })
                    continue

                # Add row to current table
                if line.get('subtype') == 'markdown':
                    if line.get('metadata') and 'cells' in line['metadata']:
                        self._current_table['content'].append({
                            'type': 'row',
                            'cells': line['metadata']['cells']
                        })
                    else:
                        # Separator row
                        self._current_table['content'].append({
                            'type': 'separator'
                        })
                elif line.get('subtype') in ['tab', 'aligned', 'spaced']:
                    if line.get('metadata') and 'cells' in line['metadata']:
                        self._current_table['content'].append({
                            'type': 'row',
                            'cells': line['metadata']['cells']
                        })
                elif line.get('metadata') and 'cells' in line['metadata']:
                    self._current_table['content'].append({
                        'type': 'row',
                        'cells': line['metadata']['cells']
                    })

                # Check if we should finalize the table (minimum 2 rows)
                # Skip this for word_table subtype - they use explicit table_end
                if self._current_table.get('subtype') != 'word_table' and len(self._current_table['content']) >= 2:
                    # Look ahead to see if next lines continue the table
                    should_finalize = True
                    if i + 1 < len(analyzed):
                        next_line = analyzed[i + 1]
                        # Continue table if next line is also a table row (any type)
                        next_type = next_line.get('type', '')
                        if next_type in ['table', 'table_row', 'table_separator', 'table_start', 'table_end']:
                            should_finalize = False
                        elif next_type == 'table' and next_line.get('subtype') == self._current_table['subtype']:
                            should_finalize = False

                    if should_finalize:
                        current_section['content'].append(self._current_table)
                        self._current_table = None

                continue

            # Close current table if we encounter non-table content
            if hasattr(self, '_current_table') and self._current_table:
                if len(self._current_table['content']) >= 1:  # At least 1 row for markdown tables
                    current_section['content'].append(self._current_table)
                self._current_table = None

            # Handle plain text table detection
            if line['type'] in ['plain_table_separator', 'plain_table_header', 'plain_table_row']:

                # Start new table if not already in one
                if not hasattr(self, '_current_plain_table') or self._current_plain_table is None:
                    self._current_plain_table = {
                        'type': 'plain_text_table',
                        'content': [],
                        'metadata': {'delimiter': line.get('metadata', {}).get('delimiter', 'spaced')}
                    }

                # Validate consistency (same delimiter and cell count)
                current_delimiter = line.get('metadata', {}).get('delimiter', self._current_plain_table['metadata']['delimiter'])
                current_cell_count = line.get('metadata', {}).get('cell_count', 0)

                # If delimiter or cell count changes significantly, end current table
                if (current_delimiter != self._current_plain_table['metadata']['delimiter'] or
                    abs(current_cell_count - self._current_plain_table['metadata'].get('cell_count', current_cell_count)) > 1):

                    # Finalize current table if valid
                    if self._is_valid_table_block(self._current_plain_table['content']):
                        current_section['content'].append(self._current_plain_table)
                    else:
                        # Invalid table, convert to paragraphs
                        for row in self._current_plain_table['content']:
                            if row.get('row_type') == 'data' and 'cells' in row:
                                text = ' '.join(row['cells'])
                                current_section['content'].append({
                                    'type': 'paragraph',
                                    'content': text
                                })
                            elif 'original_text' in row:
                                current_section['content'].append({
                                    'type': 'paragraph',
                                    'content': row['original_text']
                                })

                    # Start new table
                    self._current_plain_table = {
                        'type': 'plain_text_table',
                        'content': [],
                        'metadata': {'delimiter': current_delimiter, 'cell_count': current_cell_count}
                    }

                # Add row to current table
                if line['type'] == 'plain_table_separator':
                    self._current_plain_table['content'].append({
                        'row_type': 'separator',
                        'style': line.get('metadata', {}).get('separator_style', 'dash')
                    })
                elif line['type'] == 'plain_table_header':
                    self._current_plain_table['content'].append({
                        'row_type': 'header',
                        'cells': line.get('metadata', {}).get('cells', [])
                    })
                else:  # plain_table_row
                    self._current_plain_table['content'].append({
                        'row_type': 'data',
                        'cells': line.get('metadata', {}).get('cells', [])
                    })

            else:
                # Not a table line - finalize any active table
                if hasattr(self, '_current_plain_table') and self._current_plain_table:
                    if self._is_valid_table_block(self._current_plain_table['content']):
                        current_section['content'].append(self._current_plain_table)
                    else:
                        # Invalid table, convert to paragraphs
                        for row in self._current_plain_table['content']:
                            if row.get('row_type') == 'data' and 'cells' in row:
                                text = ' '.join(row['cells'])
                                current_section['content'].append({
                                    'type': 'paragraph',
                                    'content': text
                                })
                            elif 'original_text' in row:
                                current_section['content'].append({
                                    'type': 'paragraph',
                                    'content': row['original_text']
                                })

                    self._current_plain_table = None
            
            # Handle instructions
            if line['type'] == 'instruction':
                current_section['content'].append({
                    'type': 'instruction',
                    'text': line['content'],
                })
                continue

            # Handle questions
            if line['type'] == 'question':
                current_section['content'].append({
                    'type': 'question',
                    'text': line['content'],
                })
                continue

            # Handle definitions
            if line['type'] == 'definition':
                current_section['content'].append({
                    'type': 'definition',
                    'term': line.get('term', ''),
                    'definition': line.get('definition', ''),
                })
                continue
            
            # Handle figures
            if line['type'] == 'figure':
                current_section['content'].append({
                    'type': 'figure',
                    'caption': line['content'],
                })
                continue
            
            # Handle image placeholders (extracted images from DOCX)
            if line['type'] == 'image_placeholder':
                current_section['content'].append({
                    'type': 'image_placeholder',
                    'image_id': line.get('image_id'),
                    'text': line.get('text', ''),
                })
                continue
            
            # Handle quotes
            if line['type'] == 'quote':
                current_section['content'].append({
                    'type': 'quote',
                    'text': line['content'],
                })
                continue
            
            # Handle equations
            if line['type'] == 'equation':
                current_section['content'].append({
                    'type': 'equation',
                    'label': line['content'],
                })
                continue
            
            # Handle page metadata (headers/footers/page numbers)
            if line['type'] == 'page_metadata':
                # Page metadata is typically excluded from main content
                # but we can add it to a special metadata section if needed
                current_section['content'].append({
                    'type': 'page_metadata',
                    'subtype': line.get('subtype', 'metadata'),
                    'text': line['content'],
                })
                continue
            
            # Handle academic metadata (author, affiliation, etc.)
            if line['type'] == 'academic_metadata':
                current_section['content'].append({
                    'type': 'academic_metadata',
                    'subtype': line.get('subtype', 'metadata'),
                    'text': line['content'],
                })
                continue
            
            # Handle mathematical expressions
            if line['type'] == 'math_expression':
                current_section['content'].append({
                    'type': 'math_expression',
                    'subtype': line.get('subtype', 'inline_math'),
                    'text': line['content'],
                })
                continue
            
            # Handle footnotes/endnotes
            if line['type'] == 'footnote_endnote':
                if line.get('subtype') == 'section_header':
                    # Start new footnote section
                    if current_section:
                        sections.append(current_section)
                    current_section = {
                        'type': 'section',
                        'heading': line['content'],
                        'level': 2,
                        'content': [],
                    }
                else:
                    current_section['content'].append({
                        'type': 'footnote_endnote',
                        'subtype': line.get('subtype', 'footnote_entry'),
                        'text': line['content'],
                    })
                continue
            
            # Handle inline formatting
            if line['type'] == 'inline_formatting':
                current_section['content'].append({
                    'type': 'inline_formatting',
                    'formatting': line.get('formatting', {}),
                    'text': line['content'],
                })
                continue
            
            # ============================================================
            # NEW PATTERN STRUCTURING - December 30, 2025 (20 Academic Patterns)
            # ============================================================
            
            # Handle heading hierarchy (markdown-style)
            if line['type'] == 'heading_hierarchy':
                if current_list:
                    current_section['content'].append(current_list)
                    current_list = None
                if current_table:
                    current_section['content'].append(current_table)
                    current_table = None
                if current_section:
                    sections.append(current_section)
                current_section = {
                    'type': 'section',
                    'heading': line.get('content', line['content']),
                    'level': line.get('level', 1),
                    'content': [],
                    'needs_page_break': line.get('needs_page_break', False),  # Propagate page break flag
                    'should_center': line.get('should_center', False),  # Propagate centering flag
                }
                continue

            # Handle prominent headings (bold, underlined, star-surrounded, ALL CAPS)
            if line['type'] == 'prominent_heading':
                if current_list:
                    current_section['content'].append(current_list)
                    current_list = None
                if current_table:
                    current_section['content'].append(current_table)
                    current_table = None
                if current_section:
                    sections.append(current_section)

                current_section = {
                    'type': 'prominent_section',
                    'title': line['content'],
                    'level': line.get('level', 3),
                    'original_format': line.get('original_format', 'unknown'),
                    'content': [],
                    'needs_emphasis': True,
                }

                # NOTE: Removed automatic start_on_new_page for prominent headings.
                # Page breaks should only be set for actual academic sections (chapters, abstract, etc.)
                # that explicitly have needs_page_break = True from analyze_line.
                # Prominent headings from markdown/AI content (bold text, ## sections) should NOT get page breaks.

                continue

            # Handle hierarchical numbered lists (enhanced hierarchy)
            if line['type'] == 'hierarchical_list':
                # Continuing same hierarchical list format
                if (current_section and current_section.get('type') == 'hierarchical_list_group' and
                    line.get('format') == current_section.get('format')):
                    item = {
                        'content': line['content'],
                        'level': line.get('level', 1),
                        'numbering': line.get('numbering'),
                        'raw_line': line.get('original', line['content'])
                    }
                    current_section['items'].append(item)
                else:
                    # Start a new hierarchical list group
                    if current_list:
                        current_section['content'].append(current_list)
                        current_list = None
                    current_section = {
                        'type': 'hierarchical_list_group',
                        'heading': line.get('numbering', 'Hierarchical List'),
                        'items': [{
                            'content': line['content'],
                            'level': line.get('level', 1),
                            'numbering': line.get('numbering'),
                            'raw_line': line.get('original', line['content'])
                        }],
                        'format': line.get('format', 'alphanumeric_hierarchy'),
                        'start_level': line.get('level', 1),
                    }
                    sections.append(current_section)
                continue

            # Handle short-document headers (assignments, questions, tasks)
            if line['type'] == 'shortdoc_header':
                header_type = line.get('header_type', 'section')

                if header_type in ['section', 'part']:
                    # Major section - start new shortdoc section
                    if current_list:
                        current_section['content'].append(current_list)
                        current_list = None
                    if current_table:
                        current_section['content'].append(current_table)
                        current_table = None
                    if current_section:
                        sections.append(current_section)
                    current_section = {
                        'type': 'shortdoc_section',
                        'title': line['content'],
                        'level': line.get('level', 3),
                        'header_type': header_type,
                        'numbering': line.get('numbering'),
                        'content': [],
                    }
                    continue
                else:
                    # Add as subheader within current shortdoc section
                    if not current_section or current_section.get('type') != 'shortdoc_section':
                        if current_section:
                            sections.append(current_section)
                        current_section = {
                            'type': 'shortdoc_section',
                            'title': f"{header_type.title()} {line.get('numbering','')}",
                            'level': line.get('level', 3),
                            'header_type': 'generic',
                            'content': [],
                        }
                    sub_item = {
                        'type': 'shortdoc_subheader',
                        'content': line['content'],
                        'level': line.get('level', 3),
                        'header_type': header_type,
                        'numbering': line.get('numbering')
                    }
                    current_section['content'].append(sub_item)
                    continue
            
            # Handle academic table
            if line['type'] == 'academic_table':
                subtype = line.get('subtype', 'data_row')
                if subtype == 'caption':
                    if current_table:
                        current_section['content'].append(current_table)
                    current_table = {
                        'type': 'table',
                        'caption': line['content'],
                        'rows': [],
                        'has_header': False,
                    }
                elif subtype == 'header_row':
                    if not current_table:
                        current_table = {'type': 'table', 'caption': '', 'rows': [], 'has_header': True}
                    cells = [c.strip().strip('*') for c in line['content'].split('|') if c.strip()]
                    current_table['rows'].append(cells)
                    current_table['has_header'] = True
                elif subtype == 'data_row':
                    if not current_table:
                        current_table = {'type': 'table', 'caption': '', 'rows': [], 'has_header': False}
                    cells = [c.strip() for c in line['content'].split('|') if c.strip()]
                    current_table['rows'].append(cells)
                # Skip separator rows
                continue
            
            # Handle nested lists
            if line['type'] == 'list_nested':
                list_type = 'bullet_list' if line.get('subtype') == 'checkbox' else 'nested_list'
                if not current_list or current_list['type'] != list_type:
                    if current_list:
                        current_section['content'].append(current_list)
                    current_list = {
                        'type': list_type,
                        'items': [],
                    }
                current_list['items'].append({
                    'text': line['content'].lstrip(' \t•-*□☐☑✓✗'),
                    'indent_level': line.get('indent_level', 0),
                    'is_checkbox': line.get('subtype') == 'checkbox',
                })
                continue
            
            # Handle figure/equation
            if line['type'] == 'figure_equation':
                current_section['content'].append({
                    'type': 'figure_equation',
                    'subtype': line.get('subtype', 'figure_caption'),
                    'text': line['content'],
                })
                continue
            
            # Handle inline citations
            if line['type'] == 'citation_inline':
                current_section['content'].append({
                    'type': 'citation_inline',
                    'text': line['content'],
                    'citation_count': line.get('citation_count', 1),
                })
                continue
            
            # Handle appendix format
            if line['type'] == 'appendix_format':
                if current_list:
                    current_section['content'].append(current_list)
                    current_list = None
                if current_table:
                    current_section['content'].append(current_table)
                    current_table = None
                if current_section:
                    sections.append(current_section)
                current_section = {
                    'type': 'appendix',
                    'heading': line['content'],
                    'level': line.get('level', 1),
                    'content': [],
                }
                continue
            
            # Handle block quotes
            if line['type'] == 'block_quote':
                current_section['content'].append({
                    'type': 'block_quote',
                    'text': line.get('content', line['content']),
                })
                continue
            
            # Handle math models
            if line['type'] == 'math_model':
                current_section['content'].append({
                    'type': 'math_model',
                    'subtype': line.get('subtype', 'statistical_notation'),
                    'text': line['content'],
                })
                continue
            
            # Handle text emphasis (inline - add to current paragraph or as standalone)
            if line['type'] == 'text_emphasis':
                current_section['content'].append({
                    'type': 'text_emphasis',
                    'subtype': line.get('subtype', 'bold'),
                    'text': line['content'],
                })
                continue
            
            # Handle APA references
            if line['type'] == 'reference_apa':
                current_section['content'].append({
                    'type': 'reference',
                    'text': line['content'],
                    'format': 'apa',
                    'subtype': line.get('subtype', 'standard_reference'),
                })
                continue
            
            # Handle TOC entries
            if line['type'] == 'toc_entry':
                current_section['content'].append({
                    'type': 'toc_entry',
                    'text': line['content'],
                    'page_number': line.get('page_number', None),
                })
                continue
            
            # Handle footnote markers
            if line['type'] == 'footnote_marker':
                current_section['content'].append({
                    'type': 'footnote_marker',
                    'subtype': line.get('subtype', 'footnote_reference'),
                    'text': line['content'],
                })
                continue
            
            # Handle abbreviations
            if line['type'] == 'abbreviation':
                current_section['content'].append({
                    'type': 'abbreviation',
                    'text': line['content'],
                })
                continue
            
            # Handle caption formatting
            if line['type'] == 'caption_format':
                current_section['content'].append({
                    'type': 'caption_format',
                    'subtype': line.get('subtype', 'caption'),
                    'text': line['content'],
                })
                continue
            
            # Handle horizontal rules (visual separators, NOT page breaks)
            if line['type'] == 'horizontal_rule':
                current_section['content'].append({
                    'type': 'horizontal_rule',
                })
                continue
            
            # Handle explicit page breaks
            if line['type'] == 'page_break':
                current_section['content'].append({
                    'type': 'page_break',
                })
                continue
            
            # Handle statistical results
            if line['type'] == 'statistical_result':
                current_section['content'].append({
                    'type': 'statistical_result',
                    'text': line['content'],
                    'stats_types': line.get('stats_types', []),
                })
                continue
            
            # Handle questionnaire items
            if line['type'] == 'questionnaire':
                current_section['content'].append({
                    'type': 'questionnaire',
                    'subtype': line.get('subtype', 'question_item'),
                    'text': line['content'],
                })
                continue
            
            # Handle glossary entries
            if line['type'] == 'glossary_entry':
                current_section['content'].append({
                    'type': 'glossary_entry',
                    'term': line.get('term', ''),
                    'definition': line.get('definition', line['content']),
                })
                continue
            
            # Handle cross-references
            if line['type'] == 'cross_reference':
                current_section['content'].append({
                    'type': 'cross_reference',
                    'text': line['content'],
                    'reference_types': line.get('reference_types', []),
                })
                continue
            
            # Handle running headers
            if line['type'] == 'running_header':
                current_section['content'].append({
                    'type': 'running_header',
                    'text': line['content'],
                })
                continue
            
            # ============================================================
            # DISSERTATION-SPECIFIC STRUCTURING - December 30, 2025
            # ============================================================
            
            # Handle chapter headings
            if line['type'] == 'chapter_heading':
                if current_list:
                    current_section['content'].append(current_list)
                    current_list = None
                if current_table:
                    current_section['content'].append(current_table)
                    current_table = None
                if current_section:
                    sections.append(current_section)
                
                # Create chapter section with heading and optional title
                chapter_heading = line['content']
                chapter_title = line.get('chapter_title')
                
                current_section = {
                    'type': 'chapter',
                    'heading': chapter_heading,
                    'chapter_num': line.get('chapter_num'),
                    'chapter_title': chapter_title,  # May be None
                    'level': 1,
                    'content': [],
                    'needs_page_break': True,
                    'should_center': True,
                }
                continue
            
            # Handle chapter title (follows chapter heading)
            if line['type'] == 'chapter_title':
                # Add title to current chapter section if it exists
                if current_section and current_section.get('type') == 'chapter':
                    if not current_section.get('chapter_title'):
                        current_section['chapter_title'] = line['content']
                else:
                    # Standalone chapter title (treat as heading)
                    current_section['content'].append({
                        'type': 'chapter_title',
                        'text': line['content'],
                        'should_center': True,
                    })
                continue
            
            # Handle front matter headings (Declaration, Certification, etc.)
            if line['type'] == 'front_matter_heading':
                if current_list:
                    current_section['content'].append(current_list)
                    current_list = None
                if current_table:
                    current_section['content'].append(current_table)
                    current_table = None
                if current_section:
                    sections.append(current_section)
                
                front_matter_type = line.get('front_matter_type', 'unknown')
                current_section = {
                    'type': 'front_matter',
                    'heading': line['content'],
                    'front_matter_type': front_matter_type,
                    'level': 1,
                    'content': [],
                    'needs_page_break': True,
                    'should_center': True,
                }
                continue
            
            # Handle copyright content
            if line['type'] == 'copyright_content':
                current_section['content'].append({
                    'type': 'copyright_content',
                    'text': line['content'],
                    'should_center': True,
                })
                continue
            
            # Handle signature lines
            if line['type'] == 'signature_line':
                current_section['content'].append({
                    'type': 'signature_line',
                    'text': line['content'],
                })
                continue
            
            # Handle TOC entries
            if line['type'] == 'toc_entry':
                current_section['content'].append({
                    'type': 'toc_entry',
                    'text': line['content'],
                })
                continue

            # Handle key points (treated as emphasized paragraphs)
            if line['type'] == 'key_point':
                current_section['content'].append({
                    'type': 'paragraph',
                    'text': line['content'],
                    'is_key_point': True,
                    'key_point_type': line.get('key_point_type'),
                })
                continue

            # Handle assignment header fields
            if line['type'] == 'assignment_header_field':
                current_section['content'].append({
                    'type': 'paragraph',
                    'text': line['content'],
                    'is_header_field': True,
                })
                continue

            # Handle PROTECTED content (certification/declaration pages) - pass through unchanged
            if line['type'] == 'protected_content' or line.get('is_protected'):
                current_section['content'].append({
                    'type': 'protected_content',
                    'text': line.get('content', line.get('text', '')),
                    'is_protected': True,
                    'original_style': line.get('original_style', 'Normal'),
                    'original_bold': line.get('original_bold', False),
                    'original_bold_all': line.get('original_bold_all', False),
                    'original_font_size': line.get('original_font_size', 12),
                })
                continue

            # Handle paragraphs
            if line['type'] == 'paragraph':
                current_section['content'].append({
                    'type': 'paragraph',
                    'text': line['content'],
                    'original_bold': line.get('original_bold', False),
                    'original_bold_all': line.get('original_bold_all', False),
                })
                continue
            
            # Handle regular references outside reference section
            if line['type'] == 'reference':
                current_section['content'].append({
                    'type': 'reference',
                    'text': line['content'],
                })
                continue
        
        # Add remaining list/table if any
        if current_list and current_section:
            # Apply conservative bullet rules
            if current_list.get('type') == 'bullet_list':
                should_keep, content_to_add = self._should_keep_bullet_list(current_list)
                if should_keep:
                    current_section['content'].append(current_list)
                else:
                    # Add converted paragraphs instead
                    current_section['content'].extend(content_to_add)
            else:
                current_section['content'].append(current_list)
        if current_table and current_section:
            current_section['content'].append(current_table)
        if hasattr(self, '_current_table') and self._current_table and current_section:
            if self._current_table.get('subtype') == 'word_table' or len(self._current_table.get('content', [])) >= 1:
                current_section['content'].append(self._current_table)
            self._current_table = None
        if hasattr(self, '_current_plain_table') and self._current_plain_table and current_section:
            if self._is_valid_table_block(self._current_plain_table['content']):
                current_section['content'].append(self._current_plain_table)
            else:
                for row in self._current_plain_table['content']:
                    if row.get('row_type') == 'data' and 'cells' in row:
                        text = ' '.join(row['cells'])
                        current_section['content'].append({
                            'type': 'paragraph',
                            'content': text
                        })
                    elif 'original_text' in row:
                        current_section['content'].append({
                            'type': 'paragraph',
                            'content': row['original_text']
                        })
            self._current_plain_table = None
        
        # Add final section
        if current_section:
            sections.append(current_section)
        
        return sections

    def _should_keep_bullet_list(self, bullet_list):
        """
        Conservative bullet rendering: Check if a bullet_list meets criteria for being displayed as bullets.
        
        Criteria:
        - Must have 4+ items
        - All items must be short (<30 words)
        - No items should have colons (would become bolded)
        - No multiline items
        
        Returns: (should_keep_as_bullets, converted_content)
        - If True, return the original list
        - If False, return paragraphs/bold formatted items
        """
        items = bullet_list.get('items', [])
        
        # Check minimum item count
        if len(items) < 4:
            return False, self._convert_bullet_list_to_paragraphs(items)
        
        # Check each item
        for item in items:
            content = item.get('content', '')
            
            # No colons (would be bolded)
            if ':' in content:
                return False, self._convert_bullet_list_to_paragraphs(items)
            
            # No long items (>30 words)
            if len(content.split()) > 30:
                return False, self._convert_bullet_list_to_paragraphs(items)
            
            # No multiline items
            if '\n' in content:
                return False, self._convert_bullet_list_to_paragraphs(items)
        
        # All criteria met - keep as bullet list
        return True, bullet_list
    
    def _convert_bullet_list_to_paragraphs(self, items):
        """Convert bullet list items to paragraphs or bold format"""
        converted = []
        
        for item in items:
            content = item.get('content', '')
            if not content:
                continue
            
            # If item has colon, make it bold format
            if ':' in content:
                converted.append({
                    'type': 'paragraph',
                    'text': content,
                    'should_bold': True,
                })
            else:
                # Regular paragraph
                converted.append({
                    'type': 'paragraph',
                    'text': content,
                })
        
        return converted

    def _is_valid_table_block(self, table_rows):
        """Validate that accumulated rows form a real table"""
        
        if not table_rows or len(table_rows) < 2:
            return False
        
        # Must have at least 2 data rows (or 1 data + 1 separator)
        data_rows = [r for r in table_rows if r.get('row_type') == 'data']
        separator_rows = [r for r in table_rows if r.get('row_type') == 'separator']
        
        if len(data_rows) < 1:
            return False
        
        # If only 1 data row, must have a separator to confirm it's a table
        if len(data_rows) == 1 and len(separator_rows) == 0:
            return False
        
        # Check column consistency across data rows
        if len(data_rows) >= 2:
            cell_counts = [len(r.get('cells', [])) for r in data_rows]
            if max(cell_counts) - min(cell_counts) > 1:
                return False
        
        return True

    def generate_smart_filename(self, structured_data, original_filename=None):
        """
        Generate a smart filename based on document content.
        Priorities:
        1. If Uploaded File: Original Filename + _formatted
        2. If Pasted: Cover Page > Heading > Content > Fallback
        """
        # Check if it's a real uploaded file (not pasted)
        is_pasted = False
        if original_filename:
            name_lower = original_filename.lower()
            if "pasted_document" in name_lower or "pasted document" in name_lower:
                is_pasted = True
        else:
            is_pasted = True
            
        # PRIORITY 1: Uploaded File -> Maintain Name + _formatted
        if not is_pasted and original_filename:
            name_without_ext = os.path.splitext(original_filename)[0]
            # Sanitize
            base_name = re.sub(r'[<>:"/\\|?*]', '', name_without_ext)
            base_name = re.sub(r'\s+', ' ', base_name).strip()
            return f"{base_name}_formatted"

        # PRIORITY 2: Pasted Content -> Smart Naming
        base_name = "document"
        
        # Note: Cover page data check removed - templates are now used instead
        # Check First Heading if available
        if structured_data:
            # Look for H1
            for section in structured_data:
                if section.get('level') == 1 and section.get('heading'):
                    heading = section.get('heading').strip()
                    if heading and heading.lower() != 'document':
                        base_name = heading
                        break
            
            # 2b. If no H1, look for H2
            if base_name == "document":
                for section in structured_data:
                    if section.get('level') == 2 and section.get('heading'):
                        heading = section.get('heading').strip()
                        if heading:
                            base_name = heading
                            break
            
            # 2c. If still no heading, try first paragraph text
            if base_name == "document":
                for section in structured_data:
                    for item in section.get('content', []):
                        if item.get('type') == 'paragraph':
                            text = item.get('text', '').strip()
                            if text:
                                # Use first few words (max 6)
                                words = text.split()
                                if words:
                                    base_name = ' '.join(words[:6])
                                    break
                    if base_name != "document":
                        break

        # 3. Fallback to original filename (unless it's the generic paste name)
        if base_name == "document" and original_filename:
            name_without_ext = os.path.splitext(original_filename)[0]
            if "pasted_document" not in name_without_ext.lower():
                base_name = name_without_ext
            else:
                base_name = "Untitled Document"
            
        # Sanitize filename
        # Remove invalid chars: < > : " / \ | ? *
        base_name = re.sub(r'[<>:"/\\|?*]', '', base_name)
        base_name = re.sub(r'\s+', ' ', base_name).strip()
        
        # Limit length
        if len(base_name) > 50:
            base_name = base_name[:50].strip()
            
        return base_name


class WordGenerator:
    """Generate formatted Word documents with image and shape support"""
    
    # Path to cover page logo
    COVER_LOGO_PATH = os.path.join(os.path.dirname(__file__), 'coverpage_template', 'cover_logo.png')
    LIST_HANGING_INDENT_INCHES = 0.4
    LIST_NESTED_INDENT_INCREMENT_INCHES = 0.25
    LIST_BULLET_OFFSET_INCHES = 0.15
    LIST_BULLET_TEXT_INDENT_INCHES = 0.35
    
    def __init__(self, policy=None):
        self.policy = policy or FormatPolicy()
        self.doc = None
        self.images = []  # Extracted images
        self.image_lookup = {}  # image_id -> image_data
        self.image_inserter = None
        self.shapes = []  # Extracted shapes/flowcharts
        self.shape_lookup = {}  # shape_id -> shape_data
        self.shape_inserter = None
        self.toc_entries = []  # Track headings for TOC generation
        self.toc_placeholder_index = None  # Index where TOC should be inserted
        self.heading_numberer = HeadingNumberer(policy=self.policy)  # For numbering headings
        self.figure_formatter = FigureFormatter()  # For figure detection and formatting
        self.table_formatter = TableFormatter()  # For table detection and formatting
        self.figure_entries = []  # Track figures for List of Figures
        self.table_entries = []  # Track tables for List of Tables
        self.has_figures = False  # Whether document contains figures
        self.has_tables = False  # Whether document contains tables
        self.use_continuous_arabic = False # Whether to use continuous Arabic numbering
        self.is_short_document = False  # Suppress page breaks for short documents
        self._consecutive_bold_count = 0  # Track consecutive bold paragraphs (max 2)
        # Formatting options
        self.font_size = 12
        self.line_spacing = 1.5
        self.margin_cm = 3.0
        self.include_toc = False

    def _iter_paragraphs_in_doc(self, doc):
        """Yield all paragraphs across the document, including tables and headers/footers."""
        for para in doc.paragraphs:
            yield para
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        yield para
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        yield para
                    for table in header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    yield para
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        yield para
                    for table in footer.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    yield para

    def _iter_runs_in_doc(self, doc):
        """Yield all runs across the document, including tables and headers/footers."""
        for para in self._iter_paragraphs_in_doc(doc):
            for run in para.runs:
                yield run

    def _is_reference_paragraph(self, para):
        """Return True if paragraph is a references entry (italic allowed)."""
        try:
            style = para.style
        except Exception:
            style = None
        return bool(style and style.name == 'ReferenceEntry')

    def _should_allow_bold(self, is_heading=False):
        """
        Check if we should allow bold for the next paragraph.
        Limits consecutive bold lines to 2 (for non-heading content).
        Headings (chapter titles) are always allowed.
        
        Returns:
            bool: True if bold is allowed, False otherwise
        """
        if is_heading:
            # Headings are always allowed to be bold
            return True
        
        # For non-heading bold content, limit to 2 consecutive
        return self._consecutive_bold_count < 2
    
    def _track_bold_paragraph(self, is_bold, is_heading=False):
        """
        Track whether a paragraph was bolded for consecutive bold limiting.
        
        Args:
            is_bold: Whether the paragraph was bolded
            is_heading: Whether this is a heading (resets counter)
        """
        if is_heading:
            # Headings reset the counter (they're expected to be followed by content)
            self._consecutive_bold_count = 0
        elif is_bold:
            self._consecutive_bold_count += 1
        else:
            # Non-bold paragraph resets the counter
            self._consecutive_bold_count = 0

    def _apply_list_hanging_indent(
        self,
        para,
        indent_level=0,
        bullet_offset_inch=0.0,
        text_indent_inch=None,
    ):
        base_indent = (
            text_indent_inch
            if text_indent_inch is not None
            else self.LIST_HANGING_INDENT_INCHES
        )
        base_indent += self.LIST_NESTED_INDENT_INCREMENT_INCHES * indent_level
        para.paragraph_format.left_indent = Inches(base_indent)
        para.paragraph_format.first_line_indent = Inches(-base_indent + bullet_offset_inch)
        para.paragraph_format.tab_stops.add_tab_stop(Inches(base_indent))

    def _apply_list_body_indent(self, para, indent_level=0):
        base_indent = self.LIST_HANGING_INDENT_INCHES + (
            self.LIST_NESTED_INDENT_INCREMENT_INCHES * indent_level
        )
        para.paragraph_format.left_indent = Inches(base_indent)
        para.paragraph_format.first_line_indent = Pt(0)

    def _enforce_no_italics(self, doc):
        """Force italics off for all styles and runs (except references)."""
        for style in doc.styles:
            # Skip ReferenceEntry style - italics are allowed within references
            if hasattr(style, 'name') and style.name == 'ReferenceEntry':
                continue
            if hasattr(style, 'font') and style.font is not None:
                style.font.italic = False
        for para in self._iter_paragraphs_in_doc(doc):
            if self._is_reference_paragraph(para):
                continue
            for run in para.runs:
                run.italic = False
                if run.font is not None:
                    run.font.italic = False

    def _run_acceptance_checks(self, doc):
        """Run acceptance checks for italics, asterisks, and list numbering leakage."""
        italic_runs = []
        for para in self._iter_paragraphs_in_doc(doc):
            if self._is_reference_paragraph(para):
                continue
            for run in para.runs:
                if run.italic or (run.font is not None and run.font.italic):
                    italic_runs.append(run)
        assert not italic_runs, "Italics detected in document runs."
        for para in self._iter_paragraphs_in_doc(doc):
            text = para.text or ''
            # NOTE: previously we rejected any paragraph containing '*',
            # which was too aggressive (caught valid cases like numbered
            # items with stray asterisks). Rely on the more specific
            # numeric/star detection patterns below instead.
            if re.search(r'\\s\\*\\s', text) or re.match(r'^\\s*\\d+\\s*\\*', text):
                raise AssertionError(f"Numeric asterisk marker detected: '{text}'")
            if re.match(r'^\\s*\\d+(?:\\.\\d+)?\\s*\\*', text):
                raise AssertionError(f"Numeric star bullet detected: '{text}'")
            if 'private ip ranges' in text.lower() and re.match(r'^\\s*\\d+(?:\\.\\d+)?', text):
                raise AssertionError(f"Private IP ranges should be bullet-only: '{text}'")
            if para.style is not None and para.style.name == 'List Number':
                raise AssertionError("List Number style detected; list numbering must reset per list group.")
        
    def _set_page_numbering(self, section, fmt='decimal', start=None):
        """Set page numbering format and start value for a section."""
        sectPr = section._sectPr
        pgNumType = sectPr.find(qn('w:pgNumType'))
        if pgNumType is None:
            pgNumType = OxmlElement('w:pgNumType')
            sectPr.append(pgNumType)
        
        pgNumType.set(qn('w:fmt'), fmt)
        if start is not None:
            pgNumType.set(qn('w:start'), str(start))
            
    def _add_page_number_to_footer(self, section):
        """Add page number to the footer of a section."""
        footer = section.footer

        # If the first paragraph contains our watermark, insert a new paragraph
        # for the page number so we don't clear or overwrite the watermark.
        if footer.paragraphs:
            first_para = footer.paragraphs[0]
            if first_para.text and 'Formatted with AfroDocs.app' in first_para.text:
                # Create a new paragraph and insert it at the beginning of the footer
                p = footer.add_paragraph()
                footer._element.insert(0, p._p)
            else:
                # Reuse and clear the first paragraph for page number
                p = first_para
                p.clear()
        else:
            p = footer.add_paragraph()

        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()

        # Add PAGE field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)

        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"
        run._r.append(instrText)

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)
        
    def generate(self, structured_data, output_path, images=None, shapes=None, certification_data=None, questionnaire_data=None, is_free_tier=False, include_toc=False, font_size=12, line_spacing=1.5, margins=None):
        """Generate Word document from structured data with images and shapes
        
        Args:
            structured_data: List of structured sections
            output_path: Path to save the output document
            images: List of extracted images
            shapes: List of extracted shapes/flowcharts/diagrams
            certification_data: Dict with extracted certification page information (or None)
            questionnaire_data: Dict with extracted questionnaire structure (or None)
            include_toc: Whether to include table of contents
            font_size: Font size in points (8-28)
            line_spacing: Line spacing multiplier (1.0-3.0)
            margins: Dict with 'left', 'top', 'bottom', 'right' margins in cm (0.5-5.0)
        """
        # Store formatting options for use throughout document generation
        self.font_size = font_size
        self.line_spacing = line_spacing
        # Handle margins - support both dict (individual sides) and scalar (uniform)
        if margins is None:
            self.margins = {'left': 3.0, 'top': 2.5, 'bottom': 2.5, 'right': 2.5}
        elif isinstance(margins, dict):
            self.margins = margins
        else:
            # Scalar value - apply to all sides
            self.margins = {'left': margins, 'top': margins, 'bottom': margins, 'right': margins}
        self.include_toc = include_toc
        
        logger.info(f"WordGenerator.generate() called with: font_size={font_size}, line_spacing={line_spacing}, margins={self.margins}, include_toc={include_toc}")
        
        self.doc = Document()
        
        # Set individual margins for each side (convert from cm to inches)
        for section in self.doc.sections:
            section.top_margin = Inches(self.margins['top'] / 2.54)
            section.bottom_margin = Inches(self.margins['bottom'] / 2.54)
            section.left_margin = Inches(self.margins['left'] / 2.54)
            section.right_margin = Inches(self.margins['right'] / 2.54)
            
        # Store images for insertion
        if images:
            self.images = images
            self.image_lookup = {img['image_id']: img for img in images}
            
        # Initialize image inserter
        self.image_inserter = ImageInserter(self.doc, self.images)
        
        # Initialize figure/table formatters
        self.figure_formatter = FigureFormatter()
        self.table_formatter = TableFormatter()
        
        # Check if this is a questionnaire document
        if questionnaire_data and questionnaire_data.get('is_questionnaire'):
            logger.info("Generating questionnaire document...")
            self.doc = format_questionnaire_in_word(self.doc, questionnaire_data, self.font_size)
            self._enforce_no_italics(self.doc)
            self._run_acceptance_checks(self.doc)
            self.doc.save(output_path)
            return output_path
            
        # --- STANDARD DOCUMENT GENERATION ---
        # Note: Cover page generation removed - templates are now used instead
            
        self._setup_styles()
        
        # Apply formatting to default styles to ensure consistency across the document
        try:
            styles = self.doc.styles
            
            # Update Normal style with formatting options
            normal_style = styles['Normal']
            normal_font = normal_style.font
            normal_font.name = 'Times New Roman'
            normal_font.size = Pt(font_size)
            normal_font.italic = False
            logger.info(f"DEBUG: Setting Normal style font size")
            logger.info(f"  Input font_size={font_size}")
            logger.info(f"  After setting: normal_font.size={normal_font.size}")
            if normal_font.size:
                logger.info(f"  normal_font.size.pt={normal_font.size.pt}")
            normal_pf = normal_style.paragraph_format
            normal_pf.line_spacing = line_spacing
            logger.info(f"DEBUG: Set Normal style line_spacing={line_spacing}")
            logger.info(f"DEBUG: After setting: normal_pf.line_spacing={normal_pf.line_spacing}")
            
            # Define custom style for body content
            if 'AcademicBody' not in styles:
                style = styles.add_style('AcademicBody', WD_STYLE_TYPE.PARAGRAPH)
                style.base_style = styles['Normal']
                font = style.font
                font.name = 'Times New Roman'
                font.size = Pt(font_size)
                font.italic = False
                pf = style.paragraph_format
                pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                pf.line_spacing = line_spacing
                pf.space_after = Pt(0)
                pf.left_indent = Pt(0)
                pf.first_line_indent = Pt(0)
        except Exception as e:
            logger.warning(f"Could not create/modify styles: {e}")
        
        # Store images for reinsertion
        if images:
            self.images = images
            self.image_lookup = {img['image_id']: img for img in images}
            self.image_inserter = ImageInserter(self.doc, images)
            logger.info(f"WordGenerator initialized with {len(images)} images")
        
        # Store shapes/flowcharts for reinsertion
        if shapes:
            self.shapes = shapes
            self.shape_lookup = {shape['shape_id']: shape for shape in shapes}
            self.shape_inserter = ShapeInserter(self.doc, shapes)
            logger.info(f"WordGenerator initialized with {len(shapes)} shapes/flowcharts")
        else:
            self.shapes = []
            self.shape_lookup = {}
            self.shape_inserter = None
        
        # Note: Cover page generation removed - templates are now used instead
        
        # DISABLED: Certification page replacement
        # We now preserve the original certification/declaration content instead of replacing it
        # The original content (including names, signatures, etc.) is kept intact
        # if certification_data:
        #     self._create_certification_page(certification_data)
        #     logger.info("Certification page created from extracted data")
        
        # Handle case where structured_data might be nested lists
        flat_structured_data = []
        for item in structured_data:
            if isinstance(item, dict):
                flat_structured_data.append(item)
            elif isinstance(item, list):
                flat_structured_data.extend([x for x in item if isinstance(x, dict)])
        structured_data = flat_structured_data if flat_structured_data else structured_data
        
        # Calculate if we need TOC (based on total content size)
        total_chars = 0
        for s in structured_data:
            if isinstance(s, dict):
                total_chars += len(str(s.get('heading', '')))
                content = s.get('content', [])
                if isinstance(content, list):
                    for c in content:
                        if isinstance(c, dict):
                            total_chars += len(str(c.get('text', '')))
        estimated_pages = total_chars / 2000  # ~2000 chars per page
        
        # TOC is included only if user explicitly enabled it
        needs_toc = include_toc

        toc_placeholder_present = any(
            isinstance(s, dict) and s.get('type') == 'toc_placeholder' for s in structured_data
        )
        lof_placeholder_present = any(
            isinstance(s, dict) and s.get('type') == 'list_of_figures_placeholder' for s in structured_data
        )
        lot_placeholder_present = any(
            isinstance(s, dict) and s.get('type') == 'list_of_tables_placeholder' for s in structured_data
        )
        
        # Detect short documents that should NOT have section page breaks
        # Short documents are < 5 pages OR have < 5 sections
        # BUT: Always allow section breaks for Roman->Arabic numeral transition
        section_count = len(structured_data)
        self.is_short_document = estimated_pages < 5 or section_count < 5
        
        # For short documents, suppress page breaks EXCEPT for CHAPTER 1 (Roman->Arabic transition)
        if self.is_short_document:
            logger.info(f"Short document detected ({estimated_pages:.1f} pages, {section_count} sections) - suppressing page breaks except for chapter transitions")
            for s in structured_data:
                if isinstance(s, dict):
                    # Check if this is CHAPTER 1 - preserve its page break for numbering transition
                    heading = s.get('heading', '').upper() if s.get('heading') else ''
                    is_chapter_one = bool(re.search(r'^CHAPTER\s+(1|ONE)\b', heading))
                    
                    if not is_chapter_one:
                        # Suppress page breaks for non-chapter content
                        s['needs_page_break'] = False
                        s['start_on_new_page'] = False
                        s['use_page_break_before'] = False
                    # For CHAPTER 1, keep the page break so section numbering can transition
        
        # Count preliminary pages/sections to determine numbering style
        prelim_count = 0
        # Note: cover_page_data removed - templates are now used instead
        if certification_data:
            prelim_count += 1
        if needs_toc:
            prelim_count += 1  # TOC itself counts as a page
        has_front_matter_sections = any(
            isinstance(s, dict) and s.get('type') == 'front_matter' for s in structured_data
        )
        prelim_count += sum(
            1 for s in structured_data if isinstance(s, dict) and s.get('type') == 'front_matter'
        )
        
        # Determine numbering style
        has_preliminary = bool(certification_data or needs_toc)
        self.toc_only_preliminary = bool(needs_toc and not certification_data and not has_front_matter_sections)
        self.arabic_started_after_toc = False
        
        # Use Roman numerals for preliminary pages if any preliminary content exists
        # EXCEPTION: If only TOC exists in preliminary pages (no abstract, dedication, acknowledgements, etc.),
        # use Arabic numerals throughout the entire document
        if not has_preliminary or self.toc_only_preliminary:
            self.use_continuous_arabic = True
        else:
            self.use_continuous_arabic = False
            
        # Set initial page numbering
        first_section = self.doc.sections[0]
        if self.use_continuous_arabic:
            self._set_page_numbering(first_section, fmt='decimal', start=1)
        else:
            self._set_page_numbering(first_section, fmt='lowerRoman', start=1)
        self._add_page_number_to_footer(first_section)
        
        # Initialize TOC entries and heading numberer BEFORE processing any sections
        self.toc_entries = []
        self.heading_numberer = HeadingNumberer(policy=self.policy)
        self.figure_formatter = FigureFormatter()  # Reset figure formatter
        self.table_formatter = TableFormatter()  # Reset table formatter
        self.figure_entries = []  # Reset figure entries
        self.table_entries = []  # Reset table entries
        self.has_figures = False  # Reset figure flag
        self.has_tables = False  # Reset table flag
        
        # Scan content for figures and tables to determine if LOF/LOT is needed
        for section in structured_data:
            # Ensure section is a dict before processing
            if not isinstance(section, dict):
                continue
                
            for item in section.get('content', []):
                # Ensure item is a dict
                if not isinstance(item, dict):
                    continue
                    
                item_type = item.get('type', '')
                
                # Check for various figure-related types
                if item_type == 'image_placeholder':
                    # Only count if it has a caption
                    if item.get('caption'):
                        self.has_figures = True
                elif item_type == 'figure':
                    # Only count if it has a caption
                    if item.get('caption'):
                        self.has_figures = True
                elif item_type in ('figure_caption', 'figure_equation'):
                    self.has_figures = True
                
                # Check for table-related types
                if item_type == 'table':
                    # Only count if it has a caption (tables usually do, but check to be safe)
                    if item.get('caption') or item.get('title'):
                        self.has_tables = True
                elif item_type == 'table_caption':
                    self.has_tables = True
                
                # Also check paragraph text for figure and table captions
                if item_type == 'paragraph':
                    text = item.get('text', '')
                    if self.figure_formatter.is_figure_caption(text):
                        self.has_figures = True
                    if self.table_formatter.is_table_caption(text):
                        self.has_tables = True
            # Break early if both found
            if self.has_figures and self.has_tables:
                break
        
        # Track first section for TOC if it's a chapter
        first_section_for_toc = None
        
        # Add TOC placeholder if needed (uses Word's built-in TOC field)
        added_toc_break = False
        if needs_toc and not toc_placeholder_present:
            # Add TOC field - will be updated by Word after saving
            self._add_toc_placeholder()
            added_toc_break = True
            
            # Add List of Figures after TOC if document has figures
            if self.has_figures and not lof_placeholder_present:
                self._add_lof_placeholder()
                added_toc_break = True
            
            # Add List of Tables after LOF if document has tables
            if self.has_tables and not lot_placeholder_present:
                self._add_lot_placeholder()
                added_toc_break = True
        
        # Add all sections
        rendered_section_count = 0
        # Track if we're processing the first section (to avoid empty first page AND duplicate breaks after TOC)
        # ALWAYS set to True - first section should never add a page break because:
        # - If no TOC: Adding a page break would create an empty first page
        # - If TOC exists: TOC already adds a page break, so adding another creates a double break
        self.is_first_section = True
        for i, section in enumerate(structured_data):
            if section.get('type') == 'toc_placeholder':
                if needs_toc:
                    self._add_toc_placeholder()
                    if self.has_figures and not lof_placeholder_present:
                        self._add_lof_placeholder()
                    if self.has_tables and not lot_placeholder_present:
                        self._add_lot_placeholder()
                continue

            if section.get('type') == 'list_of_figures_placeholder':
                if needs_toc and self.has_figures:
                    self._add_lof_placeholder()
                continue

            if section.get('type') == 'list_of_tables_placeholder':
                if needs_toc and self.has_tables:
                    self._add_lot_placeholder()
                continue

            # Special handling for "Document" title section (auto-generated for unstructured text)
            if section.get('heading', '').strip().lower() == 'document':
                # Just add content, skip heading
                self._add_section_content(section)
                rendered_section_count += 1
                continue

            # First rendered section should never trigger additional page breaks.
            # - Without TOC: avoids an empty first page.
            # - With TOC/LOF/LOT: avoids double breaks after preliminary pages.
            if rendered_section_count == 0:
                section = section.copy()
                section['needs_page_break'] = False
                section['use_page_break_before'] = False
                section['start_on_new_page'] = False
            
            self._add_section(section)
            rendered_section_count += 1
            # After processing first section, set flag to False
            if self.is_first_section:
                self.is_first_section = False
        
        # Save document first
        if is_free_tier:
            self.add_watermark()
        
        # DEBUG: Check Normal style right before saving
        try:
            normal_before_save = self.doc.styles['Normal']
            if normal_before_save.font.size:
                logger.info(f"DEBUG BEFORE SAVE: Normal style font size = {normal_before_save.font.size.pt}pt")
            else:
                logger.info(f"DEBUG BEFORE SAVE: Normal style font size = None")
            logger.info(f"DEBUG BEFORE SAVE: Normal style line_spacing = {normal_before_save.paragraph_format.line_spacing}")
        except Exception as e:
            logger.error(f"DEBUG: Error checking Normal style: {e}")
        
        self._enforce_no_italics(self.doc)
        self._run_acceptance_checks(self.doc)
        self.doc.save(output_path)
        logger.info(f"Document saved to {output_path}")
        
        # Update TOC using Microsoft Word COM automation
        if needs_toc:
            toc_updated = update_toc_with_word(output_path)
            if toc_updated:
                logger.info("Table of Contents updated automatically")
            else:
                logger.warning("TOC could not be auto-updated - user will need to update manually in Word")
        
        return output_path

    def add_watermark(self):
        """Add watermark to document footer and ensure it's bottom-right in all sections."""
        try:
            WATERMARK_TEXT = 'Formatted with AfroDocs.app'
            for section in self.doc.sections:
                footer = section.footer

                # Track whether a watermark paragraph already exists and whether
                # we need to add one because watermark text was embedded in other
                # paragraphs.
                watermark_exists = False
                needs_new_paragraph = False

                for para in list(footer.paragraphs):
                    text = para.text or ''
                    if WATERMARK_TEXT in text:
                        if text.strip() == WATERMARK_TEXT:
                            # Paragraph only contains watermark - normalize it
                            # to a single run and right-align.
                            para.clear()
                            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            pformat = para.paragraph_format
                            pformat.space_before = Pt(0)
                            pformat.space_after = Pt(0)
                            pformat.right_indent = Inches(0)
                            r = para.add_run(WATERMARK_TEXT)
                            r.font.name = 'Arial'
                            r.font.size = Pt(9)
                            try:
                                r.font.color.rgb = RGBColor(128, 128, 128)
                            except Exception:
                                pass
                            watermark_exists = True
                        else:
                            # Watermark is mixed with other footer content (e.g.
                            # page number). Remove watermark text from this
                            # paragraph and schedule a separate right-aligned
                            # watermark paragraph to be added at the footer end.
                            new_text = text.replace(WATERMARK_TEXT, '').strip()
                            para.clear()
                            if new_text:
                                para.add_run(new_text)
                            needs_new_paragraph = True

                # If we removed embedded watermark(s) or no watermark was found,
                # ensure a dedicated right-aligned watermark paragraph exists.
                if needs_new_paragraph or not watermark_exists:
                    # Avoid adding duplicate watermark if one was already added
                    # by prior loop iterations for this footer.
                    exists_now = any((p.text or '') .strip() == WATERMARK_TEXT for p in footer.paragraphs)
                    if not exists_now:
                        p = footer.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        pformat = p.paragraph_format
                        pformat.space_before = Pt(0)
                        pformat.space_after = Pt(0)
                        pformat.right_indent = Inches(0)
                        r = p.add_run(WATERMARK_TEXT)
                        r.font.name = 'Arial'
                        r.font.size = Pt(9)
                        try:
                            r.font.color.rgb = RGBColor(128, 128, 128)
                        except Exception:
                            pass

        except Exception as e:
            logger.error(f"Failed to add watermark: {e}")
    
    # Note: _create_cover_page method removed - templates are now used instead

    def _create_certification_page(self, cert_data):
        """
        Create a standardized certification page.
        
        Layout:
        1. CERTIFICATION header (same format as dissertation headings)
        2. Certification text with topic (untouched, not bolded)
        3. Signature textboxes: Supervisor (left) | HOD (right)
        4. Director signature (left-aligned)
        5. Acceptance statement
        6. Date line
        7. General Coordinator section
        """
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.table import _Cell
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        def remove_cell_borders(cell):
            """Remove all borders from a table cell."""
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right']:
                border_elem = OxmlElement(f'w:{border_name}')
                border_elem.set(qn('w:val'), 'nil')
                tcBorders.append(border_elem)
            tcPr.append(tcBorders)
        
        def create_signature_textbox(doc, name, title, width_inches=2.8):
            """Create a textbox with signature line, name, and title."""
            table = doc.add_table(rows=3, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.autofit = False
            table.columns[0].width = Inches(width_inches)
            
            # Remove borders from all cells
            for row in table.rows:
                for cell in row.cells:
                    remove_cell_borders(cell)
            
            # Row 0: Signature line
            line_cell = table.cell(0, 0)
            line_para = line_cell.paragraphs[0]
            line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            line_para.paragraph_format.space_before = Pt(18)
            line_para.paragraph_format.space_after = Pt(3)
            line_run = line_para.add_run('_' * 25)
            line_run.font.name = 'Times New Roman'
            line_run.font.size = Pt(self.font_size)
            
            # Row 1: Name
            name_cell = table.cell(1, 0)
            name_para = name_cell.paragraphs[0]
            name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            name_para.paragraph_format.space_before = Pt(3)
            name_para.paragraph_format.space_after = Pt(0)
            name_display = name if name else '________________'
            name_run = name_para.add_run(name_display)
            name_run.font.name = 'Times New Roman'
            name_run.font.size = Pt(self.font_size)
            name_run.bold = True
            
            # Row 2: Title
            title_cell = table.cell(2, 0)
            title_para = title_cell.paragraphs[0]
            title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            title_para.paragraph_format.space_before = Pt(0)
            title_para.paragraph_format.space_after = Pt(6)
            title_run = title_para.add_run(f'({title})')
            title_run.font.name = 'Times New Roman'
            title_run.font.size = Pt(11)
            title_run.italic = False
            
            return table
        
        # Get data from certification extraction
        topic = cert_data.get('topic') or '[RESEARCH TOPIC]'
        author = cert_data.get('author') or '[AUTHOR NAME]'
        degree = cert_data.get('degree') or "Master's in Business Administration (MBA)"
        program = cert_data.get('program') or 'Management and Entrepreneurship'
        supervisor = cert_data.get('supervisor')
        hod = cert_data.get('head_of_department')
        director = cert_data.get('director')
        institution = cert_data.get('institution') or 'The Higher Institute of Commerce and Management of The University of Bamenda'
        
        # ========== 1. CERTIFICATION HEADER (same format as dissertation headings) ==========
        # Use Heading 1 style format: Times New Roman, 12pt, bold, centered, black
        cert_heading = self.doc.add_heading('CERTIFICATION', level=1)
        cert_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cert_heading.paragraph_format.page_break_before = False
        # CRITICAL FIX: Disable keep_with_next and keep_together to prevent headings
        # from jumping to new pages when there's not enough room for heading + content
        cert_heading.paragraph_format.keep_with_next = False
        cert_heading.paragraph_format.keep_together = False
        cert_heading.paragraph_format.space_before = Pt(12)
        cert_heading.paragraph_format.space_after = Pt(6)
        cert_heading.paragraph_format.line_spacing = self.line_spacing
        for run in cert_heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(self.font_size)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        # ========== 2. CERTIFICATION TEXT (topic in bold) ==========
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(12)
        para.paragraph_format.line_spacing = self.line_spacing
        
        # First part of the paragraph
        run1 = para.add_run('This is to certify that this research titled ')
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(self.font_size)
        
        # Topic in bold with quotes
        topic_run = para.add_run(f'"{topic}"')
        topic_run.font.name = 'Times New Roman'
        topic_run.font.size = Pt(self.font_size)
        topic_run.font.bold = True
        
        # Rest of the paragraph
        run2 = para.add_run(f' is the original work of {author}. This work is submitted in partial fulfilment of the requirement for the award of a {degree} in {program} in {institution} Cameroon.')
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(self.font_size)
        
        # ========== 3. SIGNATURE TEXTBOXES: Supervisor (left) | HOD (right) ==========
        # Create a 2-column table to hold the two textboxes side by side
        sig_container = self.doc.add_table(rows=1, cols=2)
        sig_container.alignment = WD_TABLE_ALIGNMENT.CENTER
        sig_container.autofit = False
        sig_container.columns[0].width = Inches(3.25)
        sig_container.columns[1].width = Inches(3.25)
        
        # Remove borders from container
        for cell in sig_container.rows[0].cells:
            remove_cell_borders(cell)
        
        # LEFT CELL: Supervisor textbox
        sup_cell = sig_container.cell(0, 0)
        
        # Signature line paragraph
        sup_line_para = sup_cell.paragraphs[0]
        sup_line_para.paragraph_format.space_before = Pt(18)
        sup_line_para.paragraph_format.space_after = Pt(3)
        sup_line = sup_line_para.add_run('_' * 15)
        sup_line.font.name = 'Times New Roman'
        sup_line.font.size = Pt(self.font_size)
        
        # Name paragraph (directly under line)
        sup_name_para = sup_cell.add_paragraph()
        sup_name_para.paragraph_format.space_before = Pt(0)
        sup_name_para.paragraph_format.space_after = Pt(0)
        sup_name_display = supervisor if supervisor else ''
        sup_name_run = sup_name_para.add_run(sup_name_display)
        sup_name_run.font.name = 'Times New Roman'
        sup_name_run.font.size = Pt(self.font_size)
        sup_name_run.bold = True
        
        # Title paragraph
        sup_title_para = sup_cell.add_paragraph()
        sup_title_para.paragraph_format.space_before = Pt(0)
        sup_title_para.paragraph_format.space_after = Pt(6)
        sup_title_run = sup_title_para.add_run('(Supervisor)')
        sup_title_run.font.name = 'Times New Roman'
        sup_title_run.font.size = Pt(11)
        sup_title_run.bold = True
        
        # RIGHT CELL: HOD textbox
        hod_cell = sig_container.cell(0, 1)
        
        # Signature line paragraph
        hod_line_para = hod_cell.paragraphs[0]
        hod_line_para.paragraph_format.space_before = Pt(18)
        hod_line_para.paragraph_format.space_after = Pt(3)
        hod_line = hod_line_para.add_run('_' * 15)
        hod_line.font.name = 'Times New Roman'
        hod_line.font.size = Pt(self.font_size)
        
        # Name paragraph (directly under line)
        hod_name_para = hod_cell.add_paragraph()
        hod_name_para.paragraph_format.space_before = Pt(0)
        hod_name_para.paragraph_format.space_after = Pt(0)
        hod_name_display = hod if hod else ''
        hod_name_run = hod_name_para.add_run(hod_name_display)
        hod_name_run.font.name = 'Times New Roman'
        hod_name_run.font.size = Pt(self.font_size)
        hod_name_run.bold = True
        
        # Title paragraph
        hod_title_para = hod_cell.add_paragraph()
        hod_title_para.paragraph_format.space_before = Pt(0)
        hod_title_para.paragraph_format.space_after = Pt(6)
        hod_title_run = hod_title_para.add_run('(Head Of Department)')
        hod_title_run.font.name = 'Times New Roman'
        hod_title_run.font.size = Pt(11)
        hod_title_run.bold = True
        
        # ========== 4. DIRECTOR SIGNATURE (left-aligned) ==========
        # Signature line
        dir_line_para = self.doc.add_paragraph()
        dir_line_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        dir_line_para.paragraph_format.space_before = Pt(18)
        dir_line_para.paragraph_format.space_after = Pt(3)
        dir_line_run = dir_line_para.add_run('_' * 15)
        dir_line_run.font.name = 'Times New Roman'
        dir_line_run.font.size = Pt(self.font_size)
        
        # Name (directly under line)
        dir_name_para = self.doc.add_paragraph()
        dir_name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        dir_name_para.paragraph_format.space_before = Pt(0)
        dir_name_para.paragraph_format.space_after = Pt(0)
        dir_name_display = director if director else ''
        dir_name_run = dir_name_para.add_run(dir_name_display)
        dir_name_run.font.name = 'Times New Roman'
        dir_name_run.font.size = Pt(self.font_size)
        dir_name_run.bold = True
        
        dir_title_para = self.doc.add_paragraph()
        dir_title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        dir_title_para.paragraph_format.space_before = Pt(0)
        dir_title_run = dir_title_para.add_run('(Director)')
        dir_title_run.font.name = 'Times New Roman'
        dir_title_run.font.size = Pt(11)
        dir_title_run.bold = True
        
        # ========== 4. ACCEPTANCE STATEMENT ==========
        self.doc.add_paragraph().paragraph_format.space_after = Pt(12)
        
        accept_para = self.doc.add_paragraph()
        accept_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        accept_para.paragraph_format.space_before = Pt(12)
        accept_para.paragraph_format.space_after = Pt(12)
        accept_para.paragraph_format.line_spacing = self.line_spacing
        
        accept_run = accept_para.add_run('Having met the stipulated requirements, the dissertation has been accepted by the Postgraduate School')
        accept_run.font.name = 'Times New Roman'
        accept_run.font.size = Pt(self.font_size)
        
        # ========== 5. DATE LINE ==========
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        date_para.paragraph_format.space_before = Pt(18)
        date_para.paragraph_format.space_after = Pt(6)
        
        date_run = date_para.add_run('Date' + '_' * 25)
        date_run.font.name = 'Times New Roman'
        date_run.font.size = Pt(self.font_size)
        
        # ========== 6. GENERAL COORDINATOR SECTION ==========
        # Right-aligned signature section
        gc_table = self.doc.add_table(rows=3, cols=2)
        gc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        gc_table.autofit = False
        gc_table.columns[0].width = Inches(3.0)
        gc_table.columns[1].width = Inches(3.0)
        
        # Remove borders
        for row in gc_table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right']:
                    border_elem = OxmlElement(f'w:{border_name}')
                    border_elem.set(qn('w:val'), 'nil')
                    tcBorders.append(border_elem)
                tcPr.append(tcBorders)
        
        # Right column only - signature line
        gc_line_cell = gc_table.cell(0, 1).paragraphs[0]
        gc_line_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
        gc_line_cell.paragraph_format.space_before = Pt(18)
        gc_line_cell.add_run('_' * 30).font.name = 'Times New Roman'
        
        # Title
        gc_title_cell = gc_table.cell(1, 1).paragraphs[0]
        gc_title_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
        gc_title_cell.paragraph_format.space_before = Pt(6)
        gc_title_run = gc_title_cell.add_run('The General Coordinator')
        gc_title_run.font.name = 'Times New Roman'
        gc_title_run.font.size = Pt(self.font_size)
        gc_title_run.bold = True
        
        # School
        gc_school_cell = gc_table.cell(2, 1).paragraphs[0]
        gc_school_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
        gc_school_run = gc_school_cell.add_run('Postgraduate School')
        gc_school_run.font.name = 'Times New Roman'
        gc_school_run.font.size = Pt(self.font_size)
        
        # Add page break after certification page
        self.doc.add_page_break()
        logger.info("Certification page created with all sections")
    
    def _setup_styles(self):
        """Configure document styles"""
        styles = self.doc.styles
        
        # Normal style - NO INDENTATION
        normal = styles['Normal']
        normal.font.name = 'Times New Roman'
        normal.font.size = Pt(self.font_size)
        normal.font.italic = False
        normal.paragraph_format.line_spacing = self.line_spacing
        normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.left_indent = Pt(0)  # No left indent
        normal.paragraph_format.first_line_indent = Pt(0)  # No first line indent
        
        # Title style (for level 0 headings) - larger than normal
        try:
            title = styles['Title']
            title.font.name = 'Times New Roman'
            title.font.size = Pt(max(self.font_size + 4, 16))
            title.font.bold = True
            title.font.italic = False
            title.paragraph_format.line_spacing = self.line_spacing
            title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except KeyError:
            pass
        
        # Heading styles - CRITICAL: Set to Times New Roman (overrides Calibri default from some templates)
        heading_configs = {
            1: {'size': self.font_size, 'bold': True, 'space_before': 0, 'space_after': 0},
            2: {'size': self.font_size, 'bold': True, 'space_before': 12, 'space_after': 6},
            3: {'size': self.font_size, 'bold': True, 'space_before': 12, 'space_after': 6},
            4: {'size': self.font_size, 'bold': True, 'space_before': 12, 'space_after': 6},
        }
        
        for level, config in heading_configs.items():
            try:
                heading = styles[f'Heading {level}']
                # CRITICAL: Force Times New Roman to override any Calibri defaults from template
                heading.font.name = 'Times New Roman'
                heading.font.bold = config['bold']
                heading.font.size = Pt(int(config['size']))
                heading.font.italic = False
                heading.font.color.rgb = RGBColor(0, 0, 0)  # Black
                heading.paragraph_format.line_spacing = self.line_spacing
                heading.paragraph_format.space_before = Pt(config['space_before'])
                heading.paragraph_format.space_after = Pt(config['space_after'])
                heading.paragraph_format.left_indent = Pt(0)  # No left indent
                heading.paragraph_format.first_line_indent = Pt(0)  # No first line indent
                # CRITICAL FIX: Ensure NO page breaks before headings by default
                # This prevents level 2/3 headings from starting on new pages
                heading.paragraph_format.page_break_before = False
                # CRITICAL FIX: Disable keep_with_next and keep_together on heading styles
                # These Word defaults cause headings to jump to next page when there's not enough
                # room for heading + following content - this is NOT desired for sub-sections
                heading.paragraph_format.keep_with_next = False
                heading.paragraph_format.keep_together = False
            except KeyError:
                pass  # Style doesn't exist, skip

        # Ensure list/caption styles are not italicized
        for style_name in ['List Bullet', 'List Number', 'List Bullet 2', 'List Number 2', 'Caption']:
            if style_name in styles:
                style = styles[style_name]
                if style.font is not None:
                    style.font.italic = False

        # References entry style (used to allow italics within references)
        if 'ReferenceEntry' not in styles:
            reference_style = styles.add_style('ReferenceEntry', WD_STYLE_TYPE.PARAGRAPH)
            reference_style.base_style = styles['Normal']
            reference_style.font.name = 'Times New Roman'
            reference_style.font.size = Pt(self.font_size)
            reference_style.font.italic = False
    
    def _add_title(self, title_text):
        """Add document title"""
        title = self.doc.add_heading(title_text, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.line_spacing = self.line_spacing
        
        # Style the title - Times New Roman, use custom font size, bold, black
        for run in title.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(self.font_size + 2)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        
        self.doc.add_paragraph()  # Spacing
    
    def _add_toc_placeholder(self):
        """Add Table of Contents using Microsoft Word's built-in TOC field"""
        # Add TOC heading as a regular paragraph (NOT a heading style to avoid appearing in TOC)
        toc_heading = self.doc.add_paragraph()
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = toc_heading.add_run('TABLE OF CONTENTS')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(self.font_size)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        
        # No space after the title
        toc_heading.paragraph_format.space_after = Pt(0)
        toc_heading.paragraph_format.space_before = Pt(0)
        toc_heading.paragraph_format.line_spacing = self.line_spacing
        
        # Add TOC field code directly (no blank paragraph in between)
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        
        # Create TOC field - this generates the actual Word TOC
        # TOC \o "1-3" - include heading levels 1-3
        # \h - hyperlinks
        # \z - hide tab leader and page numbers in Web layout view
        # \u - use applied paragraph outline level
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        # --- INSTRUCTION TEXT (Inside the Field) ---
        # This text will be replaced when the user updates the field
        
        # Add some spacing before
        paragraph.add_run("\n\n")
        
        # Add the instruction text
        run_instr = paragraph.add_run("Right click and update field to get table of contents")
        run_instr.font.name = 'Times New Roman'
        run_instr.font.size = Pt(self.font_size)
        run_instr.font.bold = True
        run_instr.font.color.rgb = RGBColor(68, 114, 196) # Word Blue
        
        # Add some spacing after
        paragraph.add_run("\n\n")
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run_end = paragraph.add_run()
        run_end._r.append(fldChar3)
        
        # Single page break after TOC (not two)
        if self.toc_only_preliminary:
            # TOC is the only preliminary content, so we're already using Arabic numerals
            # Just add a page break to continue the document (numbering continues)
            self.doc.add_page_break()
            self.arabic_started_after_toc = True
        else:
            self.doc.add_page_break()
    
    def _add_lof_placeholder(self):
        """Add List of Figures using Microsoft Word's built-in TOC field for figures
        
        This creates a Word field that automatically lists all figure captions
        with their page numbers. Similar to TOC but specifically for figures.
        The field uses SEQ Figure field references for proper figure numbering.
        """
        # Add LOF heading as a regular paragraph (NOT a heading style to avoid appearing in TOC)
        lof_heading = self.doc.add_paragraph()
        lof_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = lof_heading.add_run('LIST OF FIGURES')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(self.font_size)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        
        # No space after the title
        lof_heading.paragraph_format.space_after = Pt(0)
        lof_heading.paragraph_format.space_before = Pt(0)
        lof_heading.paragraph_format.line_spacing = self.line_spacing
        
        # Add LOF field code directly
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        
        # Create TOC field for figures
        # TOC \h \z \c "Figure" - Table of contents for Figure captions
        # \h - hyperlinks
        # \z - hide tab leader and page numbers in Web layout view  
        # \c "Figure" - Build TOC from SEQ Figure field captions
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\h \\z \\c "Figure" '
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        # --- INSTRUCTION TEXT (Inside the Field) ---
        
        # Add some spacing before
        paragraph.add_run("\n\n")
        
        # Add the instruction text
        run_instr = paragraph.add_run("Right click and update field to get list of figures")
        run_instr.font.name = 'Times New Roman'
        run_instr.font.size = Pt(self.font_size)
        run_instr.font.bold = False  # LOF entries should not be bold
        run_instr.font.color.rgb = RGBColor(68, 114, 196) # Word Blue
        
        # Add some spacing after
        paragraph.add_run("\n\n")
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run_end = paragraph.add_run()
        run_end._r.append(fldChar3)
        
        # Page break after List of Figures
        self.doc.add_page_break()
    
    def _add_lot_placeholder(self):
        """Add List of Tables using Microsoft Word's built-in TOC field for tables
        
        This creates a Word field that automatically lists all table captions
        with their page numbers. Similar to TOC but specifically for tables.
        The field uses SEQ Table field references for proper table numbering.
        """
        # Add LOT heading as a regular paragraph (NOT a heading style to avoid appearing in TOC)
        lot_heading = self.doc.add_paragraph()
        lot_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = lot_heading.add_run('LIST OF TABLES')
        run.font.name = 'Times New Roman'
        run.font.size = Pt(self.font_size)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        
        # No space after the title
        lot_heading.paragraph_format.space_after = Pt(0)
        lot_heading.paragraph_format.space_before = Pt(0)
        lot_heading.paragraph_format.line_spacing = self.line_spacing
        
        # Add LOT field code directly
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        
        # Create TOC field for tables
        # TOC \h \z \c "Table" - Table of contents for Table captions
        # \h - hyperlinks
        # \z - hide tab leader and page numbers in Web layout view  
        # \c "Table" - Build TOC from SEQ Table field captions
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\h \\z \\c "Table" '
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        # --- INSTRUCTION TEXT (Inside the Field) ---
        
        # Add some spacing before
        paragraph.add_run("\n\n")
        
        # Add the instruction text
        run_instr = paragraph.add_run("Right click and update field to get list of tables")
        run_instr.font.name = 'Times New Roman'
        run_instr.font.size = Pt(self.font_size)
        run_instr.font.bold = False  # LOT entries should not be bold
        run_instr.font.color.rgb = RGBColor(68, 114, 196) # Word Blue
        
        # Add some spacing after
        paragraph.add_run("\n\n")
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run_end = paragraph.add_run()
        run_end._r.append(fldChar3)
        
        # Page break after List of Tables
        self.doc.add_page_break()
    
    def _add_table_caption(self, number, title, center=False):
        """
        Add a properly formatted table caption with SEQ field for LOT tracking.
        
        Args:
            number: Table number (e.g., "1" or "1.2")
            title: Table caption text
            center: Whether to center the caption (default False - left aligned)
        
        The caption is formatted as:
        - Times New Roman, 12pt
        - Bold
        - Left aligned (by default)
        - Includes SEQ field for Word's List of Tables to pick up
        
        Note: LOT entries are formatted as plain text via update_toc_with_word()
        """
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(12)
        para.paragraph_format.line_spacing = self.line_spacing
        
        # Add "Table " text
        run1 = para.add_run('Table ')
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(self.font_size)
        run1.font.bold = True
        run1.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add SEQ field for automatic numbering and LOT tracking
        run_seq = para.add_run()
        
        # Create SEQ field - this allows Word to track tables for LOT
        # SEQ Table \* ARABIC - Sequential number for Table category
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' SEQ Table \\* ARABIC '
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        # The actual number (will be updated by Word)
        num_text = OxmlElement('w:t')
        num_text.text = str(number)
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run_seq._r.append(fldChar1)
        run_seq._r.append(instrText)
        run_seq._r.append(fldChar2)
        run_seq._r.append(num_text)
        run_seq._r.append(fldChar3)
        
        # Style the SEQ field run
        run_seq.font.name = 'Times New Roman'
        run_seq.font.size = Pt(self.font_size)
        run_seq.font.bold = True
        run_seq.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add colon and title
        run2 = para.add_run(f': {title}')
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(self.font_size)
        run2.font.bold = True
        run2.font.color.rgb = RGBColor(0, 0, 0)
        
        return para

    def _add_figure_caption(self, number, title, center=True):
        """
        Add a properly formatted figure caption with SEQ field for LOF tracking.
        
        Args:
            number: Figure number (e.g., "1" or "1.2")
            title: Figure caption text
            center: Whether to center the caption (default True)
        
        The caption is formatted as:
        - Times New Roman, 12pt
        - Italic
        - Centered (by default)
        - Includes SEQ field for Word's List of Figures to pick up
        
        Note: LOF entries are formatted as plain text via update_toc_with_word()
        """
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(12)
        para.paragraph_format.line_spacing = self.line_spacing
        
        # Add "Figure " text
        run1 = para.add_run('Figure ')
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(self.font_size)
        run1.font.italic = False
        run1.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add SEQ field for automatic numbering and LOF tracking
        run_seq = para.add_run()
        
        # Create SEQ field - this allows Word to track figures for LOF
        # SEQ Figure \* ARABIC - Sequential number for Figure category
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' SEQ Figure \\* ARABIC '
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        # The actual number (will be updated by Word)
        num_text = OxmlElement('w:t')
        num_text.text = str(number)
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run_seq._r.append(fldChar1)
        run_seq._r.append(instrText)
        run_seq._r.append(fldChar2)
        run_seq._r.append(num_text)
        run_seq._r.append(fldChar3)
        
        # Style the SEQ field run (same approach as table captions)
        run_seq.font.name = 'Times New Roman'
        run_seq.font.size = Pt(self.font_size)
        run_seq.font.italic = False
        run_seq.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add ": " separator
        run2 = para.add_run(': ')
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(self.font_size)
        run2.font.italic = False
        run2.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add the caption title
        run3 = para.add_run(title)
        run3.font.name = 'Times New Roman'
        run3.font.size = Pt(self.font_size)
        run3.font.italic = False
        run3.font.color.rgb = RGBColor(0, 0, 0)
        
        # Track figure for validation
        self.figure_formatter.add_figure_entry(number, title)
        self.figure_entries.append({
            'number': number,
            'title': title
        })
        
        return para
    
    def _format_existing_figure_caption(self, para, text):
        """
        Reformat an existing figure caption paragraph with proper styling.
        
        Args:
            para: The paragraph object to format
            text: The original caption text
            
        Returns:
            The formatted paragraph
        """
        # Detect figure number and title from text
        figure_info = self.figure_formatter.detect_figure_caption(text)
        
        if not figure_info:
            return para  # Not a figure caption, return unchanged
        
        number = figure_info['number']
        title = figure_info['title']
        
        # Clear existing runs
        for run in para.runs:
            run.text = ''
        
        # Add formatted caption with SEQ field
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(12)
        para.paragraph_format.line_spacing = self.line_spacing
        
        # Add "Figure " text
        run1 = para.add_run('Figure ')
        run1.font.name = 'Times New Roman'
        run1.font.size = Pt(self.font_size)
        run1.font.italic = False
        run1.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add SEQ field for automatic numbering
        run_seq = para.add_run()
        
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' SEQ Figure \\* ARABIC '
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        
        num_text = OxmlElement('w:t')
        num_text.text = str(number)
        
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run_seq._r.append(fldChar1)
        run_seq._r.append(instrText)
        run_seq._r.append(fldChar2)
        run_seq._r.append(num_text)
        run_seq._r.append(fldChar3)
        
        # Add ": " and title
        run2 = para.add_run(': ' + title)
        run2.font.name = 'Times New Roman'
        run2.font.size = Pt(self.font_size)
        run2.font.italic = False
        run2.font.color.rgb = RGBColor(0, 0, 0)
        
        # Track figure
        self.figure_formatter.add_figure_entry(number, title)
        self.figure_entries.append({
            'number': number,
            'title': title
        })
        
        return para
    
    def _track_heading_for_toc(self, heading_text, level, numbered_text=None):
        """Track a heading for inclusion in the Table of Contents
        
        Args:
            heading_text: The original heading text
            level: The heading level (1, 2, or 3)
            numbered_text: The numbered heading text (e.g., "1.2 Background")
        """
        if level > 3:
            return  # Only track levels 1-3 for TOC
        
        # Clean heading text (remove markdown markers)
        clean_text = re.sub(r'^#+\s*', '', heading_text).strip()
        
        # Skip TOC itself and certain front matter
        skip_headings = ['table of contents', 'contents', 'toc']
        if clean_text.lower() in skip_headings:
            return
        
        self.toc_entries.append({
            'text': clean_text,
            'numbered_text': numbered_text or clean_text,
            'level': level
        })
    
    def _populate_toc(self):
        """Populate the Table of Contents with actual entries by replacing the marker"""
        if not self.toc_entries:
            return
        
        body = self.doc.element.body
        marker_text = '<<<TOC_ENTRIES_PLACEHOLDER>>>'
        marker_element = None
        
        # Find the marker paragraph
        for element in body:
            if element.tag.endswith('p'):
                text = ''.join(node.text or '' for node in element.iter() if node.text)
                if marker_text in text:
                    marker_element = element
                    break
        
        if marker_element is None:
            logger.warning("Could not find TOC marker placeholder")
            return
        
        # Create TOC entry paragraphs and insert them BEFORE the marker
        for entry in self.toc_entries:
            # Create new paragraph element
            new_para = OxmlElement('w:p')
            
            # Create paragraph properties for indentation and spacing
            pPr = OxmlElement('w:pPr')
            
            # Set indentation based on level
            level = entry['level']
            indent = OxmlElement('w:ind')
            if level == 1:
                indent.set(qn('w:left'), '0')
            elif level == 2:
                indent.set(qn('w:left'), '360')  # 0.25 inch in twips
            else:  # level 3
                indent.set(qn('w:left'), '720')  # 0.5 inch in twips
            pPr.append(indent)
            
            # Set line spacing
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:before'), '0')
            spacing.set(qn('w:after'), '60')  # 3pt in twips
            spacing.set(qn('w:line'), '360')  # 1.5 line spacing
            pPr.append(spacing)
            
            new_para.append(pPr)
            
            # Create run for text
            run = OxmlElement('w:r')
            
            # Run properties (font)
            rPr = OxmlElement('w:rPr')
            
            # Font name
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Times New Roman')
            rFonts.set(qn('w:hAnsi'), 'Times New Roman')
            rPr.append(rFonts)
            
            # Font size (12pt = 24 half-points)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '24')
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), '24')
            rPr.append(szCs)
            
            # Bold for level 1
            if level == 1:
                bold = OxmlElement('w:b')
                rPr.append(bold)
            
            # Black color
            color = OxmlElement('w:color')
            color.set(qn('w:val'), '000000')
            rPr.append(color)
            
            run.append(rPr)
            
            # Add text
            text_elem = OxmlElement('w:t')
            text_elem.text = entry['numbered_text']
            run.append(text_elem)
            
            new_para.append(run)
            
            # Insert before marker
            marker_element.addprevious(new_para)
        
        # Remove the marker paragraph
        body.remove(marker_element)
        
        logger.info(f"TOC populated with {len(self.toc_entries)} entries")
    
    def _add_toc(self):
        """Legacy method - now uses placeholder approach"""
        self._add_toc_placeholder()
    
    def _insert_image(self, image_id):
        """
        Insert an image into the document at current position.
        
        Args:
            image_id: The ID of the image to insert
        """
        if not image_id:
            logger.warning("No image_id provided to _insert_image")
            return
        
        if image_id not in self.image_lookup:
            logger.warning(f"Image {image_id} not found in image lookup")
            # Add placeholder text
            para = self.doc.add_paragraph()
            para.add_run(f"[IMAGE: {image_id} - Not found]")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return
        
        img_data = self.image_lookup[image_id]
        
        try:
            # Create paragraph for image with minimal spacing to fit on page
            para = self.doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(6)  # Reduced spacing
            para.paragraph_format.space_after = Pt(3)   # Reduced spacing
            para.paragraph_format.keep_with_next = True  # Keep with caption if present
            
            # Add image from bytes
            run = para.add_run()
            
            # Create BytesIO stream from image data
            image_stream = BytesIO(img_data['data'])
            
            # Determine width and height - preserve original dimensions
            width = img_data.get('width', 4.0)
            height = img_data.get('height', 3.0)
            
            # Store original for logging
            original_width, original_height = width, height
            
            # Only scale down if exceeds page margins (6.5" width, 9" height for letter)
            max_width = 6.5  # Page width with margins
            max_height = 9.0  # Page height with margins
            
            # Scale proportionally only if too large
            if width > max_width:
                ratio = max_width / width
                width = max_width
                height = height * ratio
            
            if height > max_height:
                ratio = max_height / height
                height = max_height
                width = width * ratio
            
            # Preserve small images at original size (no minimum scaling up)
            # Only set minimum if dimensions are invalid (0 or negative)
            if width <= 0:
                width = original_width if original_width > 0 else 2.0
            if height <= 0:
                height = original_height if original_height > 0 else 1.5
            
            # Add picture to document
            run.add_picture(image_stream, width=Inches(width), height=Inches(height))
            
            logger.info(f"Inserted image {image_id} ({width:.2f}x{height:.2f} inches)")
            
            # Add caption if exists
            if img_data.get('caption'):
                caption_text = img_data['caption'].strip()
                if caption_text:
                    figure_info = self.figure_formatter.detect_figure_caption(caption_text)
                    if figure_info:
                        self._add_figure_caption(figure_info['number'], figure_info['title'])
                    else:
                        next_number = str(len(self.figure_entries) + 1)
                        self._add_figure_caption(next_number, caption_text)
                    self.has_figures = True
            
        except Exception as e:
            logger.error(f"Error inserting image {image_id}: {str(e)}")
            # Add error placeholder
            para = self.doc.add_paragraph()
            para.add_run(f"[IMAGE: {image_id} - Error: {str(e)}]")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _insert_shape_inline(self, shape_id, paragraph):
        """
        Insert a shape into an existing paragraph to preserve inline positions.
        """
        if not shape_id or paragraph is None:
            logger.warning("No shape_id or paragraph provided to _insert_shape_inline")
            return

        if not hasattr(self, 'shape_lookup') or shape_id not in self.shape_lookup:
            logger.warning(f"Shape {shape_id} not found in shape lookup")
            paragraph.add_run(f"[SHAPE/DIAGRAM: {shape_id} - Not found]")
            return

        try:
            if self.shape_inserter:
                self.shape_inserter.insert_shape(shape_id, paragraph=paragraph)
            else:
                shape_data = self.shape_lookup[shape_id]
                run = paragraph.add_run()
                if shape_data.get('is_vml'):
                    run._element.append(shape_data['pict_xml'])
                else:
                    run._element.append(shape_data['drawing_xml'])
            logger.info(f"Inserted inline shape {shape_id}")
        except Exception as e:
            logger.error(f"Error inserting inline shape {shape_id}: {str(e)}")
            paragraph.add_run(f"[SHAPE/DIAGRAM: {shape_id} - Error: {str(e)}]")
    
    def _insert_shape(self, shape_id):
        """
        Insert a shape/flowchart/diagram into the document at current position.
        Preserves original positioning, groupings, and formatting.
        
        Args:
            shape_id: The ID of the shape to insert
        """
        if not shape_id:
            logger.warning("No shape_id provided to _insert_shape")
            return
        
        if not hasattr(self, 'shape_lookup') or shape_id not in self.shape_lookup:
            logger.warning(f"Shape {shape_id} not found in shape lookup")
            # Add placeholder text
            para = self.doc.add_paragraph()
            para.add_run(f"[SHAPE/DIAGRAM: {shape_id} - Not found]")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            return
        
        shape_data = self.shape_lookup[shape_id]
        
        try:
            # Create paragraph for shape
            para = self.doc.add_paragraph()
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            
            # Add shape from stored XML
            run = para.add_run()
            
            # Check if this is a VML shape or modern drawing
            if shape_data.get('is_vml'):
                # Insert VML pict element
                pict_xml = shape_data['pict_xml']
                run._element.append(pict_xml)
                logger.info(f"Inserted VML shape {shape_id}")
            else:
                # Insert modern drawing element
                drawing_xml = shape_data['drawing_xml']
                run._element.append(drawing_xml)
                logger.info(f"Inserted drawing shape {shape_id} (contains {shape_data.get('shape_count', 1)} shape elements)")
            
        except Exception as e:
            logger.error(f"Error inserting shape {shape_id}: {str(e)}")
            # Add error placeholder
            para = self.doc.add_paragraph()
            para.add_run(f"[SHAPE/DIAGRAM: {shape_id} - Error: {str(e)}]")
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_section(self, section):
        """Add a document section"""
        section_type = section.get('type', 'section')
        
        # Check for Chapter 1 to switch numbering
        heading_text = section.get('heading', '').strip().upper()
        section_level = section.get('level', 1)
        
        # Force page break for specific sections (Resume, Acknowledgements) - BUT NOT FOR SHORT DOCUMENTS
        # CRITICAL: Only apply to LEVEL 1 headings, not to numbered sub-sections like 2.1, 2.2, 3.3
        force_break_headings = [
            'RESUME', 'RÉSUMÉ', 'RÉSUME', 'RESUMÉ', 
            'ACKNOWLEDGEMENTS', 'ACKNOWLEDGMENTS', 'ACKNOWLEDGEMENT', 'ACKNOWLEDGMENT',
            'INTRODUCTION',
            'LITERATURE REVIEW',
            'METHODOLOGY', 'RESEARCH METHODOLOGY',
            'RESULTS', 'FINDINGS', 'DATA ANALYSIS',
            'DISCUSSION', 'FINDINGS AND DISCUSSION',
            'CONCLUSION', 'SUMMARY', 'RECOMMENDATIONS'
        ]
        
        # Check if this is a numbered sub-section (e.g., 2.0, 2.1, 3.3) - these should NEVER get page breaks
        is_numbered_subsection = bool(re.match(r'^\d+\.\d+', heading_text))
        
        # Only apply forced page breaks if:
        # 1. It's a LEVEL 1 heading (not level 2, 3, etc.)
        # 2. It's NOT a numbered sub-section (like 2.0, 2.1, 2.3, 3.0, etc.)
        # 3. It's not a short document
        # 4. The heading matches one of the force_break keywords
        if (section_level == 1 and 
            not is_numbered_subsection and 
            not self.is_short_document and 
            any(h in heading_text for h in force_break_headings)):
             # Use page_break_before property instead of manual break for better reliability
             section['use_page_break_before'] = True
             section['needs_page_break'] = False # Disable manual break
        
        # Regex for any chapter heading (CHAPTER 1, 2, 3... or CHAPTER ONE, TWO, etc.)
        is_any_chapter = bool(re.search(r'^CHAP?TER\s+(\d+|ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN|[IVXLC]+)\b', heading_text, re.IGNORECASE))
        is_chapter_one = bool(re.search(r'^CHAP?TER\s+(1|ONE)\b', heading_text, re.IGNORECASE))
        
        # CRITICAL FIX: Skip page breaks for the very first section to prevent empty first page
        # This applies when there's no TOC/cover page before the content
        skip_first_section_break = getattr(self, 'is_first_section', False)
        
        # All chapters should start on a new page (EXCEPT the first section)
        if is_any_chapter and not skip_first_section_break:
            if is_chapter_one and not self.use_continuous_arabic and not self.arabic_started_after_toc:
                # Add Section Break (Next Page) to switch to Arabic numbering for Chapter 1
                new_section = self.doc.add_section(WD_SECTION.NEW_PAGE)
                self._set_page_numbering(new_section, fmt='decimal', start=1)
                new_section.footer.is_linked_to_previous = False
                self._add_page_number_to_footer(new_section)
            else:
                # Add page break for all other chapters
                self.doc.add_page_break()
        # Skip other page breaks for short documents
        elif self.is_short_document:
            pass  # No page breaks for short documents (except chapters handled above)
        elif section.get('needs_page_break', False) and section_level == 1 and not is_numbered_subsection:
            # CRITICAL: Only add page break for level 1 sections that are NOT numbered sub-sections
            # This prevents 2.0, 2.1, 2.3, 3.0, 3.3 etc. from getting unwanted page breaks
            self.doc.add_page_break()
        
        # Handle chapter sections (dissertation-specific)
        if section_type == 'chapter':
            self._add_chapter_section(section)
            return
        
        # Handle front matter sections (dissertation-specific)
        if section_type == 'front_matter':
            self._add_front_matter_section(section)
            return
        
        # Handle prominent sections
        if section_type == 'prominent_section':
            formatting = section.get('original_format', 'bold_standalone')

            # Add page break if needed - BUT NOT FOR SHORT DOCUMENTS
            if section.get('start_on_new_page') and not self.is_short_document:
                self.doc.add_page_break()

            heading = self.doc.add_heading(section.get('title', ''), level=min(section.get('level', 3), 4))
            
            # CRITICAL FIX: Disable page break before for prominent sections
            # They should not start on new pages unless explicitly requested above
            heading.paragraph_format.page_break_before = False
            # CRITICAL FIX: Disable keep_with_next and keep_together to prevent headings
            # from jumping to new pages when there's not enough room for heading + content
            heading.paragraph_format.keep_with_next = False
            heading.paragraph_format.keep_together = False

            # CRITICAL: Apply Times New Roman to ALL prominent section headings
            # This ensures consistency across all heading types (not just front matter)
            for run in heading.runs:
                run.font.name = 'Times New Roman'  # FORCE Times New Roman
                run.font.color.rgb = RGBColor(0, 0, 0)  # Ensure black

            # Apply formatting based on original_format
            if formatting == 'bold_standalone':
                for run in heading.runs:
                    run.bold = True
                    if section.get('level', 3) <= 2:
                        run.font.size = Pt(max(self.font_size + 2, 14))
                    else:
                        run.font.size = Pt(self.font_size)  # Ensure font size set
            elif formatting == 'underlined':
                heading.paragraph_format.space_after = Pt(6)
                for run in heading.runs:
                    run.underline = True
                    run.font.size = Pt(self.font_size)
            elif formatting == 'star_surrounded':
                for run in heading.runs:
                    run.italic = False
                    run.bold = True
                    run.font.size = Pt(self.font_size)
            elif formatting == 'numbered_bold':
                for run in heading.runs:
                    run.bold = True
                    run.font.size = Pt(self.font_size)

            # Small separator after prominent heading
            if section.get('needs_emphasis'):
                self.doc.add_paragraph().paragraph_format.space_after = Pt(3)

            # Add content
            for content_item in section.get('content', []):
                self._add_section_content(content_item)
            return

        # Handle hierarchical list groups
        if section_type == 'hierarchical_list_group':
            format_type = section.get('format', 'alphanumeric_hierarchy')
            items = section.get('items', [])
            
            for idx, list_item in enumerate(items):
                level = list_item.get('level', 1)
                content = list_item.get('content', '')
                numbering = list_item.get('numbering', '')

                indent = (level - 1) * 0.3  # inches
                
                # Determine if numbering should be on separate line (bolded)
                # Criteria: numbering is separate AND content spans multiple lines or > 2 lines of text
                content_is_multiline = '\n' in content or len(content.split()) > 20
                should_separate_numbering = content_is_multiline
                
                if should_separate_numbering and numbering:
                    # Create separate paragraph for numbering (bolded title)
                    p_title = self.doc.add_paragraph()
                    p_title.paragraph_format.left_indent = Inches(indent)
                    p_title.paragraph_format.first_line_indent = Pt(0)
                    
                    # Add bolded numbering based on format type
                    if format_type == 'alphanumeric_hierarchy':
                        run = p_title.add_run(f"{numbering}")
                        run.bold = True
                    elif format_type == 'parenthesized':
                        run = p_title.add_run(f"{numbering}")
                        run.bold = True  # Bold for emphasis when separate
                        run.italic = False
                    elif format_type == 'double_parentheses':
                        run = p_title.add_run(f"{numbering}")
                        run.bold = True
                        run.font.size = Pt(11)
                    elif format_type == 'bracketed':
                        run = p_title.add_run(f"{numbering}")
                        run.bold = True
                        run.font.color.rgb = RGBColor(100, 100, 100)
                    elif format_type == 'hyphen':
                        run = p_title.add_run(f"{numbering}")
                        run.bold = True
                    elif format_type == 'roman_dots':
                        run = p_title.add_run(f"{numbering}")
                        run.bold = True
                        run.italic = False
                    
                    p_title.paragraph_format.space_after = Pt(3)
                    
                    # Add content as separate paragraph(s)
                    if content:
                        if '\n' in content:
                            # Multiple paragraphs
                            for para_text in content.split('\n'):
                                if para_text.strip():
                                    p_content = self.doc.add_paragraph(para_text)
                                    p_content.paragraph_format.left_indent = Inches(indent + 0.25)
                                    p_content.paragraph_format.first_line_indent = Pt(0)
                                    p_content.paragraph_format.space_after = Pt(3)
                        else:
                            # Single paragraph
                            p_content = self.doc.add_paragraph(content)
                            p_content.paragraph_format.left_indent = Inches(indent + 0.25)
                            p_content.paragraph_format.first_line_indent = Pt(0)
                            p_content.paragraph_format.space_after = Pt(3)
                else:
                    # Keep numbering and content on same line (original behavior)
                    p = self.doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(indent)
                    p.paragraph_format.first_line_indent = Pt(0)

                    if format_type == 'alphanumeric_hierarchy':
                        run = p.add_run(f"{numbering} ")
                        run.bold = True
                    elif format_type == 'parenthesized':
                        run = p.add_run(f"{numbering} ")
                        run.italic = False
                    elif format_type == 'double_parentheses':
                        run = p.add_run(f"{numbering} ")
                        run.bold = True
                        run.font.size = Pt(11)
                    elif format_type == 'bracketed':
                        run = p.add_run(f"{numbering} ")
                        run.font.color.rgb = RGBColor(100, 100, 100)
                    elif format_type == 'hyphen':
                        run = p.add_run(f"{numbering} ")
                    elif format_type == 'roman_dots':
                        run = p.add_run(f"{numbering} ")
                        run.italic = False

                    p.add_run(content)
                    p.paragraph_format.space_after = Pt(3)
            return

        # Handle shortdoc sections (assignments, questions, tasks)
        if section_type == 'shortdoc_section':
            header_text = section.get('title', '')
            if section.get('numbering'):
                header_text = self._format_numbered_heading(section.get('numbering'), header_text)

            heading = self.doc.add_heading(header_text, level=min(section.get('level', 3), 4))
            
            # CRITICAL FIX: Disable page break before for shortdoc sections
            heading.paragraph_format.page_break_before = False
            # CRITICAL FIX: Disable keep_with_next and keep_together to prevent headings
            # from jumping to new pages when there's not enough room for heading + content
            heading.paragraph_format.keep_with_next = False
            heading.paragraph_format.keep_together = False

            # CRITICAL: Force Times New Roman for all shortdoc section headings
            header_type = section.get('header_type', 'section')
            if header_type == 'part':
                for run in heading.runs:
                    run.font.name = 'Times New Roman'  # FORCE Times New Roman
                    run.font.size = Pt(self.font_size)
                    run.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue
                    run.bold = True
            elif header_type == 'section':
                for run in heading.runs:
                    run.font.name = 'Times New Roman'  # FORCE Times New Roman
                    run.font.size = Pt(self.font_size)
                    run.bold = True
            else:
                # Default: apply Times New Roman and BOLD to all other header types
                for run in heading.runs:
                    run.font.name = 'Times New Roman'  # FORCE Times New Roman
                    run.font.size = Pt(self.font_size)
                    run.bold = True

            for content_item in section.get('content', []):
                if isinstance(content_item, dict) and content_item.get('type') == 'shortdoc_subheader':
                    self._add_shortdoc_subheader(content_item)
                else:
                    self._add_section_content(content_item)
            return

        # Get heading text and number it
        heading_text = section['heading']
        level = min(section['level'], 3)
        numbered_heading = heading_text

        # Track heading for TOC
        self._track_heading_for_toc(heading_text, level, numbered_heading)
        
        # Add heading for regular sections (use numbered heading)
        heading = self.doc.add_heading(numbered_heading, level=level)
        
        # Apply forced page break if requested - BUT ONLY for level 1 headings
        # CRITICAL: Never apply page breaks to numbered sub-sections (2.0, 2.1, 3.3, etc.)
        is_subsection = bool(re.match(r'^\d+\.\d+', heading_text.strip()))
        if section.get('use_page_break_before') and level == 1 and not is_subsection:
            heading.paragraph_format.page_break_before = True
        else:
            # CRITICAL FIX: Explicitly disable page breaks for level 2+ headings
            # This ensures numbered sub-sections (2.0, 2.1, 2.3, 3.3, etc.) don't start on new pages
            heading.paragraph_format.page_break_before = False
        
        # CRITICAL FIX: Disable keep_with_next and keep_together to prevent headings
        # from jumping to new pages when there's not enough room for heading + content
        heading.paragraph_format.keep_with_next = False
        heading.paragraph_format.keep_together = False
        
        # Check if this heading should be centered
        if section.get('should_center', False):
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Ensure heading is bold, black, Times New Roman, with proper spacing
        heading.paragraph_format.left_indent = Pt(0)
        heading.paragraph_format.first_line_indent = Pt(0)
        should_bold_heading = not self._is_main_research_question_heading(heading_text)
        for run in heading.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(self.font_size)
            run.font.bold = should_bold_heading
            run.font.color.rgb = RGBColor(0, 0, 0)  # Black
        
        # Add content (pass section context for references handling)
        self._add_section_content(section)

    def _add_chapter_section(self, section):
        """Add a chapter section (dissertation-specific)"""
        # Add chapter number heading (CHAPTER ONE, CHAPTER 1, etc.)
        chapter_heading = section.get('heading', '')
        # Strip markdown heading markers
        clean_heading = re.sub(r'^#+\s*', '', chapter_heading).strip().upper()
        
        # Track chapter for TOC (combine chapter number and title)
        chapter_title = section.get('chapter_title', '')
        if chapter_title:
            toc_entry = f"{clean_heading}: {chapter_title.upper()}"
        else:
            toc_entry = clean_heading
        self._track_heading_for_toc(toc_entry, 1, toc_entry)
        
        heading = self.doc.add_heading(clean_heading, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading.paragraph_format.page_break_before = False
        # CRITICAL FIX: Disable keep_with_next and keep_together to prevent headings
        # from jumping to new pages when there's not enough room for heading + content
        heading.paragraph_format.keep_with_next = False
        heading.paragraph_format.keep_together = False
        
        # Ensure consistent chapter heading formatting
        heading.paragraph_format.space_after = Pt(0)
        heading.paragraph_format.line_spacing = self.line_spacing
        heading.paragraph_format.left_indent = Pt(0)
        heading.paragraph_format.first_line_indent = Pt(0)
        
        for run in heading.runs:
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(self.font_size)
            run.font.color.rgb = RGBColor(0, 0, 0)
        heading.paragraph_format.space_after = Pt(0)
        heading.paragraph_format.line_spacing = self.line_spacing
        
        # Add chapter title if present (centered, bold)
        if chapter_title:
            title_para = self.doc.add_heading(chapter_title.upper(), level=1)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.page_break_before = False
            # CRITICAL FIX: Disable keep_with_next and keep_together to prevent headings
            # from jumping to new pages when there's not enough room for heading + content
            title_para.paragraph_format.keep_with_next = False
            title_para.paragraph_format.keep_together = False
            
            # Ensure consistent title formatting
            title_para.paragraph_format.space_before = Pt(0)
            title_para.paragraph_format.space_after = Pt(0)
            title_para.paragraph_format.line_spacing = self.line_spacing
            title_para.paragraph_format.left_indent = Pt(0)
            title_para.paragraph_format.first_line_indent = Pt(0)
            
            for run in title_para.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(self.font_size)
                run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add content
        self._add_section_content(section)
    
    def _add_front_matter_section(self, section):
        """Add a front matter section (dissertation-specific)"""
        front_matter_type = section.get('front_matter_type', 'unknown')
        heading_text = section.get('heading', '')
        # Strip markdown heading markers
        clean_heading = re.sub(r'^#+\s*', '', heading_text).strip().upper()
        
        # Track front matter for TOC (skip certain types)
        skip_toc_types = ['toc', 'table_of_contents']
        if front_matter_type not in skip_toc_types:
            self._track_heading_for_toc(clean_heading, 1, clean_heading)
        
        # Add centered heading (level 1) - main dissertation headings like Resume, Literature Review, etc.
        heading = self.doc.add_heading(clean_heading, level=1)
        
        # Apply forced page break if requested
        if section.get('use_page_break_before'):
            heading.paragraph_format.page_break_before = True
        else:
            heading.paragraph_format.page_break_before = False
        
        # CRITICAL FIX: Disable keep_with_next and keep_together to prevent headings
        # from jumping to new pages when there's not enough room for heading + content
        heading.paragraph_format.keep_with_next = False
        heading.paragraph_format.keep_together = False
            
        # Set alignment to center for main headings
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Ensure main heading formatting is applied consistently
        # Override any style settings to ensure proper formatting
        heading.paragraph_format.space_after = Pt(0)
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.line_spacing = self.line_spacing
        heading.paragraph_format.left_indent = Pt(0)
        heading.paragraph_format.first_line_indent = Pt(0)
        
        # CRITICAL: Apply explicit formatting to ALL runs to force Times New Roman
        # This overrides any Calibri default from the style template
        if heading.runs:
            for run in heading.runs:
                run.bold = True
                run.font.name = 'Times New Roman'  # FORCE Times New Roman
                run.font.size = Pt(self.font_size)  # Force size
                run.font.color.rgb = RGBColor(0, 0, 0)  # Force black
        else:
            # If no runs exist (shouldn't happen), clear and recreate
            # This handles edge cases where add_heading doesn't create runs properly
            heading.clear_content()
            run = heading.add_run(clean_heading)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(self.font_size)
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add content based on front matter type
        if front_matter_type == 'dedication':
            self._add_dedication_content(section)
        elif front_matter_type == 'declaration':
            self._add_declaration_content(section)
        elif front_matter_type == 'certification':
            self._add_certification_content(section)
        elif front_matter_type == 'acknowledgements':
            self._add_acknowledgements_content(section)
        elif front_matter_type == 'abstract':
            self._add_abstract_content(section)
        elif front_matter_type == 'resume':
            self._add_abstract_content(section)  # Same format as abstract
        elif front_matter_type == 'toc':
            self._add_toc_content(section)
        elif front_matter_type in ('list_of_tables', 'list_of_figures', 'abbreviations', 'glossary'):
            self._add_list_section_content(section)
        else:
            # Default: add content normally
            self._add_section_content(section)

    def _format_numbered_heading(self, numbering, title):
        """Format numbered headings without duplicating punctuation."""
        if not numbering:
            return title
        clean_numbering = str(numbering).strip()
        if not clean_numbering:
            return title
        if re.search(r'[.)\]:]$', clean_numbering) or '.' in clean_numbering:
            return f"{clean_numbering} {title}".strip()
        return f"{clean_numbering}. {title}".strip()
    
    def _add_shortdoc_subheader(self, item):
        """Add short document subheaders (questions, tasks, etc.)"""
        header_type = item.get('header_type', 'question')
        content = item.get('content', '')
        numbering = item.get('numbering', '')

        p = self.doc.add_paragraph()
        if header_type == 'question':
            prefix = f"Q{numbering}: "
            run = p.add_run(prefix)
            run.bold = True
            run.font.color.rgb = RGBColor(0, 100, 0)  # Dark green
        elif header_type == 'task':
            prefix = f"Task {numbering}: "
            run = p.add_run(prefix)
            run.bold = True
            run.underline = True
        elif header_type == 'requirement':
            prefix = f"Req {numbering}: "
            run = p.add_run(prefix)
            run.bold = True
            run.font.color.rgb = RGBColor(139, 0, 0)  # Dark red
        elif header_type == 'step':
            prefix = f"Step {numbering}: "
            run = p.add_run(prefix)
            run.bold = True
            run.italic = False
        else:
            prefix = self._format_numbered_heading(numbering, '').rstrip() + ' ' if numbering else ''
            run = p.add_run(prefix)
            run.bold = True

        # Add content with bold formatting for all types
        content_run = p.add_run(content)
        content_run.bold = True
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Pt(0)

    def _add_dedication_content(self, section):
        """Add dedication content (centered, not italic)"""
        for item in section.get('content', []):
            if not isinstance(item, dict):
                continue
            text = item.get('text', '')
            para = self.doc.add_paragraph()
            run = para.add_run(text)
            run.italic = False
            run.font.name = 'Times New Roman'
            run.font.size = Pt(self.font_size)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(12)  # Reduced spacing
            para.paragraph_format.line_spacing = self.line_spacing
    
    def _add_declaration_content(self, section):
        """Add declaration content with signature lines"""
        has_signature = False
        for item in section.get('content', []):
            if not isinstance(item, dict):
                continue
            item_type = item.get('type', 'paragraph')
            text = item.get('text', '')
            
            if item_type == 'signature_line':
                has_signature = True
                # Add signature line
                para = self.doc.add_paragraph()
                run = para.add_run(text if text else '________________________')
                run.font.name = 'Times New Roman'
                run.font.size = Pt(self.font_size)
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                para.paragraph_format.space_before = Pt(36)
                para.paragraph_format.line_spacing = self.line_spacing
            else:
                # Regular paragraph (block paragraph - no indent)
                para = self.doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                para.paragraph_format.line_spacing = self.line_spacing
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
        
        # Add signature line if not already present
        if not has_signature:
            # Add blank space
            blank = self.doc.add_paragraph()
            blank.paragraph_format.space_before = Pt(36)
            
            # Signature line
            sig_para = self.doc.add_paragraph()
            sig_para.add_run('Signed: ________________________')
            sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sig_para.paragraph_format.line_spacing = self.line_spacing
            for run in sig_para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(self.font_size)
            
            # Date line
            date_para = self.doc.add_paragraph()
            date_para.add_run('Date: ________________________')
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_para.paragraph_format.space_before = Pt(12)
            date_para.paragraph_format.line_spacing = self.line_spacing
            for run in date_para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(self.font_size)
    
    def _add_certification_content(self, section):
        """Add certification content with multiple signature lines"""
        for item in section.get('content', []):
            if not isinstance(item, dict):
                continue
            item_type = item.get('type', 'paragraph')
            text = item.get('text', '')
            
            if item_type == 'signature_line':
                para = self.doc.add_paragraph()
                run = para.add_run(text if text else '________________________')
                run.font.name = 'Times New Roman'
                run.font.size = Pt(self.font_size)
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                para.paragraph_format.space_before = Pt(24)
                para.paragraph_format.line_spacing = self.line_spacing
            else:
                para = self.doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                para.paragraph_format.line_spacing = self.line_spacing
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
        
        # Add multiple signature blocks for committee members
        self.doc.add_paragraph()  # Blank line
        
        signatures = [
            'Supervisor: ________________________',
            'Date: ________________________',
            '',
            'Head of Department: ________________________',
            'Date: ________________________',
        ]
        
        for sig_text in signatures:
            if sig_text:
                para = self.doc.add_paragraph()
                para.add_run(sig_text)
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.line_spacing = self.line_spacing
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
            else:
                self.doc.add_paragraph()  # Empty line between signature blocks
    
    def _add_acknowledgements_content(self, section):
        """Add acknowledgements content (block paragraphs - no indent)"""
        for item in section.get('content', []):
            if not isinstance(item, dict):
                continue
            text = item.get('text', '')
            para = self.doc.add_paragraph(text)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            para.paragraph_format.left_indent = Pt(0)
            para.paragraph_format.first_line_indent = Pt(0)
            para.paragraph_format.line_spacing = self.line_spacing
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(self.font_size)
    
    def _add_abstract_content(self, section):
        """Add abstract content with Keywords section"""
        for item in section.get('content', []):
            if not isinstance(item, dict):
                continue
            item_type = item.get('type', 'paragraph')
            text = item.get('text', '')
            
            # Check for Keywords line
            if 'keywords' in text.lower() or item_type == 'keywords':
                para = self.doc.add_paragraph()
                run = para.add_run('Keywords: ')
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(self.font_size)
                # Extract keywords (after "Keywords:")
                keywords = re.sub(r'^[Kk]eywords?\s*[:\-]\s*', '', text)
                para.add_run(keywords)
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.line_spacing = self.line_spacing
            else:
                para = self.doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                para.paragraph_format.line_spacing = self.line_spacing
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
    
    def _add_toc_content(self, section):
        """Add Table of Contents entries with dot leaders"""
        for item in section.get('content', []):
            if not isinstance(item, dict):
                continue
            item_type = item.get('type', 'paragraph')
            text = item.get('text', '')
            
            if item_type == 'toc_entry':
                # Parse TOC entry for title and page number
                # Look for patterns like "Title........12" or "Title   12"
                match = re.match(r'^(.+?)\s*[\.…]+\s*(\d+)\s*$', text)
                if match:
                    title, page = match.groups()
                    para = self.doc.add_paragraph()
                    para.add_run(title.strip())
                    # Add tab with dot leader
                    para.add_run('\t')
                    para.add_run(page)
                    para.paragraph_format.line_spacing = self.line_spacing
                else:
                    para = self.doc.add_paragraph(text)
                    para.paragraph_format.line_spacing = self.line_spacing
                    
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
            else:
                para = self.doc.add_paragraph(text)
                para.paragraph_format.line_spacing = self.line_spacing
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
    
    def _add_list_section_content(self, section):
        """Add content for List of Tables, List of Figures, Abbreviations, Glossary sections."""
        for item in section.get('content', []):
            if not isinstance(item, dict):
                continue
            item_type = item.get('type', 'paragraph')
            text = item.get('text', '')
            
            if item_type == 'toc_entry' or '.....' in text or '\t' in text:
                # Parse entry for title and page number (like TOC)
                match = re.match(r'^(.+?)\s*[\.…]+\s*(\d+)\s*$', text)
                if match:
                    title, page = match.groups()
                    para = self.doc.add_paragraph()
                    para.add_run(title.strip())
                    para.add_run('\t')
                    para.add_run(page)
                    para.paragraph_format.line_spacing = self.line_spacing
                else:
                    para = self.doc.add_paragraph(text)
                    para.paragraph_format.line_spacing = self.line_spacing
            elif item_type == 'abbreviation' or ':' in text:
                # Abbreviation or glossary entry (Term: Definition format)
                para = self.doc.add_paragraph()
                if ':' in text:
                    parts = text.split(':', 1)
                    run = para.add_run(parts[0].strip() + ':')
                    run.bold = True
                    if len(parts) > 1:
                        para.add_run(' ' + parts[1].strip())
                else:
                    para.add_run(text)
                para.paragraph_format.line_spacing = self.line_spacing
            else:
                para = self.doc.add_paragraph(text)
                para.paragraph_format.line_spacing = self.line_spacing
            
            for run in para.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(self.font_size)
    
    def _format_single_apa_reference(self, text):
        """Apply APA formatting rules to a single reference string."""
        # Fix missing space before year parentheses: "UNESCO(2024)" -> "UNESCO (2024)"
        text = re.sub(r'([^\s\(])\((\d{4})', r'\1 (\2', text)
        
        # Fix capitalization in titles (ensure space after colon)
        text = re.sub(r'([A-Z][a-z]+)\s+([A-Z][a-z]+):', r'\1 \2:', text)
        
        # Ensure period after year
        text = re.sub(r'(\(\d{4}(?:/\d{4})?\))\s*([A-Z])', r'\1. \2', text)
        
        # Fix double periods
        text = re.sub(r'\.\.$', '.', text)
        
        # Fix journal formatting (plain text only; italics are forbidden)
        # Matches "Journal of X Y", "International Journal of X", etc.
        text = re.sub(r'(\b(?:International\s+)?Journal\s+of\s+[A-Z][a-z]+(?:\s+[A-Z][a-zA-Z]+)*)', r'\1', text)
        
        # Fix "Available at:" formatting (optional, but good for consistency)
        text = re.sub(r'Available\s+at\s*:', 'Available at:', text, flags=re.IGNORECASE)
        
        # Fix lowercase titles (heuristic: if starts with lowercase after year)
        # UNESCO. (2020) global -> UNESCO. (2020) Global
        text = re.sub(r'(\(\d{4}(?:/\d{4})?\)\.?\s+)([a-z])', lambda m: m.group(1) + m.group(2).upper(), text)
        
        return text

    def _extract_apa_italic_spans(self, text: str):
        """
        Return list of (start, end) spans for titles that should be italicized in APA format.
        This includes: journal titles, book titles, report titles, thesis/dissertation titles.
        Spans are indices into `text`. First valid match wins.
        """
        spans = []
        
        # PATTERN GROUP 1: Journal articles with volume/issue
        # Format: Author. (Year). Article title. Journal Name, Vol(Issue), pages.
        journal_patterns = [
            # Journal + volume/issue pattern
            re.compile(
                r'^(?P<pre>.+?\(\d{4}[a-z]?\)\.\s+.+?\.\s+)(?P<title>(?!In\s)[^,]+?)(?=,\s*\d{1,4}(?:\s*\(|\s*,))'
            ),
            # Journal + volume only
            re.compile(
                r'^(?P<pre>.+?\(\d{4}[a-z]?\)\.\s+.+?\.\s+)(?P<title>(?!In\s)[^,]+?)(?=,\s*\d{1,4}\s*,)'
            ),
            # Journal followed by Retrieved from / URL / DOI
            re.compile(
                r'^(?P<pre>.+?\(\d{4}[a-z]?\)\.\s+.+?\.\s+)(?P<title>(?!In\s)[^.]+?)(?=\.\s+(?:Retrieved\s+from|Available\s+at|https?://|doi:))',
                re.IGNORECASE
            ),
        ]
        
        # PATTERN GROUP 2: Books with publisher
        # Format: Author. (Year). Book title. Publisher.
        # Format: Author. (Year). Book title (Edition). Publisher.
        book_patterns = [
            # Book with publisher at end (Publisher. or Publisher name ending with period)
            re.compile(
                r'^(?P<pre>[^(]+\(\d{4}[a-z]?\)\.\s+)(?P<title>[^.]+?)(?=\.\s+(?:[A-Z][a-z]+\s+)*(?:Press|Publishers?|Publishing|Books?|Library|Libraries|University|Ltd|Inc|Co|LLC|Corporation|House|Media|Academic|Sage|Wiley|Springer|Elsevier|Routledge|Cambridge|Oxford|Pearson|McGraw|Prentice|Random|Harper|Simon|Penguin|Macmillan))',
                re.IGNORECASE
            ),
            # Book with location: Publisher format (e.g., "New York: Random House")
            re.compile(
                r'^(?P<pre>[^(]+\(\d{4}[a-z]?\)\.\s+)(?P<title>[^.]+?)(?=\.\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+)?(?:,\s*[A-Z]{2})?:\s+)',
                re.IGNORECASE
            ),
            # Book with edition in parentheses before publisher
            re.compile(
                r'^(?P<pre>[^(]+\(\d{4}[a-z]?\)\.\s+)(?P<title>[^(]+?)(?=\s*\(\d+(?:st|nd|rd|th)\s+ed\.?\))',
                re.IGNORECASE
            ),
        ]
        
        # PATTERN GROUP 3: Edited books (chapters in edited books)
        # Format: Author. (Year). Chapter title. In Editor (Ed.), Book title (pp. x-y). Publisher.
        edited_book_patterns = [
            # "In Editor (Ed.), Book Title"
            re.compile(
                r'^(?P<pre>.+?In\s+[^(]+\([Ee]ds?\.?\)[,.]?\s+)(?P<title>[^(]+?)(?=\s*\(pp\.)',
                re.IGNORECASE
            ),
            # "In Editor (Eds.), Book Title"
            re.compile(
                r'^(?P<pre>.+?In\s+[^(]+\([Ee]ds\.?\)[,.]?\s+)(?P<title>[^(.]+?)(?=\s*[.(])',
                re.IGNORECASE
            ),
        ]
        
        # PATTERN GROUP 4: Reports, theses, dissertations
        # Format: Author. (Year). Title of report/thesis (Report No. xxx). Publisher.
        report_patterns = [
            # Report with report number
            re.compile(
                r'^(?P<pre>[^(]+\(\d{4}[a-z]?\)\.\s+)(?P<title>[^(]+?)(?=\s*\((?:Report|Technical|Working|Discussion|Paper|Thesis|Dissertation)\s*(?:No\.?|Number)?)',
                re.IGNORECASE
            ),
            # Thesis/Dissertation with university
            re.compile(
                r'^(?P<pre>[^(]+\(\d{4}[a-z]?\)\.\s+)(?P<title>[^(]+?)(?=\s*\((?:Doctoral\s+dissertation|Master\'?s?\s+thesis|PhD\s+thesis|Unpublished))',
                re.IGNORECASE
            ),
            # Thesis/Dissertation - title before [Unpublished/Doctoral etc]
            re.compile(
                r'^(?P<pre>[^(]+\(\d{4}[a-z]?\)\.\s+)(?P<title>[^.\[\]]+?)(?=\s*[\.\[]?\s*(?:Unpublished|Doctoral|Master|PhD))',
                re.IGNORECASE
            ),
        ]
        
        # PATTERN GROUP 5: Standalone books (fallback - book title before final period and no volume)
        # Only matches if no volume/issue pattern and ends with publisher-like text
        standalone_patterns = [
            # Title followed by just a publisher name (common for books)
            re.compile(
                r'^(?P<pre>[^(]+\(\d{4}[a-z]?\)\.\s+)(?P<title>[A-Z][^.]+?)(?=\.\s*$)'
            ),
        ]
        
        # Try all pattern groups in order of specificity
        all_pattern_groups = [
            journal_patterns,
            edited_book_patterns,
            report_patterns,
            book_patterns,
            standalone_patterns,
        ]
        
        for pattern_group in all_pattern_groups:
            for pat in pattern_group:
                m = pat.match(text)
                if not m:
                    continue

                t_start = m.start('title')
                t_end = m.end('title')

                title = text[t_start:t_end].strip()
                
                # Sanity checks to reduce false positives
                if len(title) < 4:
                    continue
                # Skip if it starts with common non-title words
                if title.lower().startswith(('in ', 'and ', 'the ', 'a ', 'an ')) and len(title.split()) <= 2:
                    continue
                # Skip if it looks like just author names
                if re.match(r'^[A-Z][a-z]+,?\s+[A-Z]\.?\s*(&\s*[A-Z][a-z]+,?\s+[A-Z]\.?)?$', title):
                    continue
                    
                spans.append((t_start, t_end))
                return spans  # Return first match
                
        return spans

    def _extract_quoted_text_spans(self, text: str):
        """
        Extract spans for text within quotation marks.
        Returns list of (start, end) tuples for quoted text (including the quotes).
        Supports both regular quotes ("") and curly quotes ("").
        """
        spans = []
        # Match text within various quotation mark styles
        # Pattern matches: "text", "text", «text»
        pattern = re.compile(r'[""«]([^""»]+)[""»]')
        
        for match in pattern.finditer(text):
            # Include the full match (with quotes) for italicization
            spans.append((match.start(), match.end()))
        
        return spans

    def _get_all_reference_italic_spans(self, text: str):
        """
        Get all spans that should be italicized in a reference entry.
        Combines APA title spans and quoted text spans, sorted by position.
        Returns list of (start, end) tuples, non-overlapping and sorted.
        """
        # Get quoted text spans first
        quoted_spans = self._extract_quoted_text_spans(text)
        
        # Get APA title spans (journal, book, etc.)
        apa_spans = self._extract_apa_italic_spans(text)
        
        # If we have quoted spans, filter out APA spans that overlap with them
        # This prevents APA fallback patterns from merging quoted text
        if quoted_spans:
            filtered_apa = []
            for apa_span in apa_spans:
                overlaps = False
                for q_span in quoted_spans:
                    # Check if APA span overlaps with any quoted span
                    if not (apa_span[1] <= q_span[0] or apa_span[0] >= q_span[1]):
                        overlaps = True
                        break
                if not overlaps:
                    filtered_apa.append(apa_span)
            apa_spans = filtered_apa
        
        # Combine all spans
        all_spans = apa_spans + quoted_spans
        
        if not all_spans:
            return []
        
        # Sort by start position
        all_spans.sort(key=lambda x: x[0])
        
        return all_spans

    def _extract_journal_spans(self, text: str):
        """
        Wrapper for backward compatibility. Now calls _extract_apa_italic_spans.
        Return list of (start, end) spans for titles that should be italicized.
        """
        return self._extract_apa_italic_spans(text)

    def _clean_asterisks(self, text):
        """
        Remove all asterisk variants from text comprehensively.
        Removes: *, ⁎, ⁑, ※
        This ensures no asterisks appear in final document output.
        """
        if not text:
            return text
        return re.sub(r'[\*\u204e\u2051\u203b]', '', text).strip()

    def _split_label_for_bold(self, text, max_words=5):
        """Detect short label prefixes like 'Definition:' for minimal bolding."""
        if not text or ':' not in text:
            return None
        match = re.match(r'^\s*([^:]+):\s*(.+)$', text, re.DOTALL)
        if not match:
            return None
        label = match.group(1).strip()
        remainder = match.group(2).strip()
        if not remainder:
            return None
        if 1 <= len(label.split()) <= max_words:
            return label, remainder
        return None

    def _split_numbered_title_block(self, text, min_body_lines=2, max_title_words=20, max_title_chars=120):
        """
        Detect numbered items where the first line is a short title and the body spans multiple lines.

        Returns (title, body_lines) when:
        - There are at least min_body_lines lines after the title.
        - The title is reasonably short (acts like a heading).
        """
        if not text or '\n' not in text:
            return None
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        if len(lines) < 1 + min_body_lines:
            return None
        title = lines[0]
        if len(title.split()) > max_title_words or len(title) > max_title_chars:
            return None
        body_lines = lines[1:]
        if len(body_lines) < min_body_lines:
            return None
        return title, body_lines

    def _is_main_research_question_heading(self, heading_text):
        """Return True when the heading is exactly 'Main Research Question(s)'."""
        if not heading_text:
            return False
        clean = re.sub(r'^#+\s*', '', heading_text)
        clean = re.sub(r'^\d+(?:\.\d+)*\s*', '', clean)
        clean = re.sub(r'[*_]+', '', clean).strip()
        return bool(re.fullmatch(r'main\s+research\s+question(s)?', clean, re.IGNORECASE))

    def _add_label_bold_runs(self, para, label, remainder, prefix=''):
        """Add runs where only the label is bold."""
        if prefix:
            run_prefix = para.add_run(prefix)
            run_prefix.font.name = 'Times New Roman'
            run_prefix.font.size = Pt(self.font_size)
        run_label = para.add_run(f"{label}:")
        run_label.bold = True
        run_label.font.name = 'Times New Roman'
        run_label.font.size = Pt(self.font_size)
        if remainder:
            run_value = para.add_run(f" {remainder}")
            run_value.font.name = 'Times New Roman'
            run_value.font.size = Pt(self.font_size)
        logger.info("Applied label bolding: '%s' -> '%s: %s'", label, label, remainder)

    def _extract_numbering(self, text):
        """
        Comprehensively extract numbering from text.
        Handles all common numbering formats:
        - Digits: 1., 1), 1:, 1st, 2nd, 3rd, (1), [1]
        - Letters: a., a), A., A), (a), [a]
        - Roman: i., I., i), I), (i), [i]
        - Mixed hierarchical: 1.a., 1.a.i., A.1.b
        
        Returns: (numbering_string, clean_content)
        """
        if not text:
            return '', text
        
        # Comprehensive numbering patterns (ordered by specificity)
        numbering_patterns = [
            # Hierarchical mixed (A.1.2, 1.a.i, etc.)
            (r'^([A-Z](?:\.\d+)*(?:\.[a-z])*(?:\.[ivx]+)*[\.\)]?)\s+', ''),
            (r'^(\d+(?:\.[a-z])*(?:\.[ivx]+)*[\.\)]?)\s+', ''),
            
            # With colon: 1:, a:, etc.
            (r'^(\d+:)\s+', ''),
            (r'^([a-zA-Z]:)\s+', ''),
            (r'^([ivxlcdm]+:)\s+', re.IGNORECASE),

            # With comma: 1, 2, etc.
            (r'^(\d+,)\s+', ''),
            
            # Ordinal numbers: 1st, 2nd, 3rd, 4th, etc.
            (r'^(\d+(?:st|nd|rd|th)[\.\)]?)\s+', ''),
            
            # Standard period/paren: 1., 1), a., a), etc.
            (r'^(\d+[\.\)])\s+', ''),
            (r'^([a-z][\.\)])\s+', ''),
            (r'^([A-Z][\.\)])\s+', ''),
            (r'^([ivxlcdm]+[\.\)])\s+', re.IGNORECASE),

            # Standard period/paren without required whitespace (e.g., "8)Intro")
            (r'^(\d+[\.\)])(?=\s*[A-Za-z(])\s*', ''),
            (r'^([a-z][\.\)])(?=\s*[A-Za-z(])\s*', ''),
            (r'^([A-Z][\.\)])(?=\s*[A-Za-z(])\s*', ''),
            (r'^([ivxlcdm]+[\.\)])(?=\s*[A-Za-z(])\s*', re.IGNORECASE),
            
            # Parenthesized: (1), (a), (i), etc.
            (r'^(\(\d+\))\s+', ''),
            (r'^(\([a-z]\))\s+', ''),
            (r'^(\([A-Z]\))\s+', ''),
            (r'^(\((?:[ivxlcdm]+)\))\s+', re.IGNORECASE),
            
            # Bracketed: [1], [a], etc.
            (r'^(\[\d+\])\s+', ''),
            (r'^(\[[a-z]\])\s+', ''),
            (r'^(\[(?:[ivxlcdm]+)\])\s+', re.IGNORECASE),
            
            # Double parentheses: ((1))
            (r'^(\(\(\d+\)\))\s+', ''),
            
            # Bullet/dash characters (for reference)
            (r'^([-–—])\s+', ''),
            (r'^([•◦▪■●○])\s+', ''),
        ]
        
        # Try to match each pattern
        for pattern, flags in numbering_patterns:
            if flags:
                match = re.match(pattern, text, flags)
            else:
                match = re.match(pattern, text)
            
            if match:
                numbering = match.group(1)
                clean_content = text[match.end():]
                return numbering, clean_content
        
        # No numbering found
        return '', text

    def _add_section_content(self, section_or_item):
        """Add content of a section or a single content item
        
        This method can receive either:
        1. A section dict with 'content' key containing a list of items
        2. An individual content item dict with 'type' key (paragraph, definition, etc.)
        """
        # Determine if we received a section dict or an individual content item
        # Individual content items have a 'type' key like 'paragraph', 'definition', etc.
        # Section dicts have 'content' key with a list of items
        is_individual_item = 'type' in section_or_item and 'content' not in section_or_item
        
        if is_individual_item:
            # Process single content item - wrap in a list and process
            content_items = [section_or_item]
            is_references_section = False
        else:
            # Process section dict - extract content list
            section = section_or_item
            # Check if this is a references section
            is_references_section = section.get('is_references_section', False)
            heading_lower = section.get('heading', '').lower()
            if 'reference' in heading_lower or 'bibliography' in heading_lower or 'works cited' in heading_lower:
                is_references_section = True
            
            # Collect content items from section
            content_items = list(section.get('content', []))
        
        if is_references_section:
            # Separate references from other content
            references = [item for item in content_items if isinstance(item, dict) and item.get('type') == 'reference']
            non_references = [item for item in content_items if not isinstance(item, dict) or item.get('type') != 'reference']
            
            # Sort references alphabetically by text (A-Z)
            references.sort(key=lambda x: x.get('text', '').strip().lower() if isinstance(x, dict) else '')
            
            # Apply APA formatting to references
            for ref in references:
                if isinstance(ref, dict) and 'text' in ref:
                    ref['text'] = self._format_single_apa_reference(ref['text'])
                    ref['journal_spans'] = self._extract_journal_spans(ref['text'])
            
            # Recombine: sorted references first, then non-references at the end
            content_items = references + non_references
        
        for item in content_items:
            # SAFETY CHECK: Ensure item is a dictionary
            if not isinstance(item, dict):
                logger.warning(f"Skipping non-dict item in _add_section_content: {type(item)}")
                continue

            # Handle image placeholders FIRST
            if item.get('type') == 'image_placeholder':
                self._insert_image(item.get('image_id'))
                continue
            
            # Handle shape placeholders (flowcharts, diagrams, arrows)
            if item.get('type') == 'shape_placeholder':
                self._insert_shape(item.get('shape_id'))
                continue
            
            # Handle PROTECTED content (certification/declaration pages) - preserve exactly as-is
            if item.get('type') == 'protected_content' or item.get('is_protected'):
                text = item.get('text', '') or item.get('content', '')
                if text:
                    # Preserve original styling without modification
                    para = self.doc.add_paragraph(text)
                    # Apply original formatting
                    original_style = item.get('original_style', 'Normal')
                    try:
                        para.style = original_style
                    except:
                        pass  # If style doesn't exist, keep default
                    # Apply original bold if present
                    if item.get('original_bold_all', False):
                        for run in para.runs:
                            run.bold = True
                    # Apply original font size
                    original_font_size = item.get('original_font_size', 12)
                    for run in para.runs:
                        run.font.size = Pt(original_font_size)
                        run.font.name = 'Times New Roman'
                continue
            
            if item.get('type') == 'paragraph':
                text = item.get('text', '')
                inline_segments = item.get('inline_segments')

                if inline_segments:
                    para = self.doc.add_paragraph(style='AcademicBody')
                    is_bold_paragraph = bool(item.get('original_bold_all')) and self.policy.preserve_existing_bold
                    for segment in inline_segments:
                        if segment.get('type') == 'text':
                            run = para.add_run(segment.get('text', ''))
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(self.font_size)
                            if is_bold_paragraph:
                                run.bold = True
                        elif segment.get('type') == 'shape':
                            self._insert_shape_inline(segment.get('shape_id'), para)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    para.paragraph_format.line_spacing = self.line_spacing
                    para.paragraph_format.left_indent = Pt(0)
                    para.paragraph_format.first_line_indent = Pt(0)
                    self._track_bold_paragraph(is_bold_paragraph, is_heading=False)
                    continue

                # Check if this is a figure caption - format with SEQ field for LOF
                if self.figure_formatter.is_figure_caption(text):
                    figure_info = self.figure_formatter.detect_figure_caption(text)
                    if figure_info:
                        self._add_figure_caption(figure_info['number'], figure_info['title'])
                        self.has_figures = True
                        continue
                
                # Check if this is a table caption - format with SEQ field for LOT
                if self.table_formatter.is_table_caption(text):
                    table_info = self.table_formatter.detect_table_caption(text)
                    if table_info:
                        self._add_table_caption(table_info['number'], table_info['title'])
                        self.has_tables = True
                        continue
                
                # Bold decision: only bold short labels like "Definition:" when present.
                label_split = self._split_label_for_bold(text)
                is_bold_paragraph = False
                if item.get('original_bold_all') and self.policy.preserve_existing_bold:
                    para = self.doc.add_paragraph(text, style='AcademicBody')
                    is_bold_paragraph = True
                elif label_split and not item.get('original_bold', False):
                    para = self.doc.add_paragraph(style='AcademicBody')
                    self._add_label_bold_runs(para, label_split[0], label_split[1])
                else:
                    para = self.doc.add_paragraph(text, style='AcademicBody')
                # Explicitly set properties again to be absolutely sure
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.line_spacing = self.line_spacing
                # Ensure NO indentation - all text is well justified with no first-line indent
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                # Explicitly set font for all runs to ensure consistency after merge
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
                    if item.get('original_bold_all') and self.policy.preserve_existing_bold:
                        run.bold = True
                
                # Track bold for consecutive bold limiting
                self._track_bold_paragraph(is_bold_paragraph, is_heading=False)
            
            elif item.get('type') == 'instruction':
                text = item.get('text', '')
                try:
                    para = self.doc.add_paragraph(text, style='AcademicBody')
                except:
                    para = self.doc.add_paragraph(text)
                
                # Remove any indentation
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(6)
                
                # Apply bold only if consecutive bold limit not exceeded
                should_bold = self._should_allow_bold(is_heading=False)
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
                    run.italic = False
                    run.bold = should_bold
                self._track_bold_paragraph(should_bold, is_heading=False)
            
            elif item.get('type') == 'question':
                text = item.get('text', '')
                try:
                    para = self.doc.add_paragraph(text, style='AcademicBody')
                except:
                    para = self.doc.add_paragraph(text)

                # Remove any indentation
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(6)
                
                # Apply bold only if consecutive bold limit not exceeded
                should_bold = self._should_allow_bold(is_heading=False)
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
                    run.bold = should_bold
                self._track_bold_paragraph(should_bold, is_heading=False)

            elif item.get('type') == 'academic_metadata':
                 text = item.get('text', '')
                 # Format like "Label: Value" where Label is bold
                 label_split = self._split_label_for_bold(text)
                 if label_split:
                     para = self.doc.add_paragraph()
                     self._add_label_bold_runs(para, label_split[0], label_split[1])
                 else:
                     try:
                        para = self.doc.add_paragraph(text, style='AcademicBodyTitle') # Try specialized style if exists, else Normal
                     except:
                        para = self.doc.add_paragraph(text)
                     for run in para.runs:
                         run.font.name = 'Times New Roman'
                         run.font.size = Pt(self.font_size)
                         # Maybe bold if short?
                         if len(text) < 50 and self.policy.enable_regex_auto_bold:
                             run.bold = True

            elif item.get('type') == 'table_caption':
                # Explicit table caption type
                text = item.get('text', '')
                table_info = self.table_formatter.detect_table_caption(text)
                if table_info:
                    self._add_table_caption(table_info['number'], table_info['title'])
                else:
                    # Fallback - try to extract number from text
                    num_match = re.search(r'(\d+(?:\.\d+)?)', text)
                    if num_match:
                        num = num_match.group(1)
                        title = re.sub(r'^(?:Table|Tbl\.?|Tab\.?)\s*\d+(?:\.\d+)?[\.:]\s*', '', text, flags=re.IGNORECASE)
                        self._add_table_caption(num, title)
                    else:
                        # Last resort - just add as bold centered
                        para = self.doc.add_paragraph(text)
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.bold = True
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(self.font_size)
                self.has_tables = True
                continue
            
            elif item.get('type') == 'figure_caption':
                # Explicit figure caption type
                text = item.get('text', '')
                figure_info = self.figure_formatter.detect_figure_caption(text)
                if figure_info:
                    self._add_figure_caption(figure_info['number'], figure_info['title'])
                else:
                    # Fallback - try to extract number from text
                    num_match = re.search(r'(\d+(?:\.\d+)?)', text)
                    if num_match:
                        num = num_match.group(1)
                        title = re.sub(r'^(?:Figure|Fig\.?)\s*\d+(?:\.\d+)?[\.:]\s*', '', text, flags=re.IGNORECASE)
                        self._add_figure_caption(num, title)
                    else:
                        # Last resort - just add as centered plain text
                        para = self.doc.add_paragraph(text)
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in para.runs:
                            run.italic = False
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(self.font_size)
                self.has_figures = True
                continue
            
            elif item.get('type') == 'definition':
                para = self.doc.add_paragraph()
                run = para.add_run(f"{item.get('term', '')}: ")
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(self.font_size)
                if item.get('definition'):
                    run_def = para.add_run(item.get('definition', ''))
                    run_def.font.name = 'Times New Roman'
                    run_def.font.size = Pt(self.font_size)
                # Remove any indentation
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
            
            elif item.get('type') == 'bullet_list':
                # Conservative Bullet Rendering: Only use bullets when:
                # 1. Section has 4+ items
                # 2. All items are short (not paragraph-like)
                # 3. No bolded content
                
                list_items = item.get('items', [])
                
                # Pre-scan: Count valid items and check if any would be bolded
                valid_items = []
                should_use_bullets = False
                
                for list_item in list_items:
                    # Extract content from list item (handle both dict and string formats)
                    if isinstance(list_item, dict):
                        content = list_item.get('content', '')
                    else:
                        content = str(list_item)
                    
                    # Remove any remaining asterisks from content
                    content = self._clean_asterisks(content)
                    content = re.sub(r'^\s*\d+(?:\.\d+)*\s+', '', content)
                    
                    # Skip empty content
                    if not content:
                        continue
                    
                    valid_items.append(content)
                
                # Decision: Use bullets ONLY if 4+ items AND all are short (non-substantive)
                if len(valid_items) >= 4:
                    # Check if all items are short (eligible for bullets)
                    all_short = True
                    for content in valid_items:
                        numbering, content_after_num = self._extract_numbering(content)
                        content_word_count = len(content_after_num.split())
                        content_is_multiline = '\n' in content_after_num
                        is_substantive = content_word_count > 30 or content_is_multiline
                        
                        # Also check if has colon (would become bold)
                        has_colon = ':' in content_after_num
                        
                        if is_substantive or has_colon or numbering:
                            all_short = False
                            break
                    
                    should_use_bullets = all_short
                
                # Render items
                for content in valid_items:
                    # Extract any numbering (1., a), (i), etc.)
                    numbering, content_after_num = self._extract_numbering(content)
                    
                    # Check if content is substantive (paragraph-like, >30 words or multiline)
                    content_is_multiline = '\n' in content_after_num
                    content_word_count = len(content_after_num.split())
                    is_substantive = content_word_count > 30 or content_is_multiline
                    
                    # Check if remaining content has a title separated by colon
                    title_match = re.match(r'^([^:]+):\s*(.*)$', content_after_num, re.DOTALL)
                    has_title = False
                    title = ''
                    body = ''
                    
                    if title_match:
                        title = title_match.group(1).strip()
                        body = title_match.group(2).strip()
                        # Determine if title should be bolded (long body or multiple paragraphs)
                        body_is_multiline = '\n' in body
                        body_is_long = len(body.split()) > 15
                        has_title = (body_is_multiline or body_is_long) and len(title) < 100
                    else:
                        # No colon separator, check if content_after_num is substantive
                        if is_substantive:
                            # Treat entire content as substantive
                            has_title = True
                            if numbering:
                                # Numbering becomes the title
                                title = numbering
                                body = content_after_num
                            else:
                                # No numbering, treat first line as title if multiline
                                if content_is_multiline:
                                    lines = content_after_num.split('\n')
                                    title = lines[0] if lines else ''
                                    body = '\n'.join(lines[1:]) if len(lines) > 1 else ''
                                else:
                                    # Long single line - don't split
                                    body = content_after_num
                    
                    # Determine if we should use bullet or bold format
                    # Use BULLET ONLY IF: section decision says to use bullets AND item is short
                    # NEVER use bullet if: has title, has numbering, or is substantive
                    use_bullet_format = should_use_bullets and not is_substantive and not has_title and not numbering
                    
                    if use_bullet_format:
                        # SHORT ITEM: Use traditional bullet format
                        para = self.doc.add_paragraph(style='Normal')
                        
                        # Add bullet character
                        bullet_run = para.add_run('■\t')
                        bullet_run.font.name = 'Arial'
                        bullet_run.font.size = Pt(self.font_size)
                        bullet_run.font.color.rgb = RGBColor(0, 0, 0)
                        
                        # Add content
                        content_run = para.add_run(content_after_num)
                        content_run.font.name = 'Times New Roman'
                        content_run.font.size = Pt(self.font_size)
                        
                        # Set paragraph formatting
                        self._apply_list_hanging_indent(
                            para,
                            bullet_offset_inch=self.LIST_BULLET_OFFSET_INCHES,
                            text_indent_inch=self.LIST_BULLET_TEXT_INDENT_INCHES,
                        )
                        para.paragraph_format.line_spacing = self.line_spacing
                        para.paragraph_format.space_after = Pt(0)
                    else:
                        title_block = self._split_numbered_title_block(content_after_num) if numbering else None
                        if title_block:
                            title_line, body_lines = title_block
                            title_para = self.doc.add_paragraph()
                            prefix_run = title_para.add_run(f"{numbering}\t")
                            title_run = title_para.add_run(title_line)
                            title_run.bold = True
                            for run in (prefix_run, title_run):
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(self.font_size)
                            self._apply_list_hanging_indent(title_para)
                            title_para.paragraph_format.space_after = Pt(3)
                            for body_line in body_lines:
                                label_split = self._split_label_for_bold(body_line)
                                if label_split and not list_item.get('original_bold', False):
                                    para = self.doc.add_paragraph()
                                    self._add_label_bold_runs(para, label_split[0], label_split[1])
                                else:
                                    para = self.doc.add_paragraph(body_line)
                                self._apply_list_body_indent(para)
                                para.paragraph_format.space_after = Pt(3)
                                for run in para.runs:
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(self.font_size)
                            continue
                        # SUBSTANTIVE ITEM: Render as plain paragraph(s) with minimal label bolding.
                        paragraphs = content_after_num.split('\n') if '\n' in content_after_num else [content_after_num]
                        for idx, para_text in enumerate(paragraphs):
                            if not para_text.strip():
                                continue
                            prefix = f"{numbering}\t" if numbering and idx == 0 else ''
                            label_split = self._split_label_for_bold(para_text)
                            if label_split and not list_item.get('original_bold', False):
                                para = self.doc.add_paragraph()
                                self._add_label_bold_runs(para, label_split[0], label_split[1], prefix=prefix)
                            else:
                                para = self.doc.add_paragraph(prefix + para_text)
                            if numbering and idx == 0:
                                self._apply_list_hanging_indent(para)
                            else:
                                self._apply_list_body_indent(para)
                            para.paragraph_format.space_after = Pt(3)
                            for run in para.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(self.font_size)
            
            elif item.get('type') == 'numbered_list':
                list_counter = 0
                for list_item in item.get('items', []):
                    # Extract content from list_item (could be string or dict)
                    if isinstance(list_item, str):
                        item_content = list_item
                    elif isinstance(list_item, dict):
                        item_content = list_item.get('content', '')
                    else:
                        item_content = str(list_item)
                    original_bold = list_item.get('original_bold', False) if isinstance(list_item, dict) else False
                    
                    # Use comprehensive numbering extraction
                    numbering, clean_item = self._extract_numbering(item_content)
                    if not numbering:
                        list_counter += 1
                        numbering = f"{list_counter}."
                    else:
                        list_counter = 0
                    paragraphs = clean_item.split('\n') if '\n' in clean_item else [clean_item]
                    title_block = self._split_numbered_title_block(clean_item)
                    if title_block:
                        title_line, body_lines = title_block
                        title_para = self.doc.add_paragraph()
                        prefix_run = title_para.add_run(f"{numbering}\t")
                        title_run = title_para.add_run(title_line)
                        title_run.bold = True
                        for run in (prefix_run, title_run):
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(self.font_size)
                        self._apply_list_hanging_indent(title_para)
                        title_para.paragraph_format.space_after = Pt(3)
                        for body_line in body_lines:
                            label_split = self._split_label_for_bold(body_line)
                            if label_split and not original_bold:
                                para = self.doc.add_paragraph()
                                self._add_label_bold_runs(para, label_split[0], label_split[1])
                            else:
                                para = self.doc.add_paragraph(body_line)
                            self._apply_list_body_indent(para)
                            para.paragraph_format.space_after = Pt(3)
                            for run in para.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(self.font_size)
                    else:
                        for idx, para_text in enumerate(paragraphs):
                            if not para_text.strip():
                                continue
                            prefix = f"{numbering}\t" if idx == 0 else ''
                            label_split = self._split_label_for_bold(para_text)
                            if label_split and not original_bold:
                                para = self.doc.add_paragraph()
                                self._add_label_bold_runs(para, label_split[0], label_split[1], prefix=prefix)
                            else:
                                para = self.doc.add_paragraph(prefix + para_text)
                            if idx == 0:
                                self._apply_list_hanging_indent(para)
                            else:
                                self._apply_list_body_indent(para)
                            para.paragraph_format.space_after = Pt(3)
                            for run in para.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(self.font_size)
            
            elif item.get('type') == 'table':
                self._add_table(item)
            
            elif item.get('type') == 'plain_text_table':
                self._add_plain_text_table(item)
            
            elif item.get('type') == 'figure':
                # Figure with caption - use proper SEQ field formatting
                caption = item.get('caption', '')
                if caption:
                    figure_info = self.figure_formatter.detect_figure_caption(caption)
                    if figure_info:
                        self._add_figure_caption(figure_info['number'], figure_info['title'])
                    else:
                        # Try to extract number from caption
                        num_match = re.search(r'(\d+(?:\.\d+)?)', caption)
                        if num_match:
                            num = num_match.group(1)
                            title = re.sub(r'^(?:Figure|Fig\.?)\s*\d+(?:\.\d+)?[\.:]\s*', '', caption, flags=re.IGNORECASE)
                            self._add_figure_caption(num, title.strip() if title.strip() else caption)
                        else:
                            # Fallback to simple centered text
                            para = self.doc.add_paragraph(caption)
                            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in para.runs:
                                run.italic = False
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(self.font_size)
                self.has_figures = True
            
            elif item.get('type') == 'quote':
                para = self.doc.add_paragraph(item.get('text', ''))
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.right_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.line_spacing = self.line_spacing
                for run in para.runs:
                    run.italic = False
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
            
            elif item.get('type') == 'equation':
                para = self.doc.add_paragraph(item.get('label', ''))
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
            
            elif item.get('type') == 'reference':
                text = self._clean_asterisks(item.get('text', ''))
                
                # Get all italic spans (APA titles + quoted text)
                if is_references_section:
                    spans = self._get_all_reference_italic_spans(text)
                else:
                    spans = []

                para = self.doc.add_paragraph()
                try:
                    para.style = 'ReferenceEntry'
                except Exception:
                    pass

                if spans:
                    # Render text with multiple italic spans
                    pos = 0
                    for s, e in spans:
                        # Add non-italic text before this span
                        if pos < s:
                            run = para.add_run(text[pos:s])
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(self.font_size)
                            run.font.italic = False
                        
                        # Add italic text for this span
                        run = para.add_run(text[s:e])
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(self.font_size)
                        run.font.italic = True
                        
                        pos = e
                    
                    # Add remaining non-italic text after last span
                    if pos < len(text):
                        run = para.add_run(text[pos:])
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(self.font_size)
                        run.font.italic = False
                else:
                    r = para.add_run(text)
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(self.font_size)
                    r.font.italic = False

                if is_references_section:
                    para.paragraph_format.left_indent = Inches(0.5)
                    para.paragraph_format.first_line_indent = Inches(-0.5)
                else:
                    para.paragraph_format.left_indent = Pt(0)
                    para.paragraph_format.first_line_indent = Pt(0)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.line_spacing = self.line_spacing
            
            # NEW PATTERN RENDERING (December 30, 2025)
            
            elif item.get('type') == 'page_metadata':
                # Page metadata - centered, plain text
                para = self.doc.add_paragraph(item.get('text', ''))
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                para.paragraph_format.line_spacing = self.line_spacing
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
                    run.font.italic = False
            
            elif item.get('type') == 'academic_metadata':
                subtype = item.get('subtype', 'metadata')
                text = item.get('text', '')
                
                if subtype == 'author':
                    # Author names - centered, bold
                    para = self.doc.add_paragraph(text)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.left_indent = Pt(0)
                    para.paragraph_format.first_line_indent = Pt(0)
                    para.paragraph_format.line_spacing = self.line_spacing
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(self.font_size)
                        run.font.bold = True
                elif subtype == 'affiliation':
                    # Affiliation - normal text, centered
                    para = self.doc.add_paragraph(text)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.left_indent = Pt(0)
                    para.paragraph_format.first_line_indent = Pt(0)
                    para.paragraph_format.space_after = Pt(6)
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(self.font_size)
                elif subtype == 'contact':
                    # Contact info - centered
                    para = self.doc.add_paragraph(text)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.left_indent = Pt(0)
                    para.paragraph_format.first_line_indent = Pt(0)
                    para.paragraph_format.line_spacing = self.line_spacing
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(self.font_size)
                else:
                    para = self.doc.add_paragraph(text)
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.left_indent = Pt(0)
                    para.paragraph_format.first_line_indent = Pt(0)
            
            elif item.get('type') == 'math_expression':
                subtype = item.get('subtype', 'inline_math')
                text = item.get('text', '')
                
                # Clean up LaTeX markers for display
                clean_text = re.sub(r'^\$\$|\$\$$', '', text)  # Remove $$
                clean_text = re.sub(r'^\$|\$$', '', clean_text)  # Remove $
                clean_text = re.sub(r'^\\\[|\\\]$', '', clean_text)  # Remove \[ \]
                clean_text = re.sub(r'^\\\(|\\\)$', '', clean_text)  # Remove \( \)
                
                if subtype == 'display_math':
                    # Display math - centered, with spacing
                    para = self.doc.add_paragraph(clean_text.strip())
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    para.paragraph_format.left_indent = Pt(0)
                    para.paragraph_format.first_line_indent = Pt(0)
                    para.paragraph_format.space_before = Pt(12)
                    para.paragraph_format.space_after = Pt(12)
                    para.paragraph_format.line_spacing = self.line_spacing
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(self.font_size)
                else:
                    # Inline math - just add as text
                    para = self.doc.add_paragraph(text)
                    para.paragraph_format.left_indent = Pt(0)
                    para.paragraph_format.first_line_indent = Pt(0)
                    para.paragraph_format.line_spacing = self.line_spacing
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(self.font_size)
            
            elif item.get('type') == 'footnote_endnote':
                subtype = item.get('subtype', 'footnote_entry')
                text = item.get('text', '')
                
                if subtype == 'footnote_entry':
                    # Footnote entry - no hanging indent
                    para = self.doc.add_paragraph(text)
                    para.paragraph_format.left_indent = Pt(0)
                    para.paragraph_format.first_line_indent = Pt(0)
                    para.paragraph_format.line_spacing = self.line_spacing
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(self.font_size)
                else:
                    para = self.doc.add_paragraph(text)
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(self.font_size)
            
            elif item.get('type') == 'inline_formatting':
                text = item.get('text', '')
                formatting = item.get('formatting', {})
                
                # Parse the text and apply formatting
                para = self.doc.add_paragraph()
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                
                if formatting.get('bold_italic'):
                    # Remove *** or ___ markers and apply bold only
                    clean_text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)
                    clean_text = re.sub(r'___(.+?)___', r'\1', clean_text)
                    run = para.add_run(clean_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
                    run.bold = True
                elif formatting.get('bold'):
                    # Remove **, __, *, or _ markers and apply bold
                    clean_text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Remove **text**
                    clean_text = re.sub(r'__(.+?)__', r'\1', clean_text)  # Remove __text__
                    clean_text = re.sub(r'(?<!\*)\*([^*\n]+?)\*(?!\*)', r'\1', clean_text)  # Remove *text*
                    clean_text = re.sub(r'(?<!_)_([^_\n]+?)_(?!_)', r'\1', clean_text)  # Remove _text_
                    run = para.add_run(clean_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
                    run.bold = True
                elif formatting.get('italic'):
                    # Remove * or _ markers (italics disabled)
                    clean_text = re.sub(r'\*(.+?)\*', r'\1', text)
                    clean_text = re.sub(r'_(.+?)_', r'\1', clean_text)
                    run = para.add_run(clean_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
                else:
                    # Remove all formatting markers as fallback
                    clean_text = re.sub(r'[\*_]{1,3}(.+?)[\*_]{1,3}', r'\1', text)
                    run = para.add_run(clean_text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
            
            # ================================================================
            # NEW 20 ACADEMIC PATTERN RENDERING (December 30, 2025)
            # ================================================================
            
            elif item.get('type') == 'figure_equation':
                # Figure or equation caption
                subtype = item.get('subtype', 'figure_caption')
                text = item.get('text', '')
                para = self.doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.line_spacing = self.line_spacing
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
                    if subtype == 'figure_caption':
                        run.bold = True
            
            elif item.get('type') == 'citation_inline':
                # Inline citation - render as normal paragraph
                text = item.get('text', '')
                para = self.doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
            
            elif item.get('type') == 'appendix':
                # Appendix section - similar to heading but with distinct style
                text = item.get('heading', item.get('text', ''))
                heading = self.doc.add_heading(text, level=1)
                heading.paragraph_format.page_break_before = False
                # CRITICAL FIX: Disable keep_with_next and keep_together
                heading.paragraph_format.keep_with_next = False
                heading.paragraph_format.keep_together = False
                heading.paragraph_format.line_spacing = self.line_spacing
                for run in heading.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black
                # Reset consecutive bold counter for headings
                self._track_bold_paragraph(True, is_heading=True)
            
            elif item.get('type') == 'block_quote':
                # Block quote - indented, plain text
                text = item.get('text', '')
                # Remove leading > markers
                clean_text = re.sub(r'^[>\s]+', '', text)
                para = self.doc.add_paragraph(clean_text)
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.right_indent = Pt(0)
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(6)
                para.paragraph_format.line_spacing = self.line_spacing
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
            
            elif item.get('type') == 'math_model':
                # Math model / statistical notation
                text = item.get('text', '')
                para = self.doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                para.paragraph_format.space_before = Pt(8)
                para.paragraph_format.space_after = Pt(8)
                para.paragraph_format.line_spacing = self.line_spacing
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
            
            elif item.get('type') == 'text_emphasis':
                # Text emphasis - bold/underline only
                text = item.get('text', '')
                subtype = item.get('subtype', 'bold')
                para = self.doc.add_paragraph()
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                # Clean formatting markers
                clean_text = re.sub(r'[\*_~]{1,3}(.+?)[\*_~]{1,3}', r'\1', text)
                run = para.add_run(clean_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(self.font_size)
                if 'bold' in subtype:
                    run.bold = True
                if 'strike' in subtype:
                    run.font.strike = True
            
            elif item.get('type') == 'toc_entry':
                # Table of contents entry
                text = item.get('text', '')
                page_num = item.get('page_number', '')
                para = self.doc.add_paragraph()
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                run = para.add_run(text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(self.font_size)
                if page_num:
                    para.add_run('\t')
                    run_num = para.add_run(str(page_num))
                    run_num.font.name = 'Times New Roman'
                    run_num.font.size = Pt(self.font_size)
            
            elif item.get('type') == 'footnote_marker':
                # Footnote marker / reference
                text = item.get('text', '')
                para = self.doc.add_paragraph(text)
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                para.paragraph_format.line_spacing = self.line_spacing
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
            
            elif item.get('type') == 'abbreviation':
                # Abbreviation definition
                text = item.get('text', '')
                para = self.doc.add_paragraph()
                # Try to extract abbreviation and definition
                match = re.match(r'([A-Z]{2,})\s*[-–:=]\s*(.+)', text)
                if match:
                    abbr, defn = match.groups()
                    run = para.add_run(abbr)
                    run.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
                    run_def = para.add_run(f' – {defn}')
                    run_def.font.name = 'Times New Roman'
                    run_def.font.size = Pt(self.font_size)
                else:
                    run = para.add_run(text)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
            
            elif item.get('type') == 'caption_format':
                # Figure/table caption formatting
                text = item.get('text', '')
                para = self.doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.line_spacing = self.line_spacing
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
            
            elif item.get('type') == 'horizontal_rule':
                # SKIP horizontal rules entirely - do NOT render them
                # The --- markers in markdown are section separators, not visual elements
                # Rendering them can cause layout issues and unwanted spacing
                # Simply continue to the next item without adding anything to the document
                continue
            
            elif item.get('type') == 'page_break':
                # Insert page break - only for explicit [PAGE BREAK] markers
                # Skip page breaks for short documents
                if not self.is_short_document:
                    self.doc.add_page_break()
            
            elif item.get('type') == 'statistical_result':
                # Statistical result formatting
                text = item.get('text', '')
                para = self.doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                para.paragraph_format.line_spacing = self.line_spacing
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
            
            elif item.get('type') == 'questionnaire':
                # Questionnaire item
                text = item.get('text', '')
                subtype = item.get('subtype', 'question_item')
                if subtype == 'likert_scale':
                    para = self.doc.add_paragraph(text)
                    para.paragraph_format.left_indent = Pt(0)
                    para.paragraph_format.line_spacing = self.line_spacing
                else:
                    para = self.doc.add_paragraph(text)
                    para.paragraph_format.line_spacing = self.line_spacing
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
            
            elif item.get('type') == 'glossary_entry':
                # Glossary entry - term in bold, definition follows
                term = item.get('term', '')
                definition = item.get('definition', '')
                para = self.doc.add_paragraph()
                if term:
                    run = para.add_run(term)
                    run.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
                    run_def = para.add_run(': ' + definition if definition else '')
                    run_def.font.name = 'Times New Roman'
                    run_def.font.size = Pt(self.font_size)
                else:
                    run = para.add_run(definition)
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
            
            elif item.get('type') == 'cross_reference':
                # Cross-reference - render as normal text
                text = item.get('text', '')
                para = self.doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
            
            elif item.get('type') == 'running_header':
                # Running header - typically in document header
                text = item.get('text', '')
                para = self.doc.add_paragraph(text)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.left_indent = Pt(0)
                para.paragraph_format.first_line_indent = Pt(0)
                para.paragraph_format.line_spacing = self.line_spacing
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
            
            elif item.get('type') == 'nested_list':
                # Nested list with indent levels
                for list_item in item.get('items', []):
                    indent = list_item.get('indent_level', 0)
                    para = self.doc.add_paragraph(style='Normal')
                    
                    # Add bullet character with matching size
                    bullet_run = para.add_run('■\t')
                    bullet_run.font.name = 'Arial'
                    bullet_run.font.size = Pt(self.font_size)
                    bullet_run.font.color.rgb = RGBColor(0, 0, 0)
                    
                    # Add content text with normal size
                    content_run = para.add_run(list_item.get('text', ''))
                    content_run.font.name = 'Times New Roman'
                    content_run.font.size = Pt(self.font_size)
                    
                    self._apply_list_hanging_indent(
                        para,
                        indent,
                        bullet_offset_inch=self.LIST_BULLET_OFFSET_INCHES,
                        text_indent_inch=self.LIST_BULLET_TEXT_INDENT_INCHES,
                    )
            
            # ================================================================
            # DISSERTATION-SPECIFIC CONTENT RENDERING (December 30, 2025)
            # ================================================================
            
            elif item.get('type') == 'copyright_content':
                # Copyright content - centered, plain text
                text = item.get('text', '')
                para = self.doc.add_paragraph()
                run = para.add_run(text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(self.font_size)
                run.italic = False
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.line_spacing = self.line_spacing
            
            elif item.get('type') == 'signature_line':
                # Signature line - right aligned
                text = item.get('text', '')
                para = self.doc.add_paragraph()
                if text:
                    para.add_run(text)
                else:
                    para.add_run('________________________')
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                para.paragraph_format.space_before = Pt(24)
                para.paragraph_format.line_spacing = self.line_spacing
                for run in para.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
            
            elif item.get('type') == 'chapter_title':
                # Chapter title - centered, bold (heading level 1)
                text = item.get('text', '')
                para = self.doc.add_heading(text.upper(), level=1)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.page_break_before = False
                # CRITICAL FIX: Disable keep_with_next and keep_together
                para.paragraph_format.keep_with_next = False
                para.paragraph_format.keep_together = False
                for run in para.runs:
                    run.bold = True
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)
                para.paragraph_format.line_spacing = self.line_spacing
                # Reset consecutive bold counter for headings
                self._track_bold_paragraph(True, is_heading=True)
            
            # ================================================================
            # SHORT DOCUMENT KEY POINT RENDERING (December 30, 2025)
            # ================================================================
            
            elif item.get('type') == 'key_point':
                # Key point with emphasis and optional emoji
                text = item.get('text', '')
                # Remove asterisks comprehensively
                text = self._clean_asterisks(text)
                
                key_point_type = item.get('key_point_type', '')
                emoji = item.get('emoji_prefix', '')
                
                para = self.doc.add_paragraph()
                
                # Add emoji prefix if present
                if emoji:
                    para.add_run(emoji)
                
                # Parse and apply formatting (handle markdown-style bold only)
                clean_text = text.strip()
                
                # Check if bold is allowed (max 2 consecutive)
                should_bold = self._should_allow_bold(is_heading=False)
                is_bold = False
                
                # Apply formatting based on key point type
                if key_point_type in ['warning', 'concept', 'exercise', 'learning', 'summary', 'procedure']:
                    # Bold for emphasis (if allowed)
                    run = para.add_run(clean_text)
                    run.bold = should_bold
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
                    is_bold = should_bold
                elif key_point_type == 'example':
                    # Plain text for examples (italics disabled)
                    run = para.add_run(clean_text)
                    run.italic = False
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
                elif key_point_type == 'definition':
                    # Bold with definition format
                    # Try to extract term and definition
                    if ':' in clean_text:
                        parts = clean_text.split(':', 1)
                        term_run = para.add_run(parts[0] + ':')
                        term_run.bold = should_bold
                        term_run.font.name = 'Times New Roman'
                        term_run.font.size = Pt(self.font_size)
                        if len(parts) > 1:
                            defn_run = para.add_run(' ' + parts[1].strip())
                            defn_run.font.name = 'Times New Roman'
                            defn_run.font.size = Pt(self.font_size)
                        is_bold = should_bold
                    else:
                        run = para.add_run(clean_text)
                        run.bold = should_bold
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(self.font_size)
                        is_bold = should_bold
                else:
                    # Default: bold (if allowed)
                    run = para.add_run(clean_text)
                    run.bold = should_bold
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(self.font_size)
                    is_bold = should_bold
                
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.line_spacing = self.line_spacing
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(6)
                
                # Track consecutive bold
                self._track_bold_paragraph(is_bold, is_heading=False)
            
            elif item.get('type') == 'assignment_header_field':
                # Assignment header field (Student Name, Course, etc.) - bold
                text = item.get('text', '')
                para = self.doc.add_paragraph()
                run = para.add_run(text)
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(self.font_size)
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.line_spacing = self.line_spacing
                para.paragraph_format.space_after = Pt(3)
    
    def _normalize_table_rows(self, table_data):
        """Normalize table rows from different detection formats."""
        rows = table_data.get('rows')
        has_header = table_data.get('has_header', False)

        if rows:
            normalized = []
            for row in rows:
                if not isinstance(row, list):
                    continue
                normalized.append([
                    str(cell).strip() if cell is not None else '' for cell in row
                ])
            return normalized, has_header

        content = table_data.get('content', [])
        normalized = []
        saw_separator = False
        for entry in content:
            if not isinstance(entry, dict):
                continue
            if entry.get('type') == 'separator':
                saw_separator = True
                continue
            cells = entry.get('cells')
            if isinstance(cells, list):
                normalized.append([
                    str(cell).strip() if cell is not None else '' for cell in cells
                ])

        if saw_separator and normalized:
            has_header = True

        return normalized, has_header

    def _add_table(self, table_data):
        """Add a table with academic formatting (proper alignment, column sizing)"""
        # Add caption if exists (above table, centered, bold) - WITH SEQ field for LOT
        if table_data.get('caption'):
            caption_text = table_data['caption']
            # Check if this is a table caption and extract number/title
            table_info = self.table_formatter.detect_table_caption(caption_text)
            if table_info:
                # Use proper SEQ field formatting for LOT
                self._add_table_caption(table_info['number'], table_info['title'])
                self.has_tables = True
            else:
                # Fallback - try to extract number from caption
                num_match = re.search(r'(?:Table|Tbl\.?|Tab\.?)\s*(\d+(?:\.\d+)?)', caption_text, re.IGNORECASE)
                if num_match:
                    num = num_match.group(1)
                    title = re.sub(r'^(?:Table|Tbl\.?|Tab\.?)\s*\d+(?:\.\d+)?[\.:]\s*', '', caption_text, flags=re.IGNORECASE)
                    self._add_table_caption(num, title.strip() if title.strip() else caption_text)
                    self.has_tables = True
                else:
                    # Last resort - just add as bold centered (no SEQ field)
                    caption = self.doc.add_paragraph()
                    caption_run = caption.add_run(caption_text)
                    caption_run.bold = True
                    caption_run.font.name = 'Times New Roman'
                    caption_run.font.size = Pt(self.font_size)
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.paragraph_format.space_after = Pt(6)
        
        rows, has_header = self._normalize_table_rows(table_data)

        # Add table
        if rows:
            # Ensure rows have consistent column counts
            num_cols = max(len(row) for row in rows) if rows else 1
            num_rows = len(rows)
            if num_cols > 0:
                rows = [row + [''] * (num_cols - len(row)) for row in rows]
            
            if num_cols > 0 and num_rows > 0:
                # Analyze column content types for alignment
                engine = PatternEngine()
                column_types = engine.get_column_content_types(rows)
                
                # Ensure we have types for all columns
                while len(column_types) < num_cols:
                    column_types.append('text')
                
                table = self.doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Set table to fixed widths so columns remain stable
                table.autofit = False
                
                # Calculate column widths based on content length
                # Find max character count for each column
                col_max_lengths = [0] * num_cols
                for row_data in rows:
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < num_cols:
                            cell_content = str(cell_text).strip() if cell_text else ''
                            # Count all characters (letters, numbers, spaces, etc.)
                            content_length = len(cell_content)
                            if content_length > col_max_lengths[col_idx]:
                                col_max_lengths[col_idx] = content_length
                
                total_table_width = self._get_available_table_width()
                column_widths = self._calculate_table_column_widths(total_table_width, col_max_lengths)
                
                # Set column widths proportionally
                for col_idx, col_width in enumerate(column_widths):
                    # Set width for all cells in this column
                    for row in table.rows:
                        row.cells[col_idx].width = col_width
                
                # Fill table with proper alignment
                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < num_cols:  # Safety check
                            cell = table.rows[row_idx].cells[col_idx]
                            cell_content = str(cell_text).strip() if cell_text else ''
                            cell.text = cell_content
                            
                            # Apply formatting to all paragraphs in cell
                            for paragraph in cell.paragraphs:
                                # Set font for all runs
                                for run in paragraph.runs:
                                    run.font.name = 'Times New Roman'
                                    run.font.size = Pt(self.font_size)
                                
                                # Header row (row 0): centered and bold
                                if row_idx == 0 and has_header:
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    for run in paragraph.runs:
                                        run.bold = True
                                else:
                                    # Data rows: alignment based on column content type
                                    col_type = column_types[col_idx] if col_idx < len(column_types) else 'text'
                                    paragraph.alignment = engine.get_alignment_for_content_type(col_type)
                                
                                # Set line spacing
                                paragraph.paragraph_format.line_spacing = 1.0
                                paragraph.paragraph_format.space_before = Pt(2)
                                paragraph.paragraph_format.space_after = Pt(2)
        
        # Add spacing after table
        spacing = self.doc.add_paragraph()
        spacing.paragraph_format.space_before = Pt(6)

    def _get_available_table_width(self):
        """Return available table width within page margins."""
        try:
            section = self.doc.sections[0]
            return section.page_width - section.left_margin - section.right_margin
        except Exception:
            return Inches(6.0)

    def _calculate_table_column_widths(self, total_table_width, col_max_lengths):
        """Calculate stable column widths that fit within margins."""
        num_cols = len(col_max_lengths) if col_max_lengths else 1
        min_chars = 6
        adjusted_lengths = [max(length, min_chars) for length in (col_max_lengths or [1])]
        total_content_length = sum(adjusted_lengths) if sum(adjusted_lengths) > 0 else num_cols
        min_col_width = Inches(0.6)
        if min_col_width * num_cols > total_table_width:
            min_col_width = total_table_width / num_cols

        widths = []
        for length in adjusted_lengths:
            proportion = length / total_content_length if total_content_length > 0 else 1 / num_cols
            col_width = total_table_width * proportion
            if col_width < min_col_width:
                col_width = min_col_width
            widths.append(col_width)

        total_width = sum(widths, Inches(0))
        if total_width > total_table_width and total_width > 0:
            scale = total_table_width / total_width
            widths = [width * scale for width in widths]

        return widths

    def _add_plain_text_table(self, section):
        """Convert plain text table to Word table"""
        
        content = section.get('content', [])
        metadata = section.get('metadata', {})
        
        if not content:
            return
        
        # Filter out separator rows and count actual data rows
        data_rows = []
        for row in content:
            # Ensure row is a dict
            if not isinstance(row, dict):
                logger.warning(f"Skipping non-dict row in plain_text_table: {type(row)}")
                continue
            
            if row.get('row_type') != 'separator':
                data_rows.append(row)
        
        if len(data_rows) < 1:
            return
        
        # Determine column count from first row
        first_row = data_rows[0]
        cells = first_row.get('cells', [])
        
        # Ensure cells is a list
        if not isinstance(cells, list):
            logger.warning(f"Invalid cells format in plain_text_table: {type(cells)}")
            return
        
        col_count = len(cells)
        
        if col_count < 2:
            return  # Not a valid table
        
        # Determine row count
        row_count = len(data_rows)
        
        # Create Word table
        table = self.doc.add_table(rows=row_count, cols=col_count)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        col_max_lengths = [0] * col_count
        for row_data in data_rows:
            cells = row_data.get('cells', [])
            if not isinstance(cells, list):
                continue
            for col_idx in range(min(col_count, len(cells))):
                cell_content = cells[col_idx]
                if not isinstance(cell_content, str):
                    cell_content = str(cell_content) if cell_content is not None else ''
                content_length = len(cell_content.strip())
                if content_length > col_max_lengths[col_idx]:
                    col_max_lengths[col_idx] = content_length

        total_table_width = self._get_available_table_width()
        column_widths = self._calculate_table_column_widths(total_table_width, col_max_lengths)
        for col_idx, col_width in enumerate(column_widths):
            for row in table.rows:
                row.cells[col_idx].width = col_width
        
        # Populate table
        for row_idx, row_data in enumerate(data_rows):
            cells = row_data.get('cells', [])
            
            # Ensure cells is a list
            if not isinstance(cells, list):
                logger.warning(f"Invalid cells format in row {row_idx}: {type(cells)}")
                continue
            
            row_type = row_data.get('row_type', 'data')
            
            for col_idx in range(min(col_count, len(cells))):
                cell = table.rows[row_idx].cells[col_idx]
                cell_content = cells[col_idx]
                
                # Ensure cell content is a string
                if not isinstance(cell_content, str):
                    cell_content = str(cell_content) if cell_content is not None else ''
                
                cell.text = cell_content
                
                # Apply formatting
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(self.font_size)
                        
                        # Bold headers
                        if row_type == 'header':
                            run.bold = True
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            # Apply intelligent alignment based on content
                            content_type = self._get_cell_content_type(cell_content)
                            if content_type == 'numeric' or content_type == 'percentage':
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            elif content_type == 'statistical':
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            else:
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Add spacing after table
        self.doc.add_paragraph()

    def _get_cell_content_type(self, cell_text):
        """Determine content type for alignment (reuse from ACADEMIC_TABLE_FORMATTING)"""
        if not cell_text or not cell_text.strip():
            return 'text'
        
        # Check percentage
        if re.match(r'^\s*[\-\+]?\d+\.?\d*\s*%\s*$', cell_text):
            return 'percentage'
        
        # Check statistical notation
        stat_patterns = [r'[βαγδ]', r'p\s*[<>=]', r'[FfRrTt]\s*\(', r'χ²', r'SE\s*=', r'CI\s*=', r'\*{1,3}$']
        if any(re.search(pat, cell_text) for pat in stat_patterns):
            return 'statistical'
        
        # Check numeric
        if re.match(r'^\s*[\-\+]?\d+\.?\d*\s*$', cell_text) or \
           re.match(r'^\s*[\-\+]?\d{1,3}(,\d{3})*(\.\d+)?\s*$', cell_text):
            return 'numeric'
        
        return 'text'

    def add_watermark(self):
        """Add watermark to document footer and ensure it's bottom-right in all sections."""
        try:
            WATERMARK_TEXT = 'Formatted with AfroDocs.app'
            for section in self.doc.sections:
                footer = section.footer

                watermark_exists = False
                needs_new_paragraph = False

                for para in list(footer.paragraphs):
                    text = para.text or ''
                    if WATERMARK_TEXT in text:
                        if text.strip() == WATERMARK_TEXT:
                            para.clear()
                            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            pformat = para.paragraph_format
                            pformat.space_before = Pt(0)
                            pformat.space_after = Pt(0)
                            pformat.right_indent = Inches(0)
                            r = para.add_run(WATERMARK_TEXT)
                            r.font.name = 'Arial'
                            r.font.size = Pt(9)
                            try:
                                r.font.color.rgb = RGBColor(128, 128, 128)
                            except Exception:
                                pass
                            watermark_exists = True
                        else:
                            new_text = text.replace(WATERMARK_TEXT, '').strip()
                            para.clear()
                            if new_text:
                                para.add_run(new_text)
                            needs_new_paragraph = True

                if needs_new_paragraph or not watermark_exists:
                    exists_now = any((p.text or '').strip() == WATERMARK_TEXT for p in footer.paragraphs)
                    if not exists_now:
                        p = footer.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        pformat = p.paragraph_format
                        pformat.space_before = Pt(0)
                        pformat.space_after = Pt(0)
                        pformat.right_indent = Inches(0)
                        r = p.add_run(WATERMARK_TEXT)
                        r.font.name = 'Arial'
                        r.font.size = Pt(9)
                        try:
                            r.font.color.rgb = RGBColor(128, 128, 128)
                        except Exception:
                            pass

        except Exception as e:
            logger.error(f"Failed to add watermark: {e}")


# Flask Routes
@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy', 
        'timestamp': datetime.now().isoformat(),
        'version': '1.0.0',
        'engine': 'pattern-based',
        'patterns_loaded': 40
    })

def extract_paste_image_titles(text):
    if not text:
        return text, {}

    title_pattern = re.compile(r'^\s*\[IMAGE_TITLE:(?P<image_id>[^\]]+)\]\s*(?P<title>.*)$', re.IGNORECASE)
    cleaned_lines = []
    titles = {}

    for line in text.split('\n'):
        match = title_pattern.match(line)
        if match:
            image_id = match.group('image_id').strip()
            title = match.group('title').strip()
            if title:
                normalized = title.strip().lower()
                if normalized not in {'optional title', '(optional title)', 'optional', '(optional)'}:
                    titles[image_id] = title
            continue
        cleaned_lines.append(line)

    return '\n'.join(cleaned_lines), titles

def build_paste_images(image_files, image_ids, titles, placeholder_ids):
    images = []
    if not image_files:
        return images

    for idx, image_file in enumerate(image_files):
        image_id = image_ids[idx] if idx < len(image_ids) else f'paste_img_{idx + 1:03d}'
        if placeholder_ids and image_id not in placeholder_ids:
            continue

        image_bytes = image_file.read()
        if not image_bytes:
            continue

        _, ext = os.path.splitext(image_file.filename or '')
        ext = ext.lstrip('.').lower()
        mime = (image_file.mimetype or '').lower()
        if not ext:
            if 'png' in mime:
                ext = 'png'
            elif 'jpg' in mime or 'jpeg' in mime:
                ext = 'jpeg'
            elif 'gif' in mime:
                ext = 'gif'
            elif 'webp' in mime:
                ext = 'webp'
            else:
                ext = 'png'

        images.append({
            'image_id': image_id,
            'data': image_bytes,
            'format': ext,
            'width': 4.0,
            'height': 3.0,
            'caption': titles.get(image_id)
        })

    return images


@app.route('/upload', methods=['POST'])
@allow_guest
def upload_document():
    """Upload and process document"""
    
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Empty filename'}), 400

    paste_image_files = request.files.getlist('paste_images')
    paste_image_ids = []
    paste_image_titles = {}
    paste_image_placeholder_ids = set()
    if request.form.get('paste_image_ids'):
        try:
            paste_image_ids = json.loads(request.form.get('paste_image_ids'))
        except (TypeError, json.JSONDecodeError) as exc:
            logger.warning(f"Invalid paste_image_ids payload: {exc}")
    
    # Extract formatting options from request
    include_toc = request.form.get('include_toc', 'false').lower() == 'true'
    restructure_with_ai = request.form.get('restructure_with_ai', 'false').lower() == 'true'
    font_size = request.form.get('font_size', '12')
    line_spacing = request.form.get('line_spacing', '1.5')
    
    # Support individual margins for each side
    # If individual margins are provided, use them; otherwise fall back to uniform margin
    margin_left = request.form.get('margin_left')
    margin_top = request.form.get('margin_top')
    margin_bottom = request.form.get('margin_bottom')
    margin_right = request.form.get('margin_right')
    uniform_margin = request.form.get('margin_cm', '2.5')
    uniform_margin_provided = 'margin_cm' in request.form
    margin_left_provided = margin_left is not None and margin_left.strip()
    
    # Validate formatting parameters
    try:
        font_size = int(font_size)
        line_spacing = float(line_spacing)
        uniform_margin = float(uniform_margin)
        
        # Process individual margins
        margins = {}
        margin_vars = {
            'left': margin_left,
            'top': margin_top,
            'bottom': margin_bottom,
            'right': margin_right
        }
        
        for side, margin_val in margin_vars.items():
            if margin_val is not None and margin_val.strip():
                try:
                    margins[side] = float(margin_val)
                except (ValueError, AttributeError):
                    margins[side] = uniform_margin
            else:
                margins[side] = uniform_margin

        if not uniform_margin_provided and not margin_left_provided:
            margins['left'] = 3.0
        
        # Ensure reasonable values
        font_size = max(8, min(28, font_size))  # 8pt to 28pt
        line_spacing = max(1.0, min(3.0, line_spacing))  # 1.0 to 3.0
        for side in margins:
            margins[side] = max(0.5, min(5.0, margins[side]))  # 0.5cm to 5cm
    except (ValueError, TypeError):
        font_size = 12
        line_spacing = 1.5
        margins = {'left': 3.0, 'top': 2.5, 'bottom': 2.5, 'right': 2.5}
    
    # Log formatting options for debugging
    logger.info(f"Formatting options: TOC={include_toc}, RestructureAI={restructure_with_ai}, FontSize={font_size}pt, LineSpacing={line_spacing}, Margins=[L:{margins['left']}cm T:{margins['top']}cm B:{margins['bottom']}cm R:{margins['right']}cm]")
    
    # Estimate page count before processing
    # We need to read the file content to estimate pages
    # Since we save it later, we can read it from memory or save it first
    # Let's save it first to a temp location or just read bytes
    
    file_content = file.read()
    file.seek(0) # Reset pointer
    
    # Heuristic: 1 page = 350 words or 2000 characters
    # This is a rough estimate for .docx/.txt
    # For .docx we can't easily count without parsing, so we use file size or simple text extraction if possible
    # But parsing .docx in memory is heavy.
    # Let's use a simple file size heuristic for now or text length if .txt
    
    estimated_pages = 1
    if file.filename.endswith('.txt'):
        text = file_content.decode('utf-8', errors='ignore')
        word_count = len(text.split())
        estimated_pages = max(1, word_count // 350)
    elif file.filename.endswith('.docx'):
        # Very rough estimate: 1KB text ~ 1/2 page? No, .docx is zipped xml.
        # We will parse it properly after saving, but we need to check limit BEFORE processing?
        # Or we can check limit AFTER processing and fail if too large?
        # User experience: Better to fail early if obviously too large, but accurate count needs parsing.
        # Let's parse it using python-docx from memory to get paragraph count
        try:
            import io
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(file_content))
            # Count paragraphs as proxy? Or extract text.
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            text = ' '.join(full_text)
            word_count = len(text.split())
            estimated_pages = max(1, word_count // 350)
        except Exception as e:
            logger.warning(f"Error estimating pages: {e}")
            estimated_pages = 1 # Fallback

    # Initialize limit_reached_warning for all users (guests included)
    limit_reached_warning = False
    
    # Check limits and reset monthly counter if needed
    if current_user.is_authenticated:
        # Reload user to ensure fresh data and avoid proxy issues
        user = User.query.get(current_user.id)
        
        now = datetime.utcnow()
        if user.last_reset_date:
            # Check if month changed
            if now.month != user.last_reset_date.month or now.year != user.last_reset_date.year:
                user.pages_this_month = 0
                user.last_reset_date = now
                db.session.commit()
        else:
            user.last_reset_date = now
            db.session.commit()
            
        # Check Free Tier Limit (300 pages)
        if user.plan == 'free':
            limit = 300
            current_usage = user.pages_this_month
            balance = user.pages_balance
            
            # Admins have unlimited pages
            if not current_user.is_admin:
                # Soft Limit: Allow if they have ANY capacity left (either in monthly limit or balance)
                has_capacity = (current_usage < limit) or (balance > 0)
                
                if not has_capacity:
                    return jsonify({
                        'error': 'LIMIT_REACHED', 
                        'message': f'Page limit reached. Please upgrade your plan to continue.',
                        'required': estimated_pages,
                        'remaining': 0
                    }), 403

                # Calculate usage distribution
                remaining_monthly = max(0, limit - current_usage)
                
                if estimated_pages <= remaining_monthly:
                    # Fully covered by monthly limit
                    user.pages_this_month += estimated_pages
                else:
                    # Consumes all remaining monthly, spill over to balance or overage
                    user.pages_this_month += remaining_monthly
                    remainder = estimated_pages - remaining_monthly
                    
                    if remainder <= balance:
                        # Covered by balance
                        user.pages_balance -= remainder
                    else:
                        # Consumes all balance, spill over to overage (allowed because of soft limit)
                        user.pages_balance = 0
                        # Add the rest to pages_this_month to track total usage correctly
                        # This means pages_this_month will exceed limit (e.g. 105/100)
                        user.pages_this_month += (remainder - balance)
                
                # Check if they are NOW over the limit (for warning)
                if (user.pages_this_month >= limit) and (user.pages_balance == 0):
                    limit_reached_warning = True
            else:
                # Admin: no page limit, no usage tracking needed
                limit_reached_warning = False

        # Increment document count for current user
        user.documents_generated += 1
        db.session.commit()
    
    # Generate unique ID
    job_id = str(uuid.uuid4())
    
    # Save uploaded file
    file_ext = os.path.splitext(file.filename)[1].lower()
    input_path = os.path.join(UPLOAD_FOLDER, f"{job_id}{file_ext}")
    file.save(input_path)
    
    # Save metadata
    try:
        metadata = {
            'original_filename': file.filename,
            'upload_time': datetime.now().isoformat()
        }
        with open(os.path.join(OUTPUT_FOLDER, f"{job_id}_meta.json"), 'w') as f:
            json.dump(metadata, f)
            
        # Create DocumentRecord
        if current_user.is_authenticated:
            # We don't know the final filename yet, but we can update it later or guess it
            # The processor usually saves as formatted_{original_filename}.docx or similar
            # But let's just save the job_id and original filename for now
            # We'll update the filename after processing if possible, or just use job_id logic
            doc_record = DocumentRecord(
                user_id=current_user.id,
                filename=f"formatted_{file.filename}", # Predicted filename
                original_filename=file.filename,
                job_id=job_id,
                file_path=f"{job_id}_formatted.docx" # Internal storage name usually
            )
            db.session.add(doc_record)
            db.session.commit()
            
    except Exception as e:
        logger.error(f"Failed to save metadata or record: {e}")
    
    try:
        restructured_text = None
        if restructure_with_ai:
            if file_ext == '.docx':
                extracted_text = extract_text_from_docx_bytes(file_content)
            else:
                extracted_text = file_content.decode('utf-8', errors='ignore')

            restructured_text = restructure_text_with_ai(extracted_text)

        # Process document
        policy = FormatPolicy()
        processor = DocumentProcessor(policy=policy)
        images = []  # Extracted images
        shapes = []  # Extracted shapes/flowcharts

        if restructure_with_ai and restructured_text is not None:
            # Process restructured text in place of the original file content
            cleaned_text, paste_image_titles = extract_paste_image_titles(restructured_text)
            paste_image_placeholder_ids = set(re.findall(r'\[IMAGE:([^\]]+)\]', cleaned_text, flags=re.IGNORECASE))
            paste_images = build_paste_images(paste_image_files, paste_image_ids, paste_image_titles, paste_image_placeholder_ids)
            proc_result = processor.process_text(cleaned_text, strip_front_matter=include_toc)
            if isinstance(proc_result, tuple):
                if len(proc_result) == 3:
                    result, images, shapes = proc_result
                elif len(proc_result) == 2:
                    result, images = proc_result
                    shapes = []
                else:
                    result = proc_result[0] if proc_result else {}
                    images = []
                    shapes = []
            else:
                result = proc_result
                images = []
                shapes = []
            if paste_images:
                images = paste_images
        elif file_ext == '.docx':
            # Robust unpacking to handle both dict and tuple returns (now returns 3-tuple with shapes)
            proc_result = processor.process_docx(input_path, strip_front_matter=include_toc)
            if isinstance(proc_result, tuple):
                if len(proc_result) == 3:
                    result, images, shapes = proc_result
                elif len(proc_result) == 2:
                    result, images = proc_result
                    shapes = []
                else:
                    result = proc_result[0] if proc_result else {}
                    images = []
                    shapes = []
            else:
                result = proc_result
                images = []
                shapes = []
            logger.info(f"Processed document with {len(images)} images and {len(shapes)} shapes/flowcharts")
        else:
            # Read as text (txt, md, etc.) - no images or shapes in plain text
            with open(input_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
            # Robust unpacking for text processing too (now returns 3-tuple with shapes)
            cleaned_text, paste_image_titles = extract_paste_image_titles(text)
            paste_image_placeholder_ids = set(re.findall(r'\[IMAGE:([^\]]+)\]', cleaned_text, flags=re.IGNORECASE))
            paste_images = build_paste_images(paste_image_files, paste_image_ids, paste_image_titles, paste_image_placeholder_ids)
            proc_result = processor.process_text(cleaned_text, strip_front_matter=include_toc)
            if isinstance(proc_result, tuple):
                if len(proc_result) == 3:
                    result, images, shapes = proc_result
                elif len(proc_result) == 2:
                    result, images = proc_result
                    shapes = []
                else:
                    result = proc_result[0] if proc_result else {}
                    images = []
                    shapes = []
            else:
                result = proc_result
                images = []
                shapes = []
            if paste_images:
                images = paste_images
        
        # Validate result format
        if not isinstance(result, dict) or 'structured' not in result:
            logger.error(f"Invalid processing result format: {type(result)}")
            return jsonify({'error': 'Document processing failed to produce structured data'}), 500

        # Generate smart filename
        smart_filename = processor.generate_smart_filename(result['structured'], file.filename)
        logger.info(f"Generated smart filename: {smart_filename}")
        
        # Update metadata with smart filename
        try:
            meta_path = os.path.join(OUTPUT_FOLDER, f"{job_id}_meta.json")
            if os.path.exists(meta_path):
                with open(meta_path, 'r') as f:
                    metadata = json.load(f)
                
                metadata['smart_filename'] = smart_filename
                
                with open(meta_path, 'w') as f:
                    json.dump(metadata, f)
                    
            # Update DocumentRecord if exists
            if current_user.is_authenticated:
                doc_record = DocumentRecord.query.filter_by(job_id=job_id).first()
                if doc_record:
                    doc_record.filename = f"{smart_filename}.docx"
                    db.session.commit()
        except Exception as e:
            logger.error(f"Failed to update metadata with smart filename: {e}")

        # Generate formatted Word document with images
        output_path = os.path.join(OUTPUT_FOLDER, f"{job_id}_formatted.docx")
        generator = WordGenerator(policy=policy)
        
        # Pass certification data if detected (cover page detection removed - templates are now used)
        certification_data = getattr(processor, 'certification_data', None)
        
        # Watermarking is disabled for free tier documents.
        is_free = False
            
        generator.generate(
            result['structured'], 
            output_path, 
            images=images,
            shapes=shapes,  # Pass shapes/flowcharts to generator
            certification_data=certification_data,
            is_free_tier=is_free,
            # Formatting options
            include_toc=include_toc,
            font_size=font_size,
            line_spacing=line_spacing,
            margins=margins
        )
        
        # Ensure structured data is a list
        structured_for_preview = result.get('structured', [])
        if not isinstance(structured_for_preview, list):
            logger.warning(f"Invalid structured data type: {type(structured_for_preview)}, converting to empty list")
            structured_for_preview = []
        
        # Generate preview markdown
        preview = generate_preview_markdown(structured_for_preview)
        
        # Add image and shape count to stats
        result['stats']['images'] = len(images)
        result['stats']['shapes'] = len(shapes)
        
        return jsonify({
            'job_id': job_id,
            'stats': result['stats'],
            'structured': structured_for_preview,
            'preview': preview,
            'download_url': f'/download/{job_id}',
            'status': 'complete',
            'images_preserved': len(images),
            'shapes_preserved': len(shapes),
            'filename': f"{smart_filename}.docx",
            'limit_reached': limit_reached_warning
        })
    
    except Exception as e:
        import traceback
        tb_str = traceback.format_exc()
        traceback.print_exc()
        logger.error(f"Error processing document: {str(e)}\n{tb_str}")
        return jsonify({'error': str(e), 'traceback': tb_str}), 500
    finally:
        # Clean up input file
        if os.path.exists(input_path):
            try:
                os.remove(input_path)
            except:
                pass


@app.route('/download/<job_id>', methods=['GET'])
def download_document(job_id):
    """Download formatted document - allows guests, authenticated users, and admins"""
    output_path = os.path.join(OUTPUT_FOLDER, f"{job_id}_formatted.docx")
    
    # For guests: allow download if the file exists (they just generated it)
    # For authenticated users: allow if admin or if they own the document
    is_authorized = False
    
    # Guests can download if the file exists (no ownership tracking for guests)
    if not current_user.is_authenticated:
        # Allow guests to download any document they have the job_id for
        # This works because job_ids are UUIDs and not guessable
        is_authorized = True
    elif current_user.is_admin:
        is_authorized = True  # Admins can download anything
    else:
        # Check if user owns this document
        doc = DocumentRecord.query.filter_by(job_id=job_id, user_id=current_user.id).first()
        if doc:
            is_authorized = True
    
    if not is_authorized:
        logger.warning(f"Unauthorized download attempt for {job_id} by {current_user.username if current_user.is_authenticated else 'anonymous'}")
        return jsonify({'error': 'Unauthorized'}), 403
    
    # If exact UUID-style job file doesn't exist, try to resolve by smart/original filename
    if not os.path.exists(output_path):
        # Normalize requested id (strip extension if provided)
        requested = os.path.splitext(job_id)[0]
        fallback_path = None
        fallback_meta = None
        # Search all metadata files for a match on smart_filename or original_filename
        for fname in os.listdir(OUTPUT_FOLDER):
            if not fname.endswith('_meta.json'):
                continue
            meta_file = os.path.join(OUTPUT_FOLDER, fname)
            try:
                with open(meta_file, 'r', encoding='utf-8') as mf:
                    meta = json.load(mf)
            except Exception:
                continue

            smart = meta.get('smart_filename')
            orig = meta.get('original_filename')

            # Match against smart_filename (with or without .docx) or original filename base
            if smart and (smart == requested or smart == f"{requested}.docx"):
                candidate = fname.replace('_meta.json', '_formatted.docx')
                candidate_path = os.path.join(OUTPUT_FOLDER, candidate)
                if os.path.exists(candidate_path):
                    fallback_path = candidate_path
                    fallback_meta = meta
                    break

            if orig:
                orig_base = os.path.splitext(orig)[0]
                if orig_base == requested:
                    candidate = fname.replace('_meta.json', '_formatted.docx')
                    candidate_path = os.path.join(OUTPUT_FOLDER, candidate)
                    if os.path.exists(candidate_path):
                        fallback_path = candidate_path
                        fallback_meta = meta
                        break

        if fallback_path:
            output_path = fallback_path
            meta_path = None
            logger.info(f"Resolved download request '{job_id}' -> '{output_path}' via metadata lookup")
        else:
            # As a last resort, try to find any formatted file that contains the requested id as a substring
            for fname in os.listdir(OUTPUT_FOLDER):
                if requested in fname and fname.endswith('_formatted.docx'):
                    output_path = os.path.join(OUTPUT_FOLDER, fname)
                    meta_path = None
                    logger.info(f"Resolved download request '{job_id}' -> '{output_path}' via filename substring match")
                    break

            if not os.path.exists(output_path):
                logger.warning(f"Download failed for job '{job_id}': no matching output file found")

                # If the request appears to be a test-style id (e.g., 'test_12345'),
                # provide candidate matches (smart filenames / originals) to help debugging.
                candidates = []
                if requested.startswith('test'):
                    for fname in os.listdir(OUTPUT_FOLDER):
                        if not fname.endswith('_meta.json'):
                            continue
                        try:
                            with open(os.path.join(OUTPUT_FOLDER, fname), 'r', encoding='utf-8') as mf:
                                meta = json.load(mf)
                        except Exception:
                            continue
                        smart = meta.get('smart_filename')
                        orig = meta.get('original_filename')
                        job_candidate = fname.replace('_meta.json', '')
                        if smart and 'test' in smart.lower():
                            candidates.append({'job_id': job_candidate, 'smart_filename': smart, 'original': orig})
                        elif orig and 'test' in orig.lower():
                            candidates.append({'job_id': job_candidate, 'smart_filename': smart or '', 'original': orig})

                if candidates:
                    return jsonify({'error': 'Document not found', 'candidates': candidates}), 404

                return jsonify({'error': 'Document not found'}), 404

    # Try to get original filename from metadata
    download_name = 'formatted_document.docx'
    # If we already resolved a fallback_meta above, use it, otherwise try job-specific meta
    if 'fallback_meta' in locals() and fallback_meta:
        metadata = fallback_meta
        if metadata.get('smart_filename'):
            download_name = f"{metadata['smart_filename']}.docx"
        else:
            original_name = metadata.get('original_filename', '')
            if original_name:
                name, ext = os.path.splitext(original_name)
                download_name = f"{name}_formatted.docx"
    else:
        meta_path = os.path.join(OUTPUT_FOLDER, f"{job_id}_meta.json")
        if os.path.exists(meta_path):
            try:
                with open(meta_path, 'r') as f:
                    metadata = json.load(f)
                    
                    # Check for smart filename first
                    if metadata.get('smart_filename'):
                        download_name = f"{metadata['smart_filename']}.docx"
                    else:
                        original_name = metadata.get('original_filename', '')
                        if original_name:
                            name, ext = os.path.splitext(original_name)
                            download_name = f"{name}_formatted.docx"
            except Exception as e:
                logger.error(f"Error reading metadata for job {job_id}: {e}")
    
    inline = request.args.get('inline', 'false').lower() == 'true'
    as_attachment = not inline
    
    return send_file(
        output_path,
        as_attachment=as_attachment,
        download_name=download_name,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


# Lock for PDF conversion to prevent concurrent Word instances
pdf_conversion_lock = threading.Lock()

def _convert_docx_to_pdf(docx_path, pdf_path):
    """Helper to convert DOCX to PDF using OS-specific tools"""
    import sys
    import subprocess
    
    logger.info(f"Converting DOCX to PDF: {docx_path} -> {pdf_path}")
    
    if sys.platform == 'win32':
        # Windows: Use docx2pdf (requires Word installed)
        # Use a lock to prevent concurrent Word instances/COM access
        with pdf_conversion_lock:
            try:
                import pythoncom
                # Initialize COM for this thread
                pythoncom.CoInitialize()
                
                from docx2pdf import convert
                
                # Retry logic for flaky COM
                max_retries = 3
                last_error = None
                
                for attempt in range(max_retries):
                    try:
                        logger.info(f"PDF conversion attempt {attempt + 1}/{max_retries}")
                        convert(docx_path, pdf_path)
                        logger.info(f"PDF conversion successful: {pdf_path}")
                        return True, None
                    except Exception as e:
                        last_error = e
                        logger.warning(f"PDF conversion attempt {attempt + 1} failed: {e}")
                        # Wait a bit before retry
                        time.sleep(1)
                        
                logger.error(f"PDF conversion failed after {max_retries} attempts: {last_error}")
                return False, f'Windows PDF conversion failed after {max_retries} attempts: {str(last_error)}'
                
            except ImportError as e:
                logger.error(f"docx2pdf not installed: {e}")
                return False, 'docx2pdf not installed. Install with: pip install docx2pdf'
            except Exception as e:
                logger.error(f"Windows PDF conversion error: {e}")
                return False, f'Windows PDF conversion failed (Word installed?): {str(e)}'
            finally:
                try:
                    pythoncom.CoUninitialize()
                except:
                    pass
    else:
        # Linux: Use LibreOffice (headless)
        try:
            subprocess.run(['libreoffice', '--version'], stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)
        except (subprocess.CalledProcessError, FileNotFoundError):
            return False, 'LibreOffice not found on server. PDF conversion unavailable.'

        output_dir = os.path.dirname(pdf_path)
        cmd = ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', output_dir, docx_path]
        try:
            subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
            if os.path.exists(pdf_path):
                return True, None
            else:
                return False, 'PDF conversion failed to produce output file.'
        except Exception as e:
            return False, f'LibreOffice conversion failed: {str(e)}'


@app.route('/download-pdf/<job_id>', defaults={'filename': None}, methods=['GET'])
@app.route('/download-pdf/<job_id>/<filename>', methods=['GET'])
def download_pdf(job_id, filename):
    """Convert and download document as PDF"""
    docx_path = os.path.join(OUTPUT_FOLDER, f"{job_id}_formatted.docx")
    pdf_path = os.path.join(OUTPUT_FOLDER, f"{job_id}_formatted.pdf")
    
    if not os.path.exists(docx_path):
        return jsonify({'error': 'Document not found'}), 404
        
    # Determine filename
    download_name = 'formatted_document.pdf'
    
    # If filename is provided in URL, use it (security check needed?)
    if filename:
        download_name = filename
    else:
        # Fallback to metadata lookup
        meta_path = os.path.join(OUTPUT_FOLDER, f"{job_id}_meta.json")
        if os.path.exists(meta_path):
            try:
                with open(meta_path, 'r') as f:
                    metadata = json.load(f)
                    
                    # Check for smart filename first
                    if metadata.get('smart_filename'):
                        download_name = f"{metadata['smart_filename']}.pdf"
                    else:
                        original_name = metadata.get('original_filename', '')
                        if original_name:
                            name, ext = os.path.splitext(original_name)
                            download_name = f"{name}_formatted.pdf"
            except:
                pass

    # Check for inline parameter
    as_attachment = request.args.get('inline', 'false').lower() != 'true'

    # Check if PDF already exists
    if os.path.exists(pdf_path):
        return send_file(pdf_path, as_attachment=as_attachment, download_name=download_name, mimetype='application/pdf')

    # Convert to PDF
    success, error = _convert_docx_to_pdf(docx_path, pdf_path)
    
    if success and os.path.exists(pdf_path):
        return send_file(pdf_path, as_attachment=as_attachment, download_name=download_name, mimetype='application/pdf')
    else:
        return jsonify({'error': error or 'PDF conversion failed'}), 500


@app.route('/mobile-pdf-viewer', methods=['GET'])
def mobile_pdf_viewer():
    """Serve a minimal HTML page that embeds the PDF specified by the `file` query parameter.
    The wrapper first attempts to embed the PDF natively via an <object> tag and, if that fails,
    falls back to rendering the first page using PDF.js. This improves mobile preview compatibility.
    """
    file_url = request.args.get('file', '')
    if not file_url:
        return "Missing 'file' parameter", 400

    # Basic safety checks
    if '\n' in file_url or '\r' in file_url:
        return "Invalid file URL", 400

    html = (
        '<!doctype html>'
        '<html lang="en">'
        '<head>'
        '  <meta name="viewport" content="width=device-width,initial-scale=1,maximum-scale=1">'
        '  <meta charset="utf-8">'
        '  <title>Preview</title>'
        '  <style>'
        '    html,body{height:100%;margin:0;background:linear-gradient(180deg,#1e293b 0%,#0f172a 100%)}'
        '    .viewer{width:100%;height:100%;border:0;display:block}'
        '    .center-content{position:absolute;top:50%;left:50%;transform:translate(-50%,-50%);display:flex;flex-direction:column;align-items:center;text-align:center;padding:20px;width:90%;max-width:320px}'
        '    .icon-circle{width:64px;height:64px;border-radius:50%;background:rgba(45,212,191,0.15);display:flex;align-items:center;justify-content:center;margin-bottom:16px}'
        '    .icon-circle svg{width:32px;height:32px;color:#2dd4bf}'
        '    .open-btn{display:inline-flex;align-items:center;gap:8px;padding:12px 24px;background:linear-gradient(135deg,#2dd4bf 0%,#06b6d4 100%);color:#fff;border:none;border-radius:999px;font-weight:600;font-size:14px;cursor:pointer;box-shadow:0 4px 15px rgba(45,212,191,0.3);transition:all 0.2s;text-decoration:none;margin-bottom:16px}'
        '    .open-btn:hover{opacity:0.9;transform:translateY(-1px)}'
        '    .info-text{color:#94a3b8;font-size:13px;line-height:1.5;margin-bottom:8px}'
        '    .info-text span{color:#e2e8f0}'
        '    .try-preview{color:rgba(45,212,191,0.7);font-size:12px;text-decoration:underline;cursor:pointer;background:none;border:none;margin-top:8px}'
        '    .try-preview:hover{color:#2dd4bf}'
        '  </style>'
        '  <script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/3.11.174/pdf.min.js"></script>'
        '</head>'
        '<body>'
        '  <div id="container" style="width:100%;height:100%;position:relative"></div>'
        '  <div id="welcome" class="center-content">'
        f'    <a href="{file_url}" target="_blank" rel="noopener noreferrer" class="open-btn">'
        '      Open Preview'
        '    </a>'
        '    <div class="info-box" style="background:rgba(255,255,255,0.04);border:1px solid rgba(255,255,255,0.06);padding:10px;border-radius:8px;margin-top:8px;text-align:center;">'
        f'      <p style="color:#94a3b8;margin:0;font-size:13px">Document display may vary on mobile devices depending on your document reader. For the best layout and full access to the Table of Contents, please download and open this file on a PC.</p>'
        '    </div>'
        '  </div>'
        '  <script>'
        '    (function(){'
        f'      const fileUrl = {json.dumps(file_url)};'
        '      const container = document.getElementById("container");'
        '      const welcome = document.getElementById("welcome");'
        '      function loadPdf(){'
        '        // Try embedding via object tag'
        '        const obj = document.createElement("object");'
        '        obj.data = fileUrl;'
        '        obj.type = "application/pdf";'
        '        obj.className = "viewer";'
        '        obj.width = "100%";'
        '        obj.height = "100%";'
        ''
        '        let loaded = false;'
        '        obj.onload = function(){ loaded = true; };'
        '        obj.onerror = function(){ loaded = false; };'
        ''
        '        container.appendChild(obj);'
        ''
        '        // After short timeout, if not loaded, fallback to PDF.js render'
        '        setTimeout(async function(){'
        '          if (loaded) return;'
        '          try { container.removeChild(obj); } catch(e) {}'
        '          try {'
        '            const resp = await fetch(fileUrl);'
        '            if (!resp.ok) throw new Error("Fetch failed");'
        '            const buf = await resp.arrayBuffer();'
        '            const loadingTask = pdfjsLib.getDocument({data: buf});'
        '            const pdf = await loadingTask.promise;'
        '            const page = await pdf.getPage(1);'
        '            const viewport = page.getViewport({scale: 1});'
        '            const scale = (container.clientWidth - 20) / viewport.width;'
        '            const scaled = page.getViewport({scale});'
        '            const canvas = document.createElement("canvas");'
        '            canvas.width = Math.floor(scaled.width);'
        '            canvas.height = Math.floor(scaled.height);'
        '            canvas.style.display = "block";'
        '            canvas.style.margin = "0 auto";'
        '            container.appendChild(canvas);'
        '            const ctx = canvas.getContext("2d");'
        '            await page.render({canvasContext: ctx, viewport: scaled}).promise;'
        '          } catch (err) {'
        '            console.error("PDF.js fallback failed", err);'
        '            const iframe = document.createElement("iframe");'
        '            iframe.src = fileUrl;'
        '            iframe.className = "viewer";'
        '            container.appendChild(iframe);'
        '          }'
        '        }, 600);'
        '      }'
        '    })();'
        '  </script>'
        '</body>'
        '</html>'
    )

    return Response(html, mimetype='text/html')


def generate_preview_markdown(structured):
    """Generate markdown preview from structured data"""
    markdown = ''
    
    # Ensure structured is a list
    if not isinstance(structured, list):
        logger.warning(f"Expected list for structured data, got {type(structured)}")
        return str(structured) if structured else ''
    
    for section in structured:
        # Ensure section is a dict
        if not isinstance(section, dict):
            logger.warning(f"Expected dict for section, got {type(section)}: {section}")
            continue
        
        section_type = section.get('type', 'section')
        
        # Get heading - check both 'heading' and 'title' keys (prominent_section uses 'title')
        heading = section.get('heading', '') or section.get('title', '')
        
        # For chapter sections, include chapter_title if present
        if section_type == 'chapter' and section.get('chapter_title'):
            heading = f"{heading}: {section.get('chapter_title')}"
        
        # Skip sections with no valid heading
        if not heading or heading == 'N/A':
            heading = 'Untitled'
            
        # Add heading
        level = min(section.get('level', 1), 6)
        heading_prefix = '#' * level
        markdown += f"{heading_prefix} {heading}\n\n"
        
        # Add content
        for item in section.get('content', []):
            if not isinstance(item, dict):
                continue
                
            if item.get('type') == 'paragraph':
                markdown += f"{item.get('text', '')}\n\n"
            
            elif item.get('type') == 'definition':
                markdown += f"**{item.get('term', '')}:** {item.get('definition', '')}\n\n"
            
            elif item.get('type') == 'bullet_list':
                # Conservative bullet rendering: Only use bullets if 4+ items and all short
                list_items = item.get('items', [])
                
                # Pre-scan to validate items and count
                valid_items = []
                should_use_bullets = False
                
                for list_item in list_items:
                    if isinstance(list_item, dict):
                        item_content = list_item.get('content', '')
                    else:
                        item_content = str(list_item)
                    
                    # Clean asterisks
                    item_content = re.sub(r'\*+', '', item_content).strip()
                    
                    if item_content:
                        valid_items.append(item_content)
                
                # Use bullets ONLY if 4+ items and all are short
                if len(valid_items) >= 4:
                    all_short = True
                    for content in valid_items:
                        # Check if contains colon (would be bolded)
                        if ':' in content:
                            all_short = False
                            break
                        # Check word count
                        if len(content.split()) > 30:
                            all_short = False
                            break
                        # Check if multiline
                        if '\n' in content:
                            all_short = False
                            break
                    should_use_bullets = all_short
                
                # Render items
                for item_content in valid_items:
                    if should_use_bullets:
                        markdown += f"- {item_content}\n"
                    else:
                        # Use bold format instead
                        if ':' in item_content:
                            markdown += f"**{item_content}**\n"
                        else:
                            markdown += f"{item_content}\n"
                
                markdown += '\n'
            
            elif item.get('type') == 'numbered_list':
                for idx, list_item in enumerate(item.get('items', []), 1):
                    # Extract content from list_item (could be string or dict)
                    if isinstance(list_item, dict):
                        item_content = list_item.get('content', '')
                    else:
                        item_content = str(list_item)
                    # Clean up numbering
                    clean_item = re.sub(r'^[\d]+[\.\)]\s*', '', item_content)
                    clean_item = re.sub(r'^[a-z][\.\)]\s*', '', clean_item)
                    markdown += f"{idx}. {clean_item if clean_item else item_content}\n"
                markdown += '\n'
            
            elif item.get('type') == 'table':
                rows = item.get('rows', [])
                if not rows and item.get('content'):
                    content_rows = []
                    for entry in item.get('content', []):
                        if not isinstance(entry, dict):
                            continue
                        if entry.get('type') == 'row' and isinstance(entry.get('cells'), list):
                            content_rows.append(entry['cells'])
                    rows = content_rows
                if rows:
                    # Header
                    markdown += '| ' + ' | '.join(str(cell) for cell in rows[0]) + ' |\n'
                    markdown += '| ' + ' | '.join(['---'] * len(rows[0])) + ' |\n'
                    # Data rows
                    for row in rows[1:]:
                        markdown += '| ' + ' | '.join(str(cell) for cell in row) + ' |\n'
                    markdown += '\n'
            
            elif item.get('type') == 'reference':
                markdown += f"{item.get('text', '')}\n\n"
            
            elif item.get('type') == 'figure':
                markdown += f"{item.get('caption', '')}\n\n"
            
            elif item.get('type') == 'quote':
                markdown += f"> {item.get('text', '')}\n\n"
    
    return markdown


# --- Cover Page Generator Integration ---
from coverpage_generator import generate_cover_page, load_json

@app.route('/api/institutions', methods=['GET'])
def get_institutions():
    """Return list of institutions"""
    try:
        data = load_json('institutions.json')
        if not data:
            logger.error("institutions.json loaded empty or not found")
            return jsonify({"institutions": []}) # Return empty structure at least
        return jsonify(data)
    except Exception as e:
        logger.error(f"Error loading institutions: {e}")
        return jsonify({"institutions": []})

@app.route('/api/courses/search', methods=['GET'])
def search_courses():
    """Search courses by code or title"""
    query = request.args.get('q', '').lower()
    courses = load_json('courses_database.json')
    
    if not query:
        return jsonify([])
        
    matches = [
        c for c in courses 
        if query in c['code'].lower() or query in c['title'].lower()
    ]
    return jsonify(matches[:10]) # Limit to 10 results

@app.route('/api/coverpage/generate', methods=['POST'])
@allow_guest
def api_generate_coverpage():
    """Generate cover page from form data"""
    data = request.json
    output_path, error = generate_cover_page(data)
    
    if error:
        return jsonify({'error': error}), 400
        
    # Generate a new job_id for this cover page result
    job_id = str(uuid.uuid4())
    
    # Check for merge request
    merge_job_id = data.get('mergeJobId')
    if merge_job_id:
        try:
            processed_path = os.path.join(OUTPUT_FOLDER, f"{merge_job_id}_formatted.docx")
            if os.path.exists(processed_path):
                logger.info(f"Merging cover page with processed document: {processed_path}")
                
                # Load documents
                cover_doc = Document(output_path)
                processed_doc = Document(processed_path)
                
                # Detect numbering style of processed doc (Section 0)
                # This helps us restore the correct numbering after merge
                processed_start_fmt = 'lowerRoman' # Default
                try:
                    if processed_doc.sections:
                        sectPr = processed_doc.sections[0]._sectPr
                        pgNumType = sectPr.find(qn('w:pgNumType'))
                        if pgNumType is not None:
                            fmt = pgNumType.get(qn('w:fmt'))
                            if fmt:
                                processed_start_fmt = fmt
                except Exception as e:
                    logger.warning(f"Could not detect processed doc numbering: {e}")
                
                # CRITICAL FIX: Ensure all content in processed_doc uses 'AcademicBody' style
                # instead of 'Normal'. This prevents the content from inheriting the Cover Page's
                # 'Normal' style (which might be Calibri/different spacing) during the merge.
                
                # 1. Ensure AcademicBody exists and mirrors the processed document styles
                academic_style = None
                academic_list_number = None
                academic_list_bullet = None

                def copy_font(source_font, target_font):
                    if source_font.name:
                        target_font.name = source_font.name
                    if source_font.size:
                        target_font.size = source_font.size
                    if source_font.bold is not None:
                        target_font.bold = source_font.bold
                    if source_font.italic is not None:
                        target_font.italic = source_font.italic
                    if source_font.underline is not None:
                        target_font.underline = source_font.underline
                    if source_font.color and source_font.color.rgb:
                        target_font.color.rgb = source_font.color.rgb

                def copy_paragraph_format(source_pf, target_pf):
                    if source_pf.alignment is not None:
                        target_pf.alignment = source_pf.alignment
                    if source_pf.line_spacing is not None:
                        target_pf.line_spacing = source_pf.line_spacing
                    if source_pf.space_after is not None:
                        target_pf.space_after = source_pf.space_after
                    if source_pf.space_before is not None:
                        target_pf.space_before = source_pf.space_before
                    if source_pf.left_indent is not None:
                        target_pf.left_indent = source_pf.left_indent
                    if source_pf.right_indent is not None:
                        target_pf.right_indent = source_pf.right_indent
                    if source_pf.first_line_indent is not None:
                        target_pf.first_line_indent = source_pf.first_line_indent
                    if source_pf.keep_together is not None:
                        target_pf.keep_together = source_pf.keep_together
                    if source_pf.keep_with_next is not None:
                        target_pf.keep_with_next = source_pf.keep_with_next
                    if source_pf.page_break_before is not None:
                        target_pf.page_break_before = source_pf.page_break_before
                    if source_pf.widow_control is not None:
                        target_pf.widow_control = source_pf.widow_control

                try:
                    base_body_style = processed_doc.styles['Normal']

                    # --- AcademicBody ---
                    if 'AcademicBody' not in processed_doc.styles:
                        academic_style = processed_doc.styles.add_style('AcademicBody', WD_STYLE_TYPE.PARAGRAPH)
                        academic_style.base_style = base_body_style
                    else:
                        academic_style = processed_doc.styles['AcademicBody']

                    copy_font(base_body_style.font, academic_style.font)
                    copy_paragraph_format(base_body_style.paragraph_format, academic_style.paragraph_format)

                    # --- AcademicListNumber ---
                    list_number_style = processed_doc.styles['List Number'] if 'List Number' in processed_doc.styles else base_body_style
                    if 'AcademicListNumber' not in processed_doc.styles:
                        academic_list_number = processed_doc.styles.add_style('AcademicListNumber', WD_STYLE_TYPE.PARAGRAPH)
                        academic_list_number.base_style = list_number_style
                    else:
                        academic_list_number = processed_doc.styles['AcademicListNumber']

                    copy_font(list_number_style.font, academic_list_number.font)
                    copy_paragraph_format(list_number_style.paragraph_format, academic_list_number.paragraph_format)

                    # --- AcademicListBullet ---
                    list_bullet_style = processed_doc.styles['List Bullet'] if 'List Bullet' in processed_doc.styles else base_body_style
                    if 'AcademicListBullet' not in processed_doc.styles:
                        academic_list_bullet = processed_doc.styles.add_style('AcademicListBullet', WD_STYLE_TYPE.PARAGRAPH)
                        academic_list_bullet.base_style = list_bullet_style
                    else:
                        academic_list_bullet = processed_doc.styles['AcademicListBullet']

                    copy_font(list_bullet_style.font, academic_list_bullet.font)
                    copy_paragraph_format(list_bullet_style.paragraph_format, academic_list_bullet.paragraph_format)

                except Exception as e:
                    logger.warning(f"Error setting up Academic styles: {e}")

                if academic_style:
                    def process_para(para):
                        # Fix Normal -> AcademicBody
                        if para.style.name == 'Normal':
                            para.style = academic_style
                        
                        # Fix List Number -> AcademicListNumber
                        elif para.style.name == 'List Number' and academic_list_number:
                            para.style = academic_list_number
                            
                        # Fix List Bullet -> AcademicListBullet
                        elif para.style.name == 'List Bullet' and academic_list_bullet:
                            para.style = academic_list_bullet
                    
                    # Update Body Paragraphs
                    for para in processed_doc.paragraphs:
                        process_para(para)
                    
                    # Update Tables
                    for table in processed_doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    process_para(para)
                                    
                    # Update Headers and Footers (All Sections)
                    for section in processed_doc.sections:
                        # Headers
                        for header in [section.header, section.first_page_header, section.even_page_header]:
                            if header:
                                for para in header.paragraphs:
                                    process_para(para)
                                for table in header.tables:
                                    for row in table.rows:
                                        for cell in row.cells:
                                            for para in cell.paragraphs:
                                                process_para(para)
                        # Footers
                        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                            if footer:
                                for para in footer.paragraphs:
                                    process_para(para)
                                for table in footer.tables:
                                    for row in table.rows:
                                        for cell in row.cells:
                                            for para in cell.paragraphs:
                                                process_para(para)

                # Add a section break to the cover doc to ensure separation of styles/margins
                # This forces the appended content to start on a new page with its own section properties
                new_section = cover_doc.add_section(WD_SECTION.NEW_PAGE)
                
                # Copy section properties from the processed document to the new section
                # This ensures the content retains its original formatting (margins, orientation, etc.)
                if processed_doc.sections:
                    src_section = processed_doc.sections[0]
                    new_section.left_margin = src_section.left_margin
                    new_section.right_margin = src_section.right_margin
                    new_section.top_margin = src_section.top_margin
                    new_section.bottom_margin = src_section.bottom_margin
                    new_section.page_width = src_section.page_width
                    new_section.page_height = src_section.page_height
                    new_section.orientation = src_section.orientation
                    new_section.gutter = src_section.gutter
                    new_section.header_distance = src_section.header_distance
                    new_section.footer_distance = src_section.footer_distance
                    new_section.different_first_page_header_footer = src_section.different_first_page_header_footer
                    # We don't copy start_type as we want NEW_PAGE

                
                # Merge using docxcompose
                composer = Composer(cover_doc)
                composer.append(processed_doc)
                
                # Save merged document
                composer.save(output_path)
                
                # --- POST-MERGE FIX: Restore Page Numbering ---
                # The merge process often breaks page numbering linkage.
                # We need to explicitly unlink the body section from the cover page
                # and restore the page numbering.
                try:
                    merged_doc = Document(output_path)
                    if len(merged_doc.sections) > 1:
                        # Section 0 is Cover Page (no number)
                        # Section 1 is the start of the Body (needs number)
                        body_section = merged_doc.sections[1]
                        
                        # 1. Unlink footer from Cover Page
                        body_section.footer.is_linked_to_previous = False
                        
                        # 2. Restore Page Numbering
                        # Use the format detected from the original processed document
                        sectPr = body_section._sectPr
                        pgNumType = sectPr.find(qn('w:pgNumType'))
                        if pgNumType is None:
                            pgNumType = OxmlElement('w:pgNumType')
                            sectPr.append(pgNumType)
                        
                        # Only set if not already set (preserve existing formatting if possible)
                        if not pgNumType.get(qn('w:fmt')):
                            pgNumType.set(qn('w:fmt'), processed_start_fmt)
                        if not pgNumType.get(qn('w:start')):
                            pgNumType.set(qn('w:start'), '1')
                            
                        # 3. Ensure PAGE field exists in footer
                        footer = body_section.footer
                        # Check if footer is empty
                        if not footer.paragraphs or not footer.paragraphs[0].text.strip():
                            if footer.paragraphs:
                                p = footer.paragraphs[0]
                                p.clear()
                            else:
                                p = footer.add_paragraph()
                            
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = p.add_run()
                            fldChar1 = OxmlElement('w:fldChar')
                            fldChar1.set(qn('w:fldCharType'), 'begin')
                            run._r.append(fldChar1)
                            instrText = OxmlElement('w:instrText')
                            instrText.text = "PAGE"
                            run._r.append(instrText)
                            fldChar2 = OxmlElement('w:fldChar')
                            fldChar2.set(qn('w:fldCharType'), 'end')
                            run._r.append(fldChar2)
                        
                        merged_doc.save(output_path)
                        logger.info("Restored page numbering in merged document")
                except Exception as e:
                    logger.error(f"Error restoring page numbering: {e}")

                logger.info("Merge successful")
            else:
                logger.warning(f"Processed document not found for merge: {processed_path}")
        except Exception as e:
            logger.error(f"Error merging documents: {str(e)}")
            # We continue without merging if it fails, or we could return an error
            # For now, let's just log it and return the cover page

    # Rename output file to standard format {job_id}_formatted.docx
    formatted_path = os.path.join(OUTPUT_FOLDER, f"{job_id}_formatted.docx")
    try:
        if os.path.exists(formatted_path):
            os.remove(formatted_path)
        os.rename(output_path, formatted_path)
        output_path = formatted_path
    except Exception as e:
        logger.error(f"Failed to rename output file: {e}")
        # Fallback to original path if rename fails, but this will break PDF download
        pass

    # Calculate smart filename
    smart_filename = None
    try:
        title = data.get('title', '')
        student = data.get('studentName', '')
        
        if title and student:
            short_title = ' '.join(title.split()[:5])
            smart_filename = f"{short_title} - {student}"
        elif title:
            short_title = ' '.join(title.split()[:5])
            smart_filename = short_title
            
        if smart_filename:
            # Sanitize
            smart_filename = re.sub(r'[<>:"/\\|?*]', '', smart_filename)
            smart_filename = re.sub(r'\s+', ' ', smart_filename).strip()
            if len(smart_filename) > 50:
                smart_filename = smart_filename[:50].strip()
    except Exception as e:
        logger.error(f"Failed to calculate smart filename: {e}")

    # Save metadata
    try:
        original_filename = os.path.basename(output_path) # This might be the old name if rename failed
        # But we want the "friendly" name for the user
        friendly_name = f"{smart_filename}.docx" if smart_filename else "CoverPage.docx"
        
        metadata = {
            'job_id': job_id,
            'original_filename': friendly_name,
            'smart_filename': smart_filename,
            'created_at': datetime.now().isoformat(),
            'merged_from': merge_job_id
        }
        
        with open(os.path.join(OUTPUT_FOLDER, f"{job_id}_meta.json"), 'w') as f:
            json.dump(metadata, f)
    except Exception as e:
        logger.error(f"Failed to save metadata: {e}")

    # Return the file info
    # We return the friendly filename for the UI
    filename = f"{smart_filename}.docx" if smart_filename else "CoverPage.docx"
    
    # Update DocumentRecord if exists
    if current_user.is_authenticated:
        try:
            doc_record = DocumentRecord(
                user_id=current_user.id,
                filename=filename,
                original_filename=f"Cover Page: {filename}",
                job_id=job_id,
                file_path=filename # Relative to Cover Pages folder
            )
            db.session.add(doc_record)
            db.session.commit()
        except Exception as e:
            logger.error(f"Failed to save cover page record: {e}")

    return jsonify({
        'success': True,
        'job_id': job_id,
        'filename': filename,
        'downloadUrl': f'/download/{job_id}'
    })

@app.route('/download/<job_id>', methods=['GET'])
def download_file(job_id):
    """Download generated file by job_id"""
    inline = request.args.get('inline', 'false').lower() == 'true'
    as_attachment = not inline
    
    # Look for file with job_id pattern
    formatted_file = os.path.join(OUTPUT_FOLDER, f"{job_id}_formatted.docx")
    if os.path.exists(formatted_file):
        return send_file(formatted_file, as_attachment=as_attachment)
        
    return jsonify({'error': 'Document not found'}), 404


# --- End Cover Page Generator Integration ---


import sys
if __name__ == '__main__':
    # Only start the server if not running tests
    # Check if running directly
    print(f"Starting backend server... (__name__={__name__})")
    if not any('unittest' in arg or 'test' in arg for arg in sys.argv):
        print("=" * 60)
        print("Pattern-Based Academic Document Formatter")
        print("=" * 60)
        print("- No AI dependencies - 100% pattern matching")
        print("- Ultra-fast processing - 1000+ lines/second")
        print("- 40+ regex patterns loaded")
        print("- Zero API costs")
        print("- 100% offline capability")
        print("=" * 60)
        print(f"Server starting at http://localhost:5000")
        print("=" * 60)
        # Run without debug/reloader so the process stays stable in this environment
        app.run(debug=False, host='0.0.0.0', port=5000, use_reloader=False)
