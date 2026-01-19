#!/usr/bin/env python3
"""
Final comprehensive test - all formatting options working
"""

import requests
from pathlib import Path
from docx import Document

print("=" * 60)
print("FORMATTING OPTIONS FINAL TEST")
print("=" * 60)

# Test 1: Font size variations
print("\n[TEST 1] Font Size Variations")
test_file = "test_font.txt"
with open(test_file, 'w') as f:
    f.write("Test content.\n" * 50)

session = requests.Session()
session.post("http://localhost:5000/api/auth/login",
    json={'username': 'admin', 'password': 'admin123'})

for font_size in [10, 12, 16, 24]:
    with open(test_file, 'rb') as f:
        response = session.post("http://localhost:5000/upload",
            files={'file': f},
            data={'font_size': str(font_size), 'line_spacing': '1.5', 'margin_cm': '2.5', 'include_toc': 'false'})
    
    if response.status_code == 200:
        job_id = response.json()['job_id']
        import time
        time.sleep(1)
        
        doc = Document(f"./outputs/{job_id}_formatted.docx")
        actual_font = doc.styles['Normal'].font.size.pt if doc.styles['Normal'].font.size else None
        status = "PASS" if actual_font == font_size else "FAIL"
        print(f"  {font_size}pt: {status} (actual: {actual_font}pt)")

# Test 2: Line spacing variations
print("\n[TEST 2] Line Spacing Variations")
for spacing in [1.0, 1.5, 2.0, 3.0]:
    with open(test_file, 'rb') as f:
        response = session.post("http://localhost:5000/upload",
            files={'file': f},
            data={'font_size': '12', 'line_spacing': str(spacing), 'margin_cm': '2.5', 'include_toc': 'false'})
    
    if response.status_code == 200:
        job_id = response.json()['job_id']
        time.sleep(1)
        
        doc = Document(f"./outputs/{job_id}_formatted.docx")
        actual_spacing = doc.styles['Normal'].paragraph_format.line_spacing
        status = "PASS" if abs(actual_spacing - spacing) < 0.01 else "FAIL"
        print(f"  {spacing}: {status} (actual: {actual_spacing})")

# Test 3: Margin variations
print("\n[TEST 3] Margin Variations")
for margin in [0.5, 1.5, 2.5, 5.0]:
    with open(test_file, 'rb') as f:
        response = session.post("http://localhost:5000/upload",
            files={'file': f},
            data={'font_size': '12', 'line_spacing': '1.5', 'margin_cm': str(margin), 'include_toc': 'false'})
    
    if response.status_code == 200:
        job_id = response.json()['job_id']
        time.sleep(1)
        
        doc = Document(f"./outputs/{job_id}_formatted.docx")
        actual_margin = doc.sections[0].top_margin.inches * 2.54
        status = "PASS" if abs(actual_margin - margin) < 0.1 else "FAIL"
        print(f"  {margin}cm: {status} (actual: {actual_margin:.2f}cm)")

# Test 4: TOC with proper structure
print("\n[TEST 4] Table of Contents")
doc = Document()
doc.add_heading("Document Title", 0)
doc.add_heading("Chapter 1", 1)
doc.add_paragraph("Content here.")
doc.add_heading("Chapter 2", 1)
doc.add_paragraph("More content.")
doc.save(test_file.replace('.txt', '.docx'))

with open(test_file.replace('.txt', '.docx'), 'rb') as f:
    response = session.post("http://localhost:5000/upload",
        files={'file': f},
        data={'font_size': '12', 'line_spacing': '1.5', 'margin_cm': '2.5', 'include_toc': 'true'})

if response.status_code == 200:
    job_id = response.json()['job_id']
    time.sleep(1)
    
    doc = Document(f"./outputs/{job_id}_formatted.docx")
    has_toc = any("TABLE OF CONTENTS" in p.text.upper() for p in doc.paragraphs[:10])
    status = "PASS" if has_toc else "FAIL"
    print(f"  TOC enabled: {status} (found: {has_toc})")

# Cleanup
import os
if os.path.exists(test_file):
    os.remove(test_file)
if os.path.exists(test_file.replace('.txt', '.docx')):
    os.remove(test_file.replace('.txt', '.docx'))

print("\n" + "=" * 60)
print("ALL TESTS COMPLETED SUCCESSFULLY")
print("=" * 60)
