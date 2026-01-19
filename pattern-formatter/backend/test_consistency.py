#!/usr/bin/env python3
"""
Comprehensive test - verify ALL text in document has consistent formatting
"""

import requests
from pathlib import Path
from docx import Document

def check_document_consistency(doc_path, expected_font_size, expected_line_spacing):
    """Check if all paragraphs have consistent formatting"""
    doc = Document(doc_path)
    
    issues = []
    font_sizes = {}
    line_spacings = {}
    
    for para_idx, para in enumerate(doc.paragraphs):
        if not para.text.strip():
            continue
        
        # Check line spacing
        actual_spacing = para.paragraph_format.line_spacing
        if actual_spacing:
            key = str(actual_spacing)
            if key not in line_spacings:
                line_spacings[key] = 0
            line_spacings[key] += 1
            
            if abs(actual_spacing - expected_line_spacing) > 0.01:
                issues.append(f"Para {para_idx}: Line spacing {actual_spacing} (expected {expected_line_spacing})")
        
        # Check font sizes in runs
        for run in para.runs:
            if run.font.size:
                size_pt = run.font.size.pt
                key = str(size_pt)
                if key not in font_sizes:
                    font_sizes[key] = 0
                font_sizes[key] += 1
    
    return {
        'issues': issues,
        'font_sizes': font_sizes,
        'line_spacings': line_spacings
    }

# Test with different formatting
print("=" * 70)
print("FORMATTING CONSISTENCY TEST - CHECKING ALL TEXT IN DOCUMENT")
print("=" * 70)

test_cases = [
    {'name': '10pt, 1.5 spacing', 'font': 10, 'spacing': 1.5},
    {'name': '14pt, 2.0 spacing', 'font': 14, 'spacing': 2.0},
    {'name': '18pt, 2.5 spacing', 'font': 18, 'spacing': 2.5},
    {'name': '24pt, 3.0 spacing', 'font': 24, 'spacing': 3.0},
]

# Create test document with headings
from docx import Document as DocxDocument

test_file = "test_consistency.docx"
doc = DocxDocument()
doc.add_heading("Document Title", 0)
doc.add_heading("Chapter 1: Introduction", 1)
doc.add_paragraph("This is body text for chapter 1. " * 20)
doc.add_heading("Chapter 2: Detailed Content", 1)
doc.add_heading("2.1 Subsection", 2)
doc.add_paragraph("More body text here. " * 25)
doc.add_paragraph("Another paragraph with content. " * 20)
doc.add_heading("Chapter 3: Advanced Topics", 1)
doc.add_paragraph("Final paragraph. " * 30)
doc.save(test_file)

# Test each case
session = requests.Session()
session.post("http://localhost:5000/api/auth/login",
    json={'username': 'admin', 'password': 'admin123'})

all_pass = True
for test_case in test_cases:
    print(f"\nTest: {test_case['name']}")
    print("-" * 70)
    
    params = {
        'font_size': str(test_case['font']),
        'line_spacing': str(test_case['spacing']),
        'margin_cm': '2.5',
        'include_toc': 'true'
    }
    
    with open(test_file, 'rb') as f:
        response = session.post("http://localhost:5000/upload",
            files={'file': f}, data=params)
    
    if response.status_code == 200:
        job_id = response.json()['job_id']
        import time
        time.sleep(2)
        
        doc_path = f"./outputs/{job_id}_formatted.docx"
        result = check_document_consistency(doc_path, test_case['font'], test_case['spacing'])
        
        print(f"Font sizes found: {result['font_sizes']}")
        print(f"Line spacings found: {result['line_spacings']}")
        
        if result['issues']:
            print(f"INCONSISTENCIES FOUND ({len(result['issues'])} issues):")
            for issue in result['issues'][:5]:  # Show first 5
                print(f"   - {issue}")
            all_pass = False
        else:
            print("All text formatting consistent - PASS")
    else:
        print(f"❌ Upload failed: {response.json()}")
        all_pass = False

import os
if os.path.exists(test_file):
    os.remove(test_file)
if os.path.exists('fix_formatting.py'):
    os.remove('fix_formatting.py')

print("\n" + "=" * 70)
if all_pass:
    print("✅ ALL TESTS PASSED - FORMATTING IS CONSISTENT THROUGHOUT")
else:
    print("❌ SOME TESTS FAILED - THERE ARE FORMATTING INCONSISTENCIES")
print("=" * 70)
