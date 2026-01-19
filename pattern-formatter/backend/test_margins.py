#!/usr/bin/env python3
"""
Test individual margin configuration
"""

import requests
from pathlib import Path
from docx import Document
import time
import os
import shutil

# Create a simple test document
test_file = "test_margins.docx"
doc = Document()
doc.add_heading("Margin Test Document", 0)
doc.add_paragraph("This document tests individual margin configuration." * 10)
doc.save(test_file)

# Ensure outputs directory exists
os.makedirs("outputs", exist_ok=True)
shutil.rmtree("outputs", ignore_errors=True)
os.makedirs("outputs", exist_ok=True)

session = requests.Session()
try:
    session.post("http://localhost:5000/api/auth/login",
        json={'username': 'admin', 'password': 'admin123'})
except:
    print("Auth not available")

# Test 1: Individual margins (left=1cm, top=2cm, bottom=3cm, right=4cm)
print("="*70)
print("TEST: Individual Margin Configuration")
print("="*70)

params = {
    'font_size': '12',
    'line_spacing': '1.5',
    'margin_left': '1.0',
    'margin_top': '2.0',
    'margin_bottom': '3.0',
    'margin_right': '4.0',
    'include_toc': 'false'
}

print(f"\nTest params: {params}")

with open(test_file, 'rb') as f:
    response = session.post("http://localhost:5000/upload",
        files={'file': f}, data=params)

if response.status_code == 200:
    result = response.json()
    job_id = result['job_id']
    print(f"Upload successful, job_id: {job_id}")
    
    time.sleep(2)
    
    output_path = f"./outputs/{job_id}_formatted.docx"
    if os.path.exists(output_path):
        doc = Document(output_path)
        
        # Get the first section and check margins
        section = doc.sections[0]
        
        # Convert from inches to cm for readability
        left_cm = round(section.left_margin.inches * 2.54, 2)
        top_cm = round(section.top_margin.inches * 2.54, 2)
        bottom_cm = round(section.bottom_margin.inches * 2.54, 2)
        right_cm = round(section.right_margin.inches * 2.54, 2)
        
        print(f"\nDocument Margins (in cm):")
        print(f"  Left:   {left_cm} cm (expected: 1.0 cm)")
        print(f"  Top:    {top_cm} cm (expected: 2.0 cm)")
        print(f"  Bottom: {bottom_cm} cm (expected: 3.0 cm)")
        print(f"  Right:  {right_cm} cm (expected: 4.0 cm)")
        
        # Verify margins match
        success = (
            abs(left_cm - 1.0) < 0.05 and
            abs(top_cm - 2.0) < 0.05 and
            abs(bottom_cm - 3.0) < 0.05 and
            abs(right_cm - 4.0) < 0.05
        )
        
        if success:
            print("\n✓ PASS: All margins configured correctly!")
        else:
            print("\n✗ FAIL: Margins do not match expected values")
    else:
        print(f"Output file not found at {output_path}")
else:
    print(f"Upload failed: {response.status_code}")
    print(response.json())

# Test 2: Uniform margin (fallback)
print("\n" + "="*70)
print("TEST: Uniform Margin (Fallback)")
print("="*70)

params2 = {
    'font_size': '12',
    'line_spacing': '1.5',
    'margin_cm': '2.5',
    'include_toc': 'false'
}

print(f"\nTest params: {params2}")

with open(test_file, 'rb') as f:
    response = session.post("http://localhost:5000/upload",
        files={'file': f}, data=params2)

if response.status_code == 200:
    result = response.json()
    job_id = result['job_id']
    print(f"Upload successful, job_id: {job_id}")
    
    time.sleep(2)
    
    output_path = f"./outputs/{job_id}_formatted.docx"
    if os.path.exists(output_path):
        doc = Document(output_path)
        section = doc.sections[0]
        
        left_cm = round(section.left_margin.inches * 2.54, 2)
        top_cm = round(section.top_margin.inches * 2.54, 2)
        bottom_cm = round(section.bottom_margin.inches * 2.54, 2)
        right_cm = round(section.right_margin.inches * 2.54, 2)
        
        print(f"\nDocument Margins (in cm):")
        print(f"  Left:   {left_cm} cm (expected: 2.5 cm)")
        print(f"  Top:    {top_cm} cm (expected: 2.5 cm)")
        print(f"  Bottom: {bottom_cm} cm (expected: 2.5 cm)")
        print(f"  Right:  {right_cm} cm (expected: 2.5 cm)")
        
        success = all(abs(m - 2.5) < 0.05 for m in [left_cm, top_cm, bottom_cm, right_cm])
        
        if success:
            print("\n✓ PASS: Uniform margins applied correctly!")
        else:
            print("\n✗ FAIL: Margins do not match expected values")
else:
    print(f"Upload failed: {response.status_code}")

# Cleanup
if os.path.exists(test_file):
    os.remove(test_file)

print("\n" + "="*70)
print("MARGIN TESTS COMPLETE")
print("="*70)
