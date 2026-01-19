#!/usr/bin/env python3
"""Create a test document with plain text tables for testing the pattern formatter."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new document
doc = Document()

# Add title
title = doc.add_heading('Test Document: Plain Text Tables', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add introduction
doc.add_paragraph(
    "This document contains various types of plain text tables to test the table detection patterns."
)

# Section 1: Spaced columns table
doc.add_heading('Section 1: Spaced Columns Table', level=1)
doc.add_paragraph(
    "The following table uses spaced columns separated by multiple spaces:"
)
doc.add_paragraph(
    "Student Name          Subject           Grade    Status"
)
doc.add_paragraph(
    "Alice Johnson         Mathematics       A        Passed"
)
doc.add_paragraph(
    "Bob Smith             Physics           B+       Passed"
)
doc.add_paragraph(
    "Carol White           Chemistry         A-       Passed"
)

# Section 2: Pipe-delimited table
doc.add_heading('Section 2: Pipe-Delimited Table', level=1)
doc.add_paragraph(
    "The following table uses pipe delimiters:"
)
doc.add_paragraph(
    "| Month    | Sales | Growth  |"
)
doc.add_paragraph(
    "| January  | 5000  | 10.5%   |"
)
doc.add_paragraph(
    "| February | 5500  | 12.3%   |"
)
doc.add_paragraph(
    "| March    | 6200  | 18.2%   |"
)

# Section 3: Tab-separated table
doc.add_heading('Section 3: Tab-Separated Table', level=1)
doc.add_paragraph(
    "The following table uses tab separators:"
)
doc.add_paragraph(
    "Product\tPrice\tQuantity\tTotal"
)
doc.add_paragraph(
    "Laptop\t$1000\t5\t$5000"
)
doc.add_paragraph(
    "Monitor\t$300\t10\t$3000"
)
doc.add_paragraph(
    "Keyboard\t$50\t20\t$1000"
)

# Section 4: Comma-separated table
doc.add_heading('Section 4: Comma-Separated Table', level=1)
doc.add_paragraph(
    "The following table uses comma separators:"
)
doc.add_paragraph(
    "Region,Q1,Q2,Q3,Q4,Total"
)
doc.add_paragraph(
    "North,10000,12000,15000,18000,55000"
)
doc.add_paragraph(
    "South,8000,9000,10000,11000,38000"
)
doc.add_paragraph(
    "East,12000,14000,16000,18000,60000"
)
doc.add_paragraph(
    "West,9000,10000,11000,12000,42000"
)

# Section 5: Box-drawn table
doc.add_heading('Section 5: Box-Drawn Table', level=1)
doc.add_paragraph(
    "The following table uses box-drawing characters:"
)
doc.add_paragraph(
    "+--------+--------+--------+"
)
doc.add_paragraph(
    "| Header | Value1 | Value2 |"
)
doc.add_paragraph(
    "+--------+--------+--------+"
)
doc.add_paragraph(
    "| Row 1  | 100    | 200    |"
)
doc.add_paragraph(
    "| Row 2  | 300    | 400    |"
)
doc.add_paragraph(
    "+--------+--------+--------+"
)

# Section 6: Key-value pairs
doc.add_heading('Section 6: Key-Value Pairs Table', level=1)
doc.add_paragraph(
    "The following table uses key-value format:"
)
doc.add_paragraph(
    "Name: John Doe"
)
doc.add_paragraph(
    "Age: 30"
)
doc.add_paragraph(
    "Position: Senior Developer"
)
doc.add_paragraph(
    "Department: Engineering"
)

# Section 7: Numbered rows
doc.add_heading('Section 7: Numbered Rows Table', level=1)
doc.add_paragraph(
    "The following table uses numbered rows:"
)
doc.add_paragraph(
    "1. Apple         $0.99    Fruit"
)
doc.add_paragraph(
    "2. Banana        $0.59    Fruit"
)
doc.add_paragraph(
    "3. Carrot        $0.79    Vegetable"
)
doc.add_paragraph(
    "4. Broccoli      $1.29    Vegetable"
)

# Add footer
footer_para = doc.add_paragraph("\n\nEnd of document - Test file created for plain text table pattern testing.")
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save the document
output_path = "c:\\Users\\user\\Desktop\\PATTERN\\pattern-formatter\\backend\\test_document_with_tables.docx"
doc.save(output_path)
print(f"Test document created: {output_path}")
