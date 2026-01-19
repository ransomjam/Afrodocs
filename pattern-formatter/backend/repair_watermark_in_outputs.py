"""Repair watermark placement in existing .docx files in outputs/.

Behavior:
- For each .docx in outputs (ignore temp ~$ files), open file.
- For each footer paragraph containing 'Formatted with AfroDocs.app':
  - If paragraph contains only the watermark text: set alignment RIGHT and format run.
  - If paragraph contains other text (e.g., page number + watermark), remove watermark text from that paragraph and create a new paragraph at end of footer with watermark (right-aligned).
- Save repaired file as the same path, create a backup copy with .bak.docx extension before saving.

Run with: python repair_watermark_in_outputs.py
"""
import os
import shutil
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT_DIR = os.path.join(os.path.dirname(__file__), 'outputs')
WATERMARK = 'Formatted with AfroDocs.app'

def repair_file(path):
    print('Processing', path)
    doc = Document(path)
    modified = False
    for si, section in enumerate(doc.sections):
        footer = section.footer
        # Collect paragraphs to add watermark to at end if split needed
        to_add = []
        for para in list(footer.paragraphs):
            text = para.text or ''
            if WATERMARK in text:
                if text.strip() == WATERMARK:
                    # Simple case: single paragraph with watermark - set right alignment and styling
                    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for run in para.runs:
                        run.text = WATERMARK
                        run.font.name = 'Arial'
                        run.font.size = Pt(9)
                        try:
                            run.font.color.rgb = RGBColor(128,128,128)
                        except Exception:
                            pass
                    modified = True
                else:
                    # Complex case: watermark is in same paragraph as page number or other text
                    new_text = text.replace(WATERMARK, '').strip()
                    # Replace paragraph runs with cleaned text
                    para.clear()
                    if new_text:
                        para.add_run(new_text)
                    # Schedule watermark to be added as its own paragraph at footer end
                    to_add.append((section,))
                    modified = True
        # Add watermark paragraphs for scheduled items
        for _ in to_add:
            p = footer.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = p.add_run(WATERMARK)
            r.font.name = 'Arial'
            r.font.size = Pt(9)
            try:
                r.font.color.rgb = RGBColor(128,128,128)
            except Exception:
                pass
            modified = True
    if modified:
        # Backup original
        bak = path + '.bak.docx'
        if not os.path.exists(bak):
            shutil.copy2(path, bak)
            print('  backup saved to', bak)
        doc.save(path)
        print('  saved repaired file')
    else:
        print('  no watermark issues found')

if __name__ == '__main__':
    files = [f for f in os.listdir(OUT_DIR) if f.endswith('.docx') and not f.startswith('~$')]
    for f in files:
        repair_file(os.path.join(OUT_DIR, f))
    print('Done')
