#!/usr/bin/env python3
"""
Test document upload with formatting options to verify the fix
"""

import requests
from pathlib import Path
from docx import Document
import time
import os
import shutil

# Create a simple test document
test_file = "test_error_fix.docx"
doc = Document()
doc.add_heading("Test Document", 0)
doc.add_paragraph("This is a test document to verify the error fix. " * 10)
doc.add_heading("Section 1", 1)
doc.add_paragraph("More test content here. " * 15)
doc.save(test_file)

# Ensure outputs directory exists
os.makedirs("outputs", exist_ok=True)
shutil.rmtree("outputs", ignore_errors=True)
os.makedirs("outputs", exist_ok=True)

session = requests.Session()
try:
    session.post("http://localhost:5000/api/auth/login",
        json={'username': 'admin', 'password': 'admin123'})
except:
    print("Auth not available - continuing without login")

# Test 1: Upload with uniform margin
print("="*70)
print("TEST 1: Upload with Uniform Margin")
print("="*70)

params = {
    'font_size': '12',
    'line_spacing': '1.5',
    'margin_cm': '2.5',
    'include_toc': 'false'
}

print(f"Uploading with params: {params}")

try:
    with open(test_file, 'rb') as f:
        response = session.post("http://localhost:5000/upload",
            files={'file': f}, data=params, timeout=30)
    
    if response.status_code == 200:
        result = response.json()
        print(f"✓ SUCCESS: Upload completed")
        print(f"  Job ID: {result.get('job_id')}")
        print(f"  Stats: {result.get('stats')}")
    else:
        print(f"✗ FAILED: Status {response.status_code}")
        print(f"  Response: {response.json()}")
except Exception as e:
    print(f"✗ ERROR: {e}")

# Test 2: Upload with individual margins
print("\n" + "="*70)
print("TEST 2: Upload with Individual Margins")
print("="*70)

params2 = {
    'font_size': '14',
    'line_spacing': '1.8',
    'margin_left': '1.5',
    'margin_top': '2.0',
    'margin_bottom': '2.0',
    'margin_right': '2.5',
    'include_toc': 'false'
}

print(f"Uploading with params: {params2}")

try:
    with open(test_file, 'rb') as f:
        response = session.post("http://localhost:5000/upload",
            files={'file': f}, data=params2, timeout=30)
    
    if response.status_code == 200:
        result = response.json()
        print(f"✓ SUCCESS: Upload completed")
        print(f"  Job ID: {result.get('job_id')}")
        print(f"  Stats: {result.get('stats')}")
    else:
        print(f"✗ FAILED: Status {response.status_code}")
        print(f"  Response: {response.json()}")
except Exception as e:
    print(f"✗ ERROR: {e}")

# Cleanup
if os.path.exists(test_file):
    os.remove(test_file)

print("\n" + "="*70)
print("TEST COMPLETE")
print("="*70)
