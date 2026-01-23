import os
import json
import re
from copy import deepcopy
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Paths
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, 'data')

# Templates directory - check multiple possible locations for Docker and local compatibility
# In Docker: /app/Cover Pages (same level as backend code)
# Locally: ../Cover Pages (parent of backend folder)
def get_templates_dir():
    """Find the Cover Pages templates directory."""
    possible_paths = [
        os.path.join(BASE_DIR, 'Cover Pages'),  # Docker: /app/Cover Pages
        os.path.join(os.path.dirname(BASE_DIR), 'Cover Pages'),  # Local: parent/Cover Pages
        '/Cover Pages',  # Docker root fallback
    ]
    for path in possible_paths:
        if os.path.exists(path):
            logger.info(f"Found templates directory at: {path}")
            return path
    # Default fallback
    logger.warning(f"Templates directory not found. Checked: {possible_paths}")
    return os.path.join(os.path.dirname(BASE_DIR), 'Cover Pages')

TEMPLATES_DIR = get_templates_dir()
OUTPUT_DIR = os.path.join(BASE_DIR, 'outputs', 'Cover Pages')

# Ensure output directory exists
os.makedirs(OUTPUT_DIR, exist_ok=True)

def load_json(filename):
    filepath = os.path.join(DATA_DIR, filename)
    if os.path.exists(filepath):
        with open(filepath, 'r', encoding='utf-8') as f:
            return json.load(f)
    return []

def get_template_path(document_type, university='uba'):
    """
    Map document type and university to template file.
    Supports University of Bamenda (uba), University of Buea (ub), and NPUI (npui).
    
    IMPORTANT: Institution folder names MUST match exactly with those on disk, including all spaces.
    See INSTITUTION_TEMPLATES_GUIDE.md for adding new institutions.
    DO NOT remove spaces around underscores in folder names.
    
    CRITICAL: When adding new institutions:
    1. Create folder: pattern-formatter/Cover Pages/[Exact Folder Name]/
    2. Add all template files to the folder
    3. Add mapping here with EXACT folder name (copy-paste from disk)
    4. Update institutions.json with institution data
    5. See INSTITUTION_TEMPLATES_GUIDE.md for complete checklist
    """
    mapping = {
        'Assignment': 'Assignments Cover Page Template.docx',
        'Thesis': 'Dissertation Cover Page Template.docx',
        'Dissertation': 'Dissertation Cover Page Template.docx',
        'Research Proposal': 'Research Proposal Cover Page Template.docx',
        'Internship Report': 'Internship Cover Page Template.docx',
        'Project Report': 'Internship Cover Page Template.docx',
        'Research Paper': 'Assignments Cover Page Template.docx',
        'Lab Report': 'Assignments Cover Page Template.docx',
        'Term Paper': 'Assignments Cover Page Template.docx',
    }
    
    filename = mapping.get(document_type, 'Assignments Cover Page Template.docx')
    
    # Map institution ID to folder name
    # NOTE: Folder names must match EXACTLY with folders in: pattern-formatter/Cover Pages/
    # Include all spaces, underscores, and special characters exactly as they appear on disk
    institution_mapping = {
        'uba': 'Cover Pages _ University of Bamenda',
        'Bamenda': 'Cover Pages _ University of Bamenda',
        'ub': 'Cover Page _ University of Buea',
        'Buea': 'Cover Page _ University of Buea',
        'npui': 'Cover Pages _ National University Institute (NPUI)',
        'NPUI': 'Cover Pages _ National University Institute (NPUI)',
        'bust': 'Cover Page _ BUST',
        'BUST': 'Cover Page _ BUST',
        'cucb': 'Cover Page _ Catholic University',
        'Catholic University': 'Cover Page _ Catholic University',
    }
    
    # Determine university folder based on ID
    # Bamenda is default fallback for backward compatibility
    university_folder = institution_mapping.get(university, 'Cover Pages _ University of Bamenda')
    
    return os.path.join(TEMPLATES_DIR, university_folder, filename)

def get_all_placeholders(doc):
    """
    Scan document for all {{...}} placeholders, including those with newlines.
    Returns a list of unique placeholder strings found in the doc.
    """
    text_content = []
    
    # Paragraphs
    for p in doc.paragraphs:
        text_content.append(p.text)
        
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text_content.append(p.text)
                    
    # Text Boxes
    if doc.element.body is not None:
        for txbx in doc.element.body.iter(qn('w:txbxContent')):
            for p in txbx.iter(qn('w:p')):
                para_text = ""
                for r in p.iter(qn('w:r')):
                    for t in r.iter(qn('w:t')):
                        if t.text:
                            para_text += t.text
                if para_text:
                    text_content.append(para_text)
                            
    full_text = "\n".join(text_content)
    # Find all {{...}} patterns, allowing for newlines and whitespace inside
    # The regex matches {{ followed by any char (including newline) until }}
    matches = re.findall(r'\{\{[^}]+\}\}', full_text, re.DOTALL)
    return list(set(matches))

def sanitize_input(value):
    """
    Sanitize input values to handle corrupted/garbage characters.
    Removes control characters, excessive whitespace, and validates UTF-8.
    Also detects and filters obvious garbage patterns.
    """
    if not isinstance(value, str):
        value = str(value) if value else ""
    
    # Remove control characters (except newline, tab)
    value = ''.join(ch for ch in value if ch.isprintable() or ch in '\n\t')
    
    # Strip excessive whitespace
    value = value.strip()
    
    # Detect common garbage patterns:
    # Very short strings with no vowels OR all repeating patterns
    if value and len(value) <= 10:
        # Count vowels
        vowels = sum(1 for ch in value.lower() if ch in 'aeiou')
        
        # If no vowels and short, likely garbage (e.g., 'hjh', 'bf')
        if vowels == 0 and len(value) < 5:
            return ""
        
        # Check for repeating pattern (e.g., 'uiui', 'aba')
        if len(value) >= 4:
            # Check if it's a repeating 2-char pattern
            if len(value) >= 4 and all(value[i] == value[i % 2] for i in range(len(value))):
                return ""
            
            # Check if value is mostly alternating between 2 chars
            if len(value) >= 4:
                unique_chars = len(set(value))
                if unique_chars <= 2:
                    return ""
    
    return value

def replace_text_in_paragraph(paragraph, replacements):
    """
    Replace text in a paragraph, handling placeholders split across multiple runs.
    Completely removes placeholder and replaces with value.
    """
    # First, reconstruct the full paragraph text to find split placeholders
    full_text = paragraph.text
    
    # Check which replacement keys are in the full text
    matched_keys = [key for key in replacements.keys() if key in full_text]
    
    if not matched_keys:
        return  # Nothing to replace
    
    # Process replacements in order
    for key in matched_keys:
        value = replacements[key]
        full_text = full_text.replace(key, str(value))
    
    # Now we need to update the paragraph with the new text
    # Strategy: Remove all runs and create a new single run with the replacement text
    
    # Store formatting properties from first run for reuse
    first_run_properties = None
    if paragraph.runs:
        first_run_properties = paragraph.runs[0]._element.find(qn('w:rPr'))
    
    # Remove all runs from paragraph
    paragraph_element = paragraph._element
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    
    # Create new run with replaced text
    new_r = OxmlElement('w:r')
    
    # Copy formatting from original first run if available
    if first_run_properties is not None:
        new_r.append(deepcopy(first_run_properties))
    
    # Add text element with replaced content
    t = OxmlElement('w:t')
    t.text = full_text
    if full_text.startswith(' ') or full_text.endswith(' '):
        t.set(qn('xml:space'), 'preserve')
    new_r.append(t)
    
    # Append new run to paragraph
    paragraph_element.append(new_r)

def replace_in_textboxes(doc, replacements):
    """
    Replace text in text boxes (shapes) by iterating over XML.
    Handles placeholders split across multiple runs.
    Replaces ALL occurrences of ALL keys in the textbox.
    """
    if doc.element.body is None:
        return

    for txbx in doc.element.body.iter(qn('w:txbxContent')):
        for p in txbx.iter(qn('w:p')):
            # 1. Reconstruct full paragraph text
            full_text = ""
            xml_runs = []
            for r in p.iter(qn('w:r')):
                xml_runs.append(r)
                for t in r.iter(qn('w:t')):
                    if t.text:
                        full_text += t.text
            
            if not xml_runs:
                continue
            
            # 2. Apply ALL replacements to the full text
            original_text = full_text
            new_text = full_text
            replacements_made = []
            
            for key, value in replacements.items():
                if key in new_text:
                    replacement_value = str(value)
                    new_text = new_text.replace(key, replacement_value)
                    replacements_made.append(f"'{key}' -> '{replacement_value}'")
            
            # 3. If any replacements were made, update the paragraph
            if replacements_made:
                print(f"DEBUG: Replaced in textbox: {', '.join(replacements_made)}")
                
                # Remove ALL runs in the paragraph
                p_element = xml_runs[0].getparent()
                for run_element in xml_runs:
                    p_element.remove(run_element)
                
                # Create a single new run with the replaced text
                new_r = OxmlElement('w:r')
                
                # Copy formatting properties from first original run if available
                if xml_runs[0].find(qn('w:rPr')) is not None:
                    new_r.append(deepcopy(xml_runs[0].find(qn('w:rPr'))))
                
                # Handle newlines in replacement
                if '\n' in new_text:
                    lines = new_text.split('\n')
                    for i, line in enumerate(lines):
                        if i > 0:
                            br = OxmlElement('w:br')
                            new_r.append(br)
                        t = OxmlElement('w:t')
                        t.text = line
                        if line.startswith(' ') or line.endswith(' '):
                            t.set(qn('xml:space'), 'preserve')
                        new_r.append(t)
                else:
                    t = OxmlElement('w:t')
                    t.text = new_text
                    if new_text.startswith(' ') or new_text.endswith(' '):
                        t.set(qn('xml:space'), 'preserve')
                    new_r.append(t)
                
                p_element.append(new_r)

def apply_times_new_roman_to_fields(doc, document_type):
    """
    Apply Times New Roman font to specific fields in dissertations.
    For dissertations, ensures department and school/faculty names are in Times New Roman.
    """
    if document_type not in ['Dissertation', 'Thesis']:
        return
    
    # Fields that should always be Times New Roman in dissertations
    target_fields = ['{{DEPARMENT}}', '{{Deparment}}', '{{SCHOOL/FACULTY}}', '{{School/Faculty}}']
    
    # Process all paragraphs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.name != 'Times New Roman':
                # Check if this run contains dissertation fields
                if any(field in run.text or field.replace('{{', '').replace('}}', '') in run.text for field in target_fields):
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
    
    # Process textboxes
    if doc.element.body is not None:
        for txbx in doc.element.body.iter(qn('w:txbxContent')):
            for p in txbx.iter(qn('w:p')):
                for r in p.iter(qn('w:r')):
                    for t in r.iter(qn('w:t')):
                        if t.text and any(field in t.text for field in target_fields):
                            # Apply Times New Roman to this run
                            rPr = r.find(qn('w:rPr'))
                            if rPr is None:
                                rPr = OxmlElement('w:rPr')
                                r.insert(0, rPr)
                            
                            # Set font to Times New Roman
                            rFonts = rPr.find(qn('w:rFonts'))
                            if rFonts is None:
                                rFonts = OxmlElement('w:rFonts')
                                rPr.append(rFonts)
                            rFonts.set(qn('w:ascii'), 'Times New Roman')
                            rFonts.set(qn('w:hAnsi'), 'Times New Roman')

def preserve_submission_statement(doc, document_type):
    """
    Mark submission statement text to preserve from modifications.
    For dissertations, the submission statement should NOT be modified.
    """
    if document_type not in ['Dissertation', 'Thesis']:
        return
    
    # Find and mark paragraphs containing submission statement
    submission_keywords = ['submitted', 'presented', 'dissertation']
    
    for paragraph in doc.paragraphs:
        para_lower = paragraph.text.lower()
        # Check if this is a submission statement paragraph
        if any(keyword in para_lower for keyword in submission_keywords):
            # Mark with a custom property to prevent modification
            # This is informational - the replace functions respect placeholders
            if not any(placeholder in paragraph.text for placeholder in ['{{', '}}'] if paragraph.text):
                # Add metadata that this should not be modified
                pass  # Statement is already safe - only {{}} placeholders get replaced

def apply_dissertation_formatting(doc, document_type):
    """
    Apply all formatting fixes for dissertation documents.
    """
    if document_type not in ['Dissertation', 'Thesis']:
        return
    
    apply_times_new_roman_to_fields(doc, document_type)
    preserve_submission_statement(doc, document_type)

def replace_placeholders(doc, replacements):
    """
    Replace placeholders in the document (paragraphs, tables, and text boxes).
    """
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        replace_text_in_paragraph(paragraph, replacements)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_text_in_paragraph(paragraph, replacements)
                    
    # Replace in text boxes
    replace_in_textboxes(doc, replacements)

def generate_cover_page(data):
    """
    Generate a cover page based on the provided data.
    """
    try:
        document_type = data.get('documentType', 'Assignment')
        # Get institution ID from data, with fallback to old 'university' parameter for backward compatibility
        institution = data.get('institution', data.get('university', 'uba'))
        template_path = get_template_path(document_type, institution)
        
        logger.info(f"DEBUG: Generating cover page for institution='{institution}', document_type='{document_type}'")
        logger.info(f"DEBUG: Template path: {template_path}")
        
        if not os.path.exists(template_path):
            logger.error(f"Template not found: {template_path}")
            return None, f"Template for {document_type} not found."
        
        logger.info(f"DEBUG: Template found, loading document...")
        doc = Document(template_path)
        
        # Sanitize all incoming data to handle corrupted characters
        sanitized_data = {}
        for key, value in data.items():
            if isinstance(value, str):
                sanitized_data[key] = sanitize_input(value)
            else:
                sanitized_data[key] = value
        data = sanitized_data
        
        # Force document defaults to Times New Roman to ensure consistency after merge
        try:
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
        except Exception as e:
            logger.warning(f"Could not set default font: {e}")

        # Standardize date
        date_obj = datetime.now()
        if data.get('date'):
            try:
                date_obj = datetime.strptime(data.get('date'), '%Y-%m-%d')
            except:
                pass
        
        date_str = date_obj.strftime('%B %d, %Y')
        month_year = date_obj.strftime('%B %Y')
        year_str = date_obj.strftime('%Y')
        academic_year = f"{date_obj.year}/{date_obj.year + 1}" if date_obj.month >= 9 else f"{date_obj.year - 1}/{date_obj.year}"
        
        # Helper to get value or empty string
        def get_val(k): return str(data.get(k, '') or '')

        # Infer degree based on faculty
        faculty = get_val('faculty')
        degree = "Bachelor of Science" # Default
        if "Arts" in faculty:
            degree = "Bachelor of Arts"
        elif "Technology" in faculty:
            degree = "Bachelor of Technology"
        elif "Engineering" in faculty or "Polytechnic" in faculty:
            degree = "Bachelor of Engineering"
        elif "Education" in faculty:
            degree = "Bachelor of Education"
        elif "Commerce" in faculty or "Management" in faculty:
            degree = "Bachelor of Science"
            
        # Override for Thesis/Dissertation if needed (usually Masters)
        if document_type in ['Thesis', 'Dissertation']:
            degree = "Master of Science" # Simplified assumption
            if "Arts" in faculty: degree = "Master of Arts"

        # Base values map - use exact case for placeholders
        values_map = {
            'STUDENT_NAME': (get_val('studentName') or '').upper(),  # Matches {{ STUDENT_NAME }} - Always uppercase
            'Matricule Number': get_val('studentId'),  # Matches {{ Matricule Number }}
            'COURSE_CODE': get_val('courseCode'),
            'COURSE_TITLE': get_val('courseTitle'),
            'DEPARTMENT': (get_val('departmentCustom') or get_val('department') or '').upper(),  # Matches {{ DEPARTMENT }} - Always uppercase
            'Deparment': (get_val('departmentCustom') or get_val('department') or '').upper(),  # Matches {{ Deparment }} (typo in template) - Always uppercase
            'SCHOO/FACULTY': (get_val('facultyCustom') or get_val('faculty') or '').upper(),  # Matches {{ SCHOO/FACULTY }} - Always uppercase
            'School/Faculty': (get_val('facultyCustom') or get_val('faculty') or '').upper(),  # Matches {{ School/Faculty }} - Always uppercase
            'LECTURER\'S NAME': get_val('instructor'),  # Matches {{ LECTURER'S NAME }}
            'LEVEL': get_val('levelCustom') or get_val('level'),  # Matches {{ LEVEL }}
            'Month and Year': month_year,  # Matches {{ Month and Year }}
            'degree_selected': degree,  # Matches {{ degree_selected }}
            # Additional textbox fields
            'PROJECT TOPIC': (get_val('title') or '').upper(),  # Textbox field - Always uppercase
            'REPORT TITLE': (get_val('title') or '').upper(),  # Textbox field - Always uppercase
            'ASSIGNMENT_TITLE': (get_val('title') or '').upper(),  # Textbox field - Always uppercase
            'ACADEMIC YEAR': academic_year,  # Textbox field
            'Supervisor\'s Name': get_val('supervisor') or get_val('academicSupervisor'),
            'Supervisor': get_val('supervisor') or get_val('academicSupervisor'),
            'Field Supervisor\'s name': get_val('fieldSupervisor'),
            'Field Supervisor': get_val('fieldSupervisor'),
            'institution': (get_val('institutionCustom') or get_val('institution') or '').upper(),  # Always uppercase
            'title': (get_val('title') or '').upper(),  # Always uppercase
            'date': date_str,
            'assignmentNumber': get_val('assignmentNumber'),
        }

        # Scan document for actual placeholders to match case-sensitive
        found_placeholders = get_all_placeholders(doc)
        print(f"DEBUG: Found placeholders: {found_placeholders}")
        
        replacements = {}
        
        for ph in found_placeholders:
            # Extract the placeholder name (remove {{ and }})
            clean_key = ph.replace('{{', '').replace('}}', '').strip()
            
            val = ""
            
            # Direct lookup in values_map first
            if clean_key in values_map:
                val = values_map[clean_key]
            # Fallback for case-insensitive matching on common variations
            else:
                # Normalize key for fuzzy matching
                lower_key = clean_key.lower().strip()
                
                if 'student' in lower_key and 'name' in lower_key:
                    val = values_map.get('STUDENT_NAME', '')
                elif 'matricule' in lower_key or 'id' in lower_key:
                    val = values_map.get('Matricule Number', '')
                elif 'degree' in lower_key:
                    val = values_map.get('degree_selected', '')
                elif 'deparment' in lower_key or 'department' in lower_key:
                    # Prefer exact match from values_map, else empty
                    val = values_map.get('Deparment', '') or values_map.get('DEPARTMENT', '')
                elif 'faculty' in lower_key or 'school' in lower_key:
                    val = values_map.get('School/Faculty', '') or values_map.get('SCHOO/FACULTY', '')
                elif 'month' in lower_key and 'year' in lower_key:
                    val = values_map.get('Month and Year', '')
                elif 'year' in lower_key or 'academic' in lower_key:
                    val = values_map.get('Month and Year', '')
                elif 'course' in lower_key:
                    if 'code' in lower_key:
                        val = values_map.get('COURSE_CODE', '')
                    else:
                        val = values_map.get('COURSE_TITLE', '')
                elif 'lecturer' in lower_key or 'instructor' in lower_key:
                    val = values_map.get('LECTURER\'S NAME', '')
                elif 'level' in lower_key:
                    val = values_map.get('LEVEL', '')
                elif 'supervisor' in lower_key:
                    # Handle supervisor fields - check for academic/field supervisor
                    if 'field' in lower_key:
                        val = values_map.get('Field Supervisor\'s name', '') or values_map.get('Field Supervisor', '')
                    else:
                        val = values_map.get('Supervisor\'s Name', '') or values_map.get('Supervisor', '')
                elif 'title' in lower_key or 'topic' in lower_key:
                    val = values_map.get('title', '')
            
            # Add to replacements
            replacements[ph] = val
            if val:
                print(f"DEBUG: Mapped '{ph}' -> '{val}'")
            else:
                print(f"DEBUG: Mapped '{ph}' -> '' (Empty)")
        
        # Add raw placeholder keys as well for fallback matching
        replacements.update({
            '{{ STUDENT_NAME }}': values_map.get('STUDENT_NAME', ''),
            '{{ Matricule Number }}': values_map.get('Matricule Number', ''),
            '{{ COURSE_CODE }}': values_map.get('COURSE_CODE', ''),
            '{{ COURSE_TITLE }}': values_map.get('COURSE_TITLE', ''),
            '{{ DEPARTMENT }}': values_map.get('DEPARTMENT', ''),
            '{{ Deparment }}': values_map.get('Deparment', ''),
            '{{ SCHOO/FACULTY }}': values_map.get('SCHOO/FACULTY', ''),
            '{{ School/Faculty }}': values_map.get('School/Faculty', ''),
            '{{ LECTURER\'S NAME }}': values_map.get('LECTURER\'S NAME', ''),
            '{{ LEVEL }}': values_map.get('LEVEL', ''),
            '{{ Month and Year }}': values_map.get('Month and Year', ''),
            '{{ degree_selected }}': values_map.get('degree_selected', ''),
        })
        
        replace_placeholders(doc, replacements)
        
        # Apply dissertation-specific formatting
        apply_dissertation_formatting(doc, document_type)
        
        # Generate output filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        safe_title = "".join([c for c in data.get('title', 'cover_page') if c.isalpha() or c.isdigit() or c==' ']).rstrip()
        filename = f"CoverPage_{safe_title}_{timestamp}.docx"
        output_path = os.path.join(OUTPUT_DIR, filename)
        
        doc.save(output_path)
        
        return output_path, None
        
    except Exception as e:
        logger.error(f"Error generating cover page: {str(e)}")
        return None, str(e)
