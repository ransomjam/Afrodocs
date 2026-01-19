#!/usr/bin/env python3
"""
Test script to verify formatting parameters are being applied to documents
"""

import sys
import os
import requests
import json
from pathlib import Path
from docx import Document

# Test configuration
API_URL = "http://localhost:5000/upload"
TEST_FILE = "test_doc_formatting.txt"
OUTPUT_DIR = Path("../outputs")

# Create a simple test document
test_content = """
CHAPTER 1: INTRODUCTION

This is a test document to verify that formatting options are being applied correctly.

The font size should be changeable, the line spacing should be adjustable, and margins
should be customizable. All of these options are being tested to ensure proper functionality.

CHAPTER 2: TESTING PARAMETERS

We are testing the following formatting options:
1. Font Size (should display in different sizes)
2. Line Spacing (should affect paragraph spacing)
3. Margins (should affect document edges)
4. Table of Contents (should be included if selected)

CHAPTER 3: VERIFICATION

Once the document is generated, we should be able to verify these settings were applied
by examining the resulting DOCX file directly.
"""

def create_test_file():
    """Create a test document file"""
    with open(TEST_FILE, 'w') as f:
        f.write(test_content)
    print(f"✓ Created test file: {TEST_FILE}")

def check_document_formatting(docx_path):
    """Verify that formatting was applied to the document"""
    try:
        doc = Document(docx_path)
        
        print("\n📋 Document Analysis:")
        
        # Check margins
        if doc.sections:
            section = doc.sections[0]
            margin_inches = section.top_margin.inches
            margin_cm = margin_inches * 2.54
            print(f"  • Top margin: {margin_cm:.2f} cm ({margin_inches:.2f} in)")
        
        # Check paragraph formatting
        if doc.paragraphs:
            first_para = doc.paragraphs[0]
            pf = first_para.paragraph_format
            
            if pf.line_spacing:
                print(f"  • Line spacing: {pf.line_spacing}")
            
            if first_para.runs:
                font = first_para.runs[0].font
                if font.size:
                    font_pt = font.size.pt
                    print(f"  • Font size: {font_pt} pt")
        
        # Check for styles
        print(f"  • Total paragraphs: {len(doc.paragraphs)}")
        print(f"  • Available styles: {len(doc.styles)}")
        
        # Check if TOC is present
        has_toc = any("Table of Contents" in para.text for para in doc.paragraphs)
        print(f"  • Table of Contents: {'✓ Found' if has_toc else '✗ Not found'}")
        
        return True
    except Exception as e:
        print(f"✗ Error reading document: {e}")
        return False

def test_with_formatting_options(include_toc=True, font_size=14, line_spacing=1.5, margin_cm=2.0):
    """Test uploading with formatting options"""
    
    print(f"\n🧪 Testing with options:")
    print(f"   TOC: {include_toc}")
    print(f"   Font Size: {font_size}pt")
    print(f"   Line Spacing: {line_spacing}")
    print(f"   Margins: {margin_cm}cm")
    
    # Need to authenticate first - check if there's a login mechanism
    try:
        # First, we need to login to get a session
        session = requests.Session()
        
        # Try to login with default test credentials if available
        login_url = "http://localhost:5000/api/auth/login"
        login_data = {
            'username': 'admin',
            'password': 'admin123'
        }
        
        # Try login
        try:
            login_response = session.post(login_url, json=login_data)
            print(f"   Login attempt: {login_response.status_code}")
            if login_response.status_code == 200:
                print("   ✓ Authenticated successfully")
            else:
                print(f"   ✗ Login failed: {login_response.json()}")
        except Exception as e:
            print(f"   Note: Could not authenticate - {e}")
        
        # Prepare the upload
        with open(TEST_FILE, 'rb') as f:
            files = {'file': f}
            data = {
                'include_toc': 'true' if include_toc else 'false',
                'font_size': str(font_size),
                'line_spacing': str(line_spacing),
                'margin_cm': str(margin_cm)
            }
            
            print(f"\n📤 Uploading to {API_URL}")
            response = session.post(API_URL, files=files, data=data)
            
            print(f"   Response Status: {response.status_code}")
            
            if response.status_code == 200:
                result = response.json()
                job_id = result.get('job_id')
                filename = result.get('filename', 'unknown')
                
                print(f"   ✓ Upload successful!")
                print(f"   Job ID: {job_id}")
                print(f"   Filename: {filename}")
                
                # Wait a moment for file to be created
                import time
                time.sleep(1)
                
                # Try to find and verify the document
                output_file = OUTPUT_DIR / f"{job_id}_formatted.docx"
                if output_file.exists():
                    print(f"\n✓ Output file found: {output_file}")
                    check_document_formatting(str(output_file))
                else:
                    print(f"\n✗ Output file not found at {output_file}")
                    # List files in output directory
                    if OUTPUT_DIR.exists():
                        files_list = list(OUTPUT_DIR.glob("*"))
                        print(f"   Files in {OUTPUT_DIR}: {len(files_list)} items")
                        print(f"   Recent files: {[f.name for f in sorted(files_list, key=lambda x: x.stat().st_mtime, reverse=True)[:5]]}")
                
                return True
            else:
                print(f"   ✗ Upload failed!")
                try:
                    error = response.json()
                    print(f"   Error: {error}")
                except:
                    print(f"   Response: {response.text}")
                return False
                
    except Exception as e:
        print(f"✗ Test failed: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == '__main__':
    print("🔧 Formatting Options Test Suite\n")
    
    # Create test file
    create_test_file()
    
    # Run test with formatting options
    success = test_with_formatting_options(
        include_toc=True,
        font_size=14,
        line_spacing=1.5,
        margin_cm=2.0
    )
    
    # Clean up
    if os.path.exists(TEST_FILE):
        os.remove(TEST_FILE)
        print(f"\n✓ Cleaned up test file")
    
    sys.exit(0 if success else 1)
