from pattern_formatter_backend import DocumentProcessor, WordGenerator
from docx import Document

fp = r'c:/Users/user/Desktop/Afrodocs_dev/Afrodocs/Samples/Jam _ sample project with figures.docx'
proc = DocumentProcessor()
# Process
with open(fp,'rb') as f:
    proc_res = proc.process_docx(f)
# Unpack
if isinstance(proc_res, tuple):
    result, images, shapes = proc_res
else:
    result = proc_res
structured = result.get('structured', []) if isinstance(result, dict) else []

# Generate output file
out = 'verify_refs_output.docx'
wg = WordGenerator()
wg.generate(structured, out, images=images, shapes=shapes)
print('Generated', out)

# Inspect output docx paragraphs around References heading
doc = Document(out)
paras = [p.text for p in doc.paragraphs]
# find index of References heading (case-insensitive)
idx = None
for i,p in enumerate(doc.paragraphs):
    if p.text and p.text.strip().lower() in ('references','reference','bibliography','works cited'):
        idx = i
        break
if idx is None:
    print('References heading not found in generated doc')
else:
    print('References heading at paragraph index', idx)
    # print next 10 paragraphs with indent info
    for j in range(idx+1, min(len(doc.paragraphs), idx+31)):
        p = doc.paragraphs[j]
        if not p.text.strip():
            continue
        left = p.paragraph_format.left_indent
        first = p.paragraph_format.first_line_indent
        print(f'Para {j}: len={len(p.text)} left={left} first_line={first} text="{p.text[:120]}"')

print('Done')
