#!/usr/bin/env python3
"""
Check the formatting of the generated document
"""

from docx import Document
from pathlib import Path

docx_path = Path("./outputs/5ed5ff69-de26-47cd-b0f5-78e57b6a7a3a_formatted.docx")

if docx_path.exists():
    doc = Document(str(docx_path))
    
    print("📋 Document Formatting Analysis:")
    print("=" * 50)
    
    # Check margins
    if doc.sections:
        section = doc.sections[0]
        margin_inches = section.top_margin.inches
        margin_cm = margin_inches * 2.54
        print(f"\nMargins:")
        print(f"  • Top margin: {margin_cm:.2f} cm ({margin_inches:.2f} in)")
        print(f"  • Bottom margin: {section.bottom_margin.inches * 2.54:.2f} cm")
        print(f"  • Left margin: {section.left_margin.inches * 2.54:.2f} cm")
        print(f"  • Right margin: {section.right_margin.inches * 2.54:.2f} cm")
    
    # Check paragraph formatting
    print(f"\nBody Text Formatting:")
    found_body_para = False
    for i, para in enumerate(doc.paragraphs[:10]):
        if para.text.strip() and not para.text.startswith("test_doc"):
            pf = para.paragraph_format
            
            if pf.line_spacing:
                print(f"  • Line spacing (para {i}): {pf.line_spacing}")
                found_body_para = True
                break
    
    if doc.paragraphs:
        for i, para in enumerate(doc.paragraphs[:5]):
            if para.runs:
                font = para.runs[0].font
                if font.size:
                    font_pt = font.size.pt
                    print(f"  • Font size (para {i}): {font_pt} pt")
                    found_body_para = True
                    break
    
    # Check for TOC
    has_toc = any("Table of Contents" in para.text for para in doc.paragraphs)
    print(f"\n  • Table of Contents: {'✓ Found' if has_toc else '✗ Not found'}")
    
    # Check styles
    print(f"\nDocument Structure:")
    print(f"  • Total paragraphs: {len(doc.paragraphs)}")
    print(f"  • Total tables: {len(doc.tables)}")
    
    # Check default styles
    print(f"\nStyle Information:")
    try:
        normal_style = doc.styles['Normal']
        if normal_style.font.size:
            print(f"  • Normal style font size: {normal_style.font.size.pt} pt")
        if normal_style.paragraph_format.line_spacing:
            print(f"  • Normal style line spacing: {normal_style.paragraph_format.line_spacing}")
    except:
        pass
    
    print("\n" + "=" * 50)
else:
    print(f"❌ File not found: {docx_path}")
