#!/usr/bin/env python3
"""
Simple test to check what formatting is actually being applied
"""

import requests
from pathlib import Path
from docx import Document
import time
import os
import shutil

# Create a simple test document
test_file = "test_format.docx"
doc = Document()
doc.add_heading("Test Document", 0)
doc.add_paragraph("Body paragraph 1. " * 30)
doc.add_heading("Section 1", 1)
doc.add_paragraph("Body paragraph 2. " * 25)
doc.save(test_file)

# Ensure outputs directory exists
os.makedirs("outputs", exist_ok=True)
shutil.rmtree("outputs", ignore_errors=True)
os.makedirs("outputs", exist_ok=True)

# Test formatting with 18pt font
session = requests.Session()
try:
    session.post("http://localhost:5000/api/auth/login",
        json={'username': 'admin', 'password': 'admin123'})
except:
    print("Auth failed - trying without login")

params = {
    'font_size': '18',
    'line_spacing': '2.5',
    'margin_cm': '2.5',
    'include_toc': 'false'
}

print(f"Uploading test document with params: {params}")
with open(test_file, 'rb') as f:
    response = session.post("http://localhost:5000/upload",
        files={'file': f}, data=params)

if response.status_code == 200:
    result = response.json()
    job_id = result['job_id']
    print(f"Upload successful, job_id: {job_id}")
    
    time.sleep(3)
    
    output_path = f"./outputs/{job_id}_formatted.docx"
    if os.path.exists(output_path):
        print(f"\nAnalyzing output: {output_path}")
        doc = Document(output_path)
        
        print(f"\nTotal paragraphs: {len(doc.paragraphs)}")
        print("\nFirst 5 paragraphs analysis:")
        
        for idx, para in enumerate(doc.paragraphs[:5]):
            text_preview = para.text[:50] if para.text else "[empty]"
            line_spacing = para.paragraph_format.line_spacing
            print(f"\nParagraph {idx}:")
            print(f"  Text: {text_preview}")
            print(f"  Line spacing: {line_spacing}")
            
            if para.runs:
                for run_idx, run in enumerate(para.runs[:3]):
                    if run.font.size:
                        print(f"  Run {run_idx} font size: {run.font.size.pt}pt")
                    else:
                        print(f"  Run {run_idx} font size: [inherited]")
    else:
        print(f"Output file not found at {output_path}")
        print(f"Available files in outputs: {os.listdir('outputs')}")
else:
    print(f"Upload failed: {response.status_code}")
    print(response.json())

# Cleanup
if os.path.exists(test_file):
    os.remove(test_file)
