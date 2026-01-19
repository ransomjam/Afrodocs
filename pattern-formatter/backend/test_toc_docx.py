#!/usr/bin/env python3
"""
Test with a properly formatted Word document
"""

import requests
from pathlib import Path
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

# Create a test DOCX file with proper structure
test_file = "test_structured.docx"
doc = Document()

# Add title
doc.add_heading("Test Document", 0)

# Add sections with headings
doc.add_heading("Chapter 1: Introduction", 1)
doc.add_paragraph("This is the introduction section with some content.")
doc.add_paragraph("It has multiple paragraphs.")

doc.add_heading("Chapter 2: Main Content", 1)
doc.add_heading("2.1 Subsection", 2)
doc.add_paragraph("This is subsection content.")

doc.add_heading("Chapter 3: Conclusion", 1)
doc.add_paragraph("Final thoughts and conclusions.")

doc.save(test_file)

# Test with TOC
params = {
    'include_toc': 'true',
    'font_size': '14',
    'line_spacing': '1.5',
    'margin_cm': '2.0'
}

# Login
session = requests.Session()
login_response = session.post(
    "http://localhost:5000/api/auth/login",
    json={'username': 'admin', 'password': 'admin123'}
)

print(f"Login: {login_response.status_code}")

# Upload
with open(test_file, 'rb') as f:
    files = {'file': f}
    response = session.post("http://localhost:5000/upload", files=files, data=params)

if response.status_code == 200:
    result = response.json()
    job_id = result['job_id']
    print(f"Upload successful: Job ID {job_id}")
    
    import time
    time.sleep(2)
    
    # Check document
    doc_file = Path(f"./outputs/{job_id}_formatted.docx")
    if doc_file.exists():
        doc_result = Document(str(doc_file))
        
        # List all paragraphs to see TOC
        print(f"\nDocument contains {len(doc_result.paragraphs)} paragraphs")
        print("\nFirst 20 paragraphs:")
        for i, para in enumerate(doc_result.paragraphs[:20]):
            text = para.text[:80] if para.text else "(empty)"
            style = para.style.name if para.style else "None"
            print(f"  {i}: [{style}] {text}")
        
        # Check for TOC
        has_toc = any("TABLE OF CONTENTS" in para.text.upper() for para in doc_result.paragraphs[:30])
        print(f"\nTable of Contents found: {has_toc}")
    else:
        print(f"Output file not found")
else:
    print(f"Upload failed: {response.json()}")

import os
if os.path.exists(test_file):
    os.remove(test_file)
